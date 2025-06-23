#! /usr/bin/env python
# -*- coding: utf 8 -*-
"""
/***************************************************************************
        pyArchInit Plugin  - A QGIS plugin to manage archaeological dataset
                             stored in Postgres
                             -------------------
    begin                : 2007-12-01
    copyright            : (C) 2008 by Luca Mandolesi; Enzo Cocca <enzo.ccc@gmail.com>
    email                : mandoluca at gmail.com
 ***************************************************************************/
/***************************************************************************
 *                                                                          *
 *   This program is free software; you can redistribute it and/or modify   *
 *   it under the terms of the GNU General Public License as published by   *
 *   the Free Software Foundation; either version 2 of the License, or      *
 *   (at your option) any later version.                                    *
 ***************************************************************************/
"""
from __future__ import absolute_import

import ast
import csv
import json
import tempfile
from datetime import datetime

import math
import platform
import re
import sqlite3 as sq
import sys
import time


import numpy as np
import urllib.parse
import pyvista as pv
import vtk
from qgis.PyQt.QtGui import QDesktopServices,QImage


from bs4 import BeautifulSoup

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import nsdecls

from pyvistaqt import QtInteractor
import functools
from collections import OrderedDict, Counter, defaultdict
from datetime import date
from xml.etree.ElementTree import ElementTree as ET

import cv2
import matplotlib
import pandas as pd
import requests
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml

from langchain.chat_models import ChatOpenAI
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

from langchain.agents import AgentType, Tool, initialize_agent
from langchain.memory import ConversationBufferMemory

from langchain.schema import SystemMessage

matplotlib.use('QT5Agg')  # Assicurati di chiamare use() prima di importare FigureCanvas
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from qgis.PyQt import QtCore, QtGui, QtWidgets

from qgis.PyQt.QtGui import QFont,QKeySequence,QStandardItemModel,QStandardItem
from qgis.core import *
from qgis.gui import QgsMapCanvas, QgsMapToolPan
from qgis.PyQt.QtSql import QSqlDatabase, QSqlTableModel


from .Interactive_matrix import *
from ..modules.report.archeo_analysis import ArchaeologicalAnalysis
from ..modules.report.validation_tools import ArchaeologicalValidators
from ..modules.utility.report_generator import ReportGenerator
from ..modules.utility.VideoPlayer import VideoPlayerWindow
from ..modules.utility.pyarchinit_media_utility import *
from ..modules.utility.response_sql import ResponseSQL
from ..modules.utility.textTosql import *
from ..modules.db.pyarchinit_conn_strings import Connection
from ..modules.db.pyarchinit_db_manager import Pyarchinit_db_management
from ..modules.db.pyarchinit_utility import Utility
from ..modules.gis.pyarchinit_pyqgis import Pyarchinit_pyqgis, Order_layer_v2
from ..modules.utility.delegateComboBox import ComboBoxDelegate
from ..modules.utility.pyarchinit_error_check import Error_check
from ..modules.utility.pyarchinit_exp_USsheet_pdf import generate_US_pdf
from ..modules.utility.pyarchinit_print_utility import Print_utility
from ..modules.utility.settings import Settings
from ..modules.utility.skatch_gpt_US import GPTWindow
from ..searchLayers import SearchLayers
from ..gui.imageViewer import ImageViewer
from ..gui.pyarchinitConfigDialog import pyArchInitDialog_Config
from ..gui.sortpanelmain import SortPanelMain
from sqlalchemy import create_engine, MetaData, Table, select, update, and_

MAIN_DIALOG_CLASS, _ = loadUiType(
    os.path.join(os.path.dirname(__file__), os.pardir, 'gui', 'ui', 'US_USM.ui'))

#from ..modules.utility.screen_adaptative import ScreenAdaptive

class CollapsibleSection(QWidget):
    def __init__(self, title, parent=None):
        super().__init__(parent)
        self.setLayout(QVBoxLayout())

        # Header button
        self.toggle_button = QPushButton(f"▼ {title}")
        self.toggle_button.setStyleSheet("""
            QPushButton {
                text-align: left;
                padding: 5px;
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
            }
        """)
        self.toggle_button.clicked.connect(self.toggle_content)

        # Content widget
        self.content = QWidget()
        self.content_layout = QVBoxLayout(self.content)

        # Add to main layout
        self.layout().addWidget(self.toggle_button)
        self.layout().addWidget(self.content)

        # Initialize as expanded
        self.is_expanded = True

    def toggle_content(self):
        self.is_expanded = not self.is_expanded
        self.content.setVisible(self.is_expanded)
        self.toggle_button.setText(f"{'▼' if self.is_expanded else '▶'} {self.toggle_button.text()[2:]}")

    def add_widget(self, widget):
        self.content_layout.addWidget(widget)

    def add_layout(self, layout):
        self.content_layout.addLayout(layout)


class ReportGeneratorDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setWindowTitle('Generatore di Report')
        self.setModal(True)
        self.resize(500, 400)

        # Main layout
        main_layout = QVBoxLayout(self)

        # Language Section
        language_section = CollapsibleSection("Selezione Lingua")
        language_layout = QHBoxLayout()
        language_label = QLabel("Lingua output:")
        self.language_combo = QComboBox()
        self.language_combo.addItems([
            'Italiano',
            'English (UK)',
            'English (US)',
            'Español',
            'Français',
            'Deutsch',
            'العربية',
            'Ελληνικά',  # Greek
            'Русский',  # Russian
            'Português'  # Portuguese
        ])
        language_layout.addWidget(language_label)
        language_layout.addWidget(self.language_combo)
        language_section.add_layout(language_layout)
        main_layout.addWidget(language_section)

        # Tables Section
        tables_section = CollapsibleSection("Selezione Tabelle")
        self.combo_box = CheckableComboBox()
        self.TABLES_NAMES = [
            'site_table', 'us_table', 'inventario_materiali_table',
            'pottery_table', 'periodizzazione_table','struttura_table','tomba_table',

        ]
        for table_name in self.TABLES_NAMES:
            self.combo_box.add_item(table_name)
        tables_section.add_widget(QLabel("Seleziona le tabelle:"))
        tables_section.add_widget(self.combo_box)
        main_layout.addWidget(tables_section)

        # Filters Section
        filters_section = CollapsibleSection("Filtri")

        # Year filter
        self.year_input = QLineEdit()
        self.year_input.setPlaceholderText("Inserisci l'anno di scavo (opzionale)")
        filters_section.add_widget(QLabel("Anno di scavo:"))
        filters_section.add_widget(self.year_input)

        # US range filter
        filters_section.add_widget(QLabel("Range di US (se non si inserisce l'anno di scavo):"))
        range_layout = QHBoxLayout()
        self.us_start_input = QLineEdit()
        self.us_end_input = QLineEdit()
        self.us_start_input.setPlaceholderText("US iniziale")
        self.us_end_input.setPlaceholderText("US finale")
        range_layout.addWidget(QLabel("Da:"))
        range_layout.addWidget(self.us_start_input)
        range_layout.addWidget(QLabel("A:"))
        range_layout.addWidget(self.us_end_input)
        filters_section.add_layout(range_layout)

        main_layout.addWidget(filters_section)

        # Buttons Section
        button_layout = QHBoxLayout()
        self.validate_button = QPushButton("Verifica Dati Mancanti")
        self.validate_button.clicked.connect(self.validate_data)
        button_layout.addWidget(self.validate_button)
        button_layout.addStretch()
        self.buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        button_layout.addWidget(self.buttons)
        main_layout.addLayout(button_layout)




    def get_selected_language(self):
        """Get the selected output language"""
        return self.language_combo.currentText()

    def validate_data(self):
        """Esegue la validazione dei dati"""
        try:
            # Crea un'istanza temporanea di GenerateReportThread per la validazione
            report_thread = GenerateReportThread(
                custom_prompt="",
                descriptions_text="",
                api_key=self.parent.apikey_gpt(),
                selected_model="",
                selected_tables=self.get_selected_tables(),
                analysis_steps=[],
                agent=None,
                us_data=self.get_us_data(),
                materials_data=self.get_materials_data(),
                pottery_data=self.get_pottery_data(),
                site_data={},
                py_dialog=self,
                output_language=self.get_selected_language()
            )

            missing_data_report = []

            # Esegui le validazioni per le tabelle selezionate
            if 'us_table' in self.get_selected_tables():
                us_validation = report_thread.validate_us()
                if not us_validation['valid']:
                    missing_data_report.append(us_validation['message'])

            if 'inventario_materiali_table' in self.get_selected_tables():
                materials_validation = report_thread.validate_materials()
                if not materials_validation['valid']:
                    missing_data_report.append(materials_validation['message'])

            if 'pottery_table' in self.get_selected_tables():
                pottery_validation = report_thread.validate_pottery()
                if not pottery_validation['valid']:
                    missing_data_report.append(pottery_validation['message'])

            # Mostra i risultati
            if missing_data_report:
                msg = "REPORT DATI MANCANTI\n\n" + "\n\n".join(missing_data_report)
                QMessageBox.warning(self, "Dati Mancanti", msg)
            else:
                QMessageBox.information(self, "Verifica Completata", "Tutti i dati sono completi!")

        except Exception as e:
            QMessageBox.critical(self, "Errore", f"Errore durante la validazione: {str(e)}")

    def get_us_data(self):
        """Recupera i dati delle US dal database"""
        selected_tables = self.get_selected_tables()
        if 'us_table' not in selected_tables:
            return []

        conn = Connection()
        records, _ = ReportGenerator.read_data_from_db(conn.conn_str(), 'us_table')

        # Applica i filtri
        year_filter = self.get_year_filter()
        us_start, us_end = self.get_us_range()

        if year_filter:
            records = [r for r in records if str(getattr(r, 'anno_scavo', '')) == year_filter]
        elif us_start and us_end:
            records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

        return [{
            'id_us': getattr(r, 'id_us', ''),
            'us': getattr(r, 'us', ''),
            'area': getattr(r, 'area', ''),
            'settore': getattr(r, 'settore', ''),
            'd_stratigrafica': getattr(r, 'd_stratigrafica', ''),
            'descrizione': getattr(r, 'descrizione', ''),
            'interpretazione': getattr(r, 'interpretazione', ''),
            'rapporti': getattr(r, 'rapporti', '')
        } for r in records]

    def get_materials_data(self):
        """Recupera i dati dei materiali dal database"""
        selected_tables = self.get_selected_tables()
        if 'inventario_materiali_table' not in selected_tables:
            return []

        conn = Connection()
        records, _ = ReportGenerator.read_data_from_db(conn.conn_str(), 'inventario_materiali_table')

        # Applica i filtri
        year_filter = self.get_year_filter()
        us_start, us_end = self.get_us_range()

        if year_filter:
            records = [r for r in records if str(getattr(r, 'years', '')) == year_filter]
        elif us_start and us_end:
            records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

        return [{
            'id_invmat': getattr(r, 'id_invmat', ''),
            'numero_inventario': getattr(r, 'numero_inventario', ''),
            'tipo_reperto': getattr(r, 'tipo_reperto', ''),
            'definizione': getattr(r, 'definizione', ''),
            'descrizione': getattr(r, 'descrizione', ''),
            'datazione': getattr(r, 'datazione', ''),
            'stato_conservazione': getattr(r, 'stato_conservazione', ''),
            'area': getattr(r, 'area', ''),
            'us': getattr(r, 'us', '')
        } for r in records]

    def get_pottery_data(self):
        """Recupera i dati della ceramica dal database"""
        selected_tables = self.get_selected_tables()
        if 'pottery_table' not in selected_tables:
            return []

        conn = Connection()
        records, _ = ReportGenerator.read_data_from_db(conn.conn_str(), 'pottery_table')

        # Applica i filtri
        year_filter = self.get_year_filter()
        us_start, us_end = self.get_us_range()

        if year_filter:
            records = [r for r in records if str(getattr(r, 'anno', '')) == year_filter]
        elif us_start and us_end:
            records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

        return [{
            'id_rep': getattr(r, 'id_rep', ''),
            'id_number': getattr(r, 'id_number', ''),
            'fabric': getattr(r, 'fabric', ''),
            'form': getattr(r, 'form', ''),
            'area': getattr(r, 'area', ''),
            'us': getattr(r, 'us', '')
        } for r in records]

    def get_selected_tables(self):
        """Get list of checked tables"""
        return self.combo_box.items_checked()

    def get_year_filter(self):
        """Get year filter value"""
        return self.year_input.text().strip()

    def get_us_range(self):
        """Get US range values"""
        us_start = self.us_start_input.text().strip()
        us_end = self.us_end_input.text().strip()
        return us_start, us_end


class CheckableComboBox(QComboBox):
    def __init__(self):
        super().__init__()
        self.setModel(QStandardItemModel(self))
        self.view().pressed.connect(self.handle_item_pressed)
        self.setStyleSheet("QComboBox { combobox-popup: 0; }")

    def add_item(self, text):
        """Add a checkable item to the combo box"""
        item = QStandardItem(text)
        item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
        item.setData(Qt.Unchecked, Qt.CheckStateRole)
        self.model().appendRow(item)

    def items_checked(self):
        """Get list of checked items"""
        checked_items = []
        for index in range(self.count()):
            item = self.model().item(index)
            if item.data(Qt.CheckStateRole) == Qt.Checked:
                checked_items.append(item.text())
        return checked_items

    def handle_item_pressed(self, index):
        """Handle item check/uncheck"""
        item = self.model().itemFromIndex(index)
        if item.data(Qt.CheckStateRole) == Qt.Checked:
            item.setData(Qt.Unchecked, Qt.CheckStateRole)
        else:
            item.setData(Qt.Checked, Qt.CheckStateRole)


class ReportDialog(QDialog):
    def __init__(self, content="", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Report Preview")
        self.setModal(True)
        self.resize(1000, 800)

        # Create main vertical layout
        self.main_layout = QVBoxLayout(self)

        # Create splitter for resizable sections
        self.splitter = QSplitter(Qt.Vertical)

        # Report Preview Section
        self.report_widget = QWidget()
        report_layout = QVBoxLayout(self.report_widget)

        report_header = QLabel("Report Preview")
        report_header.setStyleSheet("font-weight: bold; font-size: 14px; padding: 5px;")
        report_layout.addWidget(report_header)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setFont(QFont("Arial", 15))
        self.text_edit.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
        """)
        # Abilita la selezione del testo e i link
        self.text_edit.setTextInteractionFlags(
            Qt.TextSelectableByMouse |
            Qt.TextSelectableByKeyboard |
            Qt.LinksAccessibleByMouse
        )

        # Connetti l'evento mousePressEvent
        self.text_edit.mousePressEvent = self.handle_mouse_press

        if content:
            self.text_edit.setHtml(content)
        report_layout.addWidget(self.text_edit)


        # Terminal Section
        self.terminal_widget = QWidget()
        terminal_layout = QVBoxLayout(self.terminal_widget)

        terminal_header = QLabel("Process Terminal")
        terminal_header.setStyleSheet("font-weight: bold; font-size: 14px; padding: 5px;")
        terminal_layout.addWidget(terminal_header)

        self.terminal = QTextEdit()
        self.terminal.setReadOnly(True)
        self.terminal.setFont(QFont("Courier", 13))
        self.terminal.setStyleSheet("""
            QTextEdit {
                background-color: #1a1a1a;
                color: #ffffff;
                border: 1px solid #333;
                border-radius: 4px;
            }
        """)
        terminal_layout.addWidget(self.terminal)

        # Add widgets to splitter
        self.splitter.addWidget(self.report_widget)
        self.splitter.addWidget(self.terminal_widget)

        # Set initial sizes (2/3 for report, 1/3 for terminal)
        self.splitter.setSizes([500, 300])

        # Add splitter to main layout
        self.main_layout.addWidget(self.splitter)

        # Add buttons
        button_layout = QHBoxLayout()

        self.save_button = QPushButton("Save Report")
        self.save_button.clicked.connect(self.save_report)
        button_layout.addWidget(self.save_button)

        self.copy_button = QPushButton("Copy to Clipboard")
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        button_layout.addWidget(self.copy_button)

        self.close_button = QPushButton("Close")
        self.close_button.clicked.connect(self.reject)
        button_layout.addWidget(self.close_button)

        self.main_layout.addLayout(button_layout)

    def update_content(self, new_content):
        """Update the report content"""
        try:
            if new_content:
                self.text_edit.clear()
                self.text_edit.setHtml(new_content)
                scrollbar = self.text_edit.verticalScrollBar()
                scrollbar.setValue(scrollbar.maximum())
        except Exception as e:
            print(f"Errore nell'aggiornamento del contenuto: {str(e)}")

    def handle_mouse_press(self, event):
        """Gestisce il click del mouse nel text edit"""
        if event.button() == Qt.LeftButton:
            cursor = self.text_edit.cursorForPosition(event.pos())
            format = cursor.charFormat()
            if format.isAnchor():
                url = format.anchorHref()
                if url:
                    try:
                        # Aggiungi log per debug
                        print(f"URL clicked: {url}")

                        # Se l'URL è già un file path completo
                        if os.path.exists(url):
                            qurl = QUrl.fromLocalFile(url)
                            QDesktopServices.openUrl(qurl)
                            return

                        # Se l'URL inizia con file://
                        if url.startswith('file://'):
                            path = url[7:]
                            if os.path.exists(path):
                                qurl = QUrl.fromLocalFile(path)
                                QDesktopServices.openUrl(qurl)
                                return

                        # Prova a creare un QUrl direttamente
                        qurl = QUrl(url)
                        if qurl.isLocalFile() and os.path.exists(qurl.toLocalFile()):
                            QDesktopServices.openUrl(qurl)
                            return

                        # Se niente funziona, prova a costruire un QUrl dal path locale
                        qurl = QUrl.fromLocalFile(os.path.abspath(url))
                        QDesktopServices.openUrl(qurl)

                    except Exception as e:
                        print(f"Error opening URL: {str(e)}")
                        # Mostra un messaggio di errore all'utente
                        QMessageBox.warning(
                            self,
                            "Errore",
                            f"Impossibile aprire l'immagine:\n{url}\n\nErrore: {str(e)}"
                        )

        # Chiamata al metodo originale per gestire altri eventi del mouse
        QTextEdit.mousePressEvent(self.text_edit, event)

    def log_to_terminal(self, message, msg_type="info"):
        """Add a new log message to the terminal"""
        colors = {
            'info': '#FFFFFF',  # bianco per info
            'warning': '#f1c40f',  # Giallo per warning
            'error': '#e74c3c',  # Rosso per errori
            'success': '#27ae60',  # Verde scuro per successi
            'step': '#3498db',  # Blu per gli step
            'validation': '#2ecc71'  # Verde per validazione
        }
        color = colors.get(msg_type, colors["info"])
        timestamp = QDateTime.currentDateTime().toString("hh:mm:ss")

        formatted_message = f'<span style="color: {color};">[{timestamp}][{msg_type.upper()}]</span> {message}'
        self.terminal.append(formatted_message)

        scrollbar = self.terminal.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def add_toc(self, doc):
        """
        Inserisce un 'campo TOC' nel documento, visibile come indice.
        Quando apri il doc in Word, fai tasto destro > "Aggiorna campo"
        o vai su "Riferimenti > Aggiorna sommario" per visualizzare l'indice.
        """


        # Aggiungi la pagina dell'indice
        doc.add_paragraph()
        indice_heading = doc.add_heading('INDICE', level=1)
        indice_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Aggiungi il campo TOC
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # Crea il campo TOC
        fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        run._r.append(fldChar)

        instrText = parse_xml(r'<w:instrText {} >TOC \o "1-3" \h \z</w:instrText>'.format(nsdecls('w')))
        run._r.append(instrText)

        fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        run._r.append(fldChar)

        # Aggiungi un'interruzione di pagina dopo l'indice
        doc.add_page_break()
        # Testo segnaposto (visibile finché non aggiorni il campo in Word)
        run2 = paragraph.add_run("Indice (vai in riferimenti e clicca su aggiorna sommario).")
        run2.italic = True



    def save_report(self):
        """Save the report content to a file"""

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            "",
            "Word Documents (*.docx);;Text Files (*.txt);;All Files (*)"
        )

        if not file_path:
            return

        try:
            if file_path.endswith('.docx'):
                directory = os.path.dirname(file_path)
                if not os.access(directory, os.W_OK):
                    self.log_to_terminal(f"Non hai i permessi di scrittura nella cartella: {directory}")
                    return

                # Se il file esiste già, verifica se è scrivibile
                if os.path.exists(file_path):
                    if not os.access(file_path, os.W_OK):
                        self.log_to_terminal(f"Non hai i permessi di scrittura sul file: {file_path}")
                        return
                doc = Document()

                # Imposta lo stile del documento
                sections = doc.sections
                for section in sections:
                    section.page_margin_top = Inches(2)
                    section.page_margin_bottom = Inches(2)
                    section.page_margin_left = Inches(2.5)
                    section.page_margin_right = Inches(2.5)

                # Aggiungi intestazione
                header = sections[0].header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                current_date = datetime.now().strftime("%d/%m/%Y")
                header_run = header_para.add_run(f'Report di scavo: {current_date}')
                header_run.font.size = Pt(11)

                # Aggiungi il titolo principale
                title = doc.add_heading('REPORT DI SCAVO ARCHEOLOGICO', level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # # Aggiungi informazioni sul sito
                # site_info = doc.add_paragraph()
                # site_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # site_info.add_run('Area 1 - Settore 7\n').bold = True
                # site_info.add_run('Al-Khutm\n').bold = True

                # Aggiungi il titolo dell'indice
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run('')
                font = run.font
                font.size = Pt(16)
                font.bold = True

                doc.add_paragraph()  # Spazio vuoto

                self.add_toc(doc)
                doc.add_paragraph()  # Spazio vuoto
                # Aggiungi l'interruzione di pagina
                doc.add_page_break()

                # Contatore per le figure
                figure_counter = 1

                # Ottieni il contenuto HTML dal QTextEdit
                html_content = self.text_edit.toHtml()

                # Parse HTML content
                soup = BeautifulSoup(html_content, 'html.parser')

                # Process the HTML content
                self.process_html_content(soup, doc, figure_counter)

                try:
                    doc.save(file_path)
                    QMessageBox.information(self, "Success", "Report saved successfully!")
                except PermissionError as pe:
                    QMessageBox.critical(self, "Error", f"Errore di permessi durante il salvataggio: {str(pe)}")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Errore durante il salvataggio: {str(e)}")

            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.text_edit.toPlainText())

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving report: {str(e)}")
            print(f"Detailed error: {str(e)}")  # Per debug

    def process_html_content(self, soup, doc, figure_counter):
        """Process HTML content and convert it to Word document format"""

        # Dictionary to track which images have been used
        if not hasattr(self, 'used_images_in_report'):
            self.used_images_in_report = {}

        # Find all content divs
        content_divs = soup.find_all('div', style=lambda s: s and 'font-family: Arial' in s)

        if not content_divs:
            # Fallback to processing the entire body
            content_divs = [soup.body]

        for content_div in content_divs:
            self.process_html_element(content_div, doc, figure_counter)

    def process_html_element(self, element, doc, figure_counter):
        """Process an HTML element and its children recursively"""

        # Skip script and style elements
        if element.name in ['script', 'style']:
            return

        # Process element based on its type
        if element.name == 'h2':
            # Main section heading
            heading_text = element.get_text().strip()
            heading = doc.add_heading(heading_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.style = 'Heading 1'

        elif element.name == 'h3':
            # Subsection heading
            heading_text = element.get_text().strip()
            heading = doc.add_heading(heading_text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            #heading.style = 'Heading 2'

        elif element.name == 'h4':
            # Sub-subsection heading
            heading_text = element.get_text().strip()
            heading = doc.add_heading(heading_text, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.style = 'Heading 3'

        elif element.name == 'h5':
            # Sub-sub-subsection heading
            heading_text = element.get_text().strip()
            heading = doc.add_heading(heading_text, level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.style = 'Heading 4'

        elif element.name == 'h6':
            # Lowest level heading
            heading_text = element.get_text().strip()
            heading = doc.add_heading(heading_text, level=5)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.style = 'Heading 5'

        elif element.name == 'table':
            # Process table
            self.convert_html_table_to_docx(element, doc)

        elif element.name == 'img':
            # Process image
            self.process_image(element, doc, figure_counter)
            figure_counter += 1

        elif element.name == 'div' and element.find('img'):
            # Process image container div
            img = element.find('img')
            self.process_image(img, doc, figure_counter)

            # Process caption if present
            caption_div = element.find('div', style=lambda s: s and 'font-style: italic' in s)
            if caption_div:
                caption_text = caption_div.get_text().strip()
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"Figura {figure_counter}: {caption_text.replace('Immagine ', '')}")
                caption_run.italic = True
                caption_run.font.size = Pt(9)
                caption_run.font.name = 'Calibri'

                # Add space after image
                doc.add_paragraph()

            figure_counter += 1

        elif element.name in ['p', 'div'] and element.get_text().strip():
            # Process paragraph text
            text = element.get_text().strip()

            # Skip if this is just a separator line (======)
            if all(c == '=' for c in text):
                return

            # Process markdown formatting (bold and italic)
            # Check if text contains markdown formatting
            if '**' in text or '*' in text:
                # First handle bold formatting
                if '**' in text:
                    # Split the text by ** markers
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                else:
                    parts = [text]  # No bold formatting, treat as a single part

                # Create a new paragraph
                p = doc.add_paragraph()

                # Process each part for bold formatting
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # This is a bold part, add it with bold formatting
                        bold_text = part[2:-2]  # Remove ** markers

                        # Check for italic formatting within bold text
                        if '*' in bold_text and not bold_text.startswith('*') and not bold_text.endswith('*'):
                            # Split by italic markers
                            italic_parts = re.split(r'(\*.*?\*)', bold_text)
                            for italic_part in italic_parts:
                                if italic_part.startswith('*') and italic_part.endswith('*'):
                                    # This is an italic part within bold text
                                    italic_text = italic_part[1:-1]  # Remove * markers
                                    run = p.add_run(italic_text)
                                    run.bold = True
                                    run.italic = True
                                else:
                                    # Regular bold text
                                    run = p.add_run(italic_part)
                                    run.bold = True
                        else:
                            # Regular bold text without italic
                            run = p.add_run(bold_text)
                            run.bold = True
                    else:
                        # This is a regular part, check for italic formatting
                        if '*' in part:
                            # Split by italic markers
                            italic_parts = re.split(r'(\*.*?\*)', part)
                            for italic_part in italic_parts:
                                if italic_part.startswith('*') and italic_part.endswith('*'):
                                    # This is an italic part
                                    italic_text = italic_part[1:-1]  # Remove * markers
                                    run = p.add_run(italic_text)
                                    run.italic = True
                                else:
                                    # Regular text
                                    run = p.add_run(italic_part)
                        else:
                            # Regular text without formatting
                            run = p.add_run(part)
                            run.bold = False

                # Set paragraph formatting
                paragraph_format = p.paragraph_format
                paragraph_format.space_after = Pt(10)

                # Return early since we've already added the paragraph
                return

            # Check if this is a bullet point with an image reference
            if (text.startswith('- ') or text.startswith('* ')) and '[IMMAGINE' in text:
                # Extract the image reference
                img_match = re.search(r'\[IMMAGINE[^:]*:\s+(.*?),\s*(.*?)\]', text)
                if img_match:
                    # Process the image
                    img_path = img_match.group(1).strip()
                    caption = img_match.group(2).strip()

                    try:
                        if os.path.exists(img_path):
                            # Add image to document
                            picture = doc.add_picture(img_path)

                            # Resize proportionally
                            max_width_px = 450
                            width = picture.width
                            height = picture.height
                            aspect_ratio = width / height

                            if width > max_width_px * 9525:  # Convert to EMU
                                picture.width = max_width_px * 9525
                                picture.height = int((max_width_px * 9525) / aspect_ratio)

                            # Center the image
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # Add caption
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption_para.add_run(f"Figura {figure_counter}: {caption}")
                            caption_run.italic = True
                            caption_run.font.size = Pt(9)
                            caption_run.font.name = 'Calibri'

                            # Add space after image
                            doc.add_paragraph()

                            figure_counter += 1
                        else:
                            # If image doesn't exist, add a placeholder
                            doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                    except Exception as e:
                        # If there's an error, add a placeholder
                        doc.add_paragraph(f"[Errore con l'immagine: {str(e)}]")
                else:
                    # If no image reference is found, process as a normal bullet point
                    doc.add_paragraph(text[2:].strip(), style='List Bullet')
            # Check if this is a regular bullet point
            elif text.startswith('- ') or text.startswith('* '):
                doc.add_paragraph(text[2:].strip(), style='List Bullet')
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                # Ensure text is not bold by default
                run = p.add_run(text)
                run.bold = False  # Explicitly set bold to False
                paragraph_format = p.paragraph_format
                paragraph_format.space_after = Pt(10)

        # Process children recursively
        # Check if element has children attribute before iterating
        if hasattr(element, 'children'):
            for child in element.children:
                if hasattr(child, 'name'):  # Only process tag elements
                    self.process_html_element(child, doc, figure_counter)

    def convert_html_table_to_docx(self, table_element, doc):
        """Convert an HTML table to a Word table"""

        # Get all rows
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Count columns (use the row with the most cells)
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)

        # Create Word table
        word_table = doc.add_table(rows=len(rows), cols=max_cols)

        # Try to set table style
        try:
            word_table.style = 'Table Grid'
        except Exception:
            # Apply borders manually if style fails
            try:
                for cell in word_table._tbl.findall(".//w:tc", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                    tcPr = cell.find(".//w:tcPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if tcPr is None:
                        tcPr = parse_xml(r'<w:tcPr {0}/>'.format(nsdecls('w')))
                        cell.append(tcPr)

                    # Add borders to the cell
                    tcBorders = tcPr.find(".//w:tcBorders", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if tcBorders is None:
                        tcBorders = parse_xml(r'<w:tcBorders {0}/>'.format(nsdecls('w')))
                        tcPr.append(tcBorders)

                    # Define borders for all sides
                    for side in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(".//{0}:{1}".format('w', side))
                        if border is None:
                            border = parse_xml(r'<w:{0} {1} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'.format(side, nsdecls('w')))
                            tcBorders.append(border)
            except Exception as border_error:
                print(f"Error applying borders: {str(border_error)}")

        # Fill table with data
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < max_cols:  # Ensure we don't exceed column count
                    cell_text = cell.get_text().strip()
                    # Clear existing text in the cell to avoid duplication
                    for paragraph in word_table.cell(i, j).paragraphs:
                        paragraph.clear()

                    # Add text to the cell
                    paragraph = word_table.cell(i, j).paragraphs[0]
                    run = paragraph.add_run(cell_text)

                    # Make header cells bold
                    if cell.name == 'th' or i == 0:
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a paragraph after the table for spacing
        doc.add_paragraph()

    def process_image(self, img_element, doc, figure_counter):
        """Process an image element and add it to the document"""

        src = img_element.get('src', '')
        if not src:
            return

        if src.startswith('file:///'):
            src = src[8:]

        # Modify the path to use the original image instead of the thumbnail
        if '_thumb' in src:
            original_path = src.replace('_thumb.png', '.png')
            if '/thumbnails/' in original_path:
                original_path = original_path.replace('/thumbnails/', '/')
            if os.path.exists(original_path):
                self.log_to_terminal(f"Replacing thumbnail with original: {original_path}", "info")
                src = original_path

        try:
            if os.path.exists(src):
                # Add image to document
                picture = doc.add_picture(src)

                # Resize proportionally
                max_width_px = 450
                width = picture.width
                height = picture.height
                aspect_ratio = width / height

                if width > max_width_px * 9525:  # Convert to EMU
                    picture.width = max_width_px * 9525
                    picture.height = int((max_width_px * 9525) / aspect_ratio)

                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            self.log_to_terminal(f"Errore con l'immagine {src}: {str(e)}", "error")
            print(f"Dettaglio errore: {str(e)}")  # per debug



    def copy_to_clipboard(self):
        """Copy the report content to clipboard"""
        clipboard = QApplication.clipboard()
        clipboard.setText(self.text_edit.toPlainText())
        QMessageBox.information(self, "Success", "Report copied to clipboard!")


class GenerateReportThread(QThread):
    report_generated = pyqtSignal(str)
    log_message = pyqtSignal(str, str)
    report_completed = pyqtSignal(str, dict)

    def __init__(self, custom_prompt, descriptions_text, api_key, selected_model, selected_tables, analysis_steps,
                 agent, us_data, materials_data, pottery_data, site_data, py_dialog, output_language='italiano',
                 tomba_data=None, periodizzazione_data=None, struttura_data=None):
        super().__init__()

        self.custom_prompt = custom_prompt
        self.descriptions_text = descriptions_text
        self.api_key = api_key
        self.selected_model = selected_model
        self.selected_tables = selected_tables
        self.analysis_steps = analysis_steps
        self.agent = agent
        self.us_data = us_data
        self.materials_data = materials_data
        self.pottery_data = pottery_data
        self.site_data = site_data
        self.tomba_data = tomba_data if tomba_data is not None else []
        self.periodizzazione_data = periodizzazione_data if periodizzazione_data is not None else []
        self.struttura_data = struttura_data if struttura_data is not None else []
        self.py_dialog = py_dialog
        self.output_language = output_language
        self.full_report = ""
        self.formatted_report = ""  # Inizializza qui la variabile

    def create_vector_db(self, data, table_name):
        """
        Create a vector database from the data for RAG approach.

        Args:
            data: List of data records
            table_name: Name of the table for context

        Returns:
            FAISS vector store for retrieval
        """
        if not data:
            return None

        # Convert data records to text documents
        documents = []
        for i, record in enumerate(data):
            if isinstance(record, dict):
                # Format dictionary as text
                content = f"Record {i+1} from {table_name}:\n"
                content += "\n".join(f"{k}: {v}" for k, v in record.items())
            else:
                # Format object as text
                content = f"Record {i+1} from {table_name}:\n"
                content += "\n".join(f"{k}: {getattr(record, k, '')}" for k in dir(record) 
                                    if not k.startswith('_') and not callable(getattr(record, k, None)))

            documents.append(content)

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            length_function=len,
        )
        texts = text_splitter.create_documents(documents)

        # Create vector store
        try:
            embeddings = OpenAIEmbeddings(api_key=self.api_key)
            vector_store = FAISS.from_documents(texts, embeddings)
            return vector_store
        except Exception as e:
            self.log_message.emit(f"Error creating vector database: {str(e)}", "error")
            return None

    def retrieve_relevant_data(self, vector_store, query, k=5):
        """
        Retrieve the most relevant data from the vector store based on the query.

        Args:
            vector_store: FAISS vector store
            query: Query string
            k: Number of documents to retrieve

        Returns:
            String containing the retrieved documents
        """
        if not vector_store:
            return ""

        try:
            # Retrieve relevant documents
            docs = vector_store.similarity_search(query, k=k)

            # Format the retrieved documents
            retrieved_data = "\n\n".join([doc.page_content for doc in docs])

            return retrieved_data
        except Exception as e:
            self.log_message.emit(f"Error retrieving data: {str(e)}", "error")
            return ""

    def create_rag_chain(self, vector_store, llm):
        """
        Create a RetrievalQA chain for the RAG approach.

        Args:
            vector_store: FAISS vector store
            llm: Language model

        Returns:
            RetrievalQA chain
        """
        if not vector_store:
            return None

        try:
            # Create a prompt template
            template = """
            You are an archaeological expert analyzing data from an excavation.

            Use the following pieces of context to answer the question at the end.
            If you don't know the answer, just say that you don't know, don't try to make up an answer.

            Context:
            {context}

            Question: {question}

            Answer:
            """

            prompt = PromptTemplate(
                template=template,
                input_variables=["context", "question"]
            )

            # Create the chain
            chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=vector_store.as_retriever(),
                chain_type_kwargs={"prompt": prompt}
            )

            return chain
        except Exception as e:
            self.log_message.emit(f"Error creating RAG chain: {str(e)}", "error")
            return None

    def count_tokens(self, text):
        """
        Estimate the number of tokens in a text.
        This is a simple estimation based on character count.

        Args:
            text: The text to count tokens for

        Returns:
            Estimated token count
        """
        if not text:
            return 0

        # Roughly 4 characters per token for English text
        return len(text) // 4

    def validate_us(self):
        """Validate US data using ArchaeologicalValidators"""
        context = {
            "us_data": self.us_data
        }
        return ArchaeologicalValidators.validate_us(context)

    def validate_materials(self):
        """Validate materials data using ArchaeologicalValidators"""
        context = {
            "materials_data": self.materials_data
        }
        return ArchaeologicalValidators.validate_materials(context)

    def validate_pottery(self):
        """Validate pottery data using ArchaeologicalValidators"""
        context = {
            "pottery_data": self.pottery_data
        }
        return ArchaeologicalValidators.validate_pottery(context)

    def validate_tomba(self):
        """Validate tomb data using ArchaeologicalValidators"""
        context = {
            "tomba_data": self.tomba_data
        }
        return ArchaeologicalValidators.validate_tomba(context)

    def validate_periodizzazione(self):
        """Validate periodization data using ArchaeologicalValidators"""
        context = {
            "periodizzazione_data": self.periodizzazione_data
        }
        return ArchaeologicalValidators.validate_periodizzazione(context)

    def validate_struttura(self):
        """Validate structure data using ArchaeologicalValidators"""
        context = {
            "struttura_data": self.struttura_data
        }
        return ArchaeologicalValidators.validate_struttura(context)

    def format_prompt_from_json(self, prompt_template):
        """Converte il template JSON in un prompt testuale strutturato in modo sicuro"""
        try:
            prompt = []

            # Aggiungi l'obiettivo se presente
            if 'objective' in prompt_template:
                prompt.append(f"Objective: {prompt_template['objective']}\n")

            # Aggiungi le analisi richieste se presenti
            if 'required_analysis' in prompt_template:
                prompt.append("Required Analysis:")
                for analysis in prompt_template['required_analysis']:
                    prompt.append(f"- {analysis}")
                prompt.append("")  # linea vuota per separazione

            # Aggiungi i requisiti di output se presenti
            if 'output_requirements' in prompt_template and 'sections' in prompt_template['output_requirements']:
                prompt.append("Output Requirements:")
                for section in prompt_template['output_requirements']['sections']:
                    title = section.get('title', '')
                    if title:
                        prompt.append(f"\n{title}:")

                    # Gestisci i content requirements in modo ricorsivo e sicuro
                    if 'content_requirements' in section:
                        def format_requirements(requirements, indent=2):
                            formatted = []
                            for req in requirements:
                                if isinstance(req, str):
                                    formatted.append(f"{'  ' * indent}- {req}")
                                elif isinstance(req, dict):
                                    for key, value in req.items():
                                        formatted.append(f"{'  ' * indent}- For each {key}:")
                                        for sub_req in value:
                                            formatted.append(f"{'  ' * (indent + 1)}- {sub_req}")
                            return formatted

                        prompt.extend(format_requirements(section['content_requirements']))

            # Aggiungi le specifiche di stile se presenti
            if 'style' in prompt_template:
                prompt.append("\nStyle Requirements:")
                style = prompt_template['style']
                if 'tone' in style:
                    prompt.append(f"- Tone: {style['tone']}")
                if 'language' in style:
                    prompt.append(f"- Language: {style['language']}")
                if 'structure' in style:
                    prompt.append(f"- Structure: {style['structure']}")

            # Se il prompt_template è una stringa, usala direttamente
            elif isinstance(prompt_template, str):
                return prompt_template

            # Se non c'è una struttura JSON, ma c'è un template semplice
            elif isinstance(prompt_template, dict) and 'template' in prompt_template:
                return prompt_template['template']

            return "\n".join(prompt)

        except Exception as e:
            self.log_message.emit(f"Warning: Error formatting prompt template: {str(e)}. Using simple format.",
                                  "warning")
            # Fallback a un formato semplice se qualcosa va storto
            if isinstance(prompt_template, str):
                return prompt_template
            elif isinstance(prompt_template, dict) and 'prompt_template' in prompt_template:
                return prompt_template['prompt_template']
            else:
                return str(prompt_template)

    def get_language_instructions(self):
        """Get specific instructions for the selected language"""
        instructions = {
            'Italiano': {
                'style': 'Scrivi in italiano formale e tecnico.',
                'terms': 'Usa terminologia archeologica italiana standard.',
                'format': 'Mantieni i nomi tecnici in italiano.'
            },
            'English': {
                'style': 'Write in formal and technical English.',
                'terms': 'Use standard archaeological terminology in English.',
                'format': 'Keep technical terms in English.'
            },
            'Español': {
                'style': 'Escribe en español formal y técnico.',
                'terms': 'Utiliza terminología arqueológica estándar en español.',
                'format': 'Mantén los términos técnicos en español.'
            },
            'Français': {
                'style': 'Écrivez en français formel et technique.',
                'terms': 'Utilisez la terminologie archéologique standard en français.',
                'format': 'Gardez les termes techniques en français.'
            },
            'Deutsch': {
                'style': 'Schreiben Sie in formellem und technischem Deutsch.',
                'terms': 'Verwenden Sie standardisierte archäologische Terminologie auf Deutsch.',
                'format': 'Behalten Sie technische Begriffe auf Deutsch bei.'
            }
        }
        return instructions.get(self.output_language, instructions['Italiano'])



    def format_for_widget(self, text):
        """Converte il formato immagine per la visualizzazione nel widget."""
        # Dictionary to track which images have been used
        used_images = {}

        def replace_image(match):
            """
            Gestisce la sostituzione delle immagini con il markup HTML appropriato.
            """
            # Check if match is a string or a match object
            if isinstance(match, str):
                full_match = match
            else:
                full_match = match.group(0)  # Il match completo

            # Debug print
            print(f"Match trovato: {full_match}")

            # Check if this is a list item with an image reference
            is_list_item = full_match.startswith('- [IMMAGINE')
            if is_list_item:
                # Remove the list marker
                full_match = full_match[2:].strip()

            # Estrai le informazioni con una regex più generica
            pattern_parts = re.match(r'\[IMMAGINE\s+([^:]+):\s+(.*?),\s+(.*?)\]', full_match)
            if pattern_parts:
                number = pattern_parts.group(1)
                path = pattern_parts.group(2)
                caption = pattern_parts.group(3)

                if path.startswith('file://'):
                    path = path[7:]

                # Check if this image has already been used
                image_key = f"{number}_{path}"
                if image_key in used_images:
                    print(f"Skipping duplicate image: {image_key}")
                    return ""  # Skip duplicate images

                # Mark this image as used
                used_images[image_key] = True

                # Debug print
                print(f"Processando immagine: numero={number}, path={path}, caption={caption}")

                # Carica l'immagine
                image = QImage(path)
                if not image.isNull():
                    # Ridimensiona mantenendo l'aspect ratio
                    scaled_image = image.scaled(500, 400, Qt.KeepAspectRatio, Qt.SmoothTransformation)

                    # Create a temporary directory for thumbnails if it doesn't exist
                    import os
                    thumb_dir = os.path.join(os.path.dirname(path), "thumbnails")
                    if not os.path.exists(thumb_dir):
                        os.makedirs(thumb_dir)

                    # Save the thumbnail in the thumbnails directory
                    filename = os.path.basename(path)
                    temp_path = os.path.join(thumb_dir, filename.replace('.png', '_thumb.png'))
                    scaled_image.save(temp_path)

                    return f'''
                        <div style="margin: 10px 0; text-align: center;">
                            <img src="{temp_path}"/>
                            <div style="font-style: italic; margin-top: 5px;">
                                Immagine {caption}
                            </div>
                        </div>
                        '''
            return full_match

        def convert_to_html(text):
            lines = text.split('\n')
            html_lines = []
            i = 0
            in_list = False
            list_items = []

            # Function to convert markdown table to HTML table
            def process_markdown_table(table_lines):
                html_table = ['<table style="width:100%; border-collapse: collapse; margin: 15px 0;">']

                # Process header row
                header_row = None
                separator_row = None

                for idx, line in enumerate(table_lines):
                    if idx == 0:  # First row is header
                        header_row = line
                    elif idx == 1 and re.match(r'^\|\s*[-:]+\s*\|', line):  # Second row is separator
                        separator_row = line
                        continue

                # Process header if found
                if header_row:
                    cells = [cell.strip() for cell in header_row.strip('|').split('|')]
                    html_table.append('<tr>')
                    for cell in cells:
                        html_table.append(f'<th style="border: 1px solid #000; padding: 8px; text-align: left; background-color: #f2f2f2; font-weight: bold;">{cell}</th>')
                    html_table.append('</tr>')

                # Process data rows
                start_idx = 1 if separator_row else 0
                if header_row and not separator_row:
                    start_idx = 1

                for line in table_lines[start_idx:]:
                    if re.match(r'^\|\s*[-:]+\s*\|', line):  # Skip separator lines
                        continue

                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                    html_table.append('<tr>')
                    for cell in cells:
                        html_table.append(f'<td style="border: 1px solid #000; padding: 8px; vertical-align: top;">{cell}</td>')
                    html_table.append('</tr>')

                html_table.append('</table>')
                return '\n'.join(html_table)

            # Function to process list items and create HTML list
            def process_list_items(items):
                html_list = ['<ul style="margin: 10px 0; padding-left: 20px;">']
                for item in items:
                    # Remove the list marker and trim
                    item_text = item.lstrip('- ').strip()
                    # Apply formatting to the item text
                    item_text = format_text(item_text)
                    html_list.append(f'<li style="margin: 5px 0; font-weight: normal;">{item_text}</li>')
                html_list.append('</ul>')
                return '\n'.join(html_list)

            # Function to apply text formatting (bold, italic)
            def format_text(text):
                # Convert markdown bold formatting to HTML bold
                text = re.sub(r'\*\*(.*?)\*\*', r'<span style="font-weight: bold;">\1</span>', text)
                # Convert markdown italic formatting to HTML italic
                text = re.sub(r'\*(.*?)\*', r'<span style="font-style: italic;">\1</span>', text)
                # Process image references
                text = re.sub(r'\[IMMAGINE[^:]*:.*?\]', replace_image, text)
                return text

            # Check if a line is a section title
            def is_section_title(line):
                section_titles = ["INTRODUZIONE", "METODOLOGIA DI SCAVO",
                                 "ANALISI STRATIGRAFICA E INTERPRETAZIONE",
                                 "CATALOGO DEI MATERIALI", "DESCRIZIONE DEI MATERIALI", "CONCLUSIONI"]
                return line in section_titles

            # Check if a line is a standalone section title (not part of a paragraph)
            def is_standalone_section_title(line, next_line):
                return is_section_title(line) and next_line and all(c == '=' for c in next_line.strip())

            # Track processed section titles to avoid duplicates
            processed_section_titles = set()

            while i < len(lines):
                line = lines[i].strip()

                # Process any accumulated list items before handling other elements
                if in_list and (not line.startswith('- ') or not line):
                    html_lines.append(process_list_items(list_items))
                    list_items = []
                    in_list = False

                # Skip empty lines
                if not line:
                    i += 1
                    continue

                # Gestisci i titoli delle sezioni (solo se sono seguiti da una linea di '=')
                if i + 1 < len(lines) and is_standalone_section_title(line, lines[i + 1].strip()):
                    # Check if this section title has already been processed
                    if line not in processed_section_titles:
                        html_lines.append(
                            f'<h2 style="font-size: 16pt; font-weight: bold; margin: 20px 0 10px 0;">{line}</h2>')
                        # Add to processed section titles
                        processed_section_titles.add(line)
                    i += 2  # Skip the title and the '===' line
                    continue

                # Gestisci i sottotitoli con diversi livelli (####, ###, ##)
                if line.startswith('#'):
                    heading_level = len(re.match(r'^#+', line).group(0))
                    heading_text = line.lstrip('#').strip()

                    if heading_level == 2:  # ##
                        html_lines.append(
                            f'<h3 style="font-size: 14pt; font-weight: bold; margin: 15px 0 10px 0;">{heading_text}</h3>')
                    elif heading_level == 3:  # ###
                        html_lines.append(
                            f'<h4 style="font-size: 13pt; font-weight: bold; margin: 12px 0 8px 0;">{heading_text}</h4>')
                    elif heading_level == 4:  # ####
                        html_lines.append(
                            f'<h5 style="font-size: 12pt; font-weight: bold; margin: 10px 0 6px 0;">{heading_text}</h5>')
                    else:
                        html_lines.append(
                            f'<h6 style="font-size: 11pt; font-weight: bold; margin: 8px 0 4px 0;">{heading_text}</h6>')

                    i += 1
                    continue

                # Gestisci i sottotitoli con underline (---)
                if i + 1 < len(lines) and all(c == '-' for c in lines[i + 1].strip()):
                    html_lines.append(
                        f'<h3 style="font-size: 14pt; font-weight: bold; margin: 15px 0 10px 0;">{line}</h3>')
                    i += 2  # Skip the title and the '---' line
                    continue

                # Gestisci le liste
                if line.startswith('- '):
                    # Check if this is a list item with an image reference
                    if '[IMMAGINE' in line and ']' in line:
                        # Process the image reference directly
                        processed_line = replace_image(line)
                        if processed_line:  # If the image was processed successfully
                            html_lines.append(processed_line)
                        i += 1
                        continue

                    if not in_list:
                        in_list = True
                        list_items = []

                    list_items.append(line)
                    i += 1
                    continue

                # Gestisci le tabelle markdown
                if line.startswith('|') and line.endswith('|'):
                    # Collect table lines
                    table_lines = [line]
                    j = i + 1
                    while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                        table_lines.append(lines[j].strip())
                        j += 1

                    # Process table if we have at least one row
                    if len(table_lines) >= 1:
                        html_table = process_markdown_table(table_lines)
                        html_lines.append(html_table)

                    i = j  # Skip processed table lines
                    continue

                # Gestisci le immagini in formato [IMMAGINE ...]
                if line.startswith('[IMMAGINE') and ']' in line:
                    # Process the entire image reference
                    processed_line = replace_image(line)
                    if processed_line:  # If the image was processed successfully
                        html_lines.append(processed_line)
                    i += 1
                    continue

                # Gestisci le righe normali di testo
                if line:
                    # Apply text formatting and ensure it's not bold by default
                    formatted_line = format_text(line)
                    html_lines.append(f'<div style="margin: 5px 0; font-weight: normal;">{formatted_line}</div>')

                i += 1

            # Process any remaining list items
            if list_items:
                html_lines.append(process_list_items(list_items))

            return '\n'.join(html_lines)

        # Converti il testo in HTML
        html_content = convert_to_html(text)

        return f'''
        <div style='font-family: Arial; font-size: 11pt; line-height: 1.5; padding: 10px;'>
            {html_content}
        </div>
        '''

    def run(self):
        self.formatted_report = ""
        try:
            self.log_message.emit("Starting report generation...", "info")
            self.archaeological_analysis = ArchaeologicalAnalysis()

            # Funzione helper per convertire i dati in modo sicuro
            def safe_convert_data(data):
                if data is None:
                    return {}
                if isinstance(data, tuple):
                    # Se è una tupla di dizionari
                    if data and isinstance(data[0], dict):
                        return data[0]
                    # Se è una tupla di tuple
                    elif data and isinstance(data[0], tuple):
                        try:
                            return dict(data)
                        except:
                            return {}
                    else:
                        return {}
                elif isinstance(data, dict):
                    return data
                else:
                    return {}

            def safe_convert_list(data):
                if data is None:
                    return []
                if isinstance(data, tuple):
                    return list(data)
                elif isinstance(data, list):
                    return data
                else:
                    return []

            # Converti i dati in modo sicuro
            try:
                self.site_data = safe_convert_data(self.site_data)
                self.us_data = safe_convert_list(self.us_data)
                self.materials_data = safe_convert_list(self.materials_data)
                self.pottery_data = safe_convert_list(self.pottery_data)

                self.log_message.emit("Data conversion successful", "info")
                self.log_message.emit(f"Site data type: {type(self.site_data)}", "info")
                self.log_message.emit(f"US data length: {len(self.us_data)}", "info")
                self.log_message.emit(f"Materials data length: {len(self.materials_data)}", "info")
                self.log_message.emit(f"Pottery data length: {len(self.pottery_data)}", "info")

            except Exception as conv_error:
                self.log_message.emit(f"Warning: Error converting data: {str(conv_error)}", "warning")
                # Inizializza con valori di default in caso di errore
                self.site_data = {}
                self.us_data = []
                self.materials_data = []
                self.pottery_data = []

            while True:
                step = self.archaeological_analysis.get_next_analysis_step()
                if not step:
                    break

                self.log_message.emit(f"Processing {step['section']}...", "step")

                # Gestione sicura dei required_tables
                required_tables = step.get('required_table', [])
                if required_tables is None:
                    required_tables = []
                elif isinstance(required_tables, str):
                    required_tables = [required_tables]

                # Verifica se tutte le tabelle richieste sono state selezionate
                missing_tables = [table for table in required_tables if table not in self.selected_tables]

                # Special case for CATALOGO DEI MATERIALI: check if at least one of pottery_table or inventario_materiali_table is selected
                if step['section'] == "CATALOGO DEI MATERIALI":
                    if "pottery_table" in self.selected_tables or "inventario_materiali_table" in self.selected_tables:
                        # At least one of the required tables is selected, so we can proceed
                        missing_tables = [table for table in missing_tables if table != "pottery_table" and table != "inventario_materiali_table"]

                if missing_tables and step['section'] != "CONCLUSIONI":
                    self.log_message.emit(
                        f"Warning: Missing required tables for {step['section']}: {', '.join(missing_tables)}",
                        "warning"
                    )
                    # Skip this section if required tables are missing, except for CONCLUSIONI
                    self.log_message.emit(f"Skipping section {step['section']} due to missing required tables", "warning")
                    continue

                try:
                    # Prepare context for this step
                    context = {
                        "site_data": self.site_data if "site_table" in required_tables else {},
                        "us_data": self.us_data if "us_table" in required_tables else [],
                        "materials_data": self.materials_data if "inventario_materiali_table" in required_tables else [],
                        "pottery_data": self.pottery_data if "pottery_table" in required_tables else [],
                        "tomba_data": self.tomba_data if "tomba_table" in required_tables else [],
                        "periodizzazione_data": self.periodizzazione_data if "periodizzazione_table" in required_tables else [],
                        "struttura_data": self.struttura_data if "struttura_table" in required_tables else []
                    }

                    # Log validation requirements
                    validation_tools = step.get('validation_tool', [])
                    if validation_tools:
                        if isinstance(validation_tools, str):
                            validation_tools = [validation_tools]

                        # Filter validation tools based on selected tables
                        filtered_validation_tools = []
                        for tool in validation_tools:
                            # Map validation tools to their required tables
                            tool_table_map = {
                                "validate_site_info": "site_table",
                                "validate_us": "us_table",
                                "validate_materials": "inventario_materiali_table",
                                "validate_pottery": "pottery_table",
                                "validate_tomba": "tomba_table",
                                "validate_periodizzazione": "periodizzazione_table",
                                "validate_struttura": "struttura_table"
                            }

                            required_table = tool_table_map.get(tool)
                            if required_table is None or required_table in self.selected_tables:
                                filtered_validation_tools.append(tool)
                            else:
                                self.log_message.emit(f"Skipping validation {tool} as table {required_table} is not selected", "info")

                        if not filtered_validation_tools:
                            self.log_message.emit(f"No validations to perform for {step['section']} with selected tables", "info")
                            continue

                        self.log_message.emit(f"Validation requirements for {step['section']}:", "warning")
                        for tool in filtered_validation_tools:
                            self.log_message.emit(f"- Requires validation: {tool}", "warning")

                        # Esegui le validazioni
                        self.log_message.emit(f"\nExecuting validations for {step['section']}...", "info")
                        for tool in filtered_validation_tools:
                            try:
                                validator = getattr(ArchaeologicalValidators, tool, None)

                                if validator:
                                    validation_result = validator(context)

                                    # Log del risultato della validazione
                                    if validation_result.get('valid', validation_result.get('success', False)):
                                        self.log_message.emit(f"✓ {tool}: All data validated successfully", "success")
                                    else:
                                        self.log_message.emit(f"⚠ {tool}: Validation issues found", "warning")

                                        # Log dei campi mancanti
                                        missing_fields = validation_result.get('missing_fields', [])
                                        if missing_fields:
                                            self.log_message.emit("Missing required fields:", "warning")
                                            for field in missing_fields:
                                                self.log_message.emit(f"  - {field}", "warning")

                                        # Log dei dati incompleti
                                        incomplete = validation_result.get('incomplete', [])
                                        if incomplete:
                                            self.log_message.emit("Incomplete data found:", "warning")
                                            for item in incomplete:
                                                self.log_message.emit(f"  - {item}", "warning")

                                        # Log del messaggio dettagliato
                                        message = validation_result.get('message', '')
                                        if message:
                                            self.log_message.emit("\nDetailed validation message:", "info")
                                            self.log_message.emit(message, "warning")
                                else:
                                    self.log_message.emit(f"⚠ Validation tool '{tool}' not found", "warning")

                            except Exception as validation_error:
                                self.log_message.emit(f"Error during validation with {tool}: {str(validation_error)}",
                                                      "error")
                                continue

                finally:
                        self.log_message.emit("Validation complete. Proceeding with report generation...\n", "info")

                # This second check is redundant and can be removed since we already checked above
                # and skipped the section if required tables are missing

                try:
                    # context = {
                    #     "site_data": self.site_data if "site_table" in required_tables else None,  # Aggiunto site_data
                    #     "us_data": self.us_data if "us_table" in required_tables else None,
                    #     "materials_data": self.materials_data if "inventario_materiali_table" in required_tables else None,
                    #     "pottery_data": self.pottery_data if "pottery_table" in required_tables else None
                    # }

                    # Create data summary for the prompt
                    data_summary = "Available data:\n"
                    for table_name, data in context.items():
                        if data:
                            data_summary += f"- {table_name}: {len(data)} records\n"
                            if len(data) > 0:
                                data_summary += f"  Sample fields: {list(data[0].keys())[:5]}\n"

                    # Process images only for appropriate sections
                    should_process_images = step['section'] not in ['INTRODUZIONE', 'METODOLOGIA DI SCAVO']

                    entity_ids = []
                    if should_process_images and "us_table" in required_tables and context["us_data"]:
                        self.log_message.emit(f"Processing images for US data...", "info")
                        id_list = [str(us.get('id_us')) for us in context["us_data"] if us.get('id_us')]
                        entity_ids.extend(id_list)

                    # Recupero delle immagini
                    dialog = self.py_dialog
                    images = []
                    image_context = {}

                    # Collect entity IDs for materials and pottery if needed
                    material_ids = []
                    pottery_ids = []

                    if should_process_images:
                        # Get material IDs if materials data is available
                        if "materials_data" in context and context["materials_data"]:
                            self.log_message.emit(f"Processing images for materials data...", "info")
                            material_ids = [str(mat.get('id_invmat')) for mat in context["materials_data"] if mat.get('id_invmat')]
                            self.log_message.emit(f"Collected {len(material_ids)} material IDs", "info")

                        # Get pottery IDs if pottery data is available
                        if "pottery_data" in context and context["pottery_data"]:
                            self.log_message.emit(f"Processing images for pottery data...", "info")
                            pottery_ids = [str(pot.get('id_rep')) for pot in context["pottery_data"] if pot.get('id_rep')]
                            self.log_message.emit(f"Collected {len(pottery_ids)} pottery IDs", "info")

                    # Get US images
                    if should_process_images and entity_ids:
                        try:
                            us_images = dialog.get_images_for_entities(entity_ids, self.log_message, 'US')
                            self.log_message.emit(f"Retrieved {len(us_images)} US images", "info")
                            images.extend(us_images)

                            for img in us_images:
                                us_id = img['id']
                                if us_id not in image_context:
                                    image_context[us_id] = []
                                image_context[us_id].append({
                                    'url': img['url'],
                                    'caption': img['caption'],
                                    'entity_type': 'US'
                                })
                        except Exception as e:
                            self.log_message.emit(f"Error retrieving US images: {str(e)}", "warning")

                    # Get material images
                    if should_process_images and material_ids:
                        try:
                            material_images = dialog.get_images_for_entities(material_ids, self.log_message, 'REPERTO')
                            self.log_message.emit(f"Retrieved {len(material_images)} material images", "info")
                            images.extend(material_images)

                            for img in material_images:
                                material_id = img['id']
                                if material_id not in image_context:
                                    image_context[material_id] = []
                                image_context[material_id].append({
                                    'url': img['url'],
                                    'caption': img['caption'],
                                    'entity_type': 'REPERTO'
                                })
                        except Exception as e:
                            self.log_message.emit(f"Error retrieving material images: {str(e)}", "warning")

                    # Get pottery images
                    if should_process_images and pottery_ids:
                        try:
                            pottery_images = dialog.get_images_for_entities(pottery_ids, self.log_message, 'CERAMICA')
                            self.log_message.emit(f"Retrieved {len(pottery_images)} pottery images", "info")
                            images.extend(pottery_images)

                            for img in pottery_images:
                                pottery_id = img['id']
                                if pottery_id not in image_context:
                                    image_context[pottery_id] = []
                                image_context[pottery_id].append({
                                    'url': img['url'],
                                    'caption': img['caption'],
                                    'entity_type': 'CERAMICA'
                                })
                        except Exception as e:
                            self.log_message.emit(f"Error retrieving pottery images: {str(e)}", "warning")

                    # Create base prompt
                    base_prompt = f"""Thought: {step.get('thought', 'Analyzing archaeological data')}
                    Action: {step.get('action', 'Generate report section')}
                    Section: {step['section']}

                    Observation:
                    Required validations: {', '.join(validation_tools) if validation_tools else 'None'}

                    Data Summary:
                    {data_summary}

                    Instructions:
                    {step['prompt']}

                    Additional Context:
                    {self.descriptions_text}

                    Language: {self.create_prompt(self.output_language)}"""

                    # Add section-specific data
                    if step['section'] == 'INTRODUZIONE':
                        intro_data = "\nDati disponibili per l'introduzione:\n"

                        # Add US summary
                        if context.get("us_data"):
                            total_us = len(context["us_data"])
                            intro_data += f"\nDati stratigrafici:\n- Totale US: {total_us}\n"
                            period_data = [us.get('periodo') for us in context["us_data"] if us.get('periodo')]
                            if period_data:
                                intro_data += f"- Periodi presenti: {', '.join(set(period_data))}\n"

                        # Add materials summary
                        if context.get("materials_data"):
                            total_materials = len(context["materials_data"])
                            intro_data += f"\nMateriali:\n- Totale reperti: {total_materials}\n"
                            material_types = [mat.get('tipo_materiale') for mat in context["materials_data"] if
                                              mat.get('tipo_materiale')]
                            if material_types:
                                intro_data += f"- Tipologie: {', '.join(set(material_types))}\n"

                        # Add pottery summary
                        if context.get("pottery_data"):
                            total_pottery = len(context["pottery_data"])
                            intro_data += f"\nCeramica:\n- Totale ceramica: {total_pottery}\n"
                            pottery_types = [pot.get('tipo') for pot in context["pottery_data"] if pot.get('tipo')]
                            if pottery_types:
                                intro_data += f"- Tipologie ceramiche: {', '.join(set(pottery_types))}\n"

                        base_prompt += intro_data

                    # Add image instructions if necessary
                    if should_process_images and image_context:
                        image_instructions = "\nImmagini disponibili:\n"

                        # Group images by entity type
                        us_images = {}
                        material_images = {}
                        pottery_images = {}

                        for entity_id, images_list in image_context.items():
                            for img in images_list:
                                entity_type = img.get('entity_type', 'US')  # Default to US if not specified

                                if entity_type == 'US':
                                    if entity_id not in us_images:
                                        us_images[entity_id] = []
                                    us_images[entity_id].append(img)
                                elif entity_type == 'REPERTO':
                                    if entity_id not in material_images:
                                        material_images[entity_id] = []
                                    material_images[entity_id].append(img)
                                elif entity_type == 'CERAMICA':
                                    if entity_id not in pottery_images:
                                        pottery_images[entity_id] = []
                                    pottery_images[entity_id].append(img)

                        # Add US images
                        if us_images:
                            image_instructions += "\nImmagini delle Unità Stratigrafiche:\n"
                            for us_id, images_list in us_images.items():
                                image_instructions += f"\nUS {us_id}:\n"
                                for img in images_list:
                                    image_instructions += f"[IMMAGINE US {us_id}: {img['url']}, {img['caption']}]\n"

                        # Add material images
                        if material_images:
                            image_instructions += "\nImmagini dei Materiali:\n"
                            for material_id, images_list in material_images.items():
                                image_instructions += f"\nMateriale {material_id}:\n"
                                for img in images_list:
                                    image_instructions += f"[IMMAGINE REPERTO {material_id}: {img['url']}, {img['caption']}]\n"

                        # Add pottery images
                        if pottery_images:
                            image_instructions += "\nImmagini della Ceramica:\n"
                            for pottery_id, images_list in pottery_images.items():
                                image_instructions += f"\nCeramica {pottery_id}:\n"
                                for img in images_list:
                                    image_instructions += f"[IMMAGINE CERAMICA {pottery_id}: {img['url']}, {img['caption']}]\n"

                        image_instructions += """
                        Per favore, quando menzioni un'entità che ha immagini associate, inserisci l'immagine nel testo usando questa sintassi:
                        - Per le US: [IMMAGINE US numero: percorso, caption]
                        - Per i Materiali: [IMMAGINE REPERTO numero: percorso, caption]
                        - Per la Ceramica: [IMMAGINE CERAMICA numero: percorso, caption]

                        Inserisci le immagini nei punti appropriati del testo, quando menzioni l'entità corrispondente.
                        """
                        base_prompt += image_instructions

                    # Generate the content using chunking for large datasets
                    # Check if we need to use chunking based on data size
                    use_chunking = False
                    large_data_tables = []
                    very_large_data_tables = []

                    # Check which tables have a lot of data
                    for table_name, data in context.items():
                        if isinstance(data, list):
                            if len(data) > 100:  # Threshold for "very large" data
                                use_chunking = True
                                large_data_tables.append(table_name)
                                very_large_data_tables.append(table_name)
                            elif len(data) > 50:  # Threshold for "large" data
                                use_chunking = True
                                large_data_tables.append(table_name)

                    if use_chunking and large_data_tables:
                        self.log_message.emit(f"Using RAG approach for large datasets: {', '.join(large_data_tables)}", "info")

                        # Initialize LLM for RAG
                        llm = ChatOpenAI(
                            temperature=0.0,
                            model_name=self.selected_model,
                            api_key=self.api_key,
                            max_tokens=4000
                        )

                        # Create vector databases for large tables
                        vector_stores = {}
                        for table_name in large_data_tables:
                            self.log_message.emit(f"Creating vector database for {table_name}...", "info")
                            data = context[table_name]
                            vector_store = self.create_vector_db(data, table_name)
                            if vector_store:
                                vector_stores[table_name] = vector_store
                                self.log_message.emit(f"Vector database created for {table_name}", "info")
                            else:
                                self.log_message.emit(f"Failed to create vector database for {table_name}", "warning")

                        # Process with RAG approach
                        section_results = []

                        # First, process with a summary of all data to get an overview
                        self.log_message.emit("Generating overview...", "info")
                        summary_prompt = base_prompt + "\n\nPer questa prima fase, fornisci una panoramica generale basata sui dati disponibili."

                        # Add a note about where the vector database is stored
                        self.log_message.emit("Vector databases are stored in memory during the report generation process", "info")

                        # For the overview, we'll use a simplified context with just a few examples from each large table
                        simplified_context = context.copy()
                        for table_name in large_data_tables:
                            if table_name in context and isinstance(context[table_name], list) and len(context[table_name]) > 5:
                                simplified_context[table_name] = context[table_name][:5]  # Just take the first 5 items

                        # Create a simplified prompt with the reduced context
                        simplified_data_summary = "Processing with simplified dataset for overview.\n"
                        for table_name, data in simplified_context.items():
                            if isinstance(data, list):
                                simplified_data_summary += f"- {table_name}: {len(data)} records (sample)\n"

                        simplified_prompt = base_prompt.replace(data_summary, simplified_data_summary)
                        overview_result = self.agent.run(simplified_prompt)
                        if overview_result:
                            section_results.append(overview_result)

                        # Then process each large table using RAG
                        for table_name in large_data_tables:
                            if table_name not in vector_stores:
                                self.log_message.emit(f"Skipping {table_name} as vector database creation failed", "warning")
                                continue

                            vector_store = vector_stores[table_name]
                            self.log_message.emit(f"Processing {table_name} with RAG approach...", "info")

                            # Create analysis questions based on the section
                            analysis_questions = [
                                f"What are the key patterns or trends in the {table_name} data?",
                                f"What are the most significant findings in the {table_name} data?",
                                f"How does the {table_name} data relate to the archaeological context?",
                                f"What chronological information can be derived from the {table_name} data?",
                                f"What spatial distribution patterns are evident in the {table_name} data?"
                            ]

                            # Add image-related questions if images are available
                            if should_process_images and image_context:
                                # Find relevant entity IDs for this table
                                relevant_ids = []
                                if table_name == "us_data":
                                    relevant_ids = [str(us.get('id_us')) for us in context["us_data"] if us.get('id_us')]
                                elif table_name == "materials_data":
                                    relevant_ids = [str(mat.get('id_invmat')) for mat in context["materials_data"] if mat.get('id_invmat')]
                                elif table_name == "pottery_data":
                                    relevant_ids = [str(pot.get('id_rep')) for pot in context["pottery_data"] if pot.get('id_rep')]

                                # Filter image_context to only include relevant entities
                                relevant_images = {entity_id: imgs for entity_id, imgs in image_context.items() if entity_id in relevant_ids}

                                if relevant_images:
                                    # Add image-specific questions
                                    analysis_questions.append(f"Describe the visual characteristics of the {table_name} based on the available images.")
                                    analysis_questions.append(f"What additional information can be derived from the images of the {table_name}?")

                            # Create RAG chain for this table
                            # Pass image_context to include images in the RAG approach
                            rag_chain = self.create_rag_chain(vector_store, llm)
                            if not rag_chain:
                                self.log_message.emit(f"Failed to create RAG chain for {table_name}", "warning")
                                continue

                            # Add image instructions to the RAG chain's prompt if images are available
                            if should_process_images and image_context:
                                self.log_message.emit(f"Including images in RAG approach for {table_name}...", "info")

                            # Process each analysis question
                            table_results = []
                            for i, question in enumerate(analysis_questions):
                                self.log_message.emit(f"Processing question {i+1}/{len(analysis_questions)} for {table_name}...", "info")
                                try:
                                    # Prepare the question with image instructions if available
                                    enhanced_question = question
                                    if should_process_images and image_context and i >= len(analysis_questions) - 2:  # For image-specific questions
                                        # Find relevant entity IDs for this table
                                        relevant_ids = []
                                        if table_name == "us_data":
                                            relevant_ids = [str(us.get('id_us')) for us in context["us_data"] if us.get('id_us')]
                                        elif table_name == "materials_data":
                                            relevant_ids = [str(mat.get('id_invmat')) for mat in context["materials_data"] if mat.get('id_invmat')]
                                        elif table_name == "pottery_data":
                                            relevant_ids = [str(pot.get('id_rep')) for pot in context["pottery_data"] if pot.get('id_rep')]

                                        # Filter image_context to only include relevant entities
                                        relevant_images = {entity_id: imgs for entity_id, imgs in image_context.items() if entity_id in relevant_ids}

                                        if relevant_images:
                                            # Add image instructions to the question
                                            image_instructions = "\n\nImmagini disponibili:\n"

                                            for entity_id, images_list in relevant_images.items():
                                                for img in images_list:
                                                    entity_type = img.get('entity_type', 'US')
                                                    if entity_type == 'US':
                                                        image_instructions += f"[IMMAGINE US {entity_id}: {img['url']}, {img['caption']}]\n"
                                                    elif entity_type == 'REPERTO':
                                                        image_instructions += f"[IMMAGINE REPERTO {entity_id}: {img['url']}, {img['caption']}]\n"
                                                    elif entity_type == 'CERAMICA':
                                                        image_instructions += f"[IMMAGINE CERAMICA {entity_id}: {img['url']}, {img['caption']}]\n"

                                            image_instructions += """
                                            Per favore, quando menzioni un'entità che ha immagini associate, inserisci l'immagine nel testo usando questa sintassi:
                                            - Per le US: [IMMAGINE US numero: percorso, caption]
                                            - Per i Materiali: [IMMAGINE REPERTO numero: percorso, caption]
                                            - Per la Ceramica: [IMMAGINE CERAMICA numero: percorso, caption]

                                            Inserisci le immagini nei punti appropriati del testo, quando menzioni l'entità corrispondente.
                                            """

                                            enhanced_question = f"{question}\n\n{image_instructions}"

                                    # Run the RAG chain with the enhanced question
                                    response = rag_chain.run(enhanced_question)
                                    if response:
                                        table_results.append(f"Analysis {i+1}: {question}\n{response}")
                                except Exception as e:
                                    self.log_message.emit(f"Error processing question {i+1} for {table_name}: {str(e)}", "warning")

                            # Combine results for this table
                            if table_results:
                                table_analysis = f"\n\n### Analysis of {table_name} data:\n\n" + "\n\n".join(table_results)
                                section_results.append(table_analysis)

                        # Final integration using RAG
                        if len(section_results) > 1:
                            self.log_message.emit("Integrating all analyses...", "info")

                            # Create a prompt for integration
                            integration_text = "\n\n".join(section_results)
                            integration_prompt = f"""
                            You are an archaeological expert tasked with integrating multiple analyses into a coherent report.

                            Below are separate analyses of different aspects of the archaeological data:

                            {integration_text}

                            Please integrate these analyses into a coherent, well-structured report section.
                            Eliminate repetitions, organize the information logically, and ensure a smooth flow between topics.
                            Focus on the most significant findings and patterns across all the data.

                            IMPORTANT: Preserve all image references in the format [IMMAGINE US/REPERTO/CERAMICA numero: percorso, caption].
                            These are essential for the report and must be included exactly as they appear in the analyses.
                            """

                            # Use a direct call to the model with a reduced token count to avoid exceeding limits
                            try:
                                client = OpenAI(api_key=self.api_key)
                                response = client.chat.completions.create(
                                    model=self.selected_model,
                                    messages=[{"role": "system", "content": "You are an archaeological expert."},
                                              {"role": "user", "content": integration_prompt}],
                                    max_tokens=4000,
                                    temperature=0.0
                                )
                                final_result = response.choices[0].message.content
                                if final_result:
                                    result = final_result
                                else:
                                    # If integration fails, concatenate all results
                                    result = "\n\n".join(section_results)
                            except Exception as e:
                                self.log_message.emit(f"Error during integration: {str(e)}", "warning")
                                # If integration fails, concatenate all results
                                result = "\n\n".join(section_results)
                        else:
                            result = section_results[0] if section_results else ""
                    else:
                        # For smaller datasets, process normally
                        result = self.agent.run(base_prompt)

                    if result:
                        # Post-process to remove AI's notes and thoughts
                        result = self.clean_ai_notes(result)

                        self.log_message.emit("Formattazione del risultato...", "info")
                        section_text = f"{step['section']}\n{'=' * len(step['section'])}\n{result}"
                        formatted_section = self.format_for_widget(section_text)

                        if self.formatted_report:
                            self.formatted_report += "<br><br>"
                        self.formatted_report += formatted_section
                        self.report_generated.emit(self.formatted_report)
                        self.full_report += f"\n\n{section_text}"
                        self.log_message.emit(f"Completed {step['section']}", "step")
                    else:
                        self.log_message.emit("Nessun risultato generato", "warning")

                except Exception as section_error:
                    error_message = str(section_error)
                    self.log_message.emit(
                        f"Warning: Error processing section {step['section']}: {error_message}. Attempting to continue with available data.",
                        "warning"
                    )

                    # Check if it's a token limit error
                    if "context_length_exceeded" in error_message or "max_tokens" in error_message:
                        # Try to process with RAG approach
                        try:
                            self.log_message.emit("Retrying with RAG approach due to token limit...", "info")
                            self.log_message.emit("Vector databases are stored in memory during the report generation process", "info")

                            # Initialize LLM for RAG with reduced tokens
                            llm = ChatOpenAI(
                                temperature=0.0,
                                model_name=self.selected_model,
                                api_key=self.api_key,
                                max_tokens=2000  # Reduced token count for safety
                            )

                            # Identify large tables that need RAG
                            rag_tables = []
                            for table_name, data in context.items():
                                if isinstance(data, list) and len(data) > 20:  # Lower threshold for retry
                                    rag_tables.append(table_name)

                            if not rag_tables:
                                # If no large tables, just take a small sample of all data
                                self.log_message.emit("No large tables identified, using simplified approach", "info")

                                # Create a simplified context with reduced data
                                simplified_context = {}
                                for table_name, data in context.items():
                                    if isinstance(data, list) and len(data) > 5:
                                        # Take a very small representative sample
                                        simplified_context[table_name] = data[:5]  # Just take the first 5 items
                                    else:
                                        simplified_context[table_name] = data

                                # Create a simplified prompt
                                simplified_prompt = base_prompt.replace(data_summary, "Processing with reduced dataset due to token limits.\n")
                                simplified_prompt += "\n\nAnalizza i dati disponibili e fornisci un'analisi generale. Indica chiaramente che l'analisi è basata su un campione ridotto dei dati."

                                # Process with simplified data
                                result = self.agent.run(simplified_prompt)
                            else:
                                # Use RAG approach for large tables
                                self.log_message.emit(f"Using RAG approach for tables: {', '.join(rag_tables)}", "info")

                                # Create vector databases for large tables
                                vector_stores = {}
                                for table_name in rag_tables:
                                    self.log_message.emit(f"Creating vector database for {table_name}...", "info")
                                    data = context[table_name]
                                    vector_store = self.create_vector_db(data, table_name)
                                    if vector_store:
                                        vector_stores[table_name] = vector_store
                                        self.log_message.emit(f"Vector database created for {table_name}", "info")
                                    else:
                                        self.log_message.emit(f"Failed to create vector database for {table_name}", "warning")

                                # Process with RAG approach
                                section_results = []

                                # Create a simplified context for non-RAG tables
                                simplified_context = context.copy()
                                for table_name in rag_tables:
                                    if table_name in simplified_context:
                                        # Replace with a minimal sample
                                        if isinstance(simplified_context[table_name], list) and len(simplified_context[table_name]) > 3:
                                            simplified_context[table_name] = simplified_context[table_name][:3]

                                # Generate a brief overview with simplified context
                                self.log_message.emit("Generating brief overview...", "info")
                                overview_prompt = f"""
                                You are an archaeological expert analyzing data from an excavation.

                                Please provide a very brief overview of the archaeological context based on the following information:

                                {step['prompt']}

                                Keep your response concise and focused on the main points only.
                                """

                                # Use direct API call with reduced tokens
                                client = OpenAI(api_key=self.api_key)
                                response = client.chat.completions.create(
                                    model=self.selected_model,
                                    messages=[{"role": "system", "content": "You are an archaeological expert."},
                                              {"role": "user", "content": overview_prompt}],
                                    max_tokens=1000,
                                    temperature=0.0
                                )
                                overview_result = response.choices[0].message.content
                                if overview_result:
                                    section_results.append(overview_result)

                                # Process each RAG table
                                for table_name in rag_tables:
                                    if table_name not in vector_stores:
                                        continue

                                    vector_store = vector_stores[table_name]
                                    self.log_message.emit(f"Processing {table_name} with RAG approach...", "info")

                                    # Create focused analysis questions
                                    analysis_questions = [
                                        f"What are the key patterns in the {table_name} data?",
                                        f"What are the most significant findings in the {table_name} data?"
                                    ]

                                    # Create RAG chain
                                    rag_chain = self.create_rag_chain(vector_store, llm)
                                    if not rag_chain:
                                        continue

                                    # Process questions
                                    table_results = []
                                    for i, question in enumerate(analysis_questions):
                                        try:
                                            # Prepare the question with image instructions if available
                                            enhanced_question = question
                                            if should_process_images and image_context:
                                                # Find relevant entity IDs for this table
                                                relevant_ids = []
                                                if table_name == "us_data":
                                                    relevant_ids = [str(us.get('id_us')) for us in context["us_data"] if us.get('id_us')]
                                                elif table_name == "materials_data":
                                                    relevant_ids = [str(mat.get('id_invmat')) for mat in context["materials_data"] if mat.get('id_invmat')]
                                                elif table_name == "pottery_data":
                                                    relevant_ids = [str(pot.get('id_rep')) for pot in context["pottery_data"] if pot.get('id_rep')]

                                                # Filter image_context to only include relevant entities
                                                relevant_images = {entity_id: imgs for entity_id, imgs in image_context.items() if entity_id in relevant_ids}

                                                if relevant_images:
                                                    # Add image instructions to the question
                                                    image_instructions = "\n\nImmagini disponibili:\n"

                                                    for entity_id, images_list in relevant_images.items():
                                                        for img in images_list:
                                                            entity_type = img.get('entity_type', 'US')
                                                            if entity_type == 'US':
                                                                image_instructions += f"[IMMAGINE US {entity_id}: {img['url']}, {img['caption']}]\n"
                                                            elif entity_type == 'REPERTO':
                                                                image_instructions += f"[IMMAGINE REPERTO {entity_id}: {img['url']}, {img['caption']}]\n"
                                                            elif entity_type == 'CERAMICA':
                                                                image_instructions += f"[IMMAGINE CERAMICA {entity_id}: {img['url']}, {img['caption']}]\n"

                                                    image_instructions += """
                                                    Per favore, quando menzioni un'entità che ha immagini associate, inserisci l'immagine nel testo usando questa sintassi:
                                                    - Per le US: [IMMAGINE US numero: percorso, caption]
                                                    - Per i Materiali: [IMMAGINE REPERTO numero: percorso, caption]
                                                    - Per la Ceramica: [IMMAGINE CERAMICA numero: percorso, caption]

                                                    Inserisci le immagini nei punti appropriati del testo, quando menzioni l'entità corrispondente.
                                                    """

                                                    enhanced_question = f"{question}\n\n{image_instructions}"
                                                    self.log_message.emit(f"Including images in error recovery RAG approach for {table_name}...", "info")

                                            response = rag_chain.run(enhanced_question)
                                            if response:
                                                table_results.append(response)
                                        except Exception as e:
                                            self.log_message.emit(f"Error in RAG processing: {str(e)}", "warning")

                                    # Combine results
                                    if table_results:
                                        table_analysis = f"\n\n### Analysis of {table_name} data:\n\n" + "\n\n".join(table_results)
                                        section_results.append(table_analysis)

                                # Combine all results
                                result = "\n\n".join(section_results)

                            if result:
                                # Post-process to remove AI's notes and thoughts
                                result = self.clean_ai_notes(result)

                                # Add a note about the RAG approach
                                result = "NOTA: Questa sezione è stata generata utilizzando un approccio di Retrieval Augmented Generation (RAG) a causa di limiti tecnici nella dimensione dei dati.\n\n" + result

                                self.log_message.emit("Formattazione del risultato con approccio RAG...", "info")

                                # Ensure image references are preserved in the final result
                                if should_process_images and image_context:
                                    self.log_message.emit("Ensuring image references are preserved in the final result...", "info")
                                section_text = f"{step['section']}\n{'=' * len(step['section'])}\n{result}"
                                formatted_section = self.format_for_widget(section_text)

                                if self.formatted_report:
                                    self.formatted_report += "<br><br>"
                                self.formatted_report += formatted_section
                                self.report_generated.emit(self.formatted_report)
                                self.full_report += f"\n\n{section_text}"
                                self.log_message.emit(f"Completed {step['section']} with RAG approach", "step")
                            else:
                                self.log_message.emit("Failed to generate result with RAG approach", "warning")
                        except Exception as retry_error:
                            self.log_message.emit(f"Failed retry attempt: {str(retry_error)}. Continuing with next section.", "warning")
                            continue
                    else:
                        # For other types of errors, continue with the next section
                        continue

            # Prepara i dati finali del report
            report_data = {
                'report_text': self.full_report,
                'materials_table': self.format_materials_table() if hasattr(self, 'format_materials_table') else None,
                'pottery_table': self.format_pottery_table() if hasattr(self, 'format_pottery_table') else None,
                'tomba_table': self.format_tomba_table() if hasattr(self, 'format_tomba_table') else None,
                'periodizzazione_table': self.format_periodizzazione_table() if hasattr(self, 'format_periodizzazione_table') else None,
                'struttura_table': self.format_struttura_table() if hasattr(self, 'format_struttura_table') else None
            }

            self.report_completed.emit(self.full_report, report_data)
            self.log_message.emit("Report generation completed!", "validation")

        except Exception as e:
            self.log_message.emit(f"Error during report generation: {str(e)}", "error")
            raise
    def clean_ai_notes(self, text):
        """Remove AI's notes and thoughts from the text"""
        if not text:
            return text

        # Remove notes about complete lists being available
        text = re.sub(r'--\*?Nota: L\'elenco completo .*? disponibile.*?\*?', '', text)
        text = re.sub(r'--Nota: L\'elenco completo .*? può essere fornito.*?\.', '', text)

        # Remove continuation notes
        text = re.sub(r'\.\.\. \(continua per tutte le fasi.*?\)', '', text)
        text = re.sub(r'\.\.\. \(continua per tutte le US.*?\)', '', text)

        # Remove any other AI thoughts or notes
        text = re.sub(r'--\*?Nota:.*?\*?', '', text)

        # Remove common AI response phrases
        text = re.sub(r'^Certainly!.*?integrated,?\s+', '', text)
        text = re.sub(r'^Here is an integrated,?\s+', '', text)
        text = re.sub(r'^I\'ll provide an integrated,?\s+', '', text)
        text = re.sub(r'\*\(All image references are preserved.*?layout\.\)\*', '', text)
        text = re.sub(r'\(All image references are preserved.*?layout\.\)', '', text)
        text = re.sub(r'\*\(Please insert the relevant image references.*?data\.\)\*', '', text)
        text = re.sub(r'\(Please insert the relevant image references.*?data\.\)', '', text)
        text = re.sub(r'\[IMMAGINE US/REPERTO/CERAMICA numero: percorso, caption\]', '', text)

        # Clean up any double spaces or newlines created by the removals
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'  +', ' ', text)

        return text

    def create_prompt(self, selected_language):


        language_instructions = {
            'Italiano': "Rispondi in italiano.",
            'English (UK)': "Reply in British English.",
            'English (US)': "Reply in American English.",
            'Español': "Responde en español.",
            'Français': "Réponds en français.",
            'Deutsch': "Antworte auf Deutsch.",
            'العربية': "أجب باللغة العربية. Reply in Arabic language.",
            'Ελληνικά': "Απαντήστε στα ελληνικά.",
            'Русский': "Ответьте на русском языке.",
            'Português': "Responda em português."
        }

        # Aggiungi l'istruzione della lingua alla prompt base
        full_prompt = f"{language_instructions[selected_language]}"
        return full_prompt

    def format_materials_table(self):
        """Formatta i dati dei materiali per la tabella"""
        table_data = []
        if self.materials_data:
            table_data.append(['ID','Inv. Number', 'Type', 'Definition', 'Description', 'Dating', 'Conservation'])
            for material in self.materials_data:
                table_data.append([
                    material.get('id_invmat', ''),
                    material.get('numero_inventario', ''),
                    material.get('tipo_reperto', ''),
                    material.get('definizione', ''),
                    material.get('descrizione', ''),
                    material.get('datazione', ''),
                    material.get('stato_conservazione', '')
                ])
        return table_data

    def format_pottery_table(self):
        """Formatta i dati della ceramica per la tabella"""
        table_data = []
        if self.pottery_data:
            table_data.append(['ID','ID Number', 'Form', 'Fabric', 'Ware', 'Conservation %', 'Notes'])
            for pottery in self.pottery_data:
                table_data.append([
                    pottery.get('id_rep', ''),
                    pottery.get('id_number', ''),
                    pottery.get('form', ''),
                    pottery.get('fabric', ''),
                    pottery.get('ware', ''),
                    pottery.get('percent', ''),
                    pottery.get('note', '')
                ])
        return table_data

    def format_tomba_table(self):
        """Formatta i dati delle tombe per la tabella"""
        table_data = []
        if self.tomba_data:
            table_data.append(['Scheda TAF', 'Struttura', 'Individuo', 'Area', 'Rito', 'Descrizione', 'Interpretazione'])
            for tomba in self.tomba_data:
                table_data.append([
                    tomba.get('nr_scheda_taf', ''),
                    f"{tomba.get('sigla_struttura', '')} {tomba.get('nr_struttura', '')}",
                    tomba.get('nr_individuo', ''),
                    tomba.get('area', ''),
                    tomba.get('rito', ''),
                    tomba.get('descrizione_taf', ''),
                    tomba.get('interpretazione_taf', '')
                ])
        return table_data

    def format_periodizzazione_table(self):
        """Formatta i dati della periodizzazione per la tabella"""
        table_data = []
        if self.periodizzazione_data:
            table_data.append(['Periodo', 'Fase', 'Cronologia iniziale', 'Cronologia finale', 'Descrizione'])
            for periodo in self.periodizzazione_data:
                table_data.append([
                    periodo.get('periodo', ''),
                    periodo.get('fase', ''),
                    periodo.get('cron_iniziale', ''),
                    periodo.get('cron_finale', ''),
                    periodo.get('descrizione', '')
                ])
        return table_data

    def format_struttura_table(self):
        """Formatta i dati delle strutture per la tabella"""
        table_data = []
        if self.struttura_data:
            table_data.append(['Struttura', 'Categoria', 'Tipologia', 'Definizione', 'Descrizione', 'Interpretazione'])
            for struttura in self.struttura_data:
                table_data.append([
                    f"{struttura.get('sigla_struttura', '')} {struttura.get('numero_struttura', '')}",
                    struttura.get('categoria_struttura', ''),
                    struttura.get('tipologia_struttura', ''),
                    struttura.get('definizione_struttura', ''),
                    struttura.get('descrizione', ''),
                    struttura.get('interpretazione', '')
                ])
        return table_data


    def create_materials_table(self):
        """Create a formatted table of materials"""
        headers = ['ID','Inv. Number', 'Type', 'Definition', 'Description', 'Dating', 'Conservation']
        rows = []
        for material in self.materials_data:
            rows.append([
                material.get('id_invmat', ''),
                material.get('numero_inventario', ''),
                material.get('tipo_reperto', ''),
                material.get('definizione', ''),
                material.get('descrizione', ''),
                material.get('datazione', ''),
                material.get('stato_conservazione', '')
            ])
        return {'headers': headers, 'rows': rows}

    def create_pottery_table(self):
        """Create a formatted table of pottery"""
        headers = ['ID','ID NUmber', 'Form', 'Fabric', 'Ware', 'Conservation %', 'Notes']
        rows = []
        for pottery in self.pottery_data:
            rows.append([
                pottery.get('id_rep', ''),
                pottery.get('id_number', ''),
                pottery.get('form', ''),
                pottery.get('fabric', ''),
                pottery.get('ware', ''),
                pottery.get('percent', ''),
                pottery.get('note', '')
            ])
        return {'headers': headers, 'rows': rows}

    def format_table(self, table_data):
        """Format table data into a markdown table"""
        if not table_data['headers'] or not table_data['rows']:
            return ""

        # Calculate column widths
        col_widths = [len(str(h)) for h in table_data['headers']]
        for row in table_data['rows']:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(str(cell)))

        # Create header
        header = " | ".join(str(h).ljust(w) for h, w in zip(table_data['headers'], col_widths))
        separator = "|".join("-" * w for w in col_widths)

        # Create rows
        rows = []
        for row in table_data['rows']:
            formatted_row = " | ".join(str(cell).ljust(w) for cell, w in zip(row, col_widths))
            rows.append(formatted_row)

        # Combine all parts
        return f"\n\n{header}\n{separator}\n" + "\n".join(rows)



class ProgressDialog:
    def __init__(self):
        self.progressDialog = QProgressDialog()
        self.progressDialog.setWindowTitle("Aggiornamento rapporti area e sito")
        self.progressDialog.setLabelText("Inizializzazione...")
        #self.progressDialog.setCancelButtonText("")  # Disallow cancelling
        self.progressDialog.setRange(0, 0)
        self.progressDialog.setModal(True)
        self.progressDialog.show()

    def setValue(self, value):
        self.progressDialog.setValue(value)
        if value < value +1:  # Assuming that 100 is the maximum value
            self.progressDialog.setLabelText(f"Aggiornamento in corso... {value}")
        else:
            self.progressDialog.setLabelText("Finito")
            #self.progressDialog.close()


    def closeEvent(self, event):
        self.progressDialog.close()
        event.ignore()





class pyarchinit_US(QDialog, MAIN_DIALOG_CLASS):
    ''' This class creates the main dialog for the US form'''
    L=QgsSettings().value("locale/userLocale")[0:2]
    if L=='it':
        MSG_BOX_TITLE = "PyArchInit - Scheda US"
    elif L=='en':
        MSG_BOX_TITLE = "PyArchInit - SU form"
    elif L=='de':
        MSG_BOX_TITLE = "PyArchInit - SE formular"

    DATA_LIST = []
    DATA_LIST_REC_CORR = []
    DATA_LIST_REC_TEMP = []
    REC_CORR = 0
    REC_TOT = 0
    SITO = pyArchInitDialog_Config
    if L=='it':
        STATUS_ITEMS = {"b": "Usa", "f": "Trova", "n": "Nuovo Record"}

    if L=='de':
        STATUS_ITEMS = {"b": "Aktuell ", "f": "Finden", "n": "Neuer Rekord"}

    else :
        STATUS_ITEMS = {"b": "Current", "f": "Find", "n": "New Record"}
    BROWSE_STATUS = "b"
    SORT_MODE = 'asc'
    if L=='it':
        SORTED_ITEMS = {"n": "Non ordinati", "o": "Ordinati"}
    if L=='de':
        SORTED_ITEMS = {"n": "Nicht sortiert", "o": "Sortiert"}
    else:
        SORTED_ITEMS = {"n": "Not sorted", "o": "Sorted"}
    SORT_STATUS = "n"
    SORT_ITEMS_CONVERTED = ''
    UTILITY = Utility()
    DB_MANAGER = ""
    TABLE_NAME = 'us_table'
    MAPPER_TABLE_CLASS = "US"
    HOME = os.environ['PYARCHINIT_HOME']
    PDFFOLDER = '{}{}{}'.format(HOME, os.sep, "pyarchinit_PDF_folder")
    NOME_SCHEDA = "Scheda US"
    ID_TABLE = "id_us"
    ID_SITO ="sito"

    RELATIONSHIP_TYPES = ['Cuts', 'Covers', 'Abuts', 'Fills', 'Taglia', 'Copre', 'Si appoggia a', 'Riempie',
                          'Schneidet', 'Liegt über', 'Stützt sich auf', 'Verfüllt']
    RAPP_MAP = {
        'Riempito da': 'Riempie',
        'Tagliato da': 'Taglia',
        'Coperto da': 'Copre',
        'Si appoggia a': 'Gli si appoggia',
        'Riempie': 'Riempito da',
        'Taglia': 'Tagliato da',
        'Copre': 'Coperto da',
        'Gli si appoggia': 'Si appoggia a',
        'Filled by': 'Fills',
        'Cut by': 'Cuts',
        'Covered by': 'Covers',
        'Abuts': 'Supports',
        'Fills': 'Filled by',
        'Cuts': 'Cut by',
        'Covers': 'Covered by',
        'Supports': 'Abuts',
        '>>': '<<',
        '<<': '>>',
        '>': '<',
        '<': '>',
    }
    if L=='it':
        CONVERSION_DICT = {
            ID_TABLE: ID_TABLE,
            "Sito": "sito",
            "Area": "area",
            "US": "us",
            "Definizione stratigrafica": "d_stratigrafica",
            "Definizione interpretata": "d_interpretativa",
            "Descrizione": "descrizione",
            "Interpretazione": "interpretazione",
            "Periodo Iniziale": "periodo_iniziale",
            "Periodo Finale": "periodo_finale",
            "Fase Iniziale": "fase_iniziale",
            "Fase finale": "fase_finale",
            "Attività": "attivita",
            "Anno di scavo": "anno_scavo",
            "Sigla struttura": "struttura",
            "Scavato": "scavato",
            "Codice periodo": "cont_per",
            "Tipo unità": "unita_tipo",  # nuovi campi per USM
            "Settore": "settore",
            "Quadrato-Parete": "quad_par",
            "Ambiente": "ambient",
            "Saggio": "saggio",
            "Elementi datanti": "elem_datanti",
            "Funzione statica": "funz_statica",
            "Lavorazione": "lavorazione",
            "Spessore giunti": "spess_giunti",
            "Letti di posa": "letti_posa",
            "Altezza modulo": "alt_mod",
            "Unità edile rissuntiva": "un_ed_riass",
            "Reimpiego": "reimp",
            "Posa in opera": "posa_opera",
            "Quota minima USM": "quota_min_usm",
            "Quota max USM": "quota_max_usm",
            "Consistenza legante": "cons_legante",
            "Colore legante": "col_legante",
            "Aggregati legante": "aggreg_legante",
            "Consistenza-Texture": "con_text_mat",
            "Colore materiale": "col_materiale",
            "Inclusi materiali usm": "inclusi_materiali_usm",
            "n catalogo generale" : "n_catalogo_generale",  # campi aggiunti per archeo 3.0 e allineamento ICCD
            "n catalogo interno" : "n_catalogo_interno",
            "n catalogo internazionale" : "n_catalogo_internazionale",
            "soprintendenza" : "soprintendenza",
            "quota relativa" : "quota_relativa",
            "quota abs" : "quota_abs",
            "ref tm" : "ref_tm",
            "ref ra" : "ref_ra",
            "ref n" : "ref_n",
            "posizione" : "posizione",
            "criteri distinzione" : "criteri_distinzione",
            "modo formazione" : "modo_formazione",
            #"componenti organici" : "componenti_organici",
            #"componenti inorganici" : "componenti_inorganici",
            "lunghezza max" : "lunghezza_max",
            "altezza max" : "altezza_max",
            "altezza min" : "altezza_min",
            "profondita max" : "profondita_max",
            "profondita min" : "profondita_min",
            "larghezza media" : "larghezza_media",
            "quota max abs" : "quota_max_abs",
            "quota max rel" : "quota_max_rel",
            "quota min abs" : "quota_min_abs",
            "quota min rel" : "quota_min_rel",
            "osservazioni" : "osservazioni",
            "datazione" : "datazione",
            "flottazione" : "flottazione",
            "setacciatura" : "setacciatura",
            "affidabilita" : "affidabilita",
            "direttore us" : "direttore_us",
            "responsabile us" : "responsabile_us",
            "cod ente schedatore" : "cod_ente_schedatore",
            "data rilevazione" : "data_rilevazione",
            "data rielaborazione" : "data_rielaborazione",
            "lunghezza usm" : "lunghezza_usm",
            "altezza usm" : "altezza_usm",
            "spessore usm" : "spessore_usm",
            "tecnica muraria usm" : "tecnica_muraria_usm",
            "modulo usm" : "modulo_usm",
            "campioni malta usm" : "campioni_malta_usm",
            "campioni mattone usm" : "campioni_mattone_usm",
            "campioni pietra usm" : "campioni_pietra_usm",
            "provenienza materiali usm" : "provenienza_materiali_usm",
            "criteri distinzione usm" : "criteri_distinzione_usm",
            "uso primario usm" : "uso_primario_usm"
        }
    elif L=='de':
        CONVERSION_DICT = {
            ID_TABLE: ID_TABLE,
            "Ausgrabungsstätte": "sito",
            "Areal": "area",
            "SE": "us",
            "Stratigraphische Definition": "d_stratigrafica",
            "Interpretierte Definition": "d_interpretativa",
            "Beschreibung": "descrizione",
            "Deutung": "interpretazione",
            "Zeitraum Beginnend": "periodo_iniziale",
            "Zeitraum Ende": "periodo_finale",
            "Phase Beginnend": "fase_iniziale",
            "Phase Ende": "fase_finale",
            "Aktivität": "attivita",
            "Jahr": "anno_scavo",
            "Strukturcode": "struttura",
            "Ausgegraben": "scavato",
            "Periodencode erstellen": "cont_per",
            "Einheit eingeben": "unita_tipo",  # nuovi campi per USM
            "Sektor": "settore",
            "Quadrat / Wand": "quad_par",
            "Raum": "ambient",
            "Graben": "saggio",
            "Datierungselemente": "elem_datanti",
            "Statische Funktion": "funz_statica",
            "Verarbeitung": "lavorazione",
            "Stärke der Fugen": "spess_giunti",
            "Bett": "letti_posa",
            "Höhenmodul": "alt_mod",
            "Zusammenfassung der Baueinheit": "un_ed_riass",
            "Wiederverwendung": "reimp",
            "Verlegung": "posa_opera",
            "Mindesthöhe MSE": "quota_min_usm",
            "max. Höhe MSE": "quota_max_usm",
            "Konsistenz Bindemittel": "cons_legante",
            "Kleur Bindemittel": "col_legante",
            "Aggregat Bindemittel": "aggreg_legante",
            "Konsistenz-Texture": "con_text_mat",
            "Kleur material": "col_materiale",
            "Einschlüsse material mse": "inclusi_materiali_usm",
            "n catalogo generale" : "n_catalogo_generale",  # campi aggiunti per archeo 3.0 e allineamento ICCD
            "N °. Cat. Int." : "n_catalogo_interno",
            "N °. Cat. Internat." : "n_catalogo_internazionale",
            "Landesamt" : "soprintendenza",
            "Relative Höhe" : "quota_relativa",
            "Absolute Höhe" : "quota_abs",
            "Materialformular Referenz" : "ref_tm",
            "Archäologische Funde Referenz" : "ref_ra",
            "Ref. N." : "ref_n",
            "Lange" : "posizione",
            "Unterscheidungskriterien" : "criteri_distinzione",
            "Trainingsmodus" : "modo_formazione",
            #"componenti organici" : "componenti_organici",
            #"componenti inorganici" : "componenti_inorganici",
            "Max.Länge" : "lunghezza_max",
            "Max. Höhe" : "altezza_max",
            "Min. Höhe" : "altezza_min",
            "Max. Tiefe" : "profondita_max",
            "Min. Tiefe" : "profondita_min",
            "Durchschnittliche Breite" : "larghezza_media",
            "Absolute maximale Höhe" : "quota_max_abs",
            "Relative maximale Höhe" : "quota_max_rel",
            "Absolute min Höhe" : "quota_min_abs",
            "Relative min Höhe" : "quota_min_rel",
            "Beobachtungen" : "osservazioni",
            "Datierung" : "datazione",
            "Flotation" : "flottazione",
            "Siebanlage" : "setacciatura",
            "Zuverlässigkeit" : "affidabilita",
            "SE-Direktor" : "direttore_us",
            "SE Leiter" : "responsabile_us",
            "Verfasser Firma's Code" : "cod_ente_schedatore",
            "Datum der Entdeckung" : "data_rilevazione",
            "Überarbeitetes Datum" : "data_rielaborazione",
            "MSE Länge" : "lunghezza_usm",
            "MSE Höhe" : "altezza_usm",
            "MSE Dicke" : "spessore_usm",
            "MSE Maurer-Technik" : "tecnica_muraria_usm",
            "MSE modul" : "modulo_usm",
            "Mörtelproben" : "campioni_malta_usm",
            "Ziegelsteinproben" : "campioni_mattone_usm",
            "Steinproben" : "campioni_pietra_usm",
            "Materieller Ursprung" : "provenienza_materiali_usm",
            "MSE Kriterien für die Unterscheidung" : "criteri_distinzione_usm",
            "Hauptanwendung MSE" : "uso_primario_usm"
        }
    else:
        CONVERSION_DICT = {
            ID_TABLE: ID_TABLE,
            "Site": "sito",
            "Area": "area",
            "SU": "us",
            "Stratigraphic definition": "d_stratigrafica",
            "Interpreted definition": "d_interpretativa",
            "Description": "descrizione",
            "Interpretation": "interpretazione",
            "Initial Period": "periodo_iniziale",
            "Final Period": "periodo_finale",
            "Starting Phase": "fase_iniziale",
            "Final Phase": "fase_finale",
            "Activity": "attivita",
            "Year of excavation": "anno_scavo",
            "Structure code": "struttura",
            "Excavated": "scavato",
            "Period code": "cont_per",
            "Unit type": "unita_tipo",  # nuovi campi per USM
            "Sector": "settore",
            "Square-profile": "quad_par",
            "Room": "ambient",
            "Trench": "saggio",
            "Dating elements": "elem_datanti",
            "Static function": "funz_statica",
            "Processing": "lavorazione",
            "Joint thickness": "spess_giunti",
            "Laying beds": "letti_posa",
            "Module height": "alt_mod",
            "Resurgent Building Unit": "un_ed_riass",
            "Reuse": "reimp",
            "Laying": "posa_opera",
            "Minimum WSU elevation": "quota_min_usm",
            "Max WSU elevation": "quota_max_usm",
            "Binder consistency": "cons_legante",
            "Binder color": "col_legante",
            "Binder aggregates": "aggreg_legante",
            "Consistency-Texture": "con_text_mat",
            "Material color": "col_materiale",
            "Including wsu materials": "inclusi_materiali_usm",
            "n general catalogue" : "n_catalogo_generale",  # campi aggiunti per archeo 3.0 e allineamento ICCD
            "n internal catalogue" : "n_catalogo_interno",
            "n international catalogue" : "n_catalogo_internazionale",
            "superintendence" : "soprintendenza",
            "Relative elevation" : "quota_relativa",
            "abs elevation" : "quota_abs",
            "ref tm" : "ref_tm",
            "ref ra" : "ref_ra",
            "ref n" : "ref_n",
            "position" : "posizione",
            "distinction criteria" : "criteri_distinzione",
            "formation mode" : "modo_formazione",
            #"componenti organici" : "componenti_organici",
            #"componenti inorganici" : "componenti_inorganici",
            "max length" : "lunghezza_max",
            "Max height" : "altezza_max",
            "min height" : "altezza_min",
            "Max depth" : "profondita_max",
            "Min depth" : "profondita_min",
            "average width" : "larghezza_media",
            "elevation max abs" : "quota_max_abs",
            "elevation max relative" : "quota_max_rel",
            "elevation min abs" : "quota_min_abs",
            "elevation min relative" : "quota_min_rel",
            "observation" : "osservazioni",
            "Dating" : "datazione",
            "Flotation" : "flottazione",
            "Sieving" : "setacciatura",
            "Reliability" : "affidabilita",
            "Director SU" : "direttore_us",
            "Responsible SU" : "responsabile_us",
            "Company system code" : "cod_ente_schedatore",
            "date of detection" : "data_rilevazione",
            "date reworked" : "data_rielaborazione",
            "wsu length" : "lunghezza_usm",
            "wsu height" : "altezza_usm",
            "wsu thickness" : "spessore_usm",
            "wsu masonry technique" : "tecnica_muraria_usm",
            "wsu module" : "modulo_usm",
            "wsu mortar samples" : "campioni_malta_usm",
            "wsu brick samples" : "campioni_mattone_usm",
            "wsu stone samples" : "campioni_pietra_usm",
            "wsu material source" : "provenienza_materiali_usm",
            "wsu distinction criteria" : "criteri_distinzione_usm",
            "primary use wsu" : "uso_primario_usm"
        }
    if L=='it':
        SORT_ITEMS = [
            ID_TABLE,  #0
            "Sito", #1
            "Area", #2
            'US', #3
            "Definizione stratigrafica", #4
            "Definizione interpretata",  #5
            "Descrizione",               #6
            "Interpretazione",           #7
            "Periodo Iniziale",          #8
            "Periodo Finale",            #9
            "Fase Iniziale",             #10
            "Fase Finale",               #11
            "Attività",
            "Anno di scavo",
            "Sigla struttura",
            "Scavato",
            "Codice periodo",
            "Indice di ordinamento",
            "Tipo unità",  # nuovi campi per USM
            "Settore",
            "Quadrato-Parete",
            "Ambiente",
            "Saggio",
            "Elementi datanti",
            "Funzione statica",
            "Lavorazione",
            "Spessore giunti",
            "Letti di posa",
            "Altezza modulo",
            "Unità edile rissuntiva",
            "Reimpiego",
            "Posa in opera",
            "Quota minima USM",
            "Quota max USM",
            "Consistenza legante",
            "Colore legante",
            "Aggregati legante",
            "Consistenza-Texture",
            "Colore materiale",
            "Inclusi materiali usm",
            "n catalogo generale",  #campi aggiunti per archeo 3.0 e allineamento ICCD
            "n catalogo interno",
            "n catalogo internazionale",
            "soprintendenza",
            "quota relativa",
            "quota abs",
            "ref tm",
            "ref ra",
            "ref n",
            "posizione",
            "criteri distinzione",
            "modo formazione",
            #"organici",
            #"inorganici",
            "lunghezza max",
            "altezza max",
            "altezza min",
            "profondita max",
            "profondita min",
            "larghezza media",
            "quota max abs",
            "quota max rel",
            "quota min abs",
            "quota min rel",
            "osservazioni",
            "datazione",
            "flottazione",
            "setacciatura",
            "affidabilita",
            "direttore us",
            "responsabile us",
            "cod ente schedatore",
            "data rilevazione",
            "data rielaborazione",
            "lunghezza usm",
            "altezza usm",
            "spessore usm",
            "tecnica muraria usm",
            "modulo usm",
            "campioni malta usm",
            "campioni mattone usm",
            "campioni pietra usm",
            "provenienza materiali usm",
            "criteri distinzione usm",
            "uso primario usm"]
    elif L == 'de':
        SORT_ITEMS = [
            ID_TABLE,  #0
            "Ausgrabungsstätte",
            "Areal",
            "SE",
            "Stratigraphische Definition",
            "Interpretierte Definition",
            "Beschreibung",
            "Deutung",
            "Zeitraum Beginnend",
            "Zeitraum Ende",
            "Phase Beginnend",
            "Phase Ende",
            "Aktivität",
            "Jahr",
            "Strukturcode",
            "Ausgegraben",
            "Periodencode erstellen",
            "Einheit eingeben",
            "Sektor",
            "Quadrat / Wand",
            "Raum",
            "Graben",
            "Datierungselemente",
            "Statische Funktion",
            "Verarbeitung",
            "Stärke der Fugen",
            "Bett",
            "Höhenmodul",
            "Zusammenfassung der Baueinheit",
            "Wiederverwendung",
            "Verlegung",
            "Mindesthöhe MSE",
            "max. Höhe MSE",
            "Konsistenz Bindemittel",
            "Kleur Bindemittel",
            "Aggregat Bindemittel",
            "Konsistenz-Texture",
            "Kleur material",
            "Einschlüsse material mse",
            "n catalogo generale",
            "N °. Cat. Int." ,
            "N °. Cat. Internat." ,
            "Landesamt" ,
            "Relative Höhe" ,
            "Absolute Höhe" ,
            "Materialformular Referenz",
            "Archäologische Funde Referenz",
            "Ref. N." ,
            "Lange" ,
            "Unterscheidungskriterien" ,
            "Trainingsmodus" ,
            #"componenti organici",
            #"componenti inorganici",
            "Max.Länge" ,
            "Max. Höhe" ,
            "Min. Höhe" ,
            "Max. Tiefe" ,
            "Min. Tiefe" ,
            "Durchschnittliche Breite" ,
            "Absolute maximale Höhe",
            "Relative maximale Höhe",
            "Absolute min Höhe" ,
            "Relative min Höhe",
            "Beobachtungen",
            "Datierung" ,
            "Flotation" ,
            "Siebanlage" ,
            "Zuverlässigkeit" ,
            "SE-Direktor",
            "SE Leiter",
            "Verfasser Firma's Code",
            "Datum der Entdeckung",
            "Überarbeitetes Datum",
            "MSE Länge",
            "MSE Höhe",
            "MSE Dicke",
            "MSE Maurer-Technik",
            "MSE modul" ,
            "Mörtelproben",
            "Ziegelsteinproben",
            "Steinproben",
            "Materieller Ursprung",
            "MSE Kriterien für die Unterscheidung",
            "Hauptanwendung MSE"]
    else:
        SORT_ITEMS = [
            ID_TABLE,  #0
            "Site", #1
            "Area", #2
            "SU", #3
            "Stratigraphic definition", #4
            "Interpreted definition," #5
            "Description", #6
            "Interpretation", #7
            "Initial Period", #8
            "Final Period",#9
            "Starting Phase", #10
            "Final Phase", #11
            "Activity",
            "Year of excavation",
            "Structure code",
            "Excavated",
            "Period code",
            "Sorting index",
            "Unit type", # new fields for USM
            "Sector",
            "Square-profile",
            "Room",
            "Trench",
            "Dating elements",
            "Static function",
            "Processing",
            "Joint thickness",
            "Laying beds",
            "Module height",
            "Resurgent Building Unit",
            "Reuse",
            "Laying",
            "Minimum WSU elevation",
            "Max WSU elevation",
            "Binder consistency",
            "Binder color",
            "Binder aggregates",
            "Consistency-Texture",
            "Material color",
            "Including wsu material",
            "n general catalogue", #added fields for arch 3.0 and ICCD alignment
            "n internal catalogue",
            "n international catalogue",
            "superintendence",
            "Relative elevation",
            "abs elevation",
            "ref tm",
            "ref ra",
            "ref n",
            "position",
            "distinction criteria",
            "formation mode",
            #Organic,
            #Inorganic,
            "max length",
            "Max height",
            "min height",
            "Max depth",
            "Min depth",
            "average width",
            "elevation max abs",
            "elevation max relative",
            "elevation min abs",
            "elevation min relative",
            "observation",
            "Dating",
            "Flotation",
            "Sieving",
            "Reliability",
            "Director SU",
            "Responsible SU",
            "Company system code",
            "date of detection",
            "date reworked",
            "wsu length",
            "wsu height",
            "wsu thickness",
            "wsu masonry technique",
            "wsu module",
            "wsu mortar samples",
            "wsu brick samples",
            "wsu stone samples",
            "wsu material source",
            "wsu distinction criteria",
            "primary use wsu",
    ]
    TABLE_FIELDS = [
        'sito',  # 0
        'area',  # 1
        'us',
        'd_stratigrafica',
        'd_interpretativa',
        'descrizione',
        'interpretazione',
        'periodo_iniziale',
        'fase_iniziale',
        'periodo_finale',
        'fase_finale',
        'scavato',
        'attivita',
        'anno_scavo',
        'metodo_di_scavo',
        'inclusi',
        'campioni',
        'rapporti',
        #'organici',
        #'inorganici',
        'data_schedatura',
        'schedatore',
        'formazione',
        'stato_di_conservazione',
        'colore',
        'consistenza',
        'struttura',
        'cont_per',
        'order_layer',
        'documentazione',
        'unita_tipo',  # nuovi campi per USM
        'settore',
        'quad_par',
        'ambient',
        'saggio',
        'elem_datanti',
        'funz_statica',
        'lavorazione',
        'spess_giunti',
        'letti_posa',
        'alt_mod',
        'un_ed_riass',
        'reimp',
        'posa_opera',
        'quota_min_usm',
        'quota_max_usm',
        'cons_legante',
        'col_legante',
        'aggreg_legante',
        'con_text_mat',
        'col_materiale', #48
        'inclusi_materiali_usm', #49
        'n_catalogo_generale',  # 51 campi aggiunti per archeo 3.0 e allineamento ICCD #50
        'n_catalogo_interno',  # 52
        'n_catalogo_internazionale',  # 53
        'soprintendenza',  # 54
        'quota_relativa',  # 55
        'quota_abs',  # 56
        'ref_tm',  # 57
        'ref_ra',  # 58
        'ref_n',  # 59
        'posizione',  # 60
        'criteri_distinzione',  # 61
        'modo_formazione',  # 62
        'componenti_organici',  # 63
        'componenti_inorganici',  # 64
        'lunghezza_max',  # 65
        'altezza_max',  # 66
        'altezza_min',  # 67
        'profondita_max',  # 68
        'profondita_min',  # 69
        'larghezza_media',  # 70
        'quota_max_abs',  # 71
        'quota_max_rel',  # 72
        'quota_min_abs',  # 73
        'quota_min_rel',  # 74
        'osservazioni',  # 75
        'datazione',  # 76
        'flottazione',  # 77
        'setacciatura',  # 78
        'affidabilita',  # 79
        'direttore_us',  # 80
        'responsabile_us',  # 81
        'cod_ente_schedatore',  # 82
        'data_rilevazione',  # 83
        'data_rielaborazione',  # 84
        'lunghezza_usm',  # 85
        'altezza_usm',  # 86
        'spessore_usm',  # 87
        'tecnica_muraria_usm',  # 88
        'modulo_usm',  # 89
        'campioni_malta_usm',  # 90
        'campioni_mattone_usm',  # 91
        'campioni_pietra_usm',  # 92
        'provenienza_materiali_usm',  # 93
        'criteri_distinzione_usm',  # 94
        'uso_primario_usm',
        'tipologia_opera',
        'sezione_muraria',
        'superficie_analizzata',
        'orientamento',
        'materiali_lat',
        'lavorazione_lat',
        'consistenza_lat',
        'forma_lat',
        'colore_lat',
        'impasto_lat',
        'forma_p',
        'colore_p',
        'taglio_p',
        'posa_opera_p',
        'inerti_usm',
        'tipo_legante_usm',
        'rifinitura_usm',
        'materiale_p',
        'consistenza_p',
        'rapporti2',
        'doc_usv',
        ]
    LANG = {
        "IT": ['it_IT', 'IT', 'it', 'IT_IT'],
        "EN_US": ['en_US','EN_US','en','EN'],
        "DE": ['de_DE','de','DE', 'DE_DE']
    }

    REPORT_PATH = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Report_folder")
    MATRIX_PATH = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Matrix_folder")
    BIN = '{}{}{}'.format(HOME, os.sep, "bin")
    DB_SERVER = "not defined"  ####nuovo sistema sort
    log_message = pyqtSignal(str, str)  # Aggiungi questo segnale
    def __init__(self, iface):
        super().__init__()


        self.iface = iface
        self.pyQGIS = Pyarchinit_pyqgis(iface)

        self.setupUi(self)
        self.setAcceptDrops(True)
        self.progress_dialog = None
        self.report_thread = None
        self.report_rapporti2 = None
        self.fig = None
        self.canvas = None
        self.video_player=None
        self.iconListWidget.setDragDropMode(QAbstractItemView.DragDrop)
        self.mDockWidget_2.setHidden(True)
        self.mDockWidget_export.setHidden(True)
        self.mDockWidget_3.setHidden(True)
        self.mDockWidget_4.setHidden(True)
        self.mDockWidget_h.setHidden(True)
        self.mQgsFileWidget.setHidden(True)
        self.toolButton_file_doc.setHidden(True)
        self.mDockWidget_5.setHidden(True)
        self.tableWidget_rapporti2.setHidden(True)
        self.pushButton_insert_row_rapporti2.setHidden(True)
        self.pushButton_remove_row_rapporti2.setHidden(True)
        self.pushButton_update.setHidden(True)
        self.progressBar_2.setHidden(True)
        self.progressBar_3.setHidden(True)
        self.currentLayerId = None
        self.search = SearchLayers(iface)
        # Dizionario per memorizzare le immagini in cache
        self.image_cache = OrderedDict()

        # Numero massimo di elementi nella cache
        self.cache_limit = 100


        try:
            self.on_pushButton_connect_pressed()
        except Exception as e:
            QMessageBox.warning(self, "Connection System", str(e), QMessageBox.Ok)
            # SIGNALS & SLOTS Functions
        if len(self.DATA_LIST)==0:
            self.comboBox_sito.setCurrentIndex(0)
        else:
            self.comboBox_sito.setCurrentIndex(1)



        self.comboBox_sito.currentTextChanged.connect(self.charge_periodo_iniz_list)
        self.comboBox_unita_tipo.currentTextChanged.connect(self.charge_struttura_list)
        self.comboBox_sito.currentTextChanged.connect(self.charge_struttura_list)
        self.comboBox_per_iniz.currentIndexChanged.connect(self.charge_periodo_fin_list)
        self.comboBox_per_iniz.currentIndexChanged.connect(self.charge_fase_iniz_list)
        self.comboBox_sito.currentTextChanged.connect(self.geometry_unitastratigrafiche)### rallenta molto
        self.comboBox_sito.currentIndexChanged.connect(self.geometry_unitastratigrafiche)### rallenta molto
        self.comboBox_unita_tipo.currentTextChanged.connect(self.charge_insert_ra)
        self.comboBox_sito.currentTextChanged.connect(self.charge_insert_ra)
        self.search_1.currentTextChanged.connect(self.update_filter)
        self.comboBox_per_fin.currentIndexChanged.connect(self.charge_fase_fin_list)
        self.toolButton_pdfpath.clicked.connect(self.setPathpdf)
        self.toolButton_input.clicked.connect(self.setPathdot)
        self.toolButton_output.clicked.connect(self.setPathgraphml)
        self.toolButton_file_doc.clicked.connect(self.setDoc_ref)
        self.pbnOpenpdfDirectory.clicked.connect(self.openpdfDir)
        self.progressBar.setTextVisible(True)
        sito = self.comboBox_sito.currentText()
        self.comboBox_sito.setEditText(sito)
        self.fill_fields()
        self.customize_GUI()

        self.set_sito()
        self.msg_sito()

        self.show()
        self.checkBox_query.update()
        self.checkBox_query.stateChanged.connect(self.listview_us)###anche questo

        self.tableWidget_rapporti.itemSelectionChanged.connect(self.us_t)

        self.comboBox_unita_tipo.currentTextChanged.connect(self.change_label)
        self.field.currentTextChanged.connect(self.value_check)


        self.comboBox_per_iniz.currentTextChanged.connect(self.check_v)
        self.charge_insert_ra()
        self.charge_struttura_list()
        self.tableWidget_rapporti.itemChanged.connect(self.check_listoflist)
        self.update_dating()
        # Imposta il collegamento nascosto per attivare text2sql
        self.text2sql_db_shortcut = QShortcut(QKeySequence("Ctrl+Shift+X"), self)
        self.text2sql_db_shortcut.activated.connect(self.text2sql)
        # Imposta la scorciatoia da tastiera per l'aggiornamento
        self.update_shortcut = QShortcut(QKeySequence("Ctrl+U"), self)
        self.update_shortcut.activated.connect(self.update_all_areas)

        #self.update2_shortcut = QShortcut(QKeySequence("Ctrl+P"), self)
        #self.update2_shortcut.activated.connect(self.update_rapporti_col_2)
        #self.delete_all_shortcut = QtWidgets.QAction(self)
        # Imposta la scorciatoia da tastiera per eliminare tutti i record filtrati
        self.delete_all_shortcut=QShortcut(QKeySequence('Ctrl+Shift+D'), self)
        self.delete_all_shortcut.activated.connect(self.delete_all_filtered_records)


        self.use_like_query = False
        self.new_search_shortcut = QShortcut(QKeySequence('Ctrl+Shift+N'), self)
        self.new_search_shortcut.activated.connect(self.switch_search_mode)
        self.pushButton_sketchgpt.clicked.connect(self.sketchgpt)

        self.report_rapporti=''
        self.list_rapporti=[]
        self.pushButton_report_generator.clicked.connect(self.generate_and_display_report)

        # Definizione completa degli analysis steps
        self.analysis_steps = []


        self.materials_data = []
        self.pottery_data = []
        self.us_data = []
        self.current_step = 0
        self.token_count = 0
        self.max_tokens = 4000
        #self.log_message.connect(self.process_terminal.append_message)

    def get_images_for_entities(self, entity_ids, log_signal=None, entity_type='US'):
        def log(message, level="info"):
            if log_signal:
                log_signal.emit(message, level)
        """Recupera le immagini dalla tabella mediaentity in base agli ID forniti."""
        log(f"Called get_images_for_entities with entity_ids: {entity_ids}, entity_type: {entity_type}")

        if not entity_ids:
            return []

        try:
            images = []
            conn = Connection()
            thumb_resize = conn.thumb_resize()
            thumb_resize_str = thumb_resize['thumb_resize']

            # Map entity_type to the corresponding ID table
            id_table_mapping = {
                'US': 'id_us',
                'REPERTO': 'id_invmat',
                'INVENTARIO_MATERIALI': 'id_invmat',
                'CERAMICA': 'id_rep',
                'POTTERY': 'id_rep',
                'TOMBA': 'id_tomba',
                'STRUTTURA': 'id_struttura'
            }

            # Get the appropriate ID table for the entity type
            id_table = id_table_mapping.get(entity_type, self.ID_TABLE)

            for entity_id in entity_ids:
                log(f"Using id table: {id_table} for entity_type: {entity_type}")
                # Usa la stessa logica di loadMediaPreview
                rec_list = id_table + " = " + str(entity_id)
                log(f"Called rec list: {rec_list}")

                # Try different entity types for materials and pottery
                entity_types_to_try = [entity_type]
                if entity_type == 'REPERTO':
                    entity_types_to_try.append('INVENTARIO_MATERIALI')
                elif entity_type == 'CERAMICA':
                    entity_types_to_try.append('POTTERY')
                elif entity_type == 'INVENTARIO_MATERIALI':
                    entity_types_to_try.append('REPERTO')
                elif entity_type == 'POTTERY':
                    entity_types_to_try.append('CERAMICA')

                record_us_list = []
                found_entity_type = None
                for et in entity_types_to_try:
                    # Usa la stessa logica di loadMediaPreview
                    search_dict = {
                        'id_entity': "'" + str(entity_id) + "'",
                        'entity_type': f"'{et}'"
                    }

                    result = self.DB_MANAGER.query_bool(search_dict, 'MEDIATOENTITY')
                    if result:
                        record_us_list = result
                        found_entity_type = et
                        log(f"Found {len(record_us_list)} records for entity_id {entity_id} of type {et}")
                        break

                if not record_us_list:
                    log(f"No records found for entity_id {entity_id} with any of the tried entity types: {entity_types_to_try}")  # Debug log

                for media_record in record_us_list:
                    search_dict = {'id_media': "'" + str(media_record.id_media) + "'"}
                    u = Utility()
                    search_dict = u.remove_empty_items_fr_dict(search_dict)
                    mediathumb_data = self.DB_MANAGER.query_bool(search_dict, "MEDIA_THUMB")
                    log(f"Found {len(mediathumb_data)} thumbs for media_id {media_record.id_media}")  # Debug log

                    if mediathumb_data:
                        thumb_path = str(mediathumb_data[0].path_resize)
                        images.append({
                            'id': entity_id,
                            'url': thumb_resize_str + thumb_path,
                            'caption': media_record.media_name,
                            'entity_type': found_entity_type if found_entity_type else entity_type
                        })
                        log(f"Added image with path: {thumb_resize_str + thumb_path}")  # Debug log

            log(f"Returning total of {len(images)} images for entity_type {entity_type}")  # Debug log
            return images

        except Exception as e:
            log(f"Error in get_images_for_entities: {str(e)}")
            traceback.print_exc()  # Questo mostrerà lo stack trace completo
            return []


    def count_tokens(self, text):
        """
        Estimate the number of tokens in the text.
        This is a more accurate approximation than just splitting on whitespace.
        """
        # Roughly 4 characters per token for English text
        return len(text) // 4

    def create_vector_db(self, data, table_name):
        """
        Create a vector database from the data for RAG approach.

        Args:
            data: List of data records
            table_name: Name of the table for context

        Returns:
            FAISS vector store for retrieval
        """
        if not data:
            return None

        # Convert data records to text documents
        documents = []
        for i, record in enumerate(data):
            if isinstance(record, dict):
                # Format dictionary as text
                content = f"Record {i+1} from {table_name}:\n"
                content += "\n".join(f"{k}: {v}" for k, v in record.items())
            else:
                # Format object as text
                content = f"Record {i+1} from {table_name}:\n"
                content += "\n".join(f"{k}: {getattr(record, k, '')}" for k in dir(record) 
                                    if not k.startswith('_') and not callable(getattr(record, k, None)))

            documents.append(content)

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            length_function=len,
        )
        texts = text_splitter.create_documents(documents)

        # Create vector store
        try:
            embeddings = OpenAIEmbeddings(api_key=self.api_key)
            vector_store = FAISS.from_documents(texts, embeddings)
            return vector_store
        except Exception as e:
            self.log_message.emit(f"Error creating vector database: {str(e)}", "error")
            return None

    def retrieve_relevant_data(self, vector_store, query, k=5):
        """
        Retrieve the most relevant data from the vector store based on the query.

        Args:
            vector_store: FAISS vector store
            query: Query string
            k: Number of documents to retrieve

        Returns:
            String containing the retrieved documents
        """
        if not vector_store:
            return ""

        try:
            # Retrieve relevant documents
            docs = vector_store.similarity_search(query, k=k)

            # Format the retrieved documents
            retrieved_data = "\n\n".join([doc.page_content for doc in docs])

            return retrieved_data
        except Exception as e:
            self.log_message.emit(f"Error retrieving data: {str(e)}", "error")
            return ""

    def create_rag_chain(self, vector_store, llm):
        """
        Create a RetrievalQA chain for the RAG approach.

        Args:
            vector_store: FAISS vector store
            llm: Language model

        Returns:
            RetrievalQA chain
        """
        if not vector_store:
            return None

        try:
            # Create a prompt template
            template = """
            You are an archaeological expert analyzing data from an excavation.

            Use the following pieces of context to answer the question at the end.
            If you don't know the answer, just say that you don't know, don't try to make up an answer.

            Context:
            {context}

            Question: {question}

            Answer:
            """

            prompt = PromptTemplate(
                template=template,
                input_variables=["context", "question"]
            )

            # Create the chain
            chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=vector_store.as_retriever(),
                chain_type_kwargs={"prompt": prompt}
            )

            return chain
        except Exception as e:
            self.log_message.emit(f"Error creating RAG chain: {str(e)}", "error")
            return None

    def split_data_to_fit_tokens(self, data, columns, max_tokens_per_chunk=8000):
        """
        Split data into chunks that fit within token limits.
        Uses a more efficient algorithm that prioritizes keeping related data together.

        Args:
            data: List of data records
            columns: List of columns to include
            max_tokens_per_chunk: Maximum tokens per chunk (default: 8000)

        Returns:
            List of chunks, where each chunk is a list of records
        """
        if not data:
            return []

        chunks = []
        current_chunk = []
        current_tokens = 0

        # Group records by area or other relevant field if possible
        grouped_data = {}
        group_field = None

        # Try to find a suitable grouping field
        for field in ['area', 'saggio', 'quadrato', 'ambiente']:
            if hasattr(data[0], field) or (isinstance(data[0], dict) and field in data[0]):
                group_field = field
                break

        # Group data if possible
        if group_field:
            for record in data:
                group_value = getattr(record, group_field, '') if hasattr(record, group_field) else record.get(group_field, '')
                if group_value not in grouped_data:
                    grouped_data[group_value] = []
                grouped_data[group_value].append(record)

            # Process each group
            for group_value, group_records in grouped_data.items():
                group_chunk = []
                group_tokens = 0

                for record in group_records:
                    if isinstance(record, dict):
                        record_text = "\n".join(f"{col}: {record.get(col, '')}" for col in columns if col in record)
                    else:
                        record_text = "\n".join(f"{col}: {getattr(record, col, '')}" for col in columns)

                    try:
                        tokens = self.count_tokens(record_text)
                    except AttributeError:
                        # Fallback if count_tokens method is not accessible
                        tokens = len(record_text) // 4  # Roughly 4 characters per token for English text

                    # If this single record is too large, we need to truncate it
                    if tokens > max_tokens_per_chunk:
                        # Truncate the record text to fit
                        truncated_text = record_text[:max_tokens_per_chunk*4]  # Approximate chars
                        try:
                            tokens = self.count_tokens(truncated_text)
                        except AttributeError:
                            # Fallback if count_tokens method is not accessible
                            tokens = len(truncated_text) // 4  # Roughly 4 characters per token for English text
                        self.log_message.emit(f"Warning: Record truncated to fit token limit", "warning")

                    if group_tokens + tokens <= max_tokens_per_chunk:
                        group_chunk.append(record)
                        group_tokens += tokens
                    else:
                        # This group is full, add it to chunks and start a new one
                        if group_chunk:
                            chunks.append(group_chunk)
                        group_chunk = [record]
                        group_tokens = tokens

                # Add the last group chunk
                if group_chunk:
                    chunks.append(group_chunk)
        else:
            # No grouping field, process records sequentially
            for record in data:
                if isinstance(record, dict):
                    record_text = "\n".join(f"{col}: {record.get(col, '')}" for col in columns if col in record)
                else:
                    record_text = "\n".join(f"{col}: {getattr(record, col, '')}" for col in columns)

                try:
                    tokens = self.count_tokens(record_text)
                except AttributeError:
                    # Fallback if count_tokens method is not accessible
                    tokens = len(record_text) // 4  # Roughly 4 characters per token for English text

                # If this single record is too large, we need to truncate it
                if tokens > max_tokens_per_chunk:
                    # Truncate the record text to fit
                    truncated_text = record_text[:max_tokens_per_chunk*4]  # Approximate chars
                    try:
                        tokens = self.count_tokens(truncated_text)
                    except AttributeError:
                        # Fallback if count_tokens method is not accessible
                        tokens = len(truncated_text) // 4  # Roughly 4 characters per token for English text
                    self.log_message.emit(f"Warning: Record truncated to fit token limit", "warning")

                if current_tokens + tokens <= max_tokens_per_chunk:
                    current_chunk.append(record)
                    current_tokens += tokens
                else:
                    chunks.append(current_chunk)
                    current_chunk = [record]
                    current_tokens = tokens

            if current_chunk:
                chunks.append(current_chunk)

        return chunks


    def analyze_site_context(self, report_data):
        """Analyze site context in detail"""
        return f"""
        Analisi del contesto geografico e amministrativo:
        Regione: {report_data.get('Regione', '')}
        Provincia: {report_data.get('Provincia', '')}
        Comune: {report_data.get('Comune', '')}

        Informazioni amministrative:
        Ente di riferimento: {report_data.get('Ente di riferimento', '')}
        Committenza: {report_data.get('Committenza', '')}
        Direzione scientifica: {report_data.get('Direzione scientifica', '')}

        Dettagli del cantiere:
        Nome cantiere: {report_data.get('Cantiere', '')}
        Collocazione: {report_data.get('Collocazione cantiere', '')}
        Periodo: {report_data.get('Periodo cantiere', '')}
        """

    def analyze_stratigraphy_relationships(self, us_data):
        """Analyze stratigraphic relationships"""
        relationships = {}
        for us in us_data:
            area = us.get('area', '')
            if area not in relationships:
                relationships[area] = []
            relationships[area].append({
                'us': us.get('us', ''),
                'rapporti': us.get('rapporti', ''),
                'd_interpretazione': us.get('d_interpretazione', '')
            })
        return relationships

    def analyze_stratigraphy_by_area(self, us_data):
        """Analyze stratigraphy by area"""
        areas = {}
        for us in us_data:
            area = us.get('area', '')
            if area not in areas:
                areas[area] = []
            areas[area].append(us)
        return areas

    def analyze_materials_by_type(self, material_data):
        """Analyze materials by type"""
        types = {}
        for material in material_data:
            tipo = material.get('tipo_reperto', '')
            if tipo not in types:
                types[tipo] = []
            types[tipo].append(material)
        return types

    def analyze_pottery_by_class(self, pottery_data):
        """Analyze pottery by class"""
        classes = {}
        for pottery in pottery_data:
            ware = pottery.get('ware', '')
            fabric = pottery.get('fabric', '')
            if ware not in classes:
                classes[ware] = []
            classes[ware].append(pottery)
            if fabric not in classes:
                classes[fabric] = []
            classes[fabric].append(pottery)
        return classes

    def analyze_pottery_photo(self, pottery_data):
        """Analyze pottery photos"""
        classes = {}
        for pottery in pottery_data:
            photo = pottery.get('photo', '')
            if photo not in classes:
                classes[photo] = []
            classes[photo].append(pottery)
        return classes

    def validate_us_description(self, us_data):
        """Validate US descriptions"""
        missing_us = []
        for us in us_data:
            if not us.get('descrizione') or not us.get('interpretazione'):
                missing_us.append(us.get('us'))
        return {
            'valid': len(missing_us) == 0,
            'missing': missing_us,
            'message': f"Missing descriptions for US: {', '.join(map(str, missing_us))}" if missing_us else "All US described"
        }

    def validate_materials_description(self, materials_data):
        """Validate materials descriptions"""
        missing_materials = []
        for material in materials_data:
            if not material.get('descrizione'):
                missing_materials.append(material.get('numero_inventario'))
        return {
            'valid': len(missing_materials) == 0,
            'missing': missing_materials,
            'message': f"Missing descriptions for materials: {', '.join(map(str, missing_materials))}" if missing_materials else "All materials described"
        }

    def validate_pottery_description(self, pottery_data):
        """Validate pottery descriptions"""
        missing_pottery = []
        for pottery in pottery_data:
            if not pottery.get('fabric') or not pottery.get('form'):
                missing_pottery.append(pottery.get('id_number'))
        return {
            'valid': len(missing_pottery) == 0,
            'missing': missing_pottery,
            'message': f"Missing descriptions for pottery: {', '.join(map(str, missing_pottery))}" if missing_pottery else "All pottery described"
        }

    def create_materials_table(self, materials_data):
        """Create formatted materials table"""
        headers = ['Inv. Number', 'Type', 'Definition', 'Description', 'Dating', 'Conservation']
        rows = []
        for material in materials_data:
            rows.append([
                material.get('numero_inventario', ''),
                material.get('tipo_reperto', ''),
                material.get('definizione', ''),
                material.get('descrizione', ''),
                material.get('datazione', ''),
                material.get('stato_conservazione', '')
            ])
        return {'headers': headers, 'rows': rows}

    def create_pottery_table(self, pottery_data):
        """Create formatted pottery table"""
        headers = ['ID Number', 'Form', 'Fabric', 'Ware', 'Conservation %', 'Notes']
        rows = []
        for pottery in pottery_data:
            rows.append([
                pottery.get('id_number', ''),
                pottery.get('form', ''),
                pottery.get('fabric', ''),
                pottery.get('ware', ''),
                pottery.get('percent', ''),
                pottery.get('note', '')
            ])
        return {'headers': headers, 'rows': rows}

    def format_us_data(self, us_data):
        """Format US data for the report"""
        formatted = ""
        for us in us_data:
            formatted += (
                f"IDUS: {us['id_us']}\n"
                f"US: {us['us']}\n"
                f"Area: {us['area']}\n"
                f"Settore: {us['settore']}\n"
                f"Definizione stratigrafica: {us['d_stratigrafica']}\n"
                f"Descrizione: {us['descrizione']}\n"
                f"Interpretazione: {us['interpretazione']}\n"
                f"Rapporti stratigrafici: {us['rapporti']}\n\n"
            )
        return formatted.strip()

    def format_material_data(self, material_data):
        """Format material data for the report"""
        formatted = ""
        for item in material_data:
            formatted += (
                f"Numero inventario: {item['numero_inventario']}\n"
                f"Tipo: {item['tipo_reperto']}\n"
                f"Definizione: {item['definizione']}\n"
                f"Descrizione: {item['descrizione']}\n"
                f"Datazione: {item['datazione']}\n"
                f"Stato di conservazione: {item['stato_conservazione']}\n"
                f"Contesto: Area {item['area']}, US {item['us']}\n\n"
            )
        return formatted.strip()

    def format_pottery_data(self, pottery_data):
        """Format pottery data for the report"""
        formatted = ""
        for pottery in pottery_data:
            formatted += (
                f"ID Reperto: {pottery.get('id_number', '')}\n"  # Cambiato da id_rep a id_number
                f"Numero identificativo: {pottery.get('id_number', '')}\n"
                f"Contesto: Sito {pottery.get('sito', '')}, Area {pottery.get('area', '')}, "
                f"US {pottery.get('us', '')}, Settore {pottery.get('sector', '')}\n"
                f"Documentazione fotografica: {pottery.get('photo', '')}\n"
                f"Anno: {pottery.get('anno', '')}\n\n"
                f"CARATTERISTICHE TECNICHE:\n"
                f"Impasto: {pottery.get('fabric', '')}\n"
                f"Percentuale conservata: {pottery.get('percent', '')}\n"
                f"Materiale: {pottery.get('material', '')}\n"
                f"Forma: {pottery.get('form', '')}\n"
                f"Forma specifica: {pottery.get('specific_form', '')}\n"
                f"Tipo specifico: {pottery.get('specific_shape', '')}\n"
                f"Classe ceramica: {pottery.get('ware', '')}\n"
                f"Quantità: {pottery.get('qty', '')}\n\n"
                f"Note: {pottery.get('note', '')}\n\n"
                f"-------------------\n\n"
            )
        return formatted.strip()

    def format_tomba_data(self, tomba_data):
        """Format tomb data for the report"""
        formatted = ""
        for tomba in tomba_data:
            formatted += (
                f"Scheda TAF: {tomba.get('nr_scheda_taf', '')}\n"
                f"Struttura: {tomba.get('sigla_struttura', '')} {tomba.get('nr_struttura', '')}\n"
                f"Individuo: {tomba.get('nr_individuo', '')}\n"
                f"Area: {tomba.get('area', '')}\n"
                f"Rito: {tomba.get('rito', '')}\n"
                f"Descrizione: {tomba.get('descrizione_taf', '')}\n"
                f"Interpretazione: {tomba.get('interpretazione_taf', '')}\n\n"
                f"-------------------\n\n"
            )
        return formatted.strip()

    def format_periodizzazione_data(self, periodizzazione_data):
        """Format periodization data for the report"""
        formatted = ""
        for periodo in periodizzazione_data:
            formatted += (
                f"Periodo: {periodo.get('periodo', '')}\n"
                f"Fase: {periodo.get('fase', '')}\n"
                f"Cronologia iniziale: {periodo.get('cron_iniziale', '')}\n"
                f"Cronologia finale: {periodo.get('cron_finale', '')}\n"
                f"Descrizione: {periodo.get('descrizione', '')}\n\n"
                f"-------------------\n\n"
            )
        return formatted.strip()

    def format_struttura_data(self, struttura_data):
        """Format structure data for the report"""
        formatted = ""
        for struttura in struttura_data:
            formatted += (
                f"Struttura: {struttura.get('sigla_struttura', '')} {struttura.get('numero_struttura', '')}\n"
                f"Categoria: {struttura.get('categoria_struttura', '')}\n"
                f"Tipologia: {struttura.get('tipologia_struttura', '')}\n"
                f"Definizione: {struttura.get('definizione_struttura', '')}\n"
                f"Descrizione: {struttura.get('descrizione', '')}\n"
                f"Interpretazione: {struttura.get('interpretazione', '')}\n\n"
                f"-------------------\n\n"
            )
        return formatted.strip()

    def create_analysis_tools(self, report_data, site_data, us_data, materials_data, pottery_data):
        """Create analysis tools for the LangChain agent"""
        return [
            Tool(
                name="AnalisiContestoSito",
                func=lambda x: self.analyze_site_context(report_data),
                description="Analizza il contesto geografico, amministrativo e storico del sito"
            ),
            Tool(
                name="AnalisiAreaScavo",
                func=lambda x: self.analyze_stratigraphy_by_area(us_data),
                description="Analizza la suddivisione dei settori di scavo e la loro organizzazione"
            ),
            Tool(
                name="AnalisiRelazioniStratigrafiche",
                func=lambda x: self.analyze_stratigraphy_relationships(us_data),
                description="Analizza le relazioni stratigrafiche tra le US per ogni settore"
            ),
            Tool(
                name="AnalisiInterpretativaUS",
                func=lambda x: {us['us']: us['interpretazione'] for us in us_data},
                description="Analizza le interpretazioni delle singole US"
            ),
            Tool(
                name="AnalisiDescrittivaaUS",
                func=lambda x: {us['us']: us['descrizione'] for us in us_data},
                description="Analizza le descrizioni delle singole US"
            ),
            Tool(
                name="AnalisiMaterialiPerTipo",
                func=lambda x: self.analyze_materials_by_type(materials_data),
                description="Analizza i materiali raggruppati per tipologia"
            ),
            Tool(
                name="AnalisiCeramichePerClasse",
                func=lambda x: self.analyze_pottery_by_class(pottery_data),
                description="Analizza le ceramiche raggruppate per classe"
            ),
            Tool(
                name="AnalisiConservazione",
                func=lambda x: {
                    'materiali': [m.get('stato_conservazione', '') for m in materials_data],
                },
                description="Analizza lo stato di conservazione dei reperti"
            ),
            Tool(
                name="AnalisiPhotoPottery",
                func=lambda x: self.analyze_pottery_photo(pottery_data),
                description="Considera le foto della ceramica da associare nel testo"
            )
        ]

    def create_validation_tools(self, site_data, us_data, materials_data, pottery_data):
        """Create validation tools for the LangChain agent"""
        return [
            Tool(
                name="ValidazioneUS",
                func=lambda x: self.validate_us_description(us_data),
                description="Valida la completezza delle descrizioni delle US"
            ),
            Tool(
                name="ValidazioneMateriali",
                func=lambda x: self.validate_materials_description(materials_data),
                description="Valida la completezza delle descrizioni dei materiali"
            ),
            Tool(
                name="ValidazioneCeramica",
                func=lambda x: self.validate_pottery_description(pottery_data),
                description="Valida la completezza delle descrizioni della ceramica"
            ),
            Tool(
                name="CreaTabellaRinvenimenti",
                func=lambda x: self.create_materials_table(materials_data),
                description="Crea una tabella formattata dei rinvenimenti"
            ),
            Tool(
                name="CreaTabellaRepCeramici",
                func=lambda x: self.create_pottery_table(pottery_data),
                description="Crea una tabella formattata dei reperti ceramici"
            )
        ]

    def process_site_table(self, records, current_site, report_data):
        """Process site table records and update report data"""
        site_record = next((r for r in records if getattr(r, 'sito', '') == current_site), None)
        if site_record:
            report_data['Regione'] = getattr(site_record, 'regione', '')
            report_data['Provincia'] = f"Provincia: {getattr(site_record, 'provincia', '')}"
            report_data['Comune'] = f"Comune: {getattr(site_record, 'comune', '')}"
            report_data['Cantiere'] = current_site
            report_data['Collocazione cantiere'] = current_site

    def process_us_table(self, records, year_filter, us_start, us_end, us_data):
        """Process US table records and update US data"""
        filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'anno_scavo', 'us')
        for record in filtered_records:
            us_data.append({
                'id_us': getattr(record, 'id_us', ''),
                'us': getattr(record, 'us', ''),
                'area': getattr(record, 'area', ''),
                'settore': getattr(record, 'settore', ''),
                'd_stratigrafica': getattr(record, 'd_stratigrafica', ''),
                'descrizione': getattr(record, 'descrizione', ''),
                'interpretazione': getattr(record, 'interpretazione', ''),
                'rapporti': getattr(record, 'rapporti', '')
            })

    def process_materials_table(self, records, year_filter, us_start, us_end, materials_data):
        """Process materials table records and update materials data"""
        filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'years', 'us')
        for record in filtered_records:
            materials_data.append({
                'id_invmat': getattr(record, 'id_invmat', ''),
                'numero_inventario': getattr(record, 'numero_inventario', ''),
                'tipo_reperto': getattr(record, 'tipo_reperto', ''),
                'definizione': getattr(record, 'definizione', ''),
                'descrizione': getattr(record, 'descrizione', ''),
                'datazione': getattr(record, 'datazione', ''),
                'stato_conservazione': getattr(record, 'stato_conservazione', ''),
                'classe_materiale': getattr(record, 'classe_materiale', ''),
                'area': getattr(record, 'area', ''),
                'us': getattr(record, 'us', '')
            })

    def process_pottery_table(self, records, year_filter, us_start, us_end, pottery_data):
        """Process pottery table records and update pottery data"""
        filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'anno', 'us')
        for record in filtered_records:
            pottery_data.append({
                'id_rep': getattr(record, 'id_number', ''),
                'id_number': getattr(record, 'id_number', ''),
                'sito': getattr(record, 'sito', ''),
                'area': getattr(record, 'area', ''),
                'us': getattr(record, 'us', ''),
                'sector': getattr(record, 'sector', ''),
                'photo': getattr(record, 'photo', ''),
                'anno': getattr(record, 'anno', ''),
                'fabric': getattr(record, 'fabric', ''),
                'percent': getattr(record, 'percent', ''),
                'material': getattr(record, 'material', ''),
                'form': getattr(record, 'form', ''),
                'specific_form': getattr(record, 'specific_form', ''),
                'specific_shape': getattr(record, 'specific_shape', ''),
                'ware': getattr(record, 'ware', ''),
                'qty': getattr(record, 'qty', ''),
                'note': getattr(record, 'note', '')
            })

    def filter_records(self, records, year_filter, start, end, year_field, range_field):
        """Filter records based on year or range criteria

        year_filter can be a single year or multiple years separated by commas
        e.g. "2024" or "2024, 2025, 2026" or "2024,2025,2026"
        """
        if year_filter:
            # Split the year filter by comma and strip whitespace
            years = [y.strip() for y in year_filter.split(',')]
            # Filter records where the year field matches any of the specified years
            return [r for r in records if str(getattr(r, year_field, '')) in years]
        elif start and end:
            return [r for r in records if start <= str(getattr(r, range_field, '')) <= end]
        return records

    def initialize_report_data(self):
        """Initialize the report data dictionary with empty values"""
        return {
            'Regione': '', 'Provincia': '', 'Comune': '', 'Ente di riferimento': '',
            'Committenza': '', 'Direzione scientifica': '', 'Elaborato a cura di': '',
            'Direttore cantiere': '', 'Cantiere': '', 'Tipo di indagine': '',
            'Titolo elaborato': '', 'Collocazione cantiere': '', 'Periodo cantiere': '',
            'Intervento': '', 'Progettazione lavori': '', 'Direzione lavori': '',
            'Direzione scientifica indagini archeologiche': '',
            'Esecuzione indagini archeologiche': '', 'Direzione cantiere archeologico': '',
            'Archeologi': ''
        }

    def process_table_data(self, table_name, records, current_site, year_filter,
                          us_start, us_end, report_data, us_data,
                          materials_data, pottery_data, tomba_data=None, 
                          periodizzazione_data=None, struttura_data=None):
        """Process table data and update corresponding data structures"""
        if tomba_data is None:
            tomba_data = []
        if periodizzazione_data is None:
            periodizzazione_data = []
        if struttura_data is None:
            struttura_data = []

        if table_name == 'site_table':
            self.process_site_table(records, current_site, report_data)
        elif table_name == 'us_table':
            filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'anno_scavo', 'us')
            for record in filtered_records:
                us_data.append({
                    'id_us': getattr(record, 'id_us', ''),
                    'us': getattr(record, 'us', ''),
                    'area': getattr(record, 'area', ''),
                    'settore': getattr(record, 'settore', ''),
                    'd_stratigrafica': getattr(record, 'd_stratigrafica', ''),
                    'descrizione': getattr(record, 'descrizione', ''),
                    'interpretazione': getattr(record, 'interpretazione', ''),
                    'rapporti': getattr(record, 'rapporti', '')
                })
        elif table_name == 'inventario_materiali_table':
            filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'years', 'us')
            for record in filtered_records:
                materials_data.append({
                    'id_invmat': getattr(record, 'id_invmat', ''),
                    'numero_inventario': getattr(record, 'numero_inventario', ''),
                    'tipo_reperto': getattr(record, 'tipo_reperto', ''),
                    'definizione': getattr(record, 'definizione', ''),
                    'descrizione': getattr(record, 'descrizione', ''),
                    'datazione': getattr(record, 'datazione', ''),
                    'stato_conservazione': getattr(record, 'stato_conservazione', ''),
                    'area': getattr(record, 'area', ''),
                    'us': getattr(record, 'us', '')
                })
        elif table_name == 'pottery_table':
            filtered_records = self.filter_records(records, year_filter, us_start, us_end, 'anno', 'us')
            for record in filtered_records:
                pottery_data.append({
                    'id_rep': getattr(record, 'id_rep', ''),
                    'id_number': getattr(record, 'id_number', ''),
                    'fabric': getattr(record, 'fabric', ''),
                    'form': getattr(record, 'form', ''),
                    'ware': getattr(record, 'ware', ''),
                    'percent': getattr(record, 'percent', ''),
                    'note': getattr(record, 'note', ''),
                    'area': getattr(record, 'area', ''),
                    'us': getattr(record, 'us', '')
                })
        elif table_name == 'tomba_table':
            # Filter records by site
            records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

            for record in records:
                tomba_data.append({
                    'nr_scheda_taf': getattr(record, 'nr_scheda_taf', ''),
                    'sigla_struttura': getattr(record, 'sigla_struttura', ''),
                    'nr_struttura': getattr(record, 'nr_struttura', ''),
                    'nr_individuo': getattr(record, 'nr_individuo', ''),
                    'rito': getattr(record, 'rito', ''),
                    'descrizione_taf': getattr(record, 'descrizione_taf', ''),
                    'interpretazione_taf': getattr(record, 'interpretazione_taf', ''),
                    'area': getattr(record, 'area', '')
                })

        elif table_name == 'periodizzazione_table':
            # Filter records by site
            records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

            for record in records:
                periodizzazione_data.append({
                    'periodo': getattr(record, 'periodo', ''),
                    'fase': getattr(record, 'fase', ''),
                    'cron_iniziale': getattr(record, 'cron_iniziale', ''),
                    'cron_finale': getattr(record, 'cron_finale', ''),
                    'descrizione': getattr(record, 'descrizione', '')
                })

        elif table_name == 'struttura_table':
            # Filter records by site
            records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

            for record in records:
                struttura_data.append({
                    'sigla_struttura': getattr(record, 'sigla_struttura', ''),
                    'numero_struttura': getattr(record, 'numero_struttura', ''),
                    'categoria_struttura': getattr(record, 'categoria_struttura', ''),
                    'tipologia_struttura': getattr(record, 'tipologia_struttura', ''),
                    'definizione_struttura': getattr(record, 'definizione_struttura', ''),
                    'descrizione': getattr(record, 'descrizione', ''),
                    'interpretazione': getattr(record, 'interpretazione', '')
                })
    def create_system_message(self):
        """Create the system message for the LangChain agent"""
        return SystemMessage(content="""Sei un esperto archeologo specializzato in relazioni di scavo stratigrafiche.
        Il tuo compito è generare un report dettagliato e approfondito che deve:

        1. Per la stratigrafia:
        - Descrivere le relazioni tra US in modo discorsivo e fluido
        - Interpretare la sequenza stratigrafica
        - Analizzare ogni area di scavo separatamente
        - Collegare le diverse aree in una narrazione coerente

        2. Per i materiali:
        - Creare elenchi dettagliati per tipo
        - Fornire descrizioni approfondite
        - Analizzare gli aspetti tecnici
        - Discutere il significato cronologico e culturale

        3. Per le ceramiche:
        - Catalogare per classe e tipo
        - Descrivere aspetti tecnici e morfologici
        - Analizzare impasti e decorazioni
        - Fornire confronti tipologici

        Il report deve essere lungo almeno tre pagine e utilizzare un linguaggio tecnico appropriato.""")

    def create_custom_prompt(self, formatted_us_data, formatted_material_data,
                             formatted_pottery_data, report_data,
                             formatted_tomba_data="", formatted_periodizzazione_data="",
                             formatted_struttura_data=""):
        """Create the custom prompt for the report generation"""
        return f"""
        Genera una relazione archeologica dettagliata e discorsiva in italiano basata sui seguenti dati. 
        La relazione deve essere un testo fluido e narrativo, che analizza in profondità tutte le unità 
        stratigrafiche (US) e le mette in relazione tra loro.

        ISTRUZIONI IMPORTANTI:
        - Analizza TUTTI i dati forniti senza tralasciare nulla
        - NON inserire note personali o pensieri come "Nota: L'elenco completo delle US è disponibile..."
        - NON scrivere frasi come "... (continua per tutte le fasi/US)"
        - NON menzionare limiti di token o di spazio
        - Crea un testo discorsivo che integri tutti i dati in modo fluido
        - Evita di elencare singolarmente tutte le US, ma creane una narrazione coerente
        - Evidenzia le US più significative e le loro relazioni
        - Quando descrivi le fasi, integra tutte le informazioni disponibili senza rimandare ad altri documenti

        1. INTRODUZIONE:
           - Fornisci una panoramica dettagliata della campagna di scavo basandoti su:
        {formatted_us_data}   
        {formatted_material_data}
        {formatted_pottery_data}
        {formatted_periodizzazione_data}

        2. DESCRIZIONE METODOLOGICA ED ESITO DELL'INDAGINE:
           - Inizia con una descrizione generale dell'area di scavo e della sua suddivisione in settori.
           - Di seguito, descrivi in dettaglio le unità stratigrafiche come indicato:
        {formatted_us_data}   

        3. DESCRIZIONE DELLE STRUTTURE E TOMBE:
           - Se presenti, descrivi le strutture archeologiche rinvenute:
        {formatted_struttura_data}
           - Se presenti, descrivi le tombe e i contesti funerari:
        {formatted_tomba_data}

        4. PERIODIZZAZIONE:
           - Se presenti, descrivi le fasi cronologiche identificate:
        {formatted_periodizzazione_data}

        5. DESCRIZIONE DEI MATERIALI:   
           - Se presenti inserisci l'elenco dei materiali o delle ceramiche rinvenuti in ogni unità stratigrafica.
           - Se presenti, descrivi i materiali e/o le ceramiche in dettaglio come indicato.
        {formatted_material_data}
        {formatted_pottery_data}

        6. CONCLUSIONI:
           - Sintetizza i risultati principali dello scavo.
           - Discuti delle unità stratigrafiche, strutture, tombe e materiali rinvenuti.
           - Concludi con una sintesi delle principali scoperte e la loro interpretazione cronologica.

        Dati del sito: {report_data}
        """

    def check_missing_data(self):
        """Verifica e riporta i dati mancanti in tutte le tabelle selezionate"""
        missing_data_report = []

        # Get selected tables
        dialog = ReportGeneratorDialog(self)
        #if dialog.exec_() == QDialog.Accepted:
        selected_tables = dialog.get_selected_tables()
        year_filter = dialog.get_year_filter()
        us_start, us_end = dialog.get_us_range()

        # Connect to database
        conn = Connection()
        db_url = conn.conn_str()
        current_site = str(self.comboBox_sito.currentText())
        site_data=[]
        us_data = []
        materials_data = []
        pottery_data = []

        # Initialize data structures for new tables
        tomba_data = []
        periodizzazione_data = []
        struttura_data = []

        # Fetch data
        for table_name in selected_tables:
            records, columns = ReportGenerator.read_data_from_db(db_url, table_name)

            if table_name == 'us_table':
                # Filter records
                if year_filter:
                    records = [r for r in records if str(getattr(r, 'anno_scavo', '')) == year_filter]
                elif us_start and us_end:
                    records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

                for record in records:
                    us_data.append({
                        'us': getattr(record, 'us', ''),
                        'area': getattr(record, 'area', ''),
                        'd_stratigrafica': getattr(record, 'd_stratigrafica', ''),
                        'descrizione': getattr(record, 'descrizione', ''),
                        'interpretazione': getattr(record, 'interpretazione', ''),
                        'rapporti': getattr(record, 'rapporti', '')
                    })

            elif table_name == 'inventario_materiali_table':
                # Filter records
                if year_filter:
                    records = [r for r in records if str(getattr(r, 'years', '')) == year_filter]
                elif us_start and us_end:
                    records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

                for record in records:
                    materials_data.append({
                        'id_invmat': getattr(record, 'id_invmat', ''),
                        'numero_inventario': getattr(record, 'numero_inventario', ''),
                        'tipo_reperto': getattr(record, 'tipo_reperto', ''),
                        'definizione': getattr(record, 'definizione', ''),
                        'descrizione': getattr(record, 'descrizione', ''),
                        'datazione': getattr(record, 'datazione', ''),
                        'stato_conservazione': getattr(record, 'stato_conservazione', ''),
                        'area': getattr(record, 'area', ''),
                        'us': getattr(record, 'us', '')
                    })

            elif table_name == 'pottery_table':
                # Filter records
                if year_filter:
                    records = [r for r in records if str(getattr(r, 'anno', '')) == year_filter]
                elif us_start and us_end:
                    records = [r for r in records if us_start <= str(getattr(r, 'us', '')) <= us_end]

                for record in records:
                    pottery_data.append({
                        'id_rep': getattr(record, 'id_rep', ''),
                        'id_number': getattr(record, 'id_number', ''),
                        'fabric': getattr(record, 'fabric', ''),
                        'form': getattr(record, 'form', ''),
                        'area': getattr(record, 'area', ''),
                        'us': getattr(record, 'us', '')
                    })

            elif table_name == 'tomba_table':
                # Filter records by site
                records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

                for record in records:
                    tomba_data.append({
                        'nr_scheda_taf': getattr(record, 'nr_scheda_taf', ''),
                        'sigla_struttura': getattr(record, 'sigla_struttura', ''),
                        'nr_struttura': getattr(record, 'nr_struttura', ''),
                        'nr_individuo': getattr(record, 'nr_individuo', ''),
                        'rito': getattr(record, 'rito', ''),
                        'descrizione_taf': getattr(record, 'descrizione_taf', ''),
                        'interpretazione_taf': getattr(record, 'interpretazione_taf', ''),
                        'area': getattr(record, 'area', '')
                    })

            elif table_name == 'periodizzazione_table':
                # Filter records by site
                records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

                for record in records:
                    periodizzazione_data.append({
                        'periodo': getattr(record, 'periodo', ''),
                        'fase': getattr(record, 'fase', ''),
                        'cron_iniziale': getattr(record, 'cron_iniziale', ''),
                        'cron_finale': getattr(record, 'cron_finale', ''),
                        'descrizione': getattr(record, 'descrizione', '')
                    })

            elif table_name == 'struttura_table':
                # Filter records by site
                records = [r for r in records if str(getattr(r, 'sito', '')) == current_site]

                for record in records:
                    struttura_data.append({
                        'sigla_struttura': getattr(record, 'sigla_struttura', ''),
                        'numero_struttura': getattr(record, 'numero_struttura', ''),
                        'categoria_struttura': getattr(record, 'categoria_struttura', ''),
                        'tipologia_struttura': getattr(record, 'tipologia_struttura', ''),
                        'definizione_struttura': getattr(record, 'definizione_struttura', ''),
                        'descrizione': getattr(record, 'descrizione', ''),
                        'interpretazione': getattr(record, 'interpretazione', '')
                    })

            # Create report thread instance for validation methods
            report_thread = GenerateReportThread(
                custom_prompt="",
                descriptions_text="",
                api_key="",
                selected_model="",
                selected_tables=selected_tables,
                analysis_steps=[],
                agent=None,
                us_data=us_data,
                materials_data=materials_data,
                pottery_data=pottery_data,
                site_data=site_data,
                py_dialog=self,
                tomba_data=tomba_data,
                periodizzazione_data=periodizzazione_data,
                struttura_data=struttura_data
            )

            # Check each type of data
            if 'us_table' in selected_tables:
                us_validation = report_thread.validate_us()
                if not us_validation['valid']:
                    missing_data_report.append(us_validation['message'])

            if 'inventario_materiali_table' in selected_tables:
                materials_validation = report_thread.validate_materials()
                if not materials_validation['valid']:
                    missing_data_report.append(materials_validation['message'])

            if 'pottery_table' in selected_tables:
                pottery_validation = report_thread.validate_pottery()
                if not pottery_validation['valid']:
                    missing_data_report.append(pottery_validation['message'])

            if 'tomba_table' in selected_tables:
                tomba_validation = report_thread.validate_tomba()
                if not tomba_validation['valid']:
                    missing_data_report.append(tomba_validation['message'])

            if 'periodizzazione_table' in selected_tables:
                periodizzazione_validation = report_thread.validate_periodizzazione()
                if not periodizzazione_validation['valid']:
                    missing_data_report.append(periodizzazione_validation['message'])

            if 'struttura_table' in selected_tables:
                struttura_validation = report_thread.validate_struttura()
                if not struttura_validation['valid']:
                    missing_data_report.append(struttura_validation['message'])

            # Show results
            if missing_data_report:
                msg = "REPORT DATI MANCANTI\n\n" + "\n\n".join(missing_data_report)
                QMessageBox.warning(self, "Dati Mancanti", msg)
                return False
            else:
                QMessageBox.information(self, "Verifica Completata", "Tutti i dati sono completi!")
                return True

        return None
    def generate_and_display_report(self):
        # Prima verifica la presenza di dati mancanti
        if not self.check_missing_data():
            reply = QMessageBox.question(
                self,
                "Dati Mancanti",
                "Sono stati rilevati dati mancanti. Vuoi procedere comunque con la generazione del report?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            if reply == QMessageBox.No:
                return

        # Procedi con la generazione del report
        dialog = ReportGeneratorDialog(self)


        if dialog.exec_() == QDialog.Accepted:




            # Step 2: Initialize database connection and get selected data
            conn = Connection()
            db_url = conn.conn_str()
            selected_tables = dialog.get_selected_tables()
            year_filter = dialog.get_year_filter()
            us_start, us_end = dialog.get_us_range()

            # Step 3: Prepare data structures
            report_data = self.initialize_report_data()
            current_site = str(self.comboBox_sito.currentText())
            site_data =[]
            us_data = []
            materials_data = []
            pottery_data = []
            tomba_data = []
            periodizzazione_data = []
            struttura_data = []

            # Step 4: Fetch and process data from selected tables
            for table_name in selected_tables:
                records, columns = ReportGenerator.read_data_from_db(db_url, table_name)
                self.process_table_data(
                    table_name, records, current_site, year_filter,
                    us_start, us_end, report_data, us_data,
                    materials_data, pottery_data, tomba_data,
                    periodizzazione_data, struttura_data
                )

            # Step 5: Format data for report
            formatted_us_data = self.format_us_data(us_data)
            formatted_material_data = self.format_material_data(materials_data)
            formatted_pottery_data = self.format_pottery_data(pottery_data)
            formatted_tomba_data = self.format_tomba_data(tomba_data)
            formatted_periodizzazione_data = self.format_periodizzazione_data(periodizzazione_data)
            formatted_struttura_data = self.format_struttura_data(struttura_data)

            # Crea il testo delle descrizioni combinando tutti i dati formattati
            descriptions_text = f"""
            UNITÀ STRATIGRAFICHE:
            {formatted_us_data}

            MATERIALI:
            {formatted_material_data}

            CERAMICHE:
            {formatted_pottery_data}

            TOMBE:
            {formatted_tomba_data}

            PERIODIZZAZIONE:
            {formatted_periodizzazione_data}

            STRUTTURE:
            {formatted_struttura_data}
            """
            # Step 6: Create custom prompt
            custom_prompt = self.create_custom_prompt(
                formatted_us_data, formatted_material_data,
                formatted_pottery_data, report_data,
                formatted_tomba_data, formatted_periodizzazione_data,
                formatted_struttura_data
            )

            # Step 7: Check internet connection and start report generation
            if not ReportGenerator.is_connected():
                QMessageBox.warning(self, "Connessione Assente", "Nessuna connessione a Internet rilevata.")
                return

            # Step 8: Show progress dialog
            # self.progress_dialog = QProgressDialog("Generazione del report in corso...", None, 0, 0, self)
            # self.progress_dialog.setWindowModality(Qt.WindowModal)
            # self.progress_dialog.setCancelButton(None)
            # self.progress_dialog.setRange(0, 0)
            # self.progress_dialog.show()

            try:
                # Step 9: Initialize OpenAI components
                api_key = self.apikey_gpt()
                tools = self.create_analysis_tools(report_data, site_data, us_data, materials_data, pottery_data)
                enhanced_tools = self.create_validation_tools(site_data, us_data, materials_data, pottery_data)

                # Step 10: Set up LangChain components
                llm = ChatOpenAI(
                    temperature=0.0,
                    model_name="gpt-4.1",
                    api_key=api_key,
                    max_tokens=16000,
                    streaming=True
                )

                system_message = self.create_system_message()
                memory = ConversationBufferMemory(
                    memory_key="chat_history",
                    return_messages=True,
                    system_message=system_message
                )

                agent = initialize_agent(
                    tools + enhanced_tools,  # Combine both tool sets
                    llm,
                    agent=AgentType.CHAT_CONVERSATIONAL_REACT_DESCRIPTION,
                    memory=memory,
                    system_message=system_message,
                    verbose=True
                )

                # Step 11: Create and show report dialog
                self.report_dialog = ReportDialog(parent=self)
                self.report_dialog.show()

                # Create and start the report generation thread
                self.report_thread = GenerateReportThread(
                    custom_prompt=custom_prompt,
                    descriptions_text=descriptions_text,
                    api_key=api_key,
                    selected_model="gpt-4.1",
                    selected_tables=selected_tables,
                    analysis_steps=self.analysis_steps,
                    agent=agent,
                    us_data=us_data,
                    materials_data=materials_data,
                    pottery_data=pottery_data,
                    site_data=site_data,
                    py_dialog=self,
                    output_language=dialog.get_selected_language(),
                    tomba_data=tomba_data,
                    periodizzazione_data=periodizzazione_data,
                    struttura_data=struttura_data
                )

                # Step 13: Connect signals and start thread
                self.report_thread.report_generated.connect(self.report_dialog.update_content)
                self.report_thread.log_message.connect(self.report_dialog.log_to_terminal)
                self.report_thread.report_completed.connect(self.on_report_generated)  # Aggiungi questa linea
                self.report_thread.start()


            except Exception as e:
                #self.progress_dialog.close()
                if hasattr(self, 'report_dialog'):
                    self.report_dialog.close()
                error_msg = f"Errore durante la generazione del report: {str(e)}"
                self.show_error(error_msg, 'Errore Generazione Report')
                return

    def on_report_generated(self, report_text, report_data):
        print("\n==== REPORT GENERATION PROCESS STARTED ====")
        print(f"Report text length: {len(report_text) if report_text else 0} characters")
        print(f"Report data keys: {', '.join(report_data.keys())}")

        # Check if report is empty
        if not report_text or report_text.strip() == "":
            print("WARNING: Empty report text received")
            QMessageBox.warning(
                self,
                "Report Vuoto",
                "Il report generato è vuoto. Questo potrebbe essere dovuto alla mancanza di tabelle selezionate necessarie per la generazione del report."
            )
            return

        # First show template choice dialog
        response = QMessageBox.question(
            self,
            "Scelta Template",
            "Vuoi salvare il contenuto usando il template predefinito?",
            QMessageBox.Yes | QMessageBox.No
        )

        # If user selects No or closes dialog, just return
        if response != QMessageBox.Yes:
            print("User chose not to use template, exiting")
            return

        # Get save location
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "Salva Report",
            "",
            "Word Documents (*.docx);;All Files (*)"
        )

        if not output_path:
            print("No output path selected, exiting")
            return

        if not output_path.lower().endswith('.docx'):
            output_path += '.docx'

        print(f"Output path: {output_path}")

        # Parse the report text into sections
        print("Parsing report text into sections...")
        sections = self.parse_report_into_sections(report_text)

        # Update report data with section content from the generated report
        if sections:
            print(f"Successfully parsed {len(sections)} sections")
            report_data.update(sections)
        else:
            print("WARNING: No sections parsed from report text, using fallback content")
            # Fallback to using analysis steps if parsing fails
            section_content = {
                "introduzione": ArchaeologicalAnalysis().get_introduction_step()['prompt'],
                "descrizione_metodologica_ed_esito": "Questa sezione contiene la metodologia di scavo, l'analisi stratigrafica e la descrizione dei materiali.",
                "descrizione_metodologica": ArchaeologicalAnalysis().get_methodology_step()['prompt'],
                "analisi_stratigrafica": ArchaeologicalAnalysis().get_stratigraphy_step()['prompt'],
                "descrizione_materiali": ArchaeologicalAnalysis().get_materials_step()['prompt'],
                "conclusioni": ArchaeologicalAnalysis().get_conclusions_step()['prompt']
            }
            report_data.update(section_content)
            print("Added fallback content for sections")

        # Ensure all required sections have at least some default content
        required_sections = [
            "introduzione", 
            "descrizione_metodologica_ed_esito", 
            "descrizione_metodologica", 
            "analisi_stratigrafica", 
            "descrizione_materiali", 
            "conclusioni"
        ]

        for section in required_sections:
            if section not in report_data or not report_data[section] or report_data[section].strip() == "":
                print(f"Adding default content for missing section: {section}")
                if section == "introduzione":
                    report_data[section] = "Introduzione al sito archeologico e alla campagna di scavo."
                elif section == "descrizione_metodologica_ed_esito":
                    report_data[section] = "Descrizione metodologica ed esito dell'indagine archeologica."
                elif section == "descrizione_metodologica":
                    report_data[section] = "Metodologia di scavo utilizzata durante la campagna archeologica."
                elif section == "analisi_stratigrafica":
                    report_data[section] = "Analisi stratigrafica e interpretazione delle unità stratigrafiche."
                elif section == "descrizione_materiali":
                    report_data[section] = "Catalogo e descrizione dei materiali rinvenuti durante lo scavo."
                elif section == "conclusioni":
                    report_data[section] = "Conclusioni e interpretazione complessiva del sito archeologico."

        # Add tables data
        print("Adding tables data...")
        materials_table = self._format_materials_table()
        if materials_table:
            report_data['materials_table'] = materials_table
            print(f"Added materials table with {len(materials_table)} rows")
        else:
            print("No materials table data available")

        pottery_table = self._format_pottery_table()
        if pottery_table:
            report_data['pottery_table'] = pottery_table
            print(f"Added pottery table with {len(pottery_table)} rows")
        else:
            print("No pottery table data available")

        try:
            template_path = os.path.join(self.HOME, "bin", "template_report_adarte.docx")
            print(f"Template path: {template_path}")

            if not os.path.exists(template_path):
                print(f"ERROR: Template not found at {template_path}")
                QMessageBox.warning(self, "Errore", f"Template non trovato in {template_path}!")
                return

            # Try to open the template to verify it's a valid Word document
            try:
                test_doc = Document(template_path)
                print(f"Template opened successfully, contains {len(test_doc.paragraphs)} paragraphs")
            except Exception as doc_error:
                print(f"ERROR: Failed to open template: {str(doc_error)}")
                QMessageBox.warning(self, "Errore", f"Il template non è un documento Word valido: {str(doc_error)}")
                return

            print("Saving report to template...")
            self.save_report_to_template(report_data, template_path, output_path)
            print("Report saved successfully")
            QMessageBox.information(self, "Successo", f"Report salvato in {output_path}")

        except Exception as e:
            print(f"ERROR: Exception during save process: {str(e)}")
            import traceback
            print(traceback.format_exc())
            QMessageBox.critical(self, "Errore", f"Errore nel salvataggio: {str(e)}")

        print("==== REPORT GENERATION PROCESS COMPLETED ====\n")

    def parse_report_into_sections(self, report_text):
        """Parse the report text into sections based on headings"""
        if not report_text:
            print("Warning: Empty report text passed to parse_report_into_sections")
            return {}

        sections = {}
        current_section = None
        current_content = []

        # Define the main section headings to look for
        section_headings = {
            "INTRODUZIONE": "introduzione",
            "DESCRIZIONE METODOLOGICA ED ESITO DELL'INDAGINE": "descrizione_metodologica_ed_esito",
            "METODOLOGIA DI SCAVO": "descrizione_metodologica",
            "ANALISI STRATIGRAFICA E INTERPRETAZIONE": "analisi_stratigrafica",
            "CATALOGO DEI MATERIALI": "descrizione_materiali",
            "CONCLUSIONI": "conclusioni"
        }

        # Process the report line by line
        lines = report_text.split('\n')
        print(f"Parsing report with {len(lines)} lines")

        # Print the first few lines to help debug
        print("First 10 lines of report:")
        for i in range(min(10, len(lines))):
            print(f"Line {i+1}: {lines[i]}")

        for i, line in enumerate(lines):
            # Check if this line is a main section heading
            is_heading = False
            for heading, key in section_headings.items():
                # Check for heading with or without markdown formatting
                if heading in line:
                    print(f"Found section heading: '{heading}' at line {i+1}")
                    # If we were already processing a section, save it
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                        print(f"Saved section '{current_section}' with {len(current_content)} lines")

                    # Start a new section
                    current_section = key
                    current_content = []
                    is_heading = True
                    break

            # If not a heading and we're in a section, add the line to current content
            if not is_heading and current_section:
                current_content.append(line)

        # Save the last section if there is one
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
            print(f"Saved final section '{current_section}' with {len(current_content)} lines")

        print(f"Parsed {len(sections)} sections: {', '.join(sections.keys())}")
        return sections

    def save_report_as_plain_doc(self, report_text, output_path):
        doc = Document()
        doc.add_paragraph(report_text)
        doc.save(output_path)

    def process_section_content(self, doc, section_info, content, report_data):
        """
        Process the content of a section and add it to the document.

        Args:
            doc: Word document
            section_info: Dictionary with section information
            content: Text content of the section
            report_data: Dictionary with all report data
        """
        if not content or content.strip() == "":
            print(f"WARNING: Empty content for section '{section_info['key']}', adding placeholder text")
            # Add a note that this section has no data
            doc.add_heading(section_info["heading"], level=1)
            doc.add_paragraph("Questa sezione contiene informazioni limitate a causa della mancanza di dati completi. Si consiglia di verificare i dati di input e ripetere la generazione del report con dati più completi.")

            # Add a basic description based on the section type
            if section_info["key"] == "introduzione":
                doc.add_paragraph("Introduzione al sito archeologico e alla campagna di scavo.")
            elif section_info["key"] == "descrizione_metodologica_ed_esito":
                doc.add_paragraph("Descrizione metodologica ed esito dell'indagine archeologica.")
            elif section_info["key"] == "conclusioni":
                doc.add_paragraph("Conclusioni e interpretazione complessiva del sito archeologico.")

            # If this section has subsections, process them
            if "subsections" in section_info:
                for subsection in section_info["subsections"]:
                    subsection_key = subsection["key"]
                    subsection_heading = subsection["heading"]
                    subsection_content = report_data.get(subsection_key, "")

                    # Add subsection heading
                    doc.add_heading(subsection_heading, level=2)

                    # Add default content if subsection is empty
                    if not subsection_content or subsection_content.strip() == "":
                        if subsection_key == "descrizione_metodologica":
                            doc.add_paragraph("Metodologia di scavo utilizzata durante la campagna archeologica.")
                        elif subsection_key == "analisi_stratigrafica":
                            doc.add_paragraph("Analisi stratigrafica e interpretazione delle unità stratigrafiche.")
                        elif subsection_key == "descrizione_materiali":
                            doc.add_paragraph("Catalogo e descrizione dei materiali rinvenuti durante lo scavo.")
                    else:
                        # Process subsection content
                        self._process_content_lines(doc, subsection_content.split('\n'))
            return

        # Add section heading
        doc.add_heading(section_info["heading"], level=1)

        # Check if this section has subsections
        if "subsections" in section_info:
            # Process each subsection
            for subsection in section_info["subsections"]:
                subsection_key = subsection["key"]
                subsection_heading = subsection["heading"]
                subsection_content = report_data.get(subsection_key, "")

                # Skip empty subsections
                if not subsection_content or subsection_content.strip() == "":
                    print(f"WARNING: Empty content for subsection '{subsection_key}', adding placeholder text")
                    doc.add_heading(subsection_heading, level=2)
                    doc.add_paragraph("Questa sottosezione è stata omessa a causa della mancanza di dati necessari.")
                    continue

                # Add subsection heading
                doc.add_heading(subsection_heading, level=2)

                # Process subsection content
                self._process_content_lines(doc, subsection_content.split('\n'))
        else:
            # Process content line by line for sections without subsections
            self._process_content_lines(doc, content.split('\n'))

    def _process_content_lines(self, doc, lines):
        """
        Process content lines and add them to the document.

        Args:
            doc: Word document
            lines: List of text lines
        """
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Check for markdown headings (##, --, ==)
            if (line.startswith('#') or 
                (i+1 < len(lines) and (all(c == '=' for c in lines[i+1].strip()) or 
                                      all(c == '-' for c in lines[i+1].strip())))):

                heading_text = self.clean_heading(line)
                heading_level = 2  # Default to level 2 for subheadings

                # Determine heading level based on # count or underline character
                if line.startswith('#'):
                    # Count the number of # characters at the start of the line
                    hash_count = 0
                    for char in line:
                        if char == '#':
                            hash_count += 1
                        else:
                            break

                    # Map the heading level: # -> 1, ## -> 2, etc.
                    heading_level = min(4, hash_count)  # Limit to level 4 to avoid issues with level 5
                elif i+1 < len(lines):
                    if all(c == '=' for c in lines[i+1].strip()):
                        heading_level = 2  # = underline is level 2
                    elif all(c == '-' for c in lines[i+1].strip()):
                        heading_level = 3  # - underline is level 3

                # Use try-except to handle potential heading level issues
                try:
                    doc.add_heading(heading_text, level=heading_level)
                except Exception as e:
                    print(f"Error adding heading '{heading_text}' with level {heading_level}: {str(e)}")
                    # Fallback to a safe heading level (3)
                    doc.add_heading(heading_text, level=3)

                # Skip the underline if present
                if (i+1 < len(lines) and (all(c == '=' for c in lines[i+1].strip()) or 
                                         all(c == '-' for c in lines[i+1].strip()))):
                    i += 2
                else:
                    i += 1
                continue

            # Check for images (US, RAPPORTI, FASE, MATERIALE, CERAMICA)
            img_match = re.search(r'\[IMMAGINE\s+(?:US|RAPPORTI|FASE|REPERTO|CERAMICA)\s+[^:]*:\s+(.*?),\s*(.*?)\]', line)
            if img_match:
                img_path = img_match.group(1).strip()
                caption = img_match.group(2).strip()

                # Remove file:// prefix if present
                if img_path.startswith('file://'):
                    img_path = img_path[7:]

                # Ensure the path is absolute
                if not os.path.isabs(img_path):
                    # Try to find the image in common locations
                    possible_paths = [
                        img_path,  # Original path
                        os.path.join(self.HOME, img_path),  # Relative to HOME
                        os.path.join(self.HOME, "bin", img_path),  # Relative to bin directory
                        os.path.join(os.path.dirname(os.path.abspath(__file__)), img_path),  # Relative to current file
                    ]

                    for path in possible_paths:
                        if os.path.exists(path):
                            img_path = path
                            break

                print(f"Processing image: path={img_path}, caption={caption}")

                if os.path.exists(img_path):
                    try:
                        # Load the image to verify it's valid
                        from PIL import Image
                        try:
                            img = Image.open(img_path)
                            img.verify()  # Verify it's a valid image

                            # Add the image to the document
                            doc.add_picture(img_path, width=Inches(6))
                            caption_para = doc.add_paragraph(caption)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.style = 'Caption'
                        except Exception as img_error:
                            print(f"Invalid image file {img_path}: {str(img_error)}")
                            doc.add_paragraph(f"[Immagine non valida: {caption}]")
                    except Exception as e:
                        print(f"Error adding image {img_path}: {str(e)}")
                        doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                else:
                    print(f"Image file not found: {img_path}")
                    doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                i += 1
                continue

            # Check for markdown tables
            if line.startswith('|') and line.endswith('|'):
                # Collect table lines
                table_lines = [line]
                j = i + 1
                while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                    table_lines.append(lines[j].strip())
                    j += 1

                # Process table if we have at least one row
                if len(table_lines) >= 1:
                    # Use the convert_markdown_table function to create the table
                    self.convert_markdown_table(table_lines, doc)

                    # Add a paragraph after the table for spacing
                    doc.add_paragraph()

                i = j  # Skip processed table lines
                continue

            # Regular paragraph
            if line:
                doc.add_paragraph(line)
            i += 1

    def save_report_to_template(self, report_data, template_path, output_path):
        print(f"Starting save_report_to_template with template: {template_path}, output: {output_path}")
        print(f"Report data keys: {', '.join(report_data.keys())}")

        # Check for section keys
        section_keys = ['introduzione', 'descrizione_metodologica', 'analisi_stratigrafica', 'descrizione_materiali', 'conclusioni']
        for key in section_keys:
            if key in report_data:
                content = report_data[key]
                content_preview = content[:100] + "..." if len(content) > 100 else content
                print(f"Section '{key}' found with {len(content)} characters. Preview: {content_preview}")
            else:
                print(f"WARNING: Section '{key}' NOT found in report_data")

        doc = Document(template_path)

        # Define alternative placeholder patterns to check
        alternative_placeholders = {
            "{{INTRODUZIONE}}": ["INTRODUZIONE", "Introduzione", "introduzione"],
            "{{DESCRIZIONE_METODOLOGICA_ED_ESITO_DELL_INDAGINE}}": ["DESCRIZIONE METODOLOGICA", "Descrizione Metodologica", "METODOLOGIA", "Metodologia"],
            "{{CONCLUSIONI}}": ["CONCLUSIONI", "Conclusioni", "conclusioni"]
        }

        # Check if 'Table Grid' style exists, if not create it
        table_styles = [s for s in doc.styles if 'table' in s.name.lower()]
        has_table_grid = any('table grid' in s.name.lower() for s in table_styles)

        # If 'Table Grid' style doesn't exist, use a default table style or create a basic one
        table_style_name = 'Table Grid' if has_table_grid else 'Normal Table'
        if not has_table_grid and not any('normal table' in s.name.lower() for s in table_styles):
            table_style_name = 'Table'  # Fallback to most basic table style

        print(f"Using table style: {table_style_name}")
        print(f"Document has {len(doc.paragraphs)} paragraphs")

        # Mapping for first page metadata fields
        metadata_mapping = {
            "{{SITO}}": "sito",
            "{{AREA}}": "area",
            "{{REGIONE}}": "Regione",
            "{{PROVINCIA}}": "Provincia",
            "{{COMUNE}}": "Comune",
            "{{INDIRIZZO}}": "Indirizzo",
            "{{RESPONSABILE}}": "Responsabile",
            "{{DIRETTORE}}": "Direzione scientifica",
            "{{ENTE}}": "Ente di riferimento",
            "{{COMMITTENZA}}": "Committenza",
            "{{DATA}}": "Data",
            "{{ANNO}}": "Anno",
            "{{CANTIERE}}": "Cantiere",
            "{{COLLOCAZIONE}}": "Collocazione cantiere",
            "{{PERIODO}}": "Periodo cantiere",
            "{{AUTORE}}": "Autore"
        }

        # Mapping delle sezioni del report con le chiavi in report_data
        section_mapping = {
            "{{INTRODUZIONE}}": {
                "key": "introduzione",
                "heading": "INTRODUZIONE"
            },
            "{{DESCRIZIONE_METODOLOGICA_ED_ESITO_DELL_INDAGINE}}": {
                "key": "descrizione_metodologica_ed_esito",
                "heading": "DESCRIZIONE METODOLOGICA ED ESITO DELL'INDAGINE",
                "subsections": [
                    {
                        "key": "descrizione_metodologica",
                        "heading": "METODOLOGIA DI SCAVO"
                    },
                    {
                        "key": "analisi_stratigrafica",
                        "heading": "ANALISI STRATIGRAFICA E INTERPRETAZIONE"
                    },
                    {
                        "key": "descrizione_materiali",
                        "heading": "CATALOGO DEI MATERIALI"
                    }
                ]
            },
            "{{CONCLUSIONI}}": {
                "key": "conclusioni",
                "heading": "CONCLUSIONI"
            }
        }

        # Set document properties for formatting
        from docx.shared import Pt, Cm
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.text import WD_BREAK

        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # First, replace metadata fields on the first page
        for para in doc.paragraphs:
            for placeholder, field_key in metadata_mapping.items():
                if placeholder in para.text:
                    # Get the value from report_data, or use an empty string if not found
                    value = str(report_data.get(field_key, ""))

                    # Special formatting for CANTIERE (site name)
                    if placeholder == "{{CANTIERE}}":
                        value = value.upper()  # Convert to uppercase

                        # Replace the placeholder with empty text first
                        para.text = para.text.replace(placeholder, "")

                        # Then add the formatted text as a new run
                        run = para.add_run(value)
                        run.bold = True
                        run.font.name = "Cambria"
                        run.font.size = Pt(18)
                        continue

                    # Replace the placeholder with the actual value
                    para.text = para.text.replace(placeholder, value)

                    # Apply formatting if needed
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            # Set font for all text
                            run.font.name = "Cambria"
                            run.font.size = Pt(12)

        # Ensure the introduction starts from page 5 by adding page breaks
        # First, add a section break to start a new section for the content
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Add blank pages to ensure introduction starts on page 5
        # We already have 1 page (title page), so add 3 more blank pages
        for _ in range(3):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Now process section content
        for para_idx, para in enumerate(doc.paragraphs):
            found_in_this_para = False

            # First check for exact matches
            for placeholder, section_info in section_mapping.items():
                if placeholder in para.text:
                    found_in_this_para = True
                    print(f"Found placeholder {placeholder} in paragraph {para_idx}")

                    # Get the section content
                    section_key = section_info["key"]
                    section_content = report_data.get(section_key, "")

                    # Process the section content
                    self.process_section_content(doc, section_info, section_content, report_data)

                    # Remove the placeholder paragraph
                    p = doc.paragraphs[para_idx]._element
                    p.getparent().remove(p)

                    break

            # If no exact match, check for alternative patterns
            if not found_in_this_para:
                for placeholder, alternatives in alternative_placeholders.items():
                    section_info = section_mapping.get(placeholder)
                    if section_info:
                        for alt in alternatives:
                            if alt in para.text:
                                found_in_this_para = True
                                print(f"Found alternative placeholder {alt} in paragraph {para_idx}")

                                # Get the section content
                                section_key = section_info["key"]
                                section_content = report_data.get(section_key, "")

                                # Process the section content
                                self.process_section_content(doc, section_info, section_content, report_data)

                                # Remove the placeholder paragraph
                                p = doc.paragraphs[para_idx]._element
                                p.getparent().remove(p)

                                break

                    if found_in_this_para:
                        break

        # Add tables if they exist in report_data
        if 'materials_table' in report_data and report_data['materials_table']:
            self._add_table_to_doc(doc, "TABELLA DEI MATERIALI", report_data['materials_table'])

        if 'pottery_table' in report_data and report_data['pottery_table']:
            self._add_table_to_doc(doc, "TABELLA DELLA CERAMICA", report_data['pottery_table'])

        # Save the document
        try:
            doc.save(output_path)
            print(f"Document saved successfully to {output_path}")
            return True
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return False

    def convert_markdown_table(self, table_lines, doc):
        """
        Convert a list of markdown table lines to a Word table.

        Args:
            table_lines: List of strings representing markdown table lines
            doc: Word document to add the table to

        Returns:
            Word table object
        """
        print(f"Converting markdown table with {len(table_lines)} lines")

        # Parse table data
        rows = []
        header_row = None
        separator_found = False

        for idx, table_line in enumerate(table_lines):
            # Check if this is a separator line
            if re.match(r'^\|\s*[-:]+\s*\|', table_line):
                separator_found = True
                # If we found a separator, the previous line is the header
                if idx > 0 and header_row is None:
                    header_row = [cell.strip() for cell in table_lines[idx-1].strip('|').split('|')]
                continue

            # Process normal data rows
            cells = [cell.strip() for cell in table_line.strip('|').split('|')]
            # Skip empty cells
            if all(not cell for cell in cells):
                continue
            rows.append(cells)

        # If we have rows to process
        if rows:
            print(f"Table has {len(rows)} rows and max {max(len(row) for row in rows)} columns")

            # Determine the number of columns (max width of any row)
            max_cols = max(len(row) for row in rows)

            # Create Word table
            table = doc.add_table(rows=len(rows), cols=max_cols)

            # Check if 'Table Grid' style exists, if not use a default style
            table_styles = [s.name for s in doc.styles if hasattr(s, 'name') and 'table' in s.name.lower()]
            print(f"Available table styles: {', '.join(table_styles)}")

            # Try to find the best table style
            if 'Table Grid' in table_styles:
                table_style_name = 'Table Grid'
            elif 'Grid Table' in table_styles:
                table_style_name = 'Grid Table'
            elif any('grid' in s.lower() for s in table_styles):
                # Find the first style with 'grid' in the name
                table_style_name = next(s for s in table_styles if 'grid' in s.lower())
            elif 'Normal Table' in table_styles:
                table_style_name = 'Normal Table'
            elif len(table_styles) > 0:
                # Use the first available table style
                table_style_name = table_styles[0]
            else:
                table_style_name = 'Table'  # Fallback to most basic table style

            print(f"Using table style: {table_style_name}")

            try:
                table.style = table_style_name
            except Exception as e:
                print(f"Error setting table style: {str(e)}")
                # Try a different approach if setting style fails
                try:
                    # Apply a basic grid to the table
                    for cell in table._tbl.findall(".//w:tc", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                        tcPr = cell.find(".//w:tcPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                        if tcPr is None:
                            tcPr = parse_xml(r'<w:tcPr {0}/>'.format(nsdecls('w')))
                            cell.append(tcPr)

                        # Add borders to the cell
                        tcBorders = tcPr.find(".//w:tcBorders", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                        if tcBorders is None:
                            tcBorders = parse_xml(r'<w:tcBorders {0}/>'.format(nsdecls('w')))
                            tcPr.append(tcBorders)

                        # Define borders for all sides
                        for side in ['top', 'left', 'bottom', 'right']:
                            border = tcBorders.find(".//{0}:{1}".format('w', side))
                            if border is None:
                                border = parse_xml(r'<w:{0} {1} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'.format(side, nsdecls('w')))
                                tcBorders.append(border)
                except Exception as border_error:
                    print(f"Error applying borders: {str(border_error)}")

            # Apply table properties for better formatting
            try:
                table.autofit = True

                # Add XML properties for table grid
                tbl_pr = table._element.xpath('w:tblPr')[0]
                table_width = parse_xml(r'<w:tblW {} w:type="auto"/>'.format(nsdecls('w')))
                tbl_pr.append(table_width)

                # Add table grid
                tbl_grid = table._element.xpath('w:tblGrid')[0]
                for _ in range(max_cols):
                    grid_col = parse_xml(r'<w:gridCol {} w:w="0"/>'.format(nsdecls('w')))
                    tbl_grid.append(grid_col)
            except Exception as prop_error:
                print(f"Error setting table properties: {str(prop_error)}")

            # Fill table data
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < max_cols:  # Ensure we don't exceed column count
                        # Pad row_data if needed
                        if col_idx >= len(row_data):
                            cell_text = ""
                        else:
                            cell_text = row_data[col_idx]

                        # Set cell text
                        try:
                            cell = table.rows[row_idx].cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = cell_text
                            else:
                                p = cell.add_paragraph(cell_text)

                            # Make header row bold
                            if (row_idx == 0 and separator_found) or (header_row is not None and row_idx == 0):
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                    # Center align header cells
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as cell_error:
                            print(f"Error setting cell text at row {row_idx}, col {col_idx}: {str(cell_error)}")

            # Add a paragraph after the table for spacing
            doc.add_paragraph()

            return table

        return None

    def clean_heading(self, text):
        """
        Remove markdown heading markers (##, --, etc.) from text.

        Args:
            text: Text to clean

        Returns:
            Cleaned text
        """
        # First remove any leading # characters
        text = re.sub(r'^[\s#]*', '', text)
        text = re.sub(r'^[\s*]*', '', text)
        # Then remove any trailing -- or == characters
        text = re.sub(r'[\s\-=]*$', '', text)
        # Remove any remaining markdown formatting
        text = text.replace('--##', '').replace('--#', '').replace('--', '').replace('-**', '')
        return text.strip()

        # Define alternative placeholder patterns to check
        alternative_placeholders = {
            "{{INTRODUZIONE}}": ["INTRODUZIONE", "Introduzione", "introduzione"],
            "{{DESCRIZIONE_METODOLOGICA_ED_ESITO_DELL_INDAGINE}}": ["DESCRIZIONE METODOLOGICA", "Descrizione Metodologica", "METODOLOGIA", "Metodologia"],
            "{{CONCLUSIONI}}": ["CONCLUSIONI", "Conclusioni", "conclusioni"]
        }

        for para_idx, para in enumerate(doc.paragraphs):
            found_in_this_para = False

            # First check for exact matches
            for placeholder, section_info in section_mapping.items():
                if placeholder in para.text:
                    placeholder_found = True
                    found_in_this_para = True
                    print(f"Found exact placeholder '{placeholder}' in paragraph {para_idx+1}: '{para.text}'")
                    content = report_data.get(section_info["key"], "")

                    print(f"Content for '{section_info['key']}' has {len(content)} characters")

                    # Process this section
                    self.process_section_content(doc, section_info, content, report_data)
                    break

            # If no exact match found in this paragraph, check for alternative patterns
            if not found_in_this_para:
                for placeholder, section_info in section_mapping.items():
                    alternatives = alternative_placeholders.get(placeholder, [])
                    for alt in alternatives:
                        if alt in para.text:
                            placeholder_found = True
                            found_in_this_para = True
                            print(f"Found alternative placeholder '{alt}' for '{placeholder}' in paragraph {para_idx+1}: '{para.text}'")
                            content = report_data.get(section_info["key"], "")

                            print(f"Content for '{section_info['key']}' has {len(content)} characters")

                            # Process this section
                            self.process_section_content(doc, section_info, content, report_data)
                            break
                    if found_in_this_para:
                        break

                    # Handle empty sections with a default message
                    if not content or content.strip() == "":
                        print(f"WARNING: Empty content for section '{section_info['key']}', adding placeholder text")
                        # Add a note that this section has no data
                        doc.add_heading(section_info["heading"], level=1)
                        doc.add_paragraph("Questa sezione contiene informazioni limitate a causa della mancanza di dati completi. Si consiglia di verificare i dati di input e ripetere la generazione del report con dati più completi.")

                        # Add a basic description based on the section type
                        if section_info["key"] == "introduzione":
                            doc.add_paragraph("Introduzione al sito archeologico e alla campagna di scavo.")
                        elif section_info["key"] == "descrizione_metodologica_ed_esito":
                            doc.add_paragraph("Descrizione metodologica ed esito dell'indagine archeologica.")
                        elif section_info["key"] == "conclusioni":
                            doc.add_paragraph("Conclusioni e interpretazione complessiva del sito archeologico.")

                        # If this section has subsections, process them
                        if "subsections" in section_info:
                            for subsection in section_info["subsections"]:
                                subsection_key = subsection["key"]
                                subsection_heading = subsection["heading"]
                                subsection_content = report_data.get(subsection_key, "")

                                # Add subsection heading
                                doc.add_heading(subsection_heading, level=2)

                                # Add default content if subsection is empty
                                if not subsection_content or subsection_content.strip() == "":
                                    if subsection_key == "descrizione_metodologica":
                                        doc.add_paragraph("Metodologia di scavo utilizzata durante la campagna archeologica.")
                                    elif subsection_key == "analisi_stratigrafica":
                                        doc.add_paragraph("Analisi stratigrafica e interpretazione delle unità stratigrafiche.")
                                    elif subsection_key == "descrizione_materiali":
                                        doc.add_paragraph("Catalogo e descrizione dei materiali rinvenuti durante lo scavo.")
                                else:
                                    # Process subsection content line by line
                                    subsection_lines = subsection_content.split('\n')
                                    j = 0
                                    while j < len(subsection_lines):
                                        line = subsection_lines[j].strip()

                                        # Check for markdown headings (##, --, ==)
                                        if (line.startswith('#') or 
                                            (j+1 < len(subsection_lines) and (all(c == '=' for c in subsection_lines[j+1].strip()) or 
                                                                  all(c == '-' for c in subsection_lines[j+1].strip())))):

                                            heading_text = clean_heading(line)
                                            heading_level = 3  # Default to level 3 for subheadings in subsections

                                            # Determine heading level based on # count or underline character
                                            if line.startswith('#'):
                                                heading_level = min(6, line.count('#') + 2)  # +2 because we're in a subsection
                                            elif j+1 < len(subsection_lines):
                                                if all(c == '=' for c in subsection_lines[j+1].strip()):
                                                    heading_level = 3  # = underline is level 3 in subsection
                                                elif all(c == '-' for c in subsection_lines[j+1].strip()):
                                                    heading_level = 4  # - underline is level 4 in subsection

                                            doc.add_heading(heading_text, level=heading_level)

                                            # Skip the underline if present
                                            if (j+1 < len(subsection_lines) and (all(c == '=' for c in subsection_lines[j+1].strip()) or 
                                                                       all(c == '-' for c in subsection_lines[j+1].strip()))):
                                                j += 2
                                            else:
                                                j += 1
                                            continue

                                        # Check for images
                                        img_match = re.search(r'\[IMMAGINE\s+(?:US|RAPPORTI|FASE|REPERTO|CERAMICA)\s+[^:]*:\s+(.*?),\s*(.*?)\]', line)
                                        if img_match:
                                            img_path = img_match.group(1).strip()
                                            caption = img_match.group(2).strip()

                                            # Remove file:// prefix if present
                                            if img_path.startswith('file://'):
                                                img_path = img_path[7:]

                                            print(f"Processing image: path={img_path}, caption={caption}")

                                            if os.path.exists(img_path):
                                                try:
                                                    doc.add_picture(img_path, width=Inches(6))
                                                    caption_para = doc.add_paragraph(caption)
                                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                    caption_para.style = 'Caption'
                                                except Exception as e:
                                                    print(f"Error adding image {img_path}: {str(e)}")
                                                    doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                                            else:
                                                print(f"Image file not found: {img_path}")
                                                doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                                            j += 1
                                            continue

                                        # Check for markdown tables
                                        if line.startswith('|') and line.endswith('|'):
                                            # Collect table lines
                                            table_lines = [line]
                                            k = j + 1
                                            while k < len(subsection_lines) and subsection_lines[k].strip().startswith('|') and subsection_lines[k].strip().endswith('|'):
                                                table_lines.append(subsection_lines[k].strip())
                                                k += 1

                                            # Process table if we have at least one row
                                            if len(table_lines) >= 1:
                                                # Use the convert_markdown_table function to create the table
                                                convert_markdown_table(table_lines, doc)

                                            j = k  # Skip processed table lines
                                            continue

                                        # Regular paragraph
                                        doc.add_paragraph(line)
                                        j += 1

                        continue

                    # Add section heading
                    doc.add_heading(section_info["heading"], level=1)

                    # Check if this section has subsections
                    if "subsections" in section_info:
                        # Process each subsection
                        for subsection in section_info["subsections"]:
                            subsection_key = subsection["key"]
                            subsection_heading = subsection["heading"]
                            subsection_content = report_data.get(subsection_key, "")

                            # Skip empty subsections
                            if not subsection_content or subsection_content.strip() == "":
                                print(f"WARNING: Empty content for subsection '{subsection_key}', adding placeholder text")
                                doc.add_heading(subsection_heading, level=2)
                                doc.add_paragraph("Questa sottosezione è stata omessa a causa della mancanza di dati necessari.")
                                continue

                            # Add subsection heading
                            doc.add_heading(subsection_heading, level=2)

                            # Process subsection content line by line
                            subsection_lines = subsection_content.split('\n')
                            j = 0
                            while j < len(subsection_lines):
                                line = subsection_lines[j].strip()

                                # Check for markdown headings (##, --, ==)
                                if (line.startswith('#') or 
                                    (j+1 < len(subsection_lines) and (all(c == '=' for c in subsection_lines[j+1].strip()) or 
                                                          all(c == '-' for c in subsection_lines[j+1].strip())))):

                                    heading_text = clean_heading(line)
                                    heading_level = 3  # Default to level 3 for subheadings in subsections

                                    # Determine heading level based on # count or underline character
                                    if line.startswith('#'):
                                        heading_level = min(6, line.count('#') + 2)  # +2 because we're in a subsection
                                    elif j+1 < len(subsection_lines):
                                        if all(c == '=' for c in subsection_lines[j+1].strip()):
                                            heading_level = 3  # = underline is level 3 in subsection
                                        elif all(c == '-' for c in subsection_lines[j+1].strip()):
                                            heading_level = 4  # - underline is level 4 in subsection

                                    doc.add_heading(heading_text, level=heading_level)

                                    # Skip the underline if present
                                    if (j+1 < len(subsection_lines) and (all(c == '=' for c in subsection_lines[j+1].strip()) or 
                                                                 all(c == '-' for c in subsection_lines[j+1].strip()))):
                                        j += 2
                                    else:
                                        j += 1
                                    continue

                                # Check for images (US, RAPPORTI, FASE, MATERIALE, CERAMICA)
                                img_match = re.search(r'\[IMMAGINE\s+(?:US|RAPPORTI|FASE|REPERTO|CERAMICA)\s+[^:]*:\s+(.*?),\s*(.*?)\]', line)
                                if img_match:
                                    img_path = img_match.group(1).strip()
                                    caption = img_match.group(2).strip()

                                    # Remove file:// prefix if present
                                    if img_path.startswith('file://'):
                                        img_path = img_path[7:]

                                    print(f"Processing image: path={img_path}, caption={caption}")

                                    if os.path.exists(img_path):
                                        try:
                                            doc.add_picture(img_path, width=Inches(6))
                                            caption_para = doc.add_paragraph(caption)
                                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            caption_para.style = 'Caption'
                                        except Exception as e:
                                            print(f"Error adding image {img_path}: {str(e)}")
                                            doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                                    else:
                                        print(f"Image file not found: {img_path}")
                                        doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                                    j += 1
                                    continue

                                # Check for markdown tables
                                if line.startswith('|') and line.endswith('|'):
                                    # Collect table lines
                                    table_lines = [line]
                                    k = j + 1
                                    while k < len(subsection_lines) and subsection_lines[k].strip().startswith('|') and subsection_lines[k].strip().endswith('|'):
                                        table_lines.append(subsection_lines[k].strip())
                                        k += 1

                                    # Process table if we have at least one row
                                    if len(table_lines) >= 1:
                                        # Use the convert_markdown_table function to create the table
                                        convert_markdown_table(table_lines, doc)

                                        # Add a paragraph after the table for spacing
                                        doc.add_paragraph()

                                    j = k  # Skip processed table lines
                                    continue

                                # Add regular paragraph
                                if line:
                                    doc.add_paragraph(line)
                                j += 1

                        # Skip processing the main section content since we've processed the subsections
                        continue

                    # Process content line by line for sections without subsections
                    lines = content.split('\n')
                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()

                        # Check for markdown headings (##, --, ==)
                        if (line.startswith('#') or 
                            (i+1 < len(lines) and (all(c == '=' for c in lines[i+1].strip()) or 
                                                  all(c == '-' for c in lines[i+1].strip())))):

                            heading_text = clean_heading(line)
                            heading_level = 2  # Default to level 2 for subheadings

                            # Determine heading level based on # count or underline character
                            if line.startswith('#'):
                                heading_level = min(6, line.count('#') + 1)  # +1 because # is level 1
                            elif i+1 < len(lines):
                                if all(c == '=' for c in lines[i+1].strip()):
                                    heading_level = 2  # = underline is level 2
                                elif all(c == '-' for c in lines[i+1].strip()):
                                    heading_level = 3  # - underline is level 3

                            doc.add_heading(heading_text, level=heading_level)

                            # Skip the underline if present
                            if (i+1 < len(lines) and (all(c == '=' for c in lines[i+1].strip()) or 
                                                     all(c == '-' for c in lines[i+1].strip()))):
                                i += 2
                            else:
                                i += 1
                            continue

                        # Check for images (US, RAPPORTI, FASE, MATERIALE, CERAMICA)
                        img_match = re.search(r'\[IMMAGINE\s+(?:US|RAPPORTI|FASE|REPERTO|CERAMICA)\s+[^:]*:\s+(.*?),\s*(.*?)\]', line)
                        if img_match:
                            img_path = img_match.group(1).strip()
                            caption = img_match.group(2).strip()

                            # Remove file:// prefix if present
                            if img_path.startswith('file://'):
                                img_path = img_path[7:]

                            print(f"Processing image: path={img_path}, caption={caption}")

                            if os.path.exists(img_path):
                                try:
                                    doc.add_picture(img_path, width=Inches(6))
                                    caption_para = doc.add_paragraph(caption)
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.style = 'Caption'
                                except Exception as e:
                                    print(f"Error adding image {img_path}: {str(e)}")
                                    doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                            else:
                                print(f"Image file not found: {img_path}")
                                doc.add_paragraph(f"[Immagine non disponibile: {caption}]")
                            i += 1
                            continue

                        # Check for markdown tables
                        if line.startswith('|') and line.endswith('|'):
                            # Collect table lines
                            table_lines = [line]
                            j = i + 1
                            while j < len(lines) and lines[j].strip().startswith('|') and lines[j].strip().endswith('|'):
                                table_lines.append(lines[j].strip())
                                j += 1

                            # Process table if we have at least one row
                            if len(table_lines) >= 1:
                                # Use the convert_markdown_table function to create the table
                                convert_markdown_table(table_lines, doc)

                                # Add a paragraph after the table for spacing
                                doc.add_paragraph()

                            i = j  # Skip processed table lines
                            continue

                        # Regular paragraph
                        doc.add_paragraph(line)
                        i += 1

        # Check if any placeholders were found
        if not placeholder_found:
            print("WARNING: No section placeholders found in the template. Adding sections manually.")
            print("Expected placeholders: " + ", ".join(section_mapping.keys()))

            # Add all sections manually if no placeholders were found
            for placeholder, section_info in section_mapping.items():
                content = report_data.get(section_info["key"], "")
                print(f"Adding section '{section_info['key']}' manually")

                # Add section heading
                doc.add_heading(section_info["heading"], level=1)

                # If section has content, add it
                if content and content.strip() != "":
                    # Add content as a paragraph
                    doc.add_paragraph(content)
                else:
                    # Add default content
                    if section_info["key"] == "introduzione":
                        doc.add_paragraph("Introduzione al sito archeologico e alla campagna di scavo.")
                    elif section_info["key"] == "descrizione_metodologica_ed_esito":
                        doc.add_paragraph("Descrizione metodologica ed esito dell'indagine archeologica.")
                    elif section_info["key"] == "conclusioni":
                        doc.add_paragraph("Conclusioni e interpretazione complessiva del sito archeologico.")

                # If this section has subsections, add them
                if "subsections" in section_info:
                    for subsection in section_info["subsections"]:
                        subsection_key = subsection["key"]
                        subsection_heading = subsection["heading"]
                        subsection_content = report_data.get(subsection_key, "")

                        # Add subsection heading
                        doc.add_heading(subsection_heading, level=2)

                        # If subsection has content, add it
                        if subsection_content and subsection_content.strip() != "":
                            # Add content as a paragraph
                            doc.add_paragraph(subsection_content)
                        else:
                            # Add default content
                            if subsection_key == "descrizione_metodologica":
                                doc.add_paragraph("Metodologia di scavo utilizzata durante la campagna archeologica.")
                            elif subsection_key == "analisi_stratigrafica":
                                doc.add_paragraph("Analisi stratigrafica e interpretazione delle unità stratigrafiche.")
                            elif subsection_key == "descrizione_materiali":
                                doc.add_paragraph("Catalogo e descrizione dei materiali rinvenuti durante lo scavo.")

        # Add tables for materials, pottery, tombs, periodization, and structures
        # Materials table
        if report_data.get('materials_table'):
            print(f"Adding materials table with {len(report_data['materials_table'])} rows")
            doc.add_heading("CATALOGO DEI MATERIALI", level=2)
            table = doc.add_table(rows=1, cols=len(report_data['materials_table'][0]))
            table.style = table_style_name

            # Headers
            for i, header in enumerate(report_data['materials_table'][0]):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_data in report_data['materials_table'][1:]:
                row = table.add_row().cells
                for i, val in enumerate(row_data):
                    row[i].text = str(val)

        # Pottery table
        if report_data.get('pottery_table'):
            doc.add_heading("CATALOGO DELLA CERAMICA", level=2)
            table = doc.add_table(rows=1, cols=len(report_data['pottery_table'][0]))
            table.style = table_style_name

            # Headers
            for i, header in enumerate(report_data['pottery_table'][0]):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_data in report_data['pottery_table'][1:]:
                row = table.add_row().cells
                for i, val in enumerate(row_data):
                    row[i].text = str(val)

        # Tombs table
        if report_data.get('tomba_table'):
            doc.add_heading("CATALOGO DELLE TOMBE", level=2)
            table = doc.add_table(rows=1, cols=len(report_data['tomba_table'][0]))
            table.style = table_style_name

            # Headers
            for i, header in enumerate(report_data['tomba_table'][0]):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_data in report_data['tomba_table'][1:]:
                row = table.add_row().cells
                for i, val in enumerate(row_data):
                    row[i].text = str(val)

        # Periodization table
        if report_data.get('periodizzazione_table'):
            doc.add_heading("PERIODIZZAZIONE", level=2)
            table = doc.add_table(rows=1, cols=len(report_data['periodizzazione_table'][0]))
            table.style = table_style_name

            # Headers
            for i, header in enumerate(report_data['periodizzazione_table'][0]):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_data in report_data['periodizzazione_table'][1:]:
                row = table.add_row().cells
                for i, val in enumerate(row_data):
                    row[i].text = str(val)

        # Structures table
        if report_data.get('struttura_table'):
            doc.add_heading("CATALOGO DELLE STRUTTURE", level=2)
            table = doc.add_table(rows=1, cols=len(report_data['struttura_table'][0]))
            table.style = table_style_name

            # Headers
            for i, header in enumerate(report_data['struttura_table'][0]):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Data
            for row_data in report_data['struttura_table'][1:]:
                row = table.add_row().cells
                for i, val in enumerate(row_data):
                    row[i].text = str(val)

        # Final check to ensure document has content
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        print(f"Final document has {para_count} paragraphs and {table_count} tables")

        if para_count <= 1 and table_count == 0:
            print("WARNING: Document appears to be empty or nearly empty!")
            # Add a warning paragraph to the document
            doc.add_paragraph("ATTENZIONE: Il documento sembra essere vuoto. Verifica che il report contenga dati e che il template sia corretto.")

        # Save the document
        try:
            print(f"Saving document to {output_path}")
            doc.save(output_path)
            print("Document saved successfully")
        except Exception as e:
            print(f"ERROR saving document: {str(e)}")
            raise

        # Add current date if not already set
        if "{{DATA}}" in str(doc.paragraphs):
            for para in doc.paragraphs:
                if "{{DATA}}" in para.text:
                    current_date = datetime.now().strftime("%d/%m/%Y")
                    para.text = para.text.replace("{{DATA}}", current_date)

        doc.save(output_path)

    def _format_pottery_table(self):
        """Formatta i dati della ceramica in una struttura tabellare"""
        if not self.pottery_data:
            return []

        headers = ['ID','ID Number', 'Forma', 'Impasto', 'Tipo', 'Conservazione %', 'Note']
        rows = [[
            pottery.get('id_rep', ''),
            pottery.get('id_number', ''),
            pottery.get('form', ''),
            pottery.get('fabric', ''),
            pottery.get('ware', ''),
            pottery.get('percent', ''),
            pottery.get('note', '')
        ] for pottery in self.pottery_data]

        return [headers] + rows

    def _format_materials_table(self):
        """Formatta i dati dei materiali in una struttura tabellare"""
        if not self.materials_data:
            return []

        headers = ['ID','Inv. Number', 'Tipo', 'Definizione', 'Descrizione', 'Datazione', 'Conservazione']
        rows = [[
            material.get('inv_mat', ''),
            material.get('numero_inventario', ''),
            material.get('tipo_reperto', ''),
            material.get('definizione', ''),
            material.get('descrizione', ''),
            material.get('datazione', ''),
            material.get('stato_conservazione', '')
        ] for material in self.materials_data]

        return [headers] + rows

    def _add_table_to_doc(self, doc, title, table_data):
        """Aggiunge una tabella formattata al documento"""
        if not table_data:
            return

        print(f"Adding table to document: {title} with {len(table_data)} rows")
        doc.add_heading(title, level=2)

        # Create table
        table = doc.add_table(rows=1, cols=len(table_data[0]))

        # Find the best table style
        table_styles = [s.name for s in doc.styles if hasattr(s, 'name') and 'table' in s.name.lower()]
        print(f"Available table styles: {', '.join(table_styles)}")

        # Try to find the best table style
        if 'Table Grid' in table_styles:
            table_style_name = 'Table Grid'
        elif 'Grid Table' in table_styles:
            table_style_name = 'Grid Table'
        elif any('grid' in s.lower() for s in table_styles):
            # Find the first style with 'grid' in the name
            table_style_name = next(s for s in table_styles if 'grid' in s.lower())
        elif 'Normal Table' in table_styles:
            table_style_name = 'Normal Table'
        elif len(table_styles) > 0:
            # Use the first available table style
            table_style_name = table_styles[0]
        else:
            table_style_name = 'Table'  # Fallback to most basic table style

        print(f"Using table style: {table_style_name}")

        try:
            table.style = table_style_name
        except Exception as e:
            print(f"Error setting table style: {str(e)}")
            # Try a different approach if setting style fails
            try:
                # Apply a basic grid to the table
                for cell in table._tbl.findall(".//w:tc", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                    tcPr = cell.find(".//w:tcPr", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if tcPr is None:
                        tcPr = parse_xml(r'<w:tcPr {0}/>'.format(nsdecls('w')))
                        cell.append(tcPr)

                    # Add borders to the cell
                    tcBorders = tcPr.find(".//w:tcBorders", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if tcBorders is None:
                        tcBorders = parse_xml(r'<w:tcBorders {0}/>'.format(nsdecls('w')))
                        tcPr.append(tcBorders)

                    # Define borders for all sides
                    for side in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(".//{0}:{1}".format('w', side))
                        if border is None:
                            border = parse_xml(r'<w:{0} {1} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'.format(side, nsdecls('w')))
                            tcBorders.append(border)
            except Exception as border_error:
                print(f"Error applying borders: {str(border_error)}")

        # Headers
        for i, header in enumerate(table_data[0]):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Make header bold and centered
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for row_data in table_data[1:]:
            row = table.add_row()
            for i, val in enumerate(row_data):
                if i < len(row.cells):  # Ensure we don't exceed column count
                    row.cells[i].text = str(val)

        # Add a paragraph after the table for spacing
        doc.add_paragraph()

    def sketchgpt(self):
        items = self.iconListWidget.selectedItems()
        conn = Connection()
        thumb_resize = conn.thumb_resize()
        thumb_resize_str = thumb_resize['thumb_resize']

        def process_file_path(file_path):
            return urllib.parse.unquote(file_path)

        def query_media(search_dict, table="MEDIA_THUMB"):
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            try:
                return self.DB_MANAGER.query_bool(search_dict, table)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Database query failed: {str(e)}", QMessageBox.Ok)
                return None

        selected_images = []
        for item in items:
            id_orig_item = item.text()
            search_dict = {'media_filename': f"'{id_orig_item}'"}
            res = query_media(search_dict)

            if res:
                file_path = process_file_path(os.path.join(thumb_resize_str, str(res[0].path_resize)))
                media_type = res[0].mediatype

                if media_type == 'image':
                    selected_images.append(file_path)
                elif media_type == '3d_model':
                    # Gestisci i modelli 3D se necessario
                    selected_images.append(file_path)
                elif media_type == 'video':
                    # Gestisci i video se necessario
                    selected_images.append(file_path)
            else:
                QMessageBox.warning(self, "Error", f"File not found: {id_orig_item}", QMessageBox.Ok)

        #if selected_images:
        self.gpt_window = GPTWindow(selected_images, dbmanager=self.DB_MANAGER, main_class=self)
        self.gpt_window.show()
        #else:
            #QMessageBox.warning(self, "Warning", "No valid images selected for analysis.", QMessageBox.Ok)

    def on_pushButton_trick_pressed(self):
        # Crea un oggetto QDialog
        dialog = QDialog()

        # Imposta alcune proprietà per la finestra di dialogo
        dialog.setWindowTitle("Scorciatoie da tastiera")
        dialog.setFixedWidth(400)

        # Crea un oggetto QLabel con il testo degli shortcut
        text = """
            Ctrl+Shift+X : Attiva text2sql
            Ctrl+U : Aggiorna Tablewidget rapporti (aggiunge e aggiorna area e sito)
            Ctrl+Shift+D : Elimina tutti i record filtrati
            Ctrl+Shift+N : Cambia modalità di ricerca
            """
        label = QLabel(text)

        # Imposta l'oggetto QLabel come layout della finestra di dialogo
        dialog.setLayout(QVBoxLayout())
        dialog.layout().addWidget(label)

        # Mostra la finestra di dialogo
        dialog.exec_()

    def get_input_prompt(self, label):
        if self.L == 'it':
            return QInputDialog.getText(self, "Input", f"Inserire il {label}")
        else:
            return QInputDialog.getText(self, "Input", f"Insert the {label}")

    def show_warning(self, message):
        if self.L == 'it':
            QMessageBox.warning(self, "Input", f"Sito o {message} non forniti.")
            return
        QMessageBox.warning(self, "Input", f"Site or {message} not provided.")

    def show_error(self, error, original_message):
        if self.L == 'it':
            QMessageBox.warning(self, "Error", f"Si è verificato un errore durante {original_message}: {error}",
                                QMessageBox.Ok)
            return
        QMessageBox.warning(self, "Error", f"An error occurred during {original_message}: {error}", QMessageBox.Ok)

    def update_all_areas(self):
        conn = Connection()
        sito_set = conn.sito_set()
        sito_set_str = sito_set['sito_set']


        all_areas = self.get_all_areas(sito_set_str)

        dialog = ProgressDialog()


        for i, area in enumerate(all_areas):
            self.update_rapporti_col(sito_set_str, area)
            dialog.setValue(i + 1)

        self.update_rapporti_col_2()
        dialog.setValue(len(all_areas) + 1)
    # def get_all_areas(self):
    #     conn = Connection()
    #     conn_str = conn.conn_str()
    #     metadata = MetaData()
    #     engine = create_engine(conn_str)
    #
    #     # Assuming areas are represented in a table named 'area_table'
    #     area_table = Table('us_table', metadata, autoload_with=engine)
    #
    #     with engine.connect() as connection:
    #         # Assuming the name of the area is saved in a column named 'area_name'
    #         stmt = select([area_table.c.area])
    #         result = connection.execute(stmt)
    #
    #         # Fetch all rows from the result and return only the area names
    #         all_areas = [row['area'] for row in result]
    #
    #     return all_areas
    def get_all_areas(self, sito=None):
        conn = Connection()
        conn_str = conn.conn_str()
        metadata = MetaData()
        engine = create_engine(conn_str)

        # Assuming areas are represented in a table named 'area_table'
        area_table = Table('us_table', metadata, autoload_with=engine)

        with engine.connect() as connection:
            # Base query per selezionare le aree
            stmt = select([area_table.c.area]).distinct()

            # Se è specificato un sito, filtra per quel sito
            if sito:
                stmt = stmt.where(area_table.c.sito == sito)

            result = connection.execute(stmt)
            areas = result.fetchall()


            # Fetch all rows from the result and return only the area names
            all_areas = [row['area'] for row in result]

        return all_areas

    def update_rapporti_col(self, sito, area):
        conn = Connection()
        conn_str = conn.conn_str()
        metadata = MetaData()
        engine = create_engine(conn_str)

        us_table = Table('us_table', metadata, autoload_with=engine)

        if not sito or not area:
            self.show_warning("Sito o area non specificato.")
            return  # Exit the function if site or area is not provided



        def log_error(message, error_type="ERROR", filename = self.HOME+"/error_log.txt"):
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with open(filename, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {error_type}: {message}\n")

        try:
            with engine.connect() as connection:
                # Log dell'inizio dell'operazione
                log_error("Inizio operazione di aggiornamento", "INFO")

                stmt = select([us_table.c.id_us, us_table.c.rapporti]).where(
                    us_table.c.sito == sito, us_table.c.area == area
                )

                try:
                    rows = connection.execute(stmt).fetchall()
                    log_error(f"Recuperate {len(rows)} righe da processare", "INFO")
                except Exception as e:
                    log_error(f"Errore nel recupero delle righe: {str(e)}")
                    raise

                for row in rows:
                    id, rapporti_str = row
                    log_error(f"Processing ID: {id}", "DEBUG")

                    if rapporti_str and rapporti_str != "[[]]":
                        try:
                            # Log della stringa originale
                            log_error(f"ID {id} - Stringa originale: {rapporti_str}", "DEBUG")

                            # Conversione della stringa in lista
                            rapporti_list = ast.literal_eval(rapporti_str)

                            # Aggiornamento della lista
                            updated_rapporti_list = [sublist + [area, sito] if sublist else sublist
                                                     for sublist in rapporti_list]
                            updated_rapporti_list2 = [sub[:4] for sub in updated_rapporti_list]

                            # Conversione in stringa con controllo UTF-8
                            updated_rapporti_str = self.ensure_utf8(str(updated_rapporti_list2))

                            # Log della stringa aggiornata
                            log_error(f"ID {id} - Stringa aggiornata: {updated_rapporti_str}", "DEBUG")

                            # Preparazione e esecuzione dell'update
                            update_stmt = (
                                update(us_table)
                                .where(us_table.c.id_us == id)
                                .values(rapporti=updated_rapporti_str)
                            )

                            try:
                                connection.execute(update_stmt)
                            except Exception as e:
                                error_msg = (
                                    f"Errore nell'aggiornamento dell'ID {id}:\n"
                                    f"Errore: {str(e)}\n"
                                    f"Stringa problematica: {updated_rapporti_str}"
                                )
                                log_error(error_msg)
                                continue

                        except ValueError as e:
                            error_msg = (
                                f"Errore nella conversione per ID {id}:\n"
                                f"Errore: {str(e)}\n"
                                f"Stringa originale: {rapporti_str}"
                            )
                            log_error(error_msg)
                            continue
                        except Exception as e:
                            error_msg = (
                                f"Errore generico per ID {id}:\n"
                                f"Errore: {str(e)}\n"
                                f"Stringa originale: {rapporti_str}"
                            )
                            log_error(error_msg)
                            continue

                log_error("Operazione completata", "INFO")

        except Exception as e:
            error_msg = f"Errore critico durante l'operazione: {str(e)}"
            log_error(error_msg)
            raise

        except Exception as e:
            self.show_error(e, "l'aggiornamento")

    def update_rapporti_col_2(self):
        # Inizializza la connessione
        conn = Connection()
        conn_str = conn.conn_str()
        metadata = MetaData()
        engine = create_engine(conn_str)
        sito_set = conn.sito_set()
        sito_set_str = sito_set['sito_set']

        def log_error(message, error_type="ERROR", filename=self.HOME+"/rapporti_update_log.txt"):
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with open(filename, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {error_type}: {message}\n")
        try:
            us_table = Table('us_table', metadata, autoload_with=engine)
        except Exception as e:
            error_msg = f"Errore nel caricamento della tabella: {str(e)}"
            log_error(error_msg)
            self.show_error(e, "il caricamento della tabella")
            return

        var1 = sito_set_str  # Sito

        if not var1:
            log_error("Tentativo di aggiornamento senza specificare il sito", "WARNING")
            self.show_warning("sito non specificato")
            return

        log_error(f"Inizio aggiornamento rapporti per il sito: {var1}", "INFO")

        try:
            with engine.connect() as connection:
                # Recupera tutte le righe per il sito specificato
                stmt = select([us_table]).where(us_table.c.sito == var1)
                try:
                    rows = connection.execute(stmt).fetchall()
                    log_error(f"Recuperate {len(rows)} righe da processare per il sito {var1}", "INFO")
                except Exception as e:
                    error_msg = f"Errore nel recupero delle righe per il sito {var1}: {str(e)}"
                    log_error(error_msg)
                    self.show_error(e, "il recupero dei dati")
                    return

                # Processa ogni riga
                for row_j in rows:
                    id_us, rapporti_str = row_j.id_us, row_j.rapporti

                    log_error(f"Processing US ID: {id_us}", "DEBUG")

                    if rapporti_str and rapporti_str != "[[]]":
                        try:
                            # Log della stringa originale
                            log_error(f"US {id_us} - Stringa originale: {rapporti_str}", "DEBUG")

                            rapporti_list = ast.literal_eval(rapporti_str)
                            updated_rapporti_list = []

                            for sublist in rapporti_list:
                                try:
                                    # Verifica che la sottolista abbia abbastanza elementi
                                    if len(sublist) < 3:
                                        error_msg = f"Sottolista troppo corta per US {id_us}: {sublist}"
                                        log_error(error_msg, "WARNING")
                                        continue

                                    us_id = sublist[1]
                                    current_area = sublist[2]

                                    # Trova l'area corretta
                                    try:
                                        correct_area = self.find_correct_area_for_us(us_id, var1, connection)

                                        if correct_area is None:
                                            log_error(f"Area non trovata per US {us_id} nel sito {var1}", "WARNING")
                                            updated_rapporti_list.append(sublist)
                                            continue

                                        if correct_area == current_area:
                                            updated_rapporti_list.append(sublist)
                                        else:
                                            updated_sublist = sublist.copy()
                                            updated_sublist[2] = correct_area
                                            updated_rapporti_list.append(updated_sublist)
                                            log_error(
                                                f"Aggiornata area per US {us_id}: da {current_area} a {correct_area}",
                                                "INFO")

                                    except Exception as e:
                                        error_msg = f"Errore nel trovare l'area corretta per US {us_id}: {str(e)}"
                                        log_error(error_msg)
                                        updated_rapporti_list.append(sublist)
                                        continue

                                except Exception as e:
                                    error_msg = f"Errore nel processare la sottolista per US {id_us}: {str(e)}"
                                    log_error(error_msg)
                                    continue

                            # Aggiorna il database
                            updated_rapporti_list2 = [sub[:4] for sub in updated_rapporti_list]
                            updated_rapporti_str = self.ensure_utf8(str(updated_rapporti_list2))

                            log_error(f"US {id_us} - Stringa aggiornata: {updated_rapporti_str}", "DEBUG")

                            try:
                                update_stmt = update(us_table).where(us_table.c.id_us == id_us).values(
                                    rapporti=updated_rapporti_str)
                                connection.execute(update_stmt)
                            except Exception as e:
                                error_msg = (
                                    f"Errore nell'aggiornamento dell'US {id_us}:\n"
                                    f"Errore: {str(e)}\n"
                                    f"Stringa problematica: {updated_rapporti_str}"
                                )
                                log_error(error_msg)
                                continue

                        except ValueError as e:
                            error_msg = (
                                f"Errore nella conversione per US {id_us}:\n"
                                f"Errore: {str(e)}\n"
                                f"Stringa originale: {rapporti_str}"
                            )
                            log_error(error_msg)
                            self.show_error(e, "la conversione della stringa in una lista di liste")
                            continue

                        except Exception as e:
                            error_msg = f"Errore generico per US {id_us}: {str(e)}"
                            log_error(error_msg)
                            continue

                log_error(f"Aggiornamento completato per il sito {var1}", "INFO")
                self.view_all()

        except Exception as e:
            error_msg = f"Errore critico durante l'aggiornamento: {str(e)}"
            log_error(error_msg)
            self.show_error(e, "l'aggiornamento")

    def ensure_utf8(self,s):
        """
        Assicura che una stringa sia codificata correttamente in UTF-8.

        Args:
            s (str): La stringa da codificare.

        Returns:
            str: La stringa codificata in UTF-8.
        """
        if isinstance(s, str):
            try:
                # Prova a codificare e decodificare per verificare la validità UTF-8
                return s.encode('utf-8', errors='replace').decode('utf-8')
            except UnicodeError:
                # Se c'è un errore, sostituisci i caratteri non validi
                return s.encode('ascii', errors='replace').decode('ascii')
        return str(s) if s is not None else ""

    def find_correct_area_for_us(self, us, sito, connection):
        """
        Trova l'area corretta per una data unità stratigrafica (us) e sito.

        Args:
            us (str): L'identificativo dell'unità stratigrafica.
            sito (str): Il nome del sito.
            connection: Connessione al database esistente.

        Returns:
            str: L'area corretta o None se non trovata.
        """


        def log_error(message, error_type="ERROR", filename=self.HOME+"/error_log_fetch.txt"):
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            with open(filename, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {error_type}: {message}\n")
        try:
            # Sanitizza gli input
            us = str(us).strip() if us is not None else ""
            sito = self.ensure_utf8(str(sito).strip()) if sito is not None else ""

            if not us or not sito:
                log_error(f"Input non valido: US={us}, Sito={sito}", "WARNING")
                return None

            log_error(f"Cercando area per US={us}, Sito={sito}", "DEBUG")

            # Usa la connessione esistente invece di crearne una nuova
            metadata = MetaData()
            us_table = Table('us_table', metadata, autoload_with=connection.engine)

            # Prepara la query con parametri sanitizzati
            stmt = select([us_table.c.area]).where(
                and_(
                    us_table.c.us == us,
                    us_table.c.sito == sito
                )
            )

            try:
                result = connection.execute(stmt).fetchone()

                if result and result[0]:
                    area = self.ensure_utf8(str(result[0]))
                    log_error(f"Area trovata: {area} per US={us}, Sito={sito}", "DEBUG")
                    return area
                else:
                    log_error(f"Nessuna area trovata per US={us}, Sito={sito}", "WARNING")
                    return None

            except UnicodeDecodeError as ude:
                error_msg = (
                    f"Errore di codifica nel recupero dell'area:\n"
                    f"US={us}, Sito={sito}\n"
                    f"Posizione errore: {ude.start}\n"
                    f"Oggetto: {repr(ude.object[max(0, ude.start - 10):min(len(ude.object), ude.end + 10)])}"
                )
                log_error(error_msg, "ERROR")
                return None

            except Exception as e:
                error_msg = f"Errore nell'esecuzione della query per US={us}, Sito={sito}: {str(e)}"
                log_error(error_msg, "ERROR")
                return None

        except Exception as e:
            error_msg = f"Errore critico in find_correct_area_for_us: {str(e)}"
            log_error(error_msg, "ERROR")
            return None



    def clean_comments(self,text_to_clean):
        clean_text = text_to_clean.split("##")[0].replace("\n", "")
        return clean_text

    def EM_extract_node_name(self, node_element):
        is_d4 = False
        is_d5 = False
        node_y_pos = None
        nodeshape = None
        nodeurl = None
        nodedescription = None
        nodename = ''
        fillcolor = None
        noderel2 = ''
        attrib_ = None
        for subnode in node_element.findall('.//{http://graphml.graphdrawing.org/xmlns}data'):
            attrib = subnode.attrib
            if attrib == {'{http://www.w3.org/XML/1998/namespace}space': 'preserve', 'key': 'd4'}:
                is_d4 = True
                nodeurl = subnode.text
            if attrib == {'{http://www.w3.org/XML/1998/namespace}space': 'preserve', 'key': 'd5'}:
                is_d5 = True
                nodedescription = self.clean_comments(subnode.text)
            if attrib == {'key': 'd6'}:
                for USname in subnode.findall('.//{http://www.yworks.com/xml/graphml}NodeLabel'):
                    nodename = self.check_if_empty(USname.text)
                for fill_color in subnode.findall('.//{http://www.yworks.com/xml/graphml}Fill'):
                    fillcolor = fill_color.attrib['color']
                for USshape in subnode.findall('.//{http://www.yworks.com/xml/graphml}Shape'):
                    nodeshape = USshape.attrib['type']
                for geometry in subnode.findall(
                        './{http://www.yworks.com/xml/graphml}ShapeNode/{http://www.yworks.com/xml/graphml}Geometry'):
                    # for geometry in subnode.findall('./{http://www.yworks.com/xml/graphml}Geometry'):
                    node_y_pos = geometry.attrib['y']

        if not is_d4:
            nodeurl = '--None--'
        if not is_d5:
            nodedescription = '--None--'

        return nodename, nodedescription, nodeurl, nodeshape, node_y_pos, fillcolor


    def check_if_empty(self,name):
        if name == None:
            name = "--None--"
        return name

    ###Apri file graphml
    def on_pushButton_graphml2csv_pressed(self):
        s = QgsSettings()
        sqlite_DB_path = '{}{}{}'.format(self.HOME, os.sep,
                                         "pyarchinit_DB_folder")
        dbpath = QFileDialog.getOpenFileName(
            self,
            "Set file name",
            '',
            " graphml (*.graphml)"
        )[0]
        filval = dbpath  # .split("/")[-1]
        xmltree = ET()

        xmltree.parse(filval)
        konton = xmltree

        # crea CSV FILE
        csvfile = open(sqlite_DB_path+'/graphml2csv.csv', 'w', encoding='utf-8')
        csvfile_writer = csv.writer(csvfile)

        # aggiungi intestazione al csv CSV FILE
        csvfile_writer.writerow(["site", "area", "us", "unit_type", "i_stratigrafica"])


        ###funzione per stampare il valore della stringa

        for i in konton.iter():
            # print(EM_extract_node_name(i)[0])
            m = re.match(r"(?P<l>[a-zA-Z]+)(?P<n>.+)$", str(self.EM_extract_node_name(i)[0]))
            if m is not None:

                ### AGGIUNGO AL CSV

                csv_line = ['sito', '1', m.group('n'), m.group('l'), str(self.EM_extract_node_name(i)[1])]

                if csv_line[2].count('.') == 1:
                    csv_line[3] = csv_line[3].replace('D', 'DOC')
                if csv_line[2].count('.') > 1:
                    csv_line[3] = csv_line[3].replace('D', 'Extractor')
                if csv_line[3].startswith('C'):
                    csv_line[3] = csv_line[3].replace('C', 'Combinar')
                elif not csv_line[3].startswith('C') and not csv_line[3].startswith('D') and not csv_line[3].startswith(
                        'E') and not csv_line[3].startswith('US'):
                    csv_line[3] = 'property'
                elif csv_line[2].startswith('.'):
                    csv_line[2] = csv_line[2].replace('.', '')
                print(csv_line[3], csv_line[2])
                csvfile_writer.writerow(csv_line)
            else:
                pass
            # print(EM_extract_node_name(i)[6])
        csvfile.close()

    def on_pushButton_csv2us_pressed(self):
        #s = QgsSettings()
        sqlite_DB_path = '{}{}{}'.format(self.HOME, os.sep,
                                         "pyarchinit_DB_folder")
        dbpath = QFileDialog.getOpenFileName(
            self,
            "Set file name",
            '',
            " csv pyarchinit us_table (*.csv)"
        )[0]
        filename = dbpath  # .split("/")[-1]
        s = sqlite_DB_path + '/export_csv2us.csv'
        try:
            conn = Connection()
            conn_str = conn.conn_str()
            conn_sqlite = conn.databasename()



            con = sq.connect(sqlite_DB_path + os.sep + conn_sqlite["db_name"])
            cur = con.cursor()


            with open(filename, 'r') as in_file, open(s, 'r+') as out_file:
                seen = set()  # set for fast O(1) amortized lookup
                for line in in_file:
                    if line in seen:
                        continue  # skip duplicate

                    seen.add(line)
                    out_file.write(line)
                dr = csv.DictReader(out_file)  # comma is default delimiter
                to_db = [(i['site'], i['area'], i['us'], i['unit_type'], i['i_stratigrafica']) for i in dr]

            cur.executemany(
                "INSERT INTO us_table (sito, area, us, unita_tipo,d_interpretativa ) VALUES (?, ?,?,?,?);",
                to_db)
            con.commit()
            con.close()

        except AssertionError as e:
            QMessageBox.warning(self, 'error', str(e), QMessageBox.Ok)
        else:
            QMessageBox.information(self, 'ok', 'done', QMessageBox.Ok)
        self.pushButton_view_all.click()

    def on_pushButton_fix_pressed(self):
        sito = self.comboBox_sito.currentText()
        area = self.comboBox_area.currentText()
        search_dict = {'sito': f"'{sito}'", 'area': f"'{area}'"}
        records = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
        self.list_rapporti.append(self.report_rapporti)
        #QMessageBox.information(self, 'Fix', str(self.list_rapporti), QMessageBox.Ok)

        for _ in records:
            self.checkBox_validation_rapp.setChecked(True)

            self.check_listoflist()
            self.check_inverse_relationships(self.list_rapporti)
        self.view_all()

    def check_listoflist(self):
        if self.checkBox_validation_rapp.isChecked():
            try:

                table_name = "self.tableWidget_rapporti"
                rowSelected_cmd = ("%s.selectedItems()") % (table_name)
                rowSelected = eval(rowSelected_cmd)
                rowIndex = (rowSelected[0].row())
                sito = str(self.comboBox_sito.currentText())
                area = str(self.comboBox_area.currentText())
                us_current = str(self.lineEdit_us.text())
                print(us_current)
                unit = str(self.comboBox_unita_tipo.currentText())
                us_item = self.tableWidget_rapporti.item(rowIndex, 1)
                us = str(us_item.text())
                # print(us)
                rapp_item = self.tableWidget_rapporti.item(rowIndex, 0)
                rapp = str(rapp_item.text())

                area_item = self.tableWidget_rapporti.item(rowIndex, 2)
                ar_ = str(area_item.text())
                sito_item = self.tableWidget_rapporti.item(rowIndex, 3)
                sito_ = str(sito_item.text())
                self.save_rapp()

                if rapp == 'Riempito da':
                    rapp = 'Riempie'
                elif rapp == 'Tagliato da':
                    rapp = 'Taglia'
                elif rapp == 'Coperto da':
                    rapp = 'Copre'
                elif rapp == 'Si appoggia a':
                    rapp = 'Gli si appoggia'
                elif rapp == 'Riempie':
                    rapp = 'Riempito da'
                elif rapp == 'Taglia':
                    rapp = 'Tagliato da'
                elif rapp == 'Copre':
                    rapp = 'Coperto da'
                elif rapp == 'Gli si appoggia':
                    rapp = 'Si appoggia a'
                elif rapp == 'Filled by':
                    rapp = 'Fills'
                elif rapp == 'Cut by':
                    rapp = 'Cuts'
                elif rapp == 'Covered by':
                    rapp = 'Covers'
                elif rapp == 'Abuts':
                    rapp = 'Supports'
                elif rapp == 'Fills':
                    rapp = 'Filled by'
                elif rapp == 'Cuts':
                    rapp = 'Cut by'
                elif rapp == 'Covers':
                    rapp = 'Covered by'
                elif rapp == 'Supports':
                    rapp = 'Abuts'

                elif rapp == '>>':
                    rapp = '<<'
                elif rapp == '<<':
                    rapp = '>>'
                elif rapp == '>':
                    rapp = '<'
                elif rapp == '<':
                    rapp = '>'
                search_dict = {'sito': "'" + str(sito_) + "'",
                               'area': "'" + str(ar_) + "'",
                               'us': us}
                u = Utility()
                search_dict = u.remove_empty_items_fr_dict(search_dict)
                res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)

                if bool(res):

                    items = self.tableWidget_rapporti.findItems(us, Qt.MatchExactly)
                    items_area = self.tableWidget_rapporti.findItems(ar_, Qt.MatchExactly)
                    items_sito = self.tableWidget_rapporti.findItems(sito_, Qt.MatchExactly)
                    self.on_pushButton_go_to_us_pressed()
                    self.checkBox_validation_rapp.setChecked(False)
                    items2 = self.tableWidget_rapporti.findItems(us_current, Qt.MatchExactly)
                    items_area2 = self.tableWidget_rapporti.findItems(area, Qt.MatchExactly)
                    items_sito2 = self.tableWidget_rapporti.findItems(sito, Qt.MatchExactly)
                    # QMessageBox.information(self, 'caso1', f"{str(len(items))} - {str(len(items2))}  - {str(len(items_area))} - {str(len(items_sito))} -  {str(len(items_area2))} - {str(len(items_sito2))}")
                    if str(len(items)) == '1' and str(
                            len(items2)) == '1':  # and str(len(items_area))=='1' and str(len(items_sito))=='1' and str(len(items_area2))=='3' and str(len(items_sito2))=='5':
                        try:
                            item = items2[0]
                            self.tableWidget_rapporti.setCurrentItem(item)
                            item_area = items_area2[0]
                            self.tableWidget_rapporti.setCurrentItem(item_area)
                            item_sito = items_sito2[0]
                            self.tableWidget_rapporti.setCurrentItem(item_sito)
                        except:
                            pass
                        y = self.tableWidget_rapporti.currentRow()
                        self.tableWidget_rapporti.setItem(y, 0, QtWidgets.QTableWidgetItem(rapp))
                        self.tableWidget_rapporti.setItem(y, 1, QtWidgets.QTableWidgetItem(us_current))
                        self.tableWidget_rapporti.setItem(y, 2, QtWidgets.QTableWidgetItem(area))
                        self.tableWidget_rapporti.setItem(y, 3, QtWidgets.QTableWidgetItem(sito))
                        self.save_rapp()

                        self.tableWidget_rapporti.selectRow(y)
                        self.on_pushButton_go_to_us_pressed()

                    elif str(len(items)) == '1' and str(
                            len(items2)) == '0':  # and str(len(items_area))=='1' and str(len(items_sito))=='1' and str(len(items_area2))=='0' and str(len(items_sito2))=='0':

                        self.on_pushButton_insert_row_rapporti_pressed()
                        self.tableWidget_rapporti.currentRow()
                        self.tableWidget_rapporti.setItem(0, 0, QtWidgets.QTableWidgetItem(rapp))
                        self.tableWidget_rapporti.setItem(0, 1, QtWidgets.QTableWidgetItem(us_current))
                        self.tableWidget_rapporti.setItem(0, 2, QtWidgets.QTableWidgetItem(area))
                        self.tableWidget_rapporti.setItem(0, 3, QtWidgets.QTableWidgetItem(sito))
                        self.save_rapp()
                        self.tableWidget_rapporti.selectRow(0)
                        self.on_pushButton_go_to_us_pressed()
                    else:
                        QMessageBox.warning(self, '', 'Controlla se hai duplicato una US o USM')

                elif not bool(res):

                    tf = self.unit_type_select()

                    self.DB_MANAGER.insert_number_of_us_records(sito_, ar_, us, tf)

                    self.on_pushButton_go_to_us_pressed()
                    self.on_pushButton_insert_row_rapporti_pressed()
                    self.tableWidget_rapporti.currentRow()

                    a = self.tableWidget_rapporti.setItem(0, 0, QtWidgets.QTableWidgetItem(rapp))
                    b = self.tableWidget_rapporti.setItem(0, 1, QtWidgets.QTableWidgetItem(us_current))
                    c = self.tableWidget_rapporti.setItem(0, 2, QtWidgets.QTableWidgetItem(area))
                    d = self.tableWidget_rapporti.setItem(0, 3, QtWidgets.QTableWidgetItem(sito))

                    self.save_rapp()
                    self.tableWidget_rapporti.selectRow(0)
                    self.on_pushButton_go_to_us_pressed()

            except:
                pass  # QMessageBox.warning(self, 'error', str(e), QMessageBox.Ok)


        else:
            pass  # QMessageBox.warning(self, 'error', 'Please select a rapport', QMessageBox.Ok)

    def log_message(self, message):
        """Aggiunge un messaggio alla listWidget"""
        self.listWidget_rapp.addItem(str(message))

    def check_inverse_relationships(self, unverified_list):
        if hasattr(self, '_is_processing'):
            return

        try:
            self._is_processing = True
            self.listWidget_rapp.clear()

            # Parse delle linee
            valid_lines = []
            if len(unverified_list) == 1 and isinstance(unverified_list[0], str):
                lines = unverified_list[0].split('\n')
                for line in lines:
                    if line.strip() and "Site:" in line and "SU:" in line and "relationships not verified" in line:
                        if not line.startswith("Control report"):
                            valid_lines.append(line.strip())

            if not valid_lines:
                return

            total_rapporti = len(valid_lines)
            self.progressBar_3.setMaximum(total_rapporti)
            self.progressBar_3.setValue(0)

            inverse_rapp_dict = {
                'Riempito da': 'Riempie',
                'Tagliato da': 'Taglia',
                'Coperto da': 'Copre',
                'Si appoggia a': 'Gli si appoggia',
                'Riempie': 'Riempito da',
                'Taglia': 'Tagliato da',
                'Copre': 'Coperto da',
                'Gli si appoggia': 'Si appoggia a',
                'Filled by': 'Fills',
                'Cut by': 'Cuts',
                'Covered by': 'Covers',
                'Abuts': 'Supports',
                'Fills': 'Filled by',
                'Cuts': 'Cut by',
                'Covers': 'Covered by',
                'Supports': 'Abuts',
                'Same as': 'Same as',
                'Uguale a': 'Uguale a',
                'Connected to': 'Connected to',
                'Si lega a': 'Si lega a',
                '>>': '<<',
                '<<': '>>',
                '>': '<',
                '<': '>'
            }

            old_state = self.checkBox_validation_rapp.isChecked()
            self.checkBox_validation_rapp.setChecked(False)
            fixed_count = 0

            try:
                for i, line in enumerate(valid_lines):
                    try:
                        # Parse della linea
                        self.log_message(f"\nProcesso linea {i + 1}/{total_rapporti}: {line}")

                        parts = line.split(", ")
                        sito = parts[0].split("'")[1]
                        area = parts[1].split("'")[1]

                        words = parts[2].split()
                        us_indices = [i for i, word in enumerate(words) if word == "SU:" or word== "US:"]

                        # Single clean handling of US values
                        try:
                            us_origine = str(words[us_indices[0] + 1]).strip()
                            us_target = str(words[us_indices[1] + 1]).strip()

                            if not us_origine or not us_target:
                                raise ValueError("US values cannot be empty")

                        except (IndexError, ValueError) as e:
                            self.log_message(f"Error processing US values: {str(e)}")
                            continue

                        rapp_words = words[us_indices[0] + 2:us_indices[1]]
                        rapp = " ".join(rapp_words)

                        self.log_message(f"Analisi: {us_origine} {rapp} {us_target}")

                        inverse_rapp = inverse_rapp_dict.get(rapp)
                        if not inverse_rapp:
                            self.log_message(f"Rapporto {rapp} non trovato nel dizionario")
                            continue

                        # Vai alla scheda target usando il tuo metodo
                        search_dict = {
                            'sito': "'" + str(sito) + "'",
                            'area': "'" + str(area) + "'",
                            'us': str(us_target)  # Converti a stringa
                        }

                        u = Utility()
                        search_dict = u.remove_empty_items_fr_dict(search_dict)
                        res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)

                        if res:
                            self.empty_fields()
                            self.DATA_LIST = []
                            for i in res:
                                self.DATA_LIST.append(i)
                            self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                            self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                            self.fill_fields()
                            self.BROWSE_STATUS = "b"
                            self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                            self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)

                            # Controlla se il rapporto inverso esiste già
                            rapporto_exists = False
                            for row in range(self.tableWidget_rapporti.rowCount()):
                                tipo_rapp = self.tableWidget_rapporti.item(row, 0)
                                us_rapp = self.tableWidget_rapporti.item(row, 1)

                                if (tipo_rapp and us_rapp and
                                        tipo_rapp.text() == inverse_rapp and
                                        us_rapp.text() == us_origine):
                                    rapporto_exists = True
                                    break

                            # Se non esiste, aggiungilo
                            if not rapporto_exists:
                                self.on_pushButton_insert_row_rapporti_pressed()
                                self.tableWidget_rapporti.setItem(0, 0, QtWidgets.QTableWidgetItem(inverse_rapp))
                                self.tableWidget_rapporti.setItem(0, 1, QtWidgets.QTableWidgetItem(us_origine))
                                self.tableWidget_rapporti.setItem(0, 2, QtWidgets.QTableWidgetItem(area))
                                self.tableWidget_rapporti.setItem(0, 3, QtWidgets.QTableWidgetItem(sito))
                                self.save_rapp()
                                fixed_count += 1
                                self.log_message(
                                    f"Aggiunto rapporto {inverse_rapp} {us_origine} nella scheda {us_target}")

                        # Aggiorna progress bar
                        self.progressBar_3.setValue(i + 1)

                    except Exception as e:
                        self.log_message(f"Errore: {str(e)}")
                        continue

            finally:
                self.checkBox_validation_rapp.setChecked(old_state)
                self.progressBar_3.setValue(0)

            try:# Risultati finali
                if fixed_count > 0:
                    QMessageBox.information(self, "Completato", f"Corretti {fixed_count} rapporti su {total_rapporti}")
            except Exception as e:
                self.log_message(f"Errore: {str(e)}")
                raise

            # Aggiorna la lista dei rapporti da validare
            self.on_pushButton_h_check_pressed()

        finally:
            delattr(self, '_is_processing')

    def unit_type_select(self):
        try:
            dialog = QInputDialog()
            dialog.resize(QtCore.QSize(200, 100))
            if self.L=='it':
                items_st=('US','USM','USVA','USVB','USVC','USD','CON','VSF','SF','SUS','Combinar','Extractor','DOC','property')
            else:
                items_st=('SU','WSU','USVA','USVB','USVC','USD','CON','VSF','SF','SUS','Combinar','Extractor','DOC','property')
            ID_U = dialog.getItem(self, 'Type Unit', "Insert Unit Type",items_st, 0, False)
            Unit = str(ID_U[0])
            return Unit
        except KeyError as e:
            print(str(e))

    def search_rapp(self):
        # Clear current selection.
        #self.tableWidget_rapporti.setCurrentItem(None)
        s=''
        if not s:
            # Empty string, don't search.
            return

        matching_items = self.tableWidget_rapporti.findItems('1',Qt.MatchContains)
        if matching_items:
            # We have found something.
            item = matching_items[0]  # Take the first.
            self.tableWidget_rapporti.setCurrentItem(item)



    def check_v(self):
        if self.comboBox_per_iniz.currentText() =='':
            self.checkBox_validate.setHidden(True)
        else:
            self.checkBox_validate.setHidden(False)


    def change_label(self):
        if self.comboBox_unita_tipo.currentText()=='DOC':
            self.label_5.setText('Riferimento documentazione')
            self.comboBox_def_intepret.setHidden(True)
            self.mQgsFileWidget.setGeometry(486,128,334,20)
            self.mQgsFileWidget.show()
            self.toolButton_file_doc.show()
        else:
            self.mQgsFileWidget.setHidden(True)
            self.toolButton_file_doc.setHidden(True)
            self.comboBox_def_intepret.show()
        if self.comboBox_unita_tipo.currentText()=='property':
            self.label_5.setText('Descrizione della proprietà')
        if self.comboBox_unita_tipo.currentText().startswith('USV'):
            self.label_5.setText('Descrizione della Unità Str. Virtuale')
        if self.comboBox_unita_tipo.currentText()=='CON':
            self.label_5.setText('Riferimento alla Unità Continuativa')
        if self.comboBox_unita_tipo.currentText()=='CON':
            self.label_5.setText('Riferimento alla Unità Continuativa')
        if self.comboBox_unita_tipo.currentText()=='Combinar':
            self.label_5.setText('Descrizione connettore')
        if self.comboBox_unita_tipo.currentText()=='Extractor':
            self.label_5.setText('Descrizione estrattore')
        if self.comboBox_unita_tipo.currentText()=='SUS':
            self.label_5.setText('Descrizione')
        if self.comboBox_unita_tipo.currentText()=='SF':
            self.label_5.setText('Descrizione')


    def refresh(self):

        for i in self.DATA_LIST:
            self.us_t()
        return


    def charge_insert_ra(self):
        try:
            current_sito = "'"+str(self.comboBox_sito.currentText())+"'"
            current_area = "'" + str(self.DATA_LIST[self.REC_CORR].area) + "'"
            current_us = "'" + str(self.DATA_LIST[self.REC_CORR].us) + "'"

            # Ensure current_area and current_us are not None
            if current_area is None or current_us is None:
                return

            search_dict_inv = {
                'sito': current_sito,
                'area': current_area,
                'us': current_us
            }

            inv_vl = self.DB_MANAGER.query_bool(search_dict_inv, 'INVENTARIO_MATERIALI')
            inv_vl2 = self.DB_MANAGER.query_bool(search_dict_inv, 'POTTERY')

            # Build lists using list comprehensions
            inv_list = [f"{item.n_reperto}" for item in inv_vl if
                        item.n_reperto]
            inv_list2 = [f"{item.id_number})" for item in inv_vl2 if item.id_number]

            # Sort and remove duplicates
            inv_list = sorted(set(inv_list))
            inv_list2 = sorted(set(inv_list2))

            # Update the QComboBox
            self.comboBox_ref_ra.clear()
            self.comboBox_ref_ra.addItems(inv_list + inv_list2)

            # Set the edit text based on the browse status
            browse_status = self.STATUS_ITEMS.get(self.BROWSE_STATUS, "")
            if browse_status in ["Trova", "Finden", "Find"]:
                self.comboBox_ref_ra.setEditText("")
            elif browse_status in ["Usa", "Aktuell", "Current"] and self.DATA_LIST:
                self.comboBox_ref_ra.setEditText(self.DATA_LIST[self.rec_num].ref_ra)

        except Exception as e:
            print(f"An error occurred in charge_insert_ra: {e}")

    def charge_insert_ra_pottery(self):
        try:

            us = str(self.lineEdit_us.text())

            search_dict_inv = {

                'sito': "'" + str(self.comboBox_sito.currentText()) + "'",
                'area': "'" + str(eval("self.DATA_LIST[int(self.REC_CORR)].area")) + "'",
                'us': "'" + str(eval("self.DATA_LIST[int(self.REC_CORR)].us")) + "'"
            }

            inv_vl = self.DB_MANAGER.query_bool(search_dict_inv, 'POTTERY')
            inv_list = []
            for i in range(len(inv_vl)):
                inv_list.append(str(inv_vl[i].id_number))
                inv_list.sort()
            try:
                inv_vl.remove('')
            except:
                pass
            self.comboBox_ref_ra.clear()
            self.comboBox_ref_ra.addItems(self.UTILITY.remove_dup_from_list(inv_list))
            if self.STATUS_ITEMS[self.BROWSE_STATUS] == "Trova" or "Finden" or "Find":
                self.comboBox_ref_ra.setEditText("")
            elif self.STATUS_ITEMS[self.BROWSE_STATUS] == "Usa" or "Aktuell " or "Current":
                if len(self.DATA_LIST) > 0:
                    try:
                        self.comboBox_ref_ra.setEditText(self.DATA_LIST[self.rec_num].ref_ra)
                    except:
                        pass
        except:
            pass

    def listview_us(self):
        '''
            This function is used to filter the 'Unità Stratigrafiche' table.
        '''
        if self.checkBox_query.isChecked():
            conn = Connection()
            conn_str = conn.conn_str()
            conn_sqlite = conn.databasename()
            conn_user = conn.datauser()
            conn_host = conn.datahost()
            conn_port = conn.dataport()
            port_int  = conn_port["port"]
            port_int.replace("'", "")
            conn_password = conn.datapassword()
            sito_set= conn.sito_set()
            sito_set_str = sito_set['sito_set']
            test_conn = conn_str.find('sqlite')
            if test_conn == 0:
                sqlite_DB_path = '{}{}{}'.format(self.HOME, os.sep,
                                           "pyarchinit_DB_folder")

                db = QSqlDatabase("QSQLITE")
                db.setDatabaseName(sqlite_DB_path +os.sep+ conn_sqlite["db_name"])
                db.open()
                self.model_a = QSqlTableModel(db = db)
                self.table.setModel(self.model_a)
                self.model_a.setTable("us_table")
                self.model_a.setEditStrategy(QSqlTableModel.OnManualSubmit)
                self.pushButton_submit.clicked.connect(self.submit)
                self.pushButton_revert.clicked.connect(self.model_a.revertAll)
                column_titles = {
                    "sito": "SITO",
                    "area": "Area",
                    "us": "US"}
                for n, t in column_titles.items():
                    idx = self.model_a.fieldIndex( n)
                    self.model_a.setHeaderData( idx, Qt.Horizontal, t)
                if bool (sito_set_str):
                    filter_str = "sito = '{}'".format(str(self.comboBox_sito.currentText()))
                    self.model_a.setFilter(filter_str)
                    self.model_a.select()
                else:
                    self.model_a.select()

                uri = QgsDataSourceUri()
                uri.setDatabase(sqlite_DB_path +os.sep+ conn_sqlite["db_name"])
                schema = ''
                table = 'us_table'
                geom_column = ''
                uri.setDataSource(schema, table,geom_column)
                vlayer = QgsVectorLayer(uri.uri(), table, 'spatialite')
                pr = vlayer.dataProvider()
                fi= pr.fields().names()[1:-1]

                self.field.clear()
                self.field.addItems(fi)
                try:
                    self.search_1.clearEditText()
                    self.search_1.update()
                    self.value_check()
                except:
                    pass
            else:


                db = QSqlDatabase.addDatabase("QPSQL")
                db.setHostName(conn_host["host"])
                db.setDatabaseName(conn_sqlite["db_name"])
                db.setPort(int(port_int))
                db.setUserName(conn_user['user'])
                db.setPassword(conn_password['password'])
                db.open()
                self.model_a = QSqlTableModel(db = db)
                self.table.setModel(self.model_a)
                self.model_a.setTable("us_table")
                self.model_a.setEditStrategy(QSqlTableModel.OnManualSubmit)
                self.pushButton_submit.clicked.connect(self.submit)
                self.pushButton_revert.clicked.connect(self.model_a.revertAll)
                if bool (sito_set_str):
                    filter_str = "sito = '{}'".format(str(self.comboBox_sito.currentText()))
                    self.model_a.setFilter(filter_str)
                    self.model_a.select()
                else:
                    self.model_a.select()

                uri = QgsDataSourceUri()
                uri.setConnection(conn_host["host"], conn_port["port"], conn_sqlite["db_name"], conn_user['user'], conn_password['password'])
                schema = 'public'
                table = 'us_table'
                geom_column = ''
                uri.setDataSource(schema, table,geom_column)
                vlayer = QgsVectorLayer(uri.uri(), table, 'postgres')
                pr = vlayer.dataProvider()
                fi= pr.fields().names()[1:-1]
                pr = vlayer.dataProvider()
                fi= pr.fields().names()[1:-1]

                self.field.clear()
                self.field.addItems(fi)
                try:
                    self.search_1.clearEditText()
                    self.search_1.update()
                    self.value_check()
                except:
                    pass
        else:
            self.checkBox_query.setChecked(False)
    def submit(self):
        if self.checkBox_query.isChecked():
            self.model_a.database().transaction()
            if self.model_a.submitAll():
                self.model_a.database().commit()
                if self.L=='it':
                    QMessageBox.information(self, "Record",  "record salvato")
                elif self.L=='de':
                    QMessageBox.information(self, "Datensatz",  "Datensatz gespeichert")
                else:
                    QMessageBox.information(self, "Record",  "record saved")

            else:
                self.model_a.database().rollback()
                if self.L=='it':
                    QMessageBox.warning(self, "Cached Table",
                            "Il db ha segnalato un errore: %s" % self.model_a.lastError().text())

                elif self.L=='de':
                    QMessageBox.warning(self, "Cached Table",
                            "Die Datenbank meldete einen Fehler: %s" % self.model_a.lastError().text())

                else:
                    QMessageBox.warning(self, "Cached Table",
                            "The database reported an error: %s" % self.model_a.lastError().text())

        else:
            self.checkBox_query.setChecked(False)


    def value_check(self):


        try:

            if self.field.currentTextChanged:
                sito_vl2 = self.UTILITY.tup_2_list_III(self.DB_MANAGER.group_by('us_table', self.field.currentText(),'US'))


                sito_vl2.remove('')


                self.search_1.clear()

                sito_vl2.sort()

                self.search_1.addItems(sito_vl2)
                self.search_1.update()

        except:
            pass#QMessageBox.warning(self, "Attenzione", str(e), QMessageBox.Ok)

    def update_filter(self, s):
        '''
            This function is used to filter the 'Unità Stratigrafiche' table.
        '''

        if self.checkBox_query.isChecked():
            conn = Connection()
            conn_str = conn.conn_str()
            sito_set= conn.sito_set()
            sito_set_str = sito_set['sito_set']
            test_conn = conn_str.find('sqlite')
            s_field = self.field.currentText()
            s = re.sub("[\w_][\W_] +", "", s)
            if test_conn == 0:


                try:
                    if bool(sito_set_str):
                        filter_str = "{} LIKE '%{}%' and sito = '{}'".format(s_field,s,str(self.comboBox_sito.currentText()))
                        if bool(filter_str):
                            self.model_a.setFilter(filter_str)
                            self.model_a.select()
                        else:
                            pass

                    else:
                        filter_str = "{} LIKE '%{}%'".format(s_field,s)
                        if bool(filter_str):
                            self.model_a.setFilter(filter_str)
                            self.model_a.select()
                        else:
                            pass
                except :
                    pass#QMessageBox.warning(self, "Warning", str(e), QMessageBox.Ok)
            else:
                try:
                    if bool(sito_set_str):
                        filter_str = "{} LIKE '%{}%' and sito = '{}'".format(s_field,s,str(self.comboBox_sito.currentText()))
                        if bool(filter_str):
                            self.model_a.setFilter(filter_str)
                            self.model_a.select()
                        else:
                            pass
                    else:
                        filter_str = "{} LIKE '%{}%'".format(s_field,s)
                        if bool(filter_str):
                            self.model_a.setFilter(filter_str)
                            self.model_a.select()
                        else:
                            pass
                except :
                    pass#QMessageBox.warning(self, "Warning", str(e), QMessageBox.Ok)
        else:
            self.checkBox_query.setChecked(False)

    def on_pushButton_globalsearch_pressed(self):
        '''
            This function is used to search for a specific record in the database.
        '''
        self.search.showSearchDialog()



    def format_struttura_item(self, struttura):
        return f"{struttura.sigla_struttura}-{struttura.numero_struttura}"

    def charge_struttura_list(self):
        '''
            This function charges the 'Struttura' combobox with the values from the 'Struttura' table.
        '''
        FIND_STATUS_KEYS = ["Trova", "Finden", "Find"]
        CURRENT_STATUS_KEYS = ["Usa", "Aktuell", "Current"]

        site = str(self.comboBox_sito.currentText())
        search_dict = {
            'sito': f"'{site}'"
        }
        structures_query_result = self.DB_MANAGER.query_bool(search_dict, 'STRUTTURA')
        structures_list = [self.format_struttura_item(s) for s in structures_query_result]

        self.comboBox_struttura.clear()
        self.comboBox_struttura.addItems(self.UTILITY.remove_dup_from_list(structures_list))

        current_status = self.STATUS_ITEMS[self.BROWSE_STATUS]
        if current_status in FIND_STATUS_KEYS:
            self.comboBox_struttura.setEditText("")
        elif current_status in CURRENT_STATUS_KEYS:
            #if len(self.DATA_LIST) > 0:
            try:
                self.comboBox_struttura.setEditText(self.format_struttura_item(self.DATA_LIST[self.rec_num]))
            except Exception as e:
                # You might consider logging the error messages to improve
                # debugging. Replace `print` with a logger as necessary.
                pass#QMessageBox.warning(self, 'Warning', f"Error setting edit text: {e}")

    def geometry_unitastratigrafiche(self):
        '''
            This function charges the 'Posizione' combobox with the values from the 'Unità Stratigrafiche' table.
        '''
        try:
            # Usa i valori correnti dei widget invece di accedere a DATA_LIST
            sito = str(self.comboBox_sito.currentText())
            area = str(self.comboBox_area.currentText())
            us = str(self.lineEdit_us.text())

            search_dict = {
                'scavo_s': f"'{sito}'",
                'area_s': f"'{area}'",
                'us_s': f"'{us}'"
            }

            geometry_vl = self.DB_MANAGER.query_bool(search_dict, 'PYUS')
            geometry_list = [str(item.coord) for item in geometry_vl if item.coord]

            self.comboBox_posizione.clear()
            self.comboBox_posizione.addItems(self.UTILITY.remove_dup_from_list(geometry_list))

            if self.STATUS_ITEMS[self.BROWSE_STATUS] in ["Trova", "Finden", "Find"]:
                self.comboBox_posizione.setEditText("")
            elif self.STATUS_ITEMS[self.BROWSE_STATUS] in ["Usa", "Aktuell", "Current"]:
                if hasattr(self, 'DATA_LIST') and self.DATA_LIST:
                    try:
                        self.comboBox_posizione.setEditText(self.DATA_LIST[self.rec_num].posizione)
                        self.comboBox_posizione.show()
                    except (IndexError, AttributeError) as e:
                        print(f"Errore in geometry_unitastratigrafiche: {str(e)}")
        except Exception as e:
            print(f"Errore in geometry_unitastratigrafiche: {str(e)}")

    def charge_periodo_iniz_list(self):
        '''
            This function charges the 'Periodo' combobox with the values from the 'Periodizzazione' table.
        '''
        try:

            sito = str(self.comboBox_sito.currentText())
            area = str(self.comboBox_area.currentText())
            search_dict = {
                'sito': "'" + sito + "'",
                #'area': "'" + str(eval("self.DATA_LIST[int(self.REC_CORR)].area")) + "'",
            }
            periodo_vl = self.DB_MANAGER.query_bool(search_dict, 'PERIODIZZAZIONE')
            periodo_list = []
            for i in range(len(periodo_vl)):
                periodo_list.append(str(periodo_vl[i].periodo))
            try:
                periodo_vl.remove('')
            except:
                pass
            #
            self.comboBox_per_iniz.clear()
            self.comboBox_per_iniz.addItems(self.UTILITY.remove_dup_from_list(periodo_list))
            if self.STATUS_ITEMS[self.BROWSE_STATUS] == "Trova" or "Finden" or "Find":
                self.comboBox_per_iniz.setEditText("")
            elif self.STATUS_ITEMS[self.BROWSE_STATUS] == "Usa" or "Aktuell " or "Current":
                if len(self.DATA_LIST) > 0:
                    try:
                        self.comboBox_per_iniz.setEditText(self.DATA_LIST[self.rec_num].periodo_iniziale)
                        self.comboBox_per_iniz.show()
                    except:
                        pass  # non vi sono periodi per questo scavo
        except:
            pass


    def charge_periodo_fin_list(self):
        '''
            This function charges the 'Periodo' combobox with the values from the 'Periodizzazione' table.
        '''
        try:

            sito = str(self.comboBox_sito.currentText())
            area = str(self.comboBox_area.currentText())
            search_dict = {
                'sito': "'" + sito + "'",
                #'area': "'" + str(eval("self.DATA_LIST[int(self.REC_CORR)].area")) + "'",
            }
            periodo_vl = self.DB_MANAGER.query_bool(search_dict, 'PERIODIZZAZIONE')
            periodo_list = []
            for i in range(len(periodo_vl)):
                periodo_list.append(str(periodo_vl[i].periodo))
            try:
                periodo_vl.remove('')
            except:
                pass
            self.comboBox_per_fin.clear()
            self.comboBox_per_fin.addItems(self.UTILITY.remove_dup_from_list(periodo_list))
            if self.STATUS_ITEMS[self.BROWSE_STATUS] == "Trova" or "Finden" or "Find":
                self.comboBox_per_fin.setEditText("")
            elif self.STATUS_ITEMS[self.BROWSE_STATUS] == "Usa" or "Aktuell " or "Current":
                if len(self.DATA_LIST) > 0:
                    try:
                        self.comboBox_per_fin.setEditText(self.DATA_LIST[self.rec_num].periodo_iniziale)
                    except:
                        pass
        except:
            pass  # non vi sono periodi per questo scavo
    def charge_fase_iniz_list(self):

        '''
            This function charges the 'Fase' combobox with the values from the 'Periodizzazione' table.
        '''
        try:
            search_dict = {
                'sito': "'" + str(self.comboBox_sito.currentText()) + "'",
                'periodo': "'" + str(self.comboBox_per_iniz.currentText()) + "'",
            }
            fase_list_vl = self.DB_MANAGER.query_bool(search_dict, 'PERIODIZZAZIONE')
            fase_list = []
            for i in range(len(fase_list_vl)):
                fase_list.append(str(fase_list_vl[i].fase))
            try:
                fase_list.remove('')
            except:
                pass
            self.comboBox_fas_iniz.clear()
            fase_list.sort()
            self.comboBox_fas_iniz.addItems(self.UTILITY.remove_dup_from_list(fase_list))
            if self.STATUS_ITEMS[self.BROWSE_STATUS] == "Trova" or "Finden" or "Find":
                self.comboBox_fas_iniz.setEditText("")
            else:
                self.comboBox_fas_iniz.setEditText(self.DATA_LIST[self.rec_num].fase_iniziale)
        except:
            pass
    def charge_fase_fin_list(self):
        '''
            This function charges the 'Fase' combobox with the values from the 'Periodizzazione' table.
        '''
        #if self.comboBox_fas_fin.activated:
        try:
            search_dict = {
                'sito': "'" + str(self.comboBox_sito.currentText()) + "'",
                'periodo': "'" + str(self.comboBox_per_fin.currentText()) + "'"
            }
            fase_list_vl = self.DB_MANAGER.query_bool(search_dict, 'PERIODIZZAZIONE')
            fase_list = []
            for i in range(len(fase_list_vl)):
                fase_list.append(str(fase_list_vl[i].fase))
            try:
                fase_list.remove('')
            except:
                pass
            self.comboBox_fas_fin.clear()
            fase_list.sort()
            self.comboBox_fas_fin.addItems(self.UTILITY.remove_dup_from_list(fase_list))
            if self.STATUS_ITEMS[self.BROWSE_STATUS] == "Trova" or "Finden" or "Find":
                self.comboBox_fas_fin.setEditText("")
            else:
                self.comboBox_fas_fin.setEditText(self.DATA_LIST[self.rec_num].fase_finale)
        except:
            pass

    def charge_datazione_list(self):
        '''
            This function charges the 'Datazione' combobox with the values from the 'Periodizzazione' table.
        '''
        try:
            search_dict_iniz = {
                'sito': "'" + str(self.comboBox_sito.currentText()) + "'",
                'periodo': "'" + str(self.comboBox_per_iniz.currentText()) + "'",
                'fase': "'" + str(self.comboBox_fas_iniz.currentText()) + "'"
            }
            search_dict_fin = {
                'sito': "'" + str(self.comboBox_sito.currentText()) + "'",
                'periodo': "'" + str(self.comboBox_per_fin.currentText()) + "'",
                'fase': "'" + str(self.comboBox_fas_fin.currentText()) + "'"
            }
            datazione_list_vl_iniz = self.DB_MANAGER.query_bool(search_dict_iniz, 'PERIODIZZAZIONE')
            datazione_list_vl_fin = self.DB_MANAGER.query_bool(search_dict_fin, 'PERIODIZZAZIONE')

            datazione_list_iniz = [str(item.datazione_estesa) for item in datazione_list_vl_iniz if
                                   str(item.datazione_estesa) != '']
            datazione_list_fin = [str(item.datazione_estesa) for item in datazione_list_vl_fin if
                                  str(item.datazione_estesa) != '']

            self.lineEdit_datazione.clear()
            if datazione_list_iniz:
                datazione_list_iniz.sort()
                periodo_iniziale = datazione_list_iniz[-1]
                if datazione_list_fin:
                    datazione_list_fin.sort()
                    periodo_finale = datazione_list_fin[-1]
                    if periodo_finale and str(self.comboBox_per_fin.currentText()) != '':
                        self.lineEdit_datazione.setText(f"{periodo_iniziale} / {periodo_finale}")
                        self.lineEdit_datazione.update()
                    else:
                        self.lineEdit_datazione.setText(periodo_iniziale)
                        self.lineEdit_datazione.update()
                else:
                    self.lineEdit_datazione.setText(periodo_iniziale)
                    self.lineEdit_datazione.update()
            else:
                self.lineEdit_datazione.setText("")
                self.lineEdit_datazione.update()
        except Exception as e:
            QMessageBox.warning(self, "Warning", f"An error occurred while charging 'Datazione {e}'.", QMessageBox.Ok)


    # This function should be connected to the button click event
    def update_dating(self):
        '''
            This function updates the 'Dating' field for all US records in the database.
        '''
        self.charge_datazione_list()
        try:
            updates_made = self.DB_MANAGER.update_us_dating_from_periodizzazione(self.comboBox_sito.currentText())
            if updates_made > 0:
                # Inform the user that updates have been made
                print(f"All 'Dating' fields have been updated successfully. "
                      f"Total updates made: {updates_made}")

            else:
                # Inform the user that no updates were necessary
                pass#QMessageBox.information(self, "No Updates", "No 'Dating' fields needed to be updated.",
                                        #QMessageBox.Ok)
        except Exception as e:
            print(f"An error occurred while updating 'Dating': {e}")

    def on_pushButton_draw_doc_pressed(self):
        sito = str(self.comboBox_sito.currentText())
        area = str(self.comboBox_area.currentText())
        us = str(self.lineEdit_us.text())
        table_name = "self.tableWidget_documentazione"
        rowSelected_cmd = ("%s.selectedIndexes()") % (table_name)
        rowSelected = eval(rowSelected_cmd)
        rowIndex = (rowSelected[0].row())
        tipo_doc_item = self.tableWidget_documentazione.item(rowIndex, 0)
        nome_doc_item = self.tableWidget_documentazione.item(rowIndex, 1)
        tipo_doc = str(tipo_doc_item.text())
        nome_doc = str(nome_doc_item.text())
        lista_draw_doc = [sito, area, us, tipo_doc, nome_doc]
        self.pyQGIS.charge_vector_layers_doc_from_scheda_US(lista_draw_doc)
    def save_us(self):
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.BROWSE_STATUS == "b":
            if self.data_error_check() == 0:
                if self.records_equal_check() == 1:
                    if self.L=='it':
                        self.update_if(QMessageBox.warning(self, 'Attenzione',
                                                           "Il record e' stato modificato. Vuoi salvare le modifiche? \n Clicca OK per salvare o Annulla per abortire.\n Poi riselezione l'US su cui vuoi andare",QMessageBox.Ok | QMessageBox.Cancel))
                    elif self.L=='de':
                        self.update_if(QMessageBox.warning(self, 'Error',
                                                           "Der Record wurde geändert. Möchtest du die Änderungen speichern?",
                                                           QMessageBox.Ok | QMessageBox.Cancel))

                    else:
                        self.update_if(QMessageBox.warning(self, 'Error',
                                                           "The record has been changed. Do you want to save the changes?",
                                                           QMessageBox.Ok | QMessageBox.Cancel))
                    self.SORT_STATUS = "n"
                    self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                    self.enable_button(1)
                    self.fill_fields(self.REC_CORR)

            else:
                pass


        else:
            if self.data_error_check() == 0:
                test_insert = self.insert_new_rec()
                if test_insert == 1:
                    self.empty_fields()
                    self.label_sort.setText(self.SORTED_ITEMS["n"])
                    self.charge_list()
                    self.set_sito()
                    self.charge_records_n()
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), len(self.DATA_LIST) - 1
                    self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)

                    self.setComboBoxEditable(["self.comboBox_sito"], 1)

                    self.fill_fields(self.REC_CORR)
                    self.enable_button(1)

            else:
                pass




    def selectRows(self):
        # seleziona tutte le row della tablewidget dei rapporti
        for row in range(self.tableWidget_rapporti.rowCount()):
            table_item = self.tableWidget_rapporti.item(row, 1)
            row_data = table_item.data(QtCore.Qt.UserRole)
            row_id = row_data
            self.tableWidget_rapporti.selectRow(row)

    def on_pushButton_update_pressed(self):
        sito = "'"+self.comboBox_sito.currentText()+"'"
        area = "'"+self.comboBox_area.currentText()+"'"
        search_dict = {'sito': sito, 'area': area}
        records = self.DB_MANAGER.query_bool(search_dict,
                                             self.MAPPER_TABLE_CLASS)  # carica tutti i dati di uno scavo ordinati per numero di US
        self.selectRows()
        for rec in range(len(records)):
            self.selectRows()
            self.on_pushButton_next_rec_pressed()
            self.us_t()

            self.save_rapp()
            # Calculate the progress as a percentage
            value = (float(rec) / float(len(records))) * 100
            # Convert the progress value to an integer
            int_value = int(value)
            # Update the progress bar with the integer value
            self.progressBar_2.setValue(int_value)
            QApplication.processEvents()

        self.progressBar_2.reset()
    def us_t(self):
        if self.checkBox_validate.isChecked():
            try:

                table_name = "self.tableWidget_rapporti"

                rowSelected_cmd = ("%s.selectedItems()") % (table_name)
                rowSelected = eval(rowSelected_cmd)

                for i  in rowSelected:
                    s= self.tableWidget_rapporti.rowCount()
                    self.tableWidget_rapporti2.setRowCount(s)
                    rowIndex = (i.row())
                    sito = str(self.comboBox_sito.currentText())
                    area = str(self.comboBox_area.currentText())

                    us_item = self.tableWidget_rapporti.item(rowIndex, 1)
                    sito_item = self.tableWidget_rapporti.item(rowIndex, 3)
                    area_item = self.tableWidget_rapporti.item(rowIndex, 2)
                    sito_ = str(sito_item.text())
                    area_ = str(area_item.text())
                    us_ = str(us_item.text())
                    rapp_item = self.tableWidget_rapporti.item(rowIndex, 0)
                    rapp_ = str(rapp_item.text())

                    search_dict = {'sito': "'" + sito_ + "'",
                                   'area': "'" + area_ + "'",
                                   'us': us_}
                    u = Utility()
                    search_dict = u.remove_empty_items_fr_dict(search_dict)
                    res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)



                    for p in res:
                        #

                        self.tableWidget_rapporti2.setItem(rowIndex,0,QtWidgets.QTableWidgetItem(rapp_))
                        self.tableWidget_rapporti2.setItem(rowIndex,1,QtWidgets.QTableWidgetItem(us_))
                        self.tableWidget_rapporti2.setItem(rowIndex,2,QtWidgets.QTableWidgetItem(p.unita_tipo))
                        self.tableWidget_rapporti2.setItem(rowIndex,3,QtWidgets.QTableWidgetItem(p.d_interpretativa))
                        self.tableWidget_rapporti2.setItem(rowIndex,4,QtWidgets.QTableWidgetItem(p.periodo_iniziale+'-'+p.fase_iniziale))
                        self.tableWidget_rapporti2.setItem(rowIndex,5,QtWidgets.QTableWidgetItem(area_))
                        self.tableWidget_rapporti2.setItem(rowIndex,6,QtWidgets.QTableWidgetItem(sito_))
                    self.tableWidget_rapporti2.update()

            except Exception as e:
                QMessageBox.warning(self,'',str(e))
        else:
            pass

    def on_pushButton_go_to_us_pressed(self):
        try:
            if self.BROWSE_STATUS == "b":
                if self.data_error_check() == 0:
                    if self.records_equal_check() == 1:
                        # Store the user response
                        if self.L == 'it':
                            response = QMessageBox.warning(self, 'Attenzione',
                                                           "Il record e' stato modificato. Vuoi salvare le modifiche? \n Clicca OK per salvare o Annulla per abortire.\n Poi riselezione l'US su cui vuoi andare",
                                                           QMessageBox.Ok | QMessageBox.Cancel)
                        elif self.L == 'de':
                            response = QMessageBox.warning(self, 'Error',
                                                           "Der Record wurde geändert. Möchtest du die Änderungen speichern?",
                                                           QMessageBox.Ok | QMessageBox.Cancel)
                        else:
                            response = QMessageBox.warning(self, 'Error',
                                                           "The record has been changed. Do you want to save the changes?",
                                                           QMessageBox.Ok | QMessageBox.Cancel)

                        # Check the user response and act accordingly
                        if response == QMessageBox.Ok:
                            self.update_if(response)
                            self.SORT_STATUS = "n"
                            self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                            self.enable_button(1)
                            #self.fill_fields(self.REC_CORR)
                            table_name = "self.tableWidget_rapporti"
                            rowSelected_cmd = ("%s.selectedIndexes()") % (table_name)
                            rowSelected = eval(rowSelected_cmd)
                            rowIndex = (rowSelected[0].row())
                            sito_item = self.tableWidget_rapporti.item(rowIndex, 3)
                            area_item = self.tableWidget_rapporti.item(rowIndex, 2)
                            sito_ = str(sito_item.text())
                            area_ = str(area_item.text())
                            us_item = self.tableWidget_rapporti.item(rowIndex, 1)
                            us = str(us_item.text())
                            search_dict = {'sito': "'" +sito_ + "'",
                                           'area': "'" + area_ + "'",
                                           'us': us}
                            u = Utility()
                            search_dict = u.remove_empty_items_fr_dict(search_dict)
                            res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
                            self.empty_fields()
                            self.DATA_LIST = []
                            for i in res:
                                self.DATA_LIST.append(i)
                            self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                            self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                            self.fill_fields()
                            self.BROWSE_STATUS = "b"
                            self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                            self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)

                    else:
                        pass
            else:
                if self.data_error_check() == 0:
                    test_insert = self.insert_new_rec()
                    if test_insert == 1:
                        self.empty_fields()
                        self.label_sort.setText(self.SORTED_ITEMS["n"])
                        self.charge_list()
                        self.set_sito()
                        self.charge_records()
                        self.BROWSE_STATUS = "b"
                        self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                        self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), len(self.DATA_LIST) - 1
                        self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
                        self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
                        self.setComboBoxEditable(["self.comboBox_sito"], 1)
                        self.fill_fields(self.REC_CORR)
                        self.enable_button(1)

                else:
                    pass



            table_name = "self.tableWidget_rapporti"
            rowSelected_cmd = ("%s.selectedIndexes()") % (table_name)
            rowSelected = eval(rowSelected_cmd)
            rowIndex = (rowSelected[0].row())
            sito_item = self.tableWidget_rapporti.item(rowIndex, 3)
            area_item = self.tableWidget_rapporti.item(rowIndex, 2)
            sito_ = str(sito_item.text())
            area_ = str(area_item.text())
            us_item = self.tableWidget_rapporti.item(rowIndex, 1)
            us = str(us_item.text())
            search_dict = {'sito': "'" + sito_ + "'",
                           'area': "'" + area_ + "'",
                           'us': us}
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
            if not bool(res):

                #self.DB_MANAGER.insert_number_of_us_records(sito,area,us,'US')

                if self.L=='it':
                    QMessageBox.warning(self, "ATTENZIONE", "Non e' stato trovato alcun record!", QMessageBox.Ok)

                elif self.L=='de':
                    QMessageBox.warning(self, "ACHTUNG", "kein Eintrag gefunden!", QMessageBox.Ok)
                else:
                    QMessageBox.warning(self, "Warning", "The record has not been found ", QMessageBox.Ok)
                self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                self.fill_fields(self.REC_CORR)
                self.BROWSE_STATUS = "b"
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.setComboBoxEnable(["self.comboBox_sito"], "False")
                self.setComboBoxEnable(["self.comboBox_area"], "False")
                self.setComboBoxEnable(["self.lineEdit_us"], "False")
            else:
                self.empty_fields()
                self.DATA_LIST = []
                for i in res:
                    self.DATA_LIST.append(i)
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                self.fill_fields()
                self.BROWSE_STATUS = "b"
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                if self.REC_TOT == 1:
                    if self.L=='it':
                        strings = ("E' stato trovato", self.REC_TOT, "record")
                    elif self.L=='de':
                        strings = ("Es wurde gefunden", self.REC_TOT, "record")
                    else:
                        strings = ("has been found", self.REC_TOT, "record")
                    if self.toolButtonGis.isChecked():
                        self.pyQGIS.charge_vector_layers(self.DATA_LIST)

                    if self.toolButton_usm.isChecked():
                        self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                else:
                    if self.L=='it':
                        strings = ("Sono stati trovati", self.REC_TOT, "records")
                    elif self.L=='de':
                        strings = ("Sie wurden gefunden", self.REC_TOT, "records")
                    else:
                        strings = ("Have been found", self.REC_TOT, "records")
                    if self.toolButtonGis.isChecked():
                        self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                    if self.toolButton_usm.isChecked():
                        self.pyQGIS.charge_usm_layers(self.DATA_LIST)

                self.setComboBoxEnable(["self.comboBox_sito"], "False")
                self.setComboBoxEnable(["self.comboBox_area"], "False")
                self.setComboBoxEnable(["self.lineEdit_us"], "False")
        except:
            pass


    def on_pushButton_go_to_scheda_pressed(self):
        try:
            #table_name = "self.table"
            #rowSelected_cmd = ("%s.selectedIndexes()") % (table_name)
            rowSelected = self.table.currentIndex()#eval(rowSelected_cmd)
            rowIndex = rowSelected.row()
            sito_item = self.tableWidget_rapporti.item(rowIndex, 3)
            area_item = self.tableWidget_rapporti.item(rowIndex, 2)
            sito_ = str(sito_item.text())
            area_ = str(area_item.text())
            us_item = self.tableWidget_rapporti.item(rowIndex, 1)
            us = str(us_item.text())
            search_dict = {'sito': "'" + sito_ + "'",
                           'area': "'" + area_ + "'",
                           'us': us}
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
            self.empty_fields()
            self.DATA_LIST = []
            for i in res:
                self.DATA_LIST.append(i)
            self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
            self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
            self.fill_fields()
            self.BROWSE_STATUS = "b"
            self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
            self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
        except Exception as e:
            e = str(e)
            if self.L=='it':
                QMessageBox.warning(self, "Alert", "Non hai selezionato nessuna riga. Errore python: %s " % (str(e)),
                                QMessageBox.Ok)
            elif self.L=='de':
                QMessageBox.warning(self, "ACHTUNG", "Keine Spalte ausgewält. Error python: %s " % (str(e)),
                                QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "Alert", "You didn't select any row. Python error: %s " % (str(e)),
                                QMessageBox.Ok)
    def enable_button(self, n):
        self.pushButton_list.setEnabled(n)
        self.pushButton_new_rec.setEnabled(n)
        self.pushButton_view_all.setEnabled(n)
        self.pushButton_first_rec.setEnabled(n)
        self.pushButton_last_rec.setEnabled(n)
        self.pushButton_prev_rec.setEnabled(n)
        self.pushButton_next_rec.setEnabled(n)
        self.pushButton_delete.setEnabled(n)
        self.pushButton_new_search.setEnabled(n)
        self.pushButton_search_go.setEnabled(n)
        self.pushButton_sort.setEnabled(n)
    def enable_button_search(self, n):
        # self.pushButton_connect.setEnabled(n)
        self.pushButton_new_rec.setEnabled(n)
        self.pushButton_view_all.setEnabled(n)
        self.pushButton_first_rec.setEnabled(n)
        self.pushButton_last_rec.setEnabled(n)
        self.pushButton_prev_rec.setEnabled(n)
        self.pushButton_next_rec.setEnabled(n)
        self.pushButton_delete.setEnabled(n)
        self.pushButton_save.setEnabled(n)
        self.pushButton_sort.setEnabled(n)
        self.pushButton_sort.setEnabled(n)
        self.pushButton_insert_row_rapporti.setEnabled(n)
        self.pushButton_remove_row_rapporti.setEnabled(n)
        self.pushButton_insert_row_inclusi.setEnabled(n)
        self.pushButton_remove_row_inclusi.setEnabled(n)
        self.pushButton_insert_row_campioni.setEnabled(n)
        self.pushButton_remove_row_campioni.setEnabled(n)
        self.pushButton_insert_row_organici.setEnabled(n)
        self.pushButton_remove_row_organici.setEnabled(n)
        self.pushButton_insert_row_inorganici.setEnabled(n)
        self.pushButton_remove_row_inorganici.setEnabled(n)
        self.pushButton_insert_row_documentazione.setEnabled(n)
        self.pushButton_remove_row_documentazione.setEnabled(n)

    def on_pushButton_connect_pressed(self):

        conn = Connection()
        conn_str = conn.conn_str()
        test_conn = conn_str.find('sqlite')
        if test_conn == 0:
            self.DB_SERVER = "sqlite"

        try:
            self.DB_MANAGER = Pyarchinit_db_management(conn_str)
            self.DB_MANAGER.connection()
            self.charge_records()  # charge records from DB
            # check if DB is empty
            if self.DATA_LIST:
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                self.BROWSE_STATUS = 'b'
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.label_sort.setText(self.SORTED_ITEMS["n"])
                self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                self.charge_list()
                self.fill_fields()
                self.setComboBoxEnable(["self.comboBox_area"], "False")
                self.setComboBoxEnable(["self.lineEdit_us"], "False")
                self.iconListWidget.update()
            else:
                if self.L=='it':
                    QMessageBox.warning(self,"BENVENUTO", "Benvenuto in pyArchInit " + self.NOME_SCHEDA + ". Il database e' vuoto. Premi 'Ok' e buon lavoro!",
                                        QMessageBox.Ok)
                elif self.L=='de':
                    QMessageBox.warning(self,"WILLKOMMEN","WILLKOMMEN in pyArchInit" + "SE-MSE formular"+ ". Die Datenbank ist leer. Tippe 'Ok' und aufgehts!",
                                        QMessageBox.Ok)
                else:
                    QMessageBox.warning(self,"WELCOME", "Welcome in pyArchInit" + "Samples SU-WSU" + ". The DB is empty. Push 'Ok' and Good Work!",
                                        QMessageBox.Ok)
                self.charge_list()
                self.BROWSE_STATUS = 'x'
                self.setComboBoxEnable(["self.comboBox_area"], "True")
                self.setComboBoxEnable(["self.lineEdit_us"], "True")
                self.on_pushButton_new_rec_pressed()
                self.iconListWidget.update()
        except Exception as e:
            e = str(e)
            if e.find("no such table"):
                if self.L=='it':
                    msg = "La connessione e' fallita {}. " \
                          "E' NECESSARIO RIAVVIARE QGIS oppure rilevato bug! Segnalarlo allo sviluppatore".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)

                elif self.L=='de':
                    msg = "Verbindungsfehler {}. " \
                          " QGIS neustarten oder es wurde ein bug gefunden! Fehler einsenden".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                else:
                    msg = "The connection failed {}. " \
                          "You MUST RESTART QGIS or bug detected! Report it to the developer".format(str(e))
            else:
                if self.L=='it':
                    msg = "Attenzione rilevato bug! Segnalarlo allo sviluppatore. Errore: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                elif self.L=='de':
                    msg = "ACHTUNG. Es wurde ein bug gefunden! Fehler einsenden: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                else:
                    msg = "Warning bug detected! Report it to the developer. Error: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)

    def connect_p(self):

        conn = Connection()
        conn_str = conn.conn_str()
        test_conn = conn_str.find('sqlite')
        if test_conn == 0:
            self.DB_SERVER = "sqlite"

        try:
            self.DB_MANAGER = Pyarchinit_db_management(conn_str)
            self.DB_MANAGER.connection()
            self.charge_records()  # charge records from DB
            # check if DB is empty
            if self.DATA_LIST:
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                self.BROWSE_STATUS = 'b'
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.label_sort.setText(self.SORTED_ITEMS["n"])
                self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                self.charge_list()
                self.fill_fields()
                self.setComboBoxEnable(["self.comboBox_area"], "False")
                self.setComboBoxEnable(["self.lineEdit_us"], "False")
                self.iconListWidget.update()
            else:
                if self.L=='it':
                    QMessageBox.warning(self,"BENVENUTO", "Benvenuto in pyArchInit " + self.NOME_SCHEDA + ". Il database e' vuoto. Premi 'Ok' e buon lavoro!",
                                        QMessageBox.Ok)
                elif self.L=='de':
                    QMessageBox.warning(self,"WILLKOMMEN","WILLKOMMEN in pyArchInit" + "SE-MSE formular"+ ". Die Datenbank ist leer. Tippe 'Ok' und aufgehts!",
                                        QMessageBox.Ok)
                else:
                    QMessageBox.warning(self,"WELCOME", "Welcome in pyArchInit" + "Samples SU-WSU" + ". The DB is empty. Push 'Ok' and Good Work!",
                                        QMessageBox.Ok)
                self.charge_list()
                self.BROWSE_STATUS = 'x'
                self.setComboBoxEnable(["self.comboBox_area"], "True")
                self.setComboBoxEnable(["self.lineEdit_us"], "True")
                self.on_pushButton_new_rec_pressed()
                self.iconListWidget.update()
        except Exception as e:
            e = str(e)
            if e.find("no such table"):
                if self.L=='it':
                    msg = "La connessione e' fallita {}. " \
                          "E' NECESSARIO RIAVVIARE QGIS oppure rilevato bug! Segnalarlo allo sviluppatore".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)

                elif self.L=='de':
                    msg = "Verbindungsfehler {}. " \
                          " QGIS neustarten oder es wurde ein bug gefunden! Fehler einsenden".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                else:
                    msg = "The connection failed {}. " \
                          "You MUST RESTART QGIS or bug detected! Report it to the developer".format(str(e))
            else:
                if self.L=='it':
                    msg = "Attenzione rilevato bug! Segnalarlo allo sviluppatore. Errore: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                elif self.L=='de':
                    msg = "ACHTUNG. Es wurde ein bug gefunden! Fehler einsenden: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
                else:
                    msg = "Warning bug detected! Report it to the developer. Error: ".format(str(e))
                    self.iface.messageBar().pushMessage(self.tr(msg), Qgis.Warning, 0)
    def customize_GUI(self):
        self.iconListWidget.update()
        l = QgsSettings().value("locale/userLocale", QVariant)[0:2]
        lang = ""
        for key, values in self.LANG.items():
            if values.__contains__(l):
                lang = str(key)
        lang = "'" + lang + "'"
        if not Pyarchinit_OS_Utility.checkgraphvizinstallation():
            self.pushButton_export_matrix.setEnabled(False)
            self.pushButton_export_matrix.setToolTip("Funzione disabilitata")
        self.tableWidget_rapporti.setColumnWidth(0, 120)
        self.tableWidget_rapporti.setColumnWidth(1, 80)
        self.tableWidget_rapporti.setColumnWidth(2, 30)
        self.tableWidget_rapporti.setColumnWidth(3, 50)
        self.tableWidget_rapporti2.setColumnWidth(0, 80)
        self.tableWidget_rapporti2.setColumnWidth(1, 50)
        self.tableWidget_rapporti2.setColumnWidth(2, 50)
        self.tableWidget_rapporti2.setColumnWidth(3, 200)
        self.tableWidget_rapporti2.setColumnWidth(4, 100)
        self.tableWidget_rapporti2.setColumnWidth(5, 30)
        self.tableWidget_rapporti2.setColumnWidth(6, 50)
        # self.tableWidget_rapporti2.sortItems(0,QtCore.Qt.AscendingOrder)
        # self.tableWidget_rapporti.sortItems(0,QtCore.Qt.AscendingOrder)
        self.tableWidget_documentazione.setColumnWidth(0, 150)
        self.tableWidget_documentazione.setColumnWidth(1, 300)
        self.tableWidget_rapporti.setItemDelegateForColumn(1, IntegerDelegate(self.tableWidget_rapporti))
        self.mapPreview = QgsMapCanvas(self)
        self.mapPreview.setCanvasColor(QColor(225, 225, 225))
        self.tabWidget.addTab(self.mapPreview, "Map preview")
        # media prevew system
        self.iconListWidget.setDragEnabled(True)
        self.iconListWidget.setAcceptDrops(True)
        self.iconListWidget.setDropIndicatorShown(True)

        self.iconListWidget.setLineWidth(2)
        self.iconListWidget.setMidLineWidth(2)
        #self.iconListWidget.setProperty("showDropIndicator", False)
        self.iconListWidget.setIconSize(QSize(430, 570))

        self.iconListWidget.setUniformItemSizes(True)
        self.iconListWidget.setObjectName("iconListWidget")
        self.iconListWidget.SelectionMode()
        self.iconListWidget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.iconListWidget.itemDoubleClicked.connect(self.openWide_image)
        #self.listWidget_2.itemDoubleClicked.connect(self.opentepmplatePreview)
        # comboBox customizations

        self.setComboBoxEditable(["self.comboBox_per_fin"], 1)
        self.setComboBoxEditable(["self.comboBox_fas_fin"], 1)
        self.setComboBoxEditable(["self.comboBox_per_iniz"], 1)
        self.setComboBoxEditable(["self.comboBox_fas_iniz"], 1)
        self.setComboBoxEditable(["self.comboBox_struttura"], 1)
        self.setComboBoxEditable(["self.comboBox_ref_ra"], 1)
        #self.setComboBoxEditable(["self.comboBox_datazione"],1)
        # lista tipo rapporti stratigrafici
        if self.L=='it':
            valuesRS = ["Uguale a", "Si lega a", "Copre", "Coperto da", "Riempie", "Riempito da", "Taglia", "Tagliato da", "Si appoggia a", "Gli si appoggia", ">","<","<<",">>","<->",""]
            self.delegateRS = ComboBoxDelegate()
            self.delegateRS.def_values(valuesRS)
            self.delegateRS.def_editable('False')
            self.tableWidget_rapporti.setItemDelegateForColumn(0,self.delegateRS)

        elif self.L=='en':
            valuesRS = ["Same as", "Connected to", "Covers", "Covered by", "Fills", "Filled by", "Cuts", "Cut by", "Abuts", "Supports", ">","<","<<",">>","<->",""]
            self.delegateRS = ComboBoxDelegate()
            self.delegateRS.def_values(valuesRS)
            self.delegateRS.def_editable('False')
            self.tableWidget_rapporti.setItemDelegateForColumn(0,self.delegateRS)
        elif self.L=='de':
            valuesRS = ["Entspricht", "Bindet an", "Liegt über", "Liegt unter", "Verfüllt", "Wird verfüllt durch", "Schneidet", "Wird geschnitten", "Stützt sich auf", "Wird gestüzt von", ">","<","<<",">>","<->",""]
            self.delegateRS = ComboBoxDelegate()
            self.delegateRS.def_values(valuesRS)
            self.delegateRS.def_editable('False')
            self.tableWidget_rapporti.setItemDelegateForColumn(0,self.delegateRS)
        else:
            valuesRS = ["Same as", "Connected to", "Covers", "Covered by", "Fills", "Filled by", "Cuts", "Cut by", "Abuts", "Supports", ">","<","<<",">>","<->",""]
            self.delegateRS = ComboBoxDelegate()
            self.delegateRS.def_values(valuesRS)
            self.delegateRS.def_editable('False')
            self.tableWidget_rapporti.setItemDelegateForColumn(0,self.delegateRS)
        conn = Connection()
        sito_set = conn.sito_set()
        sito_set_str = sito_set['sito_set']
        value_site = [sito_set_str]
        self.delegatesito = ComboBoxDelegate()
        self.delegatesito.def_values(value_site)
        self.delegatesito.def_editable('False')
        self.tableWidget_rapporti.setItemDelegateForColumn(3, self.delegatesito)

        # lista tipo documentazione
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.19' + "'"
        }
        tipo_di_documentazione = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesDoc = []
        if self.L=='it':
            valuesDoc.append("ICCD-Piante")
            valuesDoc.append("ICCD-Piante-Sezioni")
            valuesDoc.append("ICCD-Sezioni")
            valuesDoc.append("ICCD-Prospetti")
            valuesDoc.append("ICCD-Foto")
        elif self.L=='de':
            valuesDoc.append("Pflanzen")
            valuesDoc.append("Sektionen")
            valuesDoc.append("Prospekte")
            valuesDoc.append("Foto")
        else:
            valuesDoc.append("Maps")
            valuesDoc.append("Sections")
            valuesDoc.append("Elevations")
            valuesDoc.append("Photo")

        for i in range(len(tipo_di_documentazione)):
            valuesDoc.append(tipo_di_documentazione[i].sigla_estesa)
        #valuesDoc.sort()
        self.delegateDoc = ComboBoxDelegate()
        self.delegateDoc.def_values(valuesDoc)
        self.delegateDoc.def_editable('False')
        self.tableWidget_documentazione.setItemDelegateForColumn(0, self.delegateDoc)




        # lista colore legante usm
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '201.201' + "'"
        }
        colore = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesCol = []
        for i in range(len(colore)):
            valuesCol.append(colore[i].sigla_estesa)
        valuesCol.sort()
        self.delegateCol = ComboBoxDelegate()
        self.delegateCol.def_values(valuesCol)
        self.delegateCol.def_editable('False')
        self.tableWidget_colore_legante_usm.setItemDelegateForColumn(0, self.delegateCol)
        # lista colore materiale usm
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '201.201' + "'"
        }
        colore = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesCol = []
        for i in range(len(colore)):
            valuesCol.append(colore[i].sigla_estesa)
        valuesCol.sort()
        self.delegateCol = ComboBoxDelegate()
        self.delegateCol.def_values(valuesCol)
        self.delegateCol.def_editable('False')
        self.tableWidget_colore_materiale_usm.setItemDelegateForColumn(0, self.delegateCol)
        # lista inclusi leganti usm
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '202.202' + "'"
        }
        # inclusi = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        # valuesInclusi = []
        # for i in range(len(inclusi)):
            # valuesInclusi.append(inclusi[i].sigla_estesa)
        # valuesCol.sort()
        # self.delegateInclusi = ComboBoxDelegate()
        # self.delegateInclusi.def_values(valuesInclusi)
        # self.delegateInclusi.def_editable('False')
        # self.tableWidget_inclusi_leganti_usm.setItemDelegateForColumn(0, self.delegateInclusi)
        # # lista inclusi materiali usm
        # search_dict = {
            # 'lingua': lang,
            # 'nome_tabella': "'" + 'us_table' + "'",
            # 'tipologia_sigla': "'" + '202.202' + "'"
        # }
        inclusi = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesInclusi = []
        for i in range(len(inclusi)):
            valuesInclusi.append(inclusi[i].sigla_estesa)
        valuesCol.sort()
        self.delegateInclusi = ComboBoxDelegate()
        self.delegateInclusi.def_values(valuesInclusi)
        self.delegateInclusi.def_editable('False')
        self.tableWidget_inclusi_materiali_usm.setItemDelegateForColumn(0, self.delegateInclusi)
        # lista consistenza/texture materiale usm
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.7' + "'"
        }
        constex = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesCTX = []
        for i in range(len(constex)):
            valuesCTX.append(constex[i].sigla_estesa)
        valuesCol.sort()
        self.delegateCons = ComboBoxDelegate()
        self.delegateCons.def_values(valuesCTX)
        self.delegateCons.def_editable('False')
        self.tableWidget_consistenza_texture_mat_usm.setItemDelegateForColumn(0, self.delegateCons)
        # lista componenti organici
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.14' + "'"
        }
        comporg = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesCOG = []
        for i in range(len(comporg)):
            valuesCOG.append(comporg[i].sigla_estesa)
        valuesCOG.sort()
        self.delegateCOG = ComboBoxDelegate()
        self.delegateCOG.def_values(valuesCOG)
        self.delegateCOG.def_editable('False')
        self.tableWidget_organici.setItemDelegateForColumn(0, self.delegateCOG)
        # lista componenti inorganici
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.15' + "'"
        }
        compinorg = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesCINOG = []
        for i in range(len(compinorg)):
            valuesCINOG.append(compinorg[i].sigla_estesa)
        valuesCINOG.sort()
        self.delegateCINOG = ComboBoxDelegate()
        self.delegateCINOG.def_values(valuesCINOG)
        self.delegateCINOG.def_editable('False')
        self.tableWidget_inorganici.setItemDelegateForColumn(0, self.delegateCINOG)
        #lista campioni
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.13' + "'"
        }
        tipo_inclusi_campioni = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesINCL_CAMP = []
        for i in range(len(tipo_inclusi_campioni)):
            valuesINCL_CAMP.append(tipo_inclusi_campioni[i].sigla_estesa)
        valuesINCL_CAMP.sort()

        self.delegateINCL_CAMP = ComboBoxDelegate()
        valuesINCL_CAMP.sort()
        self.delegateINCL_CAMP.def_values(valuesINCL_CAMP)
        self.delegateINCL_CAMP.def_editable('False')
        self.tableWidget_campioni.setItemDelegateForColumn(0, self.delegateINCL_CAMP)
        # lista inclusi
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '202.202' + "'"
        }
        tipo_inclusi = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        valuesINCL = []
        for i in range(len(tipo_inclusi)):
            valuesINCL.append(tipo_inclusi[i].sigla_estesa)
        valuesINCL.sort()
        self.delegateINCL = ComboBoxDelegate()
        self.delegateINCL.def_values(valuesINCL)
        self.delegateINCL.def_editable('False')
        self.tableWidget_inclusi.setItemDelegateForColumn(0, self.delegateINCL)

    def loadMapPreview(self, mode=0):
        if mode == 0:
            """ if has geometry column load to map canvas """
            gidstr = self.ID_TABLE + " = " + str(
                eval("self.DATA_LIST[int(self.REC_CORR)]." + self.ID_TABLE))
            layerToSet = self.pyQGIS.loadMapPreview_new(gidstr)
            #QMessageBox.warning(self, "layer to set", '\n'.join([l.name() for l in layerToSet]), QMessageBox.Ok)
            self.mapPreview.setLayers(layerToSet)
            self.mapPreview.zoomToFullExtent()
        elif mode == 1:
            self.mapPreview.setLayers([])
            self.mapPreview.zoomToFullExtent()

    def dropEvent(self, event):
        mimeData = event.mimeData()
        accepted_formats = ["jpg", "jpeg", "png", "tiff", "tif", "bmp", "mp4", "avi", "mov", "mkv", "flv", "obj", "stl",
                            "ply", "fbx", "3ds"]

        if mimeData.hasUrls():
            for url in mimeData.urls():
                try:
                    path = url.toLocalFile()
                    if os.path.isfile(path):
                        filename = os.path.basename(path)
                        filetype = filename.split(".")[-1]
                        if filetype.lower() in accepted_formats:
                            self.load_and_process_image(path)
                        else:
                            QMessageBox.warning(self, "Error", f"Unsupported file type: {filetype}", QMessageBox.Ok)
                except Exception as e:
                    QMessageBox.warning(self, "Error", f"Failed to process the file: {str(e)}", QMessageBox.Ok)
        super().dropEvent(event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        event.acceptProposedAction()
    def insert_record_media(self, mediatype, filename, filetype, filepath):
        self.mediatype = mediatype
        self.filename = filename
        self.filetype = filetype
        self.filepath = filepath
        try:
            data = self.DB_MANAGER.insert_media_values(
                self.DB_MANAGER.max_num_id('MEDIA', 'id_media') + 1,
                str(self.mediatype),  # 1 - mediatyype
                str(self.filename),  # 2 - filename
                str(self.filetype),  # 3 - filetype
                str(self.filepath),  # 4 - filepath
                str('Insert description'),  # 5 - descrizione
                str("['imagine']"))  # 6 - tags
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as  e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.filename + ": Image already in the database"
                else:
                    msg = e
                #QMessageBox.warning(self, "Errore", "Warning 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0
    def insert_record_mediathumb(self, media_max_num_id, mediatype, filename, filename_thumb, filetype, filepath_thumb, filepath_resize):
        self.media_max_num_id = media_max_num_id
        self.mediatype = mediatype
        self.filename = filename
        self.filename_thumb = filename_thumb
        self.filetype = filetype
        self.filepath_thumb = filepath_thumb
        self.filepath_resize = filepath_resize
        try:
            data = self.DB_MANAGER.insert_mediathumb_values(
                self.DB_MANAGER.max_num_id('MEDIA_THUMB', 'id_media_thumb') + 1,
                str(self.media_max_num_id),  # 1 - media_max_num_id
                str(self.mediatype),  # 2 - mediatype
                str(self.filename),  # 3 - filename
                str(self.filename_thumb),  # 4 - filename_thumb
                str(self.filetype),  # 5 - filetype
                str(self.filepath_thumb),  # 6 - filepath_thumb
                str(self.filepath_resize))  # 6 - filepath_thumb
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.filename + ": thumb already present into the database"
                else:
                    msg = e
                #QMessageBox.warning(self, "Error", "warming 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0
    def insert_mediaToEntity_rec(self, id_entity, entity_type, table_name, id_media, filepath, media_name):
        """
        id_mediaToEntity,
        id_entity,
        entity_type,
        table_name,
        id_media,
        filepath,
        media_name"""
        self.id_entity = id_entity
        self.entity_type = entity_type
        self.table_name = table_name
        self.id_media = id_media
        self.filepath = filepath
        self.media_name = media_name
        try:
            data = self.DB_MANAGER.insert_media2entity_values(
                self.DB_MANAGER.max_num_id('MEDIATOENTITY', 'id_mediaToEntity') + 1,
                int(self.id_entity),  # 1 - id_entity
                str(self.entity_type),  # 2 - entity_type
                str(self.table_name),  # 3 - table_name
                int(self.id_media),  # 4 - us
                str(self.filepath),  # 5 - filepath
                str(self.media_name))  # 6 - media_name
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as  e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.ID_TABLE + " already present into the database"
                else:
                    msg = e
                QMessageBox.warning(self, "Error", "Warning 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0
    def delete_mediaToEntity_rec(self, id_entity, entity_type, table_name, id_media, filepath, media_name):
        """
        id_mediaToEntity,
        id_entity,
        entity_type,
        table_name,
        id_media,
        filepath,
        media_name"""
        self.id_entity = id_entity
        self.entity_type = entity_type
        self.table_name = table_name
        self.id_media = id_media
        self.filepath = filepath
        self.media_name = media_name
        try:
            data = self.DB_MANAGER.insert_media2entity_values(
            self.DB_MANAGER.max_num_id('MEDIATOENTITY', 'id_mediaToEntity')+1,
            int(self.id_entity),                                                    #1 - id_entity
            str(self.entity_type),                                              #2 - entity_type
            str(self.table_name),                                               #3 - table_name
            int(self.id_media),                                                     #4 - us
            str(self.filepath),                                                     #5 - filepath
            str(self.media_name))
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0

    def generate_US(self):
        #tags_list = self.table2dict('self.tableWidgetTags_US')
        record_us_list = []
        sito=self.comboBox_sito.currentText()
        area=self.comboBox_area.currentText()
        us=self.lineEdit_us.text()
        #for sing_tags in tags_list:
        search_dict = {'sito': "'"+str(sito)+"'" ,
                       'area': "'"+str(area)+"'",
                       'us': "'"+str(us)+"'"
                       }
        j = self.DB_MANAGER.query_bool(search_dict, 'US')
        record_us_list.append(j)
        #QMessageBox.information(self, 'search db', str(record_us_list))
        us_list = []
        for r in record_us_list:
            us_list.append([r[0].id_us, 'US', 'us_table'])
        #QMessageBox.information(self, "Scheda US", str(us_list), QMessageBox.Ok)
        return us_list
    def assignTags_US(self, item):
        """
        id_mediaToEntity,
        id_entity,
        entity_type,
        table_name,
        id_media,
        filepath,
        media_name
        """
        us_list = self.generate_US()
        #QMessageBox.information(self,'search db',str(us_list))
        if not us_list:
            return

        for us_data in us_list:
            id_orig_item = item.text()  # return the name of original file
            search_dict = {'filename': "'" + str(id_orig_item) + "'"}
            media_data = self.DB_MANAGER.query_bool(search_dict, 'MEDIA')
            self.insert_mediaToEntity_rec(us_data[0], us_data[1], us_data[2], media_data[0].id_media,
                                          media_data[0].filepath, media_data[0].filename)

    def load_and_process_image(self, filepath):
        media_resize_suffix = ''
        media_thumb_suffix = ''
        conn = Connection()
        thumb_path = conn.thumb_path()
        thumb_path_str = thumb_path['thumb_path']
        if thumb_path_str == '':
            if self.L == 'it':
                QMessageBox.information(self, "Info",
                                        "devi settare prima la path per salvare le thumbnail e i video. Vai in impostazioni di sistema/ path setting ")
            elif self.L == 'de':
                QMessageBox.information(self, "Info",
                                        "müssen Sie zuerst den Pfad zum Speichern der Miniaturansichten und Videos festlegen. Gehen Sie zu System-/Pfad-Einstellung")
            else:
                QMessageBox.information(self, "Message",
                                        "you must first set the path to save the thumbnails and videos. Go to system/path setting")
        else:
            filename = os.path.basename(filepath)
            filename, filetype = filename.split(".")[0], filename.split(".")[1]
            # Check the media type based on the file extension
            accepted_image_formats = ["jpg", "jpeg", "png", "tiff", "tif", "bmp"]
            accepted_video_formats = ["mp4", "avi", "mov", "mkv", "flv"]
            accepted_3d_formats = ["obj", "stl", "ply", "fbx", "3ds"]

            if filetype.lower() in accepted_image_formats:
                mediatype = 'image'
                media_thumb_suffix = '_thumb.png'
                media_resize_suffix = '.png'
            elif filetype.lower() in accepted_video_formats:
                mediatype = 'video'
                media_thumb_suffix = '_video.png'
                media_resize_suffix = '.' + filetype.lower()
            elif filetype.lower() in accepted_3d_formats:
                mediatype = '3d_model'
                media_thumb_suffix = '_3d_thumb.png'
                media_resize_suffix = '.' + filetype.lower()
            else:
                raise ValueError(f"Unrecognized media type for file {filename}.{filetype}")

            if mediatype == 'video':
                if filetype.lower() == 'mp4':
                    media_resize_suffix = '.mp4'
                elif filetype.lower() == 'avi':
                    media_resize_suffix = '.avi'
                elif filetype.lower() == 'mov':
                    media_resize_suffix = '.mov'
                elif filetype.lower() == 'mkv':
                    media_resize_suffix = '.mkv'
                elif filetype.lower() == 'flv':
                    media_resize_suffix = '.flv'

            elif mediatype == '3d_model':
                if filetype.lower() == 'obj':
                    media_resize_suffix = '.obj'
                elif filetype.lower() == 'ply':
                    media_resize_suffix = '.ply'
                elif filetype.lower() == 'fbx':
                    media_resize_suffix = '.fbx'
                elif filetype.lower() == '3ds':
                    media_resize_suffix = '.3ds'
                elif filetype.lower() == 'stl':
                    media_resize_suffix = '.stl'
            # Check and insert record in the database
            idunique_image_check = self.db_search_check('MEDIA', 'filepath', filepath)

            try:
                if bool(idunique_image_check):

                    return
                else:
                    # mediatype = 'image'
                    self.insert_record_media(mediatype, filename, filetype, filepath)
                    MU = Media_utility()
                    MUR = Media_utility_resize()
                    MU_video = Video_utility()
                    MUR_video = Video_utility_resize()
                    media_max_num_id = self.DB_MANAGER.max_num_id('MEDIA', 'id_media')
                    thumb_path = conn.thumb_path()
                    thumb_path_str = thumb_path['thumb_path']
                    thumb_resize = conn.thumb_resize()
                    thumb_resize_str = thumb_resize['thumb_resize']
                    filenameorig = filename
                    filename_thumb = str(media_max_num_id) + "_" + filename + media_thumb_suffix
                    filename_resize = str(media_max_num_id) + "_" + filename + media_resize_suffix
                    filepath_thumb = filename_thumb
                    filepath_resize = filename_resize
                    self.SORT_ITEMS_CONVERTED = []

                    try:

                        if mediatype == '3d_model':
                            self.process_3d_model(media_max_num_id, filepath, filename, thumb_path_str,
                                                  thumb_resize_str,
                                                  media_thumb_suffix, media_resize_suffix)

                        elif mediatype == 'video':
                            vcap = cv2.VideoCapture(filepath)
                            res, im_ar = vcap.read()
                            while im_ar.mean() < 1 and res:
                                res, im_ar = vcap.read()
                            im_ar = cv2.resize(im_ar, (100, 100), 0, 0, cv2.INTER_LINEAR)
                            # to save we have two options
                            outputfile = '{}.png'.format(os.path.dirname(filepath) + '/' + filename)
                            cv2.imwrite(outputfile, im_ar)
                            MU_video.resample_images(media_max_num_id, outputfile, filenameorig, thumb_path_str,
                                                     media_thumb_suffix)
                            MUR_video.resample_images(media_max_num_id, filepath, filenameorig, thumb_resize_str,
                                                      media_resize_suffix)
                        else:
                            MU.resample_images(media_max_num_id, filepath, filenameorig, thumb_path_str,
                                               media_thumb_suffix)
                            MUR.resample_images(media_max_num_id, filepath, filenameorig, thumb_resize_str,
                                                media_resize_suffix)
                    except Exception as e:
                        QMessageBox.warning(self, "Cucu", str(e), QMessageBox.Ok)
                    self.insert_record_mediathumb(media_max_num_id, mediatype, filename, filename_thumb, filetype,
                                                  filepath_thumb, filepath_resize)

                    item = QListWidgetItem(str(filenameorig))
                    item.setData(Qt.UserRole, str(media_max_num_id))
                    icon = QIcon(str(thumb_path_str) + filepath_thumb)
                    item.setIcon(icon)
                    self.iconListWidget.addItem(item)

                self.assignTags_US(item)




            except AssertionError as e:

                if self.L == 'it':
                    QMessageBox.warning(self, "Warning", "controlla che il nome del file non abbia caratteri speciali",

                                        QMessageBox.Ok)

                if self.L == 'de':

                    QMessageBox.warning(self, "Warning", "prüfen, ob der Dateiname keine Sonderzeichen enthält",
                                        QMessageBox.Ok)

                else:

                    QMessageBox.warning(self, "Warning", str(e), QMessageBox.Ok)


    def db_search_check(self, table_class, field, value):
        self.table_class = table_class
        self.field = field
        self.value = value
        search_dict = {self.field: "'" + str(self.value) + "'"}
        u = Utility()
        search_dict = u.remove_empty_items_fr_dict(search_dict)
        res = self.DB_MANAGER.query_bool(search_dict, self.table_class)
        return res
    def on_pushButton_assigntags_pressed(self):

        # Check the locale and set the button text and message box content
        L = QgsSettings().value("locale/userLocale")[0:2]
        if L == 'it':
            done_button_text = "Fatto"
            warning_title = "Attenzione"
            warning_text = "Devi selezionare una o più US"
        elif L == 'de':
            done_button_text = "Fertig"
            warning_title = "Achtung"
            warning_text = "Sie müssen eine oder mehrere US auswählen"
        else:  # Default to English
            done_button_text = "Done"
            warning_title = "Attention"
            warning_text = "You must select one or more US"
        # Check if there are selected items in the iconListWidget
        if not self.iconListWidget.selectedItems():
            QMessageBox.warning(self, warning_title, warning_text)
            return  # Exit the function if there are no selected images

        # Query all US records from the database and sort them
        all_us = self.DB_MANAGER.query('US')
        sorted_us = sorted(all_us, key=lambda x: (x.sito, x.area, x.us))

        # Create a QListWidget and populate it with sorted US records
        self.us_listwidget = QListWidget()
        header_item = QListWidgetItem("Sito - Area - US")
        header_item.setBackground(QColor('lightgrey'))
        header_item.setFlags(header_item.flags() & ~Qt.ItemIsSelectable)
        self.us_listwidget.addItem(header_item)
        for us in sorted_us:
            item_string = f"{us.sito} - {us.area} - {us.us}"
            self.us_listwidget.addItem(QListWidgetItem(item_string))

        # Set selection mode to allow multiple selections
        self.us_listwidget.setSelectionMode(QAbstractItemView.MultiSelection)

        # Create a "Done" button and connect it to the slot
        done_button = QPushButton(done_button_text)
        done_button.clicked.connect(self.on_done_selecting)

        # Create a layout and add the QListWidget and "Done" button
        layout = QVBoxLayout()
        layout.addWidget(self.us_listwidget)
        layout.addWidget(done_button)

        # Create a widget to contain the QListWidget and button, and set the layout
        self.widget = QWidget()
        self.widget.setLayout(layout)
        self.widget.show()

    def on_done_selecting(self):
        # Check the locale and set the button text and message box content
        L = QgsSettings().value("locale/userLocale")[0:2]
        if L == 'it':
            done_button_text = "Fatto"
            warning_title = "Attenzione"
            warning_text = "Devi selezionare una o più US"
        elif L == 'de':
            done_button_text = "Fertig"
            warning_title = "Achtung"
            warning_text = "Sie müssen eine oder mehrere US auswählen"
        else:  # Default to English
            done_button_text = "Done"
            warning_title = "Attention"
            warning_text = "You must select one or more US"

        # Handle the event when the "Done" button is clicked
        selected_items = self.us_listwidget.selectedItems()
        if not selected_items:
            # Show a warning message if no items are selected
            QMessageBox.warning(self, warning_title, warning_text)
        else:
            # Process the selected items
            pass  # Replace with the code to handle the selected US records

        def r_list():

            # Ottieni le US selezionate dall'utente
            selected_us = [item.text().split(' - ') for item in self.us_listwidget.selectedItems()]
            record_us_list=[]
            for sing_tags in selected_us:
                search_dict = {'sito': "'" + str(sing_tags[0]) + "'",
                               'area': "'" + str(sing_tags[1]) + "'",
                               'us': "'" + str(sing_tags[2]) + "'"
                               }
                j = self.DB_MANAGER.query_bool(search_dict, 'US')
                record_us_list.append(j)
            us_list = []
            for r in record_us_list:
                us_list.append([r[0].id_us, 'US', 'us_table'])
            # QMessageBox.information(self, "Scheda US", str(us_list), QMessageBox.Ok)
            return us_list


        #QMessageBox.information(self, 'ok', str(r_list()))
        items_selected=self.iconListWidget.selectedItems()
        for item in items_selected:
            for us_data in r_list():



                id_orig_item = item.text()  # return the name of original file
                search_dict = {'filename': "'" + str(id_orig_item) + "'"}
                media_data = self.DB_MANAGER.query_bool(search_dict, 'MEDIA')
                self.insert_mediaToEntity_rec(us_data[0], us_data[1], us_data[2], media_data[0].id_media,
                                              media_data[0].filepath, media_data[0].filename)

        self.widget.close()  # Chiude il widget dopo che l'utente ha premuto "Fatto"

    def on_pushButton_removetags_pressed(self):
        def r_id():
            sito = self.comboBox_sito.currentText()
            area = self.comboBox_area.currentText()
            us = self.lineEdit_us.text()
            record_us_list=[]
            search_dict = {'sito': "'" + str(sito) + "'",
                           'area': "'" + str(area) + "'",
                           'us': "'" + str(us) + "'"
                           }
            j = self.DB_MANAGER.query_bool(search_dict, 'US')
            record_us_list.append(j)

            a=None
            for r in record_us_list:
                a=r[0].id_us
            #QMessageBox.information(self,'ok',str(a))# QMessageBox.information(self, "Scheda US", str(us_list), QMessageBox.Ok)
            return a
        items_selected=self.iconListWidget.selectedItems()
        if not bool(items_selected):
            if self.L == 'it':

                msg = QMessageBox.warning(self, "Attenzione!!!",
                                          "devi selezionare prima l'immagine",
                                          QMessageBox.Ok)

            elif self.L == 'de':

                msg = QMessageBox.warning(self, "Warnung",
                                          "moet je eerst de afbeelding selecteren",
                                          QMessageBox.Ok)
            else:

                msg = QMessageBox.warning(self, "Warning",
                                          "you must first select an image",
                                          QMessageBox.Ok)
        else:
            if self.L == 'it':
                msg = QMessageBox.warning(self, "Warning",
                                          "Vuoi veramente cancellare i tags dalle thumbnail selezionate? \n L'azione è irreversibile",
                                          QMessageBox.Ok | QMessageBox.Cancel)
                if msg == QMessageBox.Cancel:
                    QMessageBox.warning(self, "Messaggio!!!", "Azione Annullata!")
                else:
                    #items_selected = self.iconListWidget.selectedItems()
                    for item in items_selected:
                        id_orig_item = item.text()  # return the name of original file

                        # s = self.iconListWidget.item(0, 0).text()
                        self.DB_MANAGER.remove_tags_from_db_sql_scheda(r_id(), id_orig_item)
                        row = self.iconListWidget.row(item)
                        self.iconListWidget.takeItem(row)
                    QMessageBox.warning(self, "Info", "Tags rimossi!")
            elif self.L == 'de':
                msg = QMessageBox.warning(self, "Warning",
                                          "Wollen Sie wirklich die Tags aus den ausgewählten Miniaturbildern löschen? \n Die Aktion ist unumkehrbar",
                                          QMessageBox.Ok | QMessageBox.Cancel)
                if msg == QMessageBox.Cancel:
                    QMessageBox.warning(self, "Warnung", "Azione Annullata!")
                else:
                    #items_selected = self.iconListWidget.selectedItems()
                    for item in items_selected:
                        id_orig_item = item.text()  # return the name of original file

                        # s = self.iconListWidget.item(0, 0).text()
                        self.DB_MANAGER.remove_tags_from_db_sql_scheda(r_id(), id_orig_item)
                        row = self.iconListWidget.row(item)
                        self.iconListWidget.takeItem(row)
                    QMessageBox.warning(self, "Info", "Tags entfernt")

            else:
                msg = QMessageBox.warning(self, "Warning",
                                          "Do you really want to delete the tags from the selected thumbnails? \n The action is irreversible",
                                          QMessageBox.Ok | QMessageBox.Cancel)
                if msg == QMessageBox.Cancel:
                    QMessageBox.warning(self, "Warning", "Action cancelled")
                else:
                    #items_selected = self.iconListWidget.selectedItems()
                    for item in items_selected:
                        id_orig_item = item.text()  # return the name of original file

                        #s = self.iconListWidget.item(0, 0).text()
                        self.DB_MANAGER.remove_tags_from_db_sql_scheda(r_id(),id_orig_item)
                        row = self.iconListWidget.row(item)
                        self.iconListWidget.takeItem(row)  # remove the item from the list

                    QMessageBox.warning(self, "Info", "Tags removed")

    def on_pushButton_all_images_pressed(self):
        record_us_list = self.DB_MANAGER.query('MEDIA_THUMB')

        et = {'entity_type': "'US'"}
        ser = self.DB_MANAGER.query_bool(et, 'MEDIATOENTITY')
        # Verifica se record_us_list è vuota
        if not record_us_list and not ser:
            QMessageBox.information(self, "Informazione", "Non ci sono immagini da mostrare.")
            return  # Termina la funzione

        # Inizializza la QListWidget fuori dal ciclo
        self.new_list_widget = QListWidget()
        # ##self.new_list_widget.setFixedSize(200, 300)
        self.new_list_widget.setSelectionMode(QAbstractItemView.SingleSelection)  # Permette selezioni multiple



        done_button = QPushButton("TAG")

        def update_done_button():
            if not self.new_list_widget.selectedItems():
                done_button.setHidden(True)
            else:
                done_button.setHidden(False)
                done_button.clicked.connect(self.on_done_selecting_all)

        self.new_list_widget.itemSelectionChanged.connect(update_done_button)# Aggiungi un layout per le etichette dei numeri delle pagine
        self.pageLayout = QHBoxLayout()
        self.current_page_label = QLabel()  # Creiamo l'etichetta per la pagina corrente
        self.total_pages_label = QLabel()  # Creiamo l'etichetta per il totale delle pagine

        self.pageLayout.addWidget(self.current_page_label)  # Aggiungiamo l'etichetta della pagina corrente al layout
        self.pageLayout.addWidget(self.total_pages_label)  # Aggiungiamo l'etichetta del totale delle pagine al layout

        # Aggiungi un pulsante "Indietro"
        self.prevButton = QPushButton("<<")
        self.prevButton.clicked.connect(self.go_to_previous_page)
        self.pageLayout.addWidget(self.prevButton)

        # Aggiungi le etichette dei numeri delle pagine
        self.pageLabels = []
        for i in range(1, 6):
            label = QLabel(str(i))
            label.setAlignment(Qt.AlignCenter)
            label.setMinimumWidth(30)
            label.setFrameStyle(QFrame.Panel | QFrame.Sunken)
            label.setMargin(2)
            label.mousePressEvent = functools.partial(self.on_page_label_clicked, i)
            self.pageLabels.append(label)
            self.pageLayout.addWidget(label)

        # Aggiungi un pulsante "Avanti"
        self.nextButton = QPushButton(">>")
        self.nextButton.clicked.connect(self.go_to_next_page)
        self.pageLayout.addWidget(self.nextButton)

        layout = QVBoxLayout()
        # Crea un campo di input per la ricerca
        self.search_field = QLineEdit()
        self.search_field.setPlaceholderText("Cerca...poi schiaccia invio")
        self.current_filter_text = ""

        self.page_size = 10  # Numero di immagini per pagina
        self.current_page = 1  # Pagina corrente
        self.total_pages = 0  # Numero totale di pagine

        # Aggiungi il campo di ricerca al layout sopra la QListWidget
        layout.insertWidget(0, self.search_field)

        layout.addLayout(self.pageLayout)
        layout.addWidget(self.new_list_widget)
        layout.addWidget(done_button)

        # Imposta il fattore di estensione per i widget nel layout
        # Il primo parametro è l'indice del widget e il secondo parametro è il fattore di estensione
        # In questo caso, new_list_widget ha un indice di 0 e done_button ha un indice di 1
        layout.setStretchFactor(self.new_list_widget, 5)  # new_list_widget avrà 3 volte più spazio di done_button
        layout.setStretchFactor(done_button, 1)  # done_button avrà 1/3 dello spazio di new_list_widget

        # Imposta il layout sulla tua finestra o su un altro widget
        self.setLayout(layout)

        # Crea un nuovo widget per contenere la QListWidget e il pulsante, e applica il layout
        self.widget = QWidget()
        self.widget.setLayout(layout)
        self.widget.adjustSize()
        self.widget.show()

        self.load_images()

        # Connette il campo di ricerca a una funzione di filtraggio
        self.search_field.returnPressed.connect(self.filter_items)


    def load_images(self, filter_text=None):
        conn = Connection()
        thumb_path = conn.thumb_path()
        thumb_path_str = thumb_path['thumb_path']
        u = Utility()

        # Calcola l'offset per la pagina corrente
        #offset = (self.current_page - 1) * self.page_size

        # Ottieni tutti i record delle immagini
        all_images = self.DB_MANAGER.query('MEDIA_THUMB')

        # Ottieni tutte le immagini taggate
        tagged_images = self.DB_MANAGER.query('MEDIATOENTITY')

        # Ottieni gli id_media di tutte le immagini taggate
        tagged_ids = [i.id_media for i in tagged_images]

        # Filtra tutte le immagini per ottenere solo quelle non taggate
        untagged_images = [i for i in all_images if i.id_media not in tagged_ids]

        # Inizializza l'elenco delle immagini 'US' come un duplicato delle immagini non taggate
        us_images = untagged_images[:]





        if len(all_images)>100:

            if filter_text:  # se il filtro è attivo
                filtered_images = [i for i in untagged_images if filter_text.lower() in i.media_filename.lower()]
            else:
                filtered_images = us_images
            # Calcola gli indici di inizio e fine per la pagina corrente
            start_index = (self.current_page - 1) * self.page_size
            end_index = start_index + self.page_size

            # Ottieni i record delle immagini per la pagina corrente
            self.record_us_list = filtered_images[start_index:end_index]
            # Pulisci la QListWidget prima di aggiungere le nuove immagini
            self.new_list_widget.clear()
            # Aggiungi l'intestazione alla QListWidget
            header_item = QListWidgetItem(
                "Le righe selezionate in giallo indicano immagini non taggate\n Da questo strumento solo le righe selezionate gialle posso essere taggate ")
            header_item.setBackground(QColor('lightgrey'))
            header_item.setFlags(header_item.flags() & ~Qt.ItemIsSelectable)  # rendi l'item non selezionabile
            self.new_list_widget.addItem(header_item)
            # Aggiungi le immagini alla QListWidget

            for i in self.record_us_list:
                search_dict = {'id_media': "'" + str(i.id_media) + "'"}
                u = Utility()
                search_dict = u.remove_empty_items_fr_dict(search_dict)
                mediathumb_data = self.DB_MANAGER.query_bool(search_dict, "MEDIA_THUMB")
                thumb_path = str(mediathumb_data[0].filepath)
                # Verifica se l'immagine è già in cache
                if thumb_path not in self.image_cache:
                    # Se non è in cache, carica l'immagine
                    icon = QIcon(thumb_path_str + thumb_path)

                    # Se la cache ha raggiunto il limite, rimuove l'elemento più vecchio
                    if len(self.image_cache) >= self.cache_limit:
                        self.image_cache.popitem(last=False)

                    # Aggiunge l'immagine alla cache
                    self.image_cache[thumb_path] = icon
                else:

                    icon = self.image_cache[thumb_path]


                self.image_cache.move_to_end(thumb_path)

                item = QListWidgetItem(str(i.media_filename))
                item.setData(Qt.UserRole, str(i.media_filename))
                icon = QIcon(thumb_path_str + thumb_path)
                item.setIcon(icon)

                item.setBackground(QColor("yellow"))

                self.new_list_widget.addItem(item)


        else:
            for image in all_images:
                # Crea un nuovo dizionario di ricerca per MEDIATOENTITY
                search_dict = {'id_media': "'" + str(image.id_media) + "'",
                               'entity_type': "'US'"}
                search_dict = u.remove_empty_items_fr_dict(search_dict)

                # Recupera l'elenco di 'US' associati all'immagine
                mediatoentity_data = self.DB_MANAGER.query_bool(search_dict, "MEDIATOENTITY")

                # Se l'immagine ha una o più 'US' associate, aggiungila all'elenco
                if mediatoentity_data:
                    us_images.append(image)

            if filter_text:  # se il filtro è attivo
                filtered_images = [i for i in untagged_images if filter_text.lower() in i.media_filename.lower()]
            else:
                filtered_images = us_images
            # Calcola gli indici di inizio e fine per la pagina corrente
            start_index = (self.current_page - 1) * self.page_size
            end_index = start_index + self.page_size

            # Ottieni i record delle immagini per la pagina corrente
            self.record_us_list = filtered_images[start_index:end_index]
            # Pulisci la QListWidget prima di aggiungere le nuove immagini
            self.new_list_widget.clear()
            # Aggiungi l'intestazione alla QListWidget
            header_item = QListWidgetItem(
                "Le righe selezionate in giallo indicano immagini non taggate\n Da questo strumento solo le righe selezionate gialle posso essere taggate ")
            header_item.setBackground(QColor('lightgrey'))
            header_item.setFlags(header_item.flags() & ~Qt.ItemIsSelectable)  # rendi l'item non selezionabile
            self.new_list_widget.addItem(header_item)
            # Aggiungi le immagini alla QListWidget

            for i in self.record_us_list:
                search_dict = {'id_media': "'" + str(i.id_media) + "'"}
                u = Utility()
                search_dict = u.remove_empty_items_fr_dict(search_dict)
                mediathumb_data = self.DB_MANAGER.query_bool(search_dict, "MEDIA_THUMB")
                thumb_path = str(mediathumb_data[0].filepath)
                # Verifica se l'immagine è già in cache
                if thumb_path not in self.image_cache:
                    # Se non è in cache, carica l'immagine
                    icon = QIcon(thumb_path_str + thumb_path)

                    # Se la cache ha raggiunto il limite, rimuove l'elemento più vecchio
                    if len(self.image_cache) >= self.cache_limit:
                        self.image_cache.popitem(last=False)

                    # Aggiunge l'immagine alla cache
                    self.image_cache[thumb_path] = icon
                else:
                    # Se è in cache, utilizza l'icona dalla cache
                    icon = self.image_cache[thumb_path]

                    # Aggiorna l'ordine della cache spostando l'elemento utilizzato alla fine
                self.image_cache.move_to_end(thumb_path)
                # Crea un nuovo dizionario di ricerca per MEDIATOENTITY
                search_dict = {'id_media': "'" + str(i.id_media) + "'",
                              'entity_type': "'US'"}
                search_dict = u.remove_empty_items_fr_dict(search_dict)
                #Recupera l'elenco di US associati all'immagine
                mediatoentity_data = self.DB_MANAGER.query_bool(search_dict, "MEDIATOENTITY")
                us_list = [str(g.id_entity) for g in mediatoentity_data]# Se 'entity_type' è 'US', aggiungi l'id_media a us_images
                #Rimuovi i duplicati dalla lista convertendola in un set e poi di nuovo in una lista
                us_list = list(set(us_list))
                us_list = [g.id_entity for g in mediatoentity_data if 'US' in g.entity_type]
                item = QListWidgetItem(str(i.media_filename))
                item.setData(Qt.UserRole, str(i.media_filename))
                icon = QIcon(thumb_path_str + thumb_path)
                item.setIcon(icon)
                if us_list:


                    item.setBackground(QColor("white"))



                    # Inizializza una lista vuota per i nomi delle US
                    us_names = []

                    for us_id in us_list:
                        # Crea un nuovo dizionario di ricerca per l'US
                        search_dict_us = {'id_us': us_id}
                        search_dict_us = u.remove_empty_items_fr_dict(search_dict_us)

                        # Query the US table
                        us_data = self.DB_MANAGER.query_bool(search_dict_us, "US")

                        # Se l'US esiste, aggiungi il suo nome alla lista
                        if us_data:
                            us_names.extend([str(us.us) for us in us_data])

                    # Se ci sono dei nomi US, aggiungi questi all'elemento
                    if us_names:
                        item.setText(item.text() + " - US: " + ', '.join(us_names))
                    else:
                        pass  # oppure: item.setText(item.text() + " - US: Non trovato")
                else:

                    item.setBackground(QColor("yellow"))

                # Aggiungi l'elemento alla QListWidget
                # self.new_list_widget.clear()
                self.new_list_widget.addItem(item)

            # Calcola il numero totale di pagine
            self.total_pages = math.ceil(len(filtered_images) / self.page_size)

            # Aggiorna l'aspetto delle etichette dei numeri delle pagine
            self.update_page_labels()

    def update_page_labels(self):
        # Disabilita il pulsante "Indietro" se siamo alla prima pagina
        self.prevButton.setEnabled(self.current_page > 1)

        # Disabilita il pulsante "Avanti" se siamo all'ultima pagina
        self.nextButton.setEnabled(self.current_page < self.total_pages)

        # Aggiorna l'aspetto delle etichette dei numeri delle pagine
        for label in self.pageLabels:
            page_number = int(label.text())
            label.setEnabled(page_number != self.current_page)

        # Aggiorna l'etichetta della pagina corrente e del totale delle pagine
        self.current_page_label.setText(f"Pagina corrente: {self.current_page}")
        self.total_pages_label.setText(f"Totale pagine: {self.total_pages}")

    def go_to_previous_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.load_images(self.current_filter_text)

    def go_to_next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.load_images(self.current_filter_text)

    def on_page_label_clicked(self, page, _=None):
        if page != self.current_page:
            self.current_page = page
            self.load_images(self.current_filter_text)

    def filter_items(self):
        # Ottieni il testo corrente nel campo di ricerca
        self.current_filter_text = self.search_field.text().lower()
        self.load_images(self.current_filter_text)

    def on_done_selecting_all(self):

        def r_list():
            sito = self.comboBox_sito.currentText()
            area = self.comboBox_area.currentText()
            us = self.lineEdit_us.text()
            record_us_list=[]
            #for sing_tags in selected_us:
            search_dict = {'sito': "'" + str(sito)+ "'",
                           'area': "'" + str(area) + "'",
                           'us': "'" + str(us) + "'"
                           }
            j = self.DB_MANAGER.query_bool(search_dict, 'US')
            record_us_list.append(j)
            us_list = []
            for r in record_us_list:
                us_list.append([r[0].id_us, 'US', 'us_table'])
            # QMessageBox.information(self, "Scheda US", str(us_list), QMessageBox.Ok)
            return us_list

        items_selected = self.new_list_widget.selectedItems()
        for item in items_selected:
            for us_data in r_list():
                id_orig_item = item.text()  # return the name of original file
                search_dict = {'filename': "'" + str(id_orig_item) + "'"}
                media_data = self.DB_MANAGER.query_bool(search_dict, 'MEDIA')

                # Check if media_data is not empty
                if media_data:
                    # Check if this image is already in the database
                    search_dict = {'id_media': "'" + str(media_data[0].id_media) + "'"}
                    existing_entry = self.DB_MANAGER.query_bool(search_dict, 'MEDIATOENTITY')

                    # If this image is already in the database, continue with the next item
                    if existing_entry:
                        continue

                    self.insert_mediaToEntity_rec(us_data[0], us_data[1], us_data[2], media_data[0].id_media,
                                                  media_data[0].filepath, media_data[0].filename)
                else:
                    pass
                    #QMessageBox.warning(self, "Attenzione",
                                        #"Immagine già taggata: " + str(id_orig_item))
                    # After tagging the image, update the corresponding QListWidgetItem

        # After tagging, update the iconListWidget
        self.fill_iconListWidget()
        self.update_list_widget_item(item)
    def update_list_widget_item(self,item):
        #items_selected = self.new_list_widg)et.selectedItems(
        search_dict = {'media_name': "'" + str(item.text()) + "'"}
        u = Utility()
        search_dict = u.remove_empty_items_fr_dict(search_dict)
        mediatoentity_data = self.DB_MANAGER.query_bool(search_dict, "MEDIATOENTITY")

        # Update the QListWidgetItem based on whether it matches
        if mediatoentity_data:
            item.setBackground(QColor("white"))

            # Create a new search dictionary for the US
            search_dict_us = {'id_us': "'" + str(mediatoentity_data[0].id_entity) + "'"}
            search_dict_us = u.remove_empty_items_fr_dict(search_dict_us)

            # Query the US table
            us_data = self.DB_MANAGER.query_bool(search_dict_us, "US")

            # If the US exists, add its name to the item
            if us_data:
                item.setText(item.text() + " - US: " + str(us_data[0].us))
            else:
                item.setText(item.text() + " - US: Not found")

        else:
            item.setBackground(QColor("yellow"))

    def fill_iconListWidget(self):
        #self.iconListWidget.clear()  # pulisci prima il widget
        items_selected = self.new_list_widget.selectedItems()
        for item in items_selected:
            item.text()
        # Prendi i dati dal tuo database o dalla tua fonte dati
        #data = self.DB_MANAGER.query('MEDIA_THUMB')
        search_dict = {'media_filename': "'" + str(item.text()) + "'"}
        u = Utility()
        search_dict = u.remove_empty_items_fr_dict(search_dict)
        data = self.DB_MANAGER.query_bool(search_dict, "MEDIA_THUMB")
        #QMessageBox.information(self, 'ok',str(item.text()))
        conn = Connection()

        thumb_path = conn.thumb_path()
        thumb_path_str = thumb_path['thumb_path']
        # crea un nuovo QListWidgetItem
        if data:
            list_item = QListWidgetItem(data[0].media_filename)  # utilizza il nome del file come testo dell'elemento
            list_item.setData(Qt.UserRole,data[0].media_filename)  # utilizza il nome del file come dati personalizzati dell'elemento

            # crea una QIcon con l'immagine
            #icon = QIcon(thumb_path_str + thumb_path)
            icon = QIcon(thumb_path_str + data[0].filepath)  # utilizza il percorso del file per creare l'icona
            #QMessageBox.information(self,'ok',str(thumb_path_str + data[0].filepath))
            # imposta l'icona dell'elemento
            list_item.setIcon(icon)

            # aggiungi l'elemento al QListWidget
            self.iconListWidget.addItem(list_item)


    def loadMediaPreview(self):
        self.iconListWidget.clear()
        conn = Connection()
        thumb_path = conn.thumb_path()
        thumb_path_str = thumb_path['thumb_path']
        # if mode == 0:
        # """ if has geometry column load to map canvas """
        rec_list = self.ID_TABLE + " = " + str(
            eval("self.DATA_LIST[int(self.REC_CORR)]." + self.ID_TABLE))
        search_dict = {
            'id_entity': "'" + str(eval("self.DATA_LIST[int(self.REC_CORR)]." + self.ID_TABLE)) + "'",
            'entity_type': "'US'"}
        record_us_list = self.DB_MANAGER.query_bool(search_dict, 'MEDIATOENTITY')
        for i in record_us_list:
            search_dict = {'id_media': "'" + str(i.id_media) + "'"}
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            mediathumb_data = self.DB_MANAGER.query_bool(search_dict, "MEDIA_THUMB")
            thumb_path = str(mediathumb_data[0].filepath)
            item = QListWidgetItem(str(i.media_name))
            item.setData(Qt.UserRole, str(i.media_name))
            icon = QIcon(thumb_path_str+thumb_path)
            item.setIcon(icon)
            self.iconListWidget.addItem(item)
        # elif mode == 1:
            # self.iconListWidget.clear()


    def load_and_process_3d_model(self, filepath):
        filename = os.path.basename(filepath)
        filename, filetype = filename.split(".")[0], filename.split(".")[1]
        mediatype = '3d_model'

        # Inserisci il record nel database
        self.insert_record_media(mediatype, filename, filetype, filepath)

        # Genera una thumbnail del modello 3D
        thumbnail_path = self.generate_3d_thumbnail(filepath)

        # Inserisci il record della thumbnail
        media_max_num_id = self.DB_MANAGER.max_num_id('MEDIA', 'id_media')
        self.insert_record_mediathumb(media_max_num_id, mediatype, filename, f"{filename}_thumb.png", 'png',
                                      thumbnail_path, filepath)

        # Aggiungi l'item alla lista
        item = QListWidgetItem(str(filename))
        item.setData(Qt.UserRole, str(media_max_num_id))
        icon = QIcon(thumbnail_path)
        item.setIcon(icon)
        self.iconListWidget.addItem(item)

        self.assignTags_US(item)


    def show_3d_model(self, file_path):
        mesh = pv.read(file_path)
        points = []
        measuring = False
        measurement_objects = []


        main_widget = QWidget()
        main_layout = QHBoxLayout()
        main_widget.setLayout(main_layout)

        frame = QFrame()
        layout = QVBoxLayout()

        plotter = QtInteractor(frame)

        debug_widget = QTextEdit()
        debug_widget.setReadOnly(True)

        def add_debug_message(message, important=False):
            timestamp = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
            formatted_message = f"[{timestamp}] {message}"
            if important:
                formatted_message = f"<b>{formatted_message}</b>"
            debug_widget.append(formatted_message)
            debug_widget.ensureCursorVisible()

            max_messages = 1000
            if debug_widget.document().blockCount() > max_messages:
                cursor = debug_widget.textCursor()
                cursor.movePosition(cursor.Start)
                cursor.select(cursor.LineUnderCursor)
                cursor.removeSelectedText()
                cursor.deletePreviousChar()
                cursor.movePosition(cursor.End)
                debug_widget.setTextCursor(cursor)

        def mouse_click_callback(obj, event):
            nonlocal measuring, points
            if event == "LeftButtonPressEvent" and measuring:
                x, y = plotter.interactor.GetEventPosition()
                add_debug_message(f"Evento di clic a posizione schermo: (x: {x}, y: {y})")

                picker = vtk.vtkCellPicker()
                picker.SetTolerance(10)  # Aumenta la tolleranza per migliorare il picking
                picker.Pick(x, y, 0, plotter.renderer)

                if picker.GetCellId() != -1:
                    point = np.array(picker.GetPickPosition())
                    add_debug_message(f"Punto selezionato nello spazio del modello: {point}")

                    closest_point_id = mesh.find_closest_point(point)
                    closest_point = mesh.points[closest_point_id]
                    add_debug_message(f"Punto più vicino sulla superficie della mesh: {closest_point}")

                    on_left_click(closest_point)
                else:
                    add_debug_message("Nessun punto sulla superficie trovato", important=True)

        plotter.interactor.AddObserver(vtk.vtkCommand.LeftButtonPressEvent, mouse_click_callback)

        layout.addWidget(plotter.interactor)
        plotter.clear()

        texture_file = os.path.splitext(file_path)[0] + '.jpg'
        if os.path.exists(texture_file):
            texture = pv.read_texture(texture_file)
            plotter.add_mesh(mesh, texture=texture, show_edges=False)
        else:
            plotter.add_mesh(mesh, show_edges=False)

        instructions_widget = QTextEdit()
        instructions_widget.setReadOnly(True)
        instructions_widget.hide()

        instructions = (
            "Trackball Controls:\n"
            "- Rotate: Left-click and drag\n"
            "- Pan: Right-click and drag\n"
            "- Zoom: Mouse wheel or middle-click and drag\n"
            "- Reset view: 'r'\n"
            "- Start/Stop measuring: 'o'\n"
            "- Show bounding box measures: 'm'\n"
            "- Export image: 'e'\n"
            "- Clear measurements: 'c'\n"
            "\nMain Views:\n"
            "- XY View (top): 'z'\n"
            "- YZ View (front): 'x'\n"
            "- XZ View (side): 'y'\n"
            "- ZY View (back): 'w'\n"
            "- ZX View (opposite side): 'v'\n"
            "- YX View (bottom): 'b'"
        )
        instructions_widget.setText(instructions)

        def toggle_instructions():
            if instructions_widget.isHidden():
                instructions_widget.show()
            else:
                instructions_widget.hide()

        instructions_button = QPushButton("Toggle Instructions")
        instructions_button.clicked.connect(toggle_instructions)
        layout.addWidget(instructions_button)

        frame.setLayout(layout)
        main_layout.addWidget(frame)
        main_layout.addWidget(instructions_widget)
        #main_layout.addWidget(debug_widget)

        def toggle_measure():
            nonlocal measuring, points
            measuring = not measuring
            points.clear()
            if measuring:
                add_debug_message("Misurazione iniziata", important=True)
            else:
                add_debug_message("Misurazione terminata", important=True)

        def on_left_click(picked_point):
            nonlocal points
            if not measuring:
                return

            add_debug_message(f"Punto selezionato: {picked_point}")

            if picked_point is not None:
                points.append(picked_point)
                sphere = pv.Sphere(radius=mesh.length * 0.005,
                                   center=picked_point)  # Aumenta leggermente il raggio della sfera per una miglior visibilità
                sphere_actor = plotter.add_mesh(sphere, color='red')
                measurement_objects.append(sphere_actor)

                add_debug_message(f"Punto aggiunto. Totale punti: {len(points)}")
                if len(points) == 2:
                    add_debug_message("Due punti raccolti. Misurazione in corso...", important=True)
                    measure_distance(points[0], points[1])
                    points.clear()
            else:
                add_debug_message("Nessun punto selezionato", important=True)

        def verify_coordinates(coord1, coord2):
            add_debug_message(f"Verifica delle coordinate:\nPunto1: {coord1}\nPunto2: {coord2}", important=True)

        def measure_distance(point1, point2):
            add_debug_message(f"Misurazione della distanza tra {point1} e {point2}")
            distance = np.linalg.norm(np.array(point1) - np.array(point2))

            line = pv.Line(point1, point2)
            line_actor = plotter.add_mesh(line, color='red', line_width=2)
            measurement_objects.append(line_actor)
            add_debug_message("Linea aggiunta")

            labels = plotter.add_point_labels([point1, point2], ["P1", "P2"], point_size=1, font_size=6)
            measurement_objects.append(labels)
            add_debug_message("Etichette dei punti aggiunte")

            mid_point = (np.array(point1) + np.array(point2)) / 2
            distance_label = plotter.add_point_labels([mid_point], [f"{distance:.2f} cm"], point_size=0, font_size=6)
            measurement_objects.append(distance_label)
            add_debug_message("Etichetta della distanza aggiunta")

            verify_coordinates(point1, mid_point)  # Verifica le coordinate durante la misura

            plotter.render()
            add_debug_message(f"Distanza misurata: {distance:.2f} cm", important=True)

        def clear_measurements():
            nonlocal measurement_objects, points
            for obj in measurement_objects:
                plotter.remove_actor(obj)
            measurement_objects.clear()
            points.clear()
            plotter.render()

        def export_image():
            try:
                options = QFileDialog.Options()
                file_path, _ = QFileDialog.getSaveFileName(self, "Save Image", "",
                                                           "PNG Files (*.png);;All Files (*)", options=options)
                if file_path:
                    camera = plotter.camera_position
                    width_cm, height_cm = 15, 10
                    width_inches, height_inches = width_cm / 2.54, height_cm / 2.54
                    dpi = 300
                    width_pixels, height_pixels = int(width_inches * dpi), int(height_inches * dpi)
                    plotter.screenshot(file_path, transparent_background=False,
                                       window_size=(width_pixels, height_pixels),
                                       return_img=False)
                    plotter.camera_position = camera
                    add_debug_message(f"Immagine salvata come {file_path}", important=True)
                    QMessageBox.information(self, "Success", f"Image saved {file_path} to 300 DPI (15x10 cm)")
            except Exception as e:
                add_debug_message(f'Error: {str(e)}', important=True)
                QMessageBox.warning(self, "Error", f"Error saving image: {str(e)}")



        def get_visible_faces(plotter, mesh):
            camera_position = np.array(plotter.camera_position[0])
            center = np.array(mesh.center)
            direction = camera_position - center
            normals = np.array([
                [1, 0, 0], [-1, 0, 0],
                [0, 1, 0], [0, -1, 0],
                [0, 0, 1], [0, 0, -1]
            ])
            return [i for i, normal in enumerate(normals) if np.dot(direction, normal) > 0]

        def edge_visibility(edge, visible_faces):
            edge_to_faces = {
                (0, 1): [0, 2, 4], (1, 2): [0, 1, 4], (2, 3): [0, 3, 4], (3, 0): [0, 2, 4],
                (4, 5): [1, 2, 5], (5, 6): [1, 3, 5], (6, 7): [1, 3, 5], (7, 4): [1, 2, 5],
                (0, 4): [2, 4, 5], (1, 5): [1, 4, 5], (2, 6): [1, 3, 5], (3, 7): [2, 3, 5]
            }
            return any(face in visible_faces for face in edge_to_faces[edge])

        def calculate_label_position(p1, p2, offset_factor=0.1):
            mid_point = (p1 + p2) / 2
            direction = p2 - p1
            length = np.linalg.norm(direction)
            normalized_direction = direction / length
            perpendicular = np.cross(normalized_direction, [0, 0, 1])
            if np.allclose(perpendicular, 0):
                perpendicular = np.cross(normalized_direction, [0, 1, 0])
            perpendicular = perpendicular / np.linalg.norm(perpendicular)
            return mid_point + perpendicular * (length * offset_factor)

        def create_oriented_label(plotter, position, text, direction, is_vertical=False):
            vtk_text = vtk.vtkBillboardTextActor3D()
            vtk_text.SetPosition(position)
            vtk_text.SetInput(text)
            vtk_text.GetTextProperty().SetFontSize(6)
            vtk_text.GetTextProperty().SetColor(0, 0, 0)  # Testo nero
            vtk_text.GetTextProperty().SetBackgroundColor(1, 1, 1)  # Sfondo bianco
            vtk_text.GetTextProperty().SetBackgroundOpacity(0.8)
            vtk_text.GetTextProperty().SetJustificationToCentered()
            vtk_text.GetTextProperty().SetVerticalJustificationToCentered()

            if is_vertical:
                angle = 90
            else:
                angle = np.degrees(np.arctan2(direction[1], direction[0]))
            vtk_text.SetOrientation(0, 0, angle)

            plotter.add_actor(vtk_text)
            return vtk_text

        self.last_update_time = 0
        self.update_interval = 0.5  # Secondi tra gli aggiornamenti
        bounding_box_visible = False

        def show_measures():
            nonlocal bounding_box_visible, measurement_objects
            if not bounding_box_visible:
                return

            current_time = time.time()
            if current_time - self.last_update_time < self.update_interval:
                return
            self.last_update_time = current_time

            # Rimuovi le misure esistenti
            for obj in measurement_objects:
                plotter.remove_actor(obj)
            measurement_objects = []

            bounds = mesh.bounds
            x_min, x_max, y_min, y_max, z_min, z_max = bounds
            point = np.array([
                [x_min, y_min, z_min], [x_max, y_min, z_min], [x_max, y_max, z_min], [x_min, y_max, z_min],
                [x_min, y_min, z_max], [x_max, y_min, z_max], [x_max, y_max, z_max], [x_min, y_max, z_max]
            ])

            edges = [
                (0, 1),  # Larghezza (X)
                (0, 3),  # Profondità (Y)
                (0, 4)  # Altezza (Z)
            ]

            get_visible_faces(plotter, mesh)

            for i, edge in enumerate(edges):
                #if edge_visibility(edge):
                p1, p2 = point[edge[0]], point[edge[1]]
                distance = np.linalg.norm(p2 - p1) * 100

                label_position = calculate_label_position(p1, p2)

                line = pv.Line(p1, p2)
                line_actor = plotter.add_mesh(line, color='black', line_width=0.8)
                measurement_objects.append(line_actor)

                label = f"{distance:.2f} cm"
                is_vertical = (i == 2)  # L'etichetta verticale è per l'altezza (Z)
                label_actor = create_oriented_label(plotter, label_position, label, p2 - p1, is_vertical)
                measurement_objects.append(label_actor)

            plotter.render()

        def toggle_bounding_box_measures():
            nonlocal bounding_box_visible
            bounding_box_visible = not bounding_box_visible
            if bounding_box_visible:
                show_measures()
                add_debug_message("Misure del bounding box attivate", important=True)
            else:
                for obj in measurement_objects:
                    plotter.remove_actor(obj)
                measurement_objects.clear()
                plotter.render()
                add_debug_message("Misure del bounding box disattivate", important=True)

        def camera_changed(obj, event):
            if bounding_box_visible:
                show_measures()

        plotter.iren.add_observer('InteractionEvent', camera_changed)

        def reset_view():
            plotter.reset_camera()

        def change_view(direction):
            getattr(plotter, f'view_{direction}')()

        plotter.add_key_event("o", toggle_measure)
        plotter.add_key_event("c", clear_measurements)
        plotter.add_key_event('e', export_image)
        plotter.add_key_event('m', toggle_bounding_box_measures)
        plotter.add_key_event('r', reset_view)
        plotter.add_key_event('x', lambda: change_view('yz'))
        plotter.add_key_event('y', lambda: change_view('xz'))
        plotter.add_key_event('z', lambda: change_view('xy'))
        plotter.add_key_event('w', lambda: change_view('zy'))
        plotter.add_key_event('v', lambda: change_view('zx'))
        plotter.add_key_event('b', lambda: change_view('yx'))

        plotter.add_orientation_widget(plotter.enable_trackball_style())
        plotter.enable_trackball_style()
        frame.show()
        return main_widget

    def generate_3d_thumbnail(self, filepath):

        mesh = pv.read(filepath)
        plotter = pv.Plotter(off_screen=True)
        plotter.add_mesh(mesh)
        plotter.camera_position = 'xy'

        # Genera un nome file unico per la thumbnail
        thumbnail_filename = f"{os.path.splitext(os.path.basename(filepath))[0]}_thumb.png"
        thumbnail_path = os.path.join(self.thumb_path, thumbnail_filename)

        plotter.screenshot(thumbnail_path)
        return thumbnail_path

    def process_3d_model(self, media_max_num_id, filepath, filename, thumb_path_str, thumb_resize_str,
                         media_thumb_suffix, media_resize_suffix):
        import pyvista as pv

        # Carica il modello 3D
        mesh = pv.read(filepath)

        # Genera una thumbnail
        plotter = pv.Plotter(off_screen=True)
        plotter.add_mesh(mesh)
        plotter.camera_position = 'xy'
        thumbnail_path = os.path.join(thumb_path_str, f"{media_max_num_id}_{filename}{media_thumb_suffix}")
        plotter.screenshot(thumbnail_path)

        # Copia il file originale nella cartella di resize (non possiamo ridimensionare un modello 3D come un'immagine)
        import shutil
        resize_path = os.path.join(thumb_resize_str, f"{media_max_num_id}_{filename}{media_resize_suffix}")
        shutil.copy(filepath, resize_path)
        # Controlla se esiste una texture JPG con lo stesso nome del modello
        texture_filename = os.path.splitext(filename)[0] + ".jpg"
        texture_filepath = os.path.join(os.path.dirname(filepath), texture_filename)

        if os.path.exists(texture_filepath):
            # Se la texture esiste, copiala nella cartella di resize
            texture_resize_path = os.path.join(thumb_resize_str, f"{media_max_num_id}_{texture_filename}")
            shutil.copy(texture_filepath, texture_resize_path)

        return thumbnail_path, resize_path

    def openWide_image(self):
        items = self.iconListWidget.selectedItems()
        conn = Connection()

        thumb_resize = conn.thumb_resize()
        thumb_resize_str = thumb_resize['thumb_resize']

        def process_file_path(file_path):
            return urllib.parse.unquote(file_path)

        def show_image(file_path):
            dlg = ImageViewer(self)
            dlg.show_image(file_path)
            dlg.exec_()

        def show_video(file_path):
            if self.video_player is None:
                self.video_player = VideoPlayerWindow(self, db_manager=self.DB_MANAGER,
                                                      icon_list_widget=self.iconListWidget,
                                                      main_class=self)
            self.video_player.set_video(file_path)
            self.video_player.show()

        def show_media(file_path, media_type):
            full_path = os.path.join(thumb_resize_str, file_path)
            if media_type == 'video':
                show_video(full_path)
            elif media_type == 'image':
                show_image(full_path)
            elif media_type == '3d_model':
                self.show_3d_model(file_path)
            else:
                QMessageBox.warning(self, "Error", f"Unsupported media type: {media_type}", QMessageBox.Ok)

        def query_media(search_dict, table="MEDIA_THUMB"):
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            try:
                return self.DB_MANAGER.query_bool(search_dict, table)
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Database query failed: {str(e)}", QMessageBox.Ok)
                return None

        for item in items:
            id_orig_item = item.text()
            search_dict = {'media_filename': f"'{id_orig_item}'"}
            res = query_media(search_dict)

            if res:

                file_path = process_file_path(os.path.join(thumb_resize_str, str(res[0].path_resize)))

                media_type = res[0].mediatype

                if media_type == '3d_model':
                    widget_3d = self.show_3d_model(file_path)

                    # Crea un nuovo QDialog per contenere il widget 3D
                    dialog = QDialog(self)
                    dialog.setWindowTitle("3D Model Viewer")
                    dialog_layout = QVBoxLayout()
                    dialog_layout.addWidget(widget_3d)
                    dialog.setLayout(dialog_layout)

                    # Imposta le dimensioni del dialog
                    dialog.resize(800, 600)  # Puoi modificare queste dimensioni come preferisci

                    # Mostra il dialog
                    dialog.exec_()
                else:
                    show_media(file_path, media_type)
            else:
                QMessageBox.warning(self, "Error", f"File not found: {id_orig_item}", QMessageBox.Ok)

    def charge_list(self):
        l = QgsSettings().value("locale/userLocale", QVariant)
        lang = ""
        for key, values in self.LANG.items():
            if values.__contains__(l):
                lang = str(key)
        lang = "'" + lang + "'"
        # lista sito
        sito_vl = self.UTILITY.tup_2_list_III(self.DB_MANAGER.group_by('site_table', 'sito', 'SITE'))
        try:
            sito_vl.remove('')
        except Exception as e:
            pass
        self.comboBox_sito.clear()
        self.comboBox_sito_rappcheck.clear()
        sito_vl.sort()
        self.comboBox_sito.addItems(sito_vl)
        self.comboBox_sito_rappcheck.addItems(sito_vl)

        # responsabile_vl = self.UTILITY.tup_2_list_III(self.DB_MANAGER.group_by('us_table', 'schedatore', 'US'))
        # try:
        #     responsabile_vl.remove('')
        # except:
        #     pass
        #
        # self.comboBox_schedatore.clear()
        # responsabile_vl.sort()
        # self.comboBox_schedatore.addItems(responsabile_vl)
        #
        #
        # responsabile2_vl = self.UTILITY.tup_2_list_III(self.DB_MANAGER.group_by('us_table', 'direttore_us', 'US'))
        # try:
        #     responsabile2_vl.remove('')
        # except:
        #     pass
        #
        # self.comboBox_direttore_us.clear()
        # responsabile2_vl.sort()
        # self.comboBox_direttore_us.addItems(responsabile2_vl)
        #
        #
        # responsabile3_vl = self.UTILITY.tup_2_list_III(self.DB_MANAGER.group_by('us_table', 'responsabile_us', 'US'))
        # try:
        #     responsabile3_vl.remove('')
        # except:
        #     pass
        #
        # self.comboBox_responsabile_us.clear()
        # responsabile3_vl.sort()
        # self.comboBox_responsabile_us.addItems(responsabile3_vl)
        #




        self.comboBox_settore.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.1' + "'"
        }
        settore = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        settore_vl = []
        for i in range(len(settore)):
            settore_vl.append(settore[i].sigla)
        settore_vl.sort()
        self.comboBox_settore.addItems(settore_vl)
        # lista soprintendenza
        self.comboBox_soprintendenza.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.2' + "'"
        }
        soprintendenza = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        soprintendenza_vl = []
        for i in range(len(soprintendenza)):
            soprintendenza_vl.append(soprintendenza[i].sigla_estesa)
        soprintendenza_vl.sort()
        self.comboBox_soprintendenza.addItems(soprintendenza_vl)
        # lista definizione_stratigrafica
        self.comboBox_def_strat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.3' + "'"
        }
        d_stratigrafica = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        d_stratigrafica_vl = []
        for i in range(len(d_stratigrafica)):
            d_stratigrafica_vl.append(d_stratigrafica[i].sigla_estesa)
        d_stratigrafica_vl.sort()
        self.comboBox_def_strat.addItems(d_stratigrafica_vl)
        # lista definizione interpretata
        self.comboBox_def_intepret.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.4' + "'"
        }
        d_interpretativa = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        d_interpretativa_vl = []
        for i in range(len(d_interpretativa)):
            d_interpretativa_vl.append(d_interpretativa[i].sigla_estesa)
        d_interpretativa_vl.sort()
        self.comboBox_def_intepret.addItems(d_interpretativa_vl)
        # lista funzione statica
        self.comboBox_funz_statica_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.5' + "'"
        }
        funz_statica = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        funz_statica_vl = []
        for i in range(len(funz_statica)):
            if funz_statica[i].sigla_estesa not in funz_statica_vl:
                funz_statica_vl.append(funz_statica[i].sigla_estesa)
        funz_statica_vl.sort()
        self.comboBox_funz_statica_usm.addItems(funz_statica_vl)
        #lista consistenza legante usm
        self.comboBox_consistenza_legante_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.6' + "'"
        }
        consistenza_legante_usm = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        consistenza_legante_usm_vl = []
        for i in range(len(consistenza_legante_usm)):
            if consistenza_legante_usm[i].sigla_estesa not in consistenza_legante_usm_vl:
                consistenza_legante_usm_vl.append(consistenza_legante_usm[i].sigla_estesa)
        consistenza_legante_usm_vl.sort()
        self.comboBox_consistenza_legante_usm.addItems(consistenza_legante_usm_vl)
        # lista scavato
        self.comboBox_scavato.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '203.203' + "'"
        }
        scavato = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        scavato_vl = []
        for i in range(len(scavato)):
            if scavato[i].sigla_estesa not in scavato_vl:
                scavato_vl.append(scavato[i].sigla_estesa)
        scavato_vl.sort()
        self.comboBox_scavato.addItems(scavato_vl)
        # lista metodo di scavo
        self.comboBox_metodo.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.8' + "'"
        }
        metodo = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        metodo_vl = []
        for i in range(len(metodo)):
            if metodo[i].sigla_estesa not in metodo_vl:
                metodo_vl.append(metodo[i].sigla_estesa)
        metodo_vl.sort()
        self.comboBox_metodo.addItems(metodo_vl)
        # lista formazione
        self.comboBox_formazione.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.9' + "'"
        }
        formazione = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        formazione_vl = []
        for i in range(len(formazione)):
            if formazione[i].sigla_estesa not in formazione_vl:
                formazione_vl.append(formazione[i].sigla_estesa)
        formazione_vl.sort()
        self.comboBox_formazione.addItems(formazione_vl)
        # lista modo formazione
        self.comboBox_modo_formazione.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.10' + "'"
        }
        modo_formazione = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        modo_formazione_vl = []
        for i in range(len(modo_formazione)):
            if modo_formazione[i].sigla_estesa not in modo_formazione_vl:
                modo_formazione_vl.append(modo_formazione[i].sigla_estesa)
        modo_formazione_vl.sort()
        self.comboBox_modo_formazione.addItems(modo_formazione_vl)
        # lista colore
        self.comboBox_colore.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '201.201' + "'"
        }
        colore = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        colore_vl = []
        for i in range(len(colore)):
            if colore[i].sigla_estesa not in colore_vl:
                colore_vl.append(colore[i].sigla_estesa)
        colore_vl.sort()
        self.comboBox_colore.addItems(colore_vl)
        # lista consistenza
        self.comboBox_consistenza.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.11' + "'"
        }
        consistenza = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        consistenza_vl = []
        for i in range(len(consistenza)):
            if consistenza[i].sigla_estesa not in consistenza_vl:
                consistenza_vl.append(consistenza[i].sigla_estesa)
        consistenza_vl.sort()
        self.comboBox_consistenza.addItems(consistenza_vl)
        # lista stato di conservazione
        self.comboBox_conservazione.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.12' + "'"
        }
        conservazione = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        conservazione_vl = []
        for i in range(len(conservazione)):
            if conservazione[i].sigla_estesa not in conservazione_vl:
                conservazione_vl.append(conservazione[i].sigla_estesa)
        conservazione_vl.sort()
        self.comboBox_conservazione.addItems(conservazione_vl)
        # lista schedatore
        self.comboBox_schedatore.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.16' + "'"
        }
        schedatore = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        schedatore_vl = []
        for i in range(len(schedatore)):
            if schedatore[i].sigla_estesa not in schedatore_vl:
                schedatore_vl.append(schedatore[i].sigla_estesa)
        schedatore_vl.sort()
        self.comboBox_schedatore.addItems(schedatore_vl)
        #lista direttore us
        self.comboBox_direttore_us.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.17' + "'"
        }
        direttore_us = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        direttore_us_vl = []
        for i in range(len(direttore_us)):
            if direttore_us[i].sigla_estesa not in direttore_us_vl:
                direttore_us_vl.append(direttore_us[i].sigla_estesa)
        direttore_us_vl.sort()
        self.comboBox_direttore_us.addItems(direttore_us_vl)
        # # lista responsabile us
        self.comboBox_responsabile_us.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.18' + "'"
        }
        responsabile_us = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        responsabile_us_vl = []
        for i in range(len(responsabile_us)):
            if responsabile_us[i].sigla_estesa not in responsabile_us_vl:
                responsabile_us_vl.append(responsabile_us[i].sigla_estesa)
        responsabile_us_vl.sort()
        self.comboBox_responsabile_us.addItems(responsabile_us_vl)


        # # lista tipologia_opera
        self.comboBox_tipologia_opera.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.20' + "'"
        }
        tipologia_opera = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        tipologia_opera_us_vl = []
        for i in range(len(tipologia_opera)):
            if tipologia_opera[i].sigla_estesa not in tipologia_opera_us_vl:
                tipologia_opera_us_vl.append(tipologia_opera[i].sigla_estesa)
        tipologia_opera_us_vl.sort()
        self.comboBox_tipologia_opera.addItems(tipologia_opera_us_vl)
        # lista sezione muraria
        self.comboBox_sezione_muraria.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.21' + "'"
        }
        sezione_muraria = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        sezione_muraria_us_vl = []
        for i in range(len(sezione_muraria)):
            if sezione_muraria[i].sigla_estesa not in sezione_muraria_us_vl:
                sezione_muraria_us_vl.append(sezione_muraria[i].sigla_estesa)
        sezione_muraria_us_vl.sort()
        self.comboBox_sezione_muraria.addItems(sezione_muraria_us_vl)
        # lista superficie_analizzata
        self.comboBox_superficie_analizzata.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.22' + "'"
        }
        sup_analiz = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        sup_analiz_vl = []
        for i in range(len(sezione_muraria)):
            if sup_analiz[i].sigla_estesa not in sup_analiz_vl:
                sup_analiz_vl.append(sup_analiz[i].sigla_estesa)
        sezione_muraria_us_vl.sort()
        self.comboBox_superficie_analizzata.addItems(sup_analiz_vl)
        # lista orientamento
        self.comboBox_orientamento.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.23' + "'"
        }
        orientamento = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        orientamento_us_vl = []
        for i in range(len(orientamento)):
            if orientamento[i].sigla_estesa not in orientamento_us_vl:
                orientamento_us_vl.append(orientamento[i].sigla_estesa)
        orientamento_us_vl.sort()
        self.comboBox_orientamento.addItems(orientamento_us_vl)
        # lista materiali_laterizi
        self.comboBox_materiali_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.24' + "'"
        }
        materiali_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        materiali_lat_us_vl = []
        for i in range(len(materiali_lat)):
            if materiali_lat[i].sigla_estesa not in materiali_lat_us_vl:
                materiali_lat_us_vl.append(materiali_lat[i].sigla_estesa)
        materiali_lat_us_vl.sort()
        self.comboBox_materiali_lat.addItems(materiali_lat_us_vl)
        # lista lavorazione laterizi
        self.comboBox_lavorazione_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.25' + "'"
        }
        lavorazione_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        lavorazione_lat_us_vl = []
        for i in range(len(lavorazione_lat)):
            if lavorazione_lat[i].sigla_estesa not in lavorazione_lat_us_vl:
                lavorazione_lat_us_vl.append(lavorazione_lat[i].sigla_estesa)
        lavorazione_lat_us_vl.sort()
        self.comboBox_lavorazione_lat.addItems(lavorazione_lat_us_vl)
        # lista consistenza laterizi
        self.comboBox_consistenza_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.26' + "'"
        }
        consistenza_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        consistenza_lat_us_vl = []
        for i in range(len(consistenza_lat)):
            if consistenza_lat[i].sigla_estesa not in consistenza_lat_us_vl:
                consistenza_lat_us_vl.append(consistenza_lat[i].sigla_estesa)
        consistenza_lat_us_vl.sort()
        self.comboBox_consistenza_lat.addItems(consistenza_lat_us_vl)
        # lista forma laterizi
        self.comboBox_forma_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.27' + "'"
        }
        forma_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        forma_lat_us_vl = []
        for i in range(len(forma_lat)):
            if forma_lat[i].sigla_estesa not in forma_lat_us_vl:
                forma_lat_us_vl.append(forma_lat[i].sigla_estesa)
        forma_lat_us_vl.sort()
        self.comboBox_forma_lat.addItems(forma_lat_us_vl)
        # lista colore laterizi
        self.comboBox_colore_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.28' + "'"
        }
        colore_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        colore_lat_us_vl = []
        for i in range(len(colore_lat)):
            if colore_lat[i].sigla_estesa not in colore_lat_us_vl:
                colore_lat_us_vl.append(colore_lat[i].sigla_estesa)
        colore_lat_us_vl.sort()
        self.comboBox_colore_lat.addItems(colore_lat_us_vl)
        # lista impasto laterizi
        self.comboBox_impasto_lat.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.29' + "'"
        }
        impasto_lat = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        impasto_lat_us_vl = []
        for i in range(len(impasto_lat)):
            if impasto_lat[i].sigla_estesa not in impasto_lat_us_vl:
                impasto_lat_us_vl.append(impasto_lat[i].sigla_estesa)
        impasto_lat_us_vl.sort()
        self.comboBox_impasto_lat.addItems(impasto_lat_us_vl)
        # lista materiali litici
        self.comboBox_materiale_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.30' + "'"
        }
        materiale_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        materiale_p_us_vl = []
        for i in range(len(materiale_p)):
            if materiale_p[i].sigla_estesa not in materiale_p_us_vl:
                materiale_p_us_vl.append(materiale_p[i].sigla_estesa)
        materiale_p_us_vl.sort()
        self.comboBox_materiale_p.addItems(materiale_p_us_vl)
        # lista consistenza materiali litici
        self.comboBox_consistenza_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.31' + "'"
        }
        consistenza_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        consistenza_p_us_vl = []
        for i in range(len(consistenza_p)):
            if consistenza_p[i].sigla_estesa not in consistenza_p_us_vl:
                consistenza_p_us_vl.append(consistenza_p[i].sigla_estesa)
        consistenza_p_us_vl.sort()
        self.comboBox_consistenza_p.addItems(consistenza_p_us_vl)
        # lista forma materiali litici
        self.comboBox_forma_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.32' + "'"
        }
        forma_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        forma_p_us_vl = []
        for i in range(len(forma_p)):
            if forma_p[i].sigla_estesa not in forma_p_us_vl:
                forma_p_us_vl.append(forma_p[i].sigla_estesa)
        forma_p_us_vl.sort()
        self.comboBox_forma_p.addItems(forma_p_us_vl)
        # lista colore materiali litici
        self.comboBox_colore_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.33' + "'"
        }
        colore_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        colore_p_us_vl = []
        for i in range(len(colore_p)):
            if colore_p[i].sigla_estesa not in colore_p_us_vl:
                colore_p_us_vl.append(colore_p[i].sigla_estesa)
        colore_p_us_vl.sort()
        self.comboBox_colore_p.addItems(colore_p_us_vl)
        # lista taglio materiali litici
        self.comboBox_taglio_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.34' + "'"
        }
        taglio_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        taglio_p_us_vl = []
        for i in range(len(taglio_p)):
            if taglio_p[i].sigla_estesa not in taglio_p_us_vl:
                taglio_p_us_vl.append(taglio_p[i].sigla_estesa)
        taglio_p_us_vl.sort()
        self.comboBox_taglio_p.addItems(taglio_p_us_vl)
        # lista posa opera materiali litici
        self.comboBox_posa_opera_p.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.35' + "'"
        }
        posa_opera_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        posa_opera_p_us_vl = []
        for i in range(len(posa_opera_p)):
            if posa_opera_p[i].sigla_estesa not in posa_opera_p_us_vl:
                posa_opera_p_us_vl.append(posa_opera_p[i].sigla_estesa)
        posa_opera_p_us_vl.sort()
        self.comboBox_posa_opera_p.addItems(posa_opera_p_us_vl)
        # lista posa opera materiali laterizi
        self.comboBox_posa_in_opera_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.36' + "'"
        }
        posa_opera_usm = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        posa_opera_usm_us_vl = []
        for i in range(len(posa_opera_usm)):
            if posa_opera_usm[i].sigla_estesa not in posa_opera_usm_us_vl:
                posa_opera_usm_us_vl.append(posa_opera_usm[i].sigla_estesa)
        posa_opera_usm_us_vl.sort()
        self.comboBox_posa_in_opera_usm.addItems(posa_opera_usm_us_vl)
        # lista tecniche costruttive
        self.comboBox_tecnica_muraria_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.37' + "'"
        }
        t_costruttiva = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        t_costruttiva_us_vl = []
        for i in range(len(t_costruttiva)):
            if t_costruttiva[i].sigla_estesa not in t_costruttiva_us_vl:
                t_costruttiva_us_vl.append(t_costruttiva[i].sigla_estesa)
        t_costruttiva_us_vl.sort()
        self.comboBox_tecnica_muraria_usm.addItems(t_costruttiva_us_vl)
        # lista modulo
        self.comboBox_modulo_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.38' + "'"
        }
        modulo = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        modulo_us_vl = []
        for i in range(len(modulo)):
            if modulo[i].sigla_estesa not in modulo_us_vl:
                modulo_us_vl.append(modulo[i].sigla_estesa)
        modulo_us_vl.sort()
        self.comboBox_modulo_usm.addItems(modulo_us_vl)
        # lista inerti
        self.comboBox_inerti_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.39' + "'"
        }
        inerti = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        inerti_us_vl = []
        for i in range(len(inerti)):
            if inerti[i].sigla_estesa not in inerti_us_vl:
                inerti_us_vl.append(inerti[i].sigla_estesa)
        inerti_us_vl.sort()
        self.comboBox_inerti_usm.addItems(inerti_us_vl)
        # lista tipologia legante
        self.comboBox_tipo_legante_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.40' + "'"
        }
        tipo_legante = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        tipo_legante_us_vl = []
        for i in range(len(tipo_legante)):
            if tipo_legante[i].sigla_estesa not in tipo_legante_us_vl:
                tipo_legante_us_vl.append(tipo_legante[i].sigla_estesa)
        tipo_legante_us_vl.sort()
        self.comboBox_tipo_legante_usm.addItems(tipo_legante_us_vl)
        # lista rifinitura
        self.comboBox_rifinitura_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.41' + "'"
        }
        rifinitura = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        rifinitura_us_vl = []
        for i in range(len(rifinitura)):
            if rifinitura[i].sigla_estesa not in rifinitura_us_vl:
                rifinitura_us_vl.append(rifinitura[i].sigla_estesa)
        rifinitura_us_vl.sort()
        self.comboBox_rifinitura_usm.addItems(rifinitura_us_vl)
        # lista lavorazione litica
        self.comboBox_lavorazione_usm.clear()
        search_dict = {
            'lingua': lang,
            'nome_tabella': "'" + 'us_table' + "'",
            'tipologia_sigla': "'" + '2.42' + "'"
        }
        lavorazione_p = self.DB_MANAGER.query_bool(search_dict, 'PYARCHINIT_THESAURUS_SIGLE')
        lavorazione_p_us_vl = []
        for i in range(len(lavorazione_p)):
            if lavorazione_p[i].sigla_estesa not in lavorazione_p_us_vl:
                lavorazione_p_us_vl.append(lavorazione_p[i].sigla_estesa)
        lavorazione_p_us_vl.sort()
        self.comboBox_lavorazione_usm.addItems(lavorazione_p_us_vl)

    def msg_sito(self):
        #self.model_a.database().close()
        conn = Connection()
        sito_set = conn.sito_set()
        sito_set_str = sito_set['sito_set']

        if bool(self.comboBox_sito.currentText()) and self.comboBox_sito.currentText()==sito_set_str:

            if self.L=='it':
                QMessageBox.information(self, "OK" ,"Sei connesso al sito: %s" % str(sito_set_str),QMessageBox.Ok)

            elif self.L=='de':
                QMessageBox.information(self, "OK", "Sie sind mit der archäologischen Stätte verbunden: %s" % str(sito_set_str),QMessageBox.Ok)

            else:
                QMessageBox.information(self, "OK", "You are connected to the site: %s" % str(sito_set_str),QMessageBox.Ok)
            #self.comboBox_sito.setCurrentText(sito_set_str)
        elif sito_set_str=='':
            if self.L=='it':
                msg = QMessageBox.information(self, "Attenzione" ,"Non hai settato alcun sito. Vuoi settarne uno? click Ok altrimenti Annulla per  vedere tutti i record",QMessageBox.Ok | QMessageBox.Cancel)
            elif self.L=='de':
                msg = QMessageBox.information(self, "Achtung", "Sie haben keine archäologischen Stätten eingerichtet. Klicken Sie auf OK oder Abbrechen, um alle Datensätze zu sehen",QMessageBox.Ok | QMessageBox.Cancel)
            else:
                msg = QMessageBox.information(self, "Warning" , "You have not set up any archaeological site. Do you want to set one? click Ok otherwise Cancel to see all records",QMessageBox.Ok | QMessageBox.Cancel)
            if msg == QMessageBox.Cancel:
                pass
            else:
                dlg = pyArchInitDialog_Config(self)
                dlg.charge_list()
                dlg.exec_()

    def set_sito(self):

        conn = Connection()
        sito_set = conn.sito_set()
        sito_set_str = sito_set['sito_set']
        try:
            if sito_set_str:
                search_dict = {'sito': "'" + str(sito_set_str) + "'"}  # 1 - Sito
                u = Utility()
                search_dict = u.remove_empty_items_fr_dict(search_dict)
                res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
                self.DATA_LIST = list(res)  # Convert the result to a list directly
                if self.DATA_LIST:  # Check if DATA_LIST is not empty
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.fill_fields()
                    self.BROWSE_STATUS = "b"
                    self.SORT_STATUS = "n"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    self.setComboBoxEnable(["self.comboBox_sito"], "False")
                else:
                    raise ValueError("No records found for the specified site.")
            else:
                pass#self.setComboBoxEnable(["self.comboBox_sito"], "True")
        except Exception as e:
            if self.L == 'it':
                QMessageBox.information(self, "Attenzione",
                                        f"Non esiste questo sito: '{sito_set_str}' in questa scheda. "
                                        "Per favore disattiva la 'scelta sito' dalla scheda di configurazione plugin per vedere tutti i record oppure crea la scheda.",
                                        QMessageBox.Ok)
            elif self.L == 'de':
                QMessageBox.information(self, "Warnung",
                                        f"Es gibt keine solche archäologische Stätte: '{sito_set_str}' in dieser Registerkarte. "
                                        "Bitte deaktivieren Sie die 'Site-Wahl' in der Plugin-Konfigurationsregisterkarte, um alle Datensätze zu sehen oder die Registerkarte zu erstellen.",
                                        QMessageBox.Ok)
            else:
                QMessageBox.information(self, "Warning",
                                        f"There is no such site: '{sito_set_str}' in this tab. "
                                        "Please disable the 'site choice' from the plugin configuration tab to see all records or create the tab.",
                                        QMessageBox.Ok)
    def generate_list_foto(self):
        data_list_foto = []
        for i in range(len(self.DATA_LIST)):
            conn = Connection()
            thumb_path = conn.thumb_path()
            thumb_path_str = thumb_path['thumb_path']

            if thumb_path_str=='':
                if self.L=='it':
                    QMessageBox.information(self, "Info", "devi settare prima la path per salvare le thumbnail . Vai in impostazioni di sistema/ path setting ")
                elif self.L=='de':
                    QMessageBox.information(self, "Info", "müssen Sie zuerst den Pfad zum Speichern der Miniaturansichten und Videos festlegen. Gehen Sie zu System-/Pfad-Einstellung")
                else:
                    QMessageBox.information(self, "Message", "you must first set the path to save the thumbnails and videos. Go to system/path setting")
            else:
                search_dict = {'id_entity': "'"+ str(eval("self.DATA_LIST[i].id_us"))+"'", 'entity_type' : "'US'"}
                record_doc_list = self.DB_MANAGER.query_bool(search_dict, 'MEDIAVIEW')
                for media in record_doc_list:
                    thumbnail = (thumb_path_str+media.filepath)
                    foto= (media.id_media)
                    # #sito= (media.sito)
                    # area= (media.area)
                    # us= (media.us)
                    # d_stratigrafica= ''
                    # unita_tipo = (media.unita_tipo)
                    data_list_foto.append([
                        str(self.DATA_LIST[i].sito.replace('_',' ')), #0
                        str(self.DATA_LIST[i].area), #1
                        str(self.DATA_LIST[i].us),    #2
                        str(self.DATA_LIST[i].unita_tipo),#3
                        str(self.DATA_LIST[i].d_stratigrafica),  #4
                        str(foto),#5
                        str(thumbnail)])#6
        return data_list_foto
            # #####################fine########################
    def generate_list_pdf(self):
        data_list = []
        #############inserimento nome fiel media############
        for i in range(len(self.DATA_LIST)):
            # assegnazione valori di quota mn e max
            id_us = str(self.DATA_LIST[i].id_us)
            sito = str(self.DATA_LIST[i].sito)#.replace('_',' '))
            area = str(self.DATA_LIST[i].area)
            us = str(self.DATA_LIST[i].us)
            res = self.DB_MANAGER.select_quote_from_db_sql(sito, area, us)
            quote = []
            for sing_us in res:
                sing_quota_value = str(sing_us[5])
                if sing_quota_value[0] == '-':
                    sing_quota_value = sing_quota_value[:7]
                else:
                    sing_quota_value = sing_quota_value[:6]
                sing_quota = [sing_quota_value, sing_us[4]]
                quote.append(sing_quota)
            quote.sort()
            #QMessageBox.information(self, "Message", str(quote))
            if bool(quote):
                quota_min = '%s %s' % (quote[0][0], quote[0][1])
                quota_max = '%s %s' % (quote[-1][0], quote[-1][1])

            else:
                if self.L=='it':
                    quota_min = ""
                    quota_max = ""
                elif self.L == 'de':
                    quota_min = ""
                    quota_max = ""
                else :
                    quota_min = ""
                    quota_max = ""
                # assegnazione numero di pianta
            resus = self.DB_MANAGER.select_us_from_db_sql(sito, area, us, "2")
            elenco_record = []
            for us in resus:
                elenco_record.append(us)
            if bool(elenco_record):
                sing_rec = elenco_record[0]
                elenco_piante = sing_rec[6]
                if elenco_piante != None:
                    piante = elenco_piante
                else:
                    if self.L=='it':
                        piante = "US disegnata su base GIS"
                    elif self.L=='de':
                        piante = "SE im GIS gezeichnet"
                    else:
                        piante= "SU draft on GIS"
            else:
                if self.L=='it':
                    piante = "US disegnata su base GIS"
                elif self.L=='de':
                    piante = "SE im GIS gezeichnet"
                else:
                    piante= "SU draft on GIS"
            if self.DATA_LIST[i].quota_min_usm == None:
                quota_min_usm = ""
            else:
                quota_min_usm = str(self.DATA_LIST[i].quota_min_usm)
            if self.DATA_LIST[i].quota_max_usm == None:
                quota_max_usm = ""
            else:
                quota_max_usm = str(self.DATA_LIST[i].quota_max_usm)
            #nuovi campi per Archeo3
            if not self.DATA_LIST[i].quota_relativa:
                quota_relativa = ""  # 55
            else:
                quota_relativa = str(self.DATA_LIST[i].quota_relativa)
            if not self.DATA_LIST[i].quota_abs:
                quota_abs = ""  # 56
            else:
                quota_abs = str(self.DATA_LIST[i].quota_abs)
            if not self.DATA_LIST[i].lunghezza_max:
                lunghezza_max = ""
            else:
                lunghezza_max = str(self.DATA_LIST[i].lunghezza_max)  # 65 lunghezza max
            if not self.DATA_LIST[i].altezza_max:
                altezza_max = ""
            else:
                altezza_max = str(self.DATA_LIST[i].altezza_max)  # 66 altezza max
            if not self.DATA_LIST[i].altezza_min:
                altezza_min = ""
            else:
                altezza_min = str(self.DATA_LIST[i].altezza_min)  # 67 altezza min
            if not self.DATA_LIST[i].profondita_max:
                profondita_max = ""
            else:
                profondita_max = str(self.DATA_LIST[i].profondita_max)  # 68 profondita_max
            if not self.DATA_LIST[i].profondita_min:
                profondita_min = ""
            else:
                profondita_min = str(self.DATA_LIST[i].profondita_min)  # 69 profondita min
            if not self.DATA_LIST[i].larghezza_media:
                larghezza_media = ""
            else:
                larghezza_media = str(self.DATA_LIST[i].larghezza_media)  # 70 larghezza media
            if not self.DATA_LIST[i].quota_max_abs:
                quota_max_abs = ""
            else:
                quota_max_abs = str(self.DATA_LIST[i].quota_max_abs)  # 71 quota_max_abs
            if not self.DATA_LIST[i].quota_max_rel:
                quota_max_rel = ""
            else:
                quota_max_rel = str(self.DATA_LIST[i].quota_max_rel)  # 72 quota_max_rel
            if not self.DATA_LIST[i].quota_min_abs:
                quota_min_abs = ""
            else:
                quota_min_abs = str(self.DATA_LIST[i].quota_min_abs)  # 73 quota_min_abs
            if not self.DATA_LIST[i].quota_min_rel:
                quota_min_rel = ""
            else:
                quota_min_rel = str(self.DATA_LIST[i].quota_min_rel)  # 74 quota_min_rel
            if not self.DATA_LIST[i].lunghezza_usm:
                lunghezza_usm = ""
            else:
                lunghezza_usm = str(self.DATA_LIST[i].lunghezza_usm)  # 85 lunghezza usm
            if not self.DATA_LIST[i].altezza_usm:
                altezza_usm = ""
            else:
                altezza_usm = str(self.DATA_LIST[i].altezza_usm)  # 86 altezza usm
            if not self.DATA_LIST[i].spessore_usm:
                spessore_usm = ""
            else:
                spessore_usm = str(self.DATA_LIST[i].spessore_usm)  # 87 spessore usm
            data_list.append([
                str(self.DATA_LIST[i].sito.replace('_',' ')),  # 0 - Sito
                str(self.DATA_LIST[i].area),  # 1 - Area
                int(self.DATA_LIST[i].us),  # 2 - US
                str(self.DATA_LIST[i].d_stratigrafica),  # 3 - definizione stratigrafica
                str(self.DATA_LIST[i].d_interpretativa),  # 4 - definizione intepretata
                str(self.DATA_LIST[i].descrizione),  # 5 - descrizione
                str(self.DATA_LIST[i].interpretazione),  # 6 - interpretazione
                str(self.DATA_LIST[i].periodo_iniziale),  # 7 - periodo iniziale
                str(self.DATA_LIST[i].fase_iniziale),  # 8 - fase iniziale
                str(self.DATA_LIST[i].periodo_finale),  # 9 - periodo finale iniziale
                str(self.DATA_LIST[i].fase_finale),  # 10 - fase finale
                str(self.DATA_LIST[i].scavato),  # 11 - scavato
                str(self.DATA_LIST[i].attivita),  # 12 - attivita
                str(self.DATA_LIST[i].anno_scavo),  # 13 - anno scavo
                str(self.DATA_LIST[i].metodo_di_scavo),  # 14 - metodo
                str(self.DATA_LIST[i].inclusi),  # 15 - inclusi
                str(self.DATA_LIST[i].campioni),  # 16 - campioni
                str(self.DATA_LIST[i].rapporti),            # 17 - rapporti
                #str(self.DATA_LIST[i].organici),  # organici
                #str(self.DATA_LIST[i].inorganici),  # inorganici
                str(self.DATA_LIST[i].data_schedatura),  # 18 - data schedatura
                str(self.DATA_LIST[i].schedatore),  # 19 - schedatore
                str(self.DATA_LIST[i].formazione),  # 20 - formazione
                str(self.DATA_LIST[i].stato_di_conservazione),  # 21 - conservazione
                str(self.DATA_LIST[i].colore),  # 22 - colore
                str(self.DATA_LIST[i].consistenza),  # 23 - consistenza
                str(self.DATA_LIST[i].struttura),  # 24 - struttura
                str(quota_min),  # 25 - quota_min
                str(quota_max),  # 26 - quota_max
                str(piante),  # 27 - piante CAMPO RICAVATO DA GIS CON VALORI SI/NO
                str(self.DATA_LIST[i].documentazione),  # 28 - documentazione
                #campi USM
                str(self.DATA_LIST[i].unita_tipo),  # 29 - unita tipo
                str(self.DATA_LIST[i].settore),  # 30 - settore
                str(self.DATA_LIST[i].quad_par),  # 31 quadrato
                str(self.DATA_LIST[i].ambient),  # 32 ambiente
                str(self.DATA_LIST[i].saggio),  # 33 saggio
                str(self.DATA_LIST[i].elem_datanti),  # 34 - elem_datanti
                str(self.DATA_LIST[i].funz_statica),  # 35 - funz_statica
                str(self.DATA_LIST[i].lavorazione),  # 36 lavorazione
                str(self.DATA_LIST[i].spess_giunti),  # 37 spess_giunti
                str(self.DATA_LIST[i].letti_posa),            #38 letti posa
                str(self.DATA_LIST[i].alt_mod),               #39  al modulo
                str(self.DATA_LIST[i].un_ed_riass),           #40 unita edilizia riassuntiva
                str(self.DATA_LIST[i].reimp),                 #41 reimpiego
                str(self.DATA_LIST[i].posa_opera),            #42 posa opera
                str(quota_min_usm),                           #43 quota min usm
                str(quota_max_usm),                           #44 quota max usm
                str(self.DATA_LIST[i].cons_legante),          #45 cons legante
                str(self.DATA_LIST[i].col_legante),           #46 col legante
                str(self.DATA_LIST[i].aggreg_legante),        #47 aggreg legante
                str(self.DATA_LIST[i].con_text_mat),          #48  con text mat
                str(self.DATA_LIST[i].col_materiale),         #49  col materiale
                str(self.DATA_LIST[i].inclusi_materiali_usm),  #50 inclusi materili usm
                #NUOVI CAMPI PER ARCHEO3
                str(self.DATA_LIST[i].n_catalogo_generale),  # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
                str(self.DATA_LIST[i].n_catalogo_interno),  # 52 nr catalogo interno
                str(self.DATA_LIST[i].n_catalogo_internazionale),  # 53 nr catalogo internazionale
                str(self.DATA_LIST[i].soprintendenza),  # 54 nr soprintendenza
                str(quota_relativa), #55 quota relativa
                str(quota_abs),   #56 quota assoluta
                str(self.DATA_LIST[i].ref_tm),  # 57 ref tm
                str(self.DATA_LIST[i].ref_ra),  # 58 ref ra
                str(self.DATA_LIST[i].ref_n),  # 59 ref n
                str(self.DATA_LIST[i].posizione),  # 60 posizione
                str(self.DATA_LIST[i].criteri_distinzione),  #61 criteri distinzione
                str(self.DATA_LIST[i].modo_formazione),  # 62 modo formazione
                str(self.DATA_LIST[i].componenti_organici),  # 63 componenti organici
                str(self.DATA_LIST[i].componenti_inorganici),  # 64 #  componenti inorganici
                str(lunghezza_max), #65 lunghezza max
                str(altezza_max), #66 altezza max
                str(altezza_min),  #67 altezza min
                str(profondita_max),  #68 profondita max
                str(profondita_min),  #69 profondita min
                str(larghezza_media),  #70 larghezza media
                str(quota_max_abs),   #71 quota max assoluta
                str(quota_max_rel),   #72 quota max rel
                str(quota_min_abs),   #73 quota min assoluta
                str(quota_min_rel),   #74 quota min relativa
                str(self.DATA_LIST[i].osservazioni),  # 75 osservazioni
                str(self.DATA_LIST[i].datazione), # 76 datazione
                str(self.DATA_LIST[i].flottazione),  # 77 flottazione
                str(self.DATA_LIST[i].setacciatura),  # 78 setacciatura
                str(self.DATA_LIST[i].affidabilita),  # 79 affidabilita
                str(self.DATA_LIST[i].direttore_us),  # 80 direttore us
                str(self.DATA_LIST[i].responsabile_us),  # 81 responsabile us
                str(self.DATA_LIST[i].cod_ente_schedatore),  # 82 cod ente schedatore
                str(self.DATA_LIST[i].data_rilevazione),  # 83 data rilevazione
                str(self.DATA_LIST[i].data_rielaborazione),  # 84 data rielaborazione
                str(lunghezza_usm), #85 lunghezza usm
                str(altezza_usm),  #86 altezza usm
                str(spessore_usm),  #87 spessore usm
                str(self.DATA_LIST[i].tecnica_muraria_usm),  # 88 tecnica muraria usm
                str(self.DATA_LIST[i].modulo_usm),  # 89 modulo usm
                str(self.DATA_LIST[i].campioni_malta_usm),  # 90 campioni malta usm
                str(self.DATA_LIST[i].campioni_mattone_usm),  # 91 campioni mattone usm
                str(self.DATA_LIST[i].campioni_pietra_usm),  # 92 campioni pietra usm
                str(self.DATA_LIST[i].provenienza_materiali_usm),  # 93 provenienza_materiali_usm
                str(self.DATA_LIST[i].criteri_distinzione_usm),  # 94 criteri distinzione usm
                str(self.DATA_LIST[i].uso_primario_usm),  #95 uso primario
                str(self.DATA_LIST[i].tipologia_opera),
                str(self.DATA_LIST[i].sezione_muraria),
                str(self.DATA_LIST[i].superficie_analizzata),
                str(self.DATA_LIST[i].orientamento),
                str(self.DATA_LIST[i].materiali_lat),
                str(self.DATA_LIST[i].lavorazione_lat),
                str(self.DATA_LIST[i].consistenza_lat),
                str(self.DATA_LIST[i].forma_lat),
                str(self.DATA_LIST[i].colore_lat),
                str(self.DATA_LIST[i].impasto_lat),
                str(self.DATA_LIST[i].forma_p),
                str(self.DATA_LIST[i].colore_p),
                str(self.DATA_LIST[i].taglio_p),
                str(self.DATA_LIST[i].posa_opera_p),
                str(self.DATA_LIST[i].inerti_usm),
                str(self.DATA_LIST[i].tipo_legante_usm),
                str(self.DATA_LIST[i].rifinitura_usm),
                str(self.DATA_LIST[i].materiale_p),
                str(self.DATA_LIST[i].consistenza_p),
            ])
        return data_list
    def on_pushButton_exp_tavole_pressed(self):
        conn = Connection()
        conn_str = conn.conn_str()
        # QMessageBox.warning(self, "Messaggio", str(conn_str), QMessageBox.Ok)
        PU = Print_utility(self.iface, self.DATA_LIST)
        PU.progressBarUpdated.connect(self.updateProgressBar)
        if conn_str.find("postgresql") == 1:
            PU.first_batch_try("postgres")
        else:
            PU.first_batch_try("sqlite")
    @pyqtSlot(int, int)
    def updateProgressBar(self, tav, tot):
        value = (float(tav) / float(tot)) * 100
        int_value = int(value)
        self.progressBar.setValue(int_value)
        # text = ' di '.join([str(tav), str(tot)])
        # self.countLabel.setText(text)
    def on_pushButton_print_pressed(self):
        if self.L=='it':
            if self.checkBox_s_us.isChecked():
                US_pdf_sheet = generate_US_pdf()
                data_list = self.generate_list_pdf()
                US_pdf_sheet.build_US_sheets(data_list)
                QMessageBox.warning(self, 'Ok',"Esportazione terminata Schede US",QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_us.isChecked() :
                US_index_pdf = generate_US_pdf()
                data_list = self.generate_list_pdf()
                try:
                    if bool(data_list):
                        US_index_pdf.build_index_US(data_list, data_list[0][0])
                        QMessageBox.warning(self, 'Ok',"Esportazione terminata Elenco US",QMessageBox.Ok)
                    else:
                        QMessageBox.warning(self, 'ATTENZIONE',"L'elenco US non può essere esportato devi riempire prima le schede US",QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'ATTENZIONE',str(e),QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_foto_t.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                try:
                        if bool(data_list_foto):
                            US_index_pdf.build_index_Foto(data_list_foto, data_list_foto[0][0])
                            QMessageBox.warning(self, 'Ok',"Esportazione terminata Elenco Foto",QMessageBox.Ok)
                        else:
                            QMessageBox.warning(self, 'ATTENZIONE',"L'elenco foto non può essere esportato perchè non hai immagini taggate.",QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'ATTENZIONE',str(e),QMessageBox.Ok)
            if self.checkBox_e_foto.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                try:
                        if bool(data_list_foto):
                            US_index_pdf.build_index_Foto_2(data_list_foto, data_list_foto[0][0])
                            QMessageBox.warning(self, 'Ok',"Esportazione terminata Elenco Foto senza thumbanil",QMessageBox.Ok)
                        else:
                            QMessageBox.warning(self, 'ATTENZIONE',"L'elenco foto non può essere esportato perchè non hai immagini taggate.",QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'ATTENZIONE',str(e),QMessageBox.Ok)
        elif self.L=='en':
            if self.checkBox_s_us.isChecked():
                US_pdf_sheet = generate_US_pdf()
                data_list = self.generate_list_pdf()
                US_pdf_sheet.build_US_sheets_en(data_list)
                QMessageBox.warning(self, 'Ok',"Export finished SU Forms",QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_us.isChecked() :
                US_index_pdf = generate_US_pdf()
                data_list = self.generate_list_pdf()
                try:
                    if bool(data_list):
                        US_index_pdf.build_index_US_en(data_list, data_list[0][0])
                        QMessageBox.warning(self, 'Ok',"Export finished SU List",QMessageBox.Ok)
                    else:
                        QMessageBox.warning(self, 'WARNING',"The SU list cannot be exported you have to fill in the SU tabs before",QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNING',str(e),QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_foto_t.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                #QMessageBox.information(self, 'Ok',str(data_list_foto[0][0]),QMessageBox.Ok)
                try:
                    if bool(data_list_foto):
                        US_index_pdf.build_index_Foto_en(data_list_foto, data_list_foto[0][0])
                        QMessageBox.information(self, 'Ok',"Export finished SU List",QMessageBox.Ok)
                    else:
                        QMessageBox.warning(self, 'WARNING', 'The photo list cannot be exported because you do not have tagged images.',QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNING',str(e),QMessageBox.Ok)

            if self.checkBox_e_foto.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                try:
                        if bool(data_list_foto):
                            US_index_pdf.build_index_Foto_2_en(data_list_foto, data_list_foto[0][0])
                            QMessageBox.information(self, 'Ok', "Export finished Photo List without thumbanil",QMessageBox.Ok)
                        else:
                            QMessageBox.warning(self, 'WARNING', "The photo list cannot be exported because you do not have tagged images.",QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNING',str(e),QMessageBox.Ok)
        elif self.L=='de':
            if self.checkBox_s_us.isChecked():
                US_pdf_sheet = generate_US_pdf()
                data_list = self.generate_list_pdf()
                US_pdf_sheet.build_US_sheets_de(data_list)
                QMessageBox.warning(self, "Okay", "Export beendet",QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_us.isChecked() :
                US_index_pdf = generate_US_pdf()
                data_list = self.generate_list_pdf()
                try:
                    if bool(data_list):
                        US_index_pdf.build_index_US_de(data_list, data_list[0][0])
                        QMessageBox.warning(self, "Okay", "Export beendet",QMessageBox.Ok)
                    else:
                        QMessageBox.warning(self, 'WARNUNG', 'Die SE-Liste kann nicht exportiert werden, Sie müssen zuerst die SE-Formulare ausfüllen',QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNUNG',str(e),QMessageBox.Ok)
            else:
                pass
            if self.checkBox_e_foto_t.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                try:
                        if bool(data_list_foto):
                            US_index_pdf.build_index_Foto_de(data_list_foto, data_list_foto[0][0])
                            QMessageBox.warning(self, "Okay", "Fertige Fotoliste exportieren",QMessageBox.Ok)
                        else:
                            QMessageBox.warning(self, 'WARNUNG', 'Die Fotoliste kann nicht exportiert werden, da Sie keine markierten Bilder haben.',QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNUNG',str(e),QMessageBox.Ok)
            if self.checkBox_e_foto.isChecked():
                US_index_pdf = generate_US_pdf()
                data_list_foto = self.generate_list_foto()
                try:
                        if bool(data_list_foto):
                            US_index_pdf.build_index_Foto_2_de(data_list_foto, data_list_foto[0][0])
                            QMessageBox.warning(self, 'Ok', 'Fertige Fotoliste ohne Daumenballen exportieren',QMessageBox.Ok)
                        else:
                            QMessageBox.warning(self, 'WARNUNG', 'Die Fotoliste kann nicht exportiert werden, da Sie keine markierten Bilder haben.',QMessageBox.Ok)
                except Exception as e :
                    QMessageBox.warning(self, 'WARNUNG',str(e),QMessageBox.Ok)
    def setPathpdf(self):
        s = QgsSettings()
        dbpath = QFileDialog.getOpenFileName(
            self,
            "Set file name",
            self.PDFFOLDER,
            " PDF (*.pdf)"
        )[0]

        if dbpath:
            self.lineEdit_pdf_path.setText(dbpath)
            s.setValue('',dbpath)

    def setPathdot(self):
        s = QgsSettings()
        dbpath = QFileDialog.getOpenFileName(
            self,
            "Set file name",
            self.MATRIX_PATH,
            " Dot (*.dot)"
        )[0]

        if dbpath:
            self.lineEdit_input.setText(dbpath)
            s.setValue('',dbpath)

    def setPathgraphml(self):
        s = QgsSettings()
        dbpath = QFileDialog.getSaveFileName(
            self,
            "Set file name",
            self.MATRIX_PATH,
            " Graphml (*.graphml)"
        )[0]

        if dbpath:
            self.lineEdit_output.setText(dbpath)
            s.setValue('',dbpath)

    def setDoc_ref(self):
        s = QgsSettings()
        dbpath = QFileDialog.getSaveFileName(
            self,
            "Set file name",
            self.MATRIX_PATH,
            " All files (*.*)"
        )[0]
        filename=dbpath.split("/")[-1]
        if dbpath:
            self.mQgsFileWidget.setText('DosCo\\'+filename)
            s.setValue('',dbpath)
    def list2pipe(self,x):
        lista =[]
        if isinstance(x,str) and x.startswith('[') and '], ['  and ', ' in x:

            return ', '.join(str(e) for e in eval(x)).replace("]",'').replace("['Copre',",'').replace("['Coperto da',",'').replace("['Riempie',",'').replace("['Riempito da',",'').replace("['Taglia',",'').replace("['Tagliato da',",'').replace("['Si appoggia a',",'').replace("['Gli si appoggia',",'').replace("['Si lega a',",'').replace("['Uguale a',",'').replace("'",'').replace("Copre,",'').replace("Coperto da,",'').replace("Riempie,",'').replace("Riempito da,",'').replace("Taglia,",'').replace("Tagliato da,",'').replace("Si appoggia a,",'').replace("Gli si appoggia,",'').replace("Si lega a,",'').replace("Uguale a,",'')


        elif isinstance(x,str) and x.startswith('['):
            return ', '.join(str(e) for e in eval(x)[0])
        else:
            return x
    def on_pushButton_graphml_pressed(self):

        dottoxml='{}{}{}'.format(self.BIN, os.sep, 'dottoxml.py')
        try:
            input_file = self.lineEdit_input.text()
            output_file = self.lineEdit_output.text()

            python_path = sys.exec_prefix
            python_version = sys.version[:3]

            if platform.system()=='Windows':
                cmd = 'python'
                shell=True

            elif platform.system()=='Darwin':
                cmd = '/Applications/QGIS.app/Contents/MacOS/bin/python3'
                shell=False

            else:

                cmd = 'python3'
                shell=False

            subprocess.call([cmd, dottoxml,'-f', 'Graphml',input_file, output_file], shell=shell)



            with open(output_file, 'r') as file :
                filedata = file.read()

                # Replace the target string
                filedata = filedata.replace("b'", '')
                filedata = filedata.replace("graphml>'", 'graphml>')
                # Write the file out again


            with open(output_file, 'w') as file:

                file.write(filedata)


            sito_location = str(self.comboBox_sito.currentText())
            cfg_rel_path = os.path.join(os.sep, 'pyarchinit_DB_folder', 'config.cfg')
            file_path = '{}{}'.format(self.HOME, cfg_rel_path)
            conf = open(file_path, "r")
            data = conf.read()
            settings = Settings(data)
            settings.set_configuration()
            conf.close()

            db_username = settings.USER
            host = settings.HOST
            port = settings.PORT
            database_password=settings.PASSWORD
            db_names = settings.DATABASE
            server=settings.SERVER

            if server=='postgres':
                pass
                # # Create an SQLAlchemy engine instance
                # engine = create_engine(f"postgresql://{db_username}:{database_password}@{host}:{port}/{db_names}")
                #
                # # Create a session
                # Session = sessionmaker(bind=engine)
                # session = Session()
                #
                # try:
                #     # Perform your query using SQLAlchemy ORM or Core
                #     result2 = session.execute(text("""
                #         WITH RECURSIVE cte AS (
                #                 SELECT
                #                     rowid,
                #                     SPLIT_PART(rapporti, ';', 1) AS col,
                #                     SUBSTRING(rapporti FROM POSITION(';' IN rapporti) + 1) AS rest
                #                 FROM (
                #                     SELECT
                #                         rowid,
                #                         REPLACE(REPLACE(REPLACE(rapporti, '[[', '['), ']]', ']'), '], [', '];[') AS rapporti
                #                     FROM us_table
                #                     WHERE sito = 'sito_location'
                #                 ) AS initial
                #
                #                 UNION ALL
                #
                #                 SELECT
                #                     rowid,
                #                     SPLIT_PART(rest, ';', 1),
                #                     SUBSTRING(rest FROM POSITION(';' IN rest) + 1)
                #                 FROM cte
                #                 WHERE LENGTH(rest) > 0
                #             )
                #             SELECT
                #                 STRING_AGG(CASE WHEN col LIKE '[Copre,%' OR col LIKE '[Taglia,%' OR col LIKE '[Riempie,%' OR col LIKE '[Si appoggia a,%' THEN col END, ',') AS post,
                #                 STRING_AGG(CASE WHEN col LIKE '[Coperto da,%' OR col LIKE '[Riempito da,%' OR col LIKE '[Tagliato da,%' OR col LIKE '[Gli si appoggia,%' THEN col END, ',') AS ante,
                #                 STRING_AGG(CASE WHEN col LIKE '[Si lega a,%' OR col LIKE '[Uguale a,%' THEN col END, ',') AS contemp
                #             FROM cte
                #             GROUP BY rowid
                #             ORDER BY rowid;
                #
                #     """), {'sito_location': sito_location})
                #
                #     rows2 = result2.fetchall()
                #     col_names2 = ['Rapporto Posteriore', 'Rapporto Anteriore', 'Rapporto Contemporaneo']
                #     t2 = pd.DataFrame(rows2, columns=col_names2).applymap(self.list2pipe)
                #     t2.to_excel(writer, sheet_name='Rapporti', index=False)
                #
                #     # Configure the Excel sheets' column widths
                #     worksheet1 = writer.sheets['US']
                #     worksheet1.set_column('A:A', 30)
                #     worksheet1.set_column('B:B', 30)
                #     worksheet1.set_column('C:C', 30)
                #     worksheet1.set_column('D:D', 30)
                #
                #     worksheet2 = writer.sheets['Rapporti']
                #     worksheet2.set_column('A:A', 30)
                #     worksheet2.set_column('B:B', 30)
                #     worksheet2.set_column('C:C', 30)
                #
                #     # Close the Pandas Excel writer and output the Excel file
                #     writer.close()
                #     QMessageBox.information(self, "INFO", "Conversion completed", QMessageBox.Ok)
                #
                # except Exception as e:
                #     # Handle any errors that occur during the database operations
                #     QMessageBox.warning(self, "Error", str(e), QMessageBox.Ok)
                # finally:
                #     # Ensure the database session is closed when done
                #     session.close()



            elif server=='sqlite':


                sqlite_DB_path = '{}{}{}'.format(self.HOME, os.sep,"pyarchinit_DB_folder")

                file_path_sqlite = sqlite_DB_path+os.sep+db_names
                conn = sq.connect(file_path_sqlite)
                conn.enable_load_extension(True)


                #now we can load the extension
                # depending on your OS and sqlite/spatialite version you might need to add
                # '.so' (Linux) or '.dll' (Windows) to the extension name

                #mod_spatialite (recommended)
                #conn.execute('SELECT load_extension("mod_spatialite")')
                #conn.execute('SELECT InitSpatialMetaData(1);')
                cur = conn.cursor()
                cur2 = conn.cursor()

                name_= '%s' % (sito_location+'_' +  time.strftime('%Y%m%d_') + '.xlsx')
                dump_dir=os.path.join(self.MATRIX_PATH, name_)
                writer = pd.ExcelWriter(dump_dir, engine='xlsxwriter')
                workbook  = writer.book


                cur.execute("SELECT  area, us, attivita, datazione From us_table where sito='%s' order by rowid;" % sito_location)
                rows1 = cur.fetchall()

                col_names1 = ['Area','SU','Activities','Epoch']
                t1=pd.DataFrame(rows1,columns=col_names1).applymap(self.list2pipe)
                t1.to_excel(writer, sheet_name='SU Context',index=False)

                if self.L=='it':
                    cur2.execute("""WITH cte AS
                        ( SELECT rowid ,
                       SUBSTR(rapporti,  1, INSTR(rapporti || ';', ';') -1) col,
                       SUBSTR(rapporti, INSTR(rapporti || ';', ';') + 1) rest
                       FROM (SELECT rowid, REPLACE(REPLACE(REPLACE(rapporti, '[[', '['), ']]', ']'), '], [', '];[') rapporti FROM us_table
                       WHERE sito = """+"'"+sito_location+"'"+""")
                       UNION all
                       SELECT rowid us,
                       SUBSTR(rest, 1, INSTR(rest || ';', ';')  -1),
                       SUBSTR(rest, INSTR(rest || ';', ';') + 1)   FROM cte   WHERE LENGTH(rest) > 0 )
                       SELECT
                       GROUP_CONCAT(CASE WHEN col LIKE '[''Copre'',%' OR col LIKE '[''Taglia'',%'
                       OR col LIKE '[''Riempie'',%' OR col LIKE '[''Si appoggia a'',%'  THEN col END) post,

                       GROUP_CONCAT(CASE WHEN col LIKE '[''Coperto da'',%' OR col LIKE '[''Riempito da'',%'
                       OR col LIKE '[''Tagliato da'',%' OR col LIKE '[''Gli si appoggia'',%' THEN col END) ante,

                       GROUP_CONCAT(CASE WHEN col LIKE '[''Si lega a'',%' or col LIKE '[''Uguale a'',%' THEN col END) contemp

                        FROM cte GROUP BY rowid order by rowid""")

                else:
                    cur2.execute("""WITH cte AS
                        ( SELECT rowid ,
                       SUBSTR(rapporti,  1, INSTR(rapporti || ';', ';') -1) col,
                       SUBSTR(rapporti, INSTR(rapporti || ';', ';') + 1) rest
                       FROM (SELECT rowid, REPLACE(REPLACE(REPLACE(rapporti, '[[', '['), ']]', ']'), '], [', '];[') rapporti FROM us_table
                       WHERE sito = '""" + sito_location + """')
                       UNION all
                       SELECT rowid us,
                       SUBSTR(rest, 1, INSTR(rest || ';', ';')  -1),
                       SUBSTR(rest, INSTR(rest || ';', ';') + 1)   FROM cte   WHERE LENGTH(rest) > 0 )
                       SELECT
                       GROUP_CONCAT(CASE 
                           WHEN col LIKE '[''Covers'',%' OR col LIKE '[''Cuts'',%'
                           OR col LIKE '[''Fill'',%' OR col LIKE '[''Abuts'',%'
                           THEN SUBSTR(col, INSTR(col, ',') + 1)
                       END) post,
                       GROUP_CONCAT(CASE 
                           WHEN col LIKE '[''Covered by'',%' OR col LIKE '[''Filled by'',%'
                           OR col LIKE '[''Cut by'',%' OR col LIKE '[''Supports'',%'
                           THEN SUBSTR(col, INSTR(col, ',') + 1)
                       END) ante,
                       GROUP_CONCAT(CASE 
                           WHEN col LIKE '[''Connected to'',%' or col LIKE '[''Same as'',%'
                           THEN SUBSTR(col, INSTR(col, ',') + 1)
                       END) contemp
                        FROM cte GROUP BY rowid order by rowid""")

                rows2 = cur2.fetchall()
                col_names2 = ['Posterior', 'Anterior', 'Contemporary']

                def clean_relationship(value):
                    if value is None or value == '':
                        return ''

                    # Split la stringa per virgola
                    parts = value.split(',')
                    numbers = []

                    for part in parts:
                        # Cerca un numero isolato (non parte di una parola)
                        matches = re.findall(r'\b(\d+)\b', part)
                        if matches:
                            # Prendi solo il primo numero trovato
                            num = matches[0]
                            if num not in numbers and len(num.strip()) > 0:  # Evita duplicati e stringhe vuote
                                numbers.append(num)

                    # Rimuovi eventuali '1' che seguono altri numeri
                    filtered_numbers = []
                    prev_num = None
                    for num in numbers:
                        if num != '1' or prev_num is None:
                            filtered_numbers.append(num)
                        prev_num = num

                    return ', '.join(filtered_numbers)

                # Crea DataFrame e applica la pulizia
                t2 = pd.DataFrame(rows2, columns=col_names2)
                t2 = t2.applymap(clean_relationship)

                # Gestisci il primo record
                if not pd.isna(t2.iloc[0]['Posterior']):
                    # Se il primo record ha dati nella seconda colonna ma non nella prima
                    if pd.isna(t2.iloc[0]['Anterior']) and not pd.isna(t2.iloc[0]['Posterior']):
                        # Sposta i dati nella colonna corretta
                        t2.iloc[0]['Anterior'] = t2.iloc[0]['Posterior']
                        t2.iloc[0]['Posterior'] = ''

                t2.to_excel(writer, sheet_name='Relationships', index=False)

                worksheet1 = writer.sheets['SU Context']
                worksheet1.set_column('A:A', 30, None)
                worksheet1.set_column('B:B', 30, None)
                worksheet1.set_column('C:C', 30, None)
                worksheet1.set_column('D:D', 30, None)
                worksheet1.set_column('E:E', 30, None)


                worksheet2 = writer.sheets['Relationships']
                worksheet2.set_column('A:A', 30, None)
                worksheet2.set_column('B:B', 30, None)
                worksheet2.set_column('C:C', 30, None)
                writer.close()

            else:
                pass

            QMessageBox.information(self, "INFO", "Conversion completed",
                                    QMessageBox.Ok)
        except ValueError as e:
            QMessageBox.warning(self, "Error", str(e),
                                QMessageBox.Ok)


    def openpdfDir(self):
        HOME = os.environ['PYARCHINIT_HOME']
        path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_PDF_folder")
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])

    def on_pushButton_viewmatrix_pressed(self):
        try:
            id_us_dict = {}
            for i in range(len(self.DATA_LIST)):
                id_us_dict[self.DATA_LIST[i].us] = self.DATA_LIST[i].id_us
            dlg = pyarchinit_view_Matrix(self.iface, self.DATA_LIST, id_us_dict)
            dlg.generate_matrix()
            # Visualizza l'immagine con matplotlib
            HOME = os.environ['PYARCHINIT_HOME']
            path = '{}{}{}{}'.format(HOME, os.sep, "pyarchinit_Matrix_folder/",'Harris_matrix_viewtred.dot.jpg')
            if path:
                # Legge l'immagine dal percorso del file
                image = Image.open(path)
                # Crea una figura e un canvas e li salva come attributi dell'oggetto
                self.fig = plt.figure()
                self.canvas = self.fig.add_subplot(111)

                # Visualizza l'immagine sul canvas
                self.canvas.imshow(image)
                self.canvas.axis('off')  # Rimuove gli assi

                plt.show()  # Mostra l'immagine
        except AssertionError as e:
            print(e)

    def on_pushButton_export_matrix_pressed(self):
        if self.checkBox_ED.isChecked():

            id_us_dict = {}
            for i in range(len(self.DATA_LIST)):
                id_us_dict[self.DATA_LIST[i].us] = self.DATA_LIST[i].id_us

            dlg = pyarchinit_Interactive_Matrix(self.iface, self.DATA_LIST, id_us_dict)
            data_plot=dlg.generate_matrix_2()

            # ###interactive matrix###
            # matrix_path = '{}{}{}'.format(self.HOME, os.sep, "pyarchinit_Matrix_folder")
            # filename='Harris_matrix2ED_graphml.dot'
            # hm=os.path.join(matrix_path, filename)
            # gv = pgv.AGraph(hm, strict=False, directed=False)
            # dlg.plot_matrix(gv)
            # dlg.exec_()
        if not self.checkBox_ED.isChecked():
            id_us_dict = {}
            for i in range(len(self.DATA_LIST)):
                id_us_dict[self.DATA_LIST[i].us] = self.DATA_LIST[i].id_us
            dlg = pyarchinit_Interactive_Matrix(self.iface, self.DATA_LIST, id_us_dict)
            data_plot=dlg.generate_matrix()

            ###interactive matrix###
        # if self.checkBox_IM.isChecked():    
            # matrix_path = '{}{}{}'.format(self.HOME, os.sep, "pyarchinit_Matrix_folder")
            # filename='Harris_matrix_tred.dot'
            # hm=os.path.join(matrix_path, filename)
            # gv = pgv.AGraph(hm, strict=False, directed=True)
            # dlg.plot_matrix(gv)
            # dlg.exec_()
    def launch_matrix_exp_if(self, msg):
        if msg == QMessageBox.Ok:
            self.on_pushButton_export_matrix_pressed()
        else:
            pass
    def on_pushButton_orderLayers_pressed(self):
        # QMessageBox.warning(self, 'ATTENZIONE',
        #                     """Il sistema accetta come dataset da elaborare ricerche su singolo SITO e AREA. Se state lanciando il sistema su siti o aree differenti, i dati di siti differenti saranno sovrascritti. Per terminare il sistema dopo l'Ok premere Cancel.""",
        #                     QMessageBox.Ok)
        # self.launch_matrix_exp_if(QMessageBox.warning(self, 'ATTENZIONE',
        #                                               "Si consiglia di lanciare il matrix e controllare se sono presenti paradossi stratigrafici prima di proseguire",
        #                                               QMessageBox.Ok | QMessageBox.Cancel))
        if self.L=='it':
            self.launch_order_layer_if(QMessageBox.warning(self, 'ATTENZIONE',
                                                       "Se saranno presenti paradossi stratigrafici l'order layer non andrà a buon fine",
                                                       QMessageBox.Ok))
        elif self.L=='de':
            self.launch_order_layer_if(QMessageBox.warning(self, 'ACHTUNG',
                                                       "Bist du sicher das du fortfahren möchtest? Wenn aktuell stratigraphische Paradoxa auftauchen Könnte das System zusammenbrechen!",
                                                       QMessageBox.Ok | QMessageBox.Cancel))
        else:
            self.launch_order_layer_if(QMessageBox.warning(self, 'ATTENZIONE',
                                                       "Are you sure you want to go on? If there are stratigraphic paradoxes, the system could crush!",
                                                       QMessageBox.Ok | QMessageBox.Cancel))


    def format_message(self, sing_rapp, us):
        base_msg = sing_rapp[0]
        relativity_msg_mappings = {
            'it': "relativo a: ",
            'de': "bezüglich: ",
            'en': "concerning: "
        }
        default_relativity_msg = relativity_msg_mappings.get(self.L, "concerning: ")
        return base_msg + default_relativity_msg + str(us) + " \n"


    def launch_order_layer_if(self, msg):
        if msg == QMessageBox.Ok:
            # report errori rapporti stratigrafici
            if self.L=='it':
                msg_tipo_rapp = "Manca il tipo di rapporto nell'US: \n"
                msg_nr_rapp = "Manca il numero del rapporto nell'US: \n"
                msg_paradx_rapp = "Paradosso nei rapporti: \n"
                msg_us_mancanti = "Mancano le seguenti schede US presenti nei rapporti: \n"
            elif self.L=='de':
                msg_tipo_rapp = "Der Beziehungstyp fehlt in den SE: \n"
                msg_nr_rapp = "Die Berichtsnummer fehlt in den SE: \n"
                msg_paradx_rapp = "Paradox in Beziehungen: \n"
                msg_us_mancanti = "Folgende SE-formular fehlen in den Berichten: \n"
            else:
                msg_tipo_rapp = "The relationship type is missing in the SU: \n"
                msg_nr_rapp = "The report number is missing in the SU: \n"
                msg_paradx_rapp = "Paradox in relationships: \n"
                msg_us_mancanti = "The following SU forms are missing from the reports: \n"

            msg_nr_rapp = msg_tipo_rapp = ""
            data = []
            for sing_rec in self.DATA_LIST:
                us = sing_rec.us
                area = sing_rec.area
                sito = sing_rec.sito
                rapporti_stratigrafici = eval(sing_rec.rapporti)
                for sing_rapp in rapporti_stratigrafici:
                    if len(sing_rapp) != 4:  # cambiato da 2 a 5
                        msg_nr_rapp += self.format_message(sing_rapp, us)
                    try:
                        if sing_rapp[0] in self.RELATIONSHIP_TYPES:
                            try:
                                if sing_rapp[1] != '':
                                    harris_rapp = (str(us) + str(area) + str(sito),
                                                   str(sing_rapp[1]) + str(sing_rapp[2]) + str(sing_rapp[3]))
                                    data.append(harris_rapp)
                            except Exception as e:
                                msg_nr_rapp += str(us) + " \n"
                    except:
                        msg_tipo_rapp += str(us) + " \n"
            for i in data:
                temp_tup = (i[1], i[
                    0])  # controlla che nn vi siano rapporti inversi dentro la lista DA PROBLEMI CON GLI UGUALE A E I SI LEGA A
                #QMessageBox.warning(self, "Messaggio", "Temp_tup" + str(temp_tup), QMessageBox.Ok)
                if data.count(temp_tup) != 0:
                    msg_paradx_rapp = msg_paradx_rapp + '\n' + str(i) + '\n' + str(temp_tup)
                    data.remove(i)
                    # OK
                    ## QMessageBox.warning(self, "Messaggio", "DATA LIST" + str(data), QMessageBox.Ok)
                    # Blocca il sistema di ordinamento su un sito ed area specifci in base alla ricerca eseguita sulla scheda US


            sito = self.DATA_LIST[0].sito  # self.comboBox_sito_rappcheck.currentText()
            area = self.DATA_LIST[0].area  # self.comboBox_area.currentText()



            # Crea Order_layer_v2 con il widget

            #QMessageBox.warning(None, "Messaggio", "DATA LIST" + str(order_layer_dict), QMessageBox.Ok)
            # order_number = ""
            # us = ""
            # Il tuo codice modificato:
            # Crea l'istanza
            OL = Order_layer_graph(self.DB_MANAGER, sito, area)
            order_layer_dict = OL.main_order_layer()

            if order_layer_dict == "error":
                # La finestra rimane aperta con i dettagli dell'errore
                # L'utente può leggere il log e chiudere manualmente

                # Opzionale: mostra un messaggio aggiuntivo
                QMessageBox.critical(self, "Errore Matrix",
                                     "Errore durante la generazione del matrix.\n"
                                     "Controlla i dettagli nella finestra di log.\n\n"
                                     "La finestra rimarrà aperta per permetterti di leggere i dettagli.")
                return

            # Se arriviamo qui, order_layer_dict è un dizionario valido
            # Usa il nuovo metodo per aggiornare il database
            try:
                updates_count = OL.update_database_with_order(
                    self.DB_MANAGER,
                    self.MAPPER_TABLE_CLASS,
                    self.ID_TABLE,
                    sito,
                    area
                )

                # Mostra il risultato
                QMessageBox.information(self, "Completato",
                                        f"Matrix generato con successo!\n"
                                        f"Aggiornate {updates_count} US nel database.")

                # Opzionale: aggiorna la vista
                self.fill_fields()

            except Exception as e:
                QMessageBox.critical(self, "Errore",
                                     f"Errore durante l'aggiornamento del database: {str(e)}")
            # Versione ottimizzata
            # OL = Order_layer_v2(self.DB_MANAGER, sito, area, use_graphviz=True)
            #
            # order_layer_dict = OL.main_order_layer()
            #
            # # Opzionale: esporta visualizzazione Graphviz
            # #graph_file = OL.export_graphviz_visualization()
            # #if graph_file:
            #      #QMessageBox.information(self, "Matrix Esportato", f"Matrix salvato come: {graph_file}")
            #
            # try:
            #     # Batch update ottimizzato
            #     updates_batch = []
            #
            #     for k, v in order_layer_dict.items():
            #         order_number = k
            #         us_v = v
            #
            #         for sing_us in us_v:
            #             search_dict = {
            #                 'sito': "'" + str(sito) + "'",
            #                 'area': "'" + str(area) + "'",
            #                 'us': int(sing_us)
            #             }
            #
            #             try:
            #                 records = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
            #                 if records:
            #                     updates_batch.append({
            #                         'id': int(records[0].id_us),
            #                         'order_layer': order_number
            #                     })
            #             except Exception as e:
            #                 QMessageBox.warning(self, 'Attenzione', str(e), QMessageBox.Ok)
            #
            #     # Esegue tutti gli aggiornamenti in una transazione
            #     if updates_batch:
            #         self.DB_MANAGER.begin_transaction()
            #         try:
            #             for update in updates_batch:
            #                 self.DB_MANAGER.update(
            #                     self.MAPPER_TABLE_CLASS,
            #                     self.ID_TABLE,
            #                     [update['id']],
            #                     ['order_layer'],
            #                     [update['order_layer']]
            #                 )
            #             self.DB_MANAGER.commit_transaction()
            #
            #             # Aggiorna la UI una sola volta
            #             self.on_pushButton_view_all_pressed()
            #
            #         except Exception as e:
            #             self.DB_MANAGER.rollback_transaction()
            #             QMessageBox.warning(self, 'Errore', f"Errore durante l'aggiornamento batch: {e}",
            #                                 QMessageBox.Ok)
            #
            # except Exception as e:
            #     QMessageBox.warning(self, 'Errore', str(e), QMessageBox.Ok)
            try:
                #QMessageBox.warning(self, "Messaggio", f"order{order_number} - us{us_v}", QMessageBox.Ok)
                if self.L=='it':
                    filename_tipo_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'tipo_rapporti_mancanti.txt')
                    filename_nr_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'nr_rapporti_mancanti.txt')
                    filename_paradosso_rapporti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'paradosso_rapporti.txt')
                    filename_us_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'us_mancanti.txt')
                elif self.L=='de':
                    filename_tipo_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'type_missing_relationships.txt')
                    filename_nr_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'nr_missing relashionships.txt')
                    filename_paradosso_rapporti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'relashionships_paradox.txt')
                    filename_us_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'su_missing.txt')
                else:
                    filename_tipo_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'type_missing_relationships.txt')
                    filename_nr_rapporti_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'nr_missing relashionships.txt')
                    filename_paradosso_rapporti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'relashionships_paradox.txt')
                    filename_us_mancanti = '{}{}{}'.format(self.REPORT_PATH, os.sep, 'su_missing.txt')
                self.testing(filename_tipo_rapporti_mancanti, str(msg_tipo_rapp))
                self.testing(filename_nr_rapporti_mancanti, str(msg_nr_rapp))
                self.testing(filename_paradosso_rapporti, str(msg_paradx_rapp))
                self.testing(filename_us_mancanti, str(msg_us_mancanti))
                if self.L=='it':
                    QMessageBox.warning(self, u'ATTENZIONE', u"Sistema di ordinamento Terminato", QMessageBox.Ok)
                elif self.L=='de':
                    QMessageBox.warning(self, u'ACHTUNG', "Ordnungssystem beendet", QMessageBox.Ok)
                else:
                    QMessageBox.warning(self, u'WARNING', "Sorting system Complete", QMessageBox.Ok)
            except Exception as e:
                print(f"{e}")

    def on_toolButtonPan_toggled(self):
        self.toolPan = QgsMapToolPan(self.mapPreview)
        self.mapPreview.setMapTool(self.toolPan)
    def on_pushButton_showSelectedFeatures_pressed(self):
        # field_position = self.pyQGIS.findFieldFrDict(self.ID_TABLE) #ricava la posizione del campo
        try:
            layer = self.iface.mapCanvas().currentLayer()
            fieldname = self.ID_TABLE
            if not layer:
                if self.L=='it':
                    QMessageBox.warning(self, 'ATTENZIONE', "Nessun elemento selezionato", QMessageBox.Ok)
                elif self.L=='de':
                    QMessageBox.warning(self, 'ACHTUNG', "keine Elemente ausgewählt", QMessageBox.Ok)
                else:
                    QMessageBox.warning(self, 'WARNING', "No items selected", QMessageBox.Ok)
            features_list = layer.selectedFeatures()
            field_position = ""
            for single in layer.getFeatures():
                field_position = single.fieldNameIndex(fieldname)
            id_list = []
            for feat in features_list:
                attr_list = feat.attributes()
                id_list.append(attr_list[field_position])
                # viene impostata la query per il database
            items, order_type = [self.ID_TABLE], "asc"
            self.empty_fields()
            self.DATA_LIST = []
            temp_data_list = self.DB_MANAGER.query_sort(id_list, items, order_type, self.MAPPER_TABLE_CLASS, self.ID_TABLE)
            for us in temp_data_list:
                self.DATA_LIST.append(us)
                # vengono riempiti i campi con i dati trovati
            self.fill_fields()
            self.BROWSE_STATUS = 'b'
            self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
            if type(self.REC_CORR) == "<type 'str'>":
                corr = 0
            else:
                corr = self.REC_CORR
            self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
            self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
            self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
        except Exception as e:
            QMessageBox.warning(self, 'ATTENZIONE', str(e), QMessageBox.Ok)

    def on_pushButton_sort_pressed(self):
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.check_record_state() == 1:
            pass
        else:
            dlg = SortPanelMain(self)
            dlg.insertItems(self.SORT_ITEMS)
            dlg.exec_()
            items, order_type = dlg.ITEMS, dlg.TYPE_ORDER
            self.SORT_ITEMS_CONVERTED = []
            for i in items:
                # QMessageBox.warning(self, "Messaggio",i, QMessageBox.Ok)
                self.SORT_ITEMS_CONVERTED.append(
                    self.CONVERSION_DICT[str(i)])  # apportare la modifica nellle altre schede
            self.SORT_MODE = order_type
            self.empty_fields()
            id_list = []
            for i in self.DATA_LIST:
                id_list.append(eval("i." + self.ID_TABLE))
            self.DATA_LIST = []
            temp_data_list = self.DB_MANAGER.query_sort(id_list, self.SORT_ITEMS_CONVERTED, self.SORT_MODE,
                                                        self.MAPPER_TABLE_CLASS, self.ID_TABLE)
            for i in temp_data_list:
                self.DATA_LIST.append(i)
            self.BROWSE_STATUS = 'b'
            self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
            if type(self.REC_CORR) == "<type 'str'>":
                corr = 0
            else:
                corr = self.REC_CORR
            self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
            self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
            self.SORT_STATUS = "o"
            self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
            self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
            self.fill_fields()
    def on_toolButtonGis_toggled(self):
        if self.L=='it':
            if self.toolButtonGis.isChecked():
                QMessageBox.warning(self, "Messaggio",
                                    "Modalita' GIS attiva. Da ora le tue ricerche verranno visualizzate sul GIS",
                                    QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "Messaggio",
                                    "Modalita' GIS disattivata. Da ora le tue ricerche non verranno piu' visualizzate sul GIS",
                                    QMessageBox.Ok)
        elif self.L=='de':
            if self.toolButtonGis.isChecked():
                QMessageBox.warning(self, "Message",
                                    "Modalität' GIS aktiv. Von jetzt wird Deine Untersuchung mit Gis visualisiert",
                                    QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "Message",
                                    "Modalität' GIS deaktiviert. Von jetzt an wird deine Untersuchung nicht mehr mit Gis visualisiert",
                                    QMessageBox.Ok)
        else:
            if self.toolButtonGis.isChecked():
                QMessageBox.warning(self, "Message",
                                    "GIS mode active. From now on your searches will be displayed on the GIS",
                                    QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "Message",
                                    "GIS mode disabled. From now on, your searches will no longer be displayed on the GIS.",
                                    QMessageBox.Ok)
    def on_toolButtonPreview_toggled(self):
        if self.L=='it':
            if self.toolButtonPreview.isChecked():
                QMessageBox.warning(self, "Messaggio",
                                    "Modalita' Preview US attivata. Le piante delle US saranno visualizzate nella sezione Piante",
                                    QMessageBox.Ok)
                self.tabWidget.setCurrentIndex(10)  # Set the current tab to the map preview tab
                self.loadMapPreview()
            else:
                self.loadMapPreview(1)
        elif self.L=='de':
            if self.toolButtonPreview.isChecked():
                QMessageBox.warning(self, "Message",
                                    "Modalität' Preview der aktivierten SE. Die Plana der SE werden in der Auswahl der Plana visualisiert",
                                    QMessageBox.Ok)
                self.tabWidget.setCurrentIndex(10)  # Set the current tab to the map preview tab
                self.loadMapPreview()
            else:
                self.tabWidget.setCurrentIndex(0)
                self.loadMapPreview(1)
        else:
            if self.toolButtonPreview.isChecked():
                QMessageBox.warning(self, "Message",
                                    "Preview SU mode enabled. US plants will be displayed in the Plants section",
                                    QMessageBox.Ok)
                self.tabWidget.setCurrentIndex(10)  # Set the current tab to the map preview tab
                self.loadMapPreview()
            else:
                self.tabWidget.setCurrentIndex(0)
                self.loadMapPreview(1)

    def on_pushButton_addRaster_pressed(self):
        if self.toolButtonGis.isChecked():
            self.pyQGIS.addRasterLayer()
    def on_pushButton_new_rec_pressed(self):
        conn = Connection()

        sito_set= conn.sito_set()
        sito_set_str = sito_set['sito_set']
        if bool(self.DATA_LIST):
            if self.data_error_check() == 1:
                pass

        if self.BROWSE_STATUS != "n":
            if bool(self.comboBox_sito.currentText()) and self.comboBox_sito.currentText() == sito_set_str:
                # Call the functions directly without connecting them to signals
                self.charge_periodo_iniz_list()
                self.charge_periodo_fin_list()


                try:
                    self.comboBox_fas_iniz.currentIndexChanged.disconnect()
                    self.comboBox_per_iniz.currentIndexChanged.disconnect()
                except TypeError:
                    pass  # Ignore the error if no connections exist
                self.comboBox_fas_iniz.currentIndexChanged.connect(self.charge_datazione_list)

                try:
                    self.comboBox_fas_fin.currentIndexChanged.disconnect()
                    self.comboBox_per_fin.currentIndexChanged.disconnect()
                except TypeError:
                    pass  # Ignore the error if no connections exist
                self.comboBox_fas_fin.currentIndexChanged.connect(self.charge_datazione_list)

                self.BROWSE_STATUS = "n"

                self.setComboBoxEditable(["self.comboBox_area"], 1)
                self.setComboBoxEditable(["self.comboBox_unita_tipo"], 1)
                self.setComboBoxEnable(["self.comboBox_sito"], False)
                self.setComboBoxEnable(["self.comboBox_area"], True)
                self.setComboBoxEnable(["self.lineEdit_us"], True)
                self.setComboBoxEnable(["self.comboBox_unita_tipo"], True)
                self.SORT_STATUS = "n"
                self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.set_rec_counter('', '')
                self.label_sort.setText(self.SORTED_ITEMS["n"])
                self.empty_fields_nosite()

            else:
                self.BROWSE_STATUS = "n"
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])

                self.setComboBoxEditable(["self.comboBox_sito"], 1)
                self.setComboBoxEditable(["self.comboBox_area"], 1)
                self.setComboBoxEditable(["self.comboBox_unita_tipo"], 1)
                self.setComboBoxEnable(["self.comboBox_sito"], "True")
                self.setComboBoxEnable(["self.comboBox_area"], "True")
                self.setComboBoxEnable(["self.lineEdit_us"], "True")
                self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                self.SORT_STATUS = "n"
                self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])

                self.set_rec_counter('', '')
                self.label_sort.setText(self.SORTED_ITEMS["n"])
                self.empty_fields()

            self.enable_button(0)

    def save_rapp(self):


        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.BROWSE_STATUS == "b":
            if self.data_error_check() == 0:
                if self.records_equal_check() == 1:

                    if self.update_if(QMessageBox.Ok):
                        QMessageBox.Ok


    def on_pushButton_save_pressed(self):

        # Define messages for each language
        messages = {
            'it': {
                'change_warning': "Il record e' stato modificato. Vuoi salvare le modifiche?",
                'no_changes': "Non è stata realizzata alcuna modifica.",
                'data_entry_problem': "Problema nell'inserimento dati"
            },
            'de': {
                'change_warning': "Der Record wurde geändert. Möchtest du die Änderungen speichern?",
                'no_changes': "Keine Änderung vorgenommen",
                'data_entry_problem': "Problem der Dateneingabe"
            },
            'en': {
                'change_warning': "The record has been changed. Do you want to save the changes?",
                'no_changes': "No changes have been made",
                'data_entry_problem': "Problem with data entry"
            }
        }

        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.BROWSE_STATUS == "b":
            if self.data_error_check() == 0:
                if self.records_equal_check() == 1:
                    self.update_if(QMessageBox.warning(self, 'Attenzione',
                                                       messages[self.L]['change_warning'],
                                                       QMessageBox.Ok | QMessageBox.Cancel))
                    self.empty_fields()
                    self.SORT_STATUS = "n"
                    self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                    self.enable_button(1)
                    self.update_dating()
                    self.fill_fields(self.REC_CORR)
                else:
                    QMessageBox.warning(self, "ATTENZIONE", messages[self.L]['no_changes'], QMessageBox.Ok)
        else:
            if self.data_error_check() == 0:
                test_insert = self.insert_new_rec()
                if test_insert == 1:
                    self.empty_fields()
                    self.SORT_STATUS = "n"
                    self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                    self.charge_records_n()
                    self.charge_list()
                    self.set_sito()
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), len(self.DATA_LIST) - 1
                    self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
                    self.setComboBoxEditable(["self.comboBox_sito"], 1)
                    self.setComboBoxEditable(["self.comboBox_area"], 1)
                    self.setComboBoxEditable(["self.comboBox_unita_tipo"], 1)
                    self.setComboBoxEnable(["self.comboBox_sito"], "False")
                    self.setComboBoxEnable(["self.comboBox_area"], "False")
                    self.setComboBoxEnable(["self.lineEdit_us"], "False")
                    self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                    self.update_dating()
                    self.fill_fields(self.REC_CORR)
                    self.enable_button(1)
            else:
                QMessageBox.warning(self, "ATTENZIONE", messages[self.L]['data_entry_problem'], QMessageBox.Ok)
        self.update_dating()


    def apikey_gpt(self):
        #HOME = os.environ['PYARCHINIT_HOME']
        BIN = '{}{}{}'.format(self.HOME, os.sep, "bin")
        api_key = ""
        # Verifica se il file gpt_api_key.txt esiste
        path_key = os.path.join(BIN, 'gpt_api_key.txt')
        if os.path.exists(path_key):

            # Leggi l'API Key dal file
            with open(path_key, 'r') as f:
                api_key = f.read().strip()
                try:

                    return api_key

                except:
                    reply = QMessageBox.question(None, 'Warning', 'Apikey non valida' + '\n'
                                                 + 'Clicca ok per inserire la chiave',
                                                 QMessageBox.Ok | QMessageBox.Cancel)
                    if reply == QMessageBox.Ok:

                        api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey valida:')
                        if ok:
                            # Salva la nuova API Key nel file
                            with open(path_key, 'w') as f:
                                f.write(api_key)
                                f.close()
                            with open(path_key, 'r') as f:
                                api_key = f.read().strip()
                    else:
                        return api_key


        else:
            # Chiedi all'utente di inserire una nuova API Key
            api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey:')
            if ok:
                # Salva la nuova API Key nel file
                with open(path_key, 'w') as f:
                    f.write(api_key)
                    f.close()
                with open(path_key, 'r') as f:
                    api_key = f.read().strip()

        return api_key



    def on_pushButton_rapp_check_pressed(self):
        sito_check = str(self.comboBox_sito_rappcheck.currentText())
        area_check = str(self.comboBox_area_rappcheck.currentText())
        try:
            self.rapporti_stratigrafici_check(sito_check)
            self.def_strati_to_rapporti_stratigrafici_check(sito_check)  # SPERIMENTALE
        except AssertionError as e:
            QMessageBox.critical(self, "Error", f"An error occurred while performing the check: {str(e)}",
                                 QMessageBox.Ok)
            print(f"Error: {str(e)}")
        else:
            success_message = {
                'it': "Controllo Rapporti Stratigrafici e Definizione Stratigrafica a Rapporti Stratigrafici eseguito con successo",
                'de': "Prüfen der stratigraphischen Beziehung und Definition Stratigraphische zu Stratigraphische Berichte erfolgereich durchgeführt",
                'en': "Monitoring of stratigraphic relationships and Definition Stratigraphic to Stratigraphic Reports performed successfully"
            }
            QMessageBox.information(self, "Success", success_message.get(self.L, "Message"), QMessageBox.Ok)

    def on_pushButton_h_check_pressed(self):
        self.listWidget_rapp.clear()
        sito_check = str(self.comboBox_sito.currentText())
        area_check = str(self.comboBox_area.currentText())
        models = ["gpt-4.1", "gpt-4o-mini"]
        try:
            self.rapporti_stratigrafici_check(sito_check)
            if self.checkBox_validate.isChecked():
                self.def_strati_to_rapporti_stratigrafici_check(sito_check)
                self.periodi_to_rapporti_stratigrafici_check(sito_check)
            self.automaticform_check(sito_check)

            # # Nuovo controllo per "Area non trovata"
            # us_area_non_trovata = []
            # for row in range(self.tableWidget_rapporti.rowCount()):
            #     area_value = self.tableWidget_rapporti.item(row, 2).text()
            #     if area_value == "" or None:
            #         us_value = self.tableWidget_rapporti.item(row, 1).text()
            #         us_area_non_trovata.append(us_value)
            #
            # if us_area_non_trovata:
            #     message = "Le seguenti US hanno 'Area non trovata' e potrebbero richiedere la creazione di una nuova scheda:\n"
            #     message += ", ".join(us_area_non_trovata)
            #     self.listWidget_rapp.addItem(message)
        except Exception as e:
            full_exception = traceback.format_exc()

            os.environ["OPENAI_API_KEY"] = self.apikey_gpt()
            combo = QComboBox()
            combo.addItems(models)
            selected_model, ok = QInputDialog.getItem(self, "Select Model", "Choose a model for GPT:", models, 0,
                                                      False)

            if ok and selected_model:
                client = OpenAI()

                response = client.chat.completions.create(
                    model=selected_model,
                    messages=[
                        {"role": "user",
                         "content": f"spiegami questo errore: {full_exception} e se necessario genera dei link utili per approfondire"}
                    ],
                    stream=True
                )

                combined_message = "GPT Response:\n "
                self.listWidget_rapp.addItem(combined_message)

                try:
                    end = ''

                    for chunk in response:
                        if chunk.choices[0].delta.content is not None:
                            # print(chunk.choices[0].delta.content, end="")
                            combined_message += chunk.choices[0].delta.content
                            combined_message += end
                            # Rendi i link cliccabili
                            # combined_message = re.sub(r'(https?://\S+)', r'<a href="\1">\1</a>', combined_message)

                            self.listWidget_rapp.takeItem(self.listWidget_rapp.count() - 1)
                            self.listWidget_rapp.addItem(combined_message)
                            # self.listWidget_ai.scrollToBottom()
                            QApplication.processEvents()
                except requests.exceptions.JSONDecodeError as e:
                    print("Error decoding JSON response:", e)

                    self.listWidget_rapp.addItem(e)

            elif not ok:
                self.listWidget_rapp.addItem("Model selection was canceled.")

        else:

            success_message = {
                'it': "Controllo dei Rapporti Stratigrafici, Definizione Stratigrafica a Rapporti Stratigrafici, Periodi a Rapporti Stratigrafici e Automaticform eseguito con successo",
                'de': "Prüfen der stratigraphischen Beziehung, Definition Stratigraphische zu Stratigraphische Berichte, Perioden zu Stratigraphische Berichte und Automaticform erfolgereich durchgeführt",
                'en': "Monitoring of Stratigraphic Relationships, Definition Stratigraphic to Stratigraphic Reports, Periods to Stratigraphic Reports and Automaticform performed successfully"
            }
            self.listWidget_rapp.addItem(success_message.get(self.L, "Message"))



    def data_error_check(self):
        test = 0
        EC = Error_check()
        if self.L=='it':
            if EC.data_is_empty(str(self.comboBox_sito.currentText())) == 0:
                QMessageBox.warning(self, "ATTENZIONE", "Campo Sito. \n Il campo non deve essere vuoto", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_area.currentText())) == 0:
                QMessageBox.warning(self, "ATTENZIONE", "Campo Area. \n Il campo non deve essere vuoto", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.lineEdit_us.text())) == 0:
                QMessageBox.warning(self, "ATTENZIONE", "Campo US. \n Il campo non deve essere vuoto", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_unita_tipo.currentText())) == 0:
                QMessageBox.warning(self, "ATTENZIONE", "Campo Tipo US/USM. \n Il campo non deve essere vuoto",
                                    QMessageBox.Ok)
                test = 1
            """controllo campi numerici"""
            area = self.comboBox_area.currentText()
            us = self.lineEdit_us.text()
            if area != "":
                if EC.data_is_int(area) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Area. \n Il valore deve essere di tipo numerico",
                                        QMessageBox.Ok)
                    test = 1
            if us != "" :
                if EC.data_is_int(us) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo US. \n Il valore deve essere di tipo numerico",
                                        QMessageBox.Ok)
                    test = 1
            """controllo campi numerici float"""
            #TAB USM
            spessore_usm = self.lineEdit_spessore_usm.text()
            qmin_usm = self.lineEdit_qmin_usm.text()
            qmax_usm = self.lineEdit_qmax_usm.text()
            lunghezza_usm = self.lineEdit_lunghezza_usm.text()
            altezza_usm = self.lineEdit_altezza_usm.text()
            #TAB MISURE
            quota_abs = self.lineEdit_quota_abs.text()
            quota_relativa = self.lineEdit_quota_relativa.text()
            quota_max_abs = self.lineEdit_quota_max_abs.text()
            quota_max_rel = self.lineEdit_quota_max_rel.text()
            quota_min_abs = self.lineEdit_quota_min_abs.text()
            quota_min_rel = self.lineEdit_quota_min_rel.text()
            larghezza_media = self.lineEdit_larghezza_media.text()
            lunghezza_max = self.lineEdit_lunghezza_max.text()
            profondita_min = self.lineEdit_profondita_min.text()
            profondita_max = self.lineEdit_profondita_max.text()
            altezza_max = self.lineEdit_altezza_max.text()
            altezza_min = self.lineEdit_altezza_min.text()
            if spessore_usm != "":
                if EC.data_is_float(spessore_usm) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo USM-Spessore USM. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if qmin_usm != "":
                if EC.data_is_float(qmin_usm) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo USM 3-Quota minima USM. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if qmax_usm != "":
                if EC.data_is_float(qmax_usm) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo USM 3-Quota max USM. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if lunghezza_usm != "":
                if EC.data_is_float(lunghezza_usm) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo USM 3-Lunghezza USM. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if altezza_usm != "":
                if EC.data_is_float(altezza_usm) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo USM 3-Altezza USM. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            ###########################
            if quota_abs != "":
                if EC.data_is_float(quota_abs) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Assoluta. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if quota_relativa != "":
                if EC.data_is_float(quota_relativa) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Relativa. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if quota_max_abs != "":
                if EC.data_is_float(quota_max_abs) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Massima Assoluta. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if quota_max_rel != "":
                if EC.data_is_float(quota_max_rel) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Massima Relativa. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if quota_min_abs != "":
                if EC.data_is_float(quota_min_abs) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Minima Assoluta. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if quota_min_rel != "":
                if EC.data_is_float(quota_min_rel) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Quota Minima Relativa. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if larghezza_media != "":
                if EC.data_is_float(larghezza_media) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Larghezza Media. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if lunghezza_max != "":
                if EC.data_is_float(lunghezza_max) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Lunghezza Massima. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if profondita_min != "":
                if EC.data_is_float(profondita_min) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Profondità Minima. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if profondita_max != "":
                if EC.data_is_float(profondita_max) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Profondità Massima. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if altezza_max != "":
                if EC.data_is_float(altezza_max) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Spessore. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            if altezza_min != "":
                if EC.data_is_float(altezza_min) == 0:
                    QMessageBox.warning(self, "ATTENZIONE", "Campo Spessore Minima. \n Il valore deve essere di tipo numerico. \n (Sono stati inserite lettere, virgole, accenti o caratteri non numerici.",
                                        QMessageBox.Ok)
                    test = 1
            """controllo lunghezza campo alfanumerico"""
            attivita = self.lineEdit_attivita.text()
            colore = self.comboBox_colore.currentText()
            anno_scavo = self.lineEdit_anno.text()
            formazione = self.comboBox_formazione.currentText()
            stato_conservazione = self.comboBox_conservazione.currentText()
            colore = self.comboBox_colore.currentText()
            consistenza = self.comboBox_consistenza.currentText()
            struttura = self.comboBox_struttura.currentText()
            cont_per = self.lineEdit_codice_periodo.text()
            d_interpretativa=self.comboBox_def_intepret.currentText()
            d_stratigrafica=self.comboBox_def_strat.currentText()
            if attivita != "":
                if EC.data_lenght(attivita, 3) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Attivita. \n Il valore non deve superare i 4 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
            # if anno_scavo != "":
            # if EC.data_lenght(anno_scavo,3) == 0:
            #       QMessageBox.warning(self, "ATTENZIONE", "Campo Anno. \n immettere una sola data (es. 2014)",  QMessageBox.Ok)
            #       test = 1

            if formazione != "":
                if EC.data_lenght(d_interpretativa, 254) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo definizione interpreta. \n Il valore non deve superare i 255 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1

            if formazione != "":
                if EC.data_lenght(d_stratigrafica, 254) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo definizione stratigrafica. \n Il valore non deve superare i 255 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1

            if formazione != "":
                if EC.data_lenght(formazione, 19) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Formazione. \n Il valore non deve superare i 20 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
            if stato_conservazione != "":
                if EC.data_lenght(stato_conservazione, 19) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Conservazione. \n Il valore non deve superare i 20 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
            if colore != "":
                if EC.data_lenght(colore, 19) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Colore. \n Il valore non deve superare i 20 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
            if consistenza != "":
                if EC.data_lenght(consistenza, 19) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Consistenza. \n Il valore non deve superare i 20 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
            if struttura != "":
                if EC.data_lenght(struttura, 29) == 0:
                    QMessageBox.warning(self, "ATTENZIONE",
                                        "Campo Struttura. \n Il valore non deve superare i 30 caratteri alfanumerici",
                                        QMessageBox.Ok)
                    test = 1
        elif self.L=='de':
            if EC.data_is_empty(str(self.comboBox_sito.currentText())) == 0:
                QMessageBox.warning(self, "ACHTUNG", " Feld Ausgrabungstätte. \n Das Feld darf nicht leer sein", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_area.currentText())) == 0:
                QMessageBox.warning(self, "ACHTUNG", "Feld Areal. \n Das Feld darf nicht leer sein", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.lineEdit_us.text())) == 0:
                QMessageBox.warning(self, "ACHTUNG", "Feld SE. \n Das Feld darf nicht leer sein", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_unita_tipo.currentText())) == 0:
                QMessageBox.warning(self, "ACHTUNG", "Feld SE/MSE Typ. \n Das Feld darf nicht leer sein",   QMessageBox.Ok)
                test = 1
            area = self.comboBox_area.currentText()
            us = self.lineEdit_us.text()
            attivita = self.lineEdit_attivita.text()
            colore = self.comboBox_colore.currentText()
            anno_scavo = self.lineEdit_anno.text()
            formazione = self.comboBox_formazione.currentText()
            stato_conservazione = self.comboBox_conservazione.currentText()
            colore = self.comboBox_colore.currentText()
            consistenza = self.comboBox_consistenza.currentText()
            struttura = self.comboBox_struttura.currentText()
            cont_per = self.lineEdit_codice_periodo.text()
            if area != "":
                if EC.data_is_int(area) == 0:
                    QMessageBox.warning(self, "ACHTUNG", "Feld Areal. \n Der Wert muss numerisch eingegeben werden",
                                        QMessageBox.Ok)
                    test = 1
            if us != "":
                if EC.data_is_int(us) == 0:
                    QMessageBox.warning(self, "ACHTUNG", "Feld SE. \n Der Wert muss numerisch eingegeben werden",
                                        QMessageBox.Ok)
                    test = 1
            if attivita != "":
                if EC.data_lenght(attivita, 3) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld aktiviert. \n Der Wert darf nicht mehr als 4 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
                    # if anno_scavo != "":
            # if EC.data_lenght(anno_scavo,3) == 0:
            #       QMessageBox.warning(self, "ATTENZIONE", "Campo Anno. \n immettere una sola data (es. 2014)",  QMessageBox.Ok)
            #       test = 1
            if formazione != "":
                if EC.data_lenght(formazione, 19) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld Bodenart. \n Der Wert darf nicht mehr als 20 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
            if stato_conservazione != "":
                if EC.data_lenght(stato_conservazione, 19) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld Erhaltungszustand.  Der Wert darf nicht mehr als 20 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
            if colore != "":
                if EC.data_lenght(colore, 19) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld Farbe. \n Der Wert darf nicht mehr als 20 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
            if consistenza != "":
                if EC.data_lenght(consistenza, 19) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld Konsistenz. \n Der Wert darf nicht mehr als 20 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
            if struttura != "":
                if EC.data_lenght(struttura, 29) == 0:
                    QMessageBox.warning(self, "ACHTUNG",
                                        "Feld Struktur. \n Der Wert darf nicht mehr als 30 alphanumerische Zeichen enthalten",
                                        QMessageBox.Ok)
                    test = 1
        else:
            if EC.data_is_empty(str(self.comboBox_sito.currentText())) == 0:
                QMessageBox.warning(self, "WARNING", "Site Field. \n The field must not be empty", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_area.currentText())) == 0:
                QMessageBox.warning(self, "WARNING", "Area Field. \n The field must not be empty", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.lineEdit_us.text())) == 0:
                QMessageBox.warning(self, "WARNING", "SU Field. \n The field must not be empty", QMessageBox.Ok)
                test = 1
            if EC.data_is_empty(str(self.comboBox_unita_tipo.currentText())) == 0:
                QMessageBox.warning(self, "WARNING", "SU-WSU Field. \n The field must not be empty",
                                    QMessageBox.Ok)
                test = 1
            area = self.comboBox_area.currentText()
            us = self.lineEdit_us.text()
            attivita = self.lineEdit_attivita.text()
            colore = self.comboBox_colore.currentText()
            anno_scavo = self.lineEdit_anno.text()
            formazione = self.comboBox_formazione.currentText()
            stato_conservazione = self.comboBox_conservazione.currentText()
            colore = self.comboBox_colore.currentText()
            consistenza = self.comboBox_consistenza.currentText()
            struttura = self.comboBox_struttura.currentText()
            cont_per = self.lineEdit_codice_periodo.text()
            if area != "":
                if EC.data_is_int(area) == 0:
                    QMessageBox.warning(self, "WARNING", "Area Field. \n The value must be numerical",
                                        QMessageBox.Ok)
                    test = 1
            if us != "":
                if EC.data_is_int(us) == 0:
                    QMessageBox.warning(self, "WARNING", "SU Field. \n The value must be numerical",
                                        QMessageBox.Ok)
                    test = 1
            if attivita != "":
                if EC.data_lenght(attivita, 3) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Activity Field. \n The value must not exceed 4 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
                    # if anno_scavo != "":
            # if EC.data_lenght(anno_scavo,3) == 0:
            #       QMessageBox.warning(self, "ATTENZIONE", "Campo Anno. \n immettere una sola data (es. 2014)",  QMessageBox.Ok)
            #       test = 1
            if formazione != "":
                if EC.data_lenght(formazione, 19) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Formation Field. \n The value must not exceed 20 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
            if stato_conservazione != "":
                if EC.data_lenght(stato_conservazione, 19) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Conservation Field. \n The value must not exceed 20 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
            if colore != "":
                if EC.data_lenght(colore, 19) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Color Field. \n The value must not exceed 20 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
            if consistenza != "":
                if EC.data_lenght(consistenza, 19) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Texture Field. \n The value must not exceed 20 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
            if struttura != "":
                if EC.data_lenght(struttura, 29) == 0:
                    QMessageBox.warning(self, "WARNING",
                                        "Structure Field. \n The value must not exceed 20 alphanumeric characters",
                                        QMessageBox.Ok)
                    test = 1
                # if cont_per != "":
                #   if EC.data_lenght(cont_per,199) == 0:
                #       QMessageBox.warning(self, "ATTENZIONE", "Campo codice periodo. \n Il valore non deve superare i 200 caratteri numerici",  QMessageBox.Ok)
                #       test = 1
                # PERIODIZZAZIONE CHECK
                # periodo iniz compilato e fase vuota  il blocco deve essere utilizzato meglio a partire dai signals
        """
        if self.comboBox_per_iniz.currentText() != "" and self.comboBox_fas_iniz.currentText() == "":
            QMessageBox.warning(self, "ATTENZIONE", "Campo Fase iniziale \n Specificare la Fase iniziale oltre al Periodo",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_fin.currentText()  != "" and self.comboBox_fas_fin.currentText() == "":
            QMessageBox.warning(self, "ATTENZIONE", "Campo Fase finale \n Specificare la Fase finale oltre al Periodo",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_iniz.currentText()  == "" and self.comboBox_fas_iniz.currentText() != "":
            QMessageBox.warning(self, "ATTENZIONE", "Campo Periodo iniziale \n Specificare un Periodo iniziale oltre alla Fase",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_fin.currentText()  == "" and self.comboBox_fas_fin.currentText() != "":
            QMessageBox.warning(self, "ATTENZIONE", "Campo Periodo iniziale \n Specificare un Periodo finale oltre alla Fase",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_fin.currentText()  != "" and self.comboBox_fas_fin.currentText() != "" and self.comboBox_per_iniz.currentText()  == "" and self.comboBox_fas_iniz.currentText() == "":
            QMessageBox.warning(self, "ATTENZIONE", "Campi Periodo e Fase iniziali \n Specificare un Periodo e Fase iniziali se si vuole inserire un Periodo e Fase finali",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_fin.currentText()  != "" and self.comboBox_fas_fin.currentText() != "" and self.comboBox_per_iniz.currentText()  == "" and self.comboBox_fas_iniz.currentText() == "":
            QMessageBox.warning(self, "ATTENZIONE", "Campi Periodo e Fase iniziali \n Specificare un Periodo e Fase iniziali se si vuole inserire un Periodo e Fase finali",  QMessageBox.Ok)
            test = 1
        if self.comboBox_per_iniz.currentText()  != "" and self.comboBox_fas_iniz.currentText() != "":
            search_dict = {
            'sito'  : "'"+str(self.comboBox_sito.currentText())+"'",
            'periodo'  : "'"+str(self.comboBox_per_iniz.currentText())+"'",
            }
            if  bool(self.DB_MANAGER.query_bool(search_dict, 'PERIODIZZAZIONE')) == False:
                QMessageBox.warning(self, "ATTENZIONE", "Campi Periodo e Fase iniziali \n E' stata inserita una periodizzazione inesistente",  QMessageBox.Ok)
                test = 1
        """
        return test

    def automaticform_check(self, sito_check):

        search_dict = {'sito': "'" + str(sito_check) + "'"}
        records = self.DB_MANAGER.query_bool(search_dict,
                                             self.MAPPER_TABLE_CLASS)  # carica tutti i dati di uno scavo ordinati per numero di US
        if self.L=='it':
            report_rapporti3 = 'Report controllo e conteggio delle Schede create automatcamente - Sito: %s \n' % (
                sito_check)
        elif self.L=='de':
            report_rapporti3 = 'Kontrollbericht Definition Stratigraphische zu Stratigraphische Berichte - Ausgrabungsstätte: %s \n' % (
                sito_check)
        else:
            report_rapporti3 = 'Control and count of forms automatically created - Site: %s \n' % (
                sito_check)
        for rec in range(len(records)):
            sito = "'" + str(records[rec].sito) + "'"
            area = "'" + str(records[rec].area) + "'"
            us = int(records[rec].us)
            def_stratigrafica = '"' + str(records[rec].d_stratigrafica) + '"'
            #rapporti = records[rec].rapporti  # caricati i rapporti nella variabile
            #rapporti = eval(def_stratigrafica)
            #for sing_rapp in range(len(records)):  # itera sulla serie di rapporti
            report3 = ""

            if self.L=='it':
                if def_stratigrafica.find('SCHEDA CREATA IN AUTOMATICO')  >=0:


                    report3 = 'Sito: %s, Area: %s, US: %d - %s. Da rivedere ' % (
                        sito, area, int(us), def_stratigrafica)
            else:
                if def_stratigrafica.find('FORM MADE AUTOMATIC') >= 0:


                    report3 = 'Sito: %s, Area: %s, US: %d - %s. Review it ' % (
                        sito, area, int(us), def_stratigrafica)
            if report3 != "":
                report_rapporti3 = report_rapporti3 + report3 + '\n'
                # self.listWidget_rapp.item(0).setForeground(QtCore.Qt.blue)
                # self.listWidget_rapp.item(1).setForeground(QtCore.Qt.blue)
                #self.listWidget_rapp.clear()
        self.listWidget_rapp.addItem(report_rapporti3)
        HOME = os.environ['PYARCHINIT_HOME']
        report_path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Report_folder")
        if self.L=='it':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_schedeautomatiche.txt')
        elif self.L=='de':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_def_strat_to_SE relation.txt')
        elif self.L=='en':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_strat_def_to_SU relation.txt')
        f = open(filename, "w")
        f.write(report_rapporti3)
        f.close()

    def rapporti_stratigrafici_check(self, sito_check):
        conn = Connection()
        conn_str = conn.conn_str()
        test_conn = conn_str.find('sqlite')

        us_inesistenti = []
        rapporti_mancanti = []
        aree_vuote = []
        self.listWidget_rapp.clear()

        # Dictionary for reciprocal relationship types
        conversion_dict = {
            'Covers': 'Covered by', 'Covered by': 'Covers',
            'Fills': 'Filled by', 'Filled by': 'Fills',
            'Cuts': 'Cut by', 'Cut by': 'Cuts',
            'Abuts': 'Supports', 'Supports': 'Abuts',
            'Connected to': 'Connected to', 'Same as': 'Same as',
            'Copre': 'Coperto da', 'Coperto da': 'Copre',
            'Riempie': 'Riempito da', 'Riempito da': 'Riempie',
            'Taglia': 'Tagliato da', 'Tagliato da': 'Taglia',
            'Si appoggia a': 'Gli si appoggia', 'Gli si appoggia': 'Si appoggia a',
            'Si lega a': 'Si lega a', 'Uguale a': 'Uguale a',
            'Liegt über': 'Liegt unter', 'Liegt unter': 'Liegt über',
            'Schneidet': 'Wird geschnitten', 'Wird geschnitten': 'Schneidet',
            'Verfüllt': 'Wird verfüllt durch', 'Wird verfüllt durch': 'Verfüllt',
            'Stützt sich auf': 'Wird gestüzt von', 'Wird gestüzt von': 'Stützt sich auf',
            'Bindet an': 'Bindet an', 'Entspricht': 'Entspricht',
            '>>': '<<', '<<': '>>', '<': '>', '>': '<', '<->': '<->'
        }

        # Get all records for the site
        search_dict = {'sito': "'" + str(sito_check) + "'"}
        records = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)

        # Initialize report header based on language
        if self.L == 'it':
            self.report_rapporti = f'Report controllo Rapporti Stratigrafici - Sito: {sito_check}\n'
        elif self.L == 'de':
            self.report_rapporti = f'Kontrollbericht Stratigraphische Beziehungen - Ausgrabungsstätte: {sito_check}\n'
        else:
            self.report_rapporti = f'Control report Stratigraphic relationships - Site: {sito_check}\n'

        # Process each record
        for rec in records:
            sito = "'" + str(rec.sito) + "'"
            area = "'" + str(rec.area) + "'"
            us = int(rec.us)

            try:
                rapporti = eval(rec.rapporti)
            except Exception as e:
                self.report_rapporti += f"\nErrore nel formato dei rapporti per US {us}: {str(e)}\n"
                continue

            # Check for empty areas in relationships
            for area_vuota in rapporti:
                if len(area_vuota) <= 2:
                    continue
                elif len(area_vuota) > 2:
                    if not area_vuota[2]:
                        area_vuota[2] = 9999
                        aree_vuote.append(
                            f"Nella scheda US: {us}, il Rapporto: {area_vuota[0]} US: {area_vuota[1]}, l'Area è vuota")

            # Check each relationship
            for sing_rapp in rapporti:
                if len(sing_rapp) < 4:
                    continue

                # Get the reciprocal relationship type
                rapp_converted = conversion_dict.get(sing_rapp[0], "")
                if not rapp_converted:
                    continue

                # Prepare search parameters for the related US
                try:
                    related_us = int(sing_rapp[1])
                    related_area = sing_rapp[2]

                    # Handle "Area non trovata" case
                    if related_area == "Area non trovata":
                        related_area = 0

                    # Build search dictionary based on connection type
                    if test_conn == 0:
                        serch_dict_rapp = {
                            'sito': sito,
                            'area': related_area,
                            'us': related_us
                        }
                    else:
                        serch_dict_rapp = {
                            'sito': sito,
                            'area': "'" + str(related_area) + "'",
                            'us': related_us
                        }

                    # Query for the related US
                    us_rapp = self.DB_MANAGER.query_bool(serch_dict_rapp, self.MAPPER_TABLE_CLASS)

                    # If related US doesn't exist
                    if not bool(us_rapp):
                        if related_area == 0:
                            us_inesistenti.append(f"US: {related_us}: US inesistente")
                        else:
                            us_inesistenti.append(
                                f"Sito: {sito.strip(' ')}, Area: {related_area}, US: {related_us}: Scheda US non esistente")
                    else:
                        # Check if reciprocal relationship exists
                        try:
                            rapporti_check = eval(us_rapp[0].rapporti)

                            # Create the expected reciprocal relationship
                            expected_reciprocal = [
                                rapp_converted,  # Reciprocal relationship type
                                str(us),  # Current US number
                                str(rec.area),  # Current US area
                                str(rec.sito)  # Current US site
                            ]

                            # Check if the reciprocal relationship exists
                            reciprocal_found = False
                            for check_rapp in rapporti_check:
                                if (len(check_rapp) >= 4 and
                                        check_rapp[0] == rapp_converted and
                                        str(check_rapp[1]) == str(us) and
                                        str(check_rapp[2]) == str(rec.area)):
                                    reciprocal_found = True
                                    break

                            if not reciprocal_found:
                                report_template = {
                                    'it': 'Sito: {}, Area: {}, US: {} {} US: {} Area: {}: Rapporto reciproco non trovato',
                                    'de': 'Ausgrabungsstätte: {}, Areal: {}, SE: {} {} SE: {} Area: {}: Gegenseitiger Bericht nicht gefunden',
                                    'en': 'Site: {}, Area: {}, SU: {} {} SU: {} Area: {}: Reciprocal relationship not found'
                                }
                                rapporti_mancanti.append(report_template.get(self.L, report_template['en']).format(
                                    sito.strip("'"), area.strip("'"), us, sing_rapp[0], related_us, related_area))
                        except Exception as e:
                            rapporti_mancanti.append(
                                f"Errore nella verifica dei rapporti per US {us} con US {related_us}: {str(e)}")
                except Exception as e:
                    self.report_rapporti += f"\nErrore nell'elaborazione del rapporto {sing_rapp} per US {us}: {str(e)}\n"

        # Build the final report
        if aree_vuote:
            self.report_rapporti += f"\nCi sono {len(aree_vuote)} rapporti con area vuota. Eseguire Ctrl+U per aggiornare prima di procedere con il controllo.\n\nDettagli:\n" + "\n".join(
                aree_vuote) + "\n"

        if us_inesistenti:
            self.report_rapporti += "\nUS Inesistenti:\n" + "\n".join(us_inesistenti) + "\n"

        if rapporti_mancanti:
            self.report_rapporti += "\nRapporti Mancanti o Non Verificati:\n" + "\n".join(rapporti_mancanti) + "\n"

        # Display and save the report
        self.listWidget_rapp.addItem(self.report_rapporti)

        HOME = os.environ['PYARCHINIT_HOME']
        report_path = f"{HOME}{os.sep}pyarchinit_Report_folder"

        if self.L == 'it':
            filename = f"{report_path}{os.sep}log_rapporti_US.txt"
        elif self.L == 'de':
            filename = f"{report_path}{os.sep}log_SE.txt"
        else:
            filename = f"{report_path}{os.sep}log_SU_relations.txt"

        with open(filename, "w") as f:
            f.write(self.report_rapporti)

        return self.report_rapporti

    # def rapporti_stratigrafici_check(self, sito_check):
    #     #rapporti_check = None
    #     conn = Connection()
    #     conn_str = conn.conn_str()
    #     test_conn = conn_str.find('sqlite')
    #
    #
    #     us_inesistenti = []
    #     rapporti_mancanti = []
    #     aree_vuote = []
    #     aree_vuote_2 = []
    #     sing_rapp1 =''
    #     self.listWidget_rapp.clear()
    #     conversion_dict = {'Covers': 'Covered by',
    #                        'Covered by': 'Covers',
    #                        'Fills': 'Filled by',
    #                        'Filled by': 'Fills',
    #                        'Cuts': 'Cut by',
    #                        'Cut by': 'Cuts',
    #                        'Abuts': 'Supports',
    #                        'Supports': 'Abuts',
    #                        'Connected to': 'Connected to',
    #                        'Same as': 'Same as',
    #                        'Copre': 'Coperto da',
    #                        'Coperto da': 'Copre',
    #                        'Riempie': 'Riempito da',
    #                        'Riempito da': 'Riempie',
    #                        'Taglia': 'Tagliato da',
    #                        'Tagliato da': 'Taglia',
    #                        'Si appoggia a': 'Gli si appoggia',
    #                        'Gli si appoggia': 'Si appoggia a',
    #                        'Si lega a': 'Si lega a',
    #                        'Uguale a': 'Uguale a',
    #                        'Liegt über': 'Liegt unter',
    #                        'Liegt unter': 'Liegt über',
    #                        'Schneidet': 'Wird geschnitten',
    #                        'Wird geschnitten': 'Schneidet',
    #                        'Verfüllt': 'Wird verfüllt durch',
    #                        'Wird verfüllt durch': 'Verfüllt',
    #                        'Stützt sich auf': 'Wird gestüzt von',
    #                        'Wird gestüzt von': 'Stützt sich auf',
    #                        'Bindet an': 'Bindet an',
    #                        'Entspricht': 'Entspricht',
    #                        '>>': '<<',
    #                        '<<': '>>',
    #                        '<': '>',
    #                        '>': '<',
    #                        '<->': '<->'
    #                        }
    #     search_dict = {'sito': "'" + str(sito_check) + "'"}
    #     records = self.DB_MANAGER.query_bool(search_dict,
    #                                          self.MAPPER_TABLE_CLASS)  # carica tutti i dati di uno scavo ordinati per numero di US
    #     if self.L == 'it':
    #         self.report_rapporti = 'Report controllo Rapporti Stratigrafici - Sito: %s \n' % (sito_check)
    #     elif self.L == 'de':
    #         self.report_rapporti = 'Kontrollbericht Stratigraphische Beziehungen - Ausgrabungsstätte: %s \n' % (
    #             sito_check)
    #     else:
    #         self.report_rapporti = 'Control report Stratigraphic relationships - Site: %s \n' % (sito_check)
    #     count_0 = 0
    #     count_1 = 0
    #
    #     for rec in range(len(records)):
    #         sito = "'" + str(records[rec].sito) + "'"
    #         area = "'" + str(records[rec].area) + "'"
    #         us = int(records[rec].us)
    #         rapporti = records[rec].rapporti
    #         rapporti = eval(rapporti)
    #         report = ''
    #         for area_vuota in rapporti:
    #
    #             if len(area_vuota) <= 2:
    #                 aree_vuote_2.append('')
    #
    #
    #             elif len(area_vuota)>2:
    #                 if not area_vuota[2]:
    #
    #                     area_vuota[2] = 9999
    #
    #                     aree_vuote.append(f"Nella scheda US: {us}, il Rapporto: {area_vuota[0]} US: {area_vuota[1]}, l'Area è vuota")
    #
    #
    #
    #         for sing_rapp in rapporti:
    #             if str(us).find('0') >= 0 or str(us).find('1') >= 0:
    #
    #                 if len(sing_rapp) > 2:
    #                  # Verifica se sing_rapp[2] è 'Area non trovata', in tal caso impostalo a 0
    #                     if sing_rapp[2] == "Area non trovata":
    #                         sing_rapp[2] = 0
    #
    #                 if len(sing_rapp) == 4:
    #                     rapp_converted = conversion_dict[sing_rapp[0]]
    #                     #serch_dict_rapp = {'sito': sito, 'area': "'"+str(sing_rapp[2])+"'", 'us': int(sing_rapp[1])}
    #
    #                     if test_conn==0:
    #                         serch_dict_rapp = {'sito': sito, 'area': sing_rapp[2],
    #                                            'us': int(sing_rapp[1])}
    #
    #                         us_rapp = self.DB_MANAGER.query_bool(serch_dict_rapp, self.MAPPER_TABLE_CLASS)
    #                     else:
    #                         serch_dict_rapp = {'sito': sito, 'area': "'" + str(sing_rapp[2]) + "'",
    #                                            'us': int(sing_rapp[1])}
    #                         us_rapp = self.DB_MANAGER.query_bool(serch_dict_rapp, self.MAPPER_TABLE_CLASS)
    #                     try:
    #                         int(sing_rapp[1])
    #                     except ValueError:
    #                         raise TypeError(f"Expected an integer for sing_rapp[1], got {int(sing_rapp[1])} instead")
    #
    #                     if not bool(us_rapp):
    #
    #
    #
    #
    #                         if sing_rapp[2] == 0:
    #                             us_inesistenti.append(f"US: {sing_rapp[1]}: US inesistente")
    #                         # else:
    #                         #     report_template = {
    #                         #
    #                         #         'it': 'Sito: {}, Area: {}, US: {} {} US: {} Area: {}: Scheda US non esistente',
    #                         #
    #                         #         'de': 'Ausgrabungsstätte: {}, Areal: {}, SE: {} {} SE: {} Area: {}: SE formular nicht existent',
    #                         #
    #                         #         'en': 'Site: {}, Area: {}, SU: {} {} SU: {} Area: {}: SU form not-existent'
    #                         #
    #                         #     }
    #                         #
    #                         #     us_inesistenti.append(report_template.get(self.L, report_template['en']).format(
    #                         #
    #                         #         sito, area, us, sing_rapp[0], sing_rapp[1], sing_rapp[2]))
    #                     else:
    #                         try:
    #                             rapporti_check = eval(us_rapp[0].rapporti)
    #                             us_rapp_check = str(us)
    #                             area_rapp_check = area.strip("'")
    #                             sito_rapp_check = sito.strip("'")
    #
    #                             s = [rapp_converted, us_rapp_check, area_rapp_check, sito_rapp_check]
    #
    #                             if rapporti_check.count(s) != 1:
    #                                 report_template = {
    #                                     'it': 'Sito: {}, Area: {}, US: {} {} US: {} Area: {}: Rapporto non verificato',
    #                                     'de': 'Ausgrabungsstätte: {}, Areal: {}, SE: {} {} SE: {} Area: {}: nicht geprüfter Bericht',
    #                                     'en': 'Site: {}, Area: {}, SU: {} {} SU: {} Area: {}: relationships not verified'
    #                                 }
    #                                 rapporti_mancanti.append(report_template.get(self.L, report_template['en']).format(
    #                                     sito, area, us, sing_rapp[0], sing_rapp[1], sing_rapp[2]))
    #                         except Exception as e:
    #                             rapporti_mancanti.append(f"Errore nella verifica dei rapporti per US {us}: {str(e)}")
    #
    #                         # Aggiungi i risultati raggruppati al report
    #
    #
    #
    #     if aree_vuote_2:
    #
    #
    #         self.report_rapporti += (f"\nEseguire Ctrl+U per aggiornare prima di procedere con il controllo.\n\n"
    #                                  )
    #
    #     if aree_vuote:
    #
    #
    #         self.report_rapporti += f"\nCi sono {len(aree_vuote)} rapporti con area vuota . Eseguire Ctrl+U per aggiornare prima di procedere con il controllo.\n\nDettagli:\n" + "\n".join(
    #                 aree_vuote)
    #
    #
    #     else:
    #         if us_inesistenti:
    #             self.report_rapporti += "\nUS Inesistenti:\n" + "\n".join(us_inesistenti) + "\n"
    #         if rapporti_mancanti:
    #             self.report_rapporti += "\nRapporti Mancanti o Non Verificati:\n" + "\n".join(
    #                 rapporti_mancanti) + "\n"
    #
    #
    #
    #
    #     self.listWidget_rapp.addItem(self.report_rapporti)
    #     # Costruisci il messaggio finale includendo rapporti_check
    #     # final_message = f"Count of 0: {count_0}, Count of 1: {count_1}\nRapporti Check: {rapporti_check}"
    #     # QMessageBox.information(self, 'ok', final_message)
    #
    #     HOME = os.environ['PYARCHINIT_HOME']
    #     report_path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Report_folder")
    #     if self.L == 'it':
    #         filename = '{}{}{}'.format(report_path, os.sep, 'log_rapporti_US.txt')
    #     elif self.L == 'de':
    #         filename = '{}{}{}'.format(report_path, os.sep, 'log_SE.txt')
    #     else:
    #         filename = '{}{}{}'.format(report_path, os.sep, 'log_SU_relations.txt')
    #     f = open(filename, "w")
    #     f.write(self.report_rapporti)
    #     f.close()
    #     return self.report_rapporti

    def def_strati_to_rapporti_stratigrafici_check(self, sito_check):
        conversion_dict = {'Covers':'Covered by',
                           'Covered by': 'Covers',
                           'Fills': 'Filled by',
                           'Filled by':'Fills',
                           'Cuts': 'Cut by',
                           'Cut by': 'Cuts',
                           'Abuts': 'Supports',
                           'Supports': 'Abuts',
                           'Connected to': 'Connected to',
                           'Same as':'Same as',
                           'Copre':'Coperto da',
                           'Coperto da': 'Copre',
                           'Riempie': 'Riempito da',
                           'Riempito da' : 'Riempie',
                           'Taglia': 'Tagliato da',
                           'Tagliato da': 'Taglia',
                           'Si appoggia a': 'Gli si appoggia',
                           'Gli si appoggia': 'Si appoggia a',
                           'Si lega a': 'Si lega a',
                           'Uguale a':'Uguale a',
                           'Liegt über':'Liegt unter',
                           'Liegt unter':'Liegt über',
                           'Schneidet':'Wird geschnitten',
                           'Wird geschnitten':'Schneidet',
                           'Verfüllt':'Wird verfüllt durch',
                           'Wird verfüllt durch':'Verfüllt',
                           'Stützt sich auf':'Wird gestüzt von',
                           'Wird gestüzt von':'Stützt sich auf',
                           'Bindet an':'Bindet an',
                           'Entspricht':'Entspricht',
                           '>>':'<<',
                           '<<':'>>',
                           '<':'>',
                           '>':'<',
                           '<->':'<->'
                           }
        search_dict = {'sito': "'" + str(sito_check) + "'"}
        records = self.DB_MANAGER.query_bool(search_dict,
                                             self.MAPPER_TABLE_CLASS)  # carica tutti i dati di uno scavo ordinati per numero di US
        if self.L=='it':
            report_rapporti1 = 'Report controllo Definizione Stratigrafica a Rapporti Stratigrafici - Sito: %s \n' % (
                sito_check)
        elif self.L=='de':
            report_rapporti1 = 'Kontrollbericht Definition Stratigraphische zu Stratigraphische Berichte - Ausgrabungsstätte: %s  \n' % (
                sito_check)
        else:
            report_rapporti1 = 'Control report Definition Stratigraphic to Stratigraphic Reports - Site: %s \n' % (
                sito_check)
        for rec in range(len(records)):
            sito = "'" + str(records[rec].sito) + "'"
            area = int(records[rec].area)
            us = int(records[rec].us)
            def_stratigrafica = "'" + str(records[rec].d_stratigrafica) + "'"
            rapporti = records[rec].rapporti  # caricati i rapporti nella variabile
            rapporti = eval(rapporti)


            for sing_rapp in rapporti:  # itera sulla serie di rapporti
                report = ""


                if def_stratigrafica.find('Strato') >= 0:  # Paradosso strati che tagliano o si legano
                    if sing_rapp[0] == 'Si lega a':
                        report = 'Sito: %s, Area: %s, US: %d - %s: lo strato %s US: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if def_stratigrafica.find('Riempimento') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    if sing_rapp[0] == 'Taglia' or sing_rapp[0] == 'Si lega a'or sing_rapp[0] == 'Si appoggia a' or sing_rapp[0] == 'Gli si appoggia' or sing_rapp[0] == 'Taglia':
                        report = 'Sito: %s, Area: %s, US: %d - %s: lo strato %s US: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if def_stratigrafica.find('Taglio') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    if sing_rapp[0] == 'Riempie' or sing_rapp[0] == 'Si lega a' or sing_rapp[0] == 'Si appoggia a'  or sing_rapp[0] == 'Gli si appoggia':
                        report = 'Sito: %s, Area: %s, US: %d - %s: lo strato %s US: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if report != "":
                    report_rapporti1 = report_rapporti1 + report + '\n'
                #versione inglese
                elif def_stratigrafica.find('Stra') >= 0:  # Paradosso strati che tagliano o si legano
                    if sing_rapp[0] == 'Connected to':
                        report = 'Site: %s, Area: %s, SU: %d - %s: the stratum %s SU: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if def_stratigrafica.find('Fills') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    if sing_rapp[0] == 'Cuts' or sing_rapp[0] == 'Connected to':
                        report = 'Site: %s, Area: %s, SU: %d - %s: the startum %s SU: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                # if def_stratigrafica.find('Filling') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    # if sing_rapp[0] == 'Cuts' or sing_rapp[0] == 'Connected to':
                        # report = 'Site: %s, Area: %s, SU: %d - %s: the stratum %s SU: %d: ' % (
                            # sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if report != "":
                    report_rapporti1 = report_rapporti1 + report + '\n'
                #versione tedesca
                elif def_stratigrafica.find('Stratum') >= 0:  # Paradosso strati che tagliano o si legano
                    if sing_rapp[0] == 'Bindet an':
                        report = 'Sito: %s, Area: %s, SE: %d - %s: die startum %s US: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if def_stratigrafica.find('Verfullüng') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    if sing_rapp[0] == 'Schneidet' or sing_rapp[0] == 'Bindet an':
                        report = 'Sito: %s, Area: %s, SE: %d - %s: die stratum %s US: %d: ' % (
                            sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                # if def_stratigrafica.find('Verfullüng') >= 0:  # Paradosso riempimentiche tagliano o si legano
                    # if sing_rapp[0] == 'Schneidet' or sing_rapp[0] == 'Bindet an':
                        # report = 'Sito: %s, Area: %s, SE: %d - %s: die startum %s US: %d: ' % (
                            # sito, area, int(us), def_stratigrafica, sing_rapp[0], int(sing_rapp[1]))
                if report != "":
                    report_rapporti1 = report_rapporti1 + report + '\n'

        self.listWidget_rapp.addItem(report_rapporti1)

        HOME = os.environ['PYARCHINIT_HOME']
        report_path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Report_folder")
        if self.L=='it':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_def_strat_a_rapporti_US.txt')
        elif self.L=='de':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_def_strat_to_SE relation.txt')
        elif self.L=='en':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_strat_def_to_SU relation.txt')
        f = open(filename, "w")
        f.write(report_rapporti1)
        f.close()

    def concat(self,a, b):
        return eval(f"{a}{b}")

    def report_with_phrase(self, ut, us, sing_rapp, periodo_in, fase_in, sito, area):
        replaced_str = sing_rapp[4].replace('-', '')

        if replaced_str:
            replaced_str = int(replaced_str)
        else:
            replaced_str = 0 # or any default integer value

        if sing_rapp[0] in ['Si lega a', 'Uguale a', 'Same as', 'Connected to'] and replaced_str != periodo_in:
            return 'Sito: %s, Area: %s, %s: %d -  Il periodo e fase iniziale %s: deve essere: %s corrispondente con la %s : %d: ' % (
                sito, area, ut, us, str(periodo_in) + '-' + str(fase_in), sing_rapp[0], sing_rapp[2], int(sing_rapp[1]))

        return ""

    def periodi_to_rapporti_stratigrafici_check(self, sito_check):
        conversion_dict = {'Covers':'Covered by',
                           'Covered by': 'Covers',
                           'Fills': 'Filled by',
                           'Filled by':'Fills',
                           'Cuts': 'Cut by',
                           'Cut by': 'Cuts',
                           'Abuts': 'Supports',
                           'Supports': 'Abuts',
                           'Connected to': 'Connected to',
                           'Same as':'Same as',
                           'Copre':'Coperto da',
                           'Coperto da': 'Copre',
                           'Riempie': 'Riempito da',
                           'Riempito da' : 'Riempie',
                           'Taglia': 'Tagliato da',
                           'Tagliato da': 'Taglia',
                           'Si appoggia a': 'Gli si appoggia',
                           'Gli si appoggia': 'Si appoggia a',
                           'Si lega a': 'Si lega a',
                           'Uguale a':'Uguale a',
                           'Liegt über':'Liegt unter',
                           'Liegt unter':'Liegt über',
                           'Schneidet':'Wird geschnitten',
                           'Wird geschnitten':'Schneidet',
                           'Verfüllt':'Wird verfüllt durch',
                           'Wird verfüllt durch':'Verfüllt',
                           'Stützt sich auf':'Wird gestüzt von',
                           'Wird gestüzt von':'Stützt sich auf',
                           'Bindet an':'Bindet an',
                           'Entspricht':'Entspricht',
                           '>>':'<<',
                           '<<':'>>',
                           '<':'>',
                           '>':'<',
                           '<->':'<->'
                           }
        search_dict = {'sito': "'" + str(sito_check) + "'"}
        records = self.DB_MANAGER.query_bool(search_dict,
                                             self.MAPPER_TABLE_CLASS)  # carica tutti i dati di uno scavo ordinati per numero di US
        if self.L=='it':
            report_rapporti2 = 'Report controllo Periodi/Unità Tipo a Rapporti Stratigrafici - Sito: %s \n' % (
                sito_check)
        elif self.L=='de':
            report_rapporti2 = 'Kontrollbericht Periodization/Type Unit zu Stratigraphische Berichte - Ausgrabungsstätte: %s \n' % (
                sito_check)
        else:
            report_rapporti2 = 'Control report Periodization/Type Unit to Stratigraphic Reports - Site: %s \n' % (
                sito_check)
        for rec in range(len(records)):
            sito = "'" + str(records[rec].sito) + "'"
            area = "'" + str(records[rec].area) + "'"
            us = int(records[rec].us)
            periodo_in = str(records[rec].periodo_iniziale)
            fase_in = str(records[rec].fase_iniziale)
            periodo_fin =  str(records[rec].periodo_finale)
            fase_fin =  str(records[rec].fase_iniziale)
            ut=str(records[rec].unita_tipo)
            rapporti = records[rec].rapporti  # caricati i rapporti nella variabile
            rapporti = eval(rapporti)
            rapporti2 = records[rec].rapporti2  # caricati i rapporti nella variabile
            rapporti2 = eval(rapporti2)

            for sing_rapp in rapporti2:  # itera sulla serie di rapporti
                report = ""
                report2=''
                periodo_in_value_found = any(str(periodo_in).find(str(i)) for i in range(1, 50))

                if periodo_in_value_found:
                    report = self.report_with_phrase(ut, us, sing_rapp, periodo_in, fase_in, sito, area)
                    # if not bool(sing_rapp):
                        # report2:'la table widget deve essere riempita'
                    if sing_rapp[0] == 'Si lega a' and sing_rapp[2]=='US':

                        report2 = '%s : %d - %s : %d: devono essere USM' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Connected to' and sing_rapp[2]=='SU':

                        report2 = '%s : %d - %s : %d: should be WSU' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Uguale a' and sing_rapp[2]=='USM':

                        report2 = '%s : %d - %s : %d: devono essere US' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Same as' and sing_rapp[2]=='WSU':

                        report2 = '%s : %d - %s : %d: should be SU' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Si appoggia a' and sing_rapp[2]=='US':

                        report2 = '%s : %d - %s : %d: devono essere USM' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Gli si appoggia' and sing_rapp[2]=='US':

                        report2 = '%s : %d - %s : %d: devono essere USM' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Abuts' and sing_rapp[2]=='SU':

                        report2 = '%s : %d - %s : %d: should be WSU' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[0] == 'Supports' and sing_rapp[2]=='SU':

                        report2 = '%s : %d - %s : %d: should be WSU' % (
                            ut, int(us), sing_rapp[2], int(sing_rapp[1]))

                    if sing_rapp[4]!='-':

                        if sing_rapp[0] == 'Covers' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- should be Covered by %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Covered by' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- should be Covers %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Fills' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Should be Filled by %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Filled by' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Shuld be Fills %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Cuts' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Shuld be Cut by %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Cut by' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Shuld be Cuts %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Abuts' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Shuld be Supports %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Supports' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Shuld be Abuts %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Copre' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe essere Coperto da %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Coperto da' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe Coprire %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Riempie' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe essere Riempito da %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Riempito da' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe Riempire %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Taglia' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe essere Tagliato da %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Tagliato da' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe Tagliare %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Si appoggia a' and int(sing_rapp[4].replace('-',''))<self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe essere Gli si appoggia %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1],  str(periodo_in),str(fase_in))

                        if sing_rapp[0] == 'Gli si appoggia' and int(sing_rapp[4].replace('-',''))>self.concat(int(periodo_in),int(fase_in)):

                            report2 = '%s : %d : %s- Dovrebbe Si appoggia a %s : %s-%s' % (
                                ut, int(us),str(sing_rapp[4]), sing_rapp[1], str(periodo_in),str(fase_in))
                    if sing_rapp[4]=='-':
                        if self.L=='it':
                            report2 = 'Manca la peridizzazione in %s %s'% (sing_rapp[2], sing_rapp[1])
                        else:
                            report2 = 'Missing the periodization in %s %s'% (sing_rapp[2], sing_rapp[1])



                if report2 != "":
                    self.report_rapporti2 = report_rapporti2 + report + report2+'\n'

        self.listWidget_rapp.addItem(self.report_rapporti2)



        HOME = os.environ['PYARCHINIT_HOME']
        report_path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Report_folder")
        if self.L=='it':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_per_ut_a_rapporti_US.txt')
        elif self.L=='de':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_per_ut_to_SE relation.txt')
        elif self.L=='en':
            filename = '{}{}{}'.format(report_path, os.sep, 'log_per_ut_to_SU relation.txt')
        f = open(filename, "w")
        f.write(report_rapporti2)
        f.close()

    def insert_new_rec(self):
        # TableWidget
        #Rapporti
        rapporti = self.table2dict("self.tableWidget_rapporti")
        rapporti2 = self.table2dict("self.tableWidget_rapporti2")
        #Inclusi
        inclusi = self.table2dict("self.tableWidget_inclusi")
        #Campioni
        campioni = self.table2dict("self.tableWidget_campioni")
        #organici
        organici = self.table2dict("self.tableWidget_organici")
        #inorganici
        inorganici = self.table2dict("self.tableWidget_inorganici")
        #Documentazione
        documentazione = self.table2dict("self.tableWidget_documentazione")
        #Colore legante usm
        colore_legante_usm = self.table2dict("self.tableWidget_colore_legante_usm")
        #Inclusi leganti usm
        aggreg_legante_usm = self.table2dict("self.tableWidget_inclusi_leganti_usm")
        #Consistenza texture mat_usm
        consistenza_texture_mat_usm = self.table2dict("self.tableWidget_consistenza_texture_mat_usm")
        #inclusi_materiali_usm
        inclusi_materiali_usm = self.table2dict("self.tableWidget_inclusi_materiali_usm")
        #colore_materiale_usm
        colore_materiale_usm = self.table2dict("self.tableWidget_colore_materiale_usm")
        if self.lineEditOrderLayer.text() == "":
            order_layer = 0
        else:
            order_layer = int(self.lineEditOrderLayer.text())
        ##quota min usm
        if self.lineEdit_qmin_usm.text() == "":
            qmin_usm = None
        else:
            qmin_usm = float(self.lineEdit_qmin_usm.text())
        ##quota max usm
        if self.lineEdit_qmax_usm.text() == "":
            qmax_usm = None
        else:
            qmax_usm = float(self.lineEdit_qmax_usm.text())
        ##quota relativa
        if self.lineEdit_quota_relativa.text() == "":
            quota_relativa = None
        else:
            quota_relativa = float(self.lineEdit_quota_relativa.text())
        ##quota abs
        if self.lineEdit_quota_abs.text() == "":
            quota_abs = None
        else:
            quota_abs = float(self.lineEdit_quota_abs.text())
        ##lunghezza max
        if self.lineEdit_lunghezza_max.text() == "":
            lunghezza_max = None
        else:
            lunghezza_max = float(self.lineEdit_lunghezza_max.text())
        ##altezza max
        if self.lineEdit_altezza_max.text() == "":
            altezza_max = None
        else:
            altezza_max = float(self.lineEdit_altezza_max.text())
        ##altezza min
        if self.lineEdit_altezza_min.text() == "":
            altezza_min = None
        else:
            altezza_min = float(self.lineEdit_altezza_min.text())
        ##profondita max
        if self.lineEdit_profondita_max.text() == "":
            profondita_max = None
        else:
            profondita_max = float(self.lineEdit_profondita_max.text())
        ##profondita min
        if self.lineEdit_profondita_min.text() == "":
            profondita_min = None
        else:
            profondita_min = float(self.lineEdit_profondita_min.text())
        ##larghezza media
        if self.lineEdit_larghezza_media.text() == "":
            larghezza_media = None
        else:
            larghezza_media = float(self.lineEdit_larghezza_media.text())
        ##quota max abs
        if self.lineEdit_quota_max_abs.text() == "":
            quota_max_abs = None
        else:
            quota_max_abs = float(self.lineEdit_quota_max_abs.text())
        ##quota max relativa
        if self.lineEdit_quota_max_rel.text() == "":
            quota_max_rel = None
        else:
            quota_max_rel = float(self.lineEdit_quota_max_rel.text())
        ##quota min abs
        if self.lineEdit_quota_min_abs.text() == "":
            quota_min_abs = None
        else:
            quota_min_abs = float(self.lineEdit_quota_min_abs.text())
        ##quota min relativa
        if self.lineEdit_quota_min_rel.text() == "":
            quota_min_rel = None
        else:
            quota_min_rel = float(self.lineEdit_quota_min_rel.text())
        ##lunghezza usm
        if self.lineEdit_lunghezza_usm.text() == "":
            lunghezza_usm = None
        else:
            lunghezza_usm = float(self.lineEdit_lunghezza_usm.text())
        ##altezza usm
        if self.lineEdit_altezza_usm.text() == "":
            altezza_usm = None
        else:
            altezza_usm = float(self.lineEdit_altezza_usm.text())
        ##spessore usm
        if self.lineEdit_spessore_usm.text() == "":
            spessore_usm = None
        else:
            spessore_usm = float(self.lineEdit_spessore_usm.text())
        try:
            # data
            data = self.DB_MANAGER.insert_values(
                self.DB_MANAGER.max_num_id(self.MAPPER_TABLE_CLASS, self.ID_TABLE) + 1,
                str(self.comboBox_sito.currentText()),  # 1 - Sito
                str(self.comboBox_area.currentText()),  # 2 - Area
                int(self.lineEdit_us.text()),  # 3 - US
                str(self.comboBox_def_strat.currentText()),  # 4 - Definizione stratigrafica
                str(self.comboBox_def_intepret.currentText()),  # 5 - Definizione intepretata
                str(self.textEdit_descrizione.toPlainText()),  # 6 - descrizione
                str(self.textEdit_interpretazione.toPlainText()),  # 7 - interpretazione
                str(self.comboBox_per_iniz.currentText()),  # 8 - periodo iniziale
                str(self.comboBox_fas_iniz.currentText()),  # 9 - fase iniziale
                str(self.comboBox_per_fin.currentText()),  # 10 - periodo finale iniziale
                str(self.comboBox_fas_fin.currentText()),  # 11 - fase finale
                str(self.comboBox_scavato.currentText()),  # 12 - scavato
                str(self.lineEdit_attivita.text()),  # 13 - attivita
                str(self.lineEdit_anno.text()),  # 14 - anno scavo
                str(self.comboBox_metodo.currentText()),  # 15 - metodo
                str(inclusi),  # 16 - inclusi
                str(campioni),  # 17 - campioni
                str(rapporti),  # 18 - rapporti
                #str(organici), # componenti organici
                #str(inorganici), #componenti inorganici
                str(self.lineEdit_data_schedatura.text()),  # 19 - data schedatura
                str(self.comboBox_schedatore.currentText()),  # 20 - schedatore
                str(self.comboBox_formazione.currentText()),  # 21 - formazione
                str(self.comboBox_conservazione.currentText()),  # 22 - conservazione
                str(self.comboBox_colore.currentText()),  # 23 - colore
                str(self.comboBox_consistenza.currentText()),  # 24 - consistenza
                str(self.comboBox_struttura.currentText()),  # 25 - struttura
                str(self.lineEdit_codice_periodo.text()),  # 26 - continuita  periodo
                order_layer,  # 27 - order layer
                str(documentazione),  # 28 - documentazione
                str(self.comboBox_unita_tipo.currentText()),  # 29 us_tipo            NUOVI CAMPI NUOVI CAMPI
                str(self.comboBox_settore.currentText()),  # 30 settore
                str(self.lineEdit_quadrato.text()),  # 31 quadrato
                str(self.lineEdit_ambiente.text()),  # 32 ambiente
                str(self.lineEdit_saggio.text()),  # 33 saggio
                str(self.textEdit_elementi_datanti.toPlainText()),  # 34 elementi datanti
                str(self.comboBox_funz_statica_usm.currentText()),  # 35 funzione statica
                str(self.comboBox_lavorazione_usm.currentText()),  # 36 lavorazione usm
                str(self.lineEdit_spessore_giunti_usm.text()),  # 37 spessore giunti
                str(self.lineEdit_letti_di_posa_giunti_usm.text()),  # 38 letti posa giunti usm
                str(self.lineEdit_h_modulo_c_corsi_usm.text()),  # 39 altezza modulo corsi usm
                str(self.comboBox_unita_edilizia_riassuntiva_usm.currentText()),  # 40 unita edilizia riassuntiva
                str(self.comboBox_reimpiego_usm.currentText()),  # 41 unita edilizia riassuntiva
                str(self.comboBox_posa_in_opera_usm.currentText()),  # 42 posa in opera
                qmin_usm,  # 43 quota minima
                qmax_usm,  # 44 quota massima
                str(self.comboBox_consistenza_legante_usm.currentText()),  #  1 45 consitenza legante usm
                str(colore_legante_usm),  # 2 46 colore legante usm
                str(aggreg_legante_usm),  # 47 3 aggreg legante usm
                str(consistenza_texture_mat_usm),  # 4 48 consistenza text mat
                str(colore_materiale_usm),  # 5 49 colore materiale usm
                str(inclusi_materiali_usm), # 6 50 inclusi_mat_usm
                str(self.lineEdit_n_catalogo_generale.text()), # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
                str(self.lineEdit_n_catalogo_interno.text()), # 52 nr catalogo interno
                str(self.lineEdit_n_catalogo_internazionale.text()), # 53 nr catalogo internazionale
                str(self.comboBox_soprintendenza.currentText()), # 54 nr soprintendenza
                quota_relativa, #55 quota relativa
                quota_abs, #56 quota abs
                str(self.lineEdit_ref_tm.text()),  # 57 ref tm
                str(self.comboBox_ref_ra.currentText()),  # 58 ref ra
                str(self.lineEdit_ref_n.text()),  # 59 ref n
                str(self.comboBox_posizione.currentText()),  # 60 posizione
                str(self.lineEdit_criteri_distinzione.text()),  # 61 criteri distinzione
                str(self.comboBox_modo_formazione.currentText()),  # 62 modo formazione
                str(organici),  # 63 componenti organici
                str(inorganici),  # 64 componenti inorganici
                lunghezza_max,  # 65
                altezza_max,  # 66
                altezza_min,  # 67
                profondita_max,  # 68
                profondita_min,  # 69
                larghezza_media,  # 70
                quota_max_abs,  # 71
                quota_max_rel,  # 72
                quota_min_abs,  # 73
                quota_min_rel,  # 74
                str(self.textEdit_osservazioni.toPlainText()),  # 75 osservazioni
                str(self.lineEdit_datazione.text()),  # 76 datazione
                str(self.comboBox_flottazione.currentText()),  # 77 flottazione
                str(self.comboBox_setacciatura.currentText()),  # 78 setacciatura
                str(self.comboBox_affidabilita.currentText()),  # 79 affidabilita
                str(self.comboBox_direttore_us.currentText()),  # 80 direttore us
                str(self.comboBox_responsabile_us.currentText()),  # 81 responsabile us
                str(self.lineEdit_cod_ente_schedatore.text()),  # 82 cod ente schedatore
                str(self.lineEdit_data_rilevazione.text()),  # 83 data rilevazione
                str(self.lineEdit_data_rielaborazione.text()),  # 84 data rielaborazione
                lunghezza_usm,  # 85
                altezza_usm,  # 86
                spessore_usm,  # 87
                str(self.comboBox_tecnica_muraria_usm.currentText()),  # 88 tecnica muraria usm
                str(self.comboBox_modulo_usm.currentText()),  # 89 modulo usm
                str(self.lineEdit_campioni_malta_usm.text()),  # 90 campioni malta usm
                str(self.lineEdit_campioni_mattone_usm.text()),  # 91 campioni mattone usm
                str(self.lineEdit_campioni_pietra_usm.text()),  # 92 campioni pietra usm
                str(self.lineEdit_provenienza_materiali_usm.text()),  # 93 provenienza_materiali_usm
                str(self.lineEdit_criteri_distinzione_usm.text()),  # 94 criteri distinzione usm
                str(self.comboBox_uso_primario_usm.currentText()),
                str(self.comboBox_tipologia_opera.currentText()),
                str(self.comboBox_sezione_muraria.currentText()),
                str(self.comboBox_superficie_analizzata.currentText()),
                str(self.comboBox_orientamento.currentText()),
                str(self.comboBox_materiali_lat.currentText()),
                str(self.comboBox_lavorazione_lat.currentText()),
                str(self.comboBox_consistenza_lat.currentText()),
                str(self.comboBox_forma_lat.currentText()),
                str(self.comboBox_colore_lat.currentText()),
                str(self.comboBox_impasto_lat.currentText()),
                str(self.comboBox_forma_p.currentText()),
                str(self.comboBox_colore_p.currentText()),
                str(self.comboBox_taglio_p.currentText()),
                str(self.comboBox_posa_opera_p.currentText()),
                str(self.comboBox_inerti_usm.currentText()),  # 95 uso primario usm
                str(self.comboBox_tipo_legante_usm.currentText()),  # 95 uso primario usm
                str(self.comboBox_rifinitura_usm.currentText()),  # 95 uso primario usm
                str(self.comboBox_materiale_p.currentText()),  # 95 uso primario usm
                str(self.comboBox_consistenza_p.currentText()),  # 95 uso primario usm
                str(rapporti2),
                str(self.mQgsFileWidget.text())
                )
            # todelete
            # f = open("C:\\Users\\Luca\\pyarchinit_Report_folder\\data_insert_list.txt", "w")
            # f.write(str(data))
            # f.close
            # todelete
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as e:
                e_str = str(e)
                if e_str.__contains__("IntegrityError"):
                    if self.L=='it':
                        msg = "US già presente nel database"
                        QMessageBox.warning(self, "Error", "Error: " + str(msg), QMessageBox.Ok)
                    elif self.L=='de':
                        msg = self.ID_TABLE + " bereits in der Datenbank"
                        QMessageBox.warning(self, "Error", "Error: " + str(msg), QMessageBox.Ok)
                    else:
                        msg = self.ID_TABLE + " exist in db"
                        QMessageBox.warning(self, "Error", "Error: " + str(msg), QMessageBox.Ok)
                else:
                    msg = e
                    QMessageBox.warning(self, "Error", "Error 1 \n" + str(msg), QMessageBox.Ok)
                return 0
        except Exception as e:
            QMessageBox.warning(self, "Error", "Error 2 \n" + str(e), QMessageBox.Ok)
            return 0
            # insert new row into tableWidget
    def on_pushButton_insert_row_rapporti_pressed(self):
        self.insert_new_row('self.tableWidget_rapporti')

    def on_pushButton_remove_row_rapporti_pressed(self):
        self.remove_row('self.tableWidget_rapporti')

    def on_pushButton_insert_row_rapporti2_pressed(self):

        self.insert_new_row('self.tableWidget_rapporti2')
    def on_pushButton_remove_row_rapporti2_pressed(self):

        self.remove_row('self.tableWidget_rapporti2')

    def on_pushButton_insert_row_inclusi_pressed(self):
        self.insert_new_row('self.tableWidget_inclusi')
    def on_pushButton_remove_row_inclusi_pressed(self):
        self.remove_row('self.tableWidget_inclusi')
    def on_pushButton_insert_row_campioni_pressed(self):
        self.insert_new_row('self.tableWidget_campioni')
    def on_pushButton_remove_row_campioni_pressed(self):
        self.remove_row('self.tableWidget_campioni')
    def on_pushButton_insert_row_organici_pressed(self):
        self.insert_new_row('self.tableWidget_organici')
    def on_pushButton_remove_row_organici_pressed(self):
        self.remove_row('self.tableWidget_organici')
    def on_pushButton_insert_row_inorganici_pressed(self):
        self.insert_new_row('self.tableWidget_inorganici')
    def on_pushButton_remove_row_inorganici_pressed(self):
        self.remove_row('self.tableWidget_inorganici')
    def on_pushButton_insert_row_documentazione_pressed(self):
        self.insert_new_row('self.tableWidget_documentazione')
    def on_pushButton_remove_row_documentazione_pressed(self):
        self.remove_row('self.tableWidget_documentazione')
    def on_pushButton_insert_row_inclusi_materiali_pressed(self):
        self.insert_new_row('self.tableWidget_inclusi_materiali_usm')
    def on_pushButton_remove_row_inclusi_materiali_pressed(self):
        self.remove_row('self.tableWidget_inclusi_materiali_usm')
    def on_pushButton_insert_row_inclusi_leganti_pressed(self):
        self.insert_new_row('self.tableWidget_inclusi_leganti_usm')
    def on_pushButton_remove_row_inclusi_leganti_pressed(self):
        self.remove_row('self.tableWidget_inclusi_leganti_usm')
    def on_pushButton_insert_row_colore_legante_usm_pressed(self):
        self.insert_new_row('self.tableWidget_colore_legante_usm')
    def on_pushButton_remove_row_colore_legante_usm_pressed(self):
        self.remove_row('self.tableWidget_colore_legante_usm')
    def on_pushButton_insert_row_consistenza_texture_mat_usm_pressed(self):
        self.insert_new_row('self.tableWidget_consistenza_texture_mat_usm')
    def on_pushButton_remove_row_consistenza_texture_mat_usm_pressed(self):
        self.remove_row('self.tableWidget_consistenza_texture_mat_usm')
    def on_pushButton_insert_row_colore_materiale_usm_pressed(self):
        self.insert_new_row('self.tableWidget_colore_materiale_usm')
    def on_pushButton_remove_row_colore_materiale_usm_pressed(self):
        self.remove_row('self.tableWidget_colore_materiale_usm')
    def check_record_state(self):
        ec = self.data_error_check()
        if ec == 1:
            return 1  # ci sono errori di immissione
        elif self.records_equal_check() == 1 and ec == 0:
            if self.L=='it':
                self.update_if(
                    QMessageBox.warning(self, 'Errore', "Il record e' stato modificato. Vuoi salvare le modifiche?",
                                        QMessageBox.Ok | QMessageBox.Cancel))
            elif self.L=='de':
                self.update_if(
                    QMessageBox.warning(self, 'Errore', "Der Record wurde geändert. Möchtest du die Änderungen speichern?",
                                        QMessageBox.Ok | QMessageBox.Cancel))
            else:
                self.update_if(
                    QMessageBox.warning(self, "Error", "The record has been changed. You want to save the changes?",
                                        QMessageBox.Ok | QMessageBox.Cancel))
            return 0
            # records surf functions
    def on_pushButton_view_all_pressed(self):

        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        self.empty_fields()
        self.charge_records_n()
        self.fill_fields()
        self.BROWSE_STATUS = "b"
        self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
        if type(self.REC_CORR) == "<class 'str'>":
            corr = 0
        else:
            corr = self.REC_CORR
        self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
        self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
        self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
        self.SORT_STATUS = "n"
        self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])


    # def view_all(self):
    #
    #     self.checkBox_query.setChecked(False)
    #     if self.checkBox_query.isChecked():
    #         self.model_a.database().close()
    #     self.empty_fields()
    #     self.charge_records()
    #     # Controlla se il database è vuoto
    #     if not self.DATA_LIST:
    #         # Mostra un messaggio che indica che il database è vuoto
    #
    #         self.charge_list()
    #         self.BROWSE_STATUS = 'x'
    #         self.setComboBoxEnable(["self.comboBox_area"], "True")
    #         self.setComboBoxEnable(["self.lineEdit_us"], "True")
    #         self.on_pushButton_new_rec_pressed()
    #         return#QMessageBox.warning(self, "Attenzione", "Il database è vuoto.")
    #         #return  # Esci dalla funzione se il database è vuoto
    #
    #     self.fill_fields()
    #     self.BROWSE_STATUS = "b"
    #     self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
    #     if isinstance(self.REC_CORR, str):
    #         corr = 0
    #     else:
    #         corr = self.REC_CORR
    #
    #     self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
    #     self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
    #     self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
    #     self.SORT_STATUS = "n"
    #     self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
    def view_all(self):
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()

        self.empty_fields()

        # Filtra i record per il sito selezionato
        self.charge_records_filtered_by_site()

        # Controlla se il database è vuoto
        if not self.DATA_LIST:
            # Mostra un messaggio che indica che il database è vuoto
            self.charge_list()
            self.BROWSE_STATUS = 'x'
            self.setComboBoxEnable(["self.comboBox_area"], "True")
            self.setComboBoxEnable(["self.lineEdit_us"], "True")
            self.on_pushButton_new_rec_pressed()
            return

        self.fill_fields()
        self.BROWSE_STATUS = "b"
        self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
        if isinstance(self.REC_CORR, str):
            corr = 0
        else:
            corr = self.REC_CORR

        self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
        self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
        self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
        self.SORT_STATUS = "n"
        self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])

    def charge_records_filtered_by_site(self):
        """Carica i record filtrati per il sito selezionato"""
        try:
            conn = Connection()
            sito_set = conn.sito_set()
            sito_set_str = sito_set['sito_set']

            current_site = sito_set_str

            if not current_site:
                # Se nessun sito è selezionato, carica tutti i record
                self.charge_records()
                return

            # Filtra i record per il sito corrente
            search_dict = {'sito': f"'{current_site}'"}
            search_dict = {str(k): str(v) for k, v in search_dict.items()}

            res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)

            if bool(res):
                self.DATA_LIST = []
                for i in res:
                    self.DATA_LIST.append(i)
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                return True
            else:
                self.DATA_LIST = []
                return False

        except Exception as e:
            QMessageBox.warning(self, "Errore", f"Errore nel caricamento dei dati filtrati: {str(e)}")
            # In caso di errore, carica tutti i record
            self.charge_records()
            return False

    def on_pushButton_first_rec_pressed(self):

        self.iconListWidget.update()
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.check_record_state() == 1:
            pass
        else:
            try:
                self.empty_fields()
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                self.fill_fields(0)
                self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
            except :
                pass
    def on_pushButton_last_rec_pressed(self):
        self.iconListWidget.update()
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.check_record_state() == 1:
            pass
        else:
            try:
                self.empty_fields()
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), len(self.DATA_LIST) - 1
                self.fill_fields(self.REC_CORR)
                self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
            except:
                pass
    def on_pushButton_prev_rec_pressed(self):
        self.iconListWidget.update()
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        rec_goto = int(self.lineEdit_goto.text())
        if self.check_record_state() == 1:
            pass
        else:
            self.REC_CORR = self.REC_CORR - rec_goto
        if self.REC_CORR <= -1:
            self.REC_CORR = self.REC_CORR + rec_goto
            #QMessageBox.information(self, "Warning", "you are to the first record", QMessageBox.Ok)
        else:
            try:
                self.empty_fields()
                self.fill_fields(self.REC_CORR)
                self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
            except:# Exception as e:
                pass#QMessageBox.warning(self, "Error", str(e), QMessageBox.Ok)
    def on_pushButton_next_rec_pressed(self):
        self.iconListWidget.update()
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        rec_goto = int(self.lineEdit_goto.text())



        if self.check_record_state() == 1:
            pass
        else:
            self.REC_CORR = self.REC_CORR + rec_goto
        if self.REC_CORR >= self.REC_TOT:
            self.REC_CORR = self.REC_CORR - rec_goto
            #QMessageBox.information(self, "Warning", "you are to the last record", QMessageBox.Ok)
        else:
            try:
                self.empty_fields()
                self.fill_fields(self.REC_CORR)
                self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
            except Exception as e:
                QMessageBox.warning(self, "Error", str(e), QMessageBox.Ok)


        if self.checkBox_validate.isChecked():
            self.selectRows()
            # while True:
                # try:
                    # self.empty_fields()
                    # self.fill_fields(self.REC_CORR)
                    # self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)
                # except :#Exception as e:
                    # pass#QMessageBox.warning(self, "Error", str(e), QMessageBox.Ok)
                # #continue
                # else:
                    # break
    def on_pushButton_delete_pressed(self):
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.L=='it':
            msg = QMessageBox.warning(self, "Attenzione!!!",
                                      "Vuoi veramente eliminare il record? \n L'azione è irreversibile",
                                      QMessageBox.Ok | QMessageBox.Cancel)
            if msg == QMessageBox.Cancel:
                QMessageBox.warning(self, "Messagio!!!", "Azione Annullata!")
            else:
                try:
                    id_to_delete = eval("self.DATA_LIST[self.REC_CORR]." + self.ID_TABLE)
                    self.DB_MANAGER.delete_one_record(self.TABLE_NAME, self.ID_TABLE, id_to_delete)
                    self.charge_records()  # charge records from DB
                    QMessageBox.warning(self, "Messaggio!!!", "Record eliminato!")
                    self.iconListWidget.update()
                except Exception as e:
                    QMessageBox.warning(self, "Messaggio!!!", "Tipo di errore: " + str(e))
                if not bool(self.DATA_LIST):
                    QMessageBox.warning(self, "Attenzione", "Il database è vuoto!", QMessageBox.Ok)
                    self.iconListWidget.update()
                    self.DATA_LIST = []
                    self.DATA_LIST_REC_CORR = []
                    self.DATA_LIST_REC_TEMP = []
                    self.REC_CORR = 0
                    self.REC_TOT = 0
                    self.empty_fields()
                    self.set_rec_counter(0, 0)
                    # check if DB is empty
                if bool(self.DATA_LIST):
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    self.charge_list()
                    self.fill_fields()
                    self.set_sito()
                    self.view_all()
        elif self.L=='de':
            msg = QMessageBox.warning(self, "Achtung!!!",
                                      "Willst du wirklich diesen Eintrag löschen? \n Der Vorgang ist unumkehrbar",
                                      QMessageBox.Ok | QMessageBox.Cancel)
            if msg == QMessageBox.Cancel:
                QMessageBox.warning(self, "Message!!!", "Aktion annulliert!")
            else:
                try:
                    id_to_delete = eval("self.DATA_LIST[self.REC_CORR]." + self.ID_TABLE)
                    self.DB_MANAGER.delete_one_record(self.TABLE_NAME, self.ID_TABLE, id_to_delete)
                    self.charge_records()  # charge records from DB
                    QMessageBox.warning(self, "Message!!!", "Record gelöscht!")
                except Exception as e:
                    QMessageBox.warning(self, "Messagge!!!", "Errortyp: " + str(e))
                if not bool(self.DATA_LIST):
                    QMessageBox.warning(self, "Attenzione", "Die Datenbank ist leer!", QMessageBox.Ok)
                    self.DATA_LIST = []
                    self.DATA_LIST_REC_CORR = []
                    self.DATA_LIST_REC_TEMP = []
                    self.REC_CORR = 0
                    self.REC_TOT = 0
                    self.empty_fields()
                    self.set_rec_counter(0, 0)
                    # check if DB is empty
                if bool(self.DATA_LIST):
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    self.charge_list()
                    self.fill_fields()
                    self.set_sito()
        else:
            msg = QMessageBox.warning(self, "Warning!!!",
                                      "Do you really want to break the record? \n Action is irreversible.",
                                      QMessageBox.Ok | QMessageBox.Cancel)
            if msg == QMessageBox.Cancel:
                QMessageBox.warning(self, "Message!!!", "Action deleted!")
            else:
                try:
                    id_to_delete = eval("self.DATA_LIST[self.REC_CORR]." + self.ID_TABLE)
                    self.DB_MANAGER.delete_one_record(self.TABLE_NAME, self.ID_TABLE, id_to_delete)
                    self.charge_records()  # charge records from DB
                    QMessageBox.warning(self, "Message!!!", "Record deleted!")
                except Exception as e:
                    QMessageBox.warning(self, "Message", "error type: " + str(e))
                if not bool(self.DATA_LIST):
                    QMessageBox.warning(self, "Warning", "the db is empty!", QMessageBox.Ok)
                    self.DATA_LIST = []
                    self.DATA_LIST_REC_CORR = []
                    self.DATA_LIST_REC_TEMP = []
                    self.REC_CORR = 0
                    self.REC_TOT = 0
                    self.empty_fields()
                    self.set_rec_counter(0, 0)
                    # check if DB is empty
                if bool(self.DATA_LIST):
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    self.charge_list()
                    self.fill_fields()
                    self.set_sito()
            self.SORT_STATUS = "n"
            self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
            #

    def delete_all_filtered_records(self):
        # Se non ci sono record da eliminare, emetti un avvertimento e interrompi
        if not self.DATA_LIST:
            QMessageBox.warning(self, "Warning", "No records to delete!", QMessageBox.Ok)
            return

        # Chiedi conferma all'utente prima di eliminare i record
        if self.L == 'it':
            msg = QMessageBox.warning(self, "Attenzione!!!",
                                      "Vuoi veramente eliminare i record? \n L'azione è irreversibile",
                                      QMessageBox.Ok | QMessageBox.Cancel)
        elif self.L == 'de':
            msg = QMessageBox.warning(self, "Achtung!!!",
                                      "Willst du wirklich diese Einträge löschen? \n Der Vorgang ist unumkehrbar",
                                      QMessageBox.Ok | QMessageBox.Cancel)
        else:
            msg = QMessageBox.warning(self, "Warning!!!",
                                      "Do you really want to delete the records? \n Action is irreversible.",
                                      QMessageBox.Ok | QMessageBox.Cancel)

        # Se l'utente ha annullato, interrompi
        if msg == QMessageBox.Cancel:
            QMessageBox.warning(self, "Message!!!", "Action cancelled!", QMessageBox.Ok)
            return

        # Cancella ogni record
        for record in self.DATA_LIST:
            id_to_delete = getattr(record, self.ID_TABLE)
            self.DB_MANAGER.delete_one_record(self.TABLE_NAME, self.ID_TABLE, id_to_delete)

        # Ricarica l'elenco dei record e aggiorna l'interfaccia utente
        self.charge_records()
        self.view_all()

    def on_pushButton_new_search_pressed(self):
        self.checkBox_query.setChecked(False)
        if self.checkBox_query.isChecked():
            self.model_a.database().close()
        if self.BROWSE_STATUS != "f" and self.check_record_state() == 1:
            pass
        else:
            self.enable_button_search(0)
            conn = Connection()

            sito_set= conn.sito_set()
            sito_set_str = sito_set['sito_set']
            if self.BROWSE_STATUS != "f":
                if bool(self.comboBox_sito.currentText()) and self.comboBox_sito.currentText()==sito_set_str:
                    self.BROWSE_STATUS = "f"
                    self.empty_fields_nosite()
                    self.lineEdit_data_schedatura.setText("")
                    self.lineEdit_anno.setText("")
                    self.comboBox_formazione.setEditText("")
                    self.comboBox_metodo.setEditText("")
                    #self.setComboBoxEditable(["self.comboBox_sito"], 1)
                    self.setComboBoxEditable(["self.comboBox_area"], 1)
                    self.setComboBoxEditable(["self.comboBox_unita_tipo"],1)
                    self.setComboBoxEnable(["self.comboBox_sito"], "False")
                    self.setComboBoxEnable(["self.comboBox_area"], "True")
                    self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                    self.setComboBoxEnable(["self.lineEdit_us"], "True")
                    self.setComboBoxEnable(["self.textEdit_descrizione"], "False")
                    self.setComboBoxEnable(["self.textEdit_interpretazione"], "False")
                    self.setTableEnable(
                        ["self.tableWidget_campioni",
                         "self.tableWidget_rapporti",
                         "self.tableWidget_inclusi",
                         "self.tableWidget_organici",
                         "self.tableWidget_inorganici",
                         "self.tableWidget_documentazione",
                         "self.tableWidget_inclusi_materiali_usm",
                         "self.tableWidget_colore_legante_usm",
                         "self.tableWidget_inclusi_leganti_usm",
                         "self.tableWidget_consistenza_texture_mat_usm",
                         "self.tableWidget_colore_materiale_usm","self.tableWidget_rapporti2"], "False")
                    ###
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter('', '')
                    self.label_sort.setText(self.SORTED_ITEMS["n"])
                    #self.charge_list()

                else:
                    self.BROWSE_STATUS = "f"
                    ###
                    self.lineEdit_data_schedatura.setText("")
                    self.lineEdit_anno.setText("")
                    self.comboBox_formazione.setEditText("")
                    self.comboBox_metodo.setEditText("")
                    self.setComboBoxEditable(["self.comboBox_sito"], 1)
                    self.setComboBoxEditable(["self.comboBox_area"], 1)
                    self.setComboBoxEditable(["self.comboBox_unita_tipo"], 1)
                    self.setComboBoxEnable(["self.comboBox_sito"], "True")
                    self.setComboBoxEnable(["self.comboBox_area"], "True")
                    self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                    self.setComboBoxEnable(["self.lineEdit_us"], "True")
                    self.setComboBoxEnable(["self.textEdit_descrizione"], "False")
                    self.setComboBoxEnable(["self.textEdit_interpretazione"], "False")
                    self.setTableEnable(
                        ["self.tableWidget_campioni",
                         "self.tableWidget_rapporti",
                         "self.tableWidget_inclusi",
                         "self.tableWidget_organici",
                         "self.tableWidget_inorganici",
                         "self.tableWidget_documentazione",
                         "self.tableWidget_inclusi_materiali_usm",
                         "self.tableWidget_colore_legante_usm",
                         "self.tableWidget_inclusi_leganti_usm",
                         "self.tableWidget_consistenza_texture_mat_usm",
                         "self.tableWidget_colore_materiale_usm","self.tableWidget_rapporti2"], "False")
                    ###
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter('', '')
                    self.label_sort.setText(self.SORTED_ITEMS["n"])
                    self.charge_list()
                    self.empty_fields()
    def on_pushButton_showLayer_pressed(self):
        """
        for sing_us in range(len(self.DATA_LIST)):
            sing_layer = [self.DATA_LIST[sing_us]]
            self.pyQGIS.charge_vector_layers(sing_layer)
        """
        sing_layer = [self.DATA_LIST[self.REC_CORR]]
        self.pyQGIS.charge_vector_layers(sing_layer)
        self.pyQGIS.charge_usm_layers(sing_layer)
    def on_pushButton_crea_codice_periodo_pressed(self):

        try:
            self.set_sito()
            sito = str(self.comboBox_sito.currentText())
            self.DB_MANAGER.update_cont_per(sito)
            self.empty_fields()
            #self.charge_records()
            self.fill_fields(self.REC_CORR)  # ricaricare tutti i record in uso e passare il valore REC_CORR a fill_fields
            if self.L=='it':
                QMessageBox.warning(self, "INFO", "Codice periodo aggiornato per lo scavo %s" % (sito), QMessageBox.Ok)
            elif self.L=='de':
                QMessageBox.warning(self, "INFO", "Der Zeitstellungscode wurde für die Ausgrabung hochgeladen %s" % (sito), QMessageBox.Ok)
            elif self.L=='en':
                QMessageBox.warning(self, "INFO", "Updated period code for excavation %s" % (sito), QMessageBox.Ok)
        except KeyError as e:
            if self.L=='it':
                QMessageBox.warning(self, "Attenzione", str(e), QMessageBox.Ok)
            elif self.L=='de':
                QMessageBox.warning(self, "Achtung", str(e), QMessageBox.Ok)
            elif self.L=='en':
                QMessageBox.warning(self, "Attention", str(e), QMessageBox.Ok)

    def switch_search_mode(self):
        self.use_like_query = not self.use_like_query
        if self.use_like_query:
            QMessageBox.information(self, "Search mode changed",
                                    "Query_bool_like mode is now active. Use Ctrl+Shift+N to deactivate.")

            # Disable all the fields
            self.comboBox_sito.setEnabled(False)
            self.comboBox_area.setEnabled(False)
            self.comboBox_unita_tipo.setEnabled(False)
            self.lineEdit_us.setEnabled(False)
            self.comboBox_def_strat.setEnabled(False)
            self.comboBox_def_intepret.setEnabled(False)
            self.textEdit_descrizione.setEnabled(False)
            self.textEdit_interpretazione.setEnabled(False)
            self.comboBox_per_iniz.setEnabled(False)
            self.comboBox_fas_iniz.setEnabled(False)
            self.comboBox_per_fin.setEnabled(False)
            self.comboBox_fas_fin.setEnabled(False)
            self.comboBox_scavato.setEnabled(False)
            self.lineEdit_attivita.setEnabled(False)
            self.lineEdit_anno.setEnabled(False)
            self.comboBox_metodo.setEnabled(False)
            self.lineEdit_data_schedatura.setEnabled(False)
            self.comboBox_schedatore.setEnabled(False)
            self.comboBox_formazione.setEnabled(False)
            self.comboBox_conservazione.setEnabled(False)
            self.comboBox_colore.setEnabled(False)
            self.comboBox_consistenza.setEnabled(False)
            self.comboBox_struttura.setEnabled(False)
            self.lineEdit_codice_periodo.setEnabled(False)
            self.lineEditOrderLayer.setEnabled(False)
            self.comboBox_unita_edilizia_riassuntiva_usm.setEnabled(False)
            self.comboBox_reimpiego_usm.setEnabled(False)
            self.comboBox_posa_in_opera_usm.setEnabled(False)
            self.textEdit_osservazioni.setEnabled(False)
            self.lineEdit_datazione.setEnabled(False)
            self.comboBox_flottazione.setEnabled(False)
            self.comboBox_setacciatura.setEnabled(False)
            self.comboBox_affidabilita.setEnabled(False)
            self.comboBox_direttore_us.setEnabled(False)
            self.comboBox_responsabile_us.setEnabled(False)
            self.lineEdit_cod_ente_schedatore.setEnabled(False)
            self.lineEdit_data_rilevazione.setEnabled(False)
            self.lineEdit_data_rielaborazione.setEnabled(False)
            self.comboBox_tecnica_muraria_usm.setEnabled(False)
            self.comboBox_modulo_usm.setEnabled(False)
            self.lineEdit_campioni_malta_usm.setEnabled(False)
            self.lineEdit_campioni_mattone_usm.setEnabled(False)
            self.lineEdit_campioni_pietra_usm.setEnabled(False)
            self.lineEdit_provenienza_materiali_usm.setEnabled(False)
            self.lineEdit_criteri_distinzione_usm.setEnabled(False)
            self.comboBox_uso_primario_usm.setEnabled(False)
            self.comboBox_tipologia_opera.setEnabled(False)
            self.comboBox_sezione_muraria.setEnabled(False)
            self.comboBox_superficie_analizzata.setEnabled(False)
            self.comboBox_orientamento.setEnabled(False)
            self.comboBox_materiali_lat.setEnabled(False)
            self.comboBox_lavorazione_lat.setEnabled(False)
            self.comboBox_consistenza_lat.setEnabled(False)
            self.comboBox_forma_lat.setEnabled(False)
            self.comboBox_colore_lat.setEnabled(False)
            self.comboBox_impasto_lat.setEnabled(False)
            self.comboBox_forma_p.setEnabled(False)
            self.comboBox_colore_p.setEnabled(False)
            self.comboBox_taglio_p.setEnabled(False)
            self.comboBox_posa_opera_p.setEnabled(False)
            self.comboBox_inerti_usm.setEnabled(False)
            self.comboBox_tipo_legante_usm.setEnabled(False)
            self.comboBox_rifinitura_usm.setEnabled(False)
            self.comboBox_materiale_p.setEnabled(False)
            self.comboBox_consistenza_p.setEnabled(False)

            # And then enable those specific fields you mentioned:
            self.comboBox_area.setEnabled(True)
            self.comboBox_struttura.setEnabled(True)
            self.lineEdit_quadrato.setEnabled(True)
            self.comboBox_settore.setEnabled(True)
            self.lineEdit_ambiente.setEnabled(True)
            self.lineEdit_saggio.setEnabled(True)
        else:
            #QMessageBox.information(self, "Search mode changed",
                                    #"Regular query_bool mode is now deactive. Use Ctrl+Shift+N to activate query_bool_like mode.")

            self.comboBox_sito.setEnabled(False)
            self.comboBox_area.setEnabled(False)
            self.comboBox_unita_tipo.setEnabled(True)
            self.lineEdit_us.setEnabled(False)
            self.comboBox_def_strat.setEnabled(True)
            self.comboBox_def_intepret.setEnabled(True)
            self.textEdit_descrizione.setEnabled(True)
            self.textEdit_interpretazione.setEnabled(True)
            self.comboBox_per_iniz.setEnabled(True)
            self.comboBox_fas_iniz.setEnabled(True)
            self.comboBox_per_fin.setEnabled(True)
            self.comboBox_fas_fin.setEnabled(True)
            self.comboBox_scavato.setEnabled(True)
            self.lineEdit_attivita.setEnabled(True)
            self.lineEdit_anno.setEnabled(True)
            self.comboBox_metodo.setEnabled(True)
            self.lineEdit_data_schedatura.setEnabled(True)
            self.comboBox_schedatore.setEnabled(True)
            self.comboBox_formazione.setEnabled(True)
            self.comboBox_conservazione.setEnabled(True)
            self.comboBox_colore.setEnabled(True)
            self.comboBox_consistenza.setEnabled(True)
            self.comboBox_struttura.setEnabled(True)
            self.lineEdit_codice_periodo.setEnabled(True)
            self.lineEditOrderLayer.setEnabled(True)
            self.comboBox_unita_edilizia_riassuntiva_usm.setEnabled(True)
            self.comboBox_reimpiego_usm.setEnabled(True)
            self.comboBox_posa_in_opera_usm.setEnabled(True)
            self.textEdit_osservazioni.setEnabled(True)
            self.lineEdit_datazione.setEnabled(True)
            self.comboBox_flottazione.setEnabled(True)
            self.comboBox_setacciatura.setEnabled(True)
            self.comboBox_affidabilita.setEnabled(True)
            self.comboBox_direttore_us.setEnabled(True)
            self.comboBox_responsabile_us.setEnabled(True)
            self.lineEdit_cod_ente_schedatore.setEnabled(True)
            self.lineEdit_data_rilevazione.setEnabled(True)
            self.lineEdit_data_rielaborazione.setEnabled(True)
            self.comboBox_tecnica_muraria_usm.setEnabled(True)
            self.comboBox_modulo_usm.setEnabled(True)
            self.lineEdit_campioni_malta_usm.setEnabled(True)
            self.lineEdit_campioni_mattone_usm.setEnabled(True)
            self.lineEdit_campioni_pietra_usm.setEnabled(True)
            self.lineEdit_provenienza_materiali_usm.setEnabled(True)
            self.lineEdit_criteri_distinzione_usm.setEnabled(True)
            self.comboBox_uso_primario_usm.setEnabled(True)
            self.comboBox_tipologia_opera.setEnabled(True)
            self.comboBox_sezione_muraria.setEnabled(True)
            self.comboBox_superficie_analizzata.setEnabled(True)
            self.comboBox_orientamento.setEnabled(True)
            self.comboBox_materiali_lat.setEnabled(True)
            self.comboBox_lavorazione_lat.setEnabled(True)
            self.comboBox_consistenza_lat.setEnabled(True)
            self.comboBox_forma_lat.setEnabled(True)
            self.comboBox_colore_lat.setEnabled(True)
            self.comboBox_impasto_lat.setEnabled(True)
            self.comboBox_forma_p.setEnabled(True)
            self.comboBox_colore_p.setEnabled(True)
            self.comboBox_taglio_p.setEnabled(True)
            self.comboBox_posa_opera_p.setEnabled(True)
            self.comboBox_inerti_usm.setEnabled(True)
            self.comboBox_tipo_legante_usm.setEnabled(True)
            self.comboBox_rifinitura_usm.setEnabled(True)
            self.comboBox_materiale_p.setEnabled(True)
            self.comboBox_consistenza_p.setEnabled(True)
            self.setTableEnable(
                ["self.tableWidget_campioni", "self.tableWidget_rapporti", "self.tableWidget_inclusi",
                 "self.tableWidget_organici", "self.tableWidget_inorganici", "self.tableWidget_documentazione",
                 "self.tableWidget_rapporti2"], "True")
            self.use_like_query = False
    def on_pushButton_search_go_pressed(self):


        self.checkBox_query.setChecked(False)
        if self.BROWSE_STATUS != "f":
            if self.L=='it':
                QMessageBox.warning(self, "ATTENZIONE", "Per eseguire una nuova ricerca clicca sul pulsante 'new search' ",
                                    QMessageBox.Ok)
            elif self.L=='de':
                QMessageBox.warning(self, "ACHTUNG", "Um eine neue Abfrage zu starten drücke  'new search' ",
                                    QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "WARNING", "To perform a new search click on the 'new search' button ",
                                    QMessageBox.Ok)
        else:
            # TableWidget
            if self.lineEdit_us.text() != "":
                us = int(self.lineEdit_us.text())
            else:
                us = ""
            ##qmin_usm
            if self.lineEdit_qmin_usm.text() != "":

                qmin_usm = float(self.lineEdit_qmin_usm.text())
            else:
                qmin_usm = None
            ##qmax_usm
            if self.lineEdit_qmax_usm.text() != "":
                qmax_usm = float(self.lineEdit_qmax_usm.text())
            else:
                qmax_usm = None
            #pre pyarchinit 3.0
            ##quota relativa
            if self.lineEdit_quota_relativa.text() == "":
                quota_relativa = None
            else:
                quota_relativa = float(self.lineEdit_quota_relativa.text())
            ##quota abs
            if self.lineEdit_quota_abs.text() == "":
                quota_abs = None
            else:
                quota_abs = float(self.lineEdit_quota_abs.text())
            ##lunghezza max
            if self.lineEdit_lunghezza_max.text() == "":
                lunghezza_max = None
            else:
                lunghezza_max = float(self.lineEdit_lunghezza_max.text())
            ##altezza max
            if self.lineEdit_altezza_max.text() == "":
                altezza_max = None
            else:
                altezza_max = float(self.lineEdit_altezza_max.text())
            ##altezza min
            if self.lineEdit_altezza_min.text() == "":
                altezza_min = None
            else:
                altezza_min = float(self.lineEdit_altezza_min.text())
            ##profondita max
            if self.lineEdit_profondita_max.text() == "":
                profondita_max = None
            else:
                profondita_max = float(self.lineEdit_profondita_max.text())
            ##profondita min
            if self.lineEdit_profondita_min.text() == "":
                profondita_min = None
            else:
                profondita_min = float(self.lineEdit_profondita_min.text())
            ##larghezza media
            if self.lineEdit_larghezza_media.text() == "":
                larghezza_media = None
            else:
                larghezza_media = float(self.lineEdit_larghezza_media.text())
            ##quota max abs
            if self.lineEdit_quota_max_abs.text() == "":
                quota_max_abs = None
            else:
                quota_max_abs = float(self.lineEdit_quota_max_abs.text())
            ##quota max relativa
            if self.lineEdit_quota_max_rel.text() == "":
                quota_max_rel = None
            else:
                quota_max_rel = float(self.lineEdit_quota_max_rel.text())
            ##quota min abs
            if self.lineEdit_quota_min_abs.text() == "":
                quota_min_abs = None
            else:
                quota_min_abs = float(self.lineEdit_quota_min_abs.text())
            ##quota min relativa
            if self.lineEdit_quota_min_rel.text() == "":
                quota_min_rel = None
            else:
                quota_min_rel = float(self.lineEdit_quota_min_rel.text())
            ##lunghezza usm
            if self.lineEdit_lunghezza_usm.text() == "":
                lunghezza_usm = None
            else:
                lunghezza_usm = float(self.lineEdit_lunghezza_usm.text())
            ##altezza usm
            if self.lineEdit_altezza_usm.text() == "":
                altezza_usm = None
            else:
                altezza_usm = float(self.lineEdit_altezza_usm.text())
            ##spessore usm
            if self.lineEdit_spessore_usm.text() == "":
                spessore_usm = None
            else:
                spessore_usm = float(self.lineEdit_spessore_usm.text())

            search_dict_like = {
                self.TABLE_FIELDS[0]: self.comboBox_sito.currentText(),  # 1 - Sito
                self.TABLE_FIELDS[1]: self.comboBox_area.currentText(),
                self.TABLE_FIELDS[24]: self.comboBox_struttura.currentText(),  # 22 - struttura
                self.TABLE_FIELDS[29]: self.comboBox_settore.currentText(),  # 24 - order layer
                self.TABLE_FIELDS[30]: self.lineEdit_quadrato.text(),
                self.TABLE_FIELDS[31]: self.lineEdit_ambiente.text(),  # 30 quadrato
                self.TABLE_FIELDS[32]: self.lineEdit_saggio.text()
                # 30 quadrato
            }

            search_dict = {
                self.TABLE_FIELDS[0]:  "'"+ str(self.comboBox_sito.currentText())+"'",  # 1 - Sito
                self.TABLE_FIELDS[1]:  "'" +str(self.comboBox_area.currentText()) +"'",  # 2 - Area
                self.TABLE_FIELDS[2]: us,  # 3 - US
                self.TABLE_FIELDS[3]:  "'" +str(self.comboBox_def_strat.currentText()) +"'",
                # 4 - Definizione stratigrafica
                self.TABLE_FIELDS[4]:  "'" +str(self.comboBox_def_intepret.currentText()) +"'",
                # 5 - Definizione intepretata
                self.TABLE_FIELDS[5]: "'" +str(self.textEdit_descrizione.toPlainText())+"'",  # 6 - descrizione
                self.TABLE_FIELDS[6]: "'" +str(self.textEdit_interpretazione.toPlainText())+"'",  # 7 - interpretazione
                self.TABLE_FIELDS[7]:  "'" +str(self.comboBox_per_iniz.currentText()) +"'",  # 8 - periodo iniziale
                self.TABLE_FIELDS[8]:  "'" +str(self.comboBox_fas_iniz.currentText())+"'" ,  # 9 - fase iniziale
                self.TABLE_FIELDS[9]:  "'" +str(self.comboBox_per_fin.currentText()) +"'",
                # 10 - periodo finale iniziale
                self.TABLE_FIELDS[10]:  "'" +str(self.comboBox_fas_fin.currentText())+"'" ,  # 11 - fase finale
                self.TABLE_FIELDS[11]:  "'" +str(self.comboBox_scavato.currentText()) +"'",  # 12 - scavato
                self.TABLE_FIELDS[12]:  "'" +str(self.lineEdit_attivita.text()) +"'",  # 13 - attivita
                self.TABLE_FIELDS[13]:  "'" +str(self.lineEdit_anno.text()) +"'",  # 14 - anno scavo
                self.TABLE_FIELDS[14]:  "'" +str(self.comboBox_metodo.currentText()) +"'",  # 15 - metodo
                self.TABLE_FIELDS[18]:  "'" +str(self.lineEdit_data_schedatura.text()) +"'",  # 16 - data schedatura
                self.TABLE_FIELDS[19]:  "'" +str(self.comboBox_schedatore.currentText()) +"'",  # 17 - schedatore
                self.TABLE_FIELDS[20]:  "'" +str(self.comboBox_formazione.currentText())+"'" ,  # 18 - formazione
                self.TABLE_FIELDS[21]:  "'" +str(self.comboBox_conservazione.currentText()) +"'",  # 19 - conservazione
                self.TABLE_FIELDS[22]:  "'" +str(self.comboBox_colore.currentText())+"'" ,  # 20 - colore
                self.TABLE_FIELDS[23]:  "'" +str(self.comboBox_consistenza.currentText())+"'" ,  # 21 - consistenza
                self.TABLE_FIELDS[24]:  "'" +str(self.comboBox_struttura.currentText())+"'",  # 22 - struttura
                self.TABLE_FIELDS[25]:  "'" +str(self.lineEdit_codice_periodo.text()) +"'",  # 23 - codice_periodo
                self.TABLE_FIELDS[26]:  "'" +str(self.lineEditOrderLayer.text())+"'" ,  # 24 - order layer
                self.TABLE_FIELDS[28]:  "'" +str(self.comboBox_unita_tipo.currentText()) +"'",  # 24 - order layer
                self.TABLE_FIELDS[29]:  "'" +str(self.comboBox_settore.currentText())+"'" ,  # 24 - order layer
                self.TABLE_FIELDS[30]:  "'" +str(self.lineEdit_quadrato.text())+"'" ,  # 30 quadrato
                self.TABLE_FIELDS[31]:  "'" +str(self.lineEdit_ambiente.text())+"'" ,  # 30 quadrato
                self.TABLE_FIELDS[32]:  "'" +str(self.lineEdit_saggio.text())+"'" ,  # 30 quadrato
                self.TABLE_FIELDS[33]: "'" +str(self.textEdit_elementi_datanti.toPlainText())+"'",  # 6 - descrizione
                self.TABLE_FIELDS[34]:  "'" +str(self.comboBox_funz_statica_usm.currentText())+"'" ,
                # 24 - order layer
                self.TABLE_FIELDS[35]:  "'" +str(self.comboBox_lavorazione_usm.currentText()) +"'",  # 30 quadrato
                self.TABLE_FIELDS[36]:  "'" +str(self.lineEdit_spessore_giunti_usm.text())+"'" ,  # 30 quadrato
                self.TABLE_FIELDS[37]:  "'" +str(self.lineEdit_letti_di_posa_giunti_usm.text()) +"'",
                self.TABLE_FIELDS[38]:  "'" +str(self.lineEdit_h_modulo_c_corsi_usm.text())+"'" ,
                self.TABLE_FIELDS[39]:  "'" +str(self.comboBox_unita_edilizia_riassuntiva_usm.currentText())+"'" ,
                self.TABLE_FIELDS[40]:  "'" +str(self.comboBox_reimpiego_usm.currentText()) +"'",
                self.TABLE_FIELDS[41]:  "'" +str(self.comboBox_posa_in_opera_usm.currentText())+"'" ,
                self.TABLE_FIELDS[42]: qmin_usm,
                self.TABLE_FIELDS[43]: qmax_usm,
                self.TABLE_FIELDS[44]:  "'" +str(self.comboBox_consistenza_legante_usm.currentText())+"'" ,
                self.TABLE_FIELDS[50]:  "'" +str(self.lineEdit_n_catalogo_generale.text()) +"'",
            # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
                self.TABLE_FIELDS[51]:  "'" +str(self.lineEdit_n_catalogo_interno.text())+"'" ,
            # 52 nr catalogo interno
                self.TABLE_FIELDS[52]:  "'" +str(self.lineEdit_n_catalogo_internazionale.text()) +"'",
            # 53 nr catalogo internazionale
                self.TABLE_FIELDS[53]:  "'" +str(self.comboBox_soprintendenza.currentText()) +"'",
            # 54 nr soprintendenza
                self.TABLE_FIELDS[54]:  quota_relativa,  # 55 quota relativa
                self.TABLE_FIELDS[55]:  quota_abs,  # 56 quota abs
                self.TABLE_FIELDS[56]:  "'" +str(self.lineEdit_ref_tm.text()) +"'",  # 57 ref tm
                self.TABLE_FIELDS[57]:  "'" +str(self.comboBox_ref_ra.currentText()) +"'",  # 58 ref ra
                self.TABLE_FIELDS[58]:  "'" +str(self.lineEdit_ref_n.text()) +"'",  # 59 ref n
                self.TABLE_FIELDS[59]:  "'" +str(self.comboBox_posizione.currentText())+"'" ,  # 60 posizione
                self.TABLE_FIELDS[60]:  "'" +str(self.lineEdit_criteri_distinzione.text())+"'" ,
            # 61 criteri distinzione
                self.TABLE_FIELDS[61]:  "'" +str(self.comboBox_modo_formazione.currentText()) +"'",
            # 62 modo formazione
            #    self.TABLE_FIELDS[62]:  str(self.comboBox_componenti_organici.currentText()) ,
            # 63 componenti organici
            #    self.TABLE_FIELDS[63]:  str(self.comboBox_componenti_inorganici.currentText()) ,
            # 64 componenti inorganici
                self.TABLE_FIELDS[64]:  lunghezza_max,  # 65
                self.TABLE_FIELDS[65]:  altezza_max,  # 66
                self.TABLE_FIELDS[66]:  altezza_min,  # 67
                self.TABLE_FIELDS[67]:  profondita_max,  # 68
                self.TABLE_FIELDS[68]:  profondita_min,  # 69
                self.TABLE_FIELDS[69]:  larghezza_media,  # 70
                self.TABLE_FIELDS[70]:  quota_max_abs,  # 71
                self.TABLE_FIELDS[71]:  quota_max_rel,  # 72
                self.TABLE_FIELDS[72]:  quota_min_abs,  # 73
                self.TABLE_FIELDS[73]:  quota_min_rel,  # 74
                self.TABLE_FIELDS[74]:  "'" +str(self.textEdit_osservazioni.toPlainText())+"'" ,  # 75 osservazioni
                self.TABLE_FIELDS[75]:  "'" +str(self.lineEdit_datazione.text())+"'",  # 76 datazione
                self.TABLE_FIELDS[76]:  "'" +str(self.comboBox_flottazione.currentText()) +"'",  # 77 flottazione
                self.TABLE_FIELDS[77]:  "'" +str(self.comboBox_setacciatura.currentText()) +"'",  # 78 setacciatura
                self.TABLE_FIELDS[78]:  "'" +str(self.comboBox_affidabilita.currentText()) +"'",  # 79 affidabilita
                self.TABLE_FIELDS[79]:  "'" +str(self.comboBox_direttore_us.currentText()) +"'",  # 80 direttore us
                self.TABLE_FIELDS[80]:  "'" +str(self.comboBox_responsabile_us.currentText())+"'" , # 81 responsabile us
                self.TABLE_FIELDS[81]:  "'" +str(self.lineEdit_cod_ente_schedatore.text()) +"'", # 82 cod ente schedatore
                self.TABLE_FIELDS[82]:  "'" +str(self.lineEdit_data_rilevazione.text())+"'" ,  # 83 data rilevazione
                self.TABLE_FIELDS[83]:  "'" +str(self.lineEdit_data_rielaborazione.text())+"'" , # 84 data rielaborazione
                self.TABLE_FIELDS[84]: lunghezza_usm,  # 85
                self.TABLE_FIELDS[85]: altezza_usm,  # 86
                self.TABLE_FIELDS[86]: spessore_usm,  # 87
                self.TABLE_FIELDS[87]:  "'" +str(self.comboBox_tecnica_muraria_usm.currentText())+"'" , # 88 tecnica muraria usm
                self.TABLE_FIELDS[88]:  "'" +str(self.comboBox_modulo_usm.currentText()) +"'", # 89 modulo usm
                self.TABLE_FIELDS[89]:  "'" +str(self.lineEdit_campioni_malta_usm.text())+"'" , # 90 campioni malta usm
                self.TABLE_FIELDS[90]:  "'" +str(self.lineEdit_campioni_mattone_usm.text())+"'" , # 91 campioni mattone usm
                self.TABLE_FIELDS[91]:  "'" +str(self.lineEdit_campioni_pietra_usm.text())+"'" , # 92 campioni pietra usm
                self.TABLE_FIELDS[92]:  "'" +str(self.lineEdit_provenienza_materiali_usm.text())+"'" , # 93 provenienza_materiali_usm
                self.TABLE_FIELDS[93]:  "'" +str(self.lineEdit_criteri_distinzione_usm.text()) +"'", # 94 criteri distinzione usm
                self.TABLE_FIELDS[94]:  "'" +str(self.comboBox_uso_primario_usm.currentText())+"'" ,  # 95 uso primario usm
                self.TABLE_FIELDS[95]:  "'" +str(self.comboBox_tipologia_opera.currentText())+"'" ,  # 95 uso primario usm
                self.TABLE_FIELDS[96]:  "'" +str(self.comboBox_sezione_muraria.currentText()) +"'" , # 95 uso primario usm
                self.TABLE_FIELDS[97]:  "'" +str(self.comboBox_superficie_analizzata.currentText()) +"'" , # 95 uso primario usm
                self.TABLE_FIELDS[98]:  "'" +str(self.comboBox_orientamento.currentText()) +"'" , # 95 uso primario usm
                self.TABLE_FIELDS[99]:  "'" +str(self.comboBox_materiali_lat.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[100]:  "'" +str(self.comboBox_lavorazione_lat.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[101]:  "'" +str(self.comboBox_consistenza_lat.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[102]:  "'" +str(self.comboBox_forma_lat.currentText())  +"'", # 95 uso primario usm
                self.TABLE_FIELDS[103]:  "'" +str(self.comboBox_colore_lat.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[104]:  "'" +str(self.comboBox_impasto_lat.currentText()) +"'", # 95 uso primario usm
                self.TABLE_FIELDS[105]:  "'" +str(self.comboBox_forma_p.currentText()) +"'" , # 95 uso primario usm
                self.TABLE_FIELDS[106]:  "'" +str(self.comboBox_colore_p.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[107]:  "'" +str(self.comboBox_taglio_p.currentText())  +"'", # 95 uso primario usm
                self.TABLE_FIELDS[108]:  "'" +str(self.comboBox_posa_opera_p.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[109]:  "'" +str(self.comboBox_inerti_usm.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[110]:  "'" +str(self.comboBox_tipo_legante_usm.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[111]:  "'" +str(self.comboBox_rifinitura_usm.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[112]:  "'" +str(self.comboBox_materiale_p.currentText()) +"'",  # 95 uso primario usm
                self.TABLE_FIELDS[113]:  "'" +str(self.comboBox_consistenza_p.currentText()) +"'",  # 95 uso primario usm
            }
            # Remove empty items
            u = Utility()
            search_dict = u.remove_empty_items_fr_dict(search_dict)
            search_dict_like = u.remove_empty_items_fr_dict(search_dict_like)

            if not search_dict:
                QMessageBox.warning(self, "WARNING", "No search is set", QMessageBox.Ok)
            else:
                if self.use_like_query:
                    message = "Scegli l'operatore di unione:\n\n" + \
                              "'and': restituisce solo i records dove tutte le condizioni sono vere. " + \
                              "Ad esempio, se stai cercando una 'struttura' in un 'settore' e un 'sito' specifici, scegli 'and'. " + \
                              "Se vuoi trovare la struttura 'ST1', nel settore 'SE2', del sito 'Monte Testaccio', otterrai records dove tutte queste condizioni sono soddisfatte.\n\n" + \
                              "'or': restituisce i records dove almeno una condizione è vera. " + \
                              "Ad esempio, se stai cercando records nella struttura 'ST1' o nel settore 'SE2' o nel sito 'Monte Testaccio', scegli 'or'. " + \
                              "In questo caso, otterrai i records dove una qualsiasi di queste condizioni è soddisfatta."

                    items = ['and', 'or']
                    join_operator, okPressed = QInputDialog.getItem(self, "Input", message, items, 0, False)
                    res = self.DB_MANAGER.query_bool_like(search_dict_like, 'us_table',join_operator)
                else:
                    res = self.DB_MANAGER.query_bool(search_dict, self.MAPPER_TABLE_CLASS)
                if not bool(res):
                    if self.L=='it':
                        QMessageBox.warning(self, "ATTENZIONE", "Non è stato trovato nessun record!", QMessageBox.Ok)
                    elif self.L=='de':
                        QMessageBox.warning(self, "ACHTUNG", "Keinen Record gefunden!", QMessageBox.Ok)
                    else:
                        QMessageBox.warning(self, "WARNING", "No record found!", QMessageBox.Ok)
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.setComboBoxEnable(["self.comboBox_sito"], "False")
                    self.setComboBoxEnable(["self.comboBox_area"], "False")
                    self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                    self.setComboBoxEnable(["self.lineEdit_us"], "False")
                    self.setComboBoxEnable(["self.textEdit_descrizione"], "True")
                    self.setComboBoxEnable(["self.textEdit_interpretazione"], "True")
                    self.setTableEnable(
                        ["self.tableWidget_campioni", "self.tableWidget_rapporti", "self.tableWidget_inclusi",
                         "self.tableWidget_organici", "self.tableWidget_inorganici", "self.tableWidget_documentazione","self.tableWidget_rapporti2"], "True")
                    self.fill_fields(self.REC_CORR)
                else:
                    self.DATA_LIST = []
                    for i  in res:
                        self.DATA_LIST.append(i)
                    self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                    self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                    self.fill_fields()
                    self.BROWSE_STATUS = "b"
                    self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                    self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                    if self.L=='it':
                        if self.REC_TOT == 1:
                            strings = ("E' stato trovato", self.REC_TOT, "record")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                        else:
                            strings = ("Sono stati trovati", self.REC_TOT, "records")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                    elif self.L=='de':
                        if self.REC_TOT == 1:
                            strings = ("Es wurde gefunden", self.REC_TOT, "record")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                        else:
                            strings = ("Sie wurden gefunden", self.REC_TOT, "records")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                    else:
                        if self.REC_TOT == 1:
                            strings = ("It has been found", self.REC_TOT, "record")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)
                        else:
                            strings = ("They have been found", self.REC_TOT, "records")
                            if self.toolButtonGis.isChecked():
                                self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                            if self.toolButton_usm.isChecked():
                                self.pyQGIS.charge_usm_layers(self.DATA_LIST)

                    self.setComboBoxEnable(["self.comboBox_sito"], "False")
                    self.setComboBoxEnable(["self.comboBox_area"], "False")
                    self.setComboBoxEnable(["self.lineEdit_us"], "False")
                    self.setComboBoxEnable(["self.comboBox_unita_tipo"], "True")
                    self.setTableEnable(
                        ["self.tableWidget_campioni",
                         "self.tableWidget_rapporti",
                         "self.tableWidget_inclusi",
                         "self.tableWidget_organici",
                         "self.tableWidget_inorganici",
                         "self.tableWidget_documentazione",
                         "self.tableWidget_inclusi_materiali_usm",
                         "self.tableWidget_colore_legante_usm",
                         "self.tableWidget_inclusi_leganti_usm",
                         "self.tableWidget_consistenza_texture_mat_usm",
                         "self.tableWidget_colore_materiale_usm","self.tableWidget_rapporti2"], "True")
                    self.setComboBoxEnable(["self.textEdit_descrizione"], "True")
                    self.setComboBoxEnable(["self.textEdit_interpretazione"], "True")
                    QMessageBox.warning(self, "Message", "%s %d %s" % strings, QMessageBox.Ok)
        self.use_like_query = False  # Reimposta il flag dopo la ricerca
        self.enable_button_search(1)

    def update_if(self, msg):
        # Save current record
        if msg == QMessageBox.Ok:
            test = self.update_record()
            if test == 1:

                # reload IDs
                id_list = []
                for i in self.DATA_LIST:
                    id_list.append(eval("i." + self.ID_TABLE))
                self.DATA_LIST = []

                # reload sorted data
                if self.SORT_STATUS == "n":
                    temp_data_list = self.DB_MANAGER.query_sort(id_list, [self.ID_TABLE], 'asc',
                                                                self.MAPPER_TABLE_CLASS,
                                                                self.ID_TABLE)
                else:
                    temp_data_list = self.DB_MANAGER.query_sort(id_list, self.SORT_ITEMS_CONVERTED, self.SORT_MODE,
                                                                self.MAPPER_TABLE_CLASS, self.ID_TABLE)
                for i in temp_data_list:
                    self.DATA_LIST.append(i)

                #
                # We skip the incrementing step
                #
                # check boundaries, if at end, loop to start
                # self.REC_CORR += 1
                # if self.REC_CORR >= len(self.DATA_LIST):
                #    self.REC_CORR = 0

                # other settings
                self.BROWSE_STATUS = "b"
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                self.label_sort.setText(self.SORTED_ITEMS[self.SORT_STATUS])
                self.set_rec_counter(self.REC_TOT, self.REC_CORR + 1)

                return 1
            #elif test
    def update_record(self):
        try:
            self.DB_MANAGER.update(self.MAPPER_TABLE_CLASS,
                                   self.ID_TABLE,
                                   [eval("int(self.DATA_LIST[self.REC_CORR]." + self.ID_TABLE + ")")],
                                   self.TABLE_FIELDS,
                                   self.rec_toupdate())
            return 1
        except Exception as e:
            str(e)
            save_file='{}{}{}'.format(self.HOME, os.sep,"pyarchinit_Report_folder")
            file_=os.path.join(save_file,'error_encodig_data_recover.txt')
            with open(file_, "a") as fh:
                try:
                    raise ValueError(str(e))
                except ValueError as s:
                    print(s, file=fh)
            if self.L=='it':
                QMessageBox.warning(self, "Messaggio",
                                    "Problema di encoding: sono stati inseriti accenti o caratteri non accettati dal database. Verrà fatta una copia dell'errore con i dati che puoi recuperare nella cartella pyarchinit_Report _Folder", QMessageBox.Ok)


            elif self.L=='en':
                QMessageBox.warning(self, "Message",
                                    "Encoding problem: accents or characters not accepted by the database were entered. A copy of the error will be made with the data you can retrieve in the pyarchinit_Report _Folder", QMessageBox.Ok)
            elif self.L=='de':
                QMessageBox.warning(self, "Message",
                                    "Kodierungsproblem: Es wurden Akzente oder Zeichen eingegeben, die von der Datenbank nicht akzeptiert werden. Es wird eine Kopie des Fehlers mit den Daten erstellt, die Sie im pyarchinit_Report _Ordner abrufen können", QMessageBox.Ok)
            else:
                QMessageBox.warning(self, "Messaggio",
                                    "Problema di encoding: sono stati inseriti accenti o caratteri non accettati dal database. Verrà fatta una copia dell'errore con i dati che puoi recuperare nella cartella pyarchinit_Report _Folder",
                                    QMessageBox.Ok)

            return 0
    def rec_toupdate(self):
        rec_to_update = self.UTILITY.pos_none_in_list(self.DATA_LIST_REC_TEMP)
        return rec_to_update
        # custom functions

    def charge_records(self):
        self.DATA_LIST = []
        if self.DB_SERVER == 'sqlite':
            for i in self.DB_MANAGER.query(self.MAPPER_TABLE_CLASS):
                self.DATA_LIST.append(i)
        else:
            id_list = []
            for i in self.DB_MANAGER.query(self.MAPPER_TABLE_CLASS):
                id_list.append(eval("i." + self.ID_TABLE))
            temp_data_list = self.DB_MANAGER.query_sort(id_list, [self.ID_TABLE], 'asc', self.MAPPER_TABLE_CLASS,
                                                        self.ID_TABLE)
            for i in temp_data_list:
                self.DATA_LIST.append(i)

    def charge_records_n(self):
        self.DATA_LIST = []
        if self.DB_SERVER == 'sqlite':
            for i in self.DB_MANAGER.query(self.MAPPER_TABLE_CLASS):
                self.DATA_LIST.append(i)
        else:
            id_list = []
            for i in self.DB_MANAGER.query(self.MAPPER_TABLE_CLASS):
                id_list.append(eval("i." + self.ID_TABLE))
            # Ordina in base a 'id_us' in ordine decrescente
            temp_data_list = self.DB_MANAGER.query_sort(id_list, [self.ID_TABLE], 'desc', self.MAPPER_TABLE_CLASS,
                                                        self.ID_TABLE)
            for i in temp_data_list:
                self.DATA_LIST.append(i)
            # Inverti la lista per mantenere l'ordine originale
            self.DATA_LIST.reverse()

    def datestrfdate(self):
        now = date.today()
        today = now.strftime("%d-%m-%Y")
        return today
    def yearstrfdate(self):
        now = date.today()
        year = now.strftime("%Y")
        return year
    def table2dict(self, n):
        self.tablename = n
        row = eval(self.tablename + ".rowCount()")
        col = eval(self.tablename + ".columnCount()")
        lista = []
        for r in range(row):
            sub_list = []
            for c in range(col):
                value = eval(self.tablename + ".item(r,c)")
                if value != None:
                    sub_list.append(str(value.text()))
            if bool(sub_list):
                lista.append(sub_list)
        return lista
    def tableInsertData(self, t, d):
        """Set the value into alls Grid"""
        self.table_name = t
        self.data_list = eval(d)
        self.data_list.sort()
        # column table count
        table_col_count_cmd = "{}.columnCount()".format(self.table_name)
        table_col_count = eval(table_col_count_cmd)
        # clear table
        table_clear_cmd = "{}.clearContents()".format(self.table_name)
        eval(table_clear_cmd)
        for i in range(table_col_count):
            table_rem_row_cmd = "{}.removeRow(int({}))".format(self.table_name, i)
            eval(table_rem_row_cmd)
        for row in range(len(self.data_list)):
            cmd = '{}.insertRow(int({}))'.format(self.table_name, row)
            eval(cmd)
            for col in range(len(self.data_list[row])):
                # item = self.comboBox_sito.setEditText(self.data_list[0][col]
                # item = QTableWidgetItem(self.data_list[row][col])
                # TODO SL: evauation of QTableWidget does not work porperly
                exec_str = '{}.setItem(int({}),int({}),QTableWidgetItem(self.data_list[row][col]))'.format(self.table_name, row, col)
                eval(exec_str)
        max_row_num = len(self.data_list)
        value = eval(self.table_name+".item(max_row_num,1)")
        if value == '':
            cmd = ("%s.removeRow(%d)") % (self.table_name, max_row_num)
            eval(cmd)
    def insert_new_row(self, table_name):
        """insert new row into a table based on table_name"""
        cmd = table_name + ".insertRow(0)"
        eval(cmd)
    def remove_row(self, table_name):
        """insert new row into a table based on table_name"""
        table_row_count_cmd = ("%s.rowCount()") % (table_name)
        table_row_count = eval(table_row_count_cmd)
        rowSelected_cmd = ("%s.selectedIndexes()") % (table_name)
        rowSelected = eval(rowSelected_cmd)
        rowIndex = (rowSelected[0].row())
        cmd = ("%s.removeRow(%d)") % (table_name, rowIndex)
        eval(cmd)
    def empty_fields(self):
        rapporti_row_count = self.tableWidget_rapporti.rowCount()
        rapporti_row_count2 = self.tableWidget_rapporti2.rowCount()
        campioni_row_count = self.tableWidget_campioni.rowCount()
        inclusi_row_count = self.tableWidget_inclusi.rowCount()
        organici_row_count = self.tableWidget_organici.rowCount()
        inorganici_row_count = self.tableWidget_inorganici.rowCount()
        documentazione_row_count = self.tableWidget_documentazione.rowCount()
        self.comboBox_sito.setEditText("")  # 1 - Sito
        self.comboBox_area.setEditText("")  # 2 - Area
        self.lineEdit_us.clear()  # 3 - US
        self.comboBox_def_strat.setEditText("")  # 4 - Definizione stratigrafica
        self.comboBox_def_intepret.setEditText("")  # 5 - Definizione intepretata
        self.textEdit_descrizione.clear()  # 6 - descrizione
        self.textEdit_interpretazione.clear()  # 7 - interpretazione
        self.comboBox_per_iniz.setEditText("")  # 8 - periodo iniziale
        self.comboBox_fas_iniz.setEditText("")  # 9 - fase iniziale
        self.comboBox_per_fin.setEditText("")  # 10 - periodo finale iniziale
        self.comboBox_fas_fin.setEditText("")  # 11 - fase finale
        self.comboBox_scavato.setEditText("")  # 12 - scavato
        self.lineEdit_attivita.clear()  # 13 - attivita
        if self.BROWSE_STATUS == "n":
            self.lineEdit_anno.setText(self.yearstrfdate())  # 14 - anno scavo
        else:
            self.lineEdit_anno.clear()
        self.comboBox_metodo.setEditText("")  # 15 - metodo
        for i in range(inclusi_row_count):
            self.tableWidget_inclusi.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi")  # 16 - inclusi
        for i in range(campioni_row_count):
            self.tableWidget_campioni.removeRow(0)
        self.insert_new_row("self.tableWidget_campioni")  # 17 - campioni
        for i in range(organici_row_count):
            self.tableWidget_organici.removeRow(0)
        self.insert_new_row("self.tableWidget_organici")  # organici
        for i in range(inorganici_row_count):
            self.tableWidget_inorganici.removeRow(0)
        self.insert_new_row("self.tableWidget_inorganici")  # inorganici
        for i in range(rapporti_row_count):
            self.tableWidget_rapporti.removeRow(0)
        self.insert_new_row("self.tableWidget_rapporti")                #18 - rapporti
        for i in range(documentazione_row_count):
            self.tableWidget_documentazione.removeRow(0)
        self.insert_new_row("self.tableWidget_documentazione")  # 19 - documentazione
        for i in range(rapporti_row_count2):
            self.tableWidget_rapporti2.removeRow(0)
        self.insert_new_row("self.tableWidget_rapporti2")

        colore_legante_usm_row_count = self.tableWidget_colore_legante_usm.rowCount()
        for i in range(colore_legante_usm_row_count):
            self.tableWidget_colore_legante_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_colore_legante_usm")  # 19 - aggregati
        inclusi_leganti_usm_row_count = self.tableWidget_inclusi_leganti_usm.rowCount()
        for i in range(inclusi_leganti_usm_row_count):
            self.tableWidget_inclusi_leganti_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi_leganti_usm")  # 19 - aggregati
        cont_text_mat_row_count = self.tableWidget_consistenza_texture_mat_usm.rowCount()
        for i in range(cont_text_mat_row_count):
            self.tableWidget_consistenza_texture_mat_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_consistenza_texture_mat_usm")  # 19 - colore legante usm
        aggreg_inclusi_materiale_row_count = self.tableWidget_inclusi_materiali_usm.rowCount()
        for i in range(aggreg_inclusi_materiale_row_count):
            self.tableWidget_inclusi_materiali_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi_materiali_usm")  # 19 - aggregati
        colore_materiali_usm_row_count = self.tableWidget_colore_materiale_usm.rowCount()
        for i in range(colore_materiali_usm_row_count):
            self.tableWidget_colore_materiale_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_colore_materiale_usm")  # 19 - aggregati
        if self.BROWSE_STATUS == "n":
            self.lineEdit_data_schedatura.setText(self.datestrfdate())  # 20 - data schedatura
        else:
            self.lineEdit_data_schedatura.setText("")  # 20 - data schedatura

        self.comboBox_schedatore.setEditText("")  # 21 - schedatore
        self.comboBox_formazione.setEditText("")  # 22 - formazione
        self.comboBox_conservazione.setEditText("")  # 23 - conservazione
        self.comboBox_colore.setEditText("")  # 24 - colore
        self.comboBox_consistenza.setEditText("")  # 25 - consistenza
        self.comboBox_struttura.setEditText("")  # 26 - struttura
        self.lineEdit_codice_periodo.clear()  # 27 - codice periodo
        self.lineEditOrderLayer.clear()  # 28 - order layer
        self.comboBox_unita_tipo.setEditText("")  # 29 us_tipo            NUOVI CAMPI NUOVI CAMPI
        self.comboBox_settore.setEditText("")  # 30 settore
        self.lineEdit_quadrato.clear()  # 31 quadrato
        self.lineEdit_ambiente.clear()  # 32 ambiente
        self.lineEdit_saggio.clear()  # 33 saggio
        self.textEdit_elementi_datanti.clear()  # 34 elementi datanti
        self.comboBox_funz_statica_usm.setEditText("")  # 35 funzione statica
        self.comboBox_lavorazione_usm.setEditText("")  # 36 lavorazione usm
        self.lineEdit_spessore_giunti_usm.clear()  # 37 spessore giunti
        self.lineEdit_letti_di_posa_giunti_usm.clear()  # 38 letti posa giunti usm
        self.lineEdit_h_modulo_c_corsi_usm.clear()  # 39 altezza modulo corsi usm
        self.comboBox_unita_edilizia_riassuntiva_usm.setEditText("")  # 40 unita edilizia riassuntiva
        self.comboBox_reimpiego_usm.setEditText("")  # 41 unita edilizia riassuntiva
        self.comboBox_posa_in_opera_usm.setEditText("")  # 42 posa in opera
        self.lineEdit_qmin_usm.clear()  # 3 - US
        self.lineEdit_qmax_usm.clear()  # 3 - US
        self.comboBox_consistenza_legante_usm.setEditText("")  # 45 consitenza legante usm
        self.lineEdit_n_catalogo_generale.clear()  # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
        self.lineEdit_n_catalogo_interno.clear()  # 52 nr catalogo interno
        self.lineEdit_n_catalogo_internazionale.clear()  # 53 nr catalogo internazionale
        self.comboBox_soprintendenza.setEditText("")  # 54 nr soprintendenza
        self.lineEdit_quota_relativa.clear()  # 55
        self.lineEdit_quota_abs.clear()  # 56
        self.lineEdit_ref_tm.clear()  # 57 ref tm
        self.comboBox_ref_ra.setEditText("")   # 58 ref ra
        self.lineEdit_ref_n.clear()  # 59 ref n
        self.comboBox_posizione.setEditText("")  # 60 posizione
        self.lineEdit_criteri_distinzione.clear()  # 61 criteri distinzione
        self.comboBox_modo_formazione.setEditText("")  # 62 modo formazione
        #self.comboBox_componenti_organici.setEditText("")  # 63 componenti organici
        #self.comboBox_componenti_inorganici.setEditText("")  # 64 componenti inorganici
        self.lineEdit_lunghezza_max.clear()  # 65
        self.lineEdit_altezza_max.clear()  # 66
        self.lineEdit_altezza_min.clear()  # 67
        self.lineEdit_profondita_max.clear()  # 68
        self.lineEdit_profondita_min.clear()  # 69
        self.lineEdit_larghezza_media.clear()  # 70
        self.lineEdit_quota_max_abs.clear()  # 71
        self.lineEdit_quota_max_rel.clear()  # 72
        self.lineEdit_quota_min_abs.clear()  # 73
        self.lineEdit_quota_min_rel.clear()  # 74
        self.textEdit_osservazioni.clear()  # 75 osservazioni
        self.lineEdit_datazione.clear()  # 76 datazione
        self.comboBox_flottazione.setEditText("")  # 77 flottazione
        self.comboBox_setacciatura.setEditText("")   # 78 setacciatura
        self.comboBox_affidabilita.setEditText("")   # 79 affidabilita
        self.comboBox_direttore_us.setEditText("")  # 80 direttore us
        self.comboBox_responsabile_us.setEditText("")  # 81 responsabile us
        self.lineEdit_cod_ente_schedatore.clear()  # 82 cod ente schedatore
        self.lineEdit_data_rilevazione.clear()  # 83 data rilevazione
        self.lineEdit_data_rielaborazione.clear()  # 84 data rielaborazione
        self.lineEdit_lunghezza_usm.clear()  # 85
        self.lineEdit_altezza_usm.clear()  # 86
        self.lineEdit_spessore_usm.clear()  # 87
        self.comboBox_tecnica_muraria_usm.setEditText("")  # 88 tecnica muraria usm
        self.comboBox_modulo_usm.setEditText("")  # 89 modulo usm
        self.lineEdit_campioni_malta_usm.clear()  # 90 campioni malta usm
        self.lineEdit_campioni_mattone_usm.clear()  # 91 campioni mattone usm
        self.lineEdit_campioni_pietra_usm.clear()  # 92 campioni pietra usm
        self.lineEdit_provenienza_materiali_usm.clear()  # 93 provenienza_materiali_usm
        self.lineEdit_criteri_distinzione_usm.clear()  # 94 criteri distinzione usm
        self.comboBox_uso_primario_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_tipologia_opera.setEditText("")  # 95 uso primario usm
        self.comboBox_sezione_muraria.setEditText("")  # 95 uso primario usm
        self.comboBox_superficie_analizzata.setEditText("")  # 95 uso primario usm
        self.comboBox_orientamento.setEditText("")  # 95 uso primario usm
        self.comboBox_materiali_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_lavorazione_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_consistenza_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_forma_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_colore_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_impasto_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_forma_p.setEditText("")  # 95 uso primario usm
        self.comboBox_colore_p.setEditText("")  # 95 uso primario usm
        self.comboBox_taglio_p.setEditText("")  # 95 uso primario usm
        self.comboBox_posa_opera_p.setEditText("")  # 95 uso primario usm
        self.comboBox_inerti_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_tipo_legante_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_rifinitura_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_materiale_p.setEditText("")  # 95 uso primario usm
        self.comboBox_consistenza_p.setEditText("")  # 95 uso primario usm
        self.mQgsFileWidget.clear()
    def empty_fields_nosite(self):
        rapporti_row_count = self.tableWidget_rapporti.rowCount()
        rapporti_row_count2 = self.tableWidget_rapporti2.rowCount()
        campioni_row_count = self.tableWidget_campioni.rowCount()
        inclusi_row_count = self.tableWidget_inclusi.rowCount()
        organici_row_count = self.tableWidget_organici.rowCount()
        inorganici_row_count = self.tableWidget_inorganici.rowCount()
        documentazione_row_count = self.tableWidget_documentazione.rowCount()

        self.comboBox_area.setEditText("")  # 2 - Area
        self.lineEdit_us.clear()  # 3 - US
        self.comboBox_def_strat.setEditText("")  # 4 - Definizione stratigrafica
        self.comboBox_def_intepret.setEditText("")  # 5 - Definizione intepretata
        self.textEdit_descrizione.clear()  # 6 - descrizione
        self.textEdit_interpretazione.clear()  # 7 - interpretazione
        self.comboBox_per_iniz.setEditText("")  # 8 - periodo iniziale
        self.comboBox_fas_iniz.setEditText("")  # 9 - fase iniziale
        self.comboBox_per_fin.setEditText("")  # 10 - periodo finale iniziale
        self.comboBox_fas_fin.setEditText("")  # 11 - fase finale
        self.comboBox_scavato.setEditText("")  # 12 - scavato
        self.lineEdit_attivita.clear()  # 13 - attivita
        if self.BROWSE_STATUS == "n":
            self.lineEdit_anno.setText(self.yearstrfdate())  # 14 - anno scavo
        else:
            self.lineEdit_anno.clear()
        self.comboBox_metodo.setEditText("")  # 15 - metodo
        for i in range(inclusi_row_count):
            self.tableWidget_inclusi.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi")  # 16 - inclusi
        for i in range(campioni_row_count):
            self.tableWidget_campioni.removeRow(0)
        self.insert_new_row("self.tableWidget_campioni")  # 17 - campioni
        for i in range(organici_row_count):
            self.tableWidget_organici.removeRow(0)
        self.insert_new_row("self.tableWidget_organici")  # organici
        for i in range(inorganici_row_count):
            self.tableWidget_inorganici.removeRow(0)
        self.insert_new_row("self.tableWidget_inorganici")  # inorganici
        for i in range(rapporti_row_count):
            self.tableWidget_rapporti.removeRow(0)
        self.insert_new_row("self.tableWidget_rapporti")                #18 - rapporti
        for i in range(documentazione_row_count):
            self.tableWidget_documentazione.removeRow(0)
        for i in range(rapporti_row_count2):
            self.tableWidget_rapporti2.removeRow(0)
        self.insert_new_row("self.tableWidget_rapporti2")

        self.insert_new_row("self.tableWidget_documentazione")  # 19 - documentazione
        colore_legante_usm_row_count = self.tableWidget_colore_legante_usm.rowCount()
        for i in range(colore_legante_usm_row_count):
            self.tableWidget_colore_legante_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_colore_legante_usm")  # 19 - aggregati
        inclusi_leganti_usm_row_count = self.tableWidget_inclusi_leganti_usm.rowCount()
        for i in range(inclusi_leganti_usm_row_count):
            self.tableWidget_inclusi_leganti_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi_leganti_usm")  # 19 - aggregati
        cont_text_mat_row_count = self.tableWidget_consistenza_texture_mat_usm.rowCount()
        for i in range(cont_text_mat_row_count):
            self.tableWidget_consistenza_texture_mat_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_consistenza_texture_mat_usm")  # 19 - colore legante usm
        aggreg_inclusi_materiale_row_count = self.tableWidget_inclusi_materiali_usm.rowCount()
        for i in range(aggreg_inclusi_materiale_row_count):
            self.tableWidget_inclusi_materiali_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_inclusi_materiali_usm")  # 19 - aggregati
        colore_materiali_usm_row_count = self.tableWidget_colore_materiale_usm.rowCount()
        for i in range(colore_materiali_usm_row_count):
            self.tableWidget_colore_materiale_usm.removeRow(0)
        self.insert_new_row("self.tableWidget_colore_materiale_usm")  # 19 - aggregati
        if self.BROWSE_STATUS == "n":
            self.lineEdit_data_schedatura.setText(self.datestrfdate())  # 20 - data schedatura
        else:
            self.lineEdit_data_schedatura.setText("")  # 20 - data schedatura

        self.comboBox_schedatore.setEditText("")  # 21 - schedatore
        self.comboBox_formazione.setEditText("")  # 22 - formazione
        self.comboBox_conservazione.setEditText("")  # 23 - conservazione
        self.comboBox_colore.setEditText("")  # 24 - colore
        self.comboBox_consistenza.setEditText("")  # 25 - consistenza
        self.comboBox_struttura.setEditText("")  # 26 - struttura
        self.lineEdit_codice_periodo.clear()  # 27 - codice periodo
        self.lineEditOrderLayer.clear()  # 28 - order layer
        self.comboBox_unita_tipo.setEditText("")  # 29 us_tipo            NUOVI CAMPI NUOVI CAMPI
        self.comboBox_settore.setEditText("")  # 30 settore
        self.lineEdit_quadrato.clear()  # 31 quadrato
        self.lineEdit_ambiente.clear()  # 32 ambiente
        self.lineEdit_saggio.clear()  # 33 saggio
        self.textEdit_elementi_datanti.clear()  # 34 elementi datanti
        self.comboBox_funz_statica_usm.setEditText("")  # 35 funzione statica
        self.comboBox_lavorazione_usm.setEditText("")  # 36 lavorazione usm
        self.lineEdit_spessore_giunti_usm.clear()  # 37 spessore giunti
        self.lineEdit_letti_di_posa_giunti_usm.clear()  # 38 letti posa giunti usm
        self.lineEdit_h_modulo_c_corsi_usm.clear()  # 39 altezza modulo corsi usm
        self.comboBox_unita_edilizia_riassuntiva_usm.setEditText("")  # 40 unita edilizia riassuntiva
        self.comboBox_reimpiego_usm.setEditText("")  # 41 unita edilizia riassuntiva
        self.comboBox_posa_in_opera_usm.setEditText("")  # 42 posa in opera
        self.lineEdit_qmin_usm.clear()  # 3 - US
        self.lineEdit_qmax_usm.clear()  # 3 - US
        self.comboBox_consistenza_legante_usm.setEditText("")  # 45 consitenza legante usm
        self.lineEdit_n_catalogo_generale.clear()  # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
        self.lineEdit_n_catalogo_interno.clear()  # 52 nr catalogo interno
        self.lineEdit_n_catalogo_internazionale.clear()  # 53 nr catalogo internazionale
        self.comboBox_soprintendenza.setEditText("")  # 54 nr soprintendenza
        self.lineEdit_quota_relativa.clear()  # 55
        self.lineEdit_quota_abs.clear()  # 56
        self.lineEdit_ref_tm.clear()  # 57 ref tm
        self.comboBox_ref_ra.setEditText("")   # 58 ref ra
        self.lineEdit_ref_n.clear()  # 59 ref n
        self.comboBox_posizione.setEditText("")  # 60 posizione
        self.lineEdit_criteri_distinzione.clear()  # 61 criteri distinzione
        self.comboBox_modo_formazione.setEditText("")  # 62 modo formazione
        #self.comboBox_componenti_organici.setEditText("")  # 63 componenti organici
        #self.comboBox_componenti_inorganici.setEditText("")  # 64 componenti inorganici
        self.lineEdit_lunghezza_max.clear()  # 65
        self.lineEdit_altezza_max.clear()  # 66
        self.lineEdit_altezza_min.clear()  # 67
        self.lineEdit_profondita_max.clear()  # 68
        self.lineEdit_profondita_min.clear()  # 69
        self.lineEdit_larghezza_media.clear()  # 70
        self.lineEdit_quota_max_abs.clear()  # 71
        self.lineEdit_quota_max_rel.clear()  # 72
        self.lineEdit_quota_min_abs.clear()  # 73
        self.lineEdit_quota_min_rel.clear()  # 74
        self.textEdit_osservazioni.clear()  # 75 osservazioni
        self.lineEdit_datazione.clear()  # 76 datazione
        self.comboBox_flottazione.setEditText("")   # 77 flottazione
        self.comboBox_setacciatura.setEditText("")  # 78 setacciatura
        self.comboBox_affidabilita.setEditText("")  # 79 affidabilita
        self.comboBox_direttore_us.setEditText("")  # 80 direttore us
        self.comboBox_responsabile_us.setEditText("")  # 81 responsabile us
        self.lineEdit_cod_ente_schedatore.clear()  # 82 cod ente schedatore
        self.lineEdit_data_rilevazione.clear()  # 83 data rilevazione
        self.lineEdit_data_rielaborazione.clear()  # 84 data rielaborazione
        self.lineEdit_lunghezza_usm.clear()  # 85
        self.lineEdit_altezza_usm.clear()  # 86
        self.lineEdit_spessore_usm.clear()  # 87
        self.comboBox_tecnica_muraria_usm.setEditText("")  # 88 tecnica muraria usm
        self.comboBox_modulo_usm.setEditText("")  # 89 modulo usm
        self.lineEdit_campioni_malta_usm.clear()  # 90 campioni malta usm
        self.lineEdit_campioni_mattone_usm.clear()  # 91 campioni mattone usm
        self.lineEdit_campioni_pietra_usm.clear()  # 92 campioni pietra usm
        self.lineEdit_provenienza_materiali_usm.clear()  # 93 provenienza_materiali_usm
        self.lineEdit_criteri_distinzione_usm.clear()  # 94 criteri distinzione usm
        self.comboBox_uso_primario_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_tipologia_opera.setEditText("")  # 95 uso primario usm
        self.comboBox_sezione_muraria.setEditText("")  # 95 uso primario usm
        self.comboBox_superficie_analizzata.setEditText("")  # 95 uso primario usm
        self.comboBox_orientamento.setEditText("")  # 95 uso primario usm
        self.comboBox_materiali_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_lavorazione_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_consistenza_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_forma_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_colore_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_impasto_lat.setEditText("")  # 95 uso primario usm
        self.comboBox_forma_p.setEditText("")  # 95 uso primario usm
        self.comboBox_colore_p.setEditText("")  # 95 uso primario usm
        self.comboBox_taglio_p.setEditText("")  # 95 uso primario usm
        self.comboBox_posa_opera_p.setEditText("")  # 95 uso primario usm
        self.comboBox_inerti_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_tipo_legante_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_rifinitura_usm.setEditText("")  # 95 uso primario usm
        self.comboBox_materiale_p.setEditText("")  # 95 uso primario usm
        self.comboBox_consistenza_p.setEditText("")  # 95 uso primario usm
        self.mQgsFileWidget.clear()
    def fill_fields(self, n=0):
        self.rec_num = n
        try:
            str(self.comboBox_sito.setEditText(self.DATA_LIST[self.rec_num].sito))  # 1 - Sito
            str(self.comboBox_area.setEditText(self.DATA_LIST[self.rec_num].area))  # 2 - Area
            self.lineEdit_us.setText(str(self.DATA_LIST[self.rec_num].us))  # 3 - US
            str(self.comboBox_def_strat.setEditText(self.DATA_LIST[self.rec_num].d_stratigrafica))  # 4 - Definizione stratigrafica
            str(self.comboBox_def_intepret.setEditText(self.DATA_LIST[self.rec_num].d_interpretativa))  # 5 - Definizione intepretata
            str(self.textEdit_descrizione.setText(self.DATA_LIST[self.rec_num].descrizione))  # 6 - descrizione
            str(self.textEdit_interpretazione.setText(self.DATA_LIST[self.rec_num].interpretazione))  # 7 - interpretazione
            str(self.comboBox_per_iniz.setEditText(self.DATA_LIST[self.rec_num].periodo_iniziale))  # 8 - periodo iniziale
            str(self.comboBox_fas_iniz.setEditText(self.DATA_LIST[self.rec_num].fase_iniziale))  # 9 - fase iniziale
            str(self.comboBox_per_fin.setEditText(self.DATA_LIST[self.rec_num].periodo_finale))  # 10 - periodo finale iniziale
            str(self.comboBox_fas_fin.setEditText(self.DATA_LIST[self.rec_num].fase_finale))  # 11 - fase finale
            str(self.comboBox_scavato.setEditText(self.DATA_LIST[self.rec_num].scavato))  # 12 - scavato
            str(self.lineEdit_attivita.setText(self.DATA_LIST[self.rec_num].attivita))  # 13 - attivita
            str(self.lineEdit_anno.setText(self.DATA_LIST[self.rec_num].anno_scavo))  # 14 - anno scavo
            str(self.comboBox_metodo.setEditText(self.DATA_LIST[self.rec_num].metodo_di_scavo))  # 15 - metodo
            self.tableInsertData("self.tableWidget_inclusi", self.DATA_LIST[self.rec_num].inclusi)  # 16 - inclusi
            self.tableInsertData("self.tableWidget_campioni", self.DATA_LIST[self.rec_num].campioni)  # 17 - campioni
            self.tableInsertData("self.tableWidget_organici", self.DATA_LIST[self.rec_num].componenti_organici)  # organici
            self.tableInsertData("self.tableWidget_inorganici", self.DATA_LIST[self.rec_num].componenti_inorganici)  # inorganici
            self.tableInsertData("self.tableWidget_rapporti", self.DATA_LIST[self.rec_num].rapporti)  # 18 - rapporti
            str(self.lineEdit_data_schedatura.setText(self.DATA_LIST[self.rec_num].data_schedatura))  # 19 - data schedatura
            str(self.comboBox_schedatore.setEditText(self.DATA_LIST[self.rec_num].schedatore))  # 20 - schedatore
            str(self.comboBox_formazione.setEditText(self.DATA_LIST[self.rec_num].formazione))  # 21 - formazione
            str(self.comboBox_conservazione.setEditText(self.DATA_LIST[self.rec_num].stato_di_conservazione))  # 22 - conservazione
            str(self.comboBox_colore.setEditText(self.DATA_LIST[self.rec_num].colore))  # 23 - colore
            str(self.comboBox_consistenza.setEditText(self.DATA_LIST[self.rec_num].consistenza))  # 24 - consistenza
            str(self.comboBox_struttura.setDefaultText(self.DATA_LIST[self.rec_num].struttura)) # 25 - struttura
            if not self.DATA_LIST[self.rec_num].cont_per:
                self.lineEdit_codice_periodo.setText("")
            else:
                self.lineEdit_codice_periodo.setText(str(self.DATA_LIST[self.rec_num].cont_per))  # 26 - codice periodo
            if not self.DATA_LIST[self.rec_num].order_layer:
                self.lineEditOrderLayer.setText("")
            else:
                self.lineEditOrderLayer.setText(str(self.DATA_LIST[self.rec_num].order_layer))  # 27 - order layer
            self.tableInsertData("self.tableWidget_documentazione",self.DATA_LIST[self.rec_num].documentazione)  # 28 - documentazione
            str(self.comboBox_unita_tipo.setEditText(self.DATA_LIST[self.rec_num].unita_tipo))  # 29 unita tipo
            str(self.comboBox_settore.setEditText(self.DATA_LIST[self.rec_num].settore))  # 30 - settore
            str(self.lineEdit_quadrato.setText(self.DATA_LIST[self.rec_num].quad_par))  # 31 quadrato
            str(self.lineEdit_ambiente.setText(self.DATA_LIST[self.rec_num].ambient))  # 32 ambiente
            str(self.lineEdit_saggio.setText(self.DATA_LIST[self.rec_num].saggio))  # 33 saggio
            str(self.textEdit_elementi_datanti.setText(self.DATA_LIST[self.rec_num].elem_datanti))  # 34 - elemtenti_datanti
            str(self.comboBox_funz_statica_usm.setEditText(self.DATA_LIST[self.rec_num].funz_statica))  # 35 - funz statica
            str(self.comboBox_lavorazione_usm.setEditText(self.DATA_LIST[self.rec_num].lavorazione))  # 36 lavorazione usm
            str(self.lineEdit_spessore_giunti_usm.setText(self.DATA_LIST[self.rec_num].spess_giunti))  # 37 spessore giunti usm
            str(self.lineEdit_letti_di_posa_giunti_usm.setText(self.DATA_LIST[self.rec_num].letti_posa)) #38 letti_posa
            str(self.lineEdit_h_modulo_c_corsi_usm.setText(self.DATA_LIST[self.rec_num].alt_mod)) #39 altezza modulo corsi
            str(self.comboBox_unita_edilizia_riassuntiva_usm.setEditText(self.DATA_LIST[self.rec_num].un_ed_riass)) #40 unita edilizia riassuntiva
            str(self.comboBox_reimpiego_usm.setEditText(self.DATA_LIST[self.rec_num].reimp))  #41 reimpiego
            str(self.comboBox_posa_in_opera_usm.setEditText(self.DATA_LIST[self.rec_num].posa_opera)) #42 posa opera
            if not self.DATA_LIST[self.rec_num].quota_min_usm:
                str(self.lineEdit_qmin_usm.setText(""))
            else:
                self.lineEdit_qmin_usm.setText(str(self.DATA_LIST[self.rec_num].quota_min_usm))  # 43 - qmin usm
            if not self.DATA_LIST[self.rec_num].quota_max_usm:
               str(self.lineEdit_qmax_usm.setText(""))
            else:
               self.lineEdit_qmax_usm.setText(str(self.DATA_LIST[self.rec_num].quota_max_usm))  # 44 - qmax usm
            str(self.comboBox_consistenza_legante_usm.setEditText(self.DATA_LIST[self.rec_num].cons_legante))  # 45 - cons legante
            self.tableInsertData("self.tableWidget_colore_legante_usm", self.DATA_LIST[self.rec_num].col_legante) ## 46 - col legante usm
            self.tableInsertData("self.tableWidget_inclusi_leganti_usm", self.DATA_LIST[self.rec_num].aggreg_legante) # 47 aggregati legante usm
            self.tableInsertData("self.tableWidget_consistenza_texture_mat_usm", self.DATA_LIST[self.rec_num].con_text_mat) # 48 - con text mat
            self.tableInsertData("self.tableWidget_colore_materiale_usm", self.DATA_LIST[self.rec_num].col_materiale) # 49 - col mat
            self.tableInsertData("self.tableWidget_inclusi_materiali_usm",self.DATA_LIST[self.rec_num].inclusi_materiali_usm)  # 50  inclusi materiali usm
            str(self.lineEdit_n_catalogo_generale.setText(self.DATA_LIST[self.rec_num].n_catalogo_generale))  # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
            str(self.lineEdit_n_catalogo_interno.setText(self.DATA_LIST[self.rec_num].n_catalogo_interno))  # 52 nr catalogo interno
            str(self.lineEdit_n_catalogo_internazionale.setText(self.DATA_LIST[self.rec_num].n_catalogo_internazionale))  # 53 nr catalogo internazionale
            str(self.comboBox_soprintendenza.setEditText(self.DATA_LIST[self.rec_num].soprintendenza))  # 54 nr soprintendenza
            if not self.DATA_LIST[self.rec_num].quota_relativa:
                str(self.lineEdit_quota_relativa.setText(""))                   # 55
            else:
                self.lineEdit_quota_relativa.setText(str(self.DATA_LIST[self.rec_num].quota_relativa))
            if not self.DATA_LIST[self.rec_num].quota_abs:
                str(self.lineEdit_quota_abs.setText(""))                   # 56
            else:
                self.lineEdit_quota_abs.setText(str(self.DATA_LIST[self.rec_num].quota_abs))
            str(self.lineEdit_ref_tm.setText(self.DATA_LIST[self.rec_num].ref_tm))  # 57 ref tm
            str(self.comboBox_ref_ra.setDefaultText(self.DATA_LIST[self.rec_num].ref_ra))  # 58 ref ra
            str(self.lineEdit_ref_n.setText(self.DATA_LIST[self.rec_num].ref_n))  # 59 ref n
            str(self.comboBox_posizione.setEditText(self.DATA_LIST[self.rec_num].posizione))  # 60 posizione
            str(self.lineEdit_criteri_distinzione.setText(self.DATA_LIST[self.rec_num].criteri_distinzione))  # 61 criteri distinzione
            str(self.comboBox_modo_formazione.setEditText(self.DATA_LIST[self.rec_num].modo_formazione))  # 62 modo formazione
            if not self.DATA_LIST[self.rec_num].lunghezza_max:
                str(self.lineEdit_lunghezza_max.setText(""))
            else:
                self.lineEdit_lunghezza_max.setText(str(self.DATA_LIST[self.rec_num].lunghezza_max))  # 65 lunghezza max
            if not self.DATA_LIST[self.rec_num].altezza_max:
                str(self.lineEdit_altezza_max.setText(""))
            else:
                self.lineEdit_altezza_max.setText(str(self.DATA_LIST[self.rec_num].altezza_max))  # 66 altezza max
            if not self.DATA_LIST[self.rec_num].altezza_min:
                str(self.lineEdit_altezza_min.setText(""))
            else:
                self.lineEdit_altezza_min.setText(str(self.DATA_LIST[self.rec_num].altezza_min))  # 67 altezza min
            if not self.DATA_LIST[self.rec_num].profondita_max:
                str(self.lineEdit_profondita_max.setText(""))
            else:
                self.lineEdit_profondita_max.setText(str(
                    self.DATA_LIST[self.rec_num].profondita_max))  # 68 profondita_max
            if not self.DATA_LIST[self.rec_num].profondita_min:
                str(self.lineEdit_profondita_min.setText(""))
            else:
                self.lineEdit_profondita_min.setText(str(
                    self.DATA_LIST[self.rec_num].profondita_min))  # 69 profondita min
            if not self.DATA_LIST[self.rec_num].larghezza_media:
                str(self.lineEdit_larghezza_media.setText(""))
            else:
                self.lineEdit_larghezza_media.setText(str(
                    self.DATA_LIST[self.rec_num].larghezza_media))  # 70 larghezza media
            if not self.DATA_LIST[self.rec_num].quota_max_abs:
                str(self.lineEdit_quota_max_abs.setText(""))
            else:
                self.lineEdit_quota_max_abs.setText(str(self.DATA_LIST[self.rec_num].quota_max_abs))  # 71 quota_max_abs
            if not self.DATA_LIST[self.rec_num].quota_max_rel:
                str(self.lineEdit_quota_max_rel.setText(""))
            else:
                self.lineEdit_quota_max_rel.setText(str(
                    self.DATA_LIST[self.rec_num].quota_max_rel))  # 72 quota_max_rel
            if not self.DATA_LIST[self.rec_num].quota_min_abs:
                str(self.lineEdit_quota_min_abs.setText(""))
            else:
                self.lineEdit_quota_min_abs.setText(str(self.DATA_LIST[self.rec_num].quota_min_abs))  # 73 quota_min_abs
            if not self.DATA_LIST[self.rec_num].quota_min_rel:
                str(self.lineEdit_quota_min_rel.setText(""))
            else:
                self.lineEdit_quota_min_rel.setText(str(self.DATA_LIST[self.rec_num].quota_min_rel))  # 74 quota_min_rel
            str(self.textEdit_osservazioni.setText(self.DATA_LIST[self.rec_num].osservazioni))  # 75 osservazioni
            str(self.lineEdit_datazione.setText(self.DATA_LIST[self.rec_num].datazione))  # 76 datazione
            str(self.comboBox_flottazione.setEditText(self.DATA_LIST[self.rec_num].flottazione))  # 77 flottazione
            str(self.comboBox_setacciatura.setEditText(self.DATA_LIST[self.rec_num].setacciatura))  # 78 setacciatura
            str(self.comboBox_affidabilita.setEditText(self.DATA_LIST[self.rec_num].affidabilita))        # 79 affidabilita
            str(self.comboBox_direttore_us.setEditText(self.DATA_LIST[self.rec_num].direttore_us))  # 80 direttore us
            str(self.comboBox_responsabile_us.setEditText(self.DATA_LIST[self.rec_num].responsabile_us))  # 81 responsabile us
            str(self.lineEdit_cod_ente_schedatore.setText(self.DATA_LIST[self.rec_num].cod_ente_schedatore))  # 82 cod ente schedatore
            str(self.lineEdit_data_rilevazione.setText(self.DATA_LIST[self.rec_num].data_rilevazione))  # 83 data rilevazione
            str(self.lineEdit_data_rielaborazione.setText(self.DATA_LIST[self.rec_num].data_rielaborazione))  # 84 data rielaborazione
            if not self.DATA_LIST[self.rec_num].lunghezza_usm:
                str(self.lineEdit_lunghezza_usm.setText(""))
            else:
                self.lineEdit_lunghezza_usm.setText(str(self.DATA_LIST[self.rec_num].lunghezza_usm))  # 85 lunghezza usm
            if not self.DATA_LIST[self.rec_num].altezza_usm:
                str(self.lineEdit_altezza_usm.setText(""))
            else:
                self.lineEdit_altezza_usm.setText(str(self.DATA_LIST[self.rec_num].altezza_usm))  # 86 altezza usm
            if not self.DATA_LIST[self.rec_num].spessore_usm:
                str(self.lineEdit_spessore_usm.setText(""))
            else:
                self.lineEdit_spessore_usm.setText(str(self.DATA_LIST[self.rec_num].spessore_usm))  # 87 spessore usm
            str(self.comboBox_tecnica_muraria_usm.setEditText(self.DATA_LIST[self.rec_num].tecnica_muraria_usm))  # 88 tecnica muraria usm
            str(self.comboBox_modulo_usm.setEditText(self.DATA_LIST[self.rec_num].modulo_usm))  # 89 modulo usm
            str(self.lineEdit_campioni_malta_usm.setText(self.DATA_LIST[self.rec_num].campioni_malta_usm))  # 90 campioni malta usm
            str(self.lineEdit_campioni_mattone_usm.setText(self.DATA_LIST[self.rec_num].campioni_mattone_usm))  # 91 campioni mattone usm
            str(self.lineEdit_campioni_pietra_usm.setText(self.DATA_LIST[self.rec_num].campioni_pietra_usm))  # 92 campioni pietra usm
            str(self.lineEdit_provenienza_materiali_usm.setText(self.DATA_LIST[self.rec_num].provenienza_materiali_usm))  # 93 provenienza_materiali_usm
            str(self.lineEdit_criteri_distinzione_usm.setText(self.DATA_LIST[self.rec_num].criteri_distinzione_usm))  # 94 criteri distinzione usm
            str(self.comboBox_uso_primario_usm.setEditText(self.DATA_LIST[self.rec_num].uso_primario_usm))  # 95 uso primario usm
            str(self.comboBox_tipologia_opera.setEditText(self.DATA_LIST[self.rec_num].tipologia_opera))
            str(self.comboBox_sezione_muraria.setEditText(self.DATA_LIST[self.rec_num].sezione_muraria))
            str(self.comboBox_superficie_analizzata.setEditText(self.DATA_LIST[self.rec_num].superficie_analizzata))
            str(self.comboBox_orientamento.setEditText(self.DATA_LIST[self.rec_num].orientamento))
            str(self.comboBox_materiali_lat.setEditText(self.DATA_LIST[self.rec_num].materiali_lat))
            str(self.comboBox_lavorazione_lat.setEditText(self.DATA_LIST[self.rec_num].lavorazione_lat))
            str(self.comboBox_consistenza_lat.setEditText(self.DATA_LIST[self.rec_num].consistenza_lat))
            str(self.comboBox_forma_lat.setEditText(self.DATA_LIST[self.rec_num].forma_lat))
            str(self.comboBox_colore_lat.setEditText(self.DATA_LIST[self.rec_num].colore_lat))
            str(self.comboBox_impasto_lat.setEditText(self.DATA_LIST[self.rec_num].impasto_lat))
            str(self.comboBox_forma_p.setEditText(self.DATA_LIST[self.rec_num].forma_p))
            str(self.comboBox_colore_p.setEditText(self.DATA_LIST[self.rec_num].colore_p))
            str(self.comboBox_taglio_p.setEditText(self.DATA_LIST[self.rec_num].taglio_p))
            str(self.comboBox_posa_opera_p.setEditText(self.DATA_LIST[self.rec_num].posa_opera_p))
            str(self.comboBox_inerti_usm.setEditText(self.DATA_LIST[self.rec_num].inerti_usm))
            str(self.comboBox_tipo_legante_usm.setEditText(self.DATA_LIST[self.rec_num].tipo_legante_usm))
            str(self.comboBox_rifinitura_usm.setEditText(self.DATA_LIST[self.rec_num].rifinitura_usm))
            str(self.comboBox_materiale_p.setEditText(self.DATA_LIST[self.rec_num].materiale_p))
            str(self.comboBox_consistenza_p.setEditText(self.DATA_LIST[self.rec_num].consistenza_p))
            self.tableInsertData("self.tableWidget_rapporti2", self.DATA_LIST[self.rec_num].rapporti2)
            str(self.mQgsFileWidget.setText(self.DATA_LIST[self.rec_num].doc_usv)) # 18 - rapporti
            # gestione tool
            if self.toolButtonPreview.isChecked():
                self.loadMapPreview()
            if self.toolButtonPreviewMedia.isChecked():
                self.loadMediaPreview()
        except:
            pass
    def set_rec_counter(self, t, c):
        self.rec_tot = t
        self.rec_corr = c
        self.label_rec_tot.setText(str(self.rec_tot))
        self.label_rec_corrente.setText(str(self.rec_corr))
    def set_LIST_REC_TEMP(self):

        ##Rapporti
        rapporti = self.table2dict("self.tableWidget_rapporti")
        rapporti2 = self.table2dict("self.tableWidget_rapporti2")
        ##Inclusi
        inclusi = self.table2dict("self.tableWidget_inclusi")
        ##Campioni
        campioni = self.table2dict("self.tableWidget_campioni")
        ##Organici
        organici = self.table2dict("self.tableWidget_organici")
        ##Inorganici
        inorganici = self.table2dict("self.tableWidget_inorganici")
        ##Documentazione
        documentazione = self.table2dict("self.tableWidget_documentazione")
        ##Inclusi materiali aggregati
        inclusi_mat_usm = self.table2dict("self.tableWidget_inclusi_materiali_usm")
        ##Inclusi leganti usm
        inclusi_leganti_usm = self.table2dict("self.tableWidget_inclusi_leganti_usm")
        colore_legante_usm = self.table2dict("self.tableWidget_colore_legante_usm")
        con_text_materiale_usm = self.table2dict("self.tableWidget_consistenza_texture_mat_usm")
        col_materiale_usm = self.table2dict("self.tableWidget_colore_materiale_usm")
        #list_foto = self.table2dict("self.tableWidget_foto")
        if self.lineEditOrderLayer.text() == "":
            order_layer = 0
        else:
            order_layer = self.lineEditOrderLayer.text()
        if self.lineEdit_qmin_usm.text() == "":
            qmin_usm = None
        else:
            qmin_usm = self.lineEdit_qmin_usm.text()
        if self.lineEdit_qmax_usm.text() == "":
            qmax_usm = None
        else:
            qmax_usm = self.lineEdit_qmax_usm.text()
        ##quota relativa
        if self.lineEdit_quota_relativa.text() == "":
            quota_relativa = None
        else:
            quota_relativa = self.lineEdit_quota_relativa.text()
        ##quota abs
        if self.lineEdit_quota_abs.text() == "":
            quota_abs = None
        else:
            quota_abs = self.lineEdit_quota_abs.text().replace(",", ".")
        ##lunghezza max
        if self.lineEdit_lunghezza_max.text() == "":
            lunghezza_max = None
        else:
            lunghezza_max = self.lineEdit_lunghezza_max.text()
        ##altezza max
        if self.lineEdit_altezza_max.text() == "":
            altezza_max = None
        else:
            altezza_max = self.lineEdit_altezza_max.text()
        ##altezza min
        if self.lineEdit_altezza_min.text() == "":
            altezza_min = None
        else:
            altezza_min = self.lineEdit_altezza_min.text()
        ##profondita max
        if self.lineEdit_profondita_max.text() == "":
            profondita_max = None
        else:
            profondita_max = self.lineEdit_profondita_max.text()
        ##profondita min
        if self.lineEdit_profondita_min.text() == "":
            profondita_min = None
        else:
            profondita_min = self.lineEdit_profondita_min.text()
        ##larghezza media
        if self.lineEdit_larghezza_media.text() == "":
            larghezza_media = None
        else:
            larghezza_media = self.lineEdit_larghezza_media.text()
        ##quota max abs
        if self.lineEdit_quota_max_abs.text() == "":
            quota_max_abs = None
        else:
            quota_max_abs = self.lineEdit_quota_max_abs.text()
        ##quota max relativa
        if self.lineEdit_quota_max_rel.text() == "":
            quota_max_rel = None
        else:
            quota_max_rel = self.lineEdit_quota_max_rel.text()
        ##quota min abs
        if self.lineEdit_quota_min_abs.text() == "":
            quota_min_abs = None
        else:
            quota_min_abs = self.lineEdit_quota_min_abs.text()
        ##quota min relativa
        if self.lineEdit_quota_min_rel.text() == "":
            quota_min_rel = None
        else:
            quota_min_rel = self.lineEdit_quota_min_rel.text()
        ##lunghezza usm
        if self.lineEdit_lunghezza_usm.text() == "":
            lunghezza_usm = None
        else:
            lunghezza_usm = self.lineEdit_lunghezza_usm.text()
        ##altezza usm
        if self.lineEdit_altezza_usm.text() == "":
            altezza_usm = None
        else:
            altezza_usm = self.lineEdit_altezza_usm.text()
        ##spessore usm
        if self.lineEdit_spessore_usm.text() == "":
            spessore_usm = None
        else:
            spessore_usm = self.lineEdit_spessore_usm.text()
            # data
        self.DATA_LIST_REC_TEMP = [
            str(self.comboBox_sito.currentText()),  # 1 - Sito
            str(self.comboBox_area.currentText()),  # 2 - Area
            str(self.lineEdit_us.text()),  # 3 - US
            str(self.comboBox_def_strat.currentText()),  # 4 - Definizione stratigrafica
            str(self.comboBox_def_intepret.currentText()),  # 5 - Definizione intepretata
            str(self.textEdit_descrizione.toPlainText()),  # 6 - descrizione
            str(self.textEdit_interpretazione.toPlainText()),  # 7 - interpretazione
            str(self.comboBox_per_iniz.currentText()),  # 8 - periodo iniziale
            str(self.comboBox_fas_iniz.currentText()),  # 9 - fase iniziale
            str(self.comboBox_per_fin.currentText()),  # 10 - periodo finale iniziale
            str(self.comboBox_fas_fin.currentText()),  # 11 - fase finale
            str(self.comboBox_scavato.currentText()),  # 12 - scavato
            str(self.lineEdit_attivita.text()),  # 13 - attivita
            str(self.lineEdit_anno.text()),  # 14 - anno scavo
            str(self.comboBox_metodo.currentText()),  # 15 - metodo
            str(inclusi),  # 16 - inclusi
            str(campioni),  # 17 - campioni
            str(rapporti),  # 18 - rapporti
            #str(organici),
            #str(inorganici),
            str(self.lineEdit_data_schedatura.text()),  # 19 - data schedatura
            str(self.comboBox_schedatore.currentText()),  # 20 - schedatore
            str(self.comboBox_formazione.currentText()),  # 21 - formazione
            str(self.comboBox_conservazione.currentText()),  # 22 - conservazione
            str(self.comboBox_colore.currentText()),  # 23 - colore
            str(self.comboBox_consistenza.currentText()),  # 24 - consistenza
            str(self.comboBox_struttura.currentText()),  # 25 - struttura
            str(self.lineEdit_codice_periodo.text()),  # 26 - codice periodo
            str(order_layer),  # 27 - order layer era str(order_layer)
            str(documentazione),
            str(self.comboBox_unita_tipo.currentText()),  # 29 us_tipo            NUOVI CAMPI NUOVI CAMPI
            str(self.comboBox_settore.currentText()),  # 30 settore
            str(self.lineEdit_quadrato.text()),  # 31 quadrato
            str(self.lineEdit_ambiente.text()),  # 32 ambiente
            str(self.lineEdit_saggio.text()),  # 33 saggio
            str(self.textEdit_elementi_datanti.toPlainText()),  # 34 elementi datanti
            str(self.comboBox_funz_statica_usm.currentText()),  # 35 funzione statica
            str(self.comboBox_lavorazione_usm.currentText()),  # 36 lavorazione usm
            str(self.lineEdit_spessore_giunti_usm.text()),  # 37 spessore giunti
            str(self.lineEdit_letti_di_posa_giunti_usm.text()),  # 38 letti posa giunti usm
            str(self.lineEdit_h_modulo_c_corsi_usm.text()),  # 39 altezza modulo corsi usm
            str(self.comboBox_unita_edilizia_riassuntiva_usm.currentText()),  # 40 unita edilizia riassuntiva
            str(self.comboBox_reimpiego_usm.currentText()),  # 41 unita edilizia riassuntiva
            str(self.comboBox_posa_in_opera_usm.currentText()),  # 42 posa in opera
            str(qmin_usm),  # 43 quota minima
            str(qmax_usm),  # 44 quota massima
            str(self.comboBox_consistenza_legante_usm.currentText()),  # 45 consitenza legante usm
            str(colore_legante_usm),  # 46 colore legante usm
            str(inclusi_leganti_usm),  # 47 aggregati leganti usm
            str(con_text_materiale_usm),  # 48 consistenza text mat
            str(col_materiale_usm),  # 49 colore materiale usm
            str(inclusi_mat_usm), # 50 inclusi_mat_usm
            str(self.lineEdit_n_catalogo_generale.text()), # 51 nr catalogo generale campi aggiunti per archeo 3.0 e allineamento ICCD
            str(self.lineEdit_n_catalogo_interno.text()), # 52 nr catalogo interno
            str(self.lineEdit_n_catalogo_internazionale.text()), # 53 nr catalogo internazionale
            str(self.comboBox_soprintendenza.currentText()), # 54 nr soprintendenza
            str(quota_relativa),  # 55 quota relativa
            str(quota_abs),  # 56 quota abs
            str(self.lineEdit_ref_tm.text()),  # 57 ref tm
            str(self.comboBox_ref_ra.currentText()),  # 58 ref ra
            str(self.lineEdit_ref_n.text()),  # 59 ref n
            str(self.comboBox_posizione.currentText()),  # 60 posizione
            str(self.lineEdit_criteri_distinzione.text()), # 61 criteri distinzione
            str(self.comboBox_modo_formazione.currentText()), # 62 modo formazione
            str(organici), # 63 componenti organici
            str(inorganici), # 64 componenti inorganici
            str(lunghezza_max),  # 65
            str(altezza_max),  # 66
            str(altezza_min),  # 67
            str(profondita_max),  # 68
            str(profondita_min),  # 69
            str(larghezza_media),  # 70
            str(quota_max_abs),  # 71
            str(quota_max_rel),  # 72
            str(quota_min_abs),  # 73
            str(quota_min_rel),  # 74
            str(self.textEdit_osservazioni.toPlainText()),  # 75 osservazioni
            str(self.lineEdit_datazione.text()),  # 76 datazione
            str(self.comboBox_flottazione.currentText()),  # 77 flottazione
            str(self.comboBox_setacciatura.currentText()),  # 78 setacciatura
            str(self.comboBox_affidabilita.currentText()),  # 79 affidabilita
            str(self.comboBox_direttore_us.currentText()),  # 80 direttore us
            str(self.comboBox_responsabile_us.currentText()), # 81 responsabile us
            str(self.lineEdit_cod_ente_schedatore.text()), # 82 cod ente schedatore
            str(self.lineEdit_data_rilevazione.text()),  # 83 data rilevazione
            str(self.lineEdit_data_rielaborazione.text()), # 84 data rielaborazione
            str(lunghezza_usm),  # 85
            str(altezza_usm),  # 86
            str(spessore_usm),  # 87
            str(self.comboBox_tecnica_muraria_usm.currentText()), # 88 tecnica muraria usm
            str(self.comboBox_modulo_usm.currentText()),  # 89 modulo usm
            str(self.lineEdit_campioni_malta_usm.text()), # 90 campioni malta usm
            str(self.lineEdit_campioni_mattone_usm.text()), # 91 campioni mattone usm
            str(self.lineEdit_campioni_pietra_usm.text()), # 92 campioni pietra usm
            str(self.lineEdit_provenienza_materiali_usm.text()), # 93 provenienza_materiali_usm
            str(self.lineEdit_criteri_distinzione_usm.text()), # 94 criteri distinzione usm
            str(self.comboBox_uso_primario_usm.currentText()),  # 95 uso primario usm
            str(self.comboBox_tipologia_opera.currentText()),  # 95 uso primario usm
            str(self.comboBox_sezione_muraria.currentText()),  # 95 uso primario usm
            str(self.comboBox_superficie_analizzata.currentText()),  # 95 uso primario usm
            str(self.comboBox_orientamento.currentText()),  # 95 uso primario usm
            str(self.comboBox_materiali_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_lavorazione_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_consistenza_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_forma_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_colore_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_impasto_lat.currentText()),  # 95 uso primario usm
            str(self.comboBox_forma_p.currentText()),  # 95 uso primario usm
            str(self.comboBox_colore_p.currentText()),  # 95 uso primario usm
            str(self.comboBox_taglio_p.currentText()),  # 95 uso primario usm
            str(self.comboBox_posa_opera_p.currentText()),  # 95 uso primario usm
            #str(list_foto)
            str(self.comboBox_inerti_usm.currentText()),  # 95 uso primario usm
            str(self.comboBox_tipo_legante_usm.currentText()),  # 95 uso primario usm
            str(self.comboBox_rifinitura_usm.currentText()),  # 95 uso primario usm
            str(self.comboBox_materiale_p.currentText()),  # 95 uso primario usm
            str(self.comboBox_consistenza_p.currentText()),  # 95 uso primario usm
            str(rapporti2),
            str(self.mQgsFileWidget.text()),# 18 - rapporti
        ]

    def set_LIST_REC_CORR(self):
        print(f"self.REC_CORR: {self.REC_CORR}, len(self.DATA_LIST): {len(self.DATA_LIST)}")

        if self.REC_CORR < 0 or self.REC_CORR >= len(self.DATA_LIST):
            raise IndexError("self.REC_CORR is out of range")

        self.DATA_LIST_REC_CORR = []
        for i in self.TABLE_FIELDS:
            try:
                self.DATA_LIST_REC_CORR.append(eval("unicode(self.DATA_LIST[self.REC_CORR]." + i + ")"))
            except IndexError as e:
                print(f"IndexError: {e} - self.REC_CORR: {self.REC_CORR}, len(self.DATA_LIST): {len(self.DATA_LIST)}")
                raise
            except Exception as e:
                print(f"Unexpected error: {e}")
                raise

    def records_equal_check(self):
        try:
            #self.set_sito()
            self.set_LIST_REC_TEMP()
            self.set_LIST_REC_CORR()

            if self.DATA_LIST_REC_CORR == self.DATA_LIST_REC_TEMP:
                return 0
            else:
                return 1
        except IndexError as e:
            print(f"IndexError: {e}")
            return 0
        except Exception as e:
            print(f"Unexpected error: {e}")
        return 0
    def setComboBoxEditable(self, f, n):
        field_names = f
        value = n
        for fn in field_names:
            cmd = '{}{}{}{}'.format(fn, '.setEditable(', n, ')')
            eval(cmd)
    def setComboBoxEnable(self, f, v):
        field_names = f
        value = v
        for fn in field_names:
            cmd = '{}{}{}{}'.format(fn, '.setEnabled(', v, ')')
            eval(cmd)
    def setTableEnable(self, t, v):
        tab_names = t
        value = v
        for tn in tab_names:
            cmd = '{}{}{}{}'.format(tn, '.setEnabled(', v, ')')
            eval(cmd)
    def testing(self, name_file, message):
        f = open(str(name_file), 'w')
        f.write(str(message))
        f.close()
    def on_pushButton_open_dir_pressed(self):
        HOME = os.environ['PYARCHINIT_HOME']
        path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_PDF_folder")
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    def on_pushButton_open_dir_matrix_pressed(self):
        HOME = os.environ['PYARCHINIT_HOME']
        path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_Matrix_folder")
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    def on_pushButton_open_dir_tavole_pressed(self):
        HOME = os.environ['PYARCHINIT_HOME']
        path = '{}{}{}'.format(HOME, os.sep, "pyarchinit_MAPS_folder")
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])


    def check_db(self):
        conn = Connection()
        conn_str = conn.conn_str()
        test_conn = conn_str.find('sqlite')

        if test_conn == 0:
            self.pushButton_import_ed2pyarchinit.setHidden(False)
        else:
            self.pushButton_import_ed2pyarchinit.setHidden(True)

    def cast_tipo_dati(self,valore):
        try:
            return int(valore)
        except ValueError:
            try:
                return float(valore)
            except ValueError:
                if valore == '':
                    return None
                else:
                    return str(valore)

    def on_pushButton_import_ed2pyarchinit_pressed(self):
        '''funzione valida solo per sqlite'''

        s = QgsSettings()
        dbpath = QFileDialog.getOpenFileName(
            self,
            "Set file name",
            '',
            "CSV (*.csv)"
        )[0]
        filename = dbpath  # .split("/")[-1]
        try:
            conn = Connection()
            conn.conn_str()
            conn_sqlite = conn.databasename()

            if not conn_sqlite["db_name"].endswith(".sqlite"):
                QMessageBox.warning(self, "Errore", "L'importazione dei dati è supportata solo per SQLite.")
                return

            sqlite_DB_path = '{}{}{}'.format(self.HOME, os.sep, "pyarchinit_DB_folder")
            con = sq.connect(sqlite_DB_path + os.sep + conn_sqlite["db_name"])
            cur = con.cursor()

            def converti_int(valore):
                return '0' if valore == '' else valore

            def converti_float(valore):
                return None if valore == '' else valore

            def converti_list(valore):
                return '[]' if valore == '' else valore

            with open(filename, 'r') as fin:
                first_line = fin.readline()
                delimiter = ',' if ',' in first_line else ';'
                fin.seek(0)
                dr = csv.DictReader(fin, delimiter=delimiter)

                to_db =[(i['sito'],i['area'],i['us'],i['d_stratigrafica'],i['d_interpretativa'],
                         i['descrizione'],i['interpretazione'],i['periodo_iniziale'],
                         i['fase_iniziale'],i['periodo_finale'],i['fase_finale'],
                         i['scavato'],i['attivita'],i['anno_scavo'],i['metodo_di_scavo'],
                         converti_list(i['inclusi']),converti_list(i['campioni']),
                         converti_list(i['rapporti']),i['data_schedatura'],i['schedatore'],
                         i['formazione'],i['stato_di_conservazione'],i['colore'],i['consistenza'],
                         i['struttura'],i['cont_per'],converti_int(i['order_layer']),
                         converti_list(i['documentazione']),i['unita_tipo'],
                         i['settore'],i['quad_par'],i['ambient'],
                         i['saggio'],i['elem_datanti'],i['funz_statica'],i['lavorazione'],
                         i['spess_giunti'],i['letti_posa'],i['alt_mod'],i['un_ed_riass'],
                         i['reimp'],i['posa_opera'],converti_float(i['quota_min_usm']),converti_float(i['quota_max_usm']),
                         i['cons_legante'],converti_list(i['col_legante']),converti_list(i['aggreg_legante']),
                         converti_list(i['con_text_mat']),converti_list(i['col_materiale']),converti_list(i['inclusi_materiali_usm']),
                         i['n_catalogo_generale'],i['n_catalogo_interno'],
                         i['n_catalogo_internazionale'],i['soprintendenza'],
                         converti_float(i['quota_relativa']),converti_float(i['quota_abs']),i['ref_tm'],
                         i['ref_ra'],i['ref_n'],i['posizione'],i['criteri_distinzione'],
                         i['modo_formazione'],converti_list(i['componenti_organici']),converti_list(i['componenti_inorganici']),
                         converti_float(i['lunghezza_max']),converti_float(i['altezza_max']),converti_float(i['altezza_min']),
                         converti_float(i['profondita_max']),converti_float(i['profondita_min']),converti_float(i['larghezza_media']),
                         converti_float(i['quota_max_abs']),converti_float(i['quota_max_rel']),converti_float(i['quota_min_abs']),
                         converti_float(i['quota_min_rel']),i['osservazioni'],i['datazione'],
                         i['flottazione'],i['setacciatura'],i['affidabilita'],
                         i['direttore_us'],i['responsabile_us'],i['cod_ente_schedatore'],
                         i['data_rilevazione'],i['data_rielaborazione'],converti_float(i['lunghezza_usm']),
                         converti_float(i['altezza_usm']),converti_float(i['spessore_usm']),i['tecnica_muraria_usm'],
                         i['modulo_usm'],i['campioni_malta_usm'],i['campioni_mattone_usm'],
                         i['campioni_pietra_usm'],i['provenienza_materiali_usm'],
                         i['criteri_distinzione_usm'],i['uso_primario_usm'],
                         i['tipologia_opera'],i['sezione_muraria'],i['superficie_analizzata'],
                         i['orientamento'],i['materiali_lat'],i['lavorazione_lat'],
                         i['consistenza_lat'],i['forma_lat'],i['colore_lat'],i['impasto_lat'],
                         i['forma_p'],i['colore_p'],i['taglio_p'],i['posa_opera_p'],
                         i['inerti_usm'],i['tipo_legante_usm'],i['rifinitura_usm'],
                         i['materiale_p'],i['consistenza_p'],converti_list(i['rapporti2']),i['doc_usv'])
                        for i in dr]

            try:
                cur.executemany(
                """INSERT INTO us_table (
                sito,
                area,
                us,
                d_stratigrafica,
                d_interpretativa,
                descrizione,
                interpretazione,
                periodo_iniziale,
                fase_iniziale,
                periodo_finale,
                fase_finale,
                scavato,
                attivita,
                anno_scavo,
                metodo_di_scavo,
                inclusi,
                campioni,
                rapporti,                
                data_schedatura,
                schedatore,
                formazione,
                stato_di_conservazione,
                colore,
                consistenza,
                struttura,
                cont_per,
                order_layer,
                documentazione,
                unita_tipo,
                settore,
                quad_par,
                ambient,
                saggio,
                elem_datanti,
                funz_statica,
                lavorazione,
                spess_giunti,
                letti_posa,
                alt_mod,
                un_ed_riass,
                reimp,
                posa_opera,
                quota_min_usm,
                quota_max_usm,
                cons_legante,
                col_legante,
                aggreg_legante,
                con_text_mat,
                col_materiale,
                inclusi_materiali_usm, 
                n_catalogo_generale,
                n_catalogo_interno,
                n_catalogo_internazionale,
                soprintendenza,
                quota_relativa,
                quota_abs,
                ref_tm,
                ref_ra,
                ref_n,
                posizione,
                criteri_distinzione,
                modo_formazione,
                componenti_organici,
                componenti_inorganici,
                lunghezza_max,
                altezza_max,
                altezza_min,
                profondita_max,
                profondita_min,
                larghezza_media,
                quota_max_abs,
                quota_max_rel,
                quota_min_abs,
                quota_min_rel,
                osservazioni,
                datazione,
                flottazione,
                setacciatura,
                affidabilita,
                direttore_us,
                responsabile_us,
                cod_ente_schedatore,
                data_rilevazione,
                data_rielaborazione,
                lunghezza_usm,
                altezza_usm,
                spessore_usm,
                tecnica_muraria_usm,
                modulo_usm,
                campioni_malta_usm,
                campioni_mattone_usm,
                campioni_pietra_usm,
                provenienza_materiali_usm,
                criteri_distinzione_usm,
                uso_primario_usm,
                tipologia_opera,
                sezione_muraria,
                superficie_analizzata,
                orientamento,
                materiali_lat,
                lavorazione_lat,
                consistenza_lat,
                forma_lat,
                colore_lat,
                impasto_lat,
                forma_p,
                colore_p,
                taglio_p,
                posa_opera_p,
                inerti_usm,
                tipo_legante_usm,
                rifinitura_usm,
                materiale_p,
                consistenza_p,
                rapporti2,
                doc_usv) VALUES (?,?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?, ?, ?,?,?,?,?);""", to_db)
                con.commit()
            except sq.IntegrityError as e:
                QMessageBox.critical(self, "Errore", f"Errore di integrità: {str(e)}")
            finally:
                con.close()

        except AssertionError as e:
            QMessageBox.warning(self, 'error', str(e), QMessageBox.Ok)

        else:
            QMessageBox.information(self, 'OK', 'Imported complited', QMessageBox.Ok)
        self.view_all()

    # This method is part of your main application window class
    def on_pushButton_filter_us_pressed(self):
        self.empty_fields()
        # Create and show the dialog
        filter_dialog = USFilterDialog(self.DB_MANAGER, self)
        result = filter_dialog.exec_()  # Show the dialog and wait for it to close

        if result:
            # Get the selected US IDs from the dialog
            selected_us_ids = filter_dialog.get_selected_us()

            # Sort DATA_LIST based on the selected US IDs
            sorted_data_list = sorted(
                self.DATA_LIST,
                key=lambda record: selected_us_ids.index(record.us) if record.us in selected_us_ids else -1
            )

            # Filter out any records that are not in selected_us_ids
            filtered_data_list = [record for record in sorted_data_list if record.us in selected_us_ids]

            # Update the UI with the filtered and sorted data
            if filtered_data_list:
                self.DATA_LIST = filtered_data_list  # Update the main data list with the filtered results
                self.REC_TOT, self.REC_CORR = len(self.DATA_LIST), 0
                self.set_rec_counter(len(self.DATA_LIST), self.REC_CORR + 1)
                self.DATA_LIST_REC_TEMP = self.DATA_LIST_REC_CORR = self.DATA_LIST[0]
                self.fill_fields()  # Assuming fill_fields takes a record as a parameter
                self.BROWSE_STATUS = "b"
                self.label_status.setText(self.STATUS_ITEMS[self.BROWSE_STATUS])
                if self.toolButtonGis.isChecked():
                    self.pyQGIS.charge_vector_layers(self.DATA_LIST)
                if self.toolButton_usm.isChecked():
                    self.pyQGIS.charge_usm_layers(self.DATA_LIST)
            else:
                QMessageBox.information(self, 'No Results', "No records match the selected filters.", QMessageBox.Ok)

    # In your main window or wherever the button is located
    def text2sql(self):
        dialog = SQLPromptDialog(iface=self.iface)
        dialog.exec_()


class SQLPromptDialog(QDialog):
    def __init__(self, iface, parent=None):
        super().__init__(parent)
        self.generated_sql = None
        self.data = None
        self.iface = iface

        self.setWindowTitle("SQL Query Prompt")
        # Create UI elements
        self.select_prompt_button = QPushButton("Select Prompt", self)
        self.prompt_combobox = QComboBox(self)
        BIN = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "bin")

        # Verifica se il file gpt_api_key.txt esiste
        self.path_prompt = os.path.join(BIN, 'last_five_prompts.txt')
        self.last_five_prompts = self.load_prompts_from_file()
        self.prompt_input = QTextEdit(self)

        # Mode selection group
        self.mode_group = QGroupBox("Modalità Text2SQL", self)
        self.mode_layout = QVBoxLayout()

        # Radio buttons for API/local choice
        self.api_radio = QRadioButton("Usa API (richiede chiave)", self)
        self.local_radio = QRadioButton("Usa modello locale (senza costi)", self)
        self.api_radio.setChecked(True)

        self.mode_layout.addWidget(self.api_radio)
        self.mode_layout.addWidget(self.local_radio)

        # Download model button
        model_path = os.path.join(BIN, "phi3_text2sql.gguf")
        self.download_button = QPushButton("Scarica Modello Locale", self)
        if os.path.exists(model_path):
            self.download_button.setText("Modello già scaricato")
            self.download_button.setEnabled(False)
        self.mode_layout.addWidget(self.download_button)

        self.mode_group.setLayout(self.mode_layout)

        # Buttons
        self.start_button = QPushButton("Create Query", self)
        self.start_sql_button = QPushButton("Execute Query", self)
        self.explain_button = QPushButton("Explain Query", self)
        self.explain_button.setEnabled(False)
        self.clear_button = QPushButton("Clear", self)
        self.sql_output = QTextEdit(self)
        self.prompt_input.textChanged.connect(self.handle_text_changed)

        # Add in __init__ after creating other buttons:
        self.add_to_canvas_button = QPushButton("Add to Canvas", self)
        self.add_to_canvas_button.setEnabled(False)

        self.results_output = QTextEdit(self)
        self.export_to_excel_button = QPushButton("Export to Excel", self)
        self.export_to_excel_button.setEnabled(False)
        self.create_graph_button = QPushButton("Create Graph", self)
        self.create_graph_button.setEnabled(False)

        # Layout
        layout = QVBoxLayout(self)
        layout.addWidget(self.select_prompt_button)
        layout.addWidget(self.prompt_combobox)
        layout.addWidget(self.prompt_input)
        layout.addWidget(self.mode_group)  # Add the mode selection group
        layout.addWidget(self.start_button)
        layout.addWidget(self.start_sql_button)
        layout.addWidget(self.explain_button)
        layout.addWidget(self.clear_button)
        layout.addWidget(self.sql_output)
        layout.addWidget(self.results_output)
        layout.addWidget(self.export_to_excel_button)
        layout.addWidget(self.create_graph_button)
        layout.addWidget(self.add_to_canvas_button)

        # Connect button click to method
        self.select_prompt_button.clicked.connect(self.on_select_prompt_clicked)
        self.download_button.clicked.connect(self.on_download_model_clicked)
        self.start_button.clicked.connect(self.on_start_button_clicked)
        self.start_sql_button.clicked.connect(self.on_start_sql_query_clicked)
        self.explain_button.clicked.connect(self.on_explainsql_button_clicked)
        self.clear_button.clicked.connect(self.clear_fields)
        self.export_to_excel_button.clicked.connect(self.on_export_to_excel_button_clicked)
        self.create_graph_button.clicked.connect(self.on_create_graph_button_clicked)
        self.add_to_canvas_button.clicked.connect(self.add_spatial_layer_to_canvas)

        # Connect signals to slots
        self.prompt_combobox.currentIndexChanged.connect(self.on_prompt_selected)

    def clear_fields(self):
        """Clear all text fields and disable buttons"""
        self.prompt_input.clear()
        self.sql_output.clear()
        self.results_output.clear()
        self.clear_results_table()
        self.explain_button.setEnabled(False)
        self.export_to_excel_button.setEnabled(False)
        self.create_graph_button.setEnabled(False)
        self.add_to_canvas_button.setEnabled(False)

    def clear_results_table(self):
        if hasattr(self, 'results_table'):
            self.results_table.clear()
            self.results_table.setRowCount(0)
            self.results_table.setColumnCount(0)
            self.results_table.hide()

    def on_prompt_selected(self, index):
        # Get the selected prompt text and set it in the prompt_input
        selected_prompt = self.prompt_combobox.itemText(index)
        self.prompt_input.setPlainText(selected_prompt)

    def update_prompt_ui(self):
        # Update the UI elements related to prompts
        # For example, refresh the QComboBox with the latest prompts
        self.prompt_combobox.clear()
        self.prompt_combobox.addItems(self.last_five_prompts)

    def on_select_prompt_clicked(self):
        # Populate the dropdown with the last five prompts
        self.prompt_combobox.clear()
        self.prompt_combobox.addItems(self.last_five_prompts)

    def record_prompt(self, prompt):
        self.last_five_prompts.insert(0, prompt)
        self.last_five_prompts = self.last_five_prompts[:5]
        self.save_prompts_to_file()

    def load_prompts_from_file(self):
        if os.path.exists(self.path_prompt):
            with open(self.path_prompt, 'r') as file:
                prompts = file.read().splitlines()
                return prompts[:5]  # Ensure only the last five are loaded
        else:
            # If the file doesn't exist, return an empty list and create the file
            with open(self.path_prompt, 'w') as file:
                pass  # Just create the file, no need to write anything yet
        return []

    def save_prompts_to_file(self):
        with open(self.path_prompt, 'w') as file:
            for prompt in self.last_five_prompts:
                file.write(f"{prompt}\n")

    def handle_text_changed(self):
        if self.is_sql_query(self.prompt_input.toPlainText()):
            self.explain_button.setEnabled(True)
        else:
            self.explain_button.setEnabled(False)

    @staticmethod
    def is_sql_query(query):
        keywords = ['select', 'update', 'insert', 'delete', 'create', 'drop', 'alter',
                    'truncate', 'grant', 'revoke', 'commit', 'rollback', 'savepoint', 'set', 'show']
        return any(re.search(f'^{keyword}', query, re.IGNORECASE) for keyword in keywords)

    def apikey_text2sql(self):
        BIN = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "bin")
        api_key = ""
        # Verifica se il file gpt_api_key.txt esiste
        path_key = os.path.join(BIN, 'text2sql_api_key.txt')
        if os.path.exists(path_key):
            # Leggi l'API Key dal file
            with open(path_key, 'r') as f:
                api_key = f.read().strip()
                try:
                    return api_key
                except:
                    reply = QMessageBox.question(None, 'Warning', 'Apikey non valida' + '\n'
                                                 + '\n' + 'vai al sito https://www.text2sql.ai/ per ottenere una chiave valida'
                                                 + '\n' + 'Clicca ok per inserire la chiave',
                                                 QMessageBox.Ok | QMessageBox.Cancel)
                    if reply == QMessageBox.Ok:
                        api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey valida:')
                        if ok:
                            # Salva la nuova API Key nel file
                            with open(path_key, 'w') as f:
                                f.write(api_key)
                                f.close()
                            with open(path_key, 'r') as f:
                                api_key = f.read().strip()
                    else:
                        return api_key
        else:
            # Chiedi all'utente di inserire una nuova API Key
            api_key, ok = QInputDialog.getText(None, 'Apikey text2sql',
                                               'vai al sito https://www.text2sql.ai/ per ottenere una chiave valida'
                                               + '\n' + 'Inserisci apikey:')
            if ok:
                # Salva la nuova API Key nel file
                with open(path_key, 'w') as f:
                    f.write(api_key)
                    f.close()
                with open(path_key, 'r') as f:
                    api_key = f.read().strip()

        return api_key

    def on_download_model_clicked(self):
        """Download the Phi-3 model for local use"""
        BIN = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "bin")
        model_path = os.path.join(BIN, "phi3_text2sql.gguf")

        # Check if model already exists
        if os.path.exists(model_path):
            QMessageBox.information(self, "Modello già scaricato",
                                    f"Il modello è già presente in:\n{model_path}")
            return

        # Create download dialog

        dialog = DownloadModelDialog(self)
        dialog.save_dir = BIN
        dialog.save_path = model_path
        dialog.info_label.setText(f"Il modello verrà scaricato in:\n{model_path}\n\nDimensione: circa 2 GB. "
                                  "Questo potrebbe richiedere tempo in base alla tua connessione internet.")

        # Execute dialog
        if dialog.exec_() == QDialog.Accepted:
            # Check if download was successful
            if os.path.exists(model_path):
                self.download_button.setText("Modello già scaricato")
                self.download_button.setEnabled(False)
                QMessageBox.information(self, "Download completato",
                                        f"Il modello è stato scaricato con successo in:\n{model_path}")
            else:
                QMessageBox.warning(self, "Download fallito",
                                    "Non è stato possibile scaricare il modello. Riprova più tardi.")

    def on_start_button_clicked(self):
        prompt = self.prompt_input.toPlainText()
        conn = Connection()
        con_string = conn.conn_str()
        test_conn = con_string.find('sqlite')

        if test_conn == 0:
            db_type = "sqlite"
        else:
            db_type = "postgres"

        # Determina quale modalità usare (API o locale)
        if self.api_radio.isChecked():
            # Usa l'API
            self.generated_sql = MakeSQL.make_api_request(prompt, db_type, self.apikey_text2sql())
        else:
            # Usa modello locale
            BIN = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "bin")
            model_path = os.path.join(BIN, "phi3_text2sql.gguf")

            # Controlla se il modello esiste
            if not os.path.exists(model_path):
                reply = QMessageBox.question(
                    self, 'Modello mancante',
                    f"Il modello non è stato trovato in {model_path}.\nVuoi scaricarlo ora?",
                    QMessageBox.Yes | QMessageBox.No
                )

                if reply == QMessageBox.Yes:
                    self.on_download_model_clicked()
                    # Verifica se il download è riuscito
                    if not os.path.exists(model_path):
                        return
                else:
                    return

            # Usa il modello locale

            self.generated_sql = MakeSQL.make_local_request(prompt, db_type, self)

        # Check if this is a spatial view creation
        if self.generated_sql and 'CREATE VIEW' in self.generated_sql.upper():
            spatial_keywords = ['geometry', 'the_geom', 'ST_']
            is_spatial = any(keyword.lower() in self.generated_sql.lower() for keyword in spatial_keywords)

            if is_spatial:
                # Enhance the spatial view creation SQL
                self.generated_sql = self.enhance_spatial_view_creation(self.generated_sql)

        self.sql_output.setText(self.generated_sql or "No SQL statement was generated.")

        # Record the prompt if not empty
        if prompt.strip():
            self.record_prompt(prompt)
            self.prompt_input.clear()
            self.update_prompt_ui()

    def on_explainsql_button_clicked(self):
        global tr, generated_explain
        L = QgsSettings().value("locale/userLocale")[0:2]
        prompt = self.prompt_input.toPlainText()

        # Determina quale modalità usare (API o locale)
        if self.api_radio.isChecked():
            # Usa l'API
            if L == "it":
                tr = ". Traduci in italiano"
                generated_explain = MakeSQL.explain_request(prompt + f"{tr}", self.apikey_text2sql())
            else:
                generated_explain = MakeSQL.explain_request(prompt, self.apikey_text2sql())
        else:
            # Usa modello locale
            BIN = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "bin")
            model_path = os.path.join(BIN, "phi3_text2sql.gguf")

            # Controlla se il modello esiste
            if not os.path.exists(model_path):
                QMessageBox.critical(self, "Modello mancante",
                                     f"Il modello non è stato trovato in {model_path}.\n"
                                     "Scarica il modello prima di utilizzare questa funzione.")
                return

            # Usa il modello locale

            generated_explain = MakeSQL.explain_local_request(prompt, self)

            # Traduci se necessario
            if L == "it" and generated_explain:
                # Il modello dovrebbe già gestire l'italiano, ma possiamo aggiungere
                # un prompt di traduzione se necessario
                pass

        self.sql_output.setText(generated_explain)

    # [Il resto del codice rimane invariato]

    def _find_all_geometry_columns(self, sql):
        """
        Find all geometry columns in the query
        """
        patterns = [
            # ST functions with AS alias
            r'ST_\w+\([^)]+\)\s+AS\s+(\w+)',
            # table.the_geom AS alias
            r'(?:\w+\.)?the_geom\s+AS\s+(\w+)',
            # geometry AS alias
            r'(?:\w+\.)?geometry\s+AS\s+(\w+)',
            # direct geometry columns
            r'(?:\w+\.)?(\w+_geom(?:etry)?)\b',
            # the_geom without alias
            r'(?:\w+\.)?the_geom\b',
            # geometry without alias
            r'(?:\w+\.)?geometry\b'
        ]

        found_geometries = []
        for pattern in patterns:
            matches = re.finditer(pattern, sql, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                geom_name = match.group(1) if match.group(1) else match.group(0)
                if geom_name and geom_name not in found_geometries:
                    found_geometries.append(geom_name)

        if found_geometries:
            self.results_output.setText(f"Found geometry columns: {', '.join(found_geometries)}")

            return found_geometries
        return None

    def _is_valid_spatial_query(self, sql, results):
        """
        Determine if the query results can be added to canvas
        """
        try:
            # Check if we have results
            if not results or not isinstance(results, list):
                return False, "No results to display"

            # Get column names from first result
            columns = list(results[0].keys())

            # Find any column that contains 'geom'
            geom_columns = [col for col in columns if 'geom' in col.lower()]

            if geom_columns:
                # Store the geometry columns for later use
                self.available_geom_columns = geom_columns
                return True, f"Found geometry columns: {', '.join(geom_columns)}"

            return False, "No geometry columns found in results"

        except Exception as e:
            return False, f"Error checking spatial query: {str(e)}"

    def on_start_sql_query_clicked(self):
        """
        Execute SQL query and show results
        """
        try:
            current_sql = self.sql_output.toPlainText().strip()
            if not current_sql:
                self.results_output.setText("No SQL query to execute.")
                return

            conn = Connection()
            con_string = conn.conn_str()
            is_sqlite = 'sqlite' in con_string.lower()
            self.generated_sql = current_sql

            # Controlla se è una CREATE VIEW
            is_create_view = 'CREATE' in current_sql.upper() and 'VIEW' in current_sql.upper()

            if is_create_view and not is_sqlite:
                # Estrai il nome della vista dalla query
                view_name = re.search(r'CREATE\s+(?:OR\s+REPLACE\s+)?VIEW\s+(\w+)', current_sql, re.IGNORECASE).group(1)

                # Esegui la CREATE VIEW
                ResponseSQL.execute_sql_and_display_results(con_string, current_sql)

                # Registra la geometria
                populate_sql = f"SELECT Populate_Geometry_Columns('{view_name}'::regclass);"
                ResponseSQL.execute_sql_and_display_results(con_string, populate_sql)

                # Seleziona dati dalla vista appena creata
                select_sql = f"SELECT * FROM {view_name};"
                results = ResponseSQL.execute_sql_and_display_results(con_string, select_sql)
                status_msg = f"View {view_name} created successfully."
            else:
                # Per SpatiaLite o query non-CREATE VIEW
                results = ResponseSQL.execute_sql_and_display_results(con_string, current_sql)
                status_msg = "Query executed."

            # Display results in table if we have any
            if isinstance(results, list) and results:
                if hasattr(self, 'results_table'):
                    self.results_table.clear()
                    self.results_table.setRowCount(0)
                    self.results_table.setColumnCount(0)

                column_names = list(results[0].keys())
                rows = len(results)
                cols = len(column_names)

                if not hasattr(self, 'results_table'):
                    self.results_table = QTableWidget(rows, cols)
                    self.layout().addWidget(self.results_table)
                else:
                    self.results_table.setRowCount(rows)
                    self.results_table.setColumnCount(cols)

                self.results_table.setHorizontalHeaderLabels(column_names)
                for i, row_data in enumerate(results):
                    for j, key in enumerate(column_names):
                        value = row_data.get(key, '')
                        item = QTableWidgetItem(str(value))
                        self.results_table.setItem(i, j, item)

                self.results_table.show()
                self.export_to_excel_button.setEnabled(True)

                can_add_to_canvas, spatial_msg = self._is_valid_spatial_query(current_sql, results)
                self.add_to_canvas_button.setEnabled(can_add_to_canvas)

                self.results_output.setText(
                    f"{status_msg}\n"
                    f"Query returned {rows} results.\n"
                    f"Spatial status: {spatial_msg}"
                )
            else:
                self.results_output.setText(
                    f"{status_msg}\n"
                    "No results to display."
                )
                if hasattr(self, 'results_table'):
                    self.results_table.clear()
                    self.results_table.hide()
                self.export_to_excel_button.setEnabled(False)
                self.add_to_canvas_button.setEnabled(False)

        except Exception as e:
            self.results_output.setText(f"An error occurred: {str(e)}")
            self.export_to_excel_button.setEnabled(False)
            self.add_to_canvas_button.setEnabled(False)
            if hasattr(self, 'results_table'):
                self.results_table.clear()
                self.results_table.hide()

    def _display_table_results(self, results):
        """
        Display results in table
        """
        try:
            # Clear previous results if any
            if hasattr(self, 'results_table'):
                self.results_table.clear()
                self.results_table.setRowCount(0)
                self.results_table.setColumnCount(0)

            # Get column names and data
            column_names = list(results[0].keys())
            rows = len(results)
            cols = len(column_names)

            # Create table if it doesn't exist
            if not hasattr(self, 'results_table'):
                self.results_table = QTableWidget(rows, cols)
                self.layout().addWidget(self.results_table)
            else:
                self.results_table.setRowCount(rows)
                self.results_table.setColumnCount(cols)

            # Set column headers
            self.results_table.setHorizontalHeaderLabels(column_names)

            # Populate table with data
            for i, row_data in enumerate(results):
                for j, key in enumerate(column_names):
                    value = row_data.get(key, '')
                    item = QTableWidgetItem(str(value))
                    self.results_table.setItem(i, j, item)

            # Show the table and enable export
            self.results_table.show()
            self.export_to_excel_button.setEnabled(True)

        except Exception as e:
            raise Exception(f"Error displaying table: {str(e)}")

    def _show_geometry_info(self, geometry_columns):
        """
        Show information about geometry columns and their validity
        """
        try:
            status_msg = []
            if geometry_columns:
                status_msg.append("Geometry columns found:")
                for geom in geometry_columns:
                    status_msg.append(f"- {geom}")

                # Test each geometry column
                valid_geoms = []
                for geom in geometry_columns:
                    layer = self._test_geometry_layer(geom)
                    if layer and layer.isValid():
                        valid_geoms.append(geom)

                if valid_geoms:
                    status_msg.append("\nValid geometry columns that can be added to canvas:")
                    for geom in valid_geoms:
                        status_msg.append(f"- {geom}")
                    self.add_to_canvas_button.setEnabled(True)
                else:
                    status_msg.append("\nNo valid geometry layers can be created from this query.")
                    self.add_to_canvas_button.setEnabled(False)
            else:
                status_msg.append("No geometry columns found in query.")
                self.add_to_canvas_button.setEnabled(False)

            self.results_output.setText("\n".join(status_msg))

        except Exception as e:
            raise Exception(f"Error showing geometry info: {str(e)}")

    def _test_geometry_layer(self, geom_column):
        """Test if a valid layer can be created with the given geometry column"""
        try:
            conn = Connection()
            uri = conn.conn_str()
            is_sqlite = 'sqlite' in uri.lower()

            if is_sqlite:
                if uri.startswith('sqlite:'):
                    db_path = uri.replace('sqlite:///', '')
                    uri = f"dbname='{db_path}'"
                elif not uri.startswith("dbname='"):
                    uri = f"dbname='{uri}'"

                layer_uri = f"{uri} table=\"({self.generated_sql})\" ({geom_column})"
                return QgsVectorLayer(layer_uri, "test_layer", "spatialite")

            return None
        except:
            return None

    def _prepare_spatialite_view_statements(self, sql):
        """
        Split and enhance view creation statements for SpatiaLite
        with proper geometry registration
        """
        view_match = re.search(r'CREATE VIEW\s+(?:IF NOT EXISTS\s+)?(["\w]+)', sql, re.IGNORECASE)
        if not view_match:
            return [sql]

        view_name = view_match.group(1).replace('"', '')
        statements = []

        # Check spatialite version and init
        statements.extend([
            "SELECT load_extension('mod_spatialite');",
            "SELECT InitSpatialMetadata(1);"
        ])

        # Drop existing view and create new
        statements.extend([
            f"DROP VIEW IF EXISTS {view_name};",
            sql.rstrip(';')
        ])

        # Register geometry column properly
        statements.append(
            f"SELECT RecoverGeometryColumn('{view_name}', 'the_geom', (SELECT srid FROM geometry_columns WHERE f_table_name = 'pyunitastratigrafiche'), 'MULTIPOLYGON', 'XY');"
        )

        # Update spatial metadata
        statements.append(f"""
        INSERT INTO views_geometry_columns 
        (view_name, view_geometry, view_rowid, f_table_name, f_geometry_column)
        SELECT '{view_name}', 'the_geom', 'rowid', f_table_name, f_geometry_column
        FROM geometry_columns WHERE f_table_name = 'pyunitastratigrafiche';
        """)

        return statements

    def execute_sql_statements(self, statements, con_string):
        """
        Execute SQL statements with spatialite specific handling
        """
        conn = None
        try:
            conn = sq.connect(con_string.split('=')[1])
            conn.enable_load_extension(True)
            conn.execute('SELECT load_extension("mod_spatialite")')

            results = None
            for stmt in statements:
                try:
                    cursor = conn.execute(stmt)
                    if cursor.description:
                        results = [dict(zip([col[0] for col in cursor.description], row))
                                   for row in cursor.fetchall()]
                    conn.commit()
                except sq.OperationalError as e:
                    if "already initialized" in str(e) or "geometry column exists" in str(e):
                        continue
                    raise
            return results
        finally:
            if conn:
                conn.close()

    def on_explainsql_button_clicked(self):
        global tr, generated_explain
        L=QgsSettings().value("locale/userLocale")[0:2]
        prompt = self.prompt_input.toPlainText()

        if L == "it":
            tr=". Traduci in italiano"
            generated_explain = MakeSQL.explain_request(prompt + f"{tr}", self.apikey_text2sql())
        else:
            generated_explain = MakeSQL.explain_request(prompt, self.apikey_text2sql())
        self.sql_output.setText(generated_explain)
        #QMessageBox.information(self, 'Explain', f'Generated Explain: {generated_explain}', QMessageBox.Ok)

    def populate_results_list(self, results):
        if results:
            column_names = results[0].keys() if results else []
            row, col = len(results), len(column_names)

            # Controlla se la tabella esiste già
            if not hasattr(self, 'results_table'):
                # Se non esiste, crea una nuova TableWidget
                self.results_table = QTableWidget(row, col)
                self.layout().addWidget(self.results_table)
            else:
                # Se esiste, pulisci la tabella esistente e imposta le nuove dimensioni
                self.results_table.clear()
                self.results_table.setRowCount(row)
                self.results_table.setColumnCount(col)

            # Imposta i nomi delle colonne
            self.results_table.setHorizontalHeaderLabels(column_names)

            # Popola la tabella con i risultati
            for i, row_data in enumerate(results):
                for j, value in enumerate(row_data):
                    item = QTableWidgetItem(str(value))
                    self.results_table.setItem(i, j, item)

            # Assicurati che la tabella sia visibile
            self.results_table.show()
        else:
            # Se non ci sono risultati, nascondi la tabella se esiste
            if hasattr(self, 'results_table'):
                self.results_table.hide()

    def on_export_to_excel_button_clicked(self):
        EXC = '{}{}{}'.format(pyarchinit_US.HOME, os.sep, "pyarchinit_EXCEL_folder")

        # Conta il numero di righe e colonne nella tabella
        rows = self.results_table.rowCount()
        cols = self.results_table.columnCount()

        # Raccogli i dati dalla tabella
        data = []
        for row in range(rows):
            rowdata = []
            for col in range(cols):
                item = self.results_table.item(row, col)
                if item is not None:
                    rowdata.append(item.text())
                else:
                    rowdata.append('')
            data.append(rowdata)

        # Raccogli i nomi delle colonne dalla tabella per usarli come intestazioni nel DataFrame
        column_headers = [self.results_table.horizontalHeaderItem(i).text() for i in range(cols)]

        # Crea DataFrame con i dati e le intestazioni delle colonne
        df = pd.DataFrame(data, columns=column_headers)
        exc_result = os.path.join(EXC, 'export_result_sql.xlsx')

        # Salva in Excel
        df.to_excel(exc_result, index=False)
        self.results_output.setText(f"Exported in {exc_result}")

    def on_create_graph_button_clicked(self):
        try:
            # Raccogli i dati da `results_table`
            data = []
            for i in range(self.results_table.rowCount()):
                category = self.results_table.item(i, 0).text()
                value = float(self.results_table.item(i, 1).text())
                data.append((category, value))

            # Crea la finestra di dialogo del grafico e passa i dati
            self.graph_window = GraphWindow()
            self.graph_window.plot(data)  # Assumiamo che `plot` disegni il grafico

            # Crea un nuovo QDockWidget e imposta il graph_window come suo widget
            dockWidget = QDockWidget("Graph Widget", self)
            dockWidget.setWidget(self.graph_window)  # Imposta il graph_window come widget del dock

            # Aggiungi il dock widget all'interfaccia utente di QGIS
            self.iface.addDockWidget(Qt.BottomDockWidgetArea, dockWidget)
            dockWidget.show()  # Assicurati che il dock widget sia visibile
        except Exception as e:
            self.results_output.setText(f"An error occurred: {e}")

    def on_explain_button_clicked(self):
        prompt = self.prompt_input.toPlainText()

    def _find_geometry_alias(self, sql):
        """
        Find the geometry column alias in the query
        """
        # Pattern per trovare alias di geometrie
        patterns = [
            # ST functions with AS alias
            r'ST_\w+\([^)]+\)\s+AS\s+(\w+)',
            # table.the_geom AS alias
            r'(?:\w+\.)?the_geom\s+AS\s+(\w+)',
            # geometry AS alias
            r'(?:\w+\.)?geometry\s+AS\s+(\w+)',
            # alias for subqueries containing geometry
            r'\(SELECT[^)]+the_geom[^)]+\)\s+AS\s+(\w+)',
            # direct geometry column
            r'(?:\w+\.)?(\w+_geom(?:etry)?)\b'
        ]

        for pattern in patterns:
            matches = re.finditer(pattern, sql, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                alias = match.group(1)
                if alias:
                    return alias

        # Se troviamo the_geom direttamente nella query
        if 'the_geom' in sql:
            return 'the_geom'

        return None

    # def add_spatial_layer_to_canvas(self):
    #     """
    #     Creates a temporary layer from SQL query and adds it to QGIS canvas
    #     """
    #     try:
    #         if not self.generated_sql:
    #             return
    #
    #         # Get the list of available geometry columns
    #         if hasattr(self, 'available_geom_columns') and self.available_geom_columns:
    #             geom_column, ok = QInputDialog.getItem(
    #                 self,
    #                 'Select Geometry Column',
    #                 'Choose the geometry column to use:',
    #                 self.available_geom_columns,
    #                 0,  # current index
    #                 False  # non-editable
    #             )
    #         else:
    #             # Fallback to text input if no columns detected
    #             geom_column, ok = QInputDialog.getText(
    #                 self,
    #                 'Geometry Column',
    #                 'Enter the name of the geometry column:',
    #                 QLineEdit.Normal,
    #                 'the_geom'
    #             )
    #
    #         if not ok or not geom_column:
    #             return
    #
    #         conn = Connection()
    #         uri = conn.conn_str()
    #         is_sqlite = 'sqlite' in uri.lower()
    #
    #         if not is_sqlite:
    #             # PostGIS handling
    #             if uri.startswith('PG:'):
    #                 uri = uri[3:]
    #             layer = QgsVectorLayer(
    #                 f"{uri} table=({self.generated_sql}) AS temp_view ({geom_column})",
    #                 "query_result",
    #                 "postgres"
    #             )
    #         else:
    #             # SpatiaLite handling
    #             sql_clean = self.generated_sql.strip().rstrip(';')
    #
    #             if uri.startswith('sqlite:'):
    #                 db_path = uri.replace('sqlite:///', '')
    #                 uri = f"dbname='{db_path}'"
    #             elif not uri.startswith("dbname='"):
    #                 uri = f"dbname='{uri}'"
    #
    #             # Build the URI with the selected geometry column
    #             layer_uri = f"{uri} table=\"({sql_clean})\" ({geom_column})"
    #
    #             # Debug info
    #             self.results_output.setText(
    #                 f"Creating layer with URI:\n{layer_uri}\n"
    #                 f"Using geometry column: {geom_column}"
    #             )
    #
    #             # Create layer
    #             layer = QgsVectorLayer(layer_uri, "query_result", "spatialite")
    #
    #         if layer.isValid():
    #             # Add layer to map
    #             QgsProject.instance().addMapLayer(layer)
    #             current_text = self.results_output.toPlainText()
    #             self.results_output.setText(
    #                 f"{current_text}\n\nSpatial layer added to canvas using '{geom_column}' column."
    #             )
    #
    #             # Zoom to layer
    #             self.iface.setActiveLayer(layer)
    #             self.iface.zoomToActiveLayer()
    #         else:
    #             error_message = layer.dataProvider().error().message() if layer.dataProvider() else 'No provider error available'
    #             self.results_output.setText(
    #                 f"Error: Could not create layer from query.\n"
    #                 f"Layer URI: {layer_uri if is_sqlite else layer.source()}\n"
    #                 f"Geometry column used: {geom_column}\n"
    #                 f"Is layer valid: {layer.isValid()}\n"
    #                 f"Provider error message: {error_message}"
    #             )
    #
    #     except Exception as e:
    #         self.results_output.setText(f"Error creating temporary layer: {str(e)}")

    def _connect_table_selection(self):
        """
        Connect table selection signal to selection handler
        """
        if hasattr(self, 'results_table'):
            self.results_table.itemSelectionChanged.connect(self._handle_table_selection)

    def _zoom_to_selected_if_possible(self, layer):
        """
        Try to zoom to selected features only if they have valid geometries
        """
        try:
            # Check if we have selected features
            if layer.selectedFeatureCount() == 0:
                return False

            # Get selected features
            selected_features = layer.selectedFeatures()
            has_valid_geometry = False

            for feature in selected_features:
                if feature.hasGeometry() and not feature.geometry().isEmpty():
                    has_valid_geometry = True
                    break

            if has_valid_geometry:
                self.iface.mapCanvas().zoomToSelected(layer)
                return True
            else:
                return False

        except Exception as e:
            self.results_output.setText(f"Error checking geometry: {str(e)}")
            return False

    def _handle_table_selection(self):
        """
        Handle selection change in results table
        """
        try:
            # Get selected rows
            selected_rows = self.results_table.selectedIndexes()
            if not selected_rows:
                return

            # Get active layer
            layer = self.iface.activeLayer()
            if not layer:
                self.results_output.setText("No active layer found.")
                return

            # Clear current selection
            layer.removeSelection()

            # Get selected row data
            row = selected_rows[0].row()
            row_data = {}
            for col in range(self.results_table.columnCount()):
                header = self.results_table.horizontalHeaderItem(col).text()
                item = self.results_table.item(row, col)
                if item and item.text():  # Solo valori non vuoti
                    row_data[header] = item.text()

            # Get available fields in layer
            layer_fields = [field.name() for field in layer.fields()]

            # Get matching fields in the current row
            matching_fields = []
            for field_name in row_data.keys():
                if field_name in layer_fields:
                    value = row_data[field_name]
                    try:
                        float(value)
                        matching_fields.append((field_name, value, 'numeric'))
                    except ValueError:
                        matching_fields.append((field_name, value, 'text'))

            if matching_fields:
                # Build expression using all matching fields
                expr_parts = []
                for field_name, value, field_type in matching_fields:
                    if field_type == 'numeric':
                        expr_parts.append(f"\"{field_name}\" = {value}")
                    else:
                        expr_parts.append(f"\"{field_name}\" = '{value}'")

                # Create and apply the selection expression
                expr = " AND ".join(expr_parts)
                layer.selectByExpression(expr)

                # Report selection status
                if layer.selectedFeatureCount() > 0:
                    selected_features = layer.selectedFeatureCount()

                    # Try to zoom only if we have valid geometries
                    zoomed = self._zoom_to_selected_if_possible(layer)

                    status_msg = [f"Selected {selected_features} feature(s)"]
                    status_msg.append(f"Expression used: {expr}")

                    if not zoomed:
                        status_msg.append("(No zoomable geometry available)")

                    self.results_output.setText("\n".join(status_msg))
                else:
                    # Try with individual fields
                    for field_name, value, field_type in matching_fields:
                        if field_type == 'numeric':
                            expr = f"\"{field_name}\" = {value}"
                        else:
                            expr = f"\"{field_name}\" = '{value}'"

                        layer.selectByExpression(expr)
                        if layer.selectedFeatureCount() > 0:
                            zoomed = self._zoom_to_selected_if_possible(layer)
                            status_msg = [f"Selected using field: {field_name} = {value}"]
                            if not zoomed:
                                status_msg.append("(No zoomable geometry available)")
                            self.results_output.setText("\n".join(status_msg))
                            break
                    else:
                        self.results_output.setText(
                            f"Could not find matching features.\n"
                            f"Available fields: {layer_fields}\n"
                            f"Row values: {row_data}"
                        )
            else:
                self.results_output.setText(
                    f"No matching fields between table and layer.\n"
                    f"Layer fields: {layer_fields}\n"
                    f"Table fields: {list(row_data.keys())}"
                )

        except Exception as e:
            self.results_output.setText(
                f"Error during selection: {str(e)}\n"
                f"Layer fields: {layer_fields if 'layer_fields' in locals() else 'unknown'}\n"
                f"Row data: {row_data if 'row_data' in locals() else 'unknown'}"
            )

    # def add_spatial_layer_to_canvas(self):
    #     """
    #     Creates a temporary layer from SQL query and adds it to QGIS canvas
    #     """
    #     try:
    #         if not self.generated_sql:
    #             return
    #
    #         # Get the list of available geometry columns
    #         if hasattr(self, 'available_geom_columns') and self.available_geom_columns:
    #             geom_column, ok = QInputDialog.getItem(
    #                 self,
    #                 'Select Geometry Column',
    #                 'Choose the geometry column to use:',
    #                 self.available_geom_columns,
    #                 0,  # current index
    #                 False  # non-editable
    #             )
    #         else:
    #             # Fallback to text input if no columns detected
    #             geom_column, ok = QInputDialog.getText(
    #                 self,
    #                 'Geometry Column',
    #                 'Enter the name of the geometry column:',
    #                 QLineEdit.Normal,
    #                 'the_geom'
    #             )
    #
    #         if not ok or not geom_column:
    #             return
    #
    #         conn = Connection()
    #         uri = conn.conn_str()
    #         is_sqlite = 'sqlite' in uri.lower()
    #
    #         if not is_sqlite:
    #             # PostGIS handling
    #             if uri.startswith('PG:'):
    #                 uri = uri[3:]
    #             layer = QgsVectorLayer(
    #                 f"{uri} table=({self.generated_sql}) AS temp_view ({geom_column})",
    #                 "query_result",
    #                 "postgres"
    #             )
    #         else:
    #             # SpatiaLite handling
    #             sql_clean = self.generated_sql.strip().rstrip(';')
    #
    #             if uri.startswith('sqlite:'):
    #                 db_path = uri.replace('sqlite:///', '')
    #                 uri = f"dbname='{db_path}'"
    #             elif not uri.startswith("dbname='"):
    #                 uri = f"dbname='{uri}'"
    #
    #             layer_uri = f"{uri} table=\"({sql_clean})\" ({geom_column})"
    #             layer = QgsVectorLayer(layer_uri, "query_result", "spatialite")
    #
    #         if layer.isValid():
    #             # Add layer to map
    #             QgsProject.instance().addMapLayer(layer)
    #
    #             # Make it the active layer
    #             self.iface.setActiveLayer(layer)
    #
    #             # Connect table selection handler
    #             self._connect_table_selection()
    #
    #             # Update UI
    #             current_text = self.results_output.toPlainText()
    #             self.results_output.setText(
    #                 f"{current_text}\n\nSpatial layer added to canvas using '{geom_column}' column.\n"
    #                 "Table selection is now linked to canvas."
    #             )
    #
    #             # Zoom to layer
    #             self.iface.zoomToActiveLayer()
    #         else:
    #             error_message = layer.dataProvider().error().message() if layer.dataProvider() else 'No provider error available'
    #             self.results_output.setText(
    #                 f"Error: Could not create layer from query.\n"
    #                 f"Layer URI: {layer_uri if is_sqlite else layer.source()}\n"
    #                 f"Geometry column used: {geom_column}\n"
    #                 f"Is layer valid: {layer.isValid()}\n"
    #                 f"Provider error message: {error_message}"
    #             )
    #
    #     except Exception as e:
    #         self.results_output.setText(f"Error creating temporary layer: {str(e)}")

    def _determine_geometry_type(self, sql):
        """
        Determine geometry type from SQL query
        """
        sql_upper = sql.upper()

        # Polygon functions
        if any(x in sql_upper for x in [
            'ST_BUFFER',
            'ST_ENVELOPE',
            'ST_CONVEXHULL',
            'ST_BUILDAREA',
            'ST_POLYGONIZE',
            'ST_MAKEPOLYGON'
        ]):
            return 'MULTIPOLYGON'

        # MultiPolygon functions
        if any(x in sql_upper for x in [
            'ST_UNION',
            'ST_COLLECT',
            'ST_MULTI',
            'ST_DUMP',
            'ST_POLYGONIZE'
        ]):
            return 'MULTIPOLYGON'

        # Point functions
        if any(x in sql_upper for x in [
            'ST_POINT',
            'ST_POINTONSURFACE',
            'ST_CENTROID',
            'ST_POINTN',
            'ST_STARTPOINT',
            'ST_ENDPOINT'
        ]):
            return 'POINT'

        # MultiPoint functions
        if any(x in sql_upper for x in [
            'ST_MULTIPOINT',
            'ST_GENERATEPOINTS'
        ]):
            return 'MULTIPOINT'

        # LineString functions
        if any(x in sql_upper for x in [
            'ST_LINE',
            'ST_LINESTRING',
            'ST_BOUNDARY',
            'ST_EXTERIORRING',
            'ST_INTERIORRINGN',
            'ST_OFFSETCURVE',
            'ST_MAKELINE'
        ]):
            return 'LINESTRING'

        # MultiLineString functions
        if any(x in sql_upper for x in [
            'ST_MULTILINESTRING',
            'ST_LINEMERGE'
        ]):
            return 'MULTILINESTRING'

        # Check for direct geometry column selection
        if 'SELECT' in sql_upper:
            # Try to find the source table and get its geometry type
            source_table_match = re.search(r'FROM\s+(\w+)', sql)
            if source_table_match:
                source_table = source_table_match.group(1)
                return f"(SELECT GeometryType(the_geom) FROM {source_table} LIMIT 1)"

        return 'GEOMETRY'  # default fallback

    def _create_postgis_view(self, uri, view_name, sql, geom_column):
        """
        Create a spatial view in PostGIS with proper geometry registration
        """
        try:
            # Creiamo una lista di statement SQL da eseguire in sequenza
            sql_statements = [
                f"DROP VIEW IF EXISTS {view_name};",
                f"CREATE VIEW {view_name} AS {sql};",
                f"SELECT Populate_Geometry_Columns('{view_name}'::regclass);"
            ]

            # Esegui ogni statement separatamente e raccogli i risultati
            results = []
            for stmt in sql_statements:
                result = ResponseSQL.execute_sql_and_display_results(uri, stmt)
                if isinstance(result, list):
                    results.extend(result)

            # Verifica se la vista è stata creata
            check_sql = f"""
            SELECT f_geometry_column, type, srid 
            FROM geometry_columns 
            WHERE f_table_name = '{view_name}';
            """
            geom_info = ResponseSQL.execute_sql_and_display_results(uri, check_sql)

            debug_msg = f"""
            View creation completed:
            - Statements executed: {len(sql_statements)}
            - Results: {results}
            - Geometry info: {geom_info}
            """

            return True, debug_msg

        except Exception as e:
            return False, f"Error creating view: {str(e)}"

    def _create_vector_layer_from_postgres(self, uri, view_name, geom_column):
        """
        Create vector layer from PostgreSQL view
        """
        try:
            # Rimuovi i parametri non necessari e spazi extra dal nome della vista
            view_name = view_name.strip()

            # Verifica il tipo di geometria e SRID
            check_sql = f"""
            SELECT type, srid 
            FROM geometry_columns 
            WHERE f_table_name = '{view_name}' 
            AND f_geometry_column = '{geom_column}';
            """

            conn = Connection()
            result = ResponseSQL.execute_sql_and_display_results(conn.conn_str(), check_sql)

            if result and len(result) > 0:
                geom_type = result[0]['type']
                srid = result[0]['srid']

                # Costruisci l'URI del layer
                if uri.startswith('PG:'):
                    uri = uri[3:]

                # Estrai i parametri di connessione
                conn_params = dict(param.split('=') for param in uri.split() if '=' in param)

                # Costruisci l'URI nel formato corretto per PostGIS
                layer_uri = (
                    f"dbname='{conn_params.get('dbname', '')}' "
                    f"host='{conn_params.get('host', 'localhost')}' "
                    f"port='{conn_params.get('port', '5432')}' "
                    f"user='{conn_params.get('user', '')}' "
                    f"password='{conn_params.get('password', '')}' "
                    f"sslmode='{conn_params.get('sslmode', 'prefer')}' "
                    f"key='{geom_column}' "
                    f"srid='{srid}' "
                    f"type='{geom_type}' "
                    f"table=\"{view_name}\" ({geom_column})"
                )

                # Crea il layer
                layer = QgsVectorLayer(layer_uri, view_name, "postgres")

                if layer.isValid():
                    return layer, None
                else:
                    return None, f"Layer non valido.\nURI: {layer_uri}\nErrore: {layer.dataProvider().error().message()}"
            else:
                return None, f"Impossibile trovare informazioni sulla geometria per la vista {view_name}"

        except Exception as e:
            return None, f"Errore nella creazione del layer: {str(e)}"

    def add_spatial_layer_to_canvas(self):
        try:
            if not self.generated_sql:
                return

            conn = Connection()
            uri = conn.conn_str()
            is_sqlite = 'sqlite' in uri.lower()

            # Cerca la colonna geometrica nella query
            geom_match = re.search(r'ST_\w+\([^)]+\)\s+AS\s+(\w+)', self.generated_sql, re.IGNORECASE)
            default_geom = geom_match.group(1) if geom_match else 'the_geom'

            # Get geometry column
            if hasattr(self, 'available_geom_columns') and self.available_geom_columns:
                geom_column, ok = QInputDialog.getItem(
                    self,
                    'Select Geometry Column',
                    'Choose the geometry column to use:',
                    self.available_geom_columns,
                    0,
                    False
                )
            else:
                geom_column, ok = QInputDialog.getText(
                    self,
                    'Geometry Column',
                    'Enter the name of the geometry column:',
                    QLineEdit.Normal,
                    default_geom
                )

            if not ok or not geom_column:
                return

            if is_sqlite:
                # SpatiaLite handling
                sql_clean = self.generated_sql.strip().rstrip(';')

                if uri.startswith('sqlite:'):
                    db_path = uri.replace('sqlite:///', '')
                    uri = f"dbname='{db_path}'"
                elif not uri.startswith("dbname='"):
                    uri = f"dbname='{uri}'"

                layer_uri = f"{uri} table=\"({sql_clean})\" ({geom_column})"
                layer = QgsVectorLayer(layer_uri, "query_result", "spatialite")
                error_message = None
            else:
                # PostgreSQL handling
                try:
                    # Get connection parameters
                    conn_params_sql = "SELECT current_database(), inet_server_addr(), inet_server_port(), current_user;"
                    result = ResponseSQL.execute_sql_and_display_results(uri, conn_params_sql)

                    if result and len(result) > 0:
                        db_name = result[0]['current_database']
                        host = result[0]['inet_server_addr']
                        port = result[0]['inet_server_port']
                        user = result[0]['current_user']

                        # Clean SQL
                        sql_clean = self.generated_sql.strip().rstrip(';')

                        # Build URI
                        layer_uri = (
                            f"dbname='{db_name}' "
                            f"host='{host}' "
                            f"port='{port}' "
                            f"user='{user}' "
                            f"password='postgres' "
                            f"sslmode='prefer' "
                            f"key='gid' "
                            f"checkPrimaryKeyUnicity='1' "
                            f"table=\"({sql_clean})\" ({geom_column})"
                        )

                        debug_msg = (
                            f"PostgreSQL connection details:\n"
                            f"Database: {db_name}\n"
                            f"Host: {host}\n"
                            f"Port: {port}\n"
                            f"Username: {user}\n"
                            f"SQL: {sql_clean}\n"
                            f"Complete URI: {layer_uri}\n"
                            f"Geometry column: {geom_column}"
                        )
                    else:
                        self.results_output.setText("Could not get connection parameters")
                        return

                except Exception as e:
                    self.results_output.setText(f"Error getting connection parameters: {str(e)}")
                    return

                # Create layer
                layer = QgsVectorLayer(layer_uri, "query_result", "postgres")

            if layer and layer.isValid():
                # Add layer to map
                QgsProject.instance().addMapLayer(layer)
                self.iface.setActiveLayer(layer)
                self._connect_table_selection()

                self.results_output.setText(
                    f"Spatial layer added to canvas.\n"
                    f"Using geometry column: {geom_column}\n"
                    f"Layer is valid and loaded successfully.\n"
                    f"Connection details:\n{debug_msg}"
                )

                self.iface.zoomToActiveLayer()
            else:
                error_msg = layer.dataProvider().error().message() if layer and layer.dataProvider() else 'Unknown error'
                self.results_output.setText(
                    f"Error: Could not create layer.\n"
                    f"Connection details:\n{debug_msg}\n"
                    f"Is layer valid: {layer.isValid() if layer else False}\n"
                    f"Provider error message: {error_msg}"
                )

        except Exception as e:
            self.results_output.setText(f"Error creating layer: {str(e)}")

    def _has_explicit_geom_alias(self, sql):
        """
        Check if the query already has the geometry column
        """
        # Controlla se c'è table.* nella query
        if re.search(r'\w+\.\*', sql):
            return True

        # Controlla alias espliciti
        if re.search(r'(?i)(AS\s+)?the_geom\b', sql):
            return True

        return False

    def _find_geometry_column(self, sql):
        """
        Find the geometry column in the query
        """
        patterns = [
            r'(\w+\.the_geom)',  # table.the_geom
            r'(the_geom)',  # the_geom
            r'(\w+\.geometry)',  # table.geometry
            r'(geometry)',  # geometry
            r'(ST_\w+\([^)]+\))'  # ST_* functions
        ]

        for pattern in patterns:
            match = re.search(pattern, sql, re.IGNORECASE)
            if match:
                return match.group(1)

        return None

    def _get_geometry_column_from_query(self, sql):
        """
        Determines the geometry column name from the SQL query
        """
        # Look for AS clauses that might rename geometry columns
        as_matches = re.finditer(r'(\w+\.the_geom)\s+AS\s+(\w+)', sql, re.IGNORECASE)
        for match in as_matches:
            if match.group(2):  # If we found a renamed geometry column
                return match.group(2)

        # Look for direct geometry column references
        geom_matches = re.finditer(r'(\w+\.the_geom)', sql, re.IGNORECASE)
        for match in geom_matches:
            if match.group(1):  # If we found a geometry column
                return 'the_geom'

        return 'geometry'  # Default fallback

    # def on_start_sql_query_clicked(self):
    #     """
    #     Execute SQL query and show results
    #     """
    #     try:
    #         if not self.generated_sql:
    #             self.results_output.setText("No SQL query to execute.")
    #             return
    #
    #         conn = Connection()
    #         con_string = conn.conn_str()
    #
    #         # Esegui la query e mostra i risultati
    #         results = ResponseSQL.execute_sql_and_display_results(con_string, self.generated_sql)
    #
    #         # Gestisci i risultati
    #         if isinstance(results, list):
    #             self.populate_results_list(results)
    #             self.export_to_excel_button.setEnabled(True)
    #             self.create_graph_button.setEnabled(len(results[0]) == 2)
    #         else:
    #             self.results_output.setText("Query executed successfully.")
    #             self.export_to_excel_button.setEnabled(False)
    #             self.create_graph_button.setEnabled(False)
    #
    #         # Abilita il pulsante spatial se la query contiene geometria
    #         spatial_keywords = ['the_geom', 'ST_', 'geometry']
    #         self.add_to_canvas_button.setEnabled(
    #             any(keyword.lower() in self.generated_sql.lower() for keyword in spatial_keywords)
    #         )
    #
    #     except Exception as e:
    #         self.results_output.setText(f"An error occurred: {str(e)}")
    #         self.export_to_excel_button.setEnabled(False)
    #         self.create_graph_button.setEnabled(False)

    def _extract_view_name(self):
        """Extract view name from SQL query"""
        if 'CREATE VIEW' in self.generated_sql.upper():
            match = re.search(r'CREATE VIEW\s+(?:IF NOT EXISTS\s+)?(["\w]+)', self.generated_sql, re.IGNORECASE)
            if match:
                return match.group(1).replace('"', '')
        return None



    def enhance_spatial_view_creation(self, sql_query):
        """
        Enhances a spatial query by converting SELECT to VIEW when needed
        """
        conn = Connection()
        con_string = conn.conn_str()
        is_sqlite = 'sqlite' in con_string.lower()

        # Check if it's a SELECT query without CREATE VIEW
        if sql_query.upper().strip().startswith('SELECT') and 'CREATE VIEW' not in sql_query.upper():
            # Convert to CREATE VIEW
            view_name = 'spatial_selection_view'
            sql_query = f"CREATE VIEW {view_name} AS {sql_query}"

        # Extract view name from the query
        match = re.search(r'CREATE\s+VIEW\s+(?:IF NOT EXISTS\s+)?(["\w]+)', sql_query, re.IGNORECASE)
        if not match:
            return sql_query

        view_name = match.group(1).replace('"', '')

        # Create the enhanced SQL statements
        enhanced_sql = []

        if is_sqlite:
            # For SpatiaLite
            enhanced_sql.extend([
                f"DROP VIEW IF EXISTS {view_name};",
                sql_query.strip().rstrip(';') + ';',
                f"""INSERT INTO views_geometry_columns 
                (view_name, view_geometry, view_rowid, f_table_name, f_geometry_column)
                VALUES 
                ('{view_name}', 'the_geom', 'rowid', 'pyunitastratigrafiche', 'the_geom')"""
            ])
        else:
            # For PostGIS
            enhanced_sql.extend([
                f"DROP VIEW IF EXISTS {view_name};",
                sql_query,
                f"SELECT Populate_Geometry_Columns('{view_name}'::regclass);"
            ])

        return '\n'.join(enhanced_sql) if is_sqlite else '\n'.join(enhanced_sql) + ';'






class MplCanvas(FigureCanvas):
    def __init__(self, parent=None):
        fig = Figure()
        self.axes = fig.add_subplot(111)
        super().__init__(fig)


class GraphWindow(QDockWidget):
    def __init__(self):
        super(GraphWindow, self).__init__()  # Utilizza super() per inizializzare la classe base

        # Crea un'istanza di MplCanvas
        self.canvas = MplCanvas()

        # Usa setWidget per aggiungere il canvas al QDockWidget
        self.setWidget(self.canvas)

    def plot(self, data):
        # La tua logica di plotting va qui, usando `data`
        self.canvas.figure.clear()
        ax = self.canvas.figure.add_subplot(111)
        categories, values = zip(*data)  # Se i dati sono passati come tuple di (categoria, valore)
        #ax.bar(categories, values)
        # Configura ulteriormente il grafico come desiderato
        #self.canvas.draw()

        # calculate categories totals
        total_categories = Counter(category for category, value in data)

        # calculate total values
        total_values = sum(value for category, value in data)

        # create labels and calculate percentages
        labels = [category for category, value in data]
        values = [value for category, value in data]
        percentages = [value / total_values * 100 for value in values]


        bars = ax.bar(labels, percentages)

        # rotate labels if they overlap
        self.canvas.figure.autofmt_xdate()

        # add totals and percentages on top of each bar
        for i, bar in enumerate(bars):
            yval = bar.get_height()
            total = total_categories[labels[i]]
            ax.text(bar.get_x() + bar.get_width() / 2, yval + 0.01,
                    f"{yval:.1f}%", ha='center', va='bottom')

        self.canvas.figure.canvas.draw_idle()

class USFilterDialog(QDialog):
    L = QgsSettings().value("locale/userLocale")[0:2]
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.selected_us = []
        self.us_records = []  # Store all US records
        self.initUI()

    def initUI(self):

        if self.L=='it':

            self.setWindowTitle("Filtro UUSS Records")  # Set the window title
        else:
            self.setWindowTitle("Filter SU Records")  # Set the window title

        layout = QVBoxLayout(self)

        # Create search bar
        self.search_bar = QLineEdit(self)
        self.search_bar.setPlaceholderText("Search...")
        self.search_bar.textChanged.connect(self.filter_list)  # Connect textChanged signal to filter function
        layout.addWidget(self.search_bar)

        # Create list widget
        self.list_widget = QListWidget(self)
        layout.addWidget(self.list_widget)

        # Populate list widget with checkboxes
        self.populate_list_with_us()

        # Create filter button
        filter_button = QPushButton('Filter', self)
        filter_button.clicked.connect(self.apply_filter)
        layout.addWidget(filter_button)

        # Set dialog layout
        self.setLayout(layout)

    def populate_list_with_us(self):
        # Fetch US records from the database and sort them
        self.us_records = sorted(self.db_manager.query_all('us_table'), key=lambda x: x.us)
        self.update_list_widget(self.us_records)

    def update_list_widget(self, records):
        # Clear the list widget
        self.list_widget.clear()

        # Repopulate the list widget with given records
        for us_record in records:
            list_item = QListWidgetItem(self.list_widget)

            checkbox = QCheckBox(f"{us_record.unita_tipo} {us_record.us}")

            checkbox.us = us_record.us
            self.list_widget.addItem(list_item)
            self.list_widget.setItemWidget(list_item, checkbox)

    def filter_list(self, text):
        # Filter US records based on the search text
        filtered_records = [us_record for us_record in self.us_records if str(us_record.us).startswith(text)]
        self.update_list_widget(filtered_records)


    def apply_filter(self):
        # Clear the selected US list
        self.selected_us.clear()

        # Check which checkboxes are checked and add the US IDs to the selected list
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            checkbox = self.list_widget.itemWidget(item)
            if checkbox.isChecked():
                us_id = int(checkbox.text().split(' ')[1])  # Extract the US ID from the checkbox text
                self.selected_us.append(us_id)

        print(f"Selected US IDs: {self.selected_us}")  # Debug print statement
        self.accept()

    def get_selected_us(self):
        # Return the list of selected US IDs
        return self.selected_us

class IntegerDelegate(QtWidgets.QStyledItemDelegate):
    """
        The IntegerDelegate class is a subclass of QStyledItemDelegate from the PyQt5 library. It is used to create a delegate for editing integer values in a Qt model/view framework.
        Example Usage
        # Import the necessary libraries
        from PyQt5 import QtGui, QtWidgets

        # Create an instance of the IntegerDelegate class
        delegate = IntegerDelegate()

        # Set the delegate for a specific column in a QTableView
        tableView.setItemDelegateForColumn(columnIndex, delegate)
        Code Analysis
        Main functionalities
        The main functionality of the IntegerDelegate class is to provide a custom editor for integer values in a Qt model/view framework. It creates a QLineEdit widget as the editor and sets a QIntValidator to ensure that only valid integer values can be entered.

        Methods
        createEditor(parent, option, index): This method is called when a cell in the view needs to be edited. It creates and returns a QLineEdit widget as the editor for the cell. It also sets a QIntValidator to ensure that only valid integer values can be entered.

        Fields
        None

    """

    def __init__(self, parent=None):
        super(IntegerDelegate, self).__init__(parent)

    def createEditor(self, parent, option, index):
        editor = QtWidgets.QLineEdit(parent)
        validator = QtGui.QIntValidator()
        editor.setValidator(validator)
        return editor

# if __name__ == "__main__":
#     app = QApplication(sys.argv)
#     mainWin = GPTWindow()
#     mainWin.show()
#     sys.exit(app.exec_())
