# -*- coding: utf-8 -*-
"""
AI descriptive report for the palimpsestr SEF analysis.

After the R side (palimpsestr) has fitted the Stratigraphic Entanglement Field
model, ``_gather_sef_facts`` collects the *facts* — per-phase summary, per-class
composition, per-US diagnostics, the absolute-chronology (OxCal) rows, the
``gg_*`` diagnostic figures and the R interpretive narrative. This module runs a
pipeline of **specialised LLM agents** over those facts and produces one
didactic, human-readable report **in any pyArchInit UI language**:

    1. Methodologist — explains *why* the analysis was parameterised the way it
       was: the class model (multinomial vs gaussian), the number of phases K
       (and what the diagnostics say), the noise component and intrusion
       threshold, the finds source. States caveats.
    2. Analyst — interprets the results: phase composition and chronology
       (using the OxCal absolute dates when present), residuality / intrusions,
       spatial pattern.
    3. Synthesiser — weaves the two into one cohesive report with tables,
       references the figures, and closes with conclusions.

The report is rendered to **DOCX, PDF or Markdown** with a small Markdown
renderer (headings, bold/italic, bullet/numbered lists and real tables), the
``gg_*`` figures embedded, and a guaranteed data appendix (per-phase table +
the OxCal chronology table) so the underlying numbers are always visible.

@author: Enzo Cocca <enzo.ccc@gmail.com>
"""
import os
import re
import html
import inspect

from qgis.PyQt.QtCore import QThread, pyqtSignal
from qgis.PyQt.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QFormLayout, QLabel, QComboBox,
    QCheckBox, QPushButton, QPlainTextEdit, QMessageBox, QFileDialog)

from modules.utility.llm_providers import LLMProviderManager
from modules.utility.llm_selector_widget import LLMSelectorWidget


# pyArchInit UI languages (code -> name shown in the combo + given to the LLM).
AI_LANGUAGES = [
    ("it", "Italiano"), ("en", "English"), ("de", "Deutsch"),
    ("es", "Español"), ("fr", "Français"), ("pt", "Português"),
    ("ca", "Català"), ("ro", "Română"), ("el", "Ελληνικά"),
    ("ar", "العربية"),
]

# What each gg_* figure shows — used as figure captions and given to the agents.
FIGURE_CAPTIONS = {
    "phasefield": "Campo di fase / phase field (SEF)",
    "entropy": "Entropia di assegnazione per US / assignment entropy per US",
    "energy": "Energia del campo / field energy",
    "intrusions": "Intrusioni e residualità / intrusions and residuality",
    "phase_composition": "Composizione dei reperti per fase / finds composition by phase",
    "direction": "Direzione cronologica / chronological direction",
    "outliers": "Outlier posteriori / posterior outliers",
    "unit_coherence": "Coerenza delle US / unit coherence",
}


def _model_name(idx):
    return ("multinomial", "gaussian")[idx] if idx in (0, 1) else str(idx)


def _source_name(idx):
    return ("entrambi (materiali+ceramica)", "solo materiali", "solo ceramica")[idx] \
        if idx in (0, 1, 2) else str(idx)


# --------------------------------------------------------------- facts -> text ---
def _rows_md_table(rows, columns, headers=None):
    """Render a list-of-dicts as a Markdown table (only the given columns)."""
    if not rows:
        return ""
    headers = headers or columns
    out = ["| " + " | ".join(headers) + " |",
           "| " + " | ".join("---" for _ in headers) + " |"]
    for r in rows:
        out.append("| " + " | ".join(str(r.get(c, "")) for c in columns) + " |")
    return "\n".join(out)


def _facts_block(facts):
    """Compact, data-rich description of the run, fed to every agent."""
    lines = [
        "PARAMETRI DELL'ANALISI SEF (palimpsestr):",
        "- Sito: %s" % facts.get("site", "tutti"),
        "- Backend dati: %s" % facts.get("backend", "?"),
        "- Numero di fasi richieste K = %s" % facts.get("k"),
        "- Modello di classe = %s" % facts.get("class_model"),
        "- Componente di rumore/outlier = %s" % ("attiva" if facts.get("noise") else "disattiva"),
        "- Soglia intrusioni = %s" % facts.get("threshold"),
        "- Selezione reperti = %s" % facts.get("source"),
        "- Cronologia assoluta (palimpsest_chronology) = %s"
        % ("presente, usata al posto della datazione testuale"
           if facts.get("has_chronology") else "assente (datazione testuale)"),
    ]
    ps = facts.get("phase_summary") or []
    if ps:
        lines += ["", "SINTESI PER FASE (dati reali, NON inventare):",
                  _rows_md_table(ps, ["phase", "n_finds", "pct", "mean_date", "n_us"],
                                 ["Fase", "N reperti", "%", "Data media", "N US"])]
    comp = facts.get("composition") or []
    if comp:
        lines += ["", "COMPOSIZIONE PER CLASSE (fase, classe, conteggio):",
                  _rows_md_table(comp, ["phase", "class", "Freq"],
                                 ["Fase", "Classe", "N"])]
    ch = facts.get("chronology") or []
    if ch:
        lines += ["", "CRONOLOGIA ASSOLUTA OxCal (per US; anni calendariali, "
                  "negativo = a.C.):",
                  _rows_md_table(ch, ["us", "start", "end", "lab_code", "source"],
                                 ["US", "start", "end", "lab", "source"])]
    if facts.get("r_markdown"):
        lines += ["", "NARRATIVA E DIAGNOSTICHE prodotte da palimpsestr "
                  "(fonte di verità):", facts["r_markdown"]]
    if facts.get("figures"):
        lines += ["", "FIGURE DISPONIBILI (riferiscile per nome quando utile):"]
        for key in facts["figures"]:
            lines.append("- %s: %s" % (key, FIGURE_CAPTIONS.get(key, key)))
    return "\n".join(lines)


def _methodologist_prompt(facts, lang_name):
    return (
        "Sei un metodologo dell'archeologia quantitativa, esperto del modello "
        "Stratigraphic Entanglement Field (SEF) di palimpsestr.\n\n"
        + _facts_block(facts) +
        "\n\nCompito: spiega in modo CHIARO e DIDATTICO le scelte metodologiche, "
        "giustificandole con le diagnostiche fornite:\n"
        "1. Perché il modello di classe scelto (multinomiale vs gaussiano) è "
        "adatto a questi dati e cosa cambierebbe con l'altro.\n"
        "2. Il numero di fasi K: se è appropriato alla luce di entropia, "
        "coerenza delle US ed energia; se i dati suggerirebbero un K diverso.\n"
        "3. La componente di rumore e la soglia di intrusione: cosa implicano "
        "per la lettura della residualità.\n"
        "4. La selezione dei reperti e l'uso della cronologia assoluta OxCal.\n"
        "Indica limiti e cautele. NON inventare numeri: usa solo quelli forniti. "
        "Scrivi ESCLUSIVAMENTE in " + lang_name + ".")


def _analyst_prompt(facts, lang_name):
    return (
        "Sei un archeologo stratigrafo esperto di formazione dei depositi e "
        "analisi dei palinsesti.\n\n"
        + _facts_block(facts) +
        "\n\nCompito: INTERPRETA i risultati dell'analisi SEF:\n"
        "1. Composizione e significato delle fasi (cosa rappresentano in "
        "termini di reperti e cronologia).\n"
        "2. Cronologia assoluta: discuti gli intervalli OxCal per US/fase e la "
        "direzione cronologica; collega le date alle fasi.\n"
        "3. Residualità e intrusioni: cosa dicono sui processi di formazione.\n"
        "4. Eventuale pattern spaziale.\n"
        "Fonda ogni affermazione sui dati forniti. Scrivi ESCLUSIVAMENTE in "
        + lang_name + ".")


def _synthesis_prompt(facts, lang_name, methodology, analysis):
    return (
        "Sei il redattore scientifico del report finale dell'analisi SEF "
        "(palimpsestr) per una relazione archeologica.\n\n"
        "=== NOTA METODOLOGICA ===\n" + methodology +
        "\n\n=== INTERPRETAZIONE ===\n" + analysis +
        "\n\n" + _facts_block(facts) +
        "\n\nCompito: scrivi un UNICO report coeso, ben strutturato e MOLTO "
        "CHIARO, con titoli di sezione (Markdown `#`/`##`). Requisiti:\n"
        "- includi una **tabella Markdown di sintesi per fase** (Fase, N "
        "reperti, %, data media, N US) e, se presenti, una **tabella della "
        "cronologia OxCal** per US;\n"
        "- spiega ESPLICITAMENTE e in modo argomentato: la scelta del modello "
        "(multinomiale o gaussiano) e il perché, il valore di K e perché, il "
        "significato di soglia/rumore, la selezione dei reperti;\n"
        "- interpreta i risultati (fasi, cronologia, residualità) in modo "
        "comprensibile;\n"
        "- richiama le figure pertinenti per nome quando illustrano un punto;\n"
        "- chiudi con conclusioni e raccomandazioni operative.\n"
        "Usa tabelle Markdown vere (con `|`). NON inventare numeri. Scrivi "
        "ESCLUSIVAMENTE in " + lang_name + ".")


# --------------------------------------------------------------- worker thread ---
class PalimpsestAIReportWorker(QThread):
    chunk = pyqtSignal(str)
    status = pyqtSignal(str)
    finished_ok = pyqtSignal(str)
    failed = pyqtSignal(str)

    # Universal per-call budget (accepted by every model); long reports are
    # achieved by continuing when the model is cut off.
    _PER_CALL_TOKENS = 4096

    def __init__(self, facts, config, lang_code, lang_name):
        super().__init__()
        self.facts = facts
        self.config = config
        self.lang_code = lang_code
        self.lang_name = lang_name

    @staticmethod
    def _supports_meta():
        try:
            return "meta" in inspect.signature(
                LLMProviderManager.stream_chat).parameters
        except Exception:
            return False

    def _run_agent(self, prompt, emit_live, max_rounds=8):
        supports_meta = self._supports_meta()
        if not supports_meta:
            max_rounds = 1
        messages = [{"role": "user", "content": prompt}]
        pieces = []
        for _ in range(max_rounds):
            meta = {}
            buf = []
            kwargs = dict(max_tokens=self._PER_CALL_TOKENS)
            if supports_meta:
                kwargs["meta"] = meta
            for ch in LLMProviderManager.stream_chat(
                    self.config, messages=messages, **kwargs):
                buf.append(ch)
                if emit_live:
                    self.chunk.emit(ch)
            text = "".join(buf)
            pieces.append(text)
            if not meta.get("truncated") or not text:
                break
            messages.append({"role": "assistant", "content": text})
            messages.append({"role": "user", "content":
                             "Continua ESATTAMENTE da dove ti sei interrotto, "
                             "senza ripetere il testo già scritto e senza "
                             "reintrodurre titoli o sezioni già presenti."})
        return "".join(pieces)

    def run(self):
        try:
            self.status.emit("methodology")
            method = self._run_agent(
                _methodologist_prompt(self.facts, self.lang_name),
                emit_live=False, max_rounds=4)
            self.status.emit("analysis")
            analysis = self._run_agent(
                _analyst_prompt(self.facts, self.lang_name),
                emit_live=False, max_rounds=4)
            self.status.emit("synthesis")
            final = self._run_agent(
                _synthesis_prompt(self.facts, self.lang_name, method, analysis),
                emit_live=True, max_rounds=10)
            self.finished_ok.emit(final.strip())
        except Exception as e:  # noqa: BLE001 — surfaced to the UI thread
            self.failed.emit(str(e))


# ----------------------------------------------------------- Markdown rendering ---
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)")


def _inline_runs(text):
    """Yield (text, bold, italic) tuples from a line of inline Markdown."""
    out = []
    for tok in _INLINE_RE.split(text):
        if not tok:
            continue
        if (tok.startswith("**") and tok.endswith("**")) or \
           (tok.startswith("__") and tok.endswith("__")):
            out.append((tok[2:-2], True, False))
        elif (tok.startswith("*") and tok.endswith("*")) or \
             (tok.startswith("_") and tok.endswith("_")):
            out.append((tok[1:-1], False, True))
        elif tok.startswith("`") and tok.endswith("`"):
            out.append((tok[1:-1], False, False))
        else:
            out.append((tok, False, False))
    return out


def _inline_html(text):
    """Inline Markdown -> minimal reportlab markup (escaped)."""
    parts = []
    for t, b, it in _inline_runs(text):
        s = html.escape(t)
        if b:
            s = "<b>%s</b>" % s
        if it:
            s = "<i>%s</i>" % s
        parts.append(s)
    return "".join(parts)


def _parse_markdown(md):
    """Parse a Markdown subset into block dicts: heading/para/bullet/numbered/table."""
    lines = (md or "").replace("\r\n", "\n").split("\n")
    blocks, para = [], []
    i, n, in_code = 0, len(lines), False

    def flush():
        if para:
            blocks.append({"type": "para", "text": " ".join(para).strip()})
            para.clear()

    while i < n:
        line = lines[i].rstrip()
        s = line.strip()
        if s.startswith("```"):
            flush(); in_code = not in_code; i += 1; continue
        if in_code:
            i += 1; continue
        if not s:
            flush(); i += 1; continue
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", s):
            flush(); i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            flush()
            blocks.append({"type": "heading", "level": len(m.group(1)),
                           "text": m.group(2).strip()})
            i += 1; continue
        if "|" in s and i + 1 < n and "-" in lines[i + 1] and \
           re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            flush()
            header = [c.strip() for c in s.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in
                             lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue
        if re.match(r"^[-*+]\s+", s):
            flush(); items = []
            while i < n and re.match(r"^\s*[-*+]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*+]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "bullet", "items": items})
            continue
        if re.match(r"^\d+[.)]\s+", s):
            flush(); items = []
            while i < n and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+[.)]\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "numbered", "items": items})
            continue
        para.append(s)
        i += 1
    flush()
    return blocks


def _appendix_md(facts):
    """Guaranteed data appendix so the numbers are always present."""
    parts = []
    ps = facts.get("phase_summary") or []
    if ps:
        parts += ["## Appendice — sintesi per fase",
                  _rows_md_table(ps, ["phase", "n_finds", "pct", "mean_date", "n_us"],
                                 ["Fase", "N reperti", "%", "Data media", "N US"])]
    ch = facts.get("chronology") or []
    if ch:
        parts += ["## Appendice — cronologia assoluta OxCal",
                  "Date calibrate per US (anni calendariali, negativo = a.C.):",
                  _rows_md_table(ch, ["us", "start", "end", "lab_code", "source"],
                                 ["US", "start", "end", "lab", "source"])]
    return "\n\n".join(parts)


def write_docx(md_text, facts, path, include_figs=True):
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    blocks = _parse_markdown(md_text)
    ap = _appendix_md(facts)
    if ap:
        blocks += _parse_markdown(ap)

    def runs(par, text):
        for t, b, it in _inline_runs(text):
            r = par.add_run(t)
            r.font.name = "Cambria"; r.font.size = Pt(11)
            r.bold = b; r.italic = it

    for blk in blocks:
        bt = blk["type"]
        if bt == "heading":
            doc.add_heading(blk["text"], level=min(blk["level"], 4))
        elif bt == "para":
            runs(doc.add_paragraph(), blk["text"])
        elif bt == "bullet":
            for it in blk["items"]:
                runs(doc.add_paragraph(style="List Bullet"), it)
        elif bt == "numbered":
            for it in blk["items"]:
                runs(doc.add_paragraph(style="List Number"), it)
        elif bt == "table":
            cols = len(blk["header"])
            t = doc.add_table(rows=1, cols=cols)
            try:
                t.style = "Table Grid"
            except Exception:
                pass
            for j, h in enumerate(blk["header"]):
                cell = t.rows[0].cells[j]
                cell.text = ""
                runs(cell.paragraphs[0], "**%s**" % h)
            for row in blk["rows"]:
                cells = t.add_row().cells
                for j in range(cols):
                    cells[j].text = row[j] if j < len(row) else ""

    if include_figs and facts.get("figures_dir") and facts.get("figures"):
        doc.add_heading("Figure", level=1)
        for key in facts["figures"]:
            png = os.path.join(facts["figures_dir"], key + ".png")
            if os.path.exists(png):
                try:
                    doc.add_picture(png, width=Inches(6.0))
                except Exception:
                    continue
                cap = doc.add_paragraph(FIGURE_CAPTIONS.get(key, key))
                for r in cap.runs:
                    r.font.name = "Cambria"; r.font.size = Pt(9); r.italic = True
    doc.save(path)


def write_pdf(md_text, facts, path, include_figs=True):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image,
        ListFlowable, ListItem)

    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=10.5,
                          leading=14, spaceAfter=6)
    cap = ParagraphStyle("cap", parent=body, fontSize=8, textColor=colors.grey,
                         spaceBefore=2, spaceAfter=10)
    hstyles = {1: ss["Heading1"], 2: ss["Heading2"], 3: ss["Heading3"],
               4: ss["Heading4"]}
    doc = SimpleDocTemplate(path, pagesize=A4, leftMargin=2 * cm,
                            rightMargin=2 * cm, topMargin=2 * cm,
                            bottomMargin=2 * cm)
    story = []
    blocks = _parse_markdown(md_text)
    ap = _appendix_md(facts)
    if ap:
        blocks += _parse_markdown(ap)

    for blk in blocks:
        bt = blk["type"]
        if bt == "heading":
            story.append(Paragraph(_inline_html(blk["text"]),
                                   hstyles.get(min(blk["level"], 4), hstyles[4])))
        elif bt == "para":
            story.append(Paragraph(_inline_html(blk["text"]), body))
        elif bt in ("bullet", "numbered"):
            items = [ListItem(Paragraph(_inline_html(it), body))
                     for it in blk["items"]]
            story.append(ListFlowable(
                items, bulletType="bullet" if bt == "bullet" else "1"))
            story.append(Spacer(1, 4))
        elif bt == "table":
            data = [[Paragraph(_inline_html(h), body) for h in blk["header"]]]
            for row in blk["rows"]:
                data.append([Paragraph(_inline_html(
                    row[j] if j < len(row) else ""), body)
                    for j in range(len(blk["header"]))])
            tbl = Table(data, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8E8EE")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
            story.append(tbl)
            story.append(Spacer(1, 8))

    if include_figs and facts.get("figures_dir") and facts.get("figures"):
        story.append(Paragraph("Figure", hstyles[1]))
        max_w = 16 * cm
        for key in facts["figures"]:
            png = os.path.join(facts["figures_dir"], key + ".png")
            if not os.path.exists(png):
                continue
            try:
                iw, ih = ImageReader(png).getSize()
                w = min(max_w, iw)
                h = w * ih / iw
                story.append(Image(png, width=w, height=h))
                story.append(Paragraph(FIGURE_CAPTIONS.get(key, key), cap))
            except Exception:
                continue
    doc.build(story)


# ----------------------------------------------------------------- dialog ---
class PalimpsestAIReportDialog(QDialog):
    def __init__(self, parent_dlg, facts):
        super().__init__(parent_dlg)
        self.p = parent_dlg
        self.facts = facts
        self.worker = None
        self._final_md = ""
        self.setWindowTitle("palimpsestr — Report AI (analisi descrittiva)")
        self.resize(780, 580)
        self._build_ui()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        nfig = len(self.facts.get("figures") or [])
        nph = len(self.facts.get("phase_summary") or [])
        nch = len(self.facts.get("chronology") or [])
        lay.addWidget(QLabel(
            "Relazione descrittiva dell'analisi SEF con agenti AI specializzati "
            "(metodologo, analista, redattore). Spiega le scelte di modello, K, "
            "soglia e reperti, interpreta i risultati e include tabelle e "
            "figure.<br><i>Dati disponibili: %d fasi, %d figure, %d date "
            "OxCal.</i>" % (nph, nfig, nch)))

        form = QFormLayout()
        self.llm = LLMSelectorWidget(scope="palimpsest_report", title="")
        form.addRow("Provider AI:", self.llm)
        self.combo_lang = QComboBox()
        for code, name in AI_LANGUAGES:
            self.combo_lang.addItem(name, code)
        form.addRow("Lingua del report:", self.combo_lang)
        lay.addLayout(form)

        opts = QHBoxLayout()
        self.chk_figs = QCheckBox("Includi figure")
        self.chk_figs.setChecked(nfig > 0)
        self.chk_figs.setEnabled(nfig > 0)
        opts.addWidget(self.chk_figs); opts.addStretch()
        lay.addLayout(opts)

        self.panel = QPlainTextEdit()
        self.panel.setReadOnly(True)
        self.panel.setPlaceholderText(
            "Il report apparirà qui durante la generazione.")
        lay.addWidget(self.panel)

        self.status = QLabel("")
        lay.addWidget(self.status)

        row = QHBoxLayout()
        self.btn_run = QPushButton("Genera report AI")
        self.btn_run.clicked.connect(self._run)
        self.btn_docx = QPushButton("Salva DOCX…")
        self.btn_docx.clicked.connect(self._save_docx)
        self.btn_pdf = QPushButton("Salva PDF…")
        self.btn_pdf.clicked.connect(self._save_pdf)
        self.btn_md = QPushButton("Salva Markdown…")
        self.btn_md.clicked.connect(self._save_md)
        btn_close = QPushButton("Chiudi")
        btn_close.clicked.connect(self.accept)
        for b in (self.btn_docx, self.btn_pdf, self.btn_md):
            b.setEnabled(False)
        row.addWidget(self.btn_run); row.addWidget(self.btn_docx)
        row.addWidget(self.btn_pdf); row.addWidget(self.btn_md)
        row.addWidget(btn_close)
        lay.addLayout(row)

    _STATUS = {
        "methodology": "Agente metodologo: spiego le scelte (modello, K, soglia)…",
        "analysis": "Agente analista: interpreto fasi, cronologia, intrusioni…",
        "synthesis": "Agente redattore: compongo il report finale…",
    }

    def _run(self):
        config = self.llm.get_config()
        if not config.model:
            QMessageBox.warning(self, "palimpsestr",
                                "Seleziona un modello AI nel selettore Provider.")
            return
        if config.needs_api_key and not config.api_key:
            QMessageBox.warning(self, "palimpsestr",
                                "Manca l'API key per il provider scelto.")
            return
        self.panel.setPlainText("")
        self._final_md = ""
        self.btn_run.setEnabled(False)
        for b in (self.btn_docx, self.btn_pdf, self.btn_md):
            b.setEnabled(False)
        self.worker = PalimpsestAIReportWorker(
            self.facts, config, self.combo_lang.currentData(),
            self.combo_lang.currentText())
        self.worker.status.connect(self._on_status)
        self.worker.chunk.connect(self._on_chunk)
        self.worker.finished_ok.connect(self._on_done)
        self.worker.failed.connect(self._on_failed)
        self.worker.start()

    def _on_status(self, phase):
        self.status.setText(self._STATUS.get(phase, phase))

    def _on_chunk(self, text):
        self.panel.insertPlainText(text)
        sb = self.panel.verticalScrollBar()
        sb.setValue(sb.maximum())

    def _on_done(self, md):
        self._final_md = md
        if md and not self.panel.toPlainText().strip():
            self.panel.setPlainText(md)
        self.status.setText("Report completato.")
        self.btn_run.setEnabled(True)
        for b in (self.btn_docx, self.btn_pdf, self.btn_md):
            b.setEnabled(True)

    def _on_failed(self, err):
        self.status.setText("")
        self.btn_run.setEnabled(True)
        QMessageBox.critical(self, "palimpsestr",
                             "Generazione del report AI fallita:\n%s" % err)

    # ----------------------------------------------------------------- save ---
    def _ask(self, title, flt, ext):
        fn, _ = QFileDialog.getSaveFileName(self, title, self.p.HOME, flt)
        if fn and not fn.lower().endswith(ext):
            fn += ext
        return fn

    def _save_md(self):
        if not self._final_md:
            return
        fn = self._ask("Salva report (Markdown)", "Markdown (*.md)", ".md")
        if not fn:
            return
        try:
            with open(fn, "w", encoding="utf-8") as f:
                f.write(self._final_md)
                ap = _appendix_md(self.facts)
                if ap:
                    f.write("\n\n" + ap + "\n")
        except Exception as e:
            QMessageBox.critical(self, "palimpsestr", "Salvataggio fallito:\n%s" % e)
            return
        self.status.setText("Salvato: %s" % fn)

    def _save_docx(self):
        if not self._final_md:
            return
        fn = self._ask("Salva report (DOCX)", "Word (*.docx)", ".docx")
        if not fn:
            return
        try:
            write_docx(self._final_md, self.facts, fn,
                       include_figs=self.chk_figs.isChecked())
        except Exception as e:
            QMessageBox.critical(self, "palimpsestr",
                                 "Salvataggio DOCX fallito:\n%s" % e)
            return
        self.status.setText("Salvato: %s" % fn)
        QMessageBox.information(self, "palimpsestr", "Report salvato:\n%s" % fn)

    def _save_pdf(self):
        if not self._final_md:
            return
        fn = self._ask("Salva report (PDF)", "PDF (*.pdf)", ".pdf")
        if not fn:
            return
        try:
            write_pdf(self._final_md, self.facts, fn,
                      include_figs=self.chk_figs.isChecked())
        except Exception as e:
            QMessageBox.critical(self, "palimpsestr",
                                 "Salvataggio PDF fallito:\n%s" % e)
            return
        self.status.setText("Salvato: %s" % fn)
        QMessageBox.information(self, "palimpsestr", "Report salvato:\n%s" % fn)
