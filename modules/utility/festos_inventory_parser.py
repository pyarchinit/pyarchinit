#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Parser specializzato per importare inventari di cassette da file DOCX (formato Festos)
"""
#id_tma
import re
import os
from docx import Document
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FestosInventoryParser:
    """Parser specifico per file inventario cassette archeologiche formato Festos"""
    
    def __init__(self):
        self.default_values = {
            'ogtm': 'CERAMICA',  # Tipo materiale di default
            'ldct': 'Magazzino',  # Localizzazione di default
            'stato': 'COMPLETO',
            'macc': 'ND',  # Categoria non determinata
            'dtzg': 'ND',  # Fascia cronologica non determinata
        }
        self.errors = []
        self.warnings = []
        
    def parse_file(self, file_path: str, sito: str = None) -> Tuple[int, List[str], List[str]]:
        """
        Parsa un file docx con inventario cassette e restituisce i record
        invece di salvarli direttamente nel database
        
        Args:
            file_path: Path del file docx
            sito: Nome del sito (se non specificato, cerca nel testo)
            
        Returns:
            Tuple con (lista record, lista errori, lista warning)
        """
        try:
            doc = Document(file_path)
            records = []
            current_magazzino = None
            
            # Prima passa: cerca informazioni generali
            for para in doc.paragraphs:
                text = para.text.strip()
                
                # Cerca indicazione del magazzino
                if 'Magazzino' in text or 'Magazzini' in text:
                    match = re.search(r'Magazzin[oi]\s*([\d\-]+)', text)
                    if match:
                        current_magazzino = f"Magazzino {match.group(1)}"
                    
                    # Estrai anche il sito se presente (es. "Festos - Magazzino 3-4")
                    if not sito and '-' in text:
                        parts = text.split('-')
                        if len(parts) >= 2:
                            potential_sito = parts[0].strip().replace('*', '')
                            if potential_sito and not any(char.isdigit() for char in potential_sito):
                                sito = potential_sito
                    continue
            
            # Se il sito non è stato trovato, prova a estrarlo dal nome file
            if not sito:
                file_name = os.path.basename(file_path)
                # Cerca pattern comuni nei nomi file
                for pattern in [r'^(\w+)_', r'^(\w+)-', r'^(\w+)\s']:
                    match = re.match(pattern, file_name)
                    if match:
                        sito = match.group(1)
                        break
                
                if not sito:
                    sito = 'DA DEFINIRE'
                    self.warnings.append(f"Sito non trovato nel file, impostato a: {sito}")
            
            # Seconda passa: parsa le tabelle
            table_count = 0
            for table in doc.tables:
                table_count += 1
                headers = []
                
                # Identifica gli headers
                if table.rows:
                    first_row = table.rows[0]
                    headers = [cell.text.strip() for cell in first_row.cells]
                    
                    # Verifica che sia una tabella di inventario
                    if not any('Cassett' in h for h in headers):
                        continue
                        
                # Processa le righe di dati
                for i, row in enumerate(table.rows[1:], start=2):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    # Salta righe vuote o separatori
                    if all(cell in ['', '-'] for cell in cells) or len(cells) < 3:
                        continue
                    
                    # Parsa la riga di dati
                    record_list = self._parse_row(cells, headers, sito, current_magazzino, table_count, i)
                    if record_list:
                        records.extend(record_list)
            
            # Restituisci i record invece di importarli direttamente
            return records, self.errors, self.warnings
            
        except Exception as e:
            self.errors.append(f"Errore generale nel parsing: {str(e)}")
            return [], self.errors, self.warnings
    
    def _parse_row(self, cells: List[str], headers: List[str], 
                   sito: str, magazzino: str, table_num: int, row_num: int) -> List[Dict]:
        """Parsa una singola riga della tabella"""
        if len(cells) < 3:
            return []
            
        # Identifica le colonne basandosi sugli headers
        cassette_col = 0
        desc_col = 1
        anno_col = 2
        
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if 'cassett' in header_lower:
                cassette_col = i
            elif 'descr' in header_lower:
                desc_col = i
            elif 'anno' in header_lower:
                anno_col = i
        
        # Estrai i valori
        cassette = cells[cassette_col] if cassette_col < len(cells) else ''
        descrizione = cells[desc_col] if desc_col < len(cells) else ''
        anno = cells[anno_col] if anno_col < len(cells) else ''
        
        if not cassette or cassette == '-':
            return []
            
        # Parsa range di cassette (es. "1-72")
        cassette_list = self._parse_cassette_range(cassette)
        
        records = []
        for cass_num in cassette_list:
            record = self.default_values.copy()
            
            # Campi principali
            record['sito'] = sito
            record['dscu'] = f"Cassetta {cass_num}"
            record['ldcn'] = f"{magazzino} - Cassetta {cass_num}" if magazzino else f"Cassetta {cass_num}"
            record['cassetta'] = cass_num
            record['localita'] = sito  # Usa il sito come località di default
            
            # Localizzazione
            if magazzino:
                record['ldct'] = 'Magazzino'
                record['scan'] = magazzino  # Usa come denominazione scavo
            
            # Area (estrai se presente nella descrizione)
            area = self._extract_area(descrizione)
            if area:
                record['area'] = area
            else:
                record['area'] = ''
            
            # Descrizione
            if descrizione:
                record['deso'] = descrizione
                
                # Estrai tipo materiale dalla descrizione
                tipo_materiale = self._extract_tipo_materiale(descrizione)
                record['ogtm'] = tipo_materiale
                
                # Estrai numero reperti se presente
                n_reperti = self._extract_n_reperti(descrizione)
                if n_reperti:
                    record['n_reperti'] = str(n_reperti)
                else:
                    record['n_reperti'] = ''
                    
                # Estrai categoria dalla descrizione
                categoria = self._extract_categoria(descrizione)
                record['macc'] = categoria
            
            # Anno e datazione
            if anno:
                anno_clean = self._clean_anno(anno)
                if anno_clean:
                    # Imposta fascia cronologica basata sull'anno
                    record['dtzg'] = self._get_fascia_cronologica(anno_clean, descrizione)
                    record['cronologie'] = str(anno_clean)
                    
                    # Se c'è un range di anni (es. "1965-66")
                    if '-' in anno and not re.search(r'\d{4}-\d{4}', anno):
                        match = re.search(r'(\d{4})-(\d{2})', anno)
                        if match:
                            start_year = int(match.group(1))
                            end_suffix = match.group(2)
                            end_year = int(str(start_year)[:2] + end_suffix)
                            record['cronologie'] = f"{start_year}-{end_year}"
            
            # Aggiungi riferimenti
            record['riferimenti'] = f"Tabella {table_num}, Riga {row_num}"
            
            records.append(record)
        
        return records
    
    def _parse_cassette_range(self, cassette_str: str) -> List[str]:
        """Parsa un range di cassette (es. '1-72' -> ['1', '2', ..., '72'])"""
        cassette_str = cassette_str.strip()
        
        # Gestisci lettere (es. "246d-259")
        if re.match(r'^\d+[a-z]?-\d+[a-z]?$', cassette_str):
            match = re.match(r'^(\d+)([a-z]?)-(\d+)([a-z]?)$', cassette_str)
            if match:
                start_num = int(match.group(1))
                start_letter = match.group(2)
                end_num = int(match.group(3))
                end_letter = match.group(4)
                
                # Se c'è una lettera, tratta come singola cassetta
                if start_letter or end_letter:
                    return [cassette_str]
                else:
                    # Genera il range
                    return [str(i) for i in range(start_num, end_num + 1)]
        
        # Range semplice
        if '-' in cassette_str and not re.search(r'[a-z]', cassette_str):
            try:
                parts = cassette_str.split('-')
                if len(parts) == 2:
                    start = int(parts[0].strip())
                    end = int(parts[1].strip())
                    if start <= end:
                        return [str(i) for i in range(start, end + 1)]
            except ValueError:
                pass
                
        # Cassetta singola
        return [cassette_str]
    
    def _clean_anno(self, anno_str: str) -> Optional[int]:
        """Pulisce e converte l'anno"""
        anno_str = anno_str.strip()
        
        # Cerca il primo anno valido a 4 cifre
        match = re.search(r'(\d{4})', anno_str)
        if match:
            return int(match.group(1))
            
        return None
    
    def _extract_area(self, descrizione: str) -> str:
        """Estrae l'area dalla descrizione"""
        desc_lower = descrizione.lower()
        
        # Pattern comuni per aree
        patterns = [
            r'area\s+([A-Z]+)',
            r'settore\s+([A-Z]+)',
            r'vano\s+([A-Z]+)',
            r'piazzale\s+([IVX]+)',
            r'\b([A-Z]{1,3})\b',  # Lettere singole o doppie maiuscole
        ]
        
        for pattern in patterns:
            match = re.search(pattern, descrizione)
            if match:
                return match.group(1)
        
        return ''
    
    def _extract_tipo_materiale(self, descrizione: str) -> str:
        """Estrae il tipo di materiale dalla descrizione"""
        desc_lower = descrizione.lower()
        
        # Mapping descrizione -> OGTM
        material_mapping = {
            'ceramica': 'CERAMICA',
            'ceramic': 'CERAMICA',
            'framm': 'CERAMICA',
            'vas': 'CERAMICA',
            'pitho': 'CERAMICA',
            'anfora': 'CERAMICA',
            'terracotta': 'CERAMICA',
            'oss': 'FAUNA',
            'faun': 'FAUNA',
            'conchigli': 'MALACOFAUNA',
            'malaco': 'MALACOFAUNA',
            'pietra': 'INDUSTRIA LITICA',
            'litic': 'INDUSTRIA LITICA',
            'selce': 'INDUSTRIA LITICA',
            'obsidian': 'INDUSTRIA LITICA',
            'macin': 'INDUSTRIA LITICA',
            'ferro': 'METALLO',
            'bronz': 'METALLO',
            'metal': 'METALLO',
            'rame': 'METALLO',
            'piombo': 'METALLO',
            'stucc': 'INTONACO',
            'intonac': 'INTONACO',
            'affresc': 'INTONACO',
            'blocch': 'MATERIALE DA COSTRUZIONE',
            'mattone': 'MATERIALE DA COSTRUZIONE',
            'tegol': 'MATERIALE DA COSTRUZIONE',
            'coppo': 'MATERIALE DA COSTRUZIONE',
            'calcestru': 'MATERIALE DA COSTRUZIONE',
            'astraki': 'MATERIALE DA COSTRUZIONE',
            'vetro': 'VETRO',
            'pasta vitr': 'VETRO',
            'carbone': 'CARBONI',
            'legno': 'LEGNO',
            'osso lavorat': 'INDUSTRIA SU OSSO',
            'avorio': 'INDUSTRIA SU OSSO'
        }
        
        for key, value in material_mapping.items():
            if key in desc_lower:
                return value
                
        return 'CERAMICA'  # Default
    
    def _extract_n_reperti(self, descrizione: str) -> Optional[int]:
        """Estrae il numero di reperti dalla descrizione se presente"""
        # Cerca pattern numerici
        patterns = [
            r'(\d+)\s+framm',
            r'(\d+)\s+pezz',
            r'(\d+)\s+element',
            r'(\d+)\s+repert',
            r'(\d+)\s+esemplar',
            r'n\.\s*(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, descrizione.lower())
            if match:
                return int(match.group(1))
                
        return None
    
    def _extract_categoria(self, descrizione: str) -> str:
        """Estrae la categoria dalla descrizione"""
        desc_lower = descrizione.lower()
        
        # Categorie basate sul contenuto
        if any(word in desc_lower for word in ['sigla', 'inventaria', 'cataloga']):
            return 'INV'  # Inventariato
        elif any(word in desc_lower for word in ['scelt', 'seleziona']):
            return 'SEL'  # Selezionato
        elif any(word in desc_lower for word in ['diagnostic', 'studio']):
            return 'DIAG'  # Diagnostico
        elif any(word in desc_lower for word in ['sporad', 'superf']):
            return 'SPOR'  # Sporadico
        else:
            return 'GEN'  # Generico
    
    def _get_fascia_cronologica(self, anno: int, descrizione: str) -> str:
        """Determina la fascia cronologica basata sull'anno e la descrizione"""
        desc_lower = descrizione.lower()
        
        # Cerca indicazioni cronologiche nella descrizione
        if any(word in desc_lower for word in ['minoic', 'mm', 'tm']):
            if 'medio' in desc_lower or 'mm' in desc_lower:
                return 'Medio Minoico'
            elif 'tardo' in desc_lower or 'tm' in desc_lower:
                return 'Tardo Minoico'
            else:
                return 'Minoico'
        elif any(word in desc_lower for word in ['micene', 'te']):
            return 'Miceneo'
        elif any(word in desc_lower for word in ['geometric', 'pg']):
            if 'proto' in desc_lower:
                return 'Protogeometrico'
            else:
                return 'Geometrico'
        elif any(word in desc_lower for word in ['ellenis', 'greco', 'roman']):
            if 'roman' in desc_lower:
                return 'Romano'
            else:
                return 'Ellenistico'
        elif any(word in desc_lower for word in ['bizantin', 'medieval', 'venezian']):
            if 'venezian' in desc_lower:
                return 'Veneziano'
            else:
                return 'Bizantino'
        
        # Se non ci sono indicazioni nella descrizione, usa l'anno dello scavo
        if anno and anno >= 1950:
            return f'Scavo {anno}'
        else:
            return 'ND'  # Non determinato
