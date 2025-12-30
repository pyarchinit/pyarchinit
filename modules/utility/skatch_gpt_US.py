import base64
import difflib
import json
import os
import sys
import mimetypes
import csv
import re
import cv2
import docx
import pymupdf as fitz
from openai import OpenAI
from anthropic import Anthropic
import requests
from collections import defaultdict
from qgis.PyQt.QtWidgets import QTextBrowser, QListWidgetItem, QComboBox, QProgressBar, QScrollArea, QInputDialog, QMessageBox, QFileDialog, QLabel, QMainWindow, \
    QApplication, QGridLayout, QWidget, QTextEdit, QPushButton, QListWidget, QSplitter
from qgis.PyQt.QtCore import QThread, pyqtSignal, pyqtSlot, Qt

from qgis.PyQt.QtGui import QIcon
from PIL import Image
from modules.db.pyarchinit_conn_strings import Connection
from modules.db.pyarchinit_utility import Utility
from modules.utility.pyarchinit_media_utility import Media_utility, Media_utility_resize


class Worker(QThread):
    progress_updated = pyqtSignal(int)  # signal for progress updates
    content_updated = pyqtSignal(str)  # signal for content updates
    tokens_used_updated = pyqtSignal(int, float)  # signal for tokens used updates with cost

    def __init__(self, headers, params, url, is_image=False, image_width=512, image_height=512):
        super().__init__()
        self.headers = headers
        self.params = params
        self.is_image = is_image
        self.image_width = image_width
        self.image_height = image_height
        self.url = url

    def run(self):
        try:
            client = requests.Session()

            response = client.post( headers=self.headers, json=self.params,
                                   url=self.url,stream=True)

            if response.status_code != 200:
                print(f"Error: HTTP {response.status_code} - {response.text}")
                return

            total_received_content = 0
            total_content_length = 4096  # Assuming max_tokens is the total content length
            reply = ""
            tokens_used = 0
            input_cost_per_token = 5.00 / 1_000_000  # Cost per input token in USD
            output_cost_per_token = 15.00 / 1_000_000  # Cost per output token in USD

            # Calculate image cost based on dimensions
            if self.is_image:
                tiles_x = (self.image_width + 511) // 512
                tiles_y = (self.image_height + 511) // 512
                total_tiles = tiles_x * tiles_y
                base_tokens = 85
                tile_tokens = 170 * total_tiles
                total_image_tokens = base_tokens + tile_tokens
                image_cost_per_tile = total_image_tokens * input_cost_per_token
            else:
                image_cost_per_tile = 0

            for line in response.iter_lines():
                if line:
                    decoded_line = json.loads(line.decode('utf-8').strip()[len("data: "):])
                    if 'choices' in decoded_line:
                        for choice in decoded_line['choices']:
                            if 'delta' in choice and 'content' in choice['delta']:
                                content_chunk = choice['delta']['content']
                                if content_chunk:
                                    reply += content_chunk
                                    total_received_content += len(content_chunk)
                                    tokens_used += len(content_chunk.split())
                                    progress_percentage = int(total_received_content / 100)
                                    # Calculate progress percentage
                                    self.progress_updated.emit(progress_percentage)

                                    total_cost = tokens_used * (
                                            input_cost_per_token + output_cost_per_token) + image_cost_per_tile
                                    self.tokens_used_updated.emit(tokens_used, total_cost)
                                    self.content_updated.emit(content_chunk)

            self.progress_updated.emit(100)  # Ensure progress bar reaches 100%
        except Exception as e:
            print(f"Error in worker thread: {e}")


class GPTWindow(QMainWindow):
    HOME = os.environ.get('PYARCHINIT_HOME', '')

    def __init__(self, selected_images=None,dbmanager=None, main_class=None):
        super().__init__()
        self.setWindowTitle("PyArchInit - GPT Sketch")
        self.setGeometry(100, 100, 600, 800)
        self.set_icon(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../resources/icons/gpt.png")))
        self.selected_images = selected_images or []
        # Layout principale
        layout = QGridLayout()
        self.DB_MANAGER=dbmanager
        self.mainclass=main_class
        # creazione prompt
        self.prompt_label = QTextEdit()
        self.prompt_label.setFixedHeight(150)
        self.prompt_label.setPlaceholderText(
            "Esempio di prompt:\n"
            "Se vuoi estrarre informazioni da un'immagine, scrivi:\n"
            "Estrai le informazioni dalla lavagnetta.\n e schiacci Importa immagine\n\n"
            "Se vuoi correggere un testo, scrivi:\n"
            "Fai una correzione del testo.\n"
            "e schiacci Importa Documento"
        )
        layout.addWidget(self.prompt_label, 1, 0, 1, 3)

        # Aggiungi un selettore per il modello AI
        self.model_selector = QComboBox()
        self.model_selector.addItems(["GPT-4o", "Claude Sonnet 3.5"])
        layout.addWidget(QLabel("Select AI Model:"), 0, 0)
        layout.addWidget(self.model_selector, 0, 1)

        # Bottoni per l'interazione
        self.btn_import3 = QPushButton("Analizza Immagini Selezionate")
        self.btn_import3.clicked.connect(self.analyze_selected_images)
        self.btn_import = QPushButton("Importa Immagine")
        self.btn_import.clicked.connect(self.scketchgpt)
        self.btn_import2 = QPushButton("Importa Documento")
        self.btn_import2.clicked.connect(self.docchgpt)
        self.progress = QProgressBar()
        self.token_counter = QLabel("Tokens used: 0 - Total cost: $0.0000")
        layout.addWidget(self.btn_import3, 2, 0)
        layout.addWidget(self.btn_import, 2, 1)
        layout.addWidget(self.btn_import2, 2, 2)
        layout.addWidget(self.progress, 3, 0, 1, 0)
        layout.addWidget(self.token_counter, 3, 0, 1, 1)

        # Lista per visualizzare le risposte di GPT
        self.listWidget_ai = QTextBrowser()

        # Add scrollbar to listWidget_ai
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(self.listWidget_ai)

        # Use QSplitter to make listWidget_ai resizable
        splitter = QSplitter()
        splitter.addWidget(scroll_area)
        splitter.setStretchFactor(0, 1)  # Make listWidget_ai expandable

        layout.addWidget(splitter, 4, 0, 1, 3)

        # Widget container per i layout
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)



    def analyze_selected_images(self):
        def get_image_metadata(file_path):
            with Image.open(file_path) as img:
                dpi = img.info.get('dpi', (72, 72))  # Default DPI if not found
                width, height = img.size
                return {
                    'dpi': dpi,
                    'width': width,
                    'height': height
                }

        self.listWidget_ai.clear()
        if self.selected_images:
            selected_model = self.model_selector.currentText()
            for file_path in self.selected_images:
                # Get image metadata
                metadata = get_image_metadata(file_path)
                dpi = metadata['dpi']
                width = metadata['width']
                height = metadata['height']

                # Prepare the prompt with additional requests for information
                prompt = self.prompt_label.toPlainText()
                prompt += f"\nFile Name: {os.path.basename(file_path)}\n"
                prompt += f"File Path: {file_path}\n"
                prompt += f"DPI: {dpi[0]} x {dpi[1]}\n"
                prompt += f"Dimensioni: {width} x {height}\n"

                if selected_model == "GPT-4o":
                    response = self.ask_gpt4(prompt, self.apikey_gpt(), file_path)
                else:
                    response = self.ask_claude(prompt, self.apikey_claude(), file_path)

                # Display the response in the UI
                #self.listWidget_ai.append(f"AI Response for {os.path.basename(file_path)}:")
                #self.listWidget_ai.append(response)

                # Extract URLs from the response and make them clickable
                if 'http' or 'https' or 'www' in response:
                    self.extract_and_display_links(response)

        else:
            self.listWidget_ai.setPlainText("No images selected for analysis.")

    def extract_and_display_links(self, response):
        # Use regex to find URLs in the response
        urls = re.findall(r'(https?://[^\s]+)', response)

        if urls:
            for url in urls:
                # Create a clickable link in the text browser
                self.listWidget_ai.append(f'<a href="{url}">{url}</a>')

            # Allow external links to be opened in the default browser
            self.listWidget_ai.setOpenExternalLinks(True)
        else:
            QMessageBox.information(self, "Info", "No links found in the response.")


    def set_icon(self, icon_path):
        self.setWindowIcon(QIcon(icon_path))



    def start_worker(self, headers, params, url, is_image=False):
        self.worker = Worker(headers, params, url)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.content_updated.connect(self.update_content)
        self.worker.tokens_used_updated.connect(self.update_tokens_used)
        self.worker.start()

    def apikey_claude(self):
        BIN = '{}{}{}'.format(self.HOME, os.sep, "bin")
        api_key = ""
        path_key = os.path.join(BIN, 'claude_api_key.txt')

        if os.path.exists(path_key):
            with open(path_key, 'r') as f:
                api_key = f.read().strip()
        else:
            api_key, ok = QInputDialog.getText(None, 'Claude API Key', 'Enter Claude API key:')
            if ok and api_key:
                with open(path_key, 'w') as f:
                    f.write(api_key)

        return api_key

    def apikey_gpt(self):
        # HOME = os.environ['PYARCHINIT_HOME']
        BIN = '{}{}{}'.format(self.HOME, os.sep, "bin")
        api_key = ""
        # Verifica se il file gpt_api_key.txt esiste
        path_key = os.path.join(BIN, 'gpt_api_key.txt')
        if os.path.exists(path_key):

            # Leggi l'API Key dal file
            with open(path_key, 'r') as f:
                api_key = f.read().strip()
                try:

                    return api_key

                except:
                    reply = QMessageBox.question(None, 'Warning', 'Apikey non valida' + '\n'
                                                 + 'Clicca ok per inserire la chiave',
                                                 QMessageBox.Ok | QMessageBox.Cancel)
                    if reply == QMessageBox.Ok:

                        api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey valida:')
                        if ok:
                            # Salva la nuova API Key nel file
                            with open(path_key, 'w') as f:
                                f.write(api_key)
                                f.close()
                            with open(path_key, 'r') as f:
                                api_key = f.read().strip()
                    else:
                        return api_key


        else:
            # Chiedi all'utente di inserire una nuova API Key
            api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey:')
            if ok:
                # Salva la nuova API Key nel file
                with open(path_key, 'w') as f:
                    f.write(api_key)
                    f.close()
                with open(path_key, 'r') as f:
                    api_key = f.read().strip()

        return api_key

    def ask_gpt4(self, prompt, apikey, file_path, is_image=True):
        try:
            client = OpenAI(api_key=apikey)

            if is_image:
                with open(file_path, "rb") as image_file:
                    base64_image = base64.b64encode(image_file.read()).decode('utf-8')

                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                        ]
                    }
                ]
            else:
                file_text = self.extract_text_from_file(file_path)
                messages = [
                    {"role": "system", "content": "You are an assistant for analyzing documents."},
                    {"role": "user", "content": prompt + "\n\n" + file_text}
                ]

            stream = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                max_tokens=4096,
                temperature=0.5,
                stream=True
            )

            full_response = ""
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    full_response += content
                    self.update_content(content)

            return full_response
        except Exception as e:
            QMessageBox.warning(self, "Error", f"An error occurred in ask_gpt4: {str(e)}")
            return ""

    def ask_claude(self, prompt, apikey, file_path, is_image=True):
        client = Anthropic(api_key=apikey)
        try:
            if is_image:
                # Determine the media type based on the file extension
                _, file_extension = os.path.splitext(file_path)
                file_extension = file_extension.lower()  # Normalize to lowercase

                # Map file extensions to media types
                media_type_map = {
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.png': 'image/png',
                    '.tiff': 'image/tiff',
                    '.tif': 'image/tiff'
                }

                media_type = media_type_map.get(file_extension, 'image/jpeg')  # Default to 'image/jpeg' if unknown

                with open(file_path, "rb") as image_file:
                    base64_image = base64.b64encode(image_file.read()).decode('utf-8')

                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": base64_image
                                }
                            }
                        ]
                    }
                ]
            else:
                file_text = self.extract_text_from_file(file_path)
                messages = [
                    {
                        "role": "user",
                        "content": prompt + "\n\n" + file_text
                    }
                ]

            stream = client.messages.create(
                model="claude-3-5-sonnet-20240620",
                max_tokens=4096,
                temperature=0.5,
                messages=messages,
                stream=True
            )
            full_response=''
            for event in stream:
                if hasattr(event, 'delta') and hasattr(event.delta, 'text'):
                    text_chunk = event.delta.text
                    full_response += text_chunk  # Accumulate the response
                    self.update_content(text_chunk)  # Update content in the UI
                elif hasattr(event, 'message'):
                    if event.message.content:
                        for content in event.message.content:
                            if content.type == 'text':
                                full_response += content.text  # Accumulate the response
                                self.update_content(content.text)  # Update content in the UI

            return full_response  # Return the accumulated response
        except Exception as e:
            QMessageBox.warning(self, "Error", f"An error occurred in ask_claude: {str(e)}")
            return ""




    def manual_input(self, missing_fields):
        info = {}
        for field in missing_fields:
            value, ok = QInputDialog.getText(self, "Input richiesto", f"Inserisci {field}:")
            if ok and value:
                info[field] = value.strip()
            else:
                return None
        return info

    def check_existing_record(self, info):
        search_dict = {
            'sito': "'" + info['sito'] + "'",
            'area': "'" + info['area'] + "'",
            'us': "'" + str(info['us']) + "'",
            'unita_tipo': "'" + info['unita_tipo'] + "'"
        }
        res = self.DB_MANAGER.query_bool(search_dict, self.mainclass.MAPPER_TABLE_CLASS)
        return res[0].id_us if res else None

    def is_image_associated(self, file_path, record_id):
        filename = os.path.basename(file_path)
        search_dict = {
            'id_entity': record_id,
            'entity_type': "'US'",
            'media_name': "'" + filename + "'"
        }
        res = self.DB_MANAGER.query_bool(search_dict, 'MEDIATOENTITY')
        return bool(res)

    def scketchgpt(self):
        self.listWidget_ai.clear()

        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(self, "Select Images", "", "Images (*.png *.jpg *.jpeg)")

        if not file_paths:
            self.listWidget_ai.setPlainText("Image selection was canceled.")
            return

        # Ottieni il prompt attuale
        prompt = self.prompt_label.toPlainText()
        prompt += (
            'L\'unità tipo può essere solo US o USM e il numero che segue è il numero US.'
            'Quindi se leggi ad esempio US 1, USM10, USM 1 o US23, essi saranno '
            'unità tipo e numero US che è sempre un integer come l\'area. Essi devono essere restituiti come chiave valore '
            'per riga esempio US: 1 Unità Tipo: US.')

        if not prompt:
            prompt = ("Analizza questa immagine e estrai le seguenti informazioni: Sito, Area, US (Unità Stratigrafica)"
                      ", e Unità Tipo. Fornisci queste informazioni nel formato 'Chiave: Valore', una per riga.")
            prompt += (
                "L'unità tipo può essere solo US o USM e il numero che segue è il numero US."
                "Quindi se leggi ad esempio US 1, USM10, USM 1 o US23, essi saranno "
                "unità tipo e numero US che è sempre un integer come l'area. Essi devono essere restituiti come chiave valore "
                "per riga esempio US: 1 Unità Tipo: US.")

        selected_model = self.model_selector.currentText()

        for file_path in file_paths:
            try:
                # Controlla se il prompt indica una Harris matrix
                if self.is_harris_matrix(prompt):
                    # Chiedi conferma all'utente
                    confirm = QMessageBox.question(self, "Confirm Harris Matrix",
                                                   "The loaded image appears to be a Harris matrix. Do you want to proceed with detailed analysis?",
                                                   QMessageBox.Yes | QMessageBox.No)

                    if confirm == QMessageBox.Yes:
                        # Usa il prompt dettagliato per l'analisi della Harris matrix
                        detailed_prompt = (
                            "# Prompt: Analisi Dettagliata del Diagramma di Harris Matrix\n\n"
                            "Analizza attentamente l'immagine fornita di un diagramma di Harris Matrix e identifica le seguenti relazioni stratigrafiche, distinguendo tra Unità Stratigrafiche (US) e Unità Stratigrafiche Murarie (USM):\n\n"
                            "1. Identificazione delle unità:\n"
                            "   - US: Generalmente indicate da numeri all'interno di cerchi o rettangoli.\n"
                            "   - USM: Spesso indicate da numeri preceduti da 'M' o in un formato distintivo.\n\n"
                            "2. Relazioni stratigrafiche (identifica sia la relazione diretta che l'inversa):\n"
                            "   a) Copre / Coperto da:\n"
                            "      - Linee continue verticali tra le unità.\n"
                            "      - Es: Se A copre B, allora B è coperto da A.\n\n"
                            "   b) Taglia / Tagliato da:\n"
                            "      - Linee tratteggiate o punteggiate tra le unità.\n"
                            "      - Es: Se A taglia B, allora B è tagliato da A.\n\n"
                            "   c) Riempie / Riempito da:\n"
                            "      - Linee continue con frecce rivolte verso il basso.\n"
                            "      - Es: Se A riempie B, allora B è riempito da A.\n\n"
                            "   d) Si appoggia / Gli si appoggia:\n"
                            "      - Linee continue verticali, spesso con annotazioni.\n"
                            "      - Es: Se A si appoggia a B, allora B è appoggiato da A.\n\n"

                            "   e) Uguale a:\n"
                            "      - Linee orizzontali doppie tra unità.\n"
                            "      - Relazione simmetrica: Se A è uguale a B, allora B è uguale a A.\n\n"
                            "   f) Si lega a:\n"
                            "      - Usato specificamente per USM.\n"
                            "      - Rappresentato come 'Uguale a' con linee orizzontali doppie.\n"
                            "      - Relazione simmetrica: Se USM A si lega a USM B, allora USM B si lega a USM A.\n\n"

                            "3. Sequenza temporale:\n"
                            "   - Determina la sequenza cronologica relativa basandoti sulle relazioni identificate.\n"
                            "   - Le unità più in alto nel diagramma sono generalmente più recenti.\n\n"
                            "4. Gruppi o fasi:\n"
                            "   - Identifica eventuali raggruppamenti di US/USM, spesso indicati da riquadri o colori.\n\n"
                            "5. Annotazioni o legenda:\n"
                            "   - Interpreta eventuali note, etichette o legende presenti nel diagramma.\n\n"
                            "6. Distinzione US/USM:\n"
                            "   - Evidenzia le differenze nelle relazioni applicate a US e USM.\n"
                            "   - Nota in particolare l'uso di 'Si lega a' esclusivamente per USM.\n\n"
                            "Fornisci un'analisi dettagliata che includa:\n"
                            "1. Elenco di tutte le US e USM identificate.\n"
                            "2. Descrizione di tutte le relazioni stratigrafiche, includendo sempre sia la relazione diretta che l'inversa.\n"
                            "3. Interpretazione della sequenza temporale.\n"
                            "4. Osservazioni su eventuali gruppi o fasi identificate.\n"
                            "5. Commenti su qualsiasi caratteristica distintiva o insolita nel diagramma.\n\n"
                            "Assicurati di evidenziare chiaramente la distinzione tra US e USM nelle tue osservazioni e di utilizzare correttamente la terminologia per ciascun tipo di unità."
                        )

                        # Chiama il modello appropriato con il nuovo prompt
                        if selected_model == "GPT-4o":
                            response = self.ask_gpt4(detailed_prompt, self.apikey_gpt(), file_path)
                        else:  # Questo implica che sia selezionato Claude Sonnet
                            response = self.ask_claude(detailed_prompt, self.apikey_claude(), file_path)

                        #extracted_info = self.extract_info_from_response(response)
                        extracted_info = self.process_ai_response(response)
                        QMessageBox.information(self, "Info", f"{extracted_info}")
                        # Assicurati che extracted_info_rapporti sia una lista di liste
                        #if not isinstance(extracted_info, dict) or not all(
                                #isinstance(r, list) for rels in extracted_info.values() for r in rels):
                            #raise TypeError("extracted_info_rapporti must be a dict with lists of lists as values")



                        # Crea un record per ogni US estratta
                        for us, info in extracted_info.items():
                            # Verifica se il record esiste già
                            existing_record = self.check_existing_record(info)
                            if existing_record:
                                continue  # Passa oltre se esiste

                            # Aggiungi le informazioni dei rapporti (lista di liste)
                            rapporti = info.get("rapporti", [])
                            if not isinstance(rapporti, list) or not all(isinstance(r, list) for r in rapporti):
                                raise TypeError("Each 'rapporti' must be a list of lists")

                            # Crea il nuovo record nel database
                            new_record_id = self.create_new_record_matrix(info)

                            if new_record_id is not None:
                                # Vai al record appena creato
                                self.go_to_us_record(info["sito"], info["area"], info["us"], info["unita_tipo"])
                                self.mainclass.iconListWidget.clear()
                                self.mainclass.on_pushButton_save_pressed()
                                self.mainclass.on_pushButton_view_all_pressed()

                            else:
                                QMessageBox.warning(self, "Warning",
                                                    f"Failed to create new record for US {info['us']}")




                    else:
                        # Se l'utente sceglie di non procedere, continua con il flusso normale
                        return
                else:
                    # Logica esistente per immagini non-Harris matrix
                    if selected_model == "GPT-4o":
                        response = self.ask_gpt4(prompt, self.apikey_gpt(), file_path)
                    else:  # Questo implica che sia selezionato Claude Sonnet
                        response = self.ask_claude(prompt, self.apikey_claude(), file_path)

                    extracted_info = self.extract_info_from_response(response)

                    # Controlla se il record esiste già
                    existing_record = self.check_existing_record(extracted_info)

                    if existing_record:
                        # Controlla se l'immagine è già associata
                        if self.is_image_associated(file_path, existing_record):
                            QMessageBox.information(self, "Info",
                                                    f"The record and image for {os.path.basename(file_path)} already exist.")
                        else:
                            # Associa l'immagine con il record esistente
                            self.associate_image_with_record(file_path, existing_record)
                            QMessageBox.information(self, "Success",
                                                    f"Image {os.path.basename(file_path)} associated with existing record.")
                    else:
                        # Mostra le informazioni estratte all'utente e chiedi conferma
                        if self.confirm_information(extracted_info, file_path):
                            # Crea un nuovo record nel database
                            new_record_id = self.create_new_record(extracted_info)

                            if new_record_id is not None:
                                sito = extracted_info.get('sito')
                                area = extracted_info.get('area')
                                us = extracted_info.get('us')
                                unita_tipo = extracted_info.get('unita_tipo')

                                self.go_to_us_record(sito, area, us, unita_tipo)

                                self.mainclass.iconListWidget.clear()
                                # Associa l'immagine con il nuovo record
                                self.associate_image_with_record(file_path, new_record_id)
                                self.mainclass.iconListWidget.update()

                                self.mainclass.on_pushButton_save_pressed()
                                self.mainclass.on_pushButton_view_all_pressed()
                                self.mainclass.connect_p()

                                QMessageBox.information(self, "Success",
                                                        f"New record created and image {os.path.basename(file_path)} associated successfully.")
                            else:
                                QMessageBox.warning(self, "Warning",
                                                    f"Failed to create new record for {os.path.basename(file_path)}.")
                        else:
                            QMessageBox.information(self, "Aborted",
                                                    f"Information insertion aborted for {os.path.basename(file_path)}.")


            except ValueError as e:
                QMessageBox.warning(self, "Error",
                                f"An error occurred while processing {os.path.basename(file_path)}: {str(e)}")

            QApplication.processEvents()

        self.listWidget_ai.append("Processing completed.")

    def check_existing_record_matrix(self, info):
        query = (
            "SELECT * FROM us_table WHERE "
            "sito = {sito} AND "
            "area = {area} AND "
            "us = {us}"
        )

        query_parameters = {
            'sito': "'" + info['sito'] + "'",
            'area': "'" + info['area'] + "'",
            'us': info['us']
        }

        try:
            cursor = self.db_connection.cursor()
            cursor.execute(query.format(**query_parameters))
            record = cursor.fetchone()
            return record
        except Exception as e:
            print(f"Errore nell'esecuzione della query: {e}")
            return None

    def is_harris_matrix(self, prompt):
        # Controlla se il prompt contiene "Harris matrix" o "matrix"
        if "harris matrix" in prompt.lower() or "matrix" in prompt.lower():
            return True
        return False

    def go_to_us_record(self, sito, area, us, unita_tipo):
        # Implement the logic to navigate to the US record based on the provided parameters
        search_dict = {
            'sito': "'" + sito + "'",
            'area': "'" + area + "'",
            'us': us,
            'unita_tipo': "'" + unita_tipo + "'",
        }

        # Remove empty items from the search dictionary
        u = Utility()
        search_dict = u.remove_empty_items_fr_dict(search_dict)

        # Query the database for the US record
        res = self.DB_MANAGER.query_bool(search_dict, self.mainclass.MAPPER_TABLE_CLASS)

        # Check if we found any records
        if not res:
            QMessageBox.warning(self, "Warning", "No US records found for the selected criteria.", QMessageBox.Ok)
            return

        # Assuming you want to navigate to the first record found
        self.mainclass.DATA_LIST = res
        self.mainclass.DATA_LIST_REC_TEMP = self.mainclass.DATA_LIST_REC_CORR = self.mainclass.DATA_LIST[0]

        # Fill the fields with the retrieved data
        self.mainclass.fill_fields()  # Implement this method to populate the UI fields with DATA_LIST[0]

        # Update the browse status and record counter
        self.mainclass.BROWSE_STATUS = "b"
        self.mainclass.label_status.setText(self.mainclass.STATUS_ITEMS[self.mainclass.BROWSE_STATUS])
        self.mainclass.set_rec_counter(len(self.mainclass.DATA_LIST), 1)  # Assuming you want to show the first record as current


    def image_already_associated(self, file_path, record_id):
        # Verifica se l'immagine è già associata al record
        result = self.db_search_check('MediaTable', 'filepath', file_path)
        return result is not None

    def process_ai_response(self, response_text):
        # Estrapola informazioni per ogni sezione della risposta
        us_list = self.extract_us_list(response_text)
        # Loop per creare ogni record US con il suo tipo e rapporti
        sito = us_list.get("sito")
        area = us_list.get("area")
        QMessageBox.information(self, "Information",str(us_list))
        relations = self.extract_relations_from_text(response_text,sito,area)
        QMessageBox.information(self, "Information", str(relations))
        extracted_info = {}



        for us_type, us_numbers in us_list.items():
            if us_type not in ["US", "USM"]:
                continue

            for us_number in us_numbers:
                key = f"US {us_number}"
                # Inizializza il record con le informazioni base
                extracted_info[key] = {
                    "sito": sito,
                    "area": area,
                    "us": us_number,
                    "unita_tipo": us_type,
                    "rapporti": relations.get(str(us_number), [])
                }

        return extracted_info

    def extract_us_list(self, text):
        us_list = {
            "US": [],
            "USM": [],
            "sito": None,
            "area": None
        }

        # Pattern flessibile per individuare tutte le occorrenze di US e USM
        us_pattern = re.compile(r"\bUS\s*(\d+)", re.IGNORECASE)
        usm_pattern = re.compile(r"\bUSM\s*(\d+)", re.IGNORECASE)

        # Trova tutte le US
        us_matches = us_pattern.findall(text)
        if us_matches:
            us_list["US"] = [int(us.strip()) for us in us_matches]

        # Trova tutte le USM
        usm_matches = usm_pattern.findall(text)
        if usm_matches:
            us_list["USM"] = [int(usm.strip()) for usm in usm_matches]

        # Controlla se "sito" e "area" sono presenti nel testo
        site_pattern = re.compile(r"sito:\s*([\w\s]+)", re.IGNORECASE)
        area_pattern = re.compile(r"area:\s*([\w\s]+)", re.IGNORECASE)

        site_match = site_pattern.search(text)
        area_match = area_pattern.search(text)

        if site_match:
            us_list["sito"] = site_match.group(1).strip()
        if area_match:
            us_list["area"] = area_match.group(1).strip()

        # Se "sito" e "area" non sono nel testo, chiedi input manuale
        if not us_list["sito"] or not us_list["area"]:
            missing_fields = []
            if not us_list["sito"]:
                missing_fields.append("sito")
            if not us_list["area"]:
                missing_fields.append("area")

            # Usa la funzione manual_input per richiedere le informazioni
            manual_info = self.manual_input(missing_fields)
            if manual_info:
                us_list["sito"] = manual_info.get("sito", us_list["sito"])
                us_list["area"] = manual_info.get("area", us_list["area"])

        return us_list

    def extract_relations_from_text(self,text, sito, area):
        relations = defaultdict(list)

        # Definisci i pattern per catturare i diversi tipi di relazioni stratigrafiche
        patterns = {
            "Copre / Coperto da": re.compile(r"- US (\d+) copre US (\d+(?:, US \d+)*)"),
            "Taglia / Tagliato da": re.compile(r"- US (\d+) taglia US (\d+(?:, US \d+)*)"),
            "Riempie / Riempito da": re.compile(r"- US (\d+) riempie US (\d+)"),
        }

        site_name = sito  # Nome del sito
        area_number = area  # Numero dell'area

        for rel_type, pattern in patterns.items():
            matches = pattern.findall(text)
            for match in matches:
                us_number_1 = match[0]
                us_numbers_2 = match[1].split(", ")
                for us_number_2 in us_numbers_2:
                    # Aggiungi alla US principale
                    relations[us_number_1].append([rel_type.split(' / ')[0], us_number_2, area_number, site_name])
                    # Aggiungi alla US secondaria
                    relations[us_number_2].append([rel_type.split(' / ')[1], us_number_1, area_number, site_name])

        return relations

    def extract_info_from_response_matrix(self, response):
        if not response:
            raise ValueError("Empty response received")

        info = {
            'sito': None,
            'area': None

        }

        lines = response.split('\n')
        for line in lines:
            line = line.strip()
            for key in info.keys():
                if f"{key}:" in line.lower():
                    info[key] = line.split(':', 1)[1].strip()

        if not all(info.values()):
            self.extract_missing_info_matrix(response, info)

        self.check_manual_input(info)

        return info


    def extract_info_from_response(self, response):
        if not response:
            raise ValueError("Empty response received")

        info = {
            'sito': None,
            'area': None,
            'us': None,
            'unita_tipo': None,
        }

        lines = response.split('\n')
        for line in lines:
            line = line.strip()
            for key in info.keys():
                if f"{key}:" in line.lower():
                    info[key] = line.split(':', 1)[1].strip()

        if not all(info.values()):
            self.extract_missing_info(response, info)

        self.check_manual_input(info)

        return info

    def extract_missing_info(self, response, info):
        
        response_upper = response.upper()
        if not info['sito']:
            info['sito'] = self.extract_info_generic(response, ['sito', 'site'])
        if not info['area']:
            info['area'] = self.extract_info_generic(response, ['area'])
        if not info['us']:
            info['us'] = self.extract_info_generic(response, ['US', 'us', 'unità stratigrafica', 'stratigraphic unit'])
        if not info['unita_tipo']:
            info['unita_tipo'] = self.extract_info_generic(response_upper, ['unità tipo', 'unita tipo', 'unit type'])

    def extract_missing_info_matrix(self, response, info):

        response_upper = response.upper()
        if not info['sito']:
            info['sito'] = self.extract_info_generic(response, ['sito', 'site'])
        if not info['area']:
            info['area'] = self.extract_info_generic(response, ['area'])

    def check_manual_input(self, info):
        missing_fields = [k for k, v in info.items() if v is None or v == '']
        if missing_fields:
            QMessageBox.information(
                self,
                "Informazioni Mancanti",
                f"Le seguenti informazioni sono mancanti: {', '.join(missing_fields)}. Procederemo con l'inserimento manuale."
            )
            manual_info = self.manual_input(missing_fields)
            info.update(manual_info)

    def extract_info_generic(self, text, keywords):
        for keyword in keywords:
            pattern = re.compile(f"{keyword}:?\s*(.*)", re.IGNORECASE)
            match = pattern.search(text)
            if match:
                return match.group(1).strip()
        return None



    def confirm_information(self, info, file_path):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setText("Are the following information correct?")
        msg.setWindowTitle("Confirm Information")

        details = f"Sito: {info['sito']}\n"
        details += f"Area: {info['area']}\n"
        details += f"US: {info['us']}\n"
        details += f"Unità tipo: {info['unita_tipo']}\n"
        details += f"\nImage: {file_path}"

        msg.setInformativeText(details)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        reply = msg.exec()

        if reply == QMessageBox.Yes:
            return True
        else:
            return False


    def confirm_information_matrix(self, info, file_path):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        msg.setText("Are the following information correct?")
        msg.setWindowTitle("Confirm Information")

        details = f"Sito: {info['sito']}\n"
        details += f"Area: {info['area']}\n"
        details += f"US: {info['us']}\n"
        details += f"Rapporti: {info['rapporti']}\n"
        details += f"Unità tipo: {info['unita_tipo']}\n"
        details += f"\nImage: {file_path}"

        msg.setInformativeText(details)
        msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        reply = msg.exec()

        if reply == QMessageBox.Yes:
            return True
        else:
            return False
    def create_new_record_matrix(self, info):
        try:
            sito = info['sito']
            area = info['area']
            us = info['us']
            relation = info['rapporti']
            unita_tipo = info['unita_tipo']

            # Chiamata alla funzione per inserire il nuovo record
            self.DB_MANAGER.insert_number_of_rapporti_records(sito, area, us, relation, unita_tipo)

            # Recupera l'ID del record appena inserito
            id_us = self.DB_MANAGER.max_num_id('US', 'id_us')

            return id_us
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create new record: {str(e)}")
            return None

    def create_new_record(self, info):
        try:
            sito = info['sito']
            area = info['area']
            us = info['us']
            unita_tipo = info['unita_tipo']

            # Chiamata alla funzione per inserire il nuovo record
            self.DB_MANAGER.insert_number_of_us_records(sito, area, us, unita_tipo)

            # Recupera l'ID del record appena inserito
            id_us = self.DB_MANAGER.max_num_id('US', 'id_us')

            return id_us
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create new record: {str(e)}")
            return None

    def db_search_check(self, table_class, field, value):
        self.table_class = table_class
        self.field = field
        self.value = value
        search_dict = {self.field: "'" + str(self.value) + "'"}
        u = Utility()
        search_dict = u.remove_empty_items_fr_dict(search_dict)
        res = self.DB_MANAGER.query_bool(search_dict, self.table_class)
        return res


    def insert_record_media(self, mediatype, filename, filetype, filepath):
        self.mediatype = mediatype
        self.filename = filename
        self.filetype = filetype
        self.filepath = filepath
        try:
            data = self.DB_MANAGER.insert_media_values(
                self.DB_MANAGER.max_num_id('MEDIA', 'id_media') + 1,
                str(self.mediatype),  # 1 - mediatyype
                str(self.filename),  # 2 - filename
                str(self.filetype),  # 3 - filetype
                str(self.filepath),  # 4 - filepath
                str('Insert description'),  # 5 - descrizione
                str("['imagine']"))  # 6 - tags
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as  e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.filename + ": Image already in the database"
                else:
                    msg = e
                #QMessageBox.warning(self, "Errore", "Warning 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0

    def insert_record_mediathumb(self, media_max_num_id, mediatype, filename, filename_thumb, filetype, filepath_thumb, filepath_resize):
        self.media_max_num_id = media_max_num_id
        self.mediatype = mediatype
        self.filename = filename
        self.filename_thumb = filename_thumb
        self.filetype = filetype
        self.filepath_thumb = filepath_thumb
        self.filepath_resize = filepath_resize
        try:
            data = self.DB_MANAGER.insert_mediathumb_values(
                self.DB_MANAGER.max_num_id('MEDIA_THUMB', 'id_media_thumb') + 1,
                str(self.media_max_num_id),  # 1 - media_max_num_id
                str(self.mediatype),  # 2 - mediatype
                str(self.filename),  # 3 - filename
                str(self.filename_thumb),  # 4 - filename_thumb
                str(self.filetype),  # 5 - filetype
                str(self.filepath_thumb),  # 6 - filepath_thumb
                str(self.filepath_resize))  # 6 - filepath_thumb
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.filename + ": thumb already present into the database"
                else:
                    msg = e
                #QMessageBox.warning(self, "Error", "warming 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0

    def associate_image_with_record(self, file_path, record_id):
        if record_id is None:
            QMessageBox.information(self, 'ok', str(record_id))
            return

        try:
            # Determine file type and media type
            filename = os.path.basename(file_path)
            filename, filetype = filename.split(".")[0], filename.split(".")[1]

            accepted_image_formats = ["jpg", "jpeg", "png", "tiff", "tif", "bmp"]
            if filetype.lower() in accepted_image_formats:
                mediatype = 'image'
                media_thumb_suffix = '_thumb.png'
                media_resize_suffix = '.png'
            else:
                raise ValueError(f"Unsupported file type: {filetype}")

            # Check if the image already exists in the database
            existing_media = self.db_search_check('MEDIA', 'filepath', file_path)

            if existing_media:
                # Se l'immagine esiste già, mostra un messaggio di avviso
                QMessageBox.warning(self, "Warning", f"L'immagine '{filename}' è già presente nel database.")
                return  # Esci dalla funzione se l'immagine è già presente
            else:
                # Inserisci un nuovo record in MEDIA table
                media_id = self.insert_record_media(mediatype, filename, filetype, file_path)
                if media_id == 0:
                    QMessageBox.warning(self, "Error", "Failed to insert media record.")
                    return

                # Process thumbnails
                conn = Connection()
                thumb_path = conn.thumb_path()['thumb_path']
                thumb_resize = conn.thumb_resize()['thumb_resize']

                filename_thumb = f"{media_id}_{filename}{media_thumb_suffix}"
                filename_resize = f"{media_id}_{filename}{media_resize_suffix}"
                filepath_thumb = filename_thumb
                filepath_resize = filename_resize

                # Create thumbnails and resized images
                MU = Media_utility()
                MUR = Media_utility_resize()
                MU.resample_images(media_id, file_path, filename, thumb_path, media_thumb_suffix)
                MUR.resample_images(media_id, file_path, filename, thumb_resize, media_resize_suffix)

                # Insert thumbnail record
                media_max_num_id = self.DB_MANAGER.max_num_id('MEDIA', 'id_media')
                self.insert_record_mediathumb(media_max_num_id, mediatype, filename, filename_thumb, filetype,
                                              filepath_thumb, filepath_resize)
                # Update iconListWidget
                item = QListWidgetItem(filename)
                item.setData(Qt.UserRole, str(media_id))
                icon = QIcon(os.path.join(thumb_path, filename_thumb))
                item.setIcon(icon)
                self.mainclass.iconListWidget.addItem(item)

                # Assign tags (if needed)
                self.assignTags_US(item)

                QMessageBox.information(self, "Success", f"Image successfully associated with record {record_id}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to associate image with record: {str(e)}")

    def insert_mediaToEntity_rec(self, id_entity, entity_type, table_name, id_media, filepath, media_name):
        """
        id_mediaToEntity,
        id_entity,
        entity_type,
        table_name,
        id_media,
        filepath,
        media_name"""
        self.id_entity = id_entity
        self.entity_type = entity_type
        self.table_name = table_name
        self.id_media = id_media
        self.filepath = filepath
        self.media_name = media_name
        try:
            data = self.DB_MANAGER.insert_media2entity_values(
                self.DB_MANAGER.max_num_id('MEDIATOENTITY', 'id_mediaToEntity') + 1,
                int(self.id_entity),  # 1 - id_entity
                str(self.entity_type),  # 2 - entity_type
                str(self.table_name),  # 3 - table_name
                int(self.id_media),  # 4 - us
                str(self.filepath),  # 5 - filepath
                str(self.media_name))  # 6 - media_name
            try:
                self.DB_MANAGER.insert_data_session(data)
                return 1
            except Exception as  e:
                e_str = str(e)
                if e_str.__contains__("Integrity"):
                    msg = self.ID_TABLE + " already present into the database"
                else:
                    msg = e
                QMessageBox.warning(self, "Error", "Warning 1 ! \n"+ str(msg),  QMessageBox.Ok)
                return 0
        except Exception as  e:
            QMessageBox.warning(self, "Error", "Warning 2 ! \n"+str(e),  QMessageBox.Ok)
            return 0

    def assignTags_US(self, item):
        """
        id_mediaToEntity,
        id_entity,
        entity_type,
        table_name,
        id_media,
        filepath,
        media_name
        """
        us_list = self.generate_US()
        #QMessageBox.information(self,'search db',str(us_list))
        if not us_list:
            return

        for us_data in us_list:
            id_orig_item = item.text()  # return the name of original file
            search_dict = {'filename': "'" + str(id_orig_item) + "'"}
            media_data = self.DB_MANAGER.query_bool(search_dict, 'MEDIA')
            self.insert_mediaToEntity_rec(us_data[0], us_data[1], us_data[2], media_data[0].id_media,
                                          media_data[0].filepath, media_data[0].filename)

    def generate_US(self):
        #tags_list = self.table2dict('self.tableWidgetTags_US')
        record_us_list = []
        sito=self.mainclass.comboBox_sito.currentText()
        area=self.mainclass.comboBox_area.currentText()
        us=self.mainclass.lineEdit_us.text()
        #for sing_tags in tags_list:
        search_dict = {'sito': "'"+str(sito)+"'" ,
                       'area': "'"+str(area)+"'",
                       'us': "'"+str(us)+"'"
                       }
        j = self.DB_MANAGER.query_bool(search_dict, 'US')
        record_us_list.append(j)
        #QMessageBox.information(self, 'search db', str(record_us_list))
        us_list = []
        for r in record_us_list:
            us_list.append([r[0].id_us, 'US', 'us_table'])
        #QMessageBox.information(self, "Scheda US", str(us_list), QMessageBox.Ok)
        return us_list

    def update_icon_list_widget(self, file_path, record_id):
        # This function should update the iconListWidget in US_USM
        try:
            item = QListWidgetItem(QIcon(file_path), f"Record {record_id}")
            self.mainclass.iconListWidget.addItem(item)
        except Exception as e:
            QMessageBox.warning(self, "Warning", f"Failed to update icon list: {str(e)}")

    def ask_sketch(self, prompt, apikey, file_path):
        def encode_file(file_path):
            try:
                with open(file_path, "rb") as file:
                    return base64.b64encode(file.read()).decode('utf-8')
            except FileNotFoundError:
                print(f"No file found at {file_path}. Please check the file path.")
                return None

        def get_file_type(file_path):
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                return mime_type.split('/')[0]
            return None

        def extract_video_frames(file_path, num_frames=5):
            cap = cv2.VideoCapture(file_path)
            if not cap.isOpened():
                return "Errore nell'apertura del video", []

            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0

            frames = []
            for i in range(num_frames):
                cap.set(cv2.CAP_PROP_POS_FRAMES, i * frame_count // num_frames)
                ret, frame = cap.read()
                if ret:
                    _, buffer = cv2.imencode('.jpg', frame)
                    base64_frame = base64.b64encode(buffer).decode('utf-8')
                    frames.append(base64_frame)

            cap.release()
            return f"Video duration: {duration:.2f} seconds, FPS: {fps:.2f}, Frames: {frame_count}", frames

        OpenAI.api_key = apikey
        file_type = get_file_type(file_path)

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OpenAI.api_key}"
        }

        content = [{"type": "text", "text": prompt}]

        if file_type == 'image':
            encoded_file = encode_file(file_path)
            if not encoded_file:
                return "Errore: File non trovato o non accessibile."
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{encoded_file}"
                }
            })
        elif file_type == 'video':
            video_info, frames = extract_video_frames(file_path)
            content.append({"type": "text", "text": video_info})
            for frame in frames:
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{frame}"
                    }
                })
            content.append({"type": "text",
                            "text": "Analizza queste immagini estratte dal video e fornisci un riassunto del contenuto visivo. Se riesci a identificare del testo o dei dialoghi nelle immagini, includili nel riassunto."})
        else:
            content.append({
                "type": "text",
                "text": f"Il file è di tipo {file_type if file_type else 'sconosciuto'}. Nome del file: {os.path.basename(file_path)}"
            })

        params = {
            "model": "gpt-4-vision-preview",
            "temperature": 0.5,
            "user": "my_customer",
            "max_tokens": 4096,
            "top_p": 0.5,
            "stream": True,
            "messages": [
                {
                    "role": "system",
                    "content": "Sei un assistente esperto nell'analisi di immagini e video. Per i video, analizzi una serie di fotogrammi chiave per fornire un riassunto del contenuto visivo, identificando anche eventuali testi o dialoghi visibili."
                },
                {
                    "role": "user",
                    "content": content
                }
            ],
        }

        self.start_worker(headers, params, is_image=(file_type == 'image'))

    def extract_text_from_file(self, file_path):
        file_type = os.path.splitext(file_path)[1].lower()
        if file_type == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == ".csv":
            return self.extract_text_from_csv(file_path)
        elif file_type == ".docx":
            return self.extract_text_from_docx(file_path)
        else:
            print("Unsupported file type.")
            return None

    def extract_text_from_pdf(self, file_path):
        text = ""
        try:
            document = fitz.open(file_path)
            for page_num in range(len(document)):
                page = document.load_page(page_num)
                text += page.get_text()
            return text
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            return None

    def extract_text_from_csv(self, file_path):
        text = ""
        with open(file_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                text += ' '.join(row) + "\n"
        return text

    def extract_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def save_corrected_file(self, original_file_path, original_lines, corrected_text):
        file_type = os.path.splitext(original_file_path)[1].lower()
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Corrected Document", "", f"Documents (*{file_type})")

        if not save_path:
            return

        if file_type == ".pdf":
            self.save_corrected_pdf(original_file_path, save_path, original_lines, corrected_text)
        elif file_type == ".csv":
            self.save_corrected_csv(save_path, corrected_text)
        elif file_type == ".docx":
            self.save_corrected_docx(original_file_path, save_path, corrected_text)
        else:
            print("Unsupported file type.")

    def find_closest_match(self, corrected_line, original_lines):
        matcher = difflib.SequenceMatcher(None, corrected_line, None)
        best_ratio, best_match = 0.0, None

        for original_line in original_lines:
            matcher.set_seq2(original_line)
            ratio = matcher.ratio()
            if ratio > best_ratio:
                best_ratio, best_match = ratio, original_line

        return best_match if best_ratio > 0.6 else None

    def find_closest_match_pdf(self, corrected_line, original_lines):
        # Implementa la tua logica di connessione più vicina, se necessario
        # Per esempio, confrontare la similarità del testo o un semplice match
        for original_line in original_lines:
            if corrected_line.lower() in original_line.lower():
                return original_line
        return None

    def save_new_pdf_with_corrections(self, original_file_path, save_path, corrected_lines):
        try:
            doc = fitz.open(original_file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text("text")
                original_lines = page_text.splitlines()

                for line_num, corrected_line in enumerate(corrected_lines):
                    original_line = self.find_closest_match(corrected_line, original_lines)
                    if original_line is None:
                        print(f"No match found for corrected line '{corrected_line}'")
                        continue

                    text_instances = page.search_for(original_line)
                    if not text_instances:
                        print(f"No instances of original text '{original_line}' found on page {page_num}")
                        continue

                    for inst in text_instances:
                        insert_point = fitz.Point(inst.x0, inst.y0)  # Usa le coordinate della vecchia istanza
                        print(f"Inserting: '{corrected_line}' at {insert_point}")
                        page.insert_text(insert_point, corrected_line, fontsize=12, color=(0, 0, 0), overlay=True)
                        break  # Assicurati di non sovrascrivere più volte

            doc.save(save_path)
            QMessageBox.information(self, "Success", "Corrected PDF saved successfully.")
        except Exception as e:
            print(f"Error saving PDF file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving PDF file: {e}")



    def save_corrected_csv(self, save_path, corrected_text):
        try:
            with open(save_path, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                for line in corrected_text.split('\n'):
                    writer.writerow(line.split())
            QMessageBox.information(self, "Success", "Corrected CSV saved successfully.")
        except Exception as e:
            print(f"Error saving CSV file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving CSV file: {e}")

    def save_corrected_docx(self, original_file_path, save_path, corrected_text):
        try:
            doc = docx.Document(original_file_path)
            corrected_paragraphs = corrected_text.split('\n')
            for i, paragraph in enumerate(doc.paragraphs):
                if i < len(corrected_paragraphs):
                    paragraph.text = corrected_paragraphs[i]
            doc.save(save_path)
            QMessageBox.information(self, "Success", "Corrected DOCX saved successfully.")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving DOCX file: {e}")

    def ask_doc(self, prompt, apikey, file_path):
        OpenAI.api_key = apikey
        file_text = self.extract_text_from_file(file_path)
        if file_text is None:
            return None

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OpenAI.api_key}"
        }

        params = {
            "model": "gpt-4o",
            "temperature": 0.5,
            "user": "my_customer",
            "max_tokens": 4096,
            "top_p": 0.5,
            "stream": True,
            "messages": [
                {
                    "role": "system",
                    "content": "Sono il tuo assistente personale per la lettura di documenti ed esporterò il documento nel formato di input tramite link"
                },
                {
                    "role": "user",
                    "content": prompt + "\n\n" + file_text
                }
            ],
        }


        self.start_worker(headers, params, url="https://api.openai.com/v1/chat/completions", is_image=False)


    def ask_doc_with_claude(self, prompt, apikey, file_path):
        Anthropic.api_key=apikey

        file_text = self.extract_text_from_file(file_path)
        if file_text is None:
            return None
        messages = [
            {
                "role": "user",
                "content": prompt + "\n\n" + file_text
            }
        ]

        # Prepara i parametri per il worker
        params = {
            "model": "claude-3-5-sonnet-20240620",
            "messages": messages,
            "max_tokens": 4096,
            "temperature": 0.5,
            "stream": True
        }

        headers = {
            "x-api-key": apikey,
            "anthropic-version":"2023-06-01",  # Include the version header
            "content-type" : "application/json"

        }


        # Avvia il worker
        self.start_worker(headers, params, url = "https://api.anthropic.com/v1/message",is_image=False)




    def update_progress(self, progress):
        self.progress.setValue(progress)

    def update_content(self, content):
        # Append new content to the existing text in the list widget
        current_text = self.listWidget_ai.toPlainText()
        self.listWidget_ai.setPlainText(current_text + content)

    def update_tokens_used(self, tokens_used, total_cost):
        self.token_counter.setText(f"Tokens used: {tokens_used} - Total cost: ${total_cost:.4f}")

    def docchgpt(self):
        self.listWidget_ai.clear()
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Select Document", "", "Documents (*.pdf *.csv *.docx)")
        if file_path:
            prompt = self.prompt_label.toPlainText()
            selected_model = self.model_selector.currentText()
            if selected_model == "GPT-4o":
                self.ask_doc(prompt, self.apikey_gpt(), file_path)
            else:
                self.ask_claude(prompt, self.apikey_claude(), file_path, is_image=False)
        else:
            self.listWidget_ai.setPlainText("Document selection was canceled.")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    mainWin = GPTWindow()
    mainWin.show()
    sys.exit(app.exec())
