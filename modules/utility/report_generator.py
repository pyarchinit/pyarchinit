from docx import Document
from sqlalchemy import create_engine, MetaData, Table, select
from sqlalchemy.orm import sessionmaker
from qgis.PyQt.QtWidgets import *
import socket
# OpenAI import removed to avoid pydantic conflicts - will be imported lazily in generate_report_with_openai
from modules.utility.report_text_cleaner import ReportTextCleaner
from modules.utility.llm_providers import (
    LLMConfig,
    LLMProvider,
    LLMProviderManager,
)


class ReportGenerator(QWidget):

    MAX_TOKENS = 4096

    def __init__(self):
        super().__init__()


    @staticmethod
    def read_data_from_db(db_url, table_name, sito=None):
        """
        Read all records from ``table_name``. If ``sito`` is supplied
        and the table has a ``sito`` column, only the records for that
        site are returned — this prevents loading 510 demo US when the
        user only wants the 51 for their configured site.
        """
        engine = create_engine(db_url)
        metadata = MetaData()
        table = Table(table_name, metadata, autoload_with=engine)
        Session = sessionmaker(bind=engine)
        session = Session()
        query = session.query(table)
        if sito and 'sito' in table.c:
            query = query.filter(table.c.sito == sito)
        records = query.all()
        columns = [column.name for column in table.columns]
        session.close()
        return records, columns

    @staticmethod
    def chunk_data(data, chunk_size):
        for i in range(0, len(data), chunk_size):
            yield data[i:i + chunk_size]

    def generate_report(self, prompt_completo, config: LLMConfig):
        """
        Generate a report using any supported LLM provider (OpenAI, Anthropic,
        Ollama, LM Studio).

        ``config`` is built by ``LLMSelectorWidget`` and tells us which
        provider/model/key to use. Streaming is uniform across providers
        thanks to ``LLMProviderManager.stream_chat``.
        """
        if not config.model:
            QMessageBox.warning(
                None,
                "Modello mancante",
                "Nessun modello selezionato. Apri il selettore LLM e scegline uno.",
                QMessageBox.Ok,
            )
            return "Error: no model configured"

        if config.needs_api_key and not config.api_key:
            QMessageBox.warning(
                None,
                "API key mancante",
                f"Manca l'API key per {config.provider.value}.",
                QMessageBox.Ok,
            )
            return "Error: missing API key"

        messaggio_combinato = "\n "
        try:
            for chunk in LLMProviderManager.stream_chat(
                config,
                messages=[{"role": "user", "content": prompt_completo}],
            ):
                messaggio_combinato += chunk
        except ImportError as e:
            QMessageBox.warning(
                None,
                "LLM Import Error",
                f"Libreria mancante per il provider scelto: {e}",
                QMessageBox.Ok,
            )
            return "Error: Could not load LLM library"
        except Exception as e:
            print(f"Errore nel processo di stream: {e}")
            QMessageBox.warning(
                None,
                "Errore LLM",
                f"Errore durante la generazione del report:\n{e}",
                QMessageBox.Ok,
            )

        return ReportTextCleaner.clean_report_text(messaggio_combinato)

    def generate_report_with_openai(self, prompt_completo, modello_selezionato, apikey):
        """
        Backward-compatible wrapper. Existing callers still work, but new
        code should call ``generate_report(prompt, LLMConfig(...))``.
        """
        config = LLMConfig(
            provider=LLMProvider.OPENAI,
            model=modello_selezionato,
            api_key=apikey,
        )
        return self.generate_report(prompt_completo, config)

    @staticmethod
    def is_connected():
        try:
            # Try to connect to one of the DNS servers
            sock = socket.create_connection(("1.1.1.1", 53), timeout=2)
            sock.close()  # Close the socket properly to avoid ResourceWarning
            return True
        except OSError:
            pass
        return False

    @staticmethod
    def save_report_to_file_old(report, file_path):
        # Create a new Document
        doc = Document()
        # Add the report text to the document
        doc.add_paragraph(report)
        # Save the document to the specified file path
        doc.save(file_path)

    @staticmethod
    def save_report_to_file(report, file_path):
        """
        Save report with proper formatting and cleaning
        """
        from docx.shared import Pt

        # Clean the report text first
        cleaned_report = ReportTextCleaner.clean_report_text(report)

        # Prepare for docx
        prepared = ReportTextCleaner.prepare_for_docx(cleaned_report)

        # Create document
        doc = Document()

        # Add paragraphs with proper formatting
        for para_info in prepared['paragraphs']:
            text = para_info['text']
            style = para_info['style']

            if style.startswith('heading'):
                level = int(style[-1]) if style[-1].isdigit() else 1
                para = doc.add_heading(text, level=level)
            elif style == 'list':
                para = doc.add_paragraph(text, style='List Bullet')
            else:
                para = doc.add_paragraph(text)

            # Apply Cambria font
            for run in para.runs:
                run.font.name = 'Cambria'
                run.font.size = Pt(12)

        # Save the document
        doc.save(file_path)

