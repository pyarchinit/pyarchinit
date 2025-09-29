#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Document Processor
Utilizza DocumentStyleAgent per processare e correggere documenti Word
"""

from docx import Document
from docx.shared import Pt
from typing import Optional
from .document_style_agent import DocumentStyleAgent
from .report_text_cleaner import ReportTextCleaner


class DocumentProcessor:
    """
    Processa documenti Word applicando cleaning e correzione stili
    """

    def __init__(self):
        self.agent = DocumentStyleAgent()
        self.cleaner = ReportTextCleaner()

    def process_document(self, doc_path: str, output_path: Optional[str] = None) -> dict:
        """
        Processa un documento Word correggendo stili e formattazione

        Args:
            doc_path: Path al documento da processare
            output_path: Path di output (opzionale, sovrascrive l'originale se None)

        Returns:
            Dizionario con statistiche del processing
        """
        if output_path is None:
            output_path = doc_path

        # Carica il documento
        doc = Document(doc_path)

        # Estrai e pulisci i paragrafi
        paragraphs_data = []
        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else "Normal"

            # Pulisci il testo
            if text:
                text = ReportTextCleaner.clean_report_text(text)

            paragraphs_data.append((text, style))

        # Correggi gli stili con l'agente
        corrections = self.agent.correct_document_styles(paragraphs_data)

        # Crea nuovo documento con correzioni
        doc_out = Document()
        stats = {'total': 0, 'corrected': 0, 'cleaned': 0}

        for text, original_style, correct_style in corrections:
            if not text:
                doc_out.add_paragraph()
                continue

            stats['total'] += 1
            if original_style != correct_style:
                stats['corrected'] += 1

            # Applica lo stile corretto
            para = self._add_paragraph_with_style(doc_out, text, correct_style)

            # Applica font Cambria
            self._apply_font_formatting(para)

        # Salva
        doc_out.save(output_path)

        return stats

    def _add_paragraph_with_style(self, doc, text, style):
        """
        Aggiunge un paragrafo con lo stile specificato
        """
        if style == "Title":
            return doc.add_heading(text, level=0)
        elif style == "Heading 1":
            return doc.add_heading(text, level=1)
        elif style == "Heading 2":
            return doc.add_heading(text, level=2)
        elif style == "Heading 3":
            return doc.add_heading(text, level=3)
        elif style == "List Bullet":
            # Rimuovi marker se presente
            if text.startswith(('- ', '• ')):
                text = text[2:].strip()
            return doc.add_paragraph(text, style='List Bullet')
        else:
            return doc.add_paragraph(text)

    def _apply_font_formatting(self, para):
        """
        Applica formattazione Cambria al paragrafo
        """
        for run in para.runs:
            run.font.name = 'Cambria'

            if para.style and para.style.name:
                if para.style.name in ["Title", "Heading 1"]:
                    run.font.size = Pt(16)
                elif para.style.name == "Heading 2":
                    run.font.size = Pt(14)
                elif para.style.name == "Heading 3":
                    run.font.size = Pt(13)
                else:
                    run.font.size = Pt(12)

    def process_content_to_paragraphs(self, content: str, doc: Document):
        """
        Processa contenuto testuale e lo aggiunge al documento con stili corretti

        Args:
            content: Testo da processare
            doc: Documento Word a cui aggiungere il contenuto
        """
        # Prima pulisci il contenuto
        cleaned_content = ReportTextCleaner.clean_report_text(content)

        # Splitta in paragrafi
        paragraphs = cleaned_content.split('\n')

        # Analizza con l'agente per determinare gli stili
        analysis = self.agent.analyze_document(paragraphs)

        # Aggiungi al documento con stili corretti
        for para_text, style_info in zip(paragraphs, analysis):
            if not para_text.strip():
                doc.add_paragraph()
                continue

            # Mappa lo stile suggerito
            style_map = {
                'title': lambda: doc.add_heading(para_text, level=0),
                'heading1': lambda: doc.add_heading(para_text, level=1),
                'heading2': lambda: doc.add_heading(para_text, level=2),
                'heading3': lambda: doc.add_heading(para_text, level=3),
                'list': lambda: self._add_list_item(doc, para_text),
                'normal': lambda: doc.add_paragraph(para_text)
            }

            # Applica lo stile
            suggested_style = style_info['style']
            if suggested_style in style_map:
                para = style_map[suggested_style]()
            else:
                para = doc.add_paragraph(para_text)

            # Applica font
            self._apply_font_formatting(para)

    def _add_list_item(self, doc, text):
        """
        Aggiunge un elemento di lista
        """
        # Rimuovi marker se presente
        if text.startswith(('- ', '• ')):
            text = text[2:].strip()
        return doc.add_paragraph(text, style='List Bullet')

    @staticmethod
    def quick_fix_document(doc_path: str):
        """
        Metodo statico per correzione rapida di un documento

        Args:
            doc_path: Path al documento da correggere

        Returns:
            Statistiche della correzione
        """
        processor = DocumentProcessor()
        return processor.process_document(doc_path)