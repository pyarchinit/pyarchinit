import base64
import difflib
import json
import os
import sys
import time
import csv
from qgis.PyQt.QtCore import Qt

import docx
import fitz
import openai
import requests

from qgis.PyQt.QtWidgets import QProgressBar, QScrollArea, QInputDialog, QMessageBox, QFileDialog, QLabel, QMainWindow, \
    QApplication, QGridLayout, QWidget, QTextEdit, QPushButton, QListWidget, QSplitter
from qgis.PyQt.QtCore import QThread, pyqtSignal

from qgis.PyQt.QtGui import QIcon


class Worker(QThread):
    progress_updated = pyqtSignal(int)  # signal for progress updates
    content_updated = pyqtSignal(str)  # signal for content updates
    tokens_used_updated = pyqtSignal(int, float)  # signal for tokens used updates with cost

    def __init__(self, headers, params, is_image=False, image_width=512, image_height=512):
        super().__init__()
        self.headers = headers
        self.params = params
        self.is_image = is_image
        self.image_width = image_width
        self.image_height = image_height

    def run(self):
        try:
            client = requests.Session()
            response = client.post("https://api.openai.com/v1/chat/completions", headers=self.headers, json=self.params, stream=True)

            if response.status_code != 200:
                print(f"Error: HTTP {response.status_code} - {response.text}")
                return

            total_received_content = 0
            total_content_length = 4096  # Assuming max_tokens is the total content length
            reply = ""
            tokens_used = 0
            input_cost_per_token = 5.00 / 1_000_000  # Cost per input token in USD
            output_cost_per_token = 15.00 / 1_000_000  # Cost per output token in USD

            # Calculate image cost based on dimensions
            if self.is_image:
                tiles_x = (self.image_width + 511) // 512
                tiles_y = (self.image_height + 511) // 512
                total_tiles = tiles_x * tiles_y
                base_tokens = 85
                tile_tokens = 170 * total_tiles
                total_image_tokens = base_tokens + tile_tokens
                image_cost_per_tile = total_image_tokens * input_cost_per_token
            else:
                image_cost_per_tile = 0

            for line in response.iter_lines():
                if line:
                    decoded_line = json.loads(line.decode('utf-8').strip()[len("data: "):])
                    if 'choices' in decoded_line:
                        for choice in decoded_line['choices']:
                            if 'delta' in choice and 'content' in choice['delta']:
                                content_chunk = choice['delta']['content']
                                if content_chunk:
                                    reply += content_chunk
                                    total_received_content += len(content_chunk)
                                    tokens_used += len(content_chunk.split())
                                    progress_percentage = int(total_received_content / 100)
                                    # Calculate progress percentage
                                    self.progress_updated.emit(progress_percentage)

                                    total_cost = tokens_used * (
                                    input_cost_per_token + output_cost_per_token) + image_cost_per_tile
                                    self.tokens_used_updated.emit(tokens_used, total_cost)
                                    self.content_updated.emit(content_chunk)

            self.progress_updated.emit(100)  # Ensure progress bar reaches 100%
        except Exception as e:
            print(f"Error in worker thread: {e}")


class GPTWindow(QMainWindow):
    HOME = os.environ.get('PYARCHINIT_HOME', '')

    def __init__(self, iconListWidget=None, db_manager=None):
        super().__init__()
        self.setWindowTitle("PyArchInit - GPT Sketch")
        self.setGeometry(100, 100, 600, 800)
        self.set_icon(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../resources/icons/gpt.png")))
        self.iconListWidget = iconListWidget
        self.DB_MANAGER = db_manager
        # Layout principale
        layout = QGridLayout()

        # creazione prompt
        self.prompt_label = QTextEdit()
        self.prompt_label.setFixedHeight(150)
        self.prompt_label.setPlaceholderText(
            "Esempio di prompt:\n"
            "Se vuoi estrarre informazioni da un'immagine, scrivi:\n"
            "Estrai le informazioni dalla lavagnetta.\n e schiacci Importa immagine\n\n"
            "Se vuoi correggere un testo, scrivi:\n"
            "Fai una correzione del testo.\n"
            "e schiacci Importa Documento"
        )
        layout.addWidget(self.prompt_label, 1, 0, 1, 2)

        # Bottoni per l'interazione
        self.btn_import = QPushButton("Importa Immagine")
        self.btn_import.clicked.connect(self.scketchgpt)
        self.btn_import2 = QPushButton("Importa Documento")
        self.btn_import2.clicked.connect(self.docchgpt)
        # Aggiungi il pulsante per analizzare le immagini selezionate
        self.btn_analyze_selected = QPushButton("Analizza Immagini Selezionate")
        self.btn_analyze_selected.clicked.connect(self.analyze_selected_images)

        self.progress = QProgressBar()
        self.token_counter = QLabel("Tokens used: 0 - Total cost: $0.0000")
        layout.addWidget(self.btn_import, 2, 0)
        layout.addWidget(self.btn_import2, 2, 1)
        layout.addWidget(self.progress, 3, 0, 1, 0)
        layout.addWidget(self.token_counter, 3, 0, 1, 1)
        layout.addWidget(self.btn_analyze_selected, 7, 0, 1, 2)

        # Lista per visualizzare le risposte di GPT
        self.listWidget_ai = QTextEdit()

        # Add scrollbar to listWidget_ai
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(self.listWidget_ai)

        # Use QSplitter to make listWidget_ai resizable
        splitter = QSplitter()
        splitter.addWidget(scroll_area)
        splitter.setStretchFactor(0, 1)  # Make listWidget_ai expandable

        layout.addWidget(splitter, 4, 0, 1, 2)

        # Widget container per i layout
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def set_icon(self, icon_path):
        self.setWindowIcon(QIcon(icon_path))



    def start_worker(self, headers, params, is_image=False):
        self.worker = Worker(headers, params)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.content_updated.connect(self.update_content)
        self.worker.tokens_used_updated.connect(self.update_tokens_used)
        self.worker.start()
    def apikey_gpt(self):
        # HOME = os.environ['PYARCHINIT_HOME']
        BIN = '{}{}{}'.format(self.HOME, os.sep, "bin")
        api_key = ""
        # Verifica se il file gpt_api_key.txt esiste
        path_key = os.path.join(BIN, 'gpt_api_key.txt')
        if os.path.exists(path_key):

            # Leggi l'API Key dal file
            with open(path_key, 'r') as f:
                api_key = f.read().strip()
                try:

                    return api_key

                except:
                    reply = QMessageBox.question(None, 'Warning', 'Apikey non valida' + '\n'
                                                 + 'Clicca ok per inserire la chiave',
                                                 QMessageBox.Ok | QMessageBox.Cancel)
                    if reply == QMessageBox.Ok:

                        api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey valida:')
                        if ok:
                            # Salva la nuova API Key nel file
                            with open(path_key, 'w') as f:
                                f.write(api_key)
                                f.close()
                            with open(path_key, 'r') as f:
                                api_key = f.read().strip()
                    else:
                        return api_key


        else:
            # Chiedi all'utente di inserire una nuova API Key
            api_key, ok = QInputDialog.getText(None, 'Apikey gpt', 'Inserisci apikey:')
            if ok:
                # Salva la nuova API Key nel file
                with open(path_key, 'w') as f:
                    f.write(api_key)
                    f.close()
                with open(path_key, 'r') as f:
                    api_key = f.read().strip()

        return api_key

    def analyze_selected_images(self):
        if not self.iconListWidget or not self.DB_MANAGER:
            QMessageBox.warning(self, "Attenzione", "Dati necessari non disponibili.", QMessageBox.Ok)
            return

        selected_items = self.iconListWidget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Attenzione", "Nessuna immagine selezionata.", QMessageBox.Ok)
            return

        self.listWidget_ai.clear()
        prompt = self.prompt_label.toPlainText()

        for item in selected_items:
            media_id = item.data(Qt.ItemDataRole.UserRole)
            filepath = self.DB_MANAGER.select_mediapath_from_id(media_id)
            if filepath and os.path.exists(filepath):
                gpt_response = self.ask_sketch(prompt, self.apikey_gpt(), filepath)
                self.listWidget_ai.append(f"Analisi per l'immagine {item.text()}:\n{gpt_response}\n")
            else:
                self.listWidget_ai.append(f"Errore: Impossibile trovare o accedere all'immagine {item.text()}\n")

            QApplication.processEvents()

        # Scorri automaticamente verso il basso per mostrare l'ultimo contenuto aggiunto
        self.listWidget_ai.verticalScrollBar().setValue(self.listWidget_ai.verticalScrollBar().maximum())

    def scketchgpt(self):
        self.listWidget_ai.clear()
        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(self, "Select Images", "", "Images (*.png *.jpg *.jpeg)")
        if file_paths:
            prompt = self.prompt_label.toPlainText()
            for file_path in file_paths:
                gpt_response = self.ask_sketch(prompt, self.apikey_gpt(), file_path)
                #combined_message = f"Image: {file_path}\nGPT Response:\n\n{gpt_response}"
                #self.listWidget_ai.addItem(combined_message)

                QApplication.processEvents()
        else:
            self.listWidget_ai.addItem("Image selection was canceled.")

    def ask_sketch(self, prompt, apikey, url):
        def encode_image(image_path):
            try:
                with open(image_path, "rb") as image_file:
                    return base64.b64encode(image_file.read()).decode('utf-8')
            except FileNotFoundError:
                print(f"No file found at {image_path}. Please check the file path.")
                return None

        openai.api_key = apikey
        base64_image = encode_image(url)
        # Set headers for the API request
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {openai.api_key}"
        }

        # Payload for the API request
        params = {
            "model": "gpt-4o",
            "temperature": 0.5,
            "user": "my_customer",
            "max_tokens": 4096,
            "top_p": 0.5,
            "stream": True,
            "messages": [
                {
                    "role": "system",
                    "content": "Sono un assistente che fornisce descrizioni dettagliate e collegamenti utili."
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
        }

        self.start_worker(headers, params, is_image=True)

    def extract_text_from_file(self, file_path):
        file_type = os.path.splitext(file_path)[1].lower()
        if file_type == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == ".csv":
            return self.extract_text_from_csv(file_path)
        elif file_type == ".docx":
            return self.extract_text_from_docx(file_path)
        else:
            print("Unsupported file type.")
            return None

    def extract_text_from_pdf(self, file_path):
        text = ""
        try:
            document = fitz.open(file_path)
            for page_num in range(len(document)):
                page = document.load_page(page_num)
                text += page.get_text()
            return text
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            return None

    def extract_text_from_csv(self, file_path):
        text = ""
        with open(file_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                text += ' '.join(row) + "\n"
        return text

    def extract_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def save_corrected_file(self, original_file_path, original_lines, corrected_text):
        file_type = os.path.splitext(original_file_path)[1].lower()
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Corrected Document", "", f"Documents (*{file_type})")

        if not save_path:
            return

        if file_type == ".pdf":
            self.save_corrected_pdf(original_file_path, save_path, original_lines, corrected_text)
        elif file_type == ".csv":
            self.save_corrected_csv(save_path, corrected_text)
        elif file_type == ".docx":
            self.save_corrected_docx(original_file_path, save_path, corrected_text)
        else:
            print("Unsupported file type.")

    def find_closest_match(self, corrected_line, original_lines):
        matcher = difflib.SequenceMatcher(None, corrected_line, None)
        best_ratio, best_match = 0.0, None

        for original_line in original_lines:
            matcher.set_seq2(original_line)
            ratio = matcher.ratio()
            if ratio > best_ratio:
                best_ratio, best_match = ratio, original_line

        return best_match if best_ratio > 0.6 else None

    def save_corrected_pdf(self, original_file_path, save_path, corrected_lines, original_lines):
        try:
            fontsize = 12
            vertical_padding = 5  # space between lines
            insert_y = 10  # initial y position
            max_iterations = 100  # Ad esempio per massimo 100 correzioni per pagina
            iterations = 0
            doc = fitz.open(original_file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                for line_num, corrected_line in enumerate(corrected_lines):
                    original_line = self.find_closest_match(corrected_line, original_lines)
                    if original_line is None:
                        print(f"No match found for corrected line '{corrected_line}'")
                        continue
                    text_instances = page.search_for(original_line)
                    if not text_instances:
                        print(f"No instances of original text '{original_line}' found on page {page_num}")
                        continue
                    for inst in text_instances:
                        if iterations >= max_iterations:
                            print(f"Reached maximum iterations ({max_iterations}), stopping.")
                            break

                        insert_point = fitz.Point(10, insert_y)
                        page.insert_text(insert_point, corrected_line, fontsize=fontsize, color=(1, 0, 0))
                        #page.update()

                        insert_y += fontsize + vertical_padding
                        iterations += 1
            doc.save(save_path)
            QMessageBox.information(self, "Success", "Corrected PDF saved successfully.")
        except Exception as e:
            print(f"Error saving PDF file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving PDF file: {e}")

    def save_corrected_csv(self, save_path, corrected_text):
        try:
            with open(save_path, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                for line in corrected_text.split('\n'):
                    writer.writerow(line.split())
            QMessageBox.information(self, "Success", "Corrected CSV saved successfully.")
        except Exception as e:
            print(f"Error saving CSV file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving CSV file: {e}")

    def save_corrected_docx(self, original_file_path, save_path, corrected_text):
        try:
            doc = docx.Document(original_file_path)
            corrected_paragraphs = corrected_text.split('\n')
            for i, paragraph in enumerate(doc.paragraphs):
                if i < len(corrected_paragraphs):
                    paragraph.text = corrected_paragraphs[i]
            doc.save(save_path)
            QMessageBox.information(self, "Success", "Corrected DOCX saved successfully.")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")
            QMessageBox.critical(self, "Error", f"Error saving DOCX file: {e}")

    def ask_doc(self, prompt, apikey, file_path):
        global reply, corrected_lines
        openai.api_key = apikey
        file_text = self.extract_text_from_file(file_path)
        if file_text is None:
            return None

        # Set headers for the API request
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {openai.api_key}"
        }

        # Payload for the API request
        params = {
            "model": "gpt-4o",
            "temperature": 0.5,
            "user": "my_customer",
            "max_tokens": 4096,
            "top_p": 0.5,
            "stream": True,
            "messages": [
                {
                    "role": "system",
                    "content": "Sono il tuo assistente personale per la lettura di documenti ed esporterò il documento nel formato di input tramite link"
                },
                {
                    "role": "user",
                    "content": prompt + "\n\n" + file_text
                }
            ],
        }
        self.start_worker(headers, params, is_image=False)


    def update_progress(self, progress):
        self.progress.setValue(progress)

    def update_content(self, content):
        combined_message = self.listWidget_ai.toPlainText() + content
        self.listWidget_ai.setPlainText(combined_message)

    def update_tokens_used(self, tokens_used, total_cost):
        self.token_counter.setText(f"Tokens used: {tokens_used} - Total cost: ${total_cost:.4f}")

    def docchgpt(self):
        self.listWidget_ai.clear()
        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(self, "Select Documents", "", "Documents (*.pdf *.csv *.docx)")
        if file_paths:
            prompt = self.prompt_label.toPlainText()
            for file_path in file_paths:
                self.ask_doc(prompt, self.apikey_gpt(), file_path)
                #combined_message = f"Document: {file_path}\nGPT Response:\n\n{gpt_response}"
                #self.listWidget_ai.addItem(combined_message)
                #QApplication.processEvents()
        else:
            self.listWidget_ai.setPlainText("Document selection was canceled.")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    mainWin = GPTWindow()
    mainWin.show()
    sys.exit(app.exec_())
