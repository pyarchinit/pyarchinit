import re
import docx
import pandas as pd
from typing import Dict, List, Optional, Tuple
from datetime import datetime

class TMAInventoryParser:
    """Parser specifico per file inventario cassette archeologiche"""
    
    def __init__(self):
        self.default_values = {
            'ogtm': 'CERAMICA',  # Tipo materiale di default
            'ldct': 'MAGAZZINO',  # Localizzazione di default
            'stato': 'COMPLETO',
            'determinazione': 'ND',  # Non determinato
            'dsstcc': 'A'  # Stato conservazione
        }
        
    def parse_docx_inventory(self, file_path: str, sito: str = None) -> List[Dict]:
        """
        Parsa un file docx con inventario cassette
        
        Args:
            file_path: Path del file docx
            sito: Nome del sito (se non specificato, cerca nel testo)
            
        Returns:
            Lista di dizionari con i record da importare
        """
        doc = docx.Document(file_path)
        records = []
        current_magazzino = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Cerca indicazione del magazzino
            if 'Magazzino' in text or 'Magazzini' in text:
                match = re.search(r'Magazzin[oi]\s*([\d\-]+)', text)
                if match:
                    current_magazzino = f"Magazzino {match.group(1)}"
                # Estrai anche il sito se presente
                if not sito and '-' in text:
                    parts = text.split('-')
                    if len(parts) >= 2:
                        sito = parts[0].strip().replace('*', '')
                continue
            
        # Ora parsa le tabelle
        for table in doc.tables:
            headers = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                # Prima riga con headers
                if i == 0 or any('Cassette' in cell for cell in cells):
                    headers = cells
                    continue
                
                # Salta righe vuote o separatori
                if all(cell in ['', '-'] for cell in cells) or len(cells) < 3:
                    continue
                
                # Parsa la riga di dati
                record = self._parse_row(cells, headers, sito, current_magazzino)
                if record:
                    records.extend(record if isinstance(record, list) else [record])
                    
        return records
    
    def _parse_row(self, cells: List[str], headers: List[str], 
                   sito: str, magazzino: str) -> List[Dict]:
        """Parsa una singola riga della tabella"""
        if len(cells) < 3:
            return []
            
        # Identifica le colonne
        cassette_col = 0
        desc_col = 1
        anno_col = 2
        
        # Cerca le colonne corrette basandosi sugli headers
        for i, header in enumerate(headers):
            if 'Cassett' in header:
                cassette_col = i
            elif 'Descr' in header:
                desc_col = i
            elif 'Anno' in header:
                anno_col = i
        
        cassette = cells[cassette_col] if cassette_col < len(cells) else ''
        descrizione = cells[desc_col] if desc_col < len(cells) else ''
        anno = cells[anno_col] if anno_col < len(cells) else ''
        
        if not cassette or cassette == '-':
            return []
            
        # Parsa range di cassette (es. "1-72")
        cassette_list = self._parse_cassette_range(cassette)
        
        records = []
        for cass_num in cassette_list:
            record = self.default_values.copy()
            
            # Campi principali
            record['sito'] = sito or 'DA DEFINIRE'
            record['dscu'] = f"Cassetta {cass_num}"
            record['ldct'] = magazzino or 'MAGAZZINO'
            
            # Descrizione
            record['descrizione'] = descrizione
            
            # Anno - gestisci range di anni (es. "1965-66")
            if anno:
                anno_clean = self._clean_anno(anno)
                if anno_clean:
                    record['anno'] = anno_clean
                    record['datazione'] = f"{anno_clean}"
                    
            # Aggiungi numero reperti se presente nella descrizione
            n_reperti = self._extract_n_reperti(descrizione)
            if n_reperti:
                record['n_reperti'] = n_reperti
                
            # Estrai tipo materiale dalla descrizione
            tipo_materiale = self._extract_tipo_materiale(descrizione)
            if tipo_materiale:
                record['ogtm'] = tipo_materiale
                
            records.append(record)
            
        return records
    
    def _parse_cassette_range(self, cassette_str: str) -> List[str]:
        """Parsa un range di cassette (es. '1-72' -> ['1', '2', ..., '72'])"""
        cassette_str = cassette_str.strip()
        
        # Gestisci lettere (es. "246d-259")
        if re.match(r'^\d+[a-z]?-\d+[a-z]?$', cassette_str):
            match = re.match(r'^(\d+)([a-z]?)-(\d+)([a-z]?)$', cassette_str)
            if match:
                start_num = int(match.group(1))
                start_letter = match.group(2)
                end_num = int(match.group(3))
                end_letter = match.group(4)
                
                # Se c'è una lettera, usa solo il numero singolo
                if start_letter or end_letter:
                    return [cassette_str]
                else:
                    return [str(i) for i in range(start_num, end_num + 1)]
        
        # Range semplice
        if '-' in cassette_str and not re.search(r'[a-z]', cassette_str):
            try:
                parts = cassette_str.split('-')
                if len(parts) == 2:
                    start = int(parts[0])
                    end = int(parts[1])
                    return [str(i) for i in range(start, end + 1)]
            except ValueError:
                pass
                
        # Cassetta singola
        return [cassette_str]
    
    def _clean_anno(self, anno_str: str) -> Optional[int]:
        """Pulisce e converte l'anno"""
        anno_str = anno_str.strip()
        
        # Cerca il primo anno valido
        match = re.search(r'(\d{4})', anno_str)
        if match:
            return int(match.group(1))
            
        # Gestisci anni abbreviati (es. "1965-66")
        match = re.search(r'(\d{4})-(\d{2})', anno_str)
        if match:
            return int(match.group(1))  # Prendi il primo anno
            
        return None
    
    def _extract_n_reperti(self, descrizione: str) -> Optional[int]:
        """Estrae il numero di reperti dalla descrizione se presente"""
        # Cerca pattern come "72 frammenti", "150 pezzi", etc.
        patterns = [
            r'(\d+)\s+framm',
            r'(\d+)\s+pezz',
            r'(\d+)\s+vas[io]',
            r'(\d+)\s+materiali'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, descrizione.lower())
            if match:
                return int(match.group(1))
                
        return None
    
    def _extract_tipo_materiale(self, descrizione: str) -> str:
        """Estrae il tipo di materiale dalla descrizione"""
        desc_lower = descrizione.lower()
        
        # Mapping descrizione -> OGTM
        material_mapping = {
            'ceramica': 'CERAMICA',
            'framm': 'CERAMICA',
            'vas': 'CERAMICA',
            'pitho': 'CERAMICA',
            'anfora': 'CERAMICA',
            'oss': 'FAUNA',
            'conchigli': 'MALACOFAUNA',
            'pietra': 'INDUSTRIA LITICA',
            'macin': 'INDUSTRIA LITICA',
            'ferro': 'METALLO',
            'bronz': 'METALLO',
            'stucc': 'INTONACO',
            'intonac': 'INTONACO',
            'blocch': 'MATERIALE DA COSTRUZIONE',
            'calcestru': 'MATERIALE DA COSTRUZIONE',
            'astraki': 'MATERIALE DA COSTRUZIONE'
        }
        
        for key, value in material_mapping.items():
            if key in desc_lower:
                return value
                
        return 'CERAMICA'  # Default


class TMAInventoryImportDialog:
    """Dialog per completare i dati mancanti durante l'import"""
    
    def __init__(self, parent=None):
        self.parent = parent
        self.site_default = None
        self.material_defaults = {}
        
    def get_completion_data(self, records: List[Dict]) -> Tuple[List[Dict], bool]:
        """
        Mostra un dialog per completare i dati mancanti
        
        Returns:
            Tuple di (records aggiornati, conferma import)
        """
        from qgis.PyQt.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, 
                                     QLabel, QLineEdit, QComboBox, 
                                     QPushButton, QTableWidget, QTableWidgetItem,
                                     QDialogButtonBox, QGroupBox, QFormLayout,
                                     QMessageBox)
        
        dialog = QDialog(self.parent)
        dialog.setWindowTitle("Completa dati per importazione TMA")
        dialog.resize(800, 600)
        
        layout = QVBoxLayout()
        
        # Info sui record
        info_label = QLabel(f"Trovati {len(records)} record da importare.\n"
                          "Completa i campi mancanti:")
        layout.addWidget(info_label)
        
        # Gruppo valori di default
        default_group = QGroupBox("Valori di default per tutti i record")
        default_layout = QFormLayout()
        
        # Sito
        self.site_edit = QLineEdit()
        if records and records[0].get('sito') != 'DA DEFINIRE':
            self.site_edit.setText(records[0].get('sito', ''))
        default_layout.addRow("Sito:", self.site_edit)
        
        # Localizzazione dettagliata
        self.location_edit = QLineEdit()
        self.location_edit.setPlaceholderText("es. Magazzino 3-4, Scaffale A")
        default_layout.addRow("Localizzazione:", self.location_edit)
        
        # Responsabile
        self.resp_edit = QLineEdit()
        self.resp_edit.setPlaceholderText("Nome del responsabile")
        default_layout.addRow("Responsabile:", self.resp_edit)
        
        # Data inventario
        self.date_edit = QLineEdit()
        self.date_edit.setText(datetime.now().strftime("%Y-%m-%d"))
        default_layout.addRow("Data inventario:", self.date_edit)
        
        default_group.setLayout(default_layout)
        layout.addWidget(default_group)
        
        # Tabella preview
        preview_label = QLabel("Anteprima record (primi 10):")
        layout.addWidget(preview_label)
        
        table = QTableWidget()
        table.setColumnCount(6)
        table.setHorizontalHeaderLabels(['Cassetta', 'Descrizione', 'Anno', 
                                        'Tipo materiale', 'N. reperti', 'Sito'])
        
        # Mostra primi 10 record
        preview_records = records[:10]
        table.setRowCount(len(preview_records))
        
        for i, record in enumerate(preview_records):
            table.setItem(i, 0, QTableWidgetItem(record.get('dscu', '')))
            table.setItem(i, 1, QTableWidgetItem(record.get('descrizione', '')))
            table.setItem(i, 2, QTableWidgetItem(str(record.get('anno', ''))))
            table.setItem(i, 3, QTableWidgetItem(record.get('ogtm', '')))
            table.setItem(i, 4, QTableWidgetItem(str(record.get('n_reperti', ''))))
            table.setItem(i, 5, QTableWidgetItem(record.get('sito', '')))
            
        layout.addWidget(table)
        
        # Bottoni
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)
        
        dialog.setLayout(layout)
        
        if dialog.exec():
            # Aggiorna tutti i record con i valori di default
            sito = self.site_edit.text() or 'DA DEFINIRE'
            location = self.location_edit.text()
            resp = self.resp_edit.text()
            date_inv = self.date_edit.text()
            
            for record in records:
                if sito != 'DA DEFINIRE':
                    record['sito'] = sito
                if location:
                    record['ldct'] = location
                if resp:
                    record['responsabile'] = resp
                if date_inv:
                    record['data_inventario'] = date_inv
                    
            return records, True
        else:
            return records, False