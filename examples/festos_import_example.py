#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Esempio di utilizzo del parser Festos per importare inventari archeologici
"""

import os
import sys

# Aggiungi il percorso del progetto al PYTHONPATH
project_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_path)

from modules.utility.festos_inventory_parser import FestosInventoryParser
from modules.db_manager import DBManager


def import_festos_inventory(docx_file_path):
    """
    Importa un inventario Festos da file DOCX
    
    Args:
        docx_file_path: percorso del file DOCX da importare
    """
    # Crea connessione al database
    db_manager = DBManager()
    
    try:
        # Crea il parser
        parser = FestosInventoryParser(db_manager.session)
        
        print(f"Importazione file: {os.path.basename(docx_file_path)}")
        print("-" * 50)
        
        # Esegui l'importazione
        success_count, errors, warnings = parser.parse_file(docx_file_path)
        
        # Mostra risultati
        print(f"\nRecord importati con successo: {success_count}")
        
        if warnings:
            print(f"\nAvvisi ({len(warnings)}):")
            for warning in warnings:
                print(f"  - {warning}")
        
        if errors:
            print(f"\nErrori ({len(errors)}):")
            for error in errors:
                print(f"  - {error}")
        
        # Commit delle modifiche se non ci sono errori critici
        if success_count > 0 and not errors:
            parser.commit_changes()
            print("\nDati salvati nel database con successo!")
        else:
            print("\nNessun dato salvato a causa di errori.")
            
    except Exception as e:
        print(f"Errore durante l'importazione: {str(e)}")
    finally:
        db_manager.close_connection()


def batch_import_festos(directory_path):
    """
    Importa tutti i file DOCX di inventario Festos da una directory
    
    Args:
        directory_path: percorso della directory contenente i file DOCX
    """
    # Trova tutti i file DOCX nella directory
    docx_files = [f for f in os.listdir(directory_path) 
                  if f.endswith('.docx') and not f.startswith('~')]
    
    if not docx_files:
        print(f"Nessun file DOCX trovato in: {directory_path}")
        return
    
    print(f"Trovati {len(docx_files)} file DOCX da importare")
    print("=" * 50)
    
    # Crea connessione al database
    db_manager = DBManager()
    
    total_imported = 0
    total_errors = 0
    total_warnings = 0
    
    try:
        # Crea il parser una sola volta
        parser = FestosInventoryParser(db_manager.session)
        
        for docx_file in docx_files:
            file_path = os.path.join(directory_path, docx_file)
            print(f"\nImportazione: {docx_file}")
            print("-" * 30)
            
            # Esegui l'importazione
            success_count, errors, warnings = parser.parse_file(file_path)
            
            total_imported += success_count
            total_errors += len(errors)
            total_warnings += len(warnings)
            
            print(f"  Record importati: {success_count}")
            if errors:
                print(f"  Errori: {len(errors)}")
            if warnings:
                print(f"  Avvisi: {len(warnings)}")
        
        # Commit finale di tutti i dati
        if total_imported > 0 and total_errors == 0:
            parser.commit_changes()
            print("\n" + "=" * 50)
            print("IMPORTAZIONE COMPLETATA")
            print(f"Totale record importati: {total_imported}")
            print(f"Totale avvisi: {total_warnings}")
            print("Dati salvati nel database con successo!")
        else:
            print("\n" + "=" * 50)
            print("IMPORTAZIONE INTERROTTA")
            print(f"Totale errori: {total_errors}")
            print("Nessun dato salvato a causa di errori.")
            
    except Exception as e:
        print(f"\nErrore durante l'importazione batch: {str(e)}")
    finally:
        db_manager.close_connection()


def test_parser_with_sample():
    """Test del parser con dati di esempio"""
    # Crea un file DOCX di esempio
    from docx import Document
    
    doc = Document()
    
    # Aggiungi titolo
    doc.add_heading('Festos - Magazzino 5', level=1)
    
    # Crea tabella formato cassette
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Cassette'
    header_cells[1].text = 'Descrizione'
    header_cells[2].text = 'Anno'
    
    # Aggiungi dati di esempio
    sample_data = [
        ('1-10', 'Ceramica minoica dal vano A', '1960'),
        ('11', 'Frammenti di pithos', '1960'),
        ('12-15', 'Materiali vari dal saggio I', '1961'),
        ('16', 'Ossidiana e selce', '1962'),
    ]
    
    for cassette, desc, anno in sample_data:
        row = table.add_row()
        row.cells[0].text = cassette
        row.cells[1].text = desc
        row.cells[2].text = anno
    
    # Salva il file di test
    test_file = 'test_festos_inventory.docx'
    doc.save(test_file)
    print(f"File di test creato: {test_file}")
    
    # Importa il file di test
    import_festos_inventory(test_file)
    
    # Rimuovi il file di test
    os.remove(test_file)
    print(f"\nFile di test rimosso: {test_file}")


if __name__ == "__main__":
    # Esempi di utilizzo
    
    # 1. Test con dati di esempio
    print("TEST CON DATI DI ESEMPIO")
    print("=" * 50)
    test_parser_with_sample()
    
    # 2. Importa un singolo file (decommentare e specificare il percorso)
    # import_festos_inventory("/path/to/your/festos_inventory.docx")
    
    # 3. Importa tutti i file DOCX da una directory (decommentare e specificare il percorso)
    # batch_import_festos("/path/to/directory/with/docx/files")
